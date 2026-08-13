import json, os, bcrypt, smtplib, random, io
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

import streamlit as st
import pandas as pd
from supabase import create_client, Client

# بتوع التقارير
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import arabic_reshaper
from openpyxl.styles import Font, Alignment, PatternFill

st.set_page_config(page_title="إدارة القضايا", layout="wide")

# ====== CSS القنبلة للسحابة ======
st.markdown("""
<style>
html, body, [class*="css"] {
    direction: rtl !important;
}
.main .block-container { padding-top: 2rem; padding-left: 1rem; padding-right: 1rem; max-width: 100%; }
.stApp { background-color: #0E1117; }
h1, h2, h3, h4, h5, h6 { color: white!important; text-align: center; }

.stButton>button { 
    background-color: #C9A961; color: black; font-weight: bold; 
    border-radius: 10px; width: 100%; white-space: normal !important; line-height: 1.4;
}

/* قنبلة السايدبار - تمسح اي عمودي */
section[data-testid="stSidebar"] * {
    writing-mode: horizontal-tb !important;
    text-orientation: mixed !important;
    transform: none !important;
    direction: rtl !important;
}
</style>
""", unsafe_allow_html=True)

# ====== تهيئة السيشن ستيت ======
if "page" not in st.session_state: st.session_state.page = "login"
if "user" not in st.session_state: st.session_state.user = None
if "role" not in st.session_state: st.session_state.role = None
if "RESET_CODES" not in st.session_state: st.session_state.RESET_CODES = {}

# ====== الاتصال بالسحابة ======
supabase: Client = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])

# ====== اعدادات الادمن ======
ADMIN_USERNAME = "admin"
ADMIN_DEFAULT_PASS = "admin123"

def fix_arabic(text):
    """ نسخة متعدلة للسحابة - من غير bidi """
    if not text: 
        return ""
    reshaped_text = arabic_reshaper.reshape(str(text))
    return reshaped_text

# ===== نظام اليافطة - متعدل للسحابة + RTL ثابت =====
def load_banners():
    res = supabase.table("banners").select("*").order("created_at", desc=True).execute()
    return res.data if res.data else []

def save_banner_to_db(banner_data):
    supabase.table("banners").insert(banner_data).execute()

def delete_banner_from_db(banner_id):
    supabase.table("banners").delete().eq("id", banner_id).execute()

def init_session_state():
    if "banners" not in st.session_state:
        st.session_state.banners = load_banners()
    if "banners" in st.session_state and st.session_state.banners is None:
        st.session_state.banners = []

def show_banners():
    """ يعرض اليافطات اللي لسه منتهتش ولليوزر ده بس """
    init_session_state()
    
    now = datetime.now()
    current_user = st.session_state.user["username"]
    active_banners = []
    banners_to_delete = []
    
    for b in st.session_state.banners:
        if not isinstance(b, dict) or "expire" not in b: continue
        try: expire_date = datetime.fromisoformat(b["expire"])
        except: continue
        
        if expire_date > now:
            audience = b.get("audience", "الكل")
            visible_to = b.get("visible_to", [])
            if audience == "الكل" or current_user in visible_to:
                active_banners.append(b)
        else:
            banners_to_delete.append(b["id"])
            
    for banner_id in banners_to_delete:
        delete_banner_from_db(banner_id)

    st.session_state.banners = active_banners

    for banner in active_banners:
        # عدلت direction من ltr ل rtl
        st.markdown(f"""
        <div style="
            direction: rtl !important;
            text-align: right;
            background:linear-gradient(90deg, {banner['color']}, #ffffff22); 
            padding:14px; border-radius:12px; 
            font-size:24px; font-weight:bold; color:white; margin:15px 0;
            border: 2px solid {banner['color']}; animation: pulse 2s infinite;
            white-space: normal !important; word-wrap: break-word;
            writing-mode: horizontal-tb !important;
        ">
            📢 {banner['text']}
        </div>
        <style>@keyframes pulse {{ 0% {{transform: scale(1);}} 50% {{transform: scale(1.02);}} 100% {{transform: scale(1);}} }}</style>
        """, unsafe_allow_html=True)

def banner_sidebar():
    if 'role' not in st.session_state or st.session_state.role != 'admin':
        return 
    
    init_session_state()
    users = load_users()
    
    st.sidebar.markdown("---")
    # عدلت title ل markdown عشان نجبره افقي
    st.sidebar.markdown('<h3 style="writing-mode: horizontal-tb !important; text-align: center; color: #C9A961;">📢 تحكم الادمن</h3>', unsafe_allow_html=True)
    
    with st.sidebar.form("add_banner_form"):
        banner_text = st.text_input("اكتب التهنئة")
        banner_color = st.color_picker("اللون", "#FFD700")
        duration_minutes = st.number_input("المدة بالدقايق", 1, 10080, 60)
        
        st.markdown("### 👥 الظهور لـ")
        audience_type = st.radio("اختر الجمهور", ["الكل", "اعضاء محددين"], horizontal=True, key="audience_banner")
        
        visible_to = []
        if audience_type == "اعضاء محددين":
            all_usernames = [u["username"] for u in users]
            visible_to = st.multiselect("حدد الاعضاء", all_usernames, key="visible_users_banner")

        if st.form_submit_button("اضافة يافطة"):
            if banner_text and (audience_type == "الكل" or visible_to):
                expire_time = datetime.now() + timedelta(minutes=duration_minutes)
                new_banner = {
                    "text": banner_text, 
                    "color": banner_color, 
                    "expire": expire_time.isoformat(),
                    "created_at": datetime.now().isoformat(),
                    "audience": audience_type,
                    "visible_to": visible_to
                }
                save_banner_to_db(new_banner)
                st.session_state.banners = load_banners()
                st.success("تم النشر"); st.rerun()
            else: st.error("املى كل الحقول")

    st.sidebar.markdown("### حذف اليافطات")
    for i, banner in enumerate(st.session_state.banners):
        col1, col2 = st.sidebar.columns([4,1])
        with col1: 
            audience_info = "الكل" if banner.get("audience")=="الكل" else "محدد"
            st.write(f"• {banner['text'][:20]}... ({audience_info})")
        with col2: 
            if st.button("🗑️", key=f"del_admin_{banner['id']}"):
                delete_banner_from_db(banner['id'])
                st.session_state.banners = load_banners()
                st.rerun()
# ===== نهاية اليافطة =====
import json, os, bcrypt, smtplib, random, io
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

import streamlit as st
import pandas as pd
from supabase import create_client, Client

# بتوع التقارير
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import arabic_reshaper
from openpyxl.styles import Font, Alignment, PatternFill

st.set_page_config(page_title="إدارة القضايا", layout="wide")

# ====== CSS الاساسي + قنبلة للسايدبار ======
st.markdown("""
<style>
html, body, [class*="css"] {
    direction: rtl!important;
    overflow-x: hidden!important;
}

.main.block-container { padding-top: 2rem; padding-left: 1rem; padding-right: 1rem; max-width: 100%; }
.stApp { background-color: #0E1117; }
h1, h2, h3, h4, h5, h6 { color: white!important; text-align: center; }

.stButton>button {
    background-color: #C9A961; color: black; font-weight: bold;
    border-radius: 10px; width: 100%; white-space: normal!important; line-height: 1.4;
}

.stTextInput>div>div>input,.stSelectbox>div>div>div,.stTextArea>div>div>textarea {
    color: black; background-color: white; border-radius: 8px;
    direction: rtl!important; text-align: right;
}

div[data-testid="stWidgetLabel"] p {
    color: #C9A961!important; font-size: 16px!important; font-weight: 700!important;
}

thead tr th { color: black!important; background-color: #C9A961!important; font-weight: bold; }
tbody tr td { color: black!important; background-color: white!important; }

/* قنبلة السايدبار - تمسح اي عمودي */
section[data-testid="stSidebar"] * {
    writing-mode: horizontal-tb!important;
    text-orientation: mixed!important;
    transform: none!important;
    display: block!important;
    white-space: normal!important;
    direction: rtl!important;
}
</style>
""", unsafe_allow_html=True)

# ====== تهيئة السيشن ستيت ======
if "page" not in st.session_state: st.session_state.page = "login"
if "user" not in st.session_state: st.session_state.user = None
if "role" not in st.session_state: st.session_state.role = None
if "RESET_CODES" not in st.session_state: st.session_state.RESET_CODES = {}

# ====== الاتصال بالسحابة ======
supabase: Client = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])

# ====== اعدادات الادمن ======
ADMIN_USERNAME = "admin"
ADMIN_DEFAULT_PASS = "admin123"

def fix_arabic(text):
    """ نسخة متعدلة للسحابة - من غير bidi """
    if not text:
        return ""
    reshaped_text = arabic_reshaper.reshape(str(text))
    return reshaped_text

# ===== نظام اليافطة - متعدل للسحابة + RTL ثابت =====
def load_banners():
    res = supabase.table("banners").select("*").order("created_at", desc=True).execute()
    return res.data if res.data else []

def save_banner_to_db(banner_data):
    supabase.table("banners").insert(banner_data).execute()

def delete_banner_from_db(banner_id):
    supabase.table("banners").delete().eq("id", banner_id).execute()

def init_session_state():
    if "banners" not in st.session_state:
        st.session_state.banners = load_banners()
    if "banners" in st.session_state and st.session_state.banners is None:
        st.session_state.banners = []

def show_banners():
    """ يعرض اليافطات اللي لسه منتهتش ولليوزر ده بس """
    init_session_state()

    now = datetime.now()
    current_user = st.session_state.user["username"]
    active_banners = []
    banners_to_delete = []

    for b in st.session_state.banners:
        if not isinstance(b, dict) or "expire" not in b: continue
        try: expire_date = datetime.fromisoformat(b["expire"])
        except: continue

        if expire_date > now:
            audience = b.get("audience", "الكل")
            visible_to = b.get("visible_to", [])
            if audience == "الكل" or current_user in visible_to:
                active_banners.append(b)
        else:
            banners_to_delete.append(b["id"])

    for banner_id in banners_to_delete:
        delete_banner_from_db(banner_id)

    st.session_state.banners = active_banners

    for banner in active_banners:
        # عدلت من ltr ل rtl واجبرتها افقي
        st.markdown(f"""
        <div style="
            direction: rtl!important;
            writing-mode: horizontal-tb!important;
            text-align: right;
            background:linear-gradient(90deg, {banner['color']}, #ffffff22);
            padding:14px; border-radius:12px;
            font-size:24px; font-weight:bold; color:white; margin:15px 0;
            border: 2px solid {banner['color']}; animation: pulse 2s infinite;
            white-space: normal!important; word-wrap: break-word;
        ">
            📢 {banner['text']}
        </div>
        <style>@keyframes pulse {{ 0% {{transform: scale(1);}} 50% {{transform: scale(1.02);}} 100% {{transform: scale(1);}} }}</style>
        """, unsafe_allow_html=True)

def banner_sidebar():
    if 'role' not in st.session_state or st.session_state.role!= 'admin':
        return

    init_session_state()
    users = load_users()

    st.sidebar.markdown("---")
    # عدلت title ل markdown عشان ميقلبش عمودي
    st.sidebar.markdown('<h3 style="writing-mode: horizontal-tb!important; text-align: center; color: #C9A961;">📢 تحكم الادمن</h3>', unsafe_allow_html=True)

    with st.sidebar.form("add_banner_form"):
        banner_text = st.text_input("اكتب التهنئة")
        banner_color = st.color_picker("اللون", "#FFD700")
        duration_minutes = st.number_input("المدة بالدقايق", 1, 10080, 60)

        st.markdown("### 👥 الظهور لـ")
        audience_type = st.radio("اختر الجمهور", ["الكل", "اعضاء محددين"], horizontal=True, key="audience_banner")

        visible_to = []
        if audience_type == "اعضاء محددين":
            all_usernames = [u["username"] for u in users]
            visible_to = st.multiselect("حدد الاعضاء", all_usernames, key="visible_users_banner")

        if st.form_submit_button("اضافة يافطة"):
            if banner_text and (audience_type == "الكل" or visible_to):
                expire_time = datetime.now() + timedelta(minutes=duration_minutes)
                new_banner = {
                    "text": banner_text,
                    "color": banner_color,
                    "expire": expire_time.isoformat(),
                    "created_at": datetime.now().isoformat(),
                    "audience": audience_type,
                    "visible_to": visible_to
                }
                save_banner_to_db(new_banner)
                st.session_state.banners = load_banners()
                st.success("تم النشر"); st.rerun()
            else: st.error("املى كل الحقول")

    st.sidebar.markdown("### حذف اليافطات")
    for i, banner in enumerate(st.session_state.banners):
        col1, col2 = st.sidebar.columns([4,1])
        with col1:
            audience_info = "الكل" if banner.get("audience")=="الكل" else "محدد"
            st.write(f"• {banner['text'][:20]}... ({audience_info})")
        with col2:
            if st.button("🗑️", key=f"del_admin_{banner['id']}"):
                delete_banner_from_db(banner['id'])
                st.session_state.banners = load_banners()
                st.rerun()
# ===== نهاية اليافطة =====

# ====== الاعدادات ======
SENDER_EMAIL = st.secrets.get("SENDER_EMAIL", "")
SENDER_PASSWORD = st.secrets.get("SENDER_PASSWORD", "")

if "RESET_CODES" not in st.session_state:
    st.session_state.RESET_CODES = {}

# ====== ارسال الايميل ======
def send_email(to_email, subject, body):
    if not SENDER_EMAIL:
        st.warning("مفعلتش الايميل لسه. حطه في Secrets")
        return False
    try:
        msg = MIMEText(body, "plain", "utf-8")
        msg["Subject"] = subject
        msg["From"] = SENDER_EMAIL
        msg["To"] = to_email
        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.sendmail(SENDER_EMAIL, to_email, msg.as_string())
        server.quit()
        return True
    except Exception as e:
        st.error(f"خطأ في الارسال: {e}")
        return False

# ====== اليوزرز في السحابة ======
def load_users():
    res = supabase.table("users").select("*").execute()
    users = res.data

    if not any(u["username"] == ADMIN_USERNAME for u in users):
        admin_pass = bcrypt.hashpw(ADMIN_DEFAULT_PASS.encode(), bcrypt.gensalt()).decode()
        supabase.table("users").insert({
            "username": ADMIN_USERNAME, "password": admin_pass, "role": "admin",
            "status": "active", "password_set": True, "email": ""
        }).execute()
        return load_users()
    return users

def check_login(username, password):
    users = load_users()
    for user in users:
        if user["username"] == username and user["status"] == "active":
            if not user.get("password") or not user.get("password_set", False):
                return None
            try:
                if bcrypt.checkpw(password.encode(), user["password"].encode()):
                    return user
            except Exception:
                return None
    return None

def is_admin_email(email):
    users = load_users()
    admin = next((u for u in users if u["role"] == "admin"), None)
    if not admin: return False
    return email == admin.get("email") or email == admin.get("recovery_email","")

def add_user_db(username):
    supabase.table("users").insert({
        "username": username, "password": "", "email": "", "role": "member",
        "status": "active", "password_set": False
    }).execute()

def update_user_db(user_id, new_data):
    supabase.table("users").update(new_data).eq("id", user_id).execute()

def delete_user_db(user_id):
    supabase.table("users").delete().eq("id", user_id).execute()

# ====== الصفحات ======
def login_page():
    st.markdown("<h3 style='text-align:center; color:white'>دخول السادة الاعضاء</h3>", unsafe_allow_html=True)
    tab1, tab2 = st.tabs(["تسجيل الدخول", "تفعيل حساب جديد"])

    with tab1:
        st.markdown("<p style='color:white; font-weight:bold;'>اسم المستخدم</p>", unsafe_allow_html=True)
        username = st.text_input("", key="login_user", label_visibility="collapsed")
        st.markdown("<p style='color:white; font-weight:bold;'>كلمة السر</p>", unsafe_allow_html=True)
        password = st.text_input("", type="password", key="login_pass", label_visibility="collapsed")

        if st.button("دخول", type="primary", use_container_width=True):
            user = check_login(username, password)
            if user:
                st.session_state.user = user
                st.session_state.role = user["role"]
                if user["role"] == "member" and not user.get("password_set", False):
                    st.session_state.page = "set_password"
                    st.session_state.temp_user = user["username"]
                    st.rerun()
                else:
                    st.session_state.page = "الرئيسية"
                    st.rerun()
            else:
                st.error("اسم المستخدم او كلمة السر غلط او العضوية موقوفة")

        st.markdown("---")
        st.markdown("<p style='color:white; font-weight:bold;'>نسيت بياناتك؟ استرجعها بالايميل</p>", unsafe_allow_html=True)
        admin_recover_email = st.text_input("ايميل الادمن", key="admin_recover")
        if st.button("ارسال كود للادمن", key="admin_send", use_container_width=True):
            if is_admin_email(admin_recover_email):
                code = str(random.randint(100000, 999))
                st.session_state.RESET_CODES[admin_recover_email] = {"code": code, "role": "admin"}
                body = f"كود اعادة تعيين كلمة سر الادمن: {code}"
                if send_email(admin_recover_email, "كود استرجاع الادمن", body):
                    st.success(f"تم ارسال الكود على {admin_recover_email}")
                    st.session_state.show_reset_admin = True
            else: st.error("هذا الايميل غير مسجل كادمن")

        member_recover_email = st.text_input("ايميل العضو", key="member_recover")
        if st.button("ارسال كود للعضو", key="member_send", use_container_width=True):
            users = load_users()
            found = [u for u in users if u.get("email") == member_recover_email]
            if found:
                user = found[0]
                code = str(random.randint(100000, 999999))
                st.session_state.RESET_CODES[member_recover_email] = {"code": code, "user_id": user["id"]}
                body = f"مرحبا {user['username']}\nاسم المستخدم: {user['username']}\nكود اعادة التعيين: {code}"
                if send_email(member_recover_email, "استرجاع بيانات الدخول", body):
                    st.success("تم ارسال البيانات على ايميلك")
                    st.session_state.show_reset_member = True
            else: st.error("الايميل ده مش متسجل")

        if st.session_state.get("show_reset_admin") or st.session_state.get("show_reset_member"):
            email_to_reset = admin_recover_email if st.session_state.get("show_reset_admin") else member_recover_email
            code_input = st.text_input("ادخل الكود")
            new_pass = st.text_input("كلمة السر الجديدة", type="password")
            if st.button("تأكيد وتغيير كلمة السر"):
                if st.session_state.RESET_CODES.get(email_to_reset, {}).get("code") == code_input:
                    users = load_users()
                    logged_user = None
                    if st.session_state.RESET_CODES[email_to_reset].get("role") == "admin":
                        admin = next((u for u in users if u["role"] == "admin"), None)
                        hashed = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt()).decode()
                        update_user_db(admin["id"], {"password": hashed})
                        logged_user = check_login(admin["username"], new_pass)
                    else:
                        user_id = st.session_state.RESET_CODES[email_to_reset]["user_id"]
                        hashed = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt()).decode()
                        update_user_db(user_id, {"password": hashed, "status": "active", "password_set": True})
                        logged_user = check_login(next(u["username"] for u in users if u["id"]==user_id), new_pass)

                    st.session_state.RESET_CODES.clear()
                    st.session_state.show_reset_admin = False
                    st.session_state.show_reset_member = False
                    st.session_state.user = logged_user
                    st.session_state.role = logged_user["role"]
                    st.session_state.page = "الرئيسية"
                    st.success("تم تسجيل الدخول بنجاح")
                    st.rerun()
                else: st.error("الكود غلط")

    with tab2:
        st.markdown("**تفعيل حساب عضو**")
        member_name = st.text_input("اكتب اسم العضو للتفعيل", key="new_user")
        if st.button("تفعيل الحساب", use_container_width=True):
            users = load_users()
            found_user = next((u for u in users if u.get('username') == member_name), None)
            if not found_user:
                st.error("الاسم ده مش موجود")
            elif found_user.get("password_set"):
                st.error("العضو ده مفعل بالفعل")
            else:
                st.session_state.page = "set_password"
                st.session_state.temp_user = found_user["username"]
                st.rerun()

def extract_member_page():
    st.markdown("<h2 style='text-align:center; color:#C9A961'>استخراج عضوية جديدة</h2>", unsafe_allow_html=True)
    if st.button("العودة للرئيسية"):
        st.session_state.page = "الرئيسية"; st.session_state.role = None; st.rerun()
    new_username = st.text_input("اسم المستخدم الجديد")
    if st.button("استخراج العضو", use_container_width=True, type="primary"):
        if new_username.strip():
            users = load_users()
            existing_user = next((u for u in users if u['username'] == new_username), None)
            if existing_user:
                if existing_user["status"] == "banned" or not existing_user.get("password_set"):
                    update_user_db(existing_user["id"], {"status": "active", "password": "", "password_set": False})
                    st.success(f"تم اعادة استخراج {new_username}")
                    st.rerun()
                else:
                    st.error("الاسم موجود والعضو مفعل بالفعل")
            else:
                add_user_db(new_username)
                st.success(f"تم استخراج العضو: {new_username}")
                st.rerun()

def manage_users_page():
    st.markdown("<h2 style='text-align:center; color:#C9A961'>ادارة الاعضاء</h2>", unsafe_allow_html=True)
    if st.button("العودة للرئيسية"): st.session_state.page = "الرئيسية"; st.session_state.role = None; st.rerun()
    users = load_users()
    for user in users:
        if user["role"] == "member":
            status = "مفعل" if user.get("password_set") else "غير مفعل"
            if user["status"] == "banned": status = "موقوف لمخالفة"
            with st.container(border=True):
                col1, col2 = st.columns([3,2])
                with col1:
                    st.write(f"**{user['username']}** - {user.get('email','بدون ايميل')}")
                    st.write(f"الحالة: {status}")
                with col2:
                    if user["status"] == "active":
                        if st.button("ايقاف لمخالفة قواعد", key=f"ban_{user['id']}"):
                            update_user_db(user["id"], {"status": "banned", "password": "", "password_set": False})
                            st.rerun()
                        if st.button("ايقاف لفقد البيانات", key=f"lose_{user['id']}"):
                            update_user_db(user["id"], {"password": "", "password_set": False})
                            st.rerun()
                    elif user["status"] == "banned":
                        if st.button("تنشيط", key=f"unban_{user['id']}", type="primary"):
                            update_user_db(user["id"], {"status": "active"})
                            st.success(f"تم تنشيط {user['username']}"); st.rerun()
                    else:
                        if st.button("اعادة استخراج", key=f"re_extract_{user['id']}", type="primary"):
                            update_user_db(user["id"], {"status": "active", "password": "", "password_set": False})
                            st.success(f"تم اعادة استخراج {user['username']}"); st.rerun()
                    if st.button("حذف", key=f"del_{user['id']}"):
                        delete_user_db(user['id']); st.rerun()

def recovery_settings_page():
    st.markdown("<h2 style='text-align:center; color:#C9A961'>تأكيد البريد الالكتروني</h2>", unsafe_allow_html=True)
    if st.button("العودة للرئيسية"): st.session_state.page = "الرئيسية"; st.session_state.role = None; st.rerun()
    users = load_users()
    user = next((u for u in users if u["id"] == st.session_state.user["id"]), None)
    email = st.text_input("البريد الالكتروني", value=user.get("email",""))
    recovery_email = st.text_input("ايميل استرجاع اضافي للادمن", value=user.get("recovery_email","")) if user["role"] == "admin" else user.get("recovery_email","")
    if st.button("حفظ البريد", use_container_width=True):
        new_data = {"email": email}
        if user["role"] == "admin": new_data["recovery_email"] = recovery_email
        update_user_db(user["id"], new_data)
        st.session_state.user = {**user, **new_data}
        st.success("تم حفظ البريد بنجاح")

def change_password_page():
    st.markdown("<h1 style='text-align:center; color:#C9A961'>تغيير كلمة السر</h1>", unsafe_allow_html=True)
    if st.button("العودة للرئيسية"): 
        st.session_state.page = "الرئيسية"; st.session_state.role = None; st.rerun()
    
    old_pass = st.text_input("كلمة السر القديمة", type="password")
    new_pass = st.text_input("كلمة السر الجديدة", type="password")
    
    if st.button("تغيير", use_container_width=True):
        if bcrypt.checkpw(old_pass.encode(), st.session_state.user["password"].encode()):
            hashed = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt()).decode()
            update_user_db(st.session_state.user["id"], {"password": hashed})
            st.session_state.user["password"] = hashed
            st.success("تم تغيير الباسورد"); st.rerun()
        else: 
            st.error("كلمة السر القديمة غلط")
# ===== تشغيل الصفحات =====
if st.session_state.page == "login":
    login_page()

elif st.session_state.page == "extract":
    if st.session_state.user and st.session_state.user["role"] == "admin":
        extract_member_page()
    else:
        st.session_state.page = "login"; st.rerun()

elif st.session_state.page == "manage":
    if st.session_state.user and st.session_state.user["role"] == "admin":
        manage_users_page()

        if st.button("⚙️ إدارة اليافطات", use_container_width=True):
            st.session_state.page = "banners"
            st.rerun()
    else:
        st.session_state.page = "login"; st.rerun()

elif st.session_state.page == "banners":
    st.markdown("<h2>⚙️ إدارة اليافطات</h2>", unsafe_allow_html=True)

    if st.button("العودة لإدارة الاعضاء", use_container_width=True):
        st.session_state.page = "manage"
        st.rerun()

    st.write("---")

    with st.expander("➕ اضافة يافطة جديدة"):
        st.markdown("<p style='color:white; font-weight:bold; font-size:16px;'>🏷️ اسم اليافطة</p>", unsafe_allow_html=True)
        title = st.text_input("", label_visibility="collapsed")

        st.markdown("<p style='color:white; font-weight:bold; font-size:16px;'>📝 محتوى اليافطة</p>", unsafe_allow_html=True)
        content = st.text_area("", label_visibility="collapsed")

        st.markdown("<p style='color:white; font-weight:bold; font-size:16px;'>🎨 لون اليافطة</p>", unsafe_allow_html=True)
        color_option = st.selectbox("", ["اصفر", "احمر", "اخضر", "ازرق", "برتقاني"], label_visibility="collapsed")
        colors = {"اصفر": "#FFFF00","احمر": "#FF0000","اخضر": "#00FF00","ازرق": "#00BFFF","برتقاني": "#FF8C00"}
        selected_color = colors[color_option]

        st.markdown("<p style='color:white; font-weight:bold; font-size:16px;'>🔤 حجم الخط</p>", unsafe_allow_html=True)
        font_size = st.slider("", 14, 32, 18, label_visibility="collapsed")

        if st.button("💾 حفظ اليافطة", use_container_width=True):
            if title and content:
                expire_time = datetime.now() + timedelta(days=7)
                new_banner = {
                    "text": f"<b>{title}</b><br>{content}",
                    "color": selected_color,
                    "expire": expire_time.isoformat(),
                    "created_at": datetime.now().isoformat()
                }
                save_banner_to_db(new_banner) # <--- متعدلة للسحابة
                st.session_state.banners = load_banners() # <--- ريفريش
                st.success("✅ تم اضافة اليافطة")
                st.rerun()
            else:
                st.error("لازم تكتب اسم ومحتوى اليافطة")

    st.write("---")
    st.markdown("<h3>🗑️ اليافطات الموجودة</h3>", unsafe_allow_html=True)

    banners = load_banners()
    if not banners:
        st.info("مفيش يافطات لسه")

    for banner in banners:
        col1, col2 = st.columns([4,1])
        with col1:
            st.markdown(f"<div style='background:{banner['color']}; padding:10px; border-radius:8px;'>{banner['text']}</div>", unsafe_allow_html=True)
        with col2:
            if st.button("🗑️ حذف", key=f"del_banner_{banner['id']}"):
                delete_banner_from_db(banner['id']) # <--- متعدلة للسحابة
                st.session_state.banners = load_banners() # <--- ريفريش
                st.success("تم الحذف")
                st.rerun()

elif st.session_state.page == "recovery":
    recovery_settings_page()

elif st.session_state.page == "set_password":
    set_password_page()

elif st.session_state.page == "change_pass":
    change_password_page()

elif st.session_state.page == "الرئيسية":
    init_session_state() # <--- ضفت دي
    show_banners() # <--- ضفت دي
    st.title("الرئيسية")
    st.write(f"اهلا {st.session_state.user['username']}")
    banner_sidebar()

    st.markdown("""
    <style>
   .stButton>button {
        color: white!important;
        background-color: #0d6efd!important;
        border-radius: 12px!important;
        padding: 10px!important;
        margin-bottom: 8px!important;
    }
    h1, h2, h3, p, div, label, span {
        color: white!important;
    }
    </style>
    """, unsafe_allow_html=True)

    if st.session_state.user["role"] == "admin":
        if st.button("استخراج عضوية جديدة", use_container_width=True, type="primary"):
            st.session_state.page = "extract"
            st.rerun()
        if st.button("ادارة الاعضاء", use_container_width=True):
            st.session_state.page = "manage"
            st.rerun()

    if st.button("تغيير كلمة السر"):
        st.session_state.page = "change_pass"
        st.rerun()

    if st.button("تأكيد البريد الالكتروني"):
        st.session_state.page = "recovery"
        st.rerun()

    if st.button("تسجيل الخروج"):
        st.session_state.user = None
        st.session_state.role = None # <--- ضفت دي
        st.session_state.page = "login"
        st.rerun()
# ========================
# =========================== 

import streamlit as st
import pandas as pd
import json
import os
import io
import smtplib
import secrets
import base64
import arabic_reshaper # جديد
from bidi.algorithm import get_display # جديد
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from io import BytesIO # <--- ضفناها هنا عشان نستخدمها كلها

st.set_page_config(page_title="إدارة القضايا", layout="wide", page_icon="⚖️")

# دالة عشان تظبط العربي وتوصله
def fix_arabic(text):
    if not text:
        return ""
    text = str(text)
    reshaped_text = arabic_reshaper.reshape(text)
    bidi_text = get_display(reshaped_text)
    return bidi_text


# ====== خط عربي مدمج داخل التطبيق ======
ARABIC_FONT_B64 = "AAEAAAASAQAABAAgRFNJRwAAAAEAAq6oAAAACEdERUbQ/s73AAJThAAAAg5HUE9TcfUN6gACVZQAADsmR1NVQrdzeAAAApC8AAAd6k9TLzKnOVn8AAABqAAAAGBjbWFw1eCiVwAAGxAAAAlsY3Z0IAMfAx8AADMwAAAAKGZwZ21iLvt7AAAkfAAADgxnYXNwAAAAEAACU3wAAAAIZ2x5Zk9JAgoAAD/gAAHGXmhlYWQt+yLhAAABLAAAADZoaGVhGwMV5gAAAWQAAAAkaG10eIKXoZEAAAIIAAAZCGxvY2HdbXAGAAAzWAAADIZtYXhwCa8cMgAAAYgAAAAgbmFtZYrXr44AAgZAAAAFvHBvc3QzNDt7AAIL/AAAR4BwcmVwZUCunAAAMogAAACnAAEAAAACAQblVwTxXw889QAHA+gAAAAA3ABPpwAAAADcC4+T/wf9nhfMBCAAAAAGAAIAAAAAAAAAAQAABC39hgAAF/n/B/okF8wAAQAAAAAAAAAAAAAAAAAABkIAAQAABkICWAAQBBQAKwACAHQAxgCNAAACHxT3ABEABwAEAucBkAAFAAACigJYAAAASwKKAlgAAAFeADIA3AAAAgsFAgQFBAICBIAAIAOAACAAAAAACAAAAABHT09HAMAAAP7/BC39hgAABC0CegAAAEAAAAAAAAAAAAAAACAABQKGAFEBtQA2AH///wDuAEkA/QBMASUATAAAACIA7gAxAP0ALQAA//cBAABGAP0ATAAAACIA7gAMAP3/9wAA/8YBAAAPAP0ANgAAACIA7v/zAP3/7wAA/9cA7gAOAAD/4QD9AAoA7gBJAAD/7wD9AEwBUQBJASUATAAAACIDBABGAzEARgEk//gBE//4AwQARgMxAEYBJP/4ARP/+AMEAEYDMQBGAWj/+AFX//gDBABGAzEARgEk//gBE//4AwQARgMxAEYBaP/4AVf/+AEk//gBE//4AwQARgMxAEYBaP/4ARP/+AMEAEYDMQBGAWj/+AET//gDBABGAzEARgFo//gBV//4AwQARgMxAEYBaP/4ARP/+AMEAEYDMQBGAWj/+AET//gDBABGAzEARgFo//gBV//4AwQARgMxAEYBaP/4AVf/+AMEAEYDMQBGAej/+AHh//gDBABGAzEARgEk//gBE//4AwQARgMxAEYBaP/4AVf/+AMEAEYDMQBGAWj/+AFX//gDBABGAzEARgFo//gBV//4AwQARgMxAEYBaP/4AVf/+AMEAEYDMQBGAWj/+AFX//gDBABGAzEARgFo//gBV//5AnwARAKaAEQCmv/4Anz/+AJ8AEQCmgBEApr/+AJ8//gCfABEApoARAKa//gCfP/4AnwARAKaAEQCmv/4Anz/+AJ8AEQCmgBEApr/+AJ8//gCfABEApoARAKa//gCfP/4AnwARAKaAEQCmv/4Anz/+AJ8AEQCmgBEApr/+AJ8//gCfABEApoARAKa//gCfP/4AnwARAKaAEQCmv/4Anz/+AJ8AEQCmgBEApr/+AJ8//gCfABEApoARAKa//gCfP/4AnwARAKaAEQCmv/4Anz/+AJ8AEQCmgBEApr/+AJ8//gCfABEApoARAKa//gCfP/4AnwARAKaAEQCmv/4Anz/+AJ8AEQCfABEAnz/+AJ8//gBngAiAdoAIgGeACIB2gAiAZ4AIgHaACIBoQAlAdoAIgGeACIB2gAiAZ4AIgHaACIBngAiAdoAIgGeACIB2gAiAZ4AIgHaACIBngAiAdoAIgGeACIB2gAiAZ4AIgHaACIBngAiAdoAIgGhACUB2gAiAZ4AIgHaACIBgv/sAZT/7AGC/+wBlP/sAXT/7AGG/+wBdP/sAYb/7AF0/+wBhv/sAYL/7AGU/+wBdP/sAYb/7AGC/+wBlP/sAYL/7AGU/+wBgv/sAZT/7AGC/+wBlP/sAXT/7AGG/+wBpv/sAZT/7AGC/+wBlP/sAYL/7AGU/+wBdP/sAYb/7AF0/+wBlP/sAagABQG6AAUD9QBEA/kARAKX//gCkv/4A/UARAP5AEQCl//4ApL/+AP1AEQD+QBEApf/+AKS//gD9QBEA/kARAKX//gCkv/4A/UARAP5AEQCl//4ApL/+AP1AEQD+QBEApf/+AKS//gESABEBEsARAMK//gDEv/4BEgARARLAEQDCv/4AxL/+ARIAEQESwBEAwr/+AMS//gESABEBEsARAMK//gDEv/4BEgARARLAEQDCv/4AxL/+AMBACoDBwAqAnX/+AKU//gDAQAqAwcAKgJ1//gClP/4AwEAKgMHACoCdf/4ApT/+AMBACoDBwAqAnX/+AKU//gCXABJAd0ARAGe//gB+f/4AlwASQHdAEQBnv/4Afn/+AJcAEkB3QBEAZ7/+AH5//gCXABJAd0ARAGe//gB+f/4A1AARAMxAEQBg//4AaT/+ANQAEQDMQBEAYP/+AGd//gDUABEAzEARAGD//gBpP/4A1AARAMxAEQBg//4AaT/+ANQAEQDMQBEAYP/+AGk//gDUABEAzEARAGD//gBpP/4A1AARAMxAEQBg//4AZ3/+ANQAEQDMQBEAYP/+AGd//gChwBEAYP/+AGk//gCggBEAocARAKCAEQBg//4AaT/+AKHAEQCggBEAocARAKCAEQChwBEAoIARAGD//gBpP/4Ai4AEQLgAEQBy//4AZ//+ALSAEQDCABDAcv/+AGf//gC0gBEAwgAQwHL//UBn//4AtIARAMIAEMBy//4AZ//+ALSAEQDCABDAcv/+AGf//gC0gBEAwgAQwHL//gBn//4AtIARAMIAEMBy//4AZ//+ALSAEQDCABDAcv/1QGf/+cC0gBEAwgAQwHL/9UBn//pAtIARAMIAEMBy//VAZ//5wP3AAoEKQAJA6r/+AN3//gCrwBEAwgAQwHL//gBn//4Ai4AEQLgAEQBy//4AZ//+AIuABEC4ABEAcv/9QGf//gCLgARAuAARALSAEQDCABDAcv/1QGf/+cCLgARAuAARAHL//gBn//4AtIARAMIAEMBy//VAZ//5wLSAEQDCABDAcv/1QGf/+kCUwBEAk8ARAD1//gCYgA1ANT/9gIGAKQCUwBEAlMARAJPAEQA9f/2ANT/7AIGAKQCUwBEAlMARAJPAEQA9f/4AmIANQDU//YCBgCkAlMARAJTAEQCTwBEAPX/2gJiADUA1P/kAgYApAJTAEQCTwBEAPX/+AJiADUBYP/4AgYApAJTAEQCTwBEAPX/+ADU//YCBgCkAekARAIQAEQBnf/4Acj/+AHpAEQCSgBEAkkARAEk//gBE//4AkoARAJJAEQBJP/4ARP/+AJKAEQBJP/4ARP/+AJJAEQCSgBEAkkARAJKAEQCSQBEAST/+AET//gCSgBEAkkARAEk//gBE//4AZgARAHEAEQBff/4Afz/+AGYAEQBxABEAZgARAGpAEQBi//4ARP/+AGYAEQBqQBEAYv/+AET//gCdwARAnoAEQIB//gB/P/4AncAEQJ6ABECAf/4Afz/+AH8//gBmAAuAdMAUwGnAD0BqQBEAdQAIgHUACIB1AAiAdQAIgHUACIB1AAiAdQAIgHUACIB1AAiAdQAIgHU//8B1P//AdQAAAHUAAAB1AAUAdQAFAHUACIB1AAiAdQAIgHUACIB1AAiAdQAIgHUACIB1AAiAdQAIgHUACIB1AAiAdQAIgJqAEQCrwBEAmoARAKvAEQA7v/XAP3/zwAA/7kA7v/WAP3/zgAA/7kCagBEAq8ARAFo//gBV//4AmoAQwKvAEMBJP/4ARP/+AJqAEQCrwBEAWj/+AFX//gCagBDAq8AQwFo//gBV//4AmoARAJqAEQCrwBEAWj/+AFX//gCagBEAq8ARAFo//gBV//4AmoAAgKvAD8BaP/4ARP/+AJqABQCrwA9AWj/+AET//gCagBEAq8ARAFo//gBE//4AmoARAKvAEQBaP/4ARP/+AJqAEQCrwBEAWj/+AFH//gC8QBEAzYAQwJqAEQCrwBEAST/+AET//gCagBEAq8ARAFo//gBV//4AmoARAKvAEQDBQBEAjAAHwMFAEQCIQAQAx0AUgIfAA4BaP/4ARP/+AMdAFICHwAOAWj/+AET/+wCOAA/AdQAIgHUACIB1AAiAdQAIgLrAEQCrwBEAUz/+AFg//gBmABEA/UARAP5AEQCl//4ApL/+AJcAEkB3QBEAZ7/+AH5//gCXABJAd0ARAGmAAAB+f/4AlwASQHdAEQBpgAAAfn/+ANQAEQDMQBEAYP/+AGk//gDUABEAzEARAGD//gBpP/4AekARAIQAEQBnf/4Acj/+AHpAEQCEABEAZ3/+AHI//gCSgBEAkkARAFo//gBV//4AkoARAJJAEQBHP/yARP/+AJKAEQCSQBEAST/+AET//gCUwBEAk8ARAD1//gA1P/2AgYApAP1AEQD9QBEApf/+AKS//gD9QBEA/kARAKX//gCkv/4AnwARAKYAEQCmv/4Anz/+AP1AEQD+QBEApf/+AKS//gD9QBEA/kARAKX//gCkv/4BLYAHwDS//gBWQBEANL/+ADS//gA0v/4ANL/+ADS//gA0v/4ANL/+ADS//gA0v/4AgYAIwJiACICYgAiAmL/uQJi/7kCYgAiAgYAIwJiACICBgAjAmIAIgIG//cCYgAQAgYAIwJiACICBv/XAmL/3AIG/+ECYv/6A60ARAOt//gDrQBEA63/+AZtAEQDrQBEA63/+AZtAEQCuP/sArj/7AMjAEQDNABEAsH/+AKw//gDbQBEAqH/+AKQ//gDwgBEA9MARAPCAEQD0wBEA60ARAOt//gFSv/4BrEARAaxAEQDrQBEA63/+AacAEQGR//4BUr/+AOtAEQDrf/4BUr/+AaxAEQGsQBEAvz/7AL8/+wDIwBEA3gARAMF//gCsP/4BUr/+AVK//gFSv/4BbQARAW0AEQDsQBEAuX/+AKQ//gDwgBEBBcARAPCAEQEFwBEA60ARAL8/+wC/P/sAyMARAN4AEQDBf/4ArD/+AOxAEQC5f/4A8IARAQXAEQDwgBEBBcARAUWAEQFFv/4B8UARAfjAEQEjABEBBn/+AbRAEQGs//4BuYARAbmAEQFKwBEBUkARAUrAEQFSQBEBRYARAUW//gH4wBEBIwARAQZ//gG5gBEBuYARAUrAEQFSQBEBSsARAVJAEQCmv/4Apr/+AUWAEQFFv/4BRYARASMAEQEGf/4BSsARAVJAEQFKwBEBUkARAUsAEQFMf/4BSz/+AfG//gH4ABEBSwARAUx//gFLP/4B8b/+AUsAEQFMf/4BSz/+AfgAEQH3wBEBCb/7AQr/+wEogBEBDT/+AQv//gGyf/4Bs4ARAbJ//gGRABEBcz/+AQU//gED//4BUEARAVGAEQFQQBEBUYARAKS//gFLABEBTEARAUx//gFLP/4B+AARAUsAEQFMQBEBTH/+AUs//gHQQBEBsn/+AfgAEQFLABEBTEARAUx//gFLP/4BCb/7AQr/+wEogBEBKcARAQ0//gEL//4Bs4ARAbJ//gGRABEBcz/+AQU//gED//4BUEARAVGAEQFQQBEBUYARAWsAEQFrP/4CD4ARAhG//gIUwBEBaz/+ASm/+wEnv/sBiUAHwUiAEQEr//4BrcARAZM//gFwQBEBbkARAXBAEQFuQBEBawARAWs//gFrABEBaz/+AhTAEQIUwBEBawARAWs//gHtABEB0n/+ASm/+wEnv/sBSIARASv//gFwQBEBbkARAXBAEQFuQBEBS4ARAUu//gEpABEBBL/+AQx//gGrABEBsv/+AXO//gGwQBEBUMARAUkAEQFQwBEBSQARAJ1//gEpABEBBL/+AQx//gEkwBEBJP/+AZIAEQGMP/4BAkARAOW//gFSwBEBTP/+AXqAEQF6gBEBKgARARNAEQEqABEBE0ARASTAEQEk//4BAkARAOW//gFSwBEBeoARAXqAEQEqABEBE0ARASoAEQETQBEBD4ARAQ+//gEPgBEBD7/+AQ+AEQEPv/4Bi0ARAXb//gDtABEA0H/+AXPAEQEUwBEBDIARARTAEQEMgBEBD4ARAQ+//gDtABEA0H/+AW6AEQF2//4BTAARAXPAEQEUwBEBDIARARTAEQEMgBEApwATALIAEwEOQBEBDn/+AQ5AEQEOf/4BDkARAQ5//gD7gBEBBoARALA//gClP/4A68ARAPbAEQDaP/4Azz/+AV4AEQE2f/4BhcARAROAEQEegBEBE4ARAR6AEQDbgBEA27/+AYpAEQGCP/4BZ8ARAUL//gGPgBEA24ARANu//gFnwBEBQv/+AY+AEQGPgBEA24ARANu//gFnwBEBQv/+AO4AEQC5ABEAwUARAKS//gCcf/4BSwARAUL//gFQQBEAlH/+AODAEQDpABEA4MARAOkAEQCYgAiAmL/uQJi/7kCBgAjAmIAIgJi/7kCYv+5AmIAIgJi/7kCYv+5AmIAIgJi/7kCYv+5ApoATARiAEQEYv/4Bvz/+Ab8//gF///4BuYARARiAEQEYv/4Bvz/+AX///gG5gBEBGIARARi//gG/P/4Bf//+AbmAEQD2ABEA60ARANl//gF6QBEBHcARAR3AEQDrQBEA63/+AZYAEQGR//4Bc4ARAVK//gGbQBEBm0ARAOtAEQDrf/4BUr/+AZtAEQGbQBEA60ARAOt//gCuP/sArj/7AMjAEQDNABEAsH/+AKw//gFcABEBXAARANtAEQCof/4ApD/+APCAEQD0wBEA8IARAPTAEQElgBEBJb/+AQMAEQDmf/4BjP/+AU2//gEqwBEBKsARAPxAEQD8f/4BrEARAPxAEQD8f/4BrEARAPxAEQD8f/4Avz/7AL8/+wDZwBEA3gARAMF//gC9P/4BRUARASR//gFtABEA7EARALl//gC1P/4BAYARAQXAEQEBgBEBBcARAIQAEwCIQBMA60ARAOt//gDrQBEA63/+AOt//gCuP/sArj/7AMjAEQDNABEAsH/+AKw//gDbQBEAqH/+AKQ//gC5wAiAvgAIgLnACIC+AAiAucAIgL4ACIC5wAiAvgAIgPCAEQD0wBEA8IARAPTAEQDwgBEA9MARAJ7//gCvABEAs0ARAPCAEQD0wBEAjf/+ASmAEQEpgBEBUX/7AfZACIHtQBEB6UARAYaAEQHEwBEBp8ARATBADEESgAxAgYAIwJiACICBgAjAgYAIwJiACICYgA1AM0AFQFYACsCPAAxAjwAWQI8ADACPAAtAjwAFQI8AD8CPAA3AjwALAI8ADECPAAyAPwACwD8ABYA/AAOAPwACgD8AAYA/AARAPwADAD8ABAA/AAOAPwACgDeAAMA3gADAKEAAgDJAAIBOgBMAcEAmgE8AE0BwQCEAZ0ATQHeAE0BwQAmAaQARAF+AC8BwQBWAb0AGQG0AAkBwQARAbQACAHBABIBsQAqAToATAE8AE0BnQBNAd4ATQHSAE0BwQA1AcYALwHBAC8BvAA8AcEAPAG0AAkBtAAIAa0AKwDJABsAygAmAQkAJQEyACUBDQAgAPUAGAEdAAoBFwAFARcABAEVABYBKgAlASMAGAEcAB4B0QBSAaQAGwHBAEMBwQBTAcEAGwHBAEMAzgAuAQwAKQDOAC4AzgAuAQwASgEMAEwBDABIAQ0ASAFCACgBQgAoAZ4AOwGeAB0BngA/AZ4AHQFCADUBqQCBAakAVwGpAHAA8gA1APwAOgD8ADoBqQBJAakALgGpAC4BZAAuAakAXgGpAFkBqQBbAakAWwGpAFwBfgBOAX4ATgFCAGIBfgBOAX4ATgDlAEADmQAVAPMAKQDUACgA1AAsAakANwRyACIEcgAiAZAAEwBwACAAcAAgAHAAIABwACAAcAAgBNEAMwTRADMAMgAYADIAGAEsACkDVgBEAogAJgG1ADYAvQBHAL0ARwC9AEcAvQBHAL0ARwH0AAgB9AAqAN0AAAAAAAAAAP/uAAD/lQAA/+4AAP8uAN0AAANEAEQBpP/4A9oAZgJbADAB/gAVAekARAGRACICGv/2Ahr/9gIpACIC0wAiF/kALQd+ADYHfgA2AAD/vQAA/8kAwwAJAAD/9gAA//QAAP9+AAD/sAAA/z4AAP+4AAD/yQAA/8oAAP/JAakAUgAA/9YAAP/xAAD/8QAA//kAAP/yAAAAFwAA/+sAAP/iAAD/1gDDAAIAAP/JAAD/zAAA/6sAw//1AAAAAAAAABcAAP/yAMMABQAA//UAAP/qAAD/2gDDAAkAAP/kAAD/5gAA/9kAwwALAAD/9QAAABcAAP/4AAD/7wAA/+8AwwARAAD/4QAAAAEAAAAGAAAABQAAAAwAAAARAAAAEADDADIAAP+8AAAACgAA//0AAAAaAAAAHQAAAB0AAP+3AAAADAAAABgAAP8HAAD/UQAA/z4AAP/BAAD/mQAA/5QAAP+hAAD/YAAA/8oAAP/KAAD/rwAA/2AAAP+SAAD/egAA/6QAAP+/AAD/vwAA/7cAAP/BAAD/pAAA/+sAAP+kAAD/cAAA/7IAAP9/AAD/zgAA/6EAAP+5AAD/zgAA/6EAAP+5AAD/9QAA/+QAAP/1AAD/swAA/6kAAP+aAAD/mgAA/6wAAP+vAAD/rAAA/68AAP+AAAD/VwAA/4oAAP+xAAAAAAFCADUBqQBVAAD/5AAA//kA7AAdAakAgQGpAC4BqQBZAX4ATgDyADUBQgA1AakAWwGpAF4BqQBJAX4ATgD8ADoBqQBTAakAcAAA//4DUABEAO4ATQD9AE4AAAAIAVf/+AFo//gAAP/NAAD/tgGpAHMB4gAlASAALQEgADsAAP+2Agb/xgJi/+MCBgAjAmIAIgAA/zwAAP9JAZ3/+AET//gBwQAwAcEALwDuACkA7gAWAAD/+AAA/6wAAAAAAAAAAAEAABsAAAAhAPIANQHh//gB6P/4ASoANAENABIA4gAAAP0AAAEaAAABKgAAAAABywIa//YAAAC8BHIAIgRyACIE0QAzBNEAMwAyABgAMgAYB34ANgd+ADYHfgA2AAAAIgJiADUAAP/cAAD/3wAAACIAAP/6AAAAEAAAACIAwwAAAdQAFAHUABQCOAA/AAD/0AAAAE4AAP+aAAUAAAGkAK4A7AAdASAAHQBwACAAAP/qASX/4gEl//gAAP+2AakAcAFUADgBGAAaAQAAGwEAABsBAAAbAAAAIgAA/7kAAP+5AmIANQpZAAAJDAAADYMAAAAAAAIAAAADAAAAFAADAAEAAAAUAAQJWAAAAEIAQAAFAAIAAAANACEALAAuADkAOgCgAKsAuwYEBhsG/wd/CKAIrAj+IA8gESBPJcwuQfvB/T/9j/3H/fL9+/39/nT+/P7///8AAAAAAA0AIAAsAC4AMAA6AKAAqwC7BgAGBgYeB1AIoAiiCOQgCyAQIE8lzC5B+1D70/1Q/ZL98P3z/fz+cP52/v///wXVBU0AAATpBOYEnATcBLQEcwRkAAAAAAAAAAD3zQAA/NblSuUM5MnfktbYAAAAAAAAAAAAAAbIAAAAAAAABtYAAQAAAAAAPgAAAAAAAAAAAAAAAAAAADIAOgBkAiYAAAKCAAAAAAAAAAAAAAAAAooDbAZEBsIHLAAABy4HMAc4AAAAAAVUBRcFPQVABUUFRwVoBWIFYwVbBWQFZQVcBToFSQVKBUsFbAVtBWkFagVuBW8FcAVxBXIFcwV0BTsFLAU8Am0AAQATAAcCIAAKAkIAAwAjAhQANQBBAHEAiQChALUAtwDTANUA9wEDAQ8BGwEjASsBMwE7AXsBfwJPAlcCWwLWAUMBZwFzAb0B4gHnAf0CGAI0Aj4FfAWABYQFiAWMBZAFkwWcBZ4FeQV7BXgFoAWfBaEFogWjBaQFpQWmBXYE5ATmBOgE6QTrBOwE7gTvBPEE8wVhBOAE4gU/AB8BYwV3ABYADQAQAAIAHAKEAoYCiABFAEkAKwA5AD0AJwBNAC8AjQCRAIEAhQCZAHUAeQC7AMEAwwDFAL0AvwDHAMkAywDXANkA2wDdAN8A4QDjAOUA5wD7AP8BCwETARcBJwE3AUsBTwFTAUcBVwFfAWsBbQF3AZsBnwGjAbEBqwGPAZMBtQGtAbkBlwHDAckB0AHXAesB7wHzAfUB+QILAH0CAQIDAgcCFgIaAigCKgIsAi4CMAIcAjICUwJrAkYCHgJxAnUCdwJ5BTkCjAWnBagFqQWrBawFrQWuBU0FOAWvBbAFsQWqBbIFswTKBMsFtAW1BV8FtgW3BbgFuQDNAOkE9AT1BPYE9wT4BPoE/AT+BP8FAAEHAR8BPwVMBWACDwBRAFUAWQBdAGEAZQBpAJUAnQDPANEA6wKNApEClQKZAp0CoQGDAYcBiwKlAqkCrQKxArUCuQDtAO8CvgClAKkCwgDxAsYCOAI7Al8CYwJnAiICJAJ7An8ArQLKAs4BpwCxAS8BWwFvAd0B5gJKAk4A9QImAoMAFgAYACsALAAuAC0AJwAoACoAKQAvADAAMgAxAEkASgBMAEsATQBOAFAATwBFAEYASABHAUcBSAFKAUkBXwFgAWIBYQCFAIYAiACHAIEAggCEAIMAdQB2AHgAdwB5AHoAfAB7AL8AwAC9AL4AxwDIALsAvADlAOYA1wDYAXcBeAF6AXkBjwGQAZIBkQGtAboBvAG7AbUBtgG4AbcB7wHyAfMB9ABIAEcCAQICAgMCBAIGAgUCCwIMAg4CDQJ3AngCeQJ6BSMFJQUpBSoFLwUxBS0FLgUyBTMFNAUmBScFNwXWBSIBsQGyAbQBswIsAi0CKgIrAi4CLwKGAjICMwIoAikCMAIxAnECcgJ0AnMANAAzBJUElgS0BLUEpQSmBKkEqgSnBKgEqwSsBLEEsgSzBLYEtwS4AlMCVAJWAlUElwSZBJ4ErQSvAvAC8gL1AvoDAQMDAwUDCgMPAxYDIgMkAyYDKQMvAzEDMwM3A0EDRANOA1ADUQNXA1wDYANnA5YDnwOnA6kDrQOzA7kDuwPHA8oDzgPYA9oD4wPlA+cD6wPuA/AD8gP0A/oD/AP+BAAEAgQEBAYECgQRBBMEFQQcBCIEJwQvBDEEQQRHBEwEUQRVBFYEVwRfBGQEaARxBHMEdQR3BHsEfAR9BIAEgwSHBJEEkwC5APMCNgWWBZcFmAWZBZoFmwScBJ0EnwSiBK4EsAL4AvkC+wL+AwIDBAMUAxUDFwMfAyMDJQMnAygDKgMtAzADMgPvA/ED+wP9A/8EBwQLBBIEFAQoBDAEMgRABFIEZgRnBGkEbgRyBHQCNwSFBIYEiASOBJIElASYBJoEmwShBKQC8QLzAvYC/QMAAwYDCwMQAxkDIQMsAzQDOANCA0UDTwNSA1kDXgNiA2kDlwObA6ADqAOqA64DtAO6A8kDywPPA9kD2wPkA+YD6APsA/MD9QQBBAMEBQQJBA0EFgQdBCMEKgQuBEIESARNBFMEWARgBGUEawRwBHYEeAITBH4EgQSEBIoEkASgBKMC/AL/AxgDIAMrAy4DaANvA4oDkAQIBAwEKQRqBG8EiQSPAtoC2wLcA8IDxAPUA9YD3wPhA3EDcwOSA5QDSANKAz0DPwNTA1UDowOlA7UDtwN2A3sDggOIA4YDZQOcA7EDwwPFA9UD1wPgA+IDcgN0A5MDlQNJA0sDPgNAA1QDVgOkA6YDtgO4A3cDfAODA4kDhwNmA50DsgN5A34DhQOLA3ADkQO9A1gDXQNhA3gDfQOEA7wDyAAbABkFUgVTAwcDDAMNAw4DEQMaAxsDHAM5AzoDRwNGA18DWgNbA2sDbANqA20DbgOYA5kDoQN/A4ADegOMA40DjgOPA6sDrwOwA74DvwPAA8EDzAPQA9ED0gPcA94D3QPpA+oD9gP4BB4EIQQgBBgEFwQkBCUEKwQsBEkESgRLBEMERQROBE8ERAR5BHoEYQRiBFsEXARdBG0EbASLBIwC9wMJAwgDEwMSAx4DHQM8AzUDOwNjA5oDgQOsBBsELQSCBH8EjQRUA/kEYwP3BB8D0wQQBFoEUAQaBA4EGQRZAzYDQwRGA+0C9AQPA80DogNkBF4DngLSBLkFXQVmBX8C0wWDAtQFhwWLAtcFjwLYBWsC1QWVAtkFnQLdAAEAEwAUAAcACAIgAiEACgALAkICQwJFAkQAAwAEACMAJAAmACUCFAIVADUANgA4ADcAQQBCAEQAQwBxAHIAdABzAIkAigCMAIsAoQCiAKQAowC1ALYAtwC4ANMA1ADVANYA9wD4APoA+QEDAQQBBgEFAQ8BEAESAREBGwEcAR4BHQEjASQBJgElASsBLAEuAS0BMwE0ATYBNQE7ATwBPgE9AUMBRAFGAUUBZwFoAWoBaQFzAXQBdgF1Ab0BvgHBAb8B4gHjAeUB5AHnAegB6gHpAf0B/gIAAf8CGAIZAjQCNQI+Aj8CQQJAAuwC7QLoAukC6gLrAt4C37AALCCwAFVYRVkgIEu4AA5RS7AGU1pYsDQbsChZYGYgilVYsAIlYbkIAAgAY2MjYhshIbAAWbAAQyNEsgABAENgQi2wASywIGBmLbACLCMhIyEtsAMsIGSzAxQVAEJDsBNDIGBgQrECFENCsSUDQ7ACQ1R4ILAMI7ACQ0NhZLAEUHiyAgICQ2BCsCFlHCGwAkNDsg4VAUIcILACQyNCshMBE0NgQiOwAFBYZVmyFgECQ2BCLbAELLADK7AVQ1gjISMhsBZDQyOwAFBYZVkbIGQgsMBQsAQmWrIoAQ1DRWNFsAZFWCGwAyVZUltYISMhG4pYILBQUFghsEBZGyCwOFBYIbA4WVkgsQENQ0VjRWFksChQWCGxAQ1DRWNFILAwUFghsDBZGyCwwFBYIGYgiophILAKUFhgGyCwIFBYIbAKYBsgsDZQWCGwNmAbYFlZWRuwAiWwDENjsABSWLAAS7AKUFghsAxDG0uwHlBYIbAeS2G4EABjsAxDY7gFAGJZWWRhWbABK1lZI7AAUFhlWVkgZLAWQyNCWS2wBSwgRSCwBCVhZCCwB0NQWLAHI0KwCCNCGyEhWbABYC2wBiwjISMhsAMrIGSxB2JCILAII0KwBkVYG7EBDUNFY7EBDUOwAWBFY7AFKiEgsAhDIIogirABK7EwBSWwBCZRWGBQG2FSWVgjWSFZILBAU1iwASsbIbBAWSOwAFBYZVktsAcssAlDK7IAAgBDYEItsAgssAkjQiMgsAAjQmGwAmJmsAFjsAFgsAcqLbAJLCAgRSCwDkNjuAQAYiCwAFBYsEBgWWawAWNgRLABYC2wCiyyCQ4AQ0VCKiGyAAEAQ2BCLbALLLAAQyNEsgABAENgQi2wDCwgIEUgsAErI7AAQ7AEJWAgRYojYSBkILAgUFghsAAbsDBQWLAgG7BAWVkjsABQWGVZsAMlI2FERLABYC2wDSwgIEUgsAErI7AAQ7AEJWAgRYojYSBksCRQWLAAG7BAWSOwAFBYZVmwAyUjYUREsAFgLbAOLCCwACNCsw0MAANFUFghGyMhWSohLbAPLLECAkWwZGFELbAQLLABYCAgsA9DSrAAUFggsA8jQlmwEENKsABSWCCwECNCWS2wESwgsBBiZrABYyC4BABjiiNhsBFDYCCKYCCwESNCIy2wEixLVFixBGREWSSwDWUjeC2wEyxLUVhLU1ixBGREWRshWSSwE2UjeC2wFCyxABJDVVixEhJDsAFhQrARK1mwAEOwAiVCsQ8CJUKxEAIlQrABFiMgsAMlUFixAQBDYLAEJUKKiiCKI2GwECohI7ABYSCKI2GwECohG7EBAENgsAIlQrACJWGwECohWbAPQ0ewEENHYLACYiCwAFBYsEBgWWawAWMgsA5DY7gEAGIgsABQWLBAYFlmsAFjYLEAABMjRLABQ7AAPrIBAQFDYEItsBUsALEAAkVUWLASI0IgRbAOI0KwDSOwAWBCIGC3GBgBABEAEwBCQkKKYCCwFCNCsAFhsRQIK7CLKxsiWS2wFiyxABUrLbAXLLEBFSstsBgssQIVKy2wGSyxAxUrLbAaLLEEFSstsBsssQUVKy2wHCyxBhUrLbAdLLEHFSstsB4ssQgVKy2wHyyxCRUrLbArLCMgsBBiZrABY7AGYEtUWCMgLrABXRshIVktsCwsIyCwEGJmsAFjsBZgS1RYIyAusAFxGyEhWS2wLSwjILAQYmawAWOwJmBLVFgjIC6wAXIbISFZLbAgLACwDyuxAAJFVFiwEiNCIEWwDiNCsA0jsAFgQiBgsAFhtRgYAQARAEJCimCxFAgrsIsrGyJZLbAhLLEAICstsCIssQEgKy2wIyyxAiArLbAkLLEDICstsCUssQQgKy2wJiyxBSArLbAnLLEGICstsCgssQcgKy2wKSyxCCArLbAqLLEJICstsC4sIDywAWAtsC8sIGCwGGAgQyOwAWBDsAIlYbABYLAuKiEtsDAssC8rsC8qLbAxLCAgRyAgsA5DY7gEAGIgsABQWLBAYFlmsAFjYCNhOCMgilVYIEcgILAOQ2O4BABiILAAUFiwQGBZZrABY2AjYTgbIVktsDIsALEAAkVUWLEOBkVCsAEWsDEqsQUBFUVYMFkbIlktsDMsALAPK7EAAkVUWLEOBkVCsAEWsDEqsQUBFUVYMFkbIlktsDQsIDWwAWAtsDUsALEOBkVCsAFFY7gEAGIgsABQWLBAYFlmsAFjsAErsA5DY7gEAGIgsABQWLBAYFlmsAFjsAErsAAWtAAAAAAARD4jOLE0ARUqIS2wNiwgPCBHILAOQ2O4BABiILAAUFiwQGBZZrABY2CwAENhOC2wNywuFzwtsDgsIDwgRyCwDkNjuAQAYiCwAFBYsEBgWWawAWNgsABDYbABQ2M4LbA5LLECABYlIC4gR7AAI0KwAiVJiopHI0cjYSBYYhshWbABI0KyOAEBFRQqLbA6LLAAFrAXI0KwBCWwBCVHI0cjYbEMAEKwC0MrZYouIyAgPIo4LbA7LLAAFrAXI0KwBCWwBCUgLkcjRyNhILAGI0KxDABCsAtDKyCwYFBYILBAUVizBCAFIBuzBCYFGllCQiMgsApDIIojRyNHI2EjRmCwBkOwAmIgsABQWLBAYFlmsAFjYCCwASsgiophILAEQ2BkI7AFQ2FkUFiwBENhG7AFQ2BZsAMlsAJiILAAUFiwQGBZZrABY2EjICCwBCYjRmE4GyOwCkNGsAIlsApDRyNHI2FgILAGQ7ACYiCwAFBYsEBgWWawAWNgIyCwASsjsAZDYLABK7AFJWGwBSWwAmIgsABQWLBAYFlmsAFjsAQmYSCwBCVgZCOwAyVgZFBYIRsjIVkjICCwBCYjRmE4WS2wPCywABawFyNCICAgsAUmIC5HI0cjYSM8OC2wPSywABawFyNCILAKI0IgICBGI0ewASsjYTgtsD4ssAAWsBcjQrADJbACJUcjRyNhsABUWC4gPCMhG7ACJbACJUcjRyNhILAFJbAEJUcjRyNhsAYlsAUlSbACJWG5CAAIAGNjIyBYYhshWWO4BABiILAAUFiwQGBZZrABY2AjLiMgIDyKOCMhWS2wPyywABawFyNCILAKQyAuRyNHI2EgYLAgYGawAmIgsABQWLBAYFlmsAFjIyAgPIo4LbBALCMgLkawAiVGsBdDWFAbUllYIDxZLrEwARQrLbBBLCMgLkawAiVGsBdDWFIbUFlYIDxZLrEwARQrLbBCLCMgLkawAiVGsBdDWFAbUllYIDxZIyAuRrACJUawF0NYUhtQWVggPFkusTABFCstsEMssDorIyAuRrACJUawF0NYUBtSWVggPFkusTABFCstsEQssDsriiAgPLAGI0KKOCMgLkawAiVGsBdDWFAbUllYIDxZLrEwARQrsAZDLrAwKy2wRSywABawBCWwBCYgICBGI0dhsAwjQi5HI0cjYbALQysjIDwgLiM4sTABFCstsEYssQoEJUKwABawBCWwBCUgLkcjRyNhILAGI0KxDABCsAtDKyCwYFBYILBAUVizBCAFIBuzBCYFGllCQiMgR7AGQ7ACYiCwAFBYsEBgWWawAWNgILABKyCKimEgsARDYGQjsAVDYWRQWLAEQ2EbsAVDYFmwAyWwAmIgsABQWLBAYFlmsAFjYbACJUZhOCMgPCM4GyEgIEYjR7ABKyNhOCFZsTABFCstsEcssQA6Ky6xMAEUKy2wSCyxADsrISMgIDywBiNCIzixMAEUK7AGQy6wMCstsEkssAAVIEewACNCsgABARUUEy6wNiotsEossAAVIEewACNCsgABARUUEy6wNiotsEsssQABFBOwNyotsEwssDkqLbBNLLAAFkUjIC4gRoojYTixMAEUKy2wTiywCiNCsE0rLbBPLLIAAEYrLbBQLLIAAUYrLbBRLLIBAEYrLbBSLLIBAUYrLbBTLLIAAEcrLbBULLIAAUcrLbBVLLIBAEcrLbBWLLIBAUcrLbBXLLMAAABDKy2wWCyzAAEAQystsFksswEAAEMrLbBaLLMBAQBDKy2wWyyzAAABQystsFwsswABAUMrLbBdLLMBAAFDKy2wXiyzAQEBQystsF8ssgAARSstsGAssgABRSstsGEssgEARSstsGIssgEBRSstsGMssgAASCstsGQssgABSCstsGUssgEASCstsGYssgEBSCstsGcsswAAAEQrLbBoLLMAAQBEKy2waSyzAQAARCstsGosswEBAEQrLbBrLLMAAAFEKy2wbCyzAAEBRCstsG0sswEAAUQrLbBuLLMBAQFEKy2wbyyxADwrLrEwARQrLbBwLLEAPCuwQCstsHEssQA8K7BBKy2wciywABaxADwrsEIrLbBzLLEBPCuwQCstsHQssQE8K7BBKy2wdSywABaxATwrsEIrLbB2LLEAPSsusTABFCstsHcssQA9K7BAKy2weCyxAD0rsEErLbB5LLEAPSuwQistsHossQE9K7BAKy2weyyxAT0rsEErLbB8LLEBPSuwQistsH0ssQA+Ky6xMAEUKy2wfiyxAD4rsEArLbB/LLEAPiuwQSstsIAssQA+K7BCKy2wgSyxAT4rsEArLbCCLLEBPiuwQSstsIMssQE+K7BCKy2whCyxAD8rLrEwARQrLbCFLLEAPyuwQCstsIYssQA/K7BBKy2whyyxAD8rsEIrLbCILLEBPyuwQCstsIkssQE/K7BBKy2wiiyxAT8rsEIrLbCLLLILAANFUFiwBhuyBAIDRVgjIRshWVlCK7AIZbADJFB4sQUBFUVYMFktAEu4AMhSWLEBAY5ZsAG5CAAIAGNwsQAHQrIAAQAqsQAHQrMMBQEKKrEAB0KzEQMBCiqxAAhCugNAAAEACyqxAAlCugBAAAEACyq5AAMAAESxJAGIUViwQIhYuQADAGREsSgBiFFYuAgAiFi5AAMAAERZG7EnAYhRWLoIgAABBECIY1RYuQADAABEWVlZWVmzDgMBDiq4Af+FsASNsQIARLMFZAYAREQAAAAAAAAAAAAAAAAAAAAAAAAAAAAALQAtAFMAUwKfAAAAAAKfAAAAAAAAACoAlACiAMQA+gE0AZABogG0AcQB1gHoAfoCDAIeAi4CQAJSAmQCdgKIApgCqgK6AswC2ALqAvYDCAMaAywDdAPiBDwEeASKBJwErgTABNIE5AT0BQQFFgUoBTgFSAVaBWwFfAWMBZQFnAWuBcAF0gXiBfwGFgYwBkoGXAZuBoAGkAaiBrQGxgbWBugG+gcMBxwHLgdAB1AHYgd0B4YHlgemB7gHygfcB+4ICAgiCDoIUghkCHYIiAiaCLQIzgjmCP4JGAkyCUoJYgl0CYYJlgmmCbgJygncCewJ/goQCiAKMApCClQKZgp4CooKmgqsCr4K0AriCvQLBgsgCzgLUAtoC3oLjAueC7ALwgvUC+YL+AxkDPANag3QDeQN+A4MDiAOMg5EDlYOaA54DogOmA6oDrgOyA7YDugO+g8MDx4PMA9CD1QPZA90D4YPmA+qD7wPzg/gD/IQBBAWECgQOhBMEGQQfBCUEKwQ6hEuEUARUhFsEYYRlhGmEbYRxhHWEeYR+BIKEhwSLhJGEl4SbhJ+Eo4SnhKwEsIS0hLiEvoTEhMiEzITehPYE+gT+BQIFBgUKBQ4FEoUXBRuFIAUkhSkFL4U2BToFPgVCBUYFSgVOBVIFVgVahXSFeIV8hYEFhYWJhY2FkgWWhboF4oYKBj8GcQaThpoGoIanBq2Gsga2hrqGvobDBseGzAbQhtcG3YbkBuqG8Qb3hv4HBIclB0qHbQeKh48Hk4eYB5yHoQelh6oHroezB7eHvAfAh8cHzYfUB9qH94gciDwIV4hcCGCIZQhpiG4Icoh3CHuIgAiEiIkIjYiqCM8I6gkBiQYJCgkOCRIJFokaiR6JIokpCS8JNQk7CT+JRAlICUyJUQlViVmJXYlfiX6JgImCiYcJi4mQCZSJmwmhiaeJrgmyibcJuwm/CcWJzAnSCdgJ3InhCeUJ6YoICiGKOgpZCl2KYgpmCmqKbwpzingKfIqDComKj4qWCsCK8IsIix+LOQtWi1iLWotfC2OLaAtsi3ELdYt5i32LgguGi4sLj4uUC5iLnQuhi6YLqouvC7OLuAu8i8ELxYvMC9KL2Qvfi+YL7IvzC/mMFYxUDIQMn4ykDKiMrQyxjLYMuoy/DMOMyAzMjNEM1YzZjN4M5IzrDPGM+Az8jQENBY0KDRCNFw0djSQNKo0xDTeNPg1SjWyNfw2aDaoNtw27jcANxI3JDc2N0g3WjdsN343kDeiN7Q3xjfYN+o3/DgOOCA4MjhEOFY4aDh4OIo4/DkOORo5LDk+OVA5XDmyOiw6jjrmOvY7CDsaOyw7PDtWO3A7iDugO+g78Dv4PFA8Yjx0PI48qDzCPNw87j0APRA9ID1qPcw+Vj6+PtA+5D7sPyo/fD/SP+Q/+EAMQCBAhkFYQiRCLEI+QlBCYkJ0QoZCmEKqQrxCzkM8Q7JDxEPWQ+ZD9kQGRBZEKkQ+RFBEYkR0RIZEkkSeRSpFwEXQReBF8kYERhZGKEY6RkxGXEZsRuJHRkdYR2pHfEeOR6BHskfER9ZH6Ef6SAxIHkgySEZIWkhuSIBIkkiqSMJI3kj6SRRJLklISVxJcEmISaBJqEmwScJJ1EnmSfhKEEooSjpKTEpkSnxKjkqgSrpK1ErmSvhLEkssSz5LUEtiS3RL/kxyTIRMlkyoTLpMzEzeTO5M/k0QTSJNeE3ITdxN8E4CThROLk5ITlpObE6GTqBOsk7ETtZO8E8KTx5PMk9IT1pPYk90T4ZPmE+qT7xPzE/cT+xP/lAOUB5QLlBAUFJQZFB2UIhQmlCsUL5Q0FDiUPRRBlEWUSZRNlFGUVhRalF8UY5RqFHCUdpR8lIEUhZSKFI4UkpSXFJuUn5SkFKiUrRSxlLYUupS/FMOUyBTMlNEU1ZTaFN6U4xTnlOwU8JT1FPmU/hUClQcVC5UQFRaVGZUoFSyVNZU4lTuVPpVClUaVSpVNlWyVihWOlZMVl5WcFaCVpRWpla4VshW2FbqVvxXDFccVyxXPFdIV1RXYFdsV3xXiFeUV6RXsFe8V8hX1FfgV+xX+FgEWBBYHFgoWDRYQFhMWFhYaFh4WIhYlFigWLBYwFjQWNxY6Fj4WQhZGFkkWTBZPFlIWVRZYFlwWYBZkFmgWbBZvFnIWdRZ4FnsWfhaBFoQWhxaKFo0WkBaTFpYWmRacFp8WohalFqgWqxauFrIWtha5FrwWwBbEFsgWzBbPFtIW1RbYFtsW3hbiFuUW6BbsFvAW8xb2FvkW/BcAlwaXCZcMlw+XEpcVlxiXG5celyGXJJcnlyqXLpcylzWXOJc7lz+XQpdFl0iXTJdQl1OXVpdZl1yXX5djl2eXa5dvl3OXdpd5l3yXf5eCl4WXjBePF5IXlReYF5wXnxeiF6UXqBesF7AXtBe3F7oXvRfAF8MXxhfJF8wXzxfSF9YX2hfeF+IX5RfoF+sX7hfxF/QX9xf6F/4YAhgGGAkYDBgPGBMYFhgZGB0YIRgkGCcYKhgtGDAYMxg2GDkYPRhBGEQYRxhLGE8YUhhVGFgYWxheGGEYZBhnGGoYbRhwGHMYdhh6GH4YghiGGIkYjBiPGJIYlpiZmJyYn5iimKWYqZitmLCYs5i3mLuYv5jDmMaYyZjMmM+Y0pjVmNiY25jfmOOY55jqmO2Y8JjzmPaY+Zj8mP+ZApkFmQmZDZkQmROZF5kamR2ZIJkjmSaZKZksmS+ZM5k3mTuZP5lCmUWZSJlLmU6ZUZlUmVeZWpldmWCZY5lmmWmZbJlvmXKZdZl4mXuZf5mDmYeZipmNmZCZk5mWmZmZnZmhmaWZqZmtmbCZs5m3mbuZv5nDmcaZyZnNmdGZ9hn5GfwZ/xoCGgYaChoOGhEaFBoXGhoaHRojmioaMJo1GjuaQhpImk8aVZpcGmKaaRpvmnKadZp4mnyagJqEmoiai5qOmpKalpqamp2aoJqkmqiarJqvmrKatZq5mryav5rCmsWayZrNmtGa1ZrZmt2a4Jrjmuea65rvmvKa9Zr4mvua/psBmwSbB5sLmw+bEpsVmxibG5semyGbJJsnmyqbLZswmzSbOJs7mz6bQZtEm0ibS5tOm1KbVZtYm1ubXpthm2SbZ5tqm26bcpt2m3mbfJt/m4KbhZuIm4ubjpuRm5Sbl5uam52boJujm6abqZusm6+bspu1m7ibu5u+m8GbxJvHm8qbzZvQm9Ob1pvZm9yb35vim+Wb6Jvrm+6b8Zv0m/eb/Jv/nAScEhwXHBwcIRwmHCofm6HModEh1aHaId6h4yHnogIiE6Ikoi+iQqJaomsifyKYIqIiu6LVIuUi8SMCoxijJyM8I1KjXCNzo4mjnaOho64jtKO8o76jx6PJo9qj9SP3JA4kIaQjpDOkP6RBpE0kTyRiJGQkZiRoJGokgySFJJ0knyS0pLakuKS6pLykxCTNJN8lAaUYpSqlOaVHpVYlaaWFpZwlsSXHpdkl6qYAJhCmJKYrpjSmQSZNpl0mZqZyJoMmiaaQJqEmsKbCptSm7SbvJvMm9Sb3Jvkm+yb9Jv8nAScGJwgnCicMJw4nEicUJxYnIacjpyWnOydoJ3AneyeLJ6Knt6e5p8Qn/af/qAGoA6gFqCQoJihLKE0oVKhtqJOomCkDqY8pkSmTKZUpuynfKd8p3ynlKfCp/SoJqgmqLqo1KneqvqrqKu6rAysHqwqrJKtFK2orbCurK8ir7yvxLAasICw5rFAscayFrIysp6yuLLGs3izlrOys7q0SrSQtLq1JLWgta62XrdOuFi4ZriQuQK5fLmEuZ65+Lpcumq6wLtIu+y7+rwUvHa84L04vaC9rr22vcS90r3avei98L5Gvk6+hL66vx6/Sr+Ev4y/tsAWwFTA/MG4whjCcMLgw1bDiMQSxFLEjMTcxWzFosXsxkLGWsZyxpzG9Mc8x7THxMfUx+jIAMgcyCzIlMikyLTIxMjUyODI8skCyXTJgMmYydDKDsoeyi7KQMpYymrKfMp8yoTKlsroyzDLZst+y6rL5sw2zGTM4s0ezV7NjM3czgrOSs5izqDPGM88z2zPts/00E7QotEA0SjRaNGe0dbSPtJO0l7ScNKC0qzS1NM4027TdtPA1ADUZNSI1K7U6NUm1bDV/tYc1lrWsNcA10jXZNeO16rX1NhA2GTYvtjG2M7Y1tje2ObY7tj22P7ZBtk62UzZXNls2X7Zjtme2bDZ2tpg2uTbXNuG257bttu229bcCNwY3CDcdNzY3OjdRt1O3azeCt6U3x7fJt8430rfXN9u4RbhMuMvAAAAAgBR/vECNAQgAAMABwAqQCcAAAADAgADZwACAQECVwACAgFfBAEBAgFPAAAHBgUEAAMAAxEFBhcrExEhESUhESFRAeP+ZQFT/q3+8QUv+tFIBJ8AAQA2/9sBowEVADkAIkAfJSAaGBUEAQcBSQAAAQEAWQAAAAFhAAEAAVErLwIHGCsXJzY2NyYnJiY1NDY2NzY2MzIWFRQHJiYnNjcmJiMiBgcWFxYWFzY2NzY2NzY2NxYVFAYHBgYHDgJEDhMrHg0GDREbLh0LFgwjLAwKDAgHAQYaCR0/FREgCyEOBg4HGDMWDBMHDxwTJE4bEik6JSoIEAoIBg4gERw8MwwFBSUlGBUDBQQKDwYGJBwmEgYMBAIEAQcJBAICARATFhgCBAgFAwgMAP////8CGAB/AswBBgV5BiIACLEAAbAisDUrAAAAAQBJAAAAsgKfAA4ABrMGAAEyKzMmAicmNjcWFhcHFhYGB2gDDQsEHycEEwwbBAMDBbABI3MpKQcTQSMXSpy0bwAAAAABAEwAAAEOAp8AFgAhQB4VFAIAAQFMEAsIAwFKAAEBAGEAAAAPAE4SERACBxcrISImJy4DJzY2Nx4CFxcWMzIVFRQBBjtOBQQJCg0IDiIUAwQDAQMvOQgvKCRZeqdyDh0NTpipZ0MaCDwIAAABAEwAAAE2Ap8AGQAhQB4YFwIAAQFMEQwJAwFKAAEBAGEAAAAPAE4VFBACBxcrISImJicuAyc2NjceAhcXHgIzMhUVFAEuJ1I6AwQJCg0IDiIUAwMDAwIQNTgTCBUnGyRZeqdyDh0NTpipZ0MIDAYIPAgAAAEAIgAAAdICHwAvACVAIiMMAgABAUwZAQFKAAEAAYUAAAACYQACAg8CTi8qGFUDBxgrMyY1NDY2MzI2NzY2NyYmJyYmJwcmJjU0NjcWFxYWFx4DFRQGBgcOAgciBgYiiQMFCAIdQiQkSyQUQSkoXi8lEh8NFixALV0iDicmGgoOBg8+QhcJJSwjAwgKHhkCAgIIBiZVKSpLGwIRLBYRHgseNyZZKRA7QjYLByAbAQIHBQEBAQAAAP//ADEAAACyA3YCJgADAAABBwV5ADgAzAAIsQEBsMywNSsAAP//AC0AAAEOA3YCJgAEAAABBwV5ADQAzAAIsQEBsMywNSsAAP////cAAAHSAvYCJgAGAAABBgV5/kwACLEBAbBMsDUr//8ARv8UAMYCnwAmAAMAAAEGBXsvqQAJsQEBuP+psDUrAAAA//8ATP8OAQ4CnwImAAQAAAEGBXtSowAJsQEBuP+jsDUrAAAA//8AIv8KAdICHwImAAYAAAEGBXtYnwAJsQEBuP+fsDUrAAAA//8ADAAAAM8DfAImAAMAAAEHBfUAVgDWAAixAQGw1rA1KwAA////9wAAAQ4DgQImAAQAAAEHBfUAQQDbAAixAQGw27A1KwAA////xgAAAdIDAQImAAYAAAEGBfUQWwAIsQEBsFuwNSv//wAP/xIA0gKfACYAAwAAAQYFdjmjAAmxAQG4/6OwNSsAAAD//wA2/xIBDgKfAiYABAAAAQYFdmCjAAmxAQG4/6OwNSsAAAD//wAi/xYB0gIfAiYABgAAAQYFdl2nAAmxAQG4/6ewNSsAAAD////zAAAA8gMeAiYAAwAAAQcFngA3AKgACLEBAbCosDUrAAD////vAAABDgMeAiYABAAAAQcFngAzAKgACLEBAbCosDUrAAD////XAAAB0gKeAiYABgAAAQYFnhsoAAixAQGwKLA1K///AA4AAADTA0YCJgADAAABBwXvAEEAuwAIsQECsLuwNSsAAP///+EAAAHSAskCJgAGAAABBgXvFD4ACLEBArA+sDUr//8ACgAAAQ4DRgImAAQAAAEHBe8APQC7AAixAQKwu7A1KwAA//8ASQAAAZkCnwImAAMAAAAHBXwA7gAA////7wBkAU0C8AImBh0AAAEHBXwABACCAAixAQKwgrA1KwAA//8ATAAAAYACnwImAAQAAAAHBXwA1QAA//8ASQAAAU8CzAAmAAMAAAEHBXkA1gAiAAixAQGwIrA1KwAA//8ATAAAATcCzAImAAUAAAEHBXkAvgAiAAixAQGwIrA1KwAA//8AIgAAAdIC4gImAAYAAAEHBXkAlAA4AAixAQGwOLA1KwAAAAEARgAAAtMBmAAgACtAKBEBAgEBTBkBAUoAAQIBhQACAgBhAwEAAA8ATgEADw0GBQAgASAEBxYrISImNTQ3MwYGFRQWFxYzMjY3JiYnJjU0NjcWFhUUBwYGATp2fg0kAQIsMDFOXbJFBRINEyIgEBUjQr1VVUAuBCEMOj8RESYhFTciMhEVLRJKayQiRCovAAEARgAAA0MBGAApAGFADR0aAgIDJyMiAwACAkxLsAlQWEAZAAEDAYUAAwICA3AEAQICAGIFBgIAAA8AThtAGAABAwGFAAMCA4UEAQICAGIFBgIAAA8ATllAEwEAJiUgHxkXDw0GBQApASkHBxYrISImNTQ3MwYGFRQWFxYzMj4CNzY2NzYzMhcGBgcWFjMyFRUUIyInBgYBOnZ+DSQBAiwwMU4qXVZADREREAkWEAgNFQgiPCUICGRHPbFVVUAuBCEMOj8REQoREwkNHRwPByEuDBYUCDwIRyEmAAAB//gAAAE1AOQAIABTQA8RDgIAARsXFgMCBQMAAkxLsAlQWEAUAAEAAAFwAgEAAANiBQQCAwMPA04bQBMAAQABhQIBAAADYgUEAgMDDwNOWUANAAAAIAAgFRYlFQYHGisxIjU1NDMyNzY2NzYzMhcGBgcWFjMyFRUUIyInBgYHBgYICEkfCRwbCRYMDQsYDiM5IggIXUEECAUKSwg8CBwIMDMRCBo2GhURCDwIOwYLBQ0YAAAB//gAAADhAXYAGQAlQCIDAgIBAAFMEAgCAEoAAAABYQIBAQEPAU4AAAAZABkVAwcXKzEiNTU0MzI2NyYmJyY1NDY3FhYXFhUUBgcGCAguXyQHEgsTIx4CDAkPExFbCDwIFRAXNyAyERYuEA45LEIkDzcgNwAAAP//AEb/MQLTAZgCJgAfAAABBwUlAMT/8wAJsQEBuP/zsDUrAP//AEb/MQNDARgCJgAgAAABBwUlAOL/8wAJsQEBuP/zsDUrAP////j/MQE1AOQCJgAhAAABBgUl6vMACbEBAbj/87A1KwAAAP////j/MQDhAXYCJgAiAAABBgUl1vMACbEBAbj/87A1KwAAAP//AEb/CwLTAZgCJgAfAAABBwUuAKoAEQAIsQEDsBGwNSsAAP//AEb/CwNDARgCJgAgAAABBwUuAMgAEQAIsQEDsBGwNSsAAP////j/CwF5AOQCJgXuAAABBgUu7hEACLEBA7ARsDUr////+P8LATABdgImBe0AAAEGBS7XEQAIsQEDsBGwNSv//wBG/uYC0wGYAiYAHwAAAQcFJwEFAA8ACLEBArAPsDUrAAD//wBG/uYDQwEYAiYAIAAAAQcFJwEjAA8ACLEBArAPsDUrAAD////4/uYBNQDkAiYAIQAAAQYFJysPAAixAQKwD7A1K/////j+5gDhAXYCJgAiAAABBgUnFw8ACLEBArAPsDUr//8ARv8TAtMBmAImAB8AAAEHBTMAzQAPAAixAQSwD7A1KwAA//8ARv8TA0MBGAImACAAAAEHBTMA6wAPAAixAQSwD7A1KwAA////+P8TAXkA5AImBe4AAAEGBTMRDwAIsQEEsA+wNSv////4/xMBJAF2AiYF7QAAAQYFM/oPAAixAQSwD7A1K/////gAAAE1AOQCBgAhAAD////4AAAA4QF2AgYAIgAA//8ARgAAAtMByQImAB8AAAEHBSkAkv/2AAmxAQK4//awNSsA//8ARgAAA0MByQImACAAAAEHBSkAnP/2AAmxAQK4//awNSsA////+AAAAXkBvwImBe4AAAEGBSns7AAJsQECuP/ssDUrAAAA////+AAAAQ4CNwImACIAAAEGBSmwZAAIsQECsGSwNSv//wBG/6YC0wHJAiYAHwAAACcF6AEz/eEBBwUpAJL/9gASsQECuP3hsDUrsQMCuP/2sDUr//8ARv+mA0MByQImACAAAAAnBSkAnP/2AQcF6AEz/eEAErEBArj/9rA1K7EDArj94bA1K/////j/sAF5Ab8CJgXuAAAAJgUp7OwBBwXoAET96wASsQECuP/ssDUrsQMCuP3rsDUrAAD////4/7kBDgI3AiYAIgAAACYFKbBkAQcF6AA//fQAEbEBArBksDUrsQMCuP30sDUrAAAA//8ARgAAAtMB+QImAB8AAAEHBS0AiP/fAAmxAQO4/9+wNSsA//8ARgAAA0MB+QImACAAAAEHBS0Akv/fAAmxAQO4/9+wNSsA////+AAAAXkB7wImBe4AAAEGBS3i1QAJsQEDuP/VsDUrAAAA////+AAAAT0CZwImBe0AAAEGBS3eTQAIsQEDsE2wNSv//wBGAAAC0wICAiYAHwAAAQcFLwCS//YACbEBA7j/9rA1KwD//wBGAAADQwICAiYAIAAAAQcFLwCc//YACbEBA7j/9rA1KwD////4AAABeQH4AiYF7gAAAQYFL+zsAAmxAQO4/+ywNSsAAAD////4AAABCwJwAiYAIgAAAQYFL7BkAAixAQOwZLA1K///AEYAAALTAg4CJgAfAAABBwXWAM7/4gAJsQECuP/isDUrAP//AEYAAANDAg4CJgAgAAABBwXWANj/4gAJsQECuP/isDUrAP////gAAAF5AgQCJgXuAAABBgXWKNgACbEBArj/2LA1KwAAAP////gAAAD/AnwCJgAiAAABBgXW7FAACLEBArBQsDUr//8ARgAAAtMCJwImAB8AAAEHBSYA8QAiAAixAQKwIrA1KwAA//8ARgAAA0MCJwImACAAAAEHBSYA+wAiAAixAQKwIrA1KwAA////+AAAAXkCHQImBe4AAAEGBSZLGAAIsQECsBiwNSv////4AAABJAKVAiYF7QAAAQcFJgBHAJAACLEBArCQsDUrAAD//wBGAAAC0wH7AiYAHwAAAQcFMgCyAA4ACLEBBLAOsDUrAAD//wBGAAADQwH7AiYAIAAAAQcFMgC8AA4ACLEBBLAOsDUrAAD////4AAABeQHxAiYF7gAAAQYFMgwEAAixAQSwBLA1K/////gAAAEkAmkCJgXtAAABBgUyCHwACLEBBLB8sDUr//8ARv86AtMBmAImAB8AAAEHBfIAggAEAAixAQOwBLA1KwAA//8ARv8rA0MBGAImACAAAAEHBfIAoP/1AAmxAQO4//WwNSsA////+P8rAfkA5AImBgoAAAEGBfIW9QAJsQEDuP/1sDUrAAAA////+P8rAcABdgImBgkAAAEGBfL09QAJsQEDuP/1sDUrAAAA//8ARv8xAtMCAgImAB8AAAAnBSUAxP/zAQcFLwCS//YAErEBAbj/87A1K7ECA7j/9rA1K///AEb/MQNDAgICJgAgAAAAJwUlAOL/8wEHBS8AnP/2ABKxAQG4//OwNSuxAgO4//awNSv////4/zEBNQH4AiYAIQAAACYFJerzAQYFL8TsABKxAQG4//OwNSuxAgO4/+ywNSv////4/zEBCwJwAiYAIgAAACYFJdbzAQYFL7BkABGxAQG4//OwNSuxAgOwZLA1KwD//wBG/wMC0wGYAiYAHwAAAQcFMQCl//0ACbEBA7j//bA1KwD//wBG/wMDQwEYAiYAIAAAAQcFMQDD//0ACbEBA7j//bA1KwD////4/wMBeQDkAiYF7gAAAQYFMen9AAmxAQO4//2wNSsAAAD////4/wMBLwF2AiYF7QAAAQYFMdL9AAmxAQO4//2wNSsAAAD//wBG/wMC0wHJAiYAHwAAACcFKQCS//YBBwUxAKX//QASsQECuP/2sDUrsQMDuP/9sDUr//8ARv8DA0MByQImACAAAAAnBSkAnP/2AQcFMQDD//0AErEBArj/9rA1K7EDA7j//bA1K/////j/AwF5Ab8CJgXuAAAAJgUp7OwBBgUx6f0AErEBArj/7LA1K7EDA7j//bA1K/////j/AwFGAjcCJgXtAAAAJgUp6GQBBgUx0v0AEbEBArBksDUrsQMDuP/9sDUrAP//AEb/OwLTAdYCJgAfAAAAJwUqAMT//QEHBSMAogAJABGxAQK4//2wNSuxAwGwCbA1KwD//wBG/zsDQwHWAiYAIAAAACcFIwCsAAkBBwUqAOL//QARsQEBsAmwNSuxAgK4//2wNSsA////+P87AXkBzAImBe4AAAAmBSoI/QEGBSP8/wASsQECuP/9sDUrsQMBuP//sDUr////+P87ATQCRAImBe0AAAAmBSrx/QEGBSP4dwARsQECuP/9sDUrsQMBsHewNSsA//8ARv8XAtMBmAImAB8AAAEHBi8BDAATAAixAQGwE7A1KwAA//8ARv8XA0MBGAImACAAAAEHBi8BKgATAAixAQGwE7A1KwAA////+P8XAXkA5AImBe4AAAEGBi9QEwAIsQEBsBOwNSv////4/xcBJAF2AiYF7QAAAQYGLzkTAAixAQGwE7A1K///AEYAAALTAfoCJgAfAAABBwXzANH/wAAJsQEBuP/AsDUrAP//AEYAAANDAfoCJgAgAAABBwXzANv/wAAJsQEBuP/AsDUrAP////gAAAF5AfACJgXuAAABBgXzK7YACbEBAbj/trA1KwAAAP////gAAAEkAmgCJgXtAAABBgXzJy4ACLEBAbAusDUr//8ARv8XAtMBmAImAB8AAAEHBi4A+QAxAAixAQGwMbA1KwAA//8ARv8XA0MBGAImACAAAAEHBi4BFwAxAAixAQGwMbA1KwAA////+P8XAXkA5AImBe4AAAEGBi49MQAIsQEBsDGwNSv////5/xcBJQF2ACYF7QEAAQYGLicxAAixAQGwMbA1K///AET+4QJrAX0CJgCJAAABBwY1AMAAdwAIsQEBsHewNSsAAP//AET+4gKrAWcCJgCKAAABBwY1AJgAWQAIsQEBsFmwNSsAAP////j/SAKrAVQCJgCLAAABBwUlAKsACgAIsQEBsAqwNSsAAP////j/SAJaAVQCJgCMAAABBwUlAKsACgAIsQEBsAqwNSsAAP//AET+4QJrAX0CJgCJAAABBwV1AKIAmAAIsQEDsJiwNSsAAP//AET+4gKrAWcCJgCKAAABBgV1enoACLEBA7B6sDUr////+P8iAqsBVAImAIsAAAEHBS4AkQAoAAixAQOwKLA1KwAA////+P8iAloBVAImAIwAAAEHBS4AkQAoAAixAQOwKLA1KwAA//8ARP7hAmsBfQImAIkAAAEHBTYAxgCUAAixAQSwlLA1KwAA//8ARP7iAqsBZwImAIoAAAEHBTYAngB2AAixAQSwdrA1KwAA////+P8qAqsBVAImAIsAAAEHBTMAtAAmAAixAQSwJrA1KwAA////+P8qAloBVAImAIwAAAEHBTMAtAAmAAixAQSwJrA1KwAA//8ARP7hAmsCRwImAIkAAAAnBSMAiwB6AQcFdQCiAJgAELEBAbB6sDUrsQIDsJiwNSsAAP//AET+4gKrAikCJgCKAAAAJwUjAIsAXAEGBXV6egAQsQEBsFywNSuxAgOwerA1K/////j/IgKrAi4CJgCLAAAAJgUjX2EBBwUuAJEAKAAQsQEBsGGwNSuxAgOwKLA1K/////j/IgJaAi4CJgCMAAAAJgUjX2EBBwUuAJEAKAAQsQEBsGGwNSuxAgOwKLA1K///AET+4QJrAX0CJgCJAAABBwUrALQAcwAIsQECsHOwNSsAAP//AET+4gKrAWcCJgCKAAABBwUrAIwAVQAIsQECsFWwNSsAAP////j/UgKrAVQCJgCLAAABBwUqAKsAFAAIsQECsBSwNSsAAP////j/UgJaAVQCJgCMAAABBwUqAKsAFAAIsQECsBSwNSsAAP//AET+4QJrAX0CJgCJAAABBwUoAPkApwAIsQECsKewNSsAAP//AET+4gKrAWcCJgCKAAABBwUoANEAiQAIsQECsImwNSsAAP////j+/QKrAVQCJgCLAAABBwUnAOwAJgAIsQECsCawNSsAAP////j+/QJaAVQCJgCMAAABBwUnAOwAJgAIsQECsCawNSsAAAABAET+4QJrAX0AMABGQEMPCAIEAS0BBQQuAQAFA0wAAgABBAIBaQADAAQFAwRpAAUAAAVZAAUFAGEGAQAFAFEBACspIB8cGhYUDQsAMAEwBwcWKwEiJjU0Njc2Ny4CIyIGByYmNTQ2MzIWFxYWMzcHBgYjDgIHBgYVFBYzMjY3FwYGAWh8hz02Nz4MMzgWHiskBglDNitTKipWLkIdDhwOKV5ZJTc8cHM+ejwMSX/+4X53PnsvLxgFEQ4PFwUVDCYuFg0NFgJNAgIDGSgaJlwyT1QPDSEoJQABAET+4gKrAWcAQgBTQFASCgIEAzUkAgUELi0CBgVBQAIHBgRMAAIAAQMCAWkAAwAEBQMEaQAHCAEABwBlAAUFBmEABgYPBk4BAD89MTArKSIfHRsYFhAOAEIBQgkHFisBIiYmNTQ2NzY2NyYmJyYjIgYHJjU0NjMyFhcWMzcHBgYHIgYHHgIXFhYzMhUVFCMiJicmJwYGBwYGFRQWMzI3FwYBQUNjNzMuH08uKjMRGxgfKiQPQzYdRDtoUkIdDh4PDRkOEBYfGhY8KQgIWVwdHhJbgyANDm5kVn4Ohv7iNmFAOGcqHC4QExUGChAXDxgmLRUaMAJNAQIBAgIjJhIGBAQIPAghJSVDFUwyFCwXRkwkIVMAAAAAAf/4AAACqwFUADwARkBDEwEDASUKAgQDNQEABC0DAgMGAARMAAIAAQMCAWkAAwAEAAMEaQUBAAAGYQgHAgYGDwZOAAAAPAA7FCYSFiYpFQkHHSsxIjU1NDMyNjc2NyYmJyYmIyIGByY1NDYzMhYXFhYXFjMHLgInFhYXFjMyFRUUIyYmJyYmJwYGBw4DCAheijdBIxMnEyc0Gh8qJQ5CNiBBKzRJJSZJJwgjJAwLGRQkbgkJSlsdDhsMGzYbJT9FWAg8CBUZHQcJFw4dFxAXEBcmLhkeIyUICUoBAgQBEhYFCQg8CAEVFwwlGQofDhMZDgYAAAH/+AAAAloBVAAsAEdARBMBAwELAQQDIQEABAMCAgUABEwAAwEEAQMEgAAEAAEEAH4AAgABAwIBaQAAAAVhBgEFBQ8FTgAAACwAKyMWJikVBwcbKzEiNTU0MzI2NzY2NyYnJiYjIgYHJjU0NjMyFhcWFhcWMwcmJiMiBgYHDgMICEpzKkNBGCcmKDMaHyolDkI2IEErNUklJkgnECkXFyk3LSU/RVgIPAgNDRcVBBMfIBgQFxAXJi4bICYnCAlJAgMIGBcTGQ4GAAAA//8ARP7hAmsCXgImAIkAAAEPBXkBDf9wRmYACbEBAbj/cLA1KwAAAP//AET+4gKrAkACJgCKAAABDwV5AQ3/UkZmAAmxAQG4/1KwNSsAAAD////4AAACqwJFAiYAiwAAAQ8FeQDh/1dGZgAJsQEBuP9XsDUrAAAA////+AAAAloCRQImAIwAAAEPBXkA4f9XRmYACbEBAbj/V7A1KwAAAP//AET+4QJrApgCJgCJAAABBwUmANoAkwAIsQECsJOwNSsAAP//AET+4gKrAnoCJgCKAAABBwUmANoAdQAIsQECsHWwNSsAAP////gAAAKrAn8CJgCLAAABBwUmAK4AegAIsQECsHqwNSsAAP////gAAAJaAn8CJgCMAAABBwUmAK4AegAIsQECsHqwNSsAAP//AET+4QJrAjoCJgCJAAABBgUpe2cACLEBArBnsDUr//8ARP7iAqsCHAImAIoAAAEGBSl7SQAIsQECsEmwNSv////4AAACqwIhAiYAiwAAAQYFKU9OAAixAQKwTrA1K/////gAAAJaAiECJgCMAAABBgUpT04ACLEBArBOsDUr//8ARP7hAmsCcwImAIkAAAEGBS97ZwAIsQEDsGewNSv//wBE/uICqwJVAiYAigAAAQYFL3tJAAixAQOwSbA1K/////gAAAKrAloCJgCLAAABBgUvT04ACLEBA7BOsDUr////+AAAAloCWgImAIwAAAEGBS9PTgAIsQEDsE6wNSv//wBE/uECawF9AiYAiQAAAQcFMACc/mQACbEBA7j+ZLA1KwD//wBE/uICqwFnAiYAigAAAQcFMAB0/kYACbEBA7j+RrA1KwD////4/xoCqwFUAiYAiwAAAQcFMQCMABQACLEBA7AUsDUrAAD////4/xoCWgFUAiYAjAAAAQcFMQCMABQACLEBA7AUsDUrAAD//wBE/uECawJHAiYAiQAAAQcFIwCLAHoACLEBAbB6sDUrAAD//wBE/uICqwIpAiYAigAAAQcFIwCLAFwACLEBAbBcsDUrAAD////4AAACqwIuAiYAiwAAAQYFI19hAAixAQGwYbA1K/////gAAAJaAi4CJgCMAAABBgUjX2EACLEBAbBhsDUr//8ARP7hAmsBfQImAIkAAAEHBjYAyv5RAAmxAQK4/lGwNSsA//8ARP7iAqsBZwImAIoAAAEHBjYAov4zAAmxAQK4/jOwNSsA////+P9QAqsBVAImAIsAAAEHBjcA0wAyAAixAQKwMrA1KwAA////+P9QAloBVAImAIwAAAEHBjcA0wAyAAixAQKwMrA1KwAA//8ARP7hAmsBfQImAIkAAAEHBjgA9v4fAAmxAQS4/h+wNSsA//8ARP7iAqsBZwImAIoAAAEHBjgAzv4BAAmxAQS4/gGwNSsA////+P7rAqsBVAImAIsAAAEHBjkA0//sAAmxAQS4/+ywNSsA////+P7rAloBVAImAIwAAAEHBjkA0//sAAmxAQS4/+ywNSsA//8ARP7hAmsBfQImAIkAAAEHBgcBDP23AAmxAQG4/bewNSsA//8ARP7iAqsBZwImAIoAAAEHBgcA0f2JAAmxAQG4/YmwNSsA////+P7GAqsBVAImAIsAAAEHBgcA6/0MAAmxAQG4/QywNSsA////+P71AloBVAImAIwAAAEHBgcBMv07AAmxAQG4/TuwNSsA//8ARP7hAmsCOgImAIkAAAAmBSl7ZwEHBjUAwAB3ABCxAQKwZ7A1K7EDAbB3sDUr//8ARP7iAqsCHAAmAIoAAAAmBSl7SQEHBjUAmABZABCxAQKwSbA1K7EDAbBZsDUr////+P9IAqsCIQAmAIsAAAAmBSlPTgEHBSUAqwAKABCxAQKwTrA1K7EDAbAKsDUr////+P9IAloCIQImAIwAAAAmBSlPTgEHBSUAqwAKABCxAQKwTrA1K7EDAbAKsDUrAAEAIgAAAXQBjgAdACBAHRIJBAMBSgABAQBhAgEAAA8ATgEACAYAHQEdAwcWKzMiNTQ3FhYzMjcmJicuAjU0Nx4CFxYWFRQGBwacehIZQC9DRQgVDRsaCDQJFxgIDxQTEUg/IhMUFBUWLhcwMh4QMhAWNjQTIDcXDzcgJwABACIAAAHrAY4AHgAtQCodGAIAAQFMDwkEAwFKAgEBAQBhAwQCAAAPAE4BABsaFhQIBgAeAR4FBxYrMyI1NDcWFjMyNy4CNTQ3FhYXFhYzMhUVFCMiJicGnHoSGUAvP0sZKRg1Fy4iFCocCQkoPRdbPyITFBQVLV5OEjIQVYkwHRcIPAgbGTT//wAiAAABdAJXAiYAtQAAAQcFIwARAIoACLEBAbCKsDUrAAD//wAiAAAB6wJXAiYAtgAAAQcFIwAYAIoACLEBAbCKsDUrAAD//wAiAAABdANjAiYAtQAAACcFIwARAIoBBwV3AM4A1QAQsQEBsIqwNSuxAgGw1bA1KwAA//8AIgAAAesDYwImALYAAAAnBSMAGACKAQcFdwDVANUAELEBAbCKsDUrsQIBsNWwNSsAAP//ACUAAAF3Ao8AJgC1AwABBgXWQGMACLEBArBjsDUr//8AIgAAAesCjwImALYAAAEGBdZEYwAIsQECsGOwNSv//wAiAAABdAJKAiYAtQAAAQYFKQF3AAixAQKwd7A1K///ACIAAAHrAkoCJgC2AAABBgUpCHcACLEBArB3sDUr//8AIv9FAXQBjgImALUAAAEGBSobBwAIsQECsAewNSv//wAi/0UB6wGOAiYAtgAAAQYFKiEHAAixAQKwB7A1K///ACL/pgF0AY4CJgC1AAABBwXoAJT94QAJsQECuP3hsDUrAP//ACL/pgHrAY4CJgC2AAABBwXoAJT94QAJsQECuP3hsDUrAP//ACL/OwF0AY4CJgC1AAABBgUlG/0ACbEBAbj//bA1KwAAAP//ACL/OwHrAY4CJgC2AAABBgUlIf0ACbEBAbj//bA1KwAAAP//ACL/OwF0Ao8CJgC1AAAAJgUlG/0BBgXWPWMAEbEBAbj//bA1K7ECArBjsDUrAP//ACL/OwHrAo8CJgC2AAAAJgUlIf0BBgXWRGMAEbEBAbj//bA1K7ECArBjsDUrAP//ACIAAAF0AoMCJgC1AAABBgUvAXcACLEBA7B3sDUr//8AIgAAAesCgwImALYAAAEGBS8IdwAIsQEDsHewNSv//wAiAAABdAJ6AiYAtQAAAQYFLfdgAAixAQOwYLA1K///ACIAAAHrAnoCJgC2AAABBgUt/mAACLEBA7BgsDUr//8AIgAAAXQCfAImALUAAAEHBTIAIQCPAAixAQSwj7A1KwAA//8AIgAAAesCfAImALYAAAEHBTIAKACPAAixAQSwj7A1KwAA//8AIgAAAXQCZgImALUAAAEGBaN0bgAIsQEBsG6wNSv//wAiAAAB6wJmAiYAtgAAAQYFo3tuAAixAQGwbrA1K///ACX+8AF3Ao8AJgC1AwAAJgUnXxkBBgXWQGMAELEBArAZsDUrsQMCsGOwNSsAAP//ACL+8AHrAo8CJgC2AAAAJgUnYhkBBgXWRGMAELEBArAZsDUrsQMCsGOwNSsAAP//ACL/IQF0AY4CJgC1AAABBgYvYx0ACLEBAbAdsDUr//8AIv8hAesBjgImALYAAAEGBi9pHQAIsQEBsB2wNSsAAf/s/zQBTwFVAB4AK0AoBAEAAQFMFAsFAwFKAAEAAAFZAAEBAGECAQABAFEBAAkHAB4BHgMHFisXIiYmJzcWFjMyNjcmJicmJjU0NjcWFhcWFhUUBgcGbxAxMhANHC8US2scBw8JIiAeHAYdCg0SOzM0zBIaDCEFBlpVEB8QPUkZGh8NGEwdI0QcToUlJQAAAf/s/zQBpQFVACkAOkA3JiIhCwQDAgUBAQMEAQABA0wUAQJKAAEEAQABAGUAAgIDYQADAw8DTgEAJSQfHgkHACkBKQUHFisXIiYmJzcWFjMyNjcmJicmJjU0NjcWFhcWFhcWFhcWMzIVFRQjIicOAm8QMTIQDRwvFEtrHAcPCSIgHhwECQcQFQUGBwEiLAgILyMLPlzMEhoMIQUGWlUQHxA9SRkaHw0NHxEpPRMTHw0UCDwIGEVnOAAAAP///+z/NAFPAhoCJgDTAAABBgUjAk0ACLEBAbBNsDUr////7P80AaUCGgImANQAAAEGBSMCTQAIsQEBsE2wNSv////s/zQBTwJSACYA0wAAAQYF1i4mAAixAQKwJrA1K////+z/NAGlAlIAJgDUAAABBgXWLiYACLEBArAmsDUr////7P80AU8CPgAmANMAAAEGBfMxBAAIsQEBsASwNSv////s/zQBpQI+ACYA1AAAAQYF8zEEAAixAQGwBLA1K////+z+8gFPAVUAJgDTAAABBwXoAIz9LQAJsQECuP0tsDUrAP///+z+8gGlAVUAJgDUAAABBwXoAIz9LQAJsQECuP0tsDUrAP///+z+fAFPAVUCJgDTAAABBwUl/77/PgAJsQEBuP8+sDUrAP///+z+fAGlAVUCJgDUAAABBwUl/77/PgAJsQEBuP8+sDUrAP///+z+YgFPAVUAJgDTAAABBwYu//P/fAAJsQEBuP98sDUrAP///+z+UQGlAVUAJgDUAAABBwYu/+n/awAJsQEBuP9rsDUrAP///+z+fAFPAhoCJgDTAAAAJwUl/77/PgEGBSMCTQARsQEBuP8+sDUrsQIBsE2wNSsAAAD////s/nwBpQIaAiYA1AAAACcFJf++/z4BBgUjAk0AEbEBAbj/PrA1K7ECAbBNsDUrAAAA////7P80AVACDQImANMAAAEGBSnyOgAIsQECsDqwNSv////s/zQBpQINAiYA1AAAAQYFKfI6AAixAQKwOrA1K////+z/NAFPAkYCJgDTAAABBgUv8joACLEBA7A6sDUr////7P80AaUCRgImANQAAAEGBS/yOgAIsQEDsDqwNSv////s/zQBTwI/AiYA0wAAAQYFMhJSAAixAQSwUrA1K////+z/NAGlAj8CJgDUAAABBgUyElIACLEBBLBSsDUr////7P80AU8CKQAmANMAAAEGBaNlMQAIsQEBsDGwNSv////s/zQBpQIpACYA1AAAAQYFo2UxAAixAQGwMbA1K////+z/NAFyAVUAJgDTAAABBwYPAFj+kQAJsQEBuP6RsDUrAAAB/+z/NAGlAVUALABIQEUpJSQRBAUEBQEBAgQBAAEDTBoBBEoAAQYBAAEAZQAEBAVhAAUFD00AAgIDXwADAxACTgEAKCciIQ4NDAsJBwAsASwHBxYrFyImJic3FhYzMjY3IzUzNjY3JiYnJiY1NDY3FhYXFhYXFjMyFRUUIyInDgJvEDEyEA0cLxQyURyRrwYLBAcPCSIgHhwECQcUIAQiLAgILyMLPlzMEhoMIQUGKCUwDBkNEB8QPUkZGh8NDR8RNGEjFAg8CBhFZzgAAAD////s/zQBTwJrAiYA0wAAAQYFJlFmAAixAQKwZrA1K////+z/NAGlAmsCJgDUAAABBgUmUWYACLEBArBmsDUr////7P80AU8CHgImANMAAAEHBXkAi/90AAmxAQG4/3SwNSsA////7P80AaUCHgImANQAAAEHBXkAi/90AAmxAQG4/3SwNSsA////7P80AU8CqAAmANMAAAEGBjpBOwAIsQEEsDuwNSv////s/zQBpQKoACYA1AAAAQYGOkE7AAixAQSwO7A1K////+z/NAFPAvsAJgDTAAABBwV3AL4AbQAIsQEBsG2wNSsAAP///+z/NAGlAvsCJgDUAAABBwV3AL4AbQAIsQEBsG2wNSsAAAACAAX/EAGrAVUALgA6AGNAGRQBAwE4MisDAgMGAQACA0wjFwIBSi4BAElLsBlQWEATBAECAAACAGUAAwMBYQABARADThtAGgABAAMCAQNpBAECAAACWQQBAgIAYQAAAgBRWUANMC82NC86MDomKAUHGCsFJiYnJiYnBgYjIiYmNTQ2NjMyFhc2NjcmJicmJicmJjU0NjcWFhcWFhUUBxYWFyUyNjcmJiMiBgcWFgGSDRkODRsOHUotIkIrJDkdLU8jEBgJCA4IBRgKChIeHAcdCg0SKxswFv7SIjkXGjYcGiwUEDDwDBoNDhkNHyQgNiIgLxkmHhQxHBYcDQorExQtFhofDRhMHSNEHGFLHkAeWREQERURDxEWAAAAAAIABf8QAcsBVQA4AEIAe0AiMi0aAwECLgEDARcBBQNAOzQDBAUDAQAEBUwjAQJKOAEASUuwGVBYQB0GAQQAAAQAZQACAgNhAAMDD00ABQUBYQABARAFThtAGwABAAUEAQVpBgEEAAAEAGUAAgIDYQADAw8DTllAETo5Pz05QjpCMTArKigmBwcYKwUmJicGBwYjIiYnJjU0Njc2MzIWFxYWFzY2NyYmJyYmNTQ2NxYWFxYWFxYzMhUVFCMiJwYHHgIXJTI3JiYjIgcWFgGSLC0REBIyQCNEFBQiHB8dIDobCxUKDxkJBxcQFR4eHAcdCgsSAiUpCAguJAscCyElEP7SQjAfNBkxKQ8v8CwtDhMLJSEcHB8dMA0OFBEHEAgUMB0SLhwkRRkaHw0YTB0dQBcUCDwIGEA0CikxF1khFBIgEBcAAAEARP76A8IBVABSAElARiwpFwMCA01ECAcEBQIRAQEFA0w3AQNKAAMCA4UAAQcBAAEAZQQBAgIFYQYBBQUPBU4BAExIQkEvLigmHhwODABSAVIIBxYrEyImNTQ2NjcXBhUUFjMyNjY3JiY1NDY3HgIXFjMyNjc2NzY3NjYzMhcGBgcWFjMmJicmNTQ2NxYWFxYWFRQGBgciJicHBgcGBiMiJicVFAYHBvlVYAMICCULSUEyZVAQISAdJAIFCggeHhwsEAcFGAsIEBAOCgcWDjJQLgYSDhQjHwQPBQYJCxEKQF42CQcmEisaCRIJTkRD/vpiWhQiKyEJLhlJUiE5I0dgGBopFg4hMysJBgQMEVEXFA0HDz8xFhEUNyI2DhYtERhEFxs0FwotMRMSGRAMCAQDAQETPXEjJAAAAAABAET++gQJAQQAWQCoS7AeUFhAGzwrKBcEAgNUTUdCQQgHBwcCEQEBBwNMOQEDShtAGzkBAwU8KygXBAIDVE1HQkEIBwcHAhEBAQcETFlLsB5QWEAdBQEDAgOFAAEKAQABAGUGBAICAgdiCQgCBwcPB04bQCEABQMFhQADAgOFAAEKAQABAGUGBAICAgdiCQgCBwcPB05ZQBsBAFJQTEpFRD8+ODYuLCclHRsODABZAVkLBxYrEyImNTQ2NjcXBhUUFjMyNjY3JiY1NDY3FhYXFjMyNjc2NzY2NzYzMhcGBwcWMzI3NjY3NjY3NjMyFwYGBxcWMzIVFRQjIiYnBwYGIyInDgIjIiYnFRQGBwb5VWADCAglC0lBMmVQECEgHSQCDgobKBcnEAgQBxAJCxUPCA0bCxgdGR4EDwoNFwoIDA4KDxgMEC46CAgnTRoOCDUcLSMHKC4SDBkRTkRD/vpiWhQiKyEJLhlJUiE5I0dgGBopFhNINQYGBAw5GCoQDwceVyUMCgYkHiU1BQQHKEElCRoIPAgeGBsMDxkIDAUCAhU9cSMkAAAAAf/4AAACqAEEAFEAqEuwHlBYQBomAQMFOykXFAQCA05GQQkIAwYAAgNMOAEFShtAGjgBBQcmAQMFOykXFAQCA05GQQkIAwYAAgRMWUuwHlBYQB0HAQUDBYUAAwIDhQgGBAMCAgBiCgkBCwQAAA8AThtAIQAHBQeFAAUDBYUAAwIDhQgGBAMCAgBiCgkBCwQAAA8ATllAHQEATUtEQz89NzUtKyUjGxkTEQwLBgUAUQFRDAcWKzMiJicGBiMiNTU0MzI3NjY3NjMyFwYGBxYWMzI2Nz4CNzY2MzIXBgYHFhYzMjc2Njc2Njc2MzIXBgYHFhYzMhUVFCMiJicGBgcGBiMiJw4C4hxKFA44IggIMx4RERMJFg4KDxUFCDEgFycQBxESBgkQDw4KBxkUDxYQGh0FDgoIEgkIFw4KBxkSIzIiCQknTRoDBgUHNhsuIgcqMxINDhEIPAgXDhsfEAcmLQgECQYECzY4DRQMBxBIQgcFCgYkHhksDw8HEEc3ExAIPAgeGAcNBwwPGQkLBAAAAAH/+AAAAl8BVABGAEtASCcBAwUqGBUDAgNBCQgDBAACA0w2AQVKAAUDBYUAAwIDhQYEAgICAGEHAQgDAAAPAE4BAD8+Li0mJBwaFBIMCwYFAEYBRgkHFiszIiYnBgYjIjU1NDMyNz4CNzYzMhcGBgcWFjMyNjc+Ajc2NjMyFwYGBx4CMy4DNTQ2Nx4CFRQGBgciJicGBw4C4h1JFA44IggIMx4MDQ8NCRYOCg8VBQgxIBcnEAcREQcJDxAOCgcWDyM2NiIEExQPIx8EExAMEQpAXjYDBQYpNBMMDhEIPAgXCREZFRAHJi0IBAkGBAo1OQ4UDAcQQC8PEQcPNDktCBYsEhxQUB0LLTESEhkICAoLBQAA//8ARP76A8IB1AImAPcAAAAnBSMCBAAHAQcFJQIc//MAEbEBAbAHsDUrsQIBuP/zsDUrAP//AET++gQJAdQCJgD4AAAAJwUjAgQABwEHBSUCHP/zABGxAQGwB7A1K7ECAbj/87A1KwD////4/zECqAHUAiYA+QAAACcFIwCcAAcBBwUlAJb/8wARsQEBsAewNSuxAgG4//OwNSsA////+P8xAl8B1AImAPoAAAAnBSMAnAAHAQcFJQCW//MAEbEBAbAHsDUrsQIBuP/zsDUrAP//AET++gPCAVQCJgD3AAABBwUuAgIAEQAIsQEDsBGwNSsAAP//AET++gQJAQQCJgD4AAABBwUuAgIAEQAIsQEDsBGwNSsAAP////j/CwKoAQQCJgD5AAABBgUufBEACLEBA7ARsDUr////+P8LAl8BVAImAPoAAAEGBS58EQAIsQEDsBGwNSv//wBE/voDwgIAAiYA9wAAAQcFLwH0//QACbEBA7j/9LA1KwD//wBE/voECQIAAiYA+AAAAQcFLwH0//QACbEBA7j/9LA1KwD////4AAACqAIAAiYA+QAAAQcFLwCM//QACbEBA7j/9LA1KwD////4AAACXwIAAiYA+gAAAQcFLwCM//QACbEBA7j/9LA1KwD//wBE/voDwgIAAiYA9wAAACcFLwH0//QBBwUlAhz/8wASsQEDuP/0sDUrsQQBuP/zsDUr//8ARP76BAkCAAImAPgAAAAnBS8B9P/0AQcFJQIc//MAErEBA7j/9LA1K7EEAbj/87A1K/////j/MQKoAgACJgD5AAAAJwUvAIz/9AEHBSUAlv/zABKxAQO4//SwNSuxBAG4//OwNSv////4/zECXwIAAiYA+gAAACcFLwCM//QBBwUlAJb/8wASsQEDuP/0sDUrsQQBuP/zsDUr//8ARP76A8ICAAImAPcAAAAnBS8B9P/0AQcFLgICABEAEbEBA7j/9LA1K7EEA7ARsDUrAP//AET++gQJAgACJgD4AAAAJwUvAfT/9AEHBS4CAgARABGxAQO4//SwNSuxBAOwEbA1KwD////4/wsCqAIAAiYA+QAAACcFLwCM//QBBgUufBEAEbEBA7j/9LA1K7EEA7ARsDUrAAAA////+P8LAl8CAAImAPoAAAAnBS8AjP/0AQYFLnwRABGxAQO4//SwNSuxBAOwEbA1KwAAAAACAET++gQUAWUAMgA9AElARjs1GxcEBAUrCAcDAwQRAQEDA0wAAgAFBAIFaQABBgEAAQBlBwEEBANhAAMDDwNONDMBADk3Mz00PSooHx0ODAAyATIIBxYrEyImNTQ2NjcXBhUUFjMyNjY3JiY1NDY3FhYXFzY2MzIWFhcWFRQGBwYjIicWFhUUBgcGATI3JiYjIgYHFhb5VWADCAglC0lBMmVQECEgHSQCCggXZaBCFTQ1FB5NT01YZVkBAU5EQwGagYsfVSkxh0cpSP76YloUIishCS4ZSVIhOSNHYBgaKRYOOisIg3sXJhgiGz1dHRwYChQLPXEjJAFSVjU6YVYHBwAAAAIARP76BFwBZQA8AEoAU0BQSD8mGxcFAwc1MCsqCAcGBAMRAQEEA0wAAgAHAwIHaQABCAEAAQBlCQYCAwMEYQUBBAQPBE4+PQEAQ0E9Sj5KNDIuLSgnHx0ODAA8ATwKBxYrEyImNTQ2NjcXBhUUFjMyNjY3JiY1NDY3FhYXFzY2MzIWFhcWFRQHFjMyFRUUIyImJwYGIyInFhYVFAYHBgEyNyYmIyIGBwYGBxYW+VVgAwgIJQtJQTJlUBAhIB0kAwsGF2WgQhU0NRQeKC46CAgrVBgrfkFlWQEBTkRDAZqBix9VKSFTLxYvFylI/vpiWhQiKyEJLhlJUiE5I0dgGBopFhk4IgiDexcmGCIbQSwaCDwIIxkbIRgKFAs9cSMkAVJWNTotKRQxHAcHAAL/+AAAAxsBZQA3AEEAVEBRQDopHhsYBgIDMy4MCwQFAAICTAADCAIIAwKAAAQACAMECGkKBwUDAgIAYQYBCQMAAA8ATjk4AQA+PDhBOUExMCwqIyEXFQ8OCQgANwE3CwcWKyEiJiYnBgYHBiMiNTU0MzI2NzY2NzYzMhcGBgcWFhc2NzYzMhYXFhUUBxYzMhUVFCMiJicGBgcGJzI3JiYjIgYHFgGMMGZaHgQIBSFMCAgtNRsEBwQLFQ4JCg8EDBgNYlJQQxxTISAmLD4JCStTGxxJKS4ngokeVSkwgk1ODhoQBAoFJQg8CCEvBw4HEQkZIwkGCgSAQD4uJCUbPC8cCDwIIhwSGwgJTFY1OlxaDwAAAAL/+AAAAs4BZQArADUAS0BINC4eGxgFAgMMCwQDAAICTAADBgIGAwKAAAQABgMEBmkIBQICAgBhAQcCAAAPAE4tLAEAMjAsNS01IyEXFQ8OCQgAKwErCQcWKyEiJiYnBgYHBiMiNTU0MzI2NzY2NzYzMhcGBgcWFhc2NzYzMhYXFhUUBgcGJzI3JiYjIgYHFgGMMGZaHgQIBSFMCAgtNRsEBwQLFQ4JCg8EDBgNXlJTRBxTISBOT01OgokeVSkwgk1ODhoQBAoFJQg8CCEvBw4HEQkZIwkGCgR9QEEuJCUbPV4cHExWNTpcWg///wBE/voEFAFlAiYBDwAAAQcFKgJO//0ACbECArj//bA1KwD//wBE/voEXAFlAiYBEAAAAQcFKgJO//0ACbECArj//bA1KwD////4/zsDGwFlAiYBEQAAAQcFKgEA//0ACbECArj//bA1KwD////4/zsCzgFlAiYBEgAAAQcFKgEA//0ACbECArj//bA1KwD//wBE/voEFAJ4AiYBDwAAAQcFLwH+AGwACLECA7BssDUrAAD//wBE/voEXAJ4AiYBEAAAAQcFLwH+AGwACLECA7BssDUrAAD////4AAADGwJ4AiYBEQAAAQcFLwCwAGwACLECA7BssDUrAAD////4AAACzgJ4AiYBEgAAAQcFLwCwAGwACLECA7BssDUrAAD//wBE/voEFAJMAiYBDwAAAQcFIwIOAH8ACLECAbB/sDUrAAD//wBE/voEXAJMAiYBEAAAAQcFIwIOAH8ACLECAbB/sDUrAAD////4AAADGwJMAiYBEQAAAQcFIwDAAH8ACLECAbB/sDUrAAD////4AAACzgJMAiYBEgAAAQcFIwDAAH8ACLECAbB/sDUrAAD//wBE/voEFAJMAiYBDwAAACcFIwIOAH8BBwUlAk7/8wARsQIBsH+wNSuxAwG4//OwNSsA//8ARP76BFwCTAImARAAAAAnBSUCTv/zAQcFIwIOAH8AEbECAbj/87A1K7EDAbB/sDUrAP////j/MQMbAkwCJgERAAAAJwUlAQD/8wEHBSMAwAB/ABGxAgG4//OwNSuxAwGwf7A1KwD////4/zECzgJMAiYBEgAAACcFIwDAAH8BBwUlAQD/8wARsQIBsH+wNSuxAwG4//OwNSsAAAIAKgAAAs4CnwAsADcAPUA6NS8hDwwJCAcCAwFMGRgVAwFKAAEAAwIBA2kFAQICAGEEAQAADwBOLi0CADMxLTcuNiQiACwCLAYHFishIiYmJy4CJzcWFhc2NjcmJjU0NjcWFhcHHgMXFhYXNjMyFhcWFRQGBwYnMjcmJiMiBgcWFgGMKFJBERE4ORQLJUsnDx0OICEgIwURDRsDBwkGAQIBAY1jHFMhIE5PTU6CiR5VKTGLSiZMBAUDBBggDiMJEAYTJBGw1yciJgcYNSoXGEdKPQ0VJxOWLiQlGz1eHBxMVjU6ZFoEAwACACoAAAMYAp8AQABLAEpAR0lDNCoTEAkICAIFPjk4AwACAkwkIyADAUoAAQAFAgEFaQcEAgICAGEDBgIAAA8ATkJBAgBHRUFLQko8OzY1LSsAQAJACAcWKyEiJiYnLgInNxYWFzIiFhc2NjcmJicuAicmNDU0NjcWFhcHFhYXFhYXNjMyFhcWFRQGBxYzMhUVFCMiJicGBicyNyYmIyIGBxYWAYwoUkERETg5FAs1HgIBARkpDx0OBw4IBw4KAwEdJQYQDRsKCgEFBAGMYxxTISAUEi46CAgrVBosfjeCiR5VKTGLSiZMBAUDBBggDiMLBwEECBMkEStVKyRWURwIDQUiJwgYNikXWWEKHz4glS4kJRsfNxcaCDwIIxocIUxWNTpkWgQDAAL/+AAAAoYCnwAzAD8ARUBCOicdCwQABjEsKwMCBQMAAkwVFBEDAUoAAQAGAAEGaQUCAgAAA2EHBAIDAw8DTgAAPjw1NAAzADMvLikoIR81CAcXKzEiNTU0MzYyMzY2NyYmJzQ2NxYWFwceAxcWFhc2NjMyFhcWFRQHFjMyFRUUIyImJwYGJzY2NzY2NyYmIyIGCAgWKhYQIBAgIAEfIwUPEBoBBgcFAQUEAUd9LhxTISBGLDcICC1ZF0TyHWOQLhdCKx1SLC2MCDwIARYqFLPTKSEmCBMyMhYROj8yCTJAC05TLiQlG0sxFwg8CCYcHyNOBBMRCB8WMjluAAL/+AAAAmECnwApADUAO0A4MB0KAwAEAwICAgACTBUUEAMBSgABAAQAAQRpAwEAAAJhBQECAg8CTgAANDIrKgApACkhHxUGBxcrMSI1NTQzNzM2NjcmJic0NjceAhcHHgMXFhYXNjYzMhYXFhUUBQYGNzY2NzY2NyYmIyIGCAhTAxAgECAgAR8jBgoMCBoBBQcGAQUEAUd9LhxTISD+9UOrJmOQLhdCKx1SLC2MCDwIARYqFLPTKSEmCBciJBoWCzdBNwsnSwtOUy4kJRuSMw0NTgQTEQgfFjI5bv//ACoAAALOAp8CJgEjAAABBwUvASwAbAAIsQIDsGywNSsAAP//ACoAAAMYAp8CJgEkAAABBwUvASwAbAAIsQIDsGywNSsAAP////gAAAKGAp8CJgElAAABBwUvAM4AdgAIsQIDsHawNSsAAP////gAAAJhAp8CJgEmAAABBwUvAMQAdgAIsQIDsHawNSsAAP//ACoAAALOAp8CJgEjAAABBwUjATwAfwAIsQIBsH+wNSsAAP//ACoAAAMYAp8CJgEkAAABBwUjATwAfwAIsQIBsH+wNSsAAP////gAAAKGAp8CJgElAAABBwUjAN4AiQAIsQIBsImwNSsAAP////gAAAJhAp8CJgEmAAABBwUjANQAiQAIsQIBsImwNSsAAP//ACoAAALOAp8CJgEjAAABBwUpASwAbAAIsQICsGywNSsAAP//ACoAAAMYAp8CJgEkAAABBwUpASwAbAAIsQICsGywNSsAAP////gAAAKGAp8CJgElAAABBwUpAM4AdgAIsQICsHawNSsAAP////gAAAJhAp8CJgEmAAABBwUpAMQAdgAIsQICsHawNSsAAAABAEn+4gJ1AdAANgA+QDsTAQIBNB0bGBQGBgMCNQEAAwNMAAEAAgMBAmkAAwAAA1kAAwMAYQQBAAMAUQEAMC0XFREPADYBNgUHFisBIiYmNTQ3JiYnJjU0Njc2MzIWFwcmIyIHFBYXNjcWFRQGBwYGBwYGBwYGFRQWMzI2NzY2NxcGAXROdEBRHzkRES4lIyojOiwRPCM+QlJJSmULDA8OLBpLVxcMDXJvGS4WJ0kmC4v+4jViQXBXCSwcGyIrXRwdHiQgEiQ7TgwtIRAPDRYLChQKIEAjFCkWR04CAgMMCSJNAAACAET+4gIeAXYAOwBMAEdAREYhEhEJBQIFLQEDAjcBBAM4AQAEBEwAAQAFAgEFaQAEBgEABABlAAICA2EAAwMPA04BAEVDNDIqKCUjGxkAOwE7BwcWKwEiJicmNTQ2NjcmJicmJicmJwcmNTQ2Njc2MzIWFhUUBgcWFjMyFRUUIyImJicGBhUUFjMyNjY3Fw4CAzY2NzY1NCYjIgcWFhcWFhcBTEZoGhcUMCkHDQUTHQkQERATIzUbKSgrRCcuNB1LPwoKOU89IDozXWcgMDsvDzFIOl0VHQcPKSI+OAgTCgsdE/7iNi8oMyc5OCIIDwgcJgoUDgoVGBgkGggMITgkK0kkDAkJOgkKGRYnRCs5OAQPDiIeIw4BoREcChUVHCQeBxALCyIYAAL/+AAAAa8BdgAqADYAOkA3Mx4ODQcFAQUoAQABAkwAAgAFAQIFaQMBAQEAYQQGAgAADwBOAQAyMCYkIR8XFQYEACoBKgcHFiszIjU1NDMyNycmJicmJwcmNTQ2Njc2MzIWFxYVFAYHFjMyFRUUIyImJwYGNzY2NTQmIyIHFhcWAQkJaEMSEx0JEg8QEiEzGygtK0MTFDEoKl0KCkRfJzdmqSInKiE9OSU7Cwk6CRgXHSYKFwoKFBgYIxsIDSEcHCQvTCASCToJGhwcGoQXMhccJB4dSg0AAAH/+AAAAd8BdgAqAD1AOhIBAgEhGBMDAwIDAgIEAANMAAMCAAIDAIAAAQACAwECaQAAAARhBQEEBA8ETgAAACoAKhUlJyUGBxorMSI1NTQzMjcmJjU0Njc2MzIWFwcmJiMiBxQWFxYXPgI3FhUUBwYGBwYGCAhAPRgkLiUlKCM/KhIoIxY/QSkoJzYaLTkqDRYEKhlntAg8CAMQNh8rXhwdIiUgDwglK0YWFQMFDBUQDBghDQILBhoZAAD//wBJ/uICdQLPAiYBMwAAAQcFLwAKAMMACLEBA7DDsDUrAAD//wBE/uICHgJvAiYBNAAAAQYFLyBjAAixAgOwY7A1K/////gAAAGvAm8CJgE1AAABBgUv/GMACLECA7BjsDUr////+AAAAd8CbwImATYAAAEGBS8AYwAIsQEDsGOwNSv//wBJ/uICdQKjAiYBMwAAAQcFIwAaANYACLEBAbDWsDUrAAD//wBE/uICHgJDAiYBNAAAAQYFIzB2AAixAgGwdrA1K/////gAAAGvAkMCJgE1AAABBgUjDHYACLECAbB2sDUr////+AAAAd8CQwImATYAAAEGBSMQdgAIsQEBsHawNSv//wBJ/uICdQKjAiYBMwAAACcFIwAaANYBBwY1ALwAZQAQsQEBsNawNSuxAgGwZbA1KwAA//8ARP7iAh4CQwImATQAAAAmBSMwdgEGBjV3KAAQsQIBsHawNSuxAwGwKLA1KwAA////+P8xAa8CQwImATUAAAAmBSMMdgEGBSUo8wARsQIBsHawNSuxAwG4//OwNSsA////+P8xAd8CQwImATYAAAAmBSMQdgEGBSVa8wARsQEBsHawNSuxAgG4//OwNSsA//8ARAAAAx0CpgImBekAAAEHBSMBoADZAAixAgGw2bA1KwAA//8ARAAAA0ICRAImAUwAAAEHBSMBsAB3AAixAgGwd7A1KwAA////+AAAAZQCOwImAWQAAAEGBSMBbgAIsQIBsG6wNSv////4AAABcQKVAiYBZQAAAQcFI//yAMgACLECAbDIsDUrAAD//wBEAAADHQLSAiYF6QAAAQcFLwGQAMYACLECA7DGsDUrAAD//wBEAAADQgJwAiYBTAAAAQcFLwGgAGQACLECA7BksDUrAAD////4AAABlAJnAiYBZAAAAQYFL/FbAAixAgOwW7A1K/////gAAAFqAoUCJgX8AAABBgUv4nkACLECA7B5sDUr//8ARAAAAx0B4QIGBekAAAACAEQAAANCAYcALgA8AEVAQjoxIhIEAgEsKCcDAAICTAABBgIGAQKAAAMABgEDBmkEAQICAGIFBwIAAA8ATgEANzUrKiUkHBoQDgYFAC4BLggHFishIjU0NjczBgYVFBYXFhYzMjY3JjU0NjY3NjYzMhYXFhUUBxYWMzIVFRQjIicGBjc2NyYmJyYjIgYGBxYWAUD8BwckAgItMhpCKTBZKigGCAQaUzEaMQ8POh0+HwgIfVAxncRIKAMZEhILEzAuEBQuqiA2GAwZDD09EQcJCQgqMA8rJwcwODQqKytJMwUGCDwIKBIWdxsjGTkVFiAzHxclAAAA////+AAAAZQBdgIGAWQAAP////gAAAFxAdACBgFlAAD//wBE/0UDHQHhAiYF6QAAAQcFJQHSAAcACLECAbAHsDUrAAD//wBE/zEDQgGHAiYBTAAAAQcFJQG+//MACbECAbj/87A1KwD////4/zEBlAF2AiYBZAAAAQYFJQ/zAAmxAgG4//OwNSsAAAD////4/zEBcQHQAiYBZQAAAQYFJS3zAAmxAgG4//OwNSsAAAD//wBE/0UDHQKmAiYF6QAAACcFJQHSAAcBBwUjAaAA2QAQsQIBsAewNSuxAwGw2bA1KwAA//8ARP8xA0ICRAImAUwAAAAnBSMBsAB3AQcFJQG+//MAEbECAbB3sDUrsQMBuP/zsDUrAP////j/MQGUAjsCJgFkAAAAJgUjAW4BBgUlD/MAEbECAbBusDUrsQMBuP/zsDUrAP////j/MQFxApUCJgFlAAAAJwUj//IAyAEGBSUt8wARsQIBsMiwNSuxAwG4//OwNSsAAAD//wBE/x8DHQHhAiYF6QAAAQcFLgG4ACUACLECA7AlsDUrAAD//wBE/wsDQgGHAiYBTAAAAQcFLgGkABEACLECA7ARsDUrAAD////4/wsBlAF2AiYBZAAAAQYFLvURAAixAgOwEbA1K/////j/CwFxAdACJgFlAAABBgUuExEACLECA7ARsDUr//8ARP9FAx0C0gImBekAAAAnBSUB0gAHAQcFLwGQAMYAELECAbAHsDUrsQMDsMawNSsAAP//AET/MQNCAnACJgFMAAAAJwUvAaAAZAEHBSUBvv/zABGxAgOwZLA1K7EFAbj/87A1KwD////4/zEBlAJnAiYBZAAAACYFL/FbAQYFJQ/zABGxAgOwW7A1K7EFAbj/87A1KwD////4/zEBagKFAiYF/AAAACYFL+J5AQYFJR3zABGxAgOwebA1K7EFAbj/87A1KwD//wBEAAADHQLLAiYF6QAAAQcFMgGwAN4ACLECBLDesDUrAAD//wBEAAADQgJpAiYBTAAAAQcFMgHAAHwACLECBLB8sDUrAAD////4AAABlAJgAiYBZAAAAQYFMhFzAAixAgSwc7A1K/////gAAAFqAn4CJgX8AAABBwUyAAIAkQAIsQIEsJGwNSsAAAACAET++gJUAVQAKQA5AEZAQy0BBAUVCAcDAgQSAQECA0wAAwAFBAMFaQABBgEAAQBlBwEEBAJhAAICDwJOKyoBADMxKjkrOSEfGRcODAApASkIBxYrASImNTQ2NjcXBhUUFjMyNjc2NyYmJwYGIyImNTQ2NzYzMhYXFhUUBwYGEzI2NyYmJyYjIgYHBhUUFgEKXmgDCAglC1BMNG8sLBkBBQUgQyY5RCkhISIvVxkYYS9/YBozGQ0qGBkWFCgODDr++mNZFCIrIQkuGUxPIB0cJyA3GB4eRzozYx8eVUVDTn9VKTIBWAsMLEcREiAXFhUkJwAAAv/4AAABlAF2ACQAMQA6QDcwJxgIBAEFIh4CAAECTAACAAUBAgVpAwEBAQBhBAYCAAAPAE4BAC0rISAcGRIQBwQAJAEkBwcWKzMiNTU0MzI2NyY1NDY2NzY2MzIWFxYVFAcWFjMyFRUUIyInBgY3NjcmJicmIyIGBgcWAQkJHCoQJQYIBBpUMBowDxArEDYlCQl8USJbfTktAxsREgoTMC8QKgk6CQICKS0RKyUHLzkyKyosPzMCAwg8CCARD2cVKBk8FBQfNB8wAAAAAv/4AAABcQHQACAAKwBFQEIjAQQFCwECBAgBAQIDTAADAAUEAwVpBwEEAAIBBAJpAAEBAGEGAQAADwBOIiEBACclISsiKxcVDw0GBAAgASAIBxYrMyI1NTQzMjY3JiYnBgYjIiY1NDY3NjMyFhcWFRQHDgI3MjcmJiMiBhUUFgEJCWyVPwIHBB9KJSw4IxseIDRQGBcjKFx3ajI4Ez4nHikpCToJGhwUJxMZHkMyM1oaG0tERU4iRBkgD+UZPUA3Ih8eAAIARP76ApMBXAAsADsASEBFJQgHAwIEEAEBAgJMAAMABgQDBmkAAQgBAAEAZQkHAgQEAmEFAQICDwJOLS0BAC07LTsxLygnIyEdGxIRDgwALAEsCgcWKwEiJjU0NjY3FwYVFBYzMjY3JyImJicmNTQ2NzYzMhYXFhczMhUVFCMjBgYHBhMmJiMiBgcGFRQWFxYXFgEKXmgDCAglC1BMWJQoAyNQRxY+KSEhIitQGhsGNwkJOApfS0y/FkcnFCkNDBMWGCEh/vpjWRQiKyEJLhlMT0U7NAYNCh5NM2QeHkxAP0UIPAhKdiMjAVRcXSAXFhQVHgsLBgf//wBE/voCVAIMAiYBYwAAAQcFKQC9ADkACLECArA5sDUrAAD//wBE/voCkwIWAiYBZgAAAQcFKQC9AEMACLECArBDsDUrAAD////4AAABlAIuAiYBZAAAAQYFKfFbAAixAgKwW7A1K/////gAAAFxAogCJgFlAAABBwUp/+IAtQAIsQICsLWwNSsAAP//AET++gJUAhkCJgFjAAABBwUjAM0ATAAIsQIBsEywNSsAAP//AET++gKTAiMCJgFmAAABBwUjAM0AVgAIsQIBsFawNSsAAP//AET++gJUAkUCJgFjAAABBwUvAL0AOQAIsQIDsDmwNSsAAP//AET++gKTAk8CJgFmAAABBwUvAL0AQwAIsQIDsEOwNSsAAP//AET+PwJUAgwCJgFjAAAAJwUpAL0AOQEHBSUAVv8BABGxAgKwObA1K7EEAbj/AbA1KwD//wBE/j8CkwIWAiYBZgAAACcFKQC9AEMBBwUlAFb/AQARsQICsEOwNSuxBAG4/wGwNSsA////+P8xAZQCLgImAWQAAAAmBSnxWwEGBSUP8wARsQICsFuwNSuxBAG4//OwNSsA////+P8xAXECiAImAWUAAAAnBSn/4gC1AQYFJS3zABGxAgKwtbA1K7EEAbj/87A1KwAAAAACABEAAAHyAp8AIwBFAIZAFjYBBQQ9NwIDBQ0HAgECA0wcGxgDBEpLsA1QWEAkAAUEAwMFcgcBAgMBAwIBgAAEAAMCBANpAAEBAGEGAQAADwBOG0AlAAUEAwQFA4AHAQIDAQMCAYAABAADAgQDaQABAQBhBgEAAA8ATllAFyUkAQA6ODUzLi0kRSVFCwkAIwEjCAcWKzMiJyY1NDY3FhYzMjY3JiYnLgM1NDY3FhYXBxIVFAYHBgYnIiY1NjY3NjY3JiY1NDY2MzIXByYjIgYGBxYWFxYVFAYG4kBISQoGJ3c2NWw2ChQLBg0KBh4kBRUJGicUEDd3bA0KGjERCAsDLzQYJBIaCBAFBwgXEgIRNA0JJjYGBggRMgoJDAoKO3c7IE1LPA8hJggTRh4W/pQ6HCsUCAnCDAgCEgwFCwYBER4QJBocBwUQFgsIBQkGDQ0mHQAAAgBEAAAC8QKfACkASwCjQBs8AQgHQz0CBggdEgICBSciIQMAAgRMGhcCB0pLsA1QWEAtAAgHBgYIcgABBgUGAQWACgEFAgYFAn4ABwAGAQcGaQMBAgIAYQQJAgAADwBOG0AuAAgHBgcIBoAAAQYFBgEFgAoBBQIGBQJ+AAcABgEHBmkDAQICAGEECQIAAA8ATllAHSsqAQBAPjs5NDMqSytLJSQfHhAOCQgAKQEpCwcWKyEiJicmNTQ2NzMGFRQWFjMyNjcuAyc2NjcWEhcWMzIVFRQjIiYnBgYnIiY1NjY3NjY3JiY1NDY2MzIXByYjIgYGBxYWFxYVFAYGATRcZBgYBwcjAzZhQE6CPwUJCQoHDiETBQcCLTwICC4/FU6RVQ0KGjERCAsDLzQYJBIaCBAFBwgWEwIRNA0JJjYfISExEyYTEBYxLg0PEDJjc5NhDxwNwP7waBsIPAgbGRsZwgwIAhIMBQsGAREeECQaHAcFEBYLCAUJBg0NJh0AAAAAAf/4AAAB3AKOADIALkArLikDAgQCAAFMJB0JAwBKAQEAAAJhBAMCAgIPAk4AAAAyADIsKyclFQUHFysxIjU1NDMyNzY3JicmJyYmNTQ3NjY3FhUUBgcGBgcWFhcWFRQHFjMyFRUUIyImJwYGBwYICHtaJiwVQDJKKyoUFrm7DxQWValWSYUnKgY8TgkJMFcjBxIKUwg8CBoKF0A2KyMVJRokFBZVTA8XFxoJI0cjE1I0Ni4WEjAIPAgjJgcNBi8AAf/4AAABjQKOAC4ALkArCgEAAQMCAgIAAkwlAQFKAAEAAYUAAAACYQMBAgIPAk4AAAAuAC4qFQQHGCsxIjU1NDMyNjc2NyYnJiYnBiMiJjU0NjY3PgM3FhUUBgcGBgcWFhcWFRQHBgYICEVoKCkpEjMaRS0MDRojCg0DQ2JPSCgPFBdmnzhMiykqVyp7CDwIDwsLFjgwGC4WAyUcDysjAiMwJB8QDBoXGgkrRRoTVTU0L1IqExUAAAEARAAAAsACjgAzADNAMBIBAgEBTCwlAgNKAAMBA4UAAQIBhQACAgBhBAEAAA8ATgEAGhgQDggHADMBMwUHFishIiYnJjU0NzMGBhUUFxYzMjY3JiYnJicGIyImNTQ2Njc+AzcWFhUUBwYHFhYXFhUUBgE0V2sZFQ0kAgIuL25WiUAMLSQuRw0MGyIKDQNEYExIKwgIK8xyTowpKaciIh4rIh4JFQswExUZIiI+HCYiAyUcDysjAiMwIx4SBg8LMBBVNRNXNDUtUFQAAAEAQwAAAw8CjgA9ADlANjIPAgIBPDcCAAICTCslIgMBSgABAgGFAwECAgBhBAUCAAAPAE4BADo5NTMNCwYFAD0BPQYHFishIiY1NDczBhUUFxYzMjY3JiYnJiYnJiY1NDc2Njc+AzcWFhUGBgcGBgcWFhcWFRQHFjMyFRUUIyImJwYBM3t1DSQELi9uVYpACyogGT0mKyoUCDw1Gyw3UkAICAMRF1WpVUiFKCoGPE4JCTBXI0pHRiIeERgwExUYIyA8GxQnEhUlGiUTCCIaDRMXIRsHFQoWGwkjRyMTUjQ2LhQUMAg8CCMmSQD////4AAAB3AKOAgYBdQAA////+AAAAY0CjgIGAXYAAP//AEQAAALAAqUCJgF3AAABBwYDAXX/uQAJsQECuP+5sDUrAP//AEMAAAMPAqUCJgF4AAABBwYDAXX/uQAJsQECuP+5sDUrAP////UAAAHcAqUCJgF1AAABBgYDSbkACbEBArj/ubA1KwAAAP////gAAAGNAqUCJgF2AAABBgYDU7kACbEBArj/ubA1KwAAAP//AET/CwLAAo4CJgF3AAABBwUuAPIAEQAIsQEDsBGwNSsAAP//AEP/CwMPAo4CJgF4AAABBwUuAPIAEQAIsQEDsBGwNSsAAP////j/CwHcAo4CJgF1AAABBgUu9REACLEBA7ARsDUr////+P8LAY0CjgImAXYAAAEGBS71EQAIsQEDsBGwNSv//wBEAAACwAKVAiYBdwAAAQcGCAD4AJAACLEBAbCQsDUrAAD//wBDAAADDwKVAiYBeAAAAQcGCAD4AJAACLEBAbCQsDUrAAD////4AAAB3AKVAiYBdQAAAQcGCP/MAJAACLEBAbCQsDUrAAD////4AAABjQKVAiYBdgAAAQcGCP/MAJAACLEBAbCQsDUrAAD//wBEAAACwALRAiYBdwAAAQcGBAEs/4IACbEBA7j/grA1KwD//wBDAAADDwLRAiYBeAAAAQcGBAEs/4IACbEBA7j/grA1KwD////4AAAB3ALRAiYBdQAAAQYGBACCAAmxAQO4/4KwNSsAAAD////4AAABjQLRAiYBdgAAAQYGBACCAAmxAQO4/4KwNSsAAAD//wBE/wMCwAKOAiYBdwAAAQcFMQDt//0ACbEBA7j//bA1KwD//wBD/wMDDwKOAiYBeAAAAQcFMQDt//0ACbEBA7j//bA1KwD////4/wMB3AKOAiYBdQAAAQYFMfD9AAmxAQO4//2wNSsAAAD////4/wMBjQKOAiYBdgAAAQYFMfD9AAmxAQO4//2wNSsAAAD//wBEAAACwALwAiYBdwAAAQcF+gHb/64ACbEBAbj/rrA1KwD//wBDAAADDwLxAiYBeAAAAQcF+wHO/6kACbEBAbj/qbA1KwD////VAAAB3ALxAiYBdQAAAQcF+wCM/6kACbEBAbj/qbA1KwD////nAAABjQLwAiYBdgAAAQcF+gCr/64ACbEBAbj/rrA1KwD//wBEAAACwALwAiYBdwAAACcF6AHr/+EBBwX6Adn/rgASsQECuP/hsDUrsQMBuP+usDUr//8AQwAAAw8C8QImAXgAAAAnBegB6//UAQcF+wHJ/6kAErEBArj/1LA1K7EDAbj/qbA1K////9UAAAHcAvECJgF1AAAAJwX7AIz/qQEHBegAuv/UABKxAQG4/6mwNSuxAgK4/9SwNSv////pAAABjQLwAiYBdgAAACcF+gCt/64BBwXoAKP/2AASsQEBuP+usDUrsQICuP/YsDUr//8ARAAAAsADOgImAXcAAAAnBfoB2f+uAQcGBAEs/+sAErEBAbj/rrA1K7ECA7j/67A1K///AEMAAAMPAzoCJgF4AAAAJwX7Acn/qQEHBgQBLP/rABKxAQG4/6mwNSuxAgO4/+uwNSv////VAAAB3AM6AiYBdQAAACcF+wCM/6kBBgYEAOsAErEBAbj/qbA1K7ECA7j/67A1KwAA////5wAAAY0DOgImAXYAAAAnBfoAq/+uAQYGBADrABKxAQG4/66wNSuxAgO4/+uwNSsAAAABAAoAAAPEAgcAPAAxQC4NBwIBAgFMJwEDSgADAAIBAwJnAAEBAGEEAQAADwBOAgAyKhYPDAgAPAI6BQcWKyEiJicuAjUWMzI2NjcmJiMiBgcGIyImNTQ3NjY3FhUUBw4CBwYVFBYWMzI2NzY2MzIWFxYVFAcOAwEeVFwZICEKcpN989RLD0Q5H2ZGi0lDPE4nbjcDLhxEPhIbIzIYHWFFRmMfMlMXGBZCoK6vBAYHEiMeGAcOCSwmAwQJKzJORCM6Dw0MNxILJSYPFg8IEAkDBAQDISAfKyIvChAMBgAAAAEACQAABDoCBwBDATVLsBZQWEAVNgwGAwECQDsCAAECTCYBA0o8AQBJG0uwGFBYQBU2DAYDAQJAOwIAATwBBQADTCYBA0obS7AwUFhAFTYMBgMBAkA7AgAEPAEFAANMJgEDShtAEjYMBgMBAkA8OwMFBAJMJgEDSllZWUuwFlBYQBYAAwACAQMCZwQBAQEAYQUGAgAADwBOG0uwGFBYQCAAAwACAQMCZwQBAQEAYQYBAAAPTQQBAQEFYQAFBQ8FThtLsC1QWEAeAAMAAgEDAmcAAQEAYQYBAAAPTQAEBAVhAAUFDwVOG0uwMFBYQBwAAwACAQMCZwABBgEABQEAaQAEBAVhAAUFDwVOG0AbAAECBAIBBIAAAwACAQMCZwAEBAVhAAUFDwVOWVlZWUATAQA/Pjk4MisVDgsHAEMBQgcHFiszIicmJic3FjMyNjY3JiYjIg4CIyImNTQ3NjY3FhUUBw4CBwYVFBYXFhYzMj4CMzIWFxYVBxYzMhUVFCMiJycGBPZ4KSohAQFyk33z1EsPRDknbHZuKEM8TiduNwMuHEQ+EhsSDhAqEyZnb2gnMlMXGAI1OwgIRTsFlv6eCgkfIREYBw4JLCYFBgUrMk5EIzoPDQw3EgslJg8WDwYKBQYGBAYEISAfKxMeCDwIKQIVFgAAAf/4AAADvAIHAEMAy0uwGFBYQBI2CgIAAT87AwIEBAACTCcBAkobS7AeUFhAEjYKAgABPzsDAgQEAwJMJwECShtAFjYKAgABPwMCBQMCTDsCAgUBSycBAkpZWUuwGFBYQBYAAgABAAIBZwMBAAAEYQYFAgQEDwROG0uwHlBYQCEAAgABAAIBZwAAAARhBgUCBAQPTQADAwRhBgUCBAQPBE4bQB4AAgABAAIBZwAAAAVhBgEFBQ9NAAMDBGEABAQPBE5ZWUARAAAAQwBBPj05NzEqc0UHBxgrMSI1NTQzMj4CNyYmIyIOAiMiJjU0Njc2NjcWFBUUBgcOAgcGFRQWFjMyPgIzMhYXFhUHFjMyFRUUIyInDgMICGrk2rg/D0Q5J2x3bSlDO0ZHI0cjAhAdHEQ+ExsjMxcmZ3BnJzJUFxcBMEAJCUVAPLXd7Qg8CAIGDAosJgUGBSsyNWErFR8JBwwFEisNCyUmDxYPCBAJBAYEISAeLBMeCDwILAoRCwYAAf/4AAADRAIHADsANkAzCgEAAQMBAwACTCcBAkoCAQNJAAIAAQACAWcAAAADYQQBAwMPA04AAAA7ADkxKnNFBQcYKzEiNTU0MzI+AjcmJiMiDgIjIiY1NDY3NjY3FhQVFAYHDgIHBhUUFhYzMj4CMzIWFxYVFAcOAwgIauTauD8PRDknbHdtKUM7RkcjRyMCEB0cRD4TGyMzFyZncGcnMlQXFxU8t93tCDwIAgYMCiwmBQYFKzI1YSsVHwkHDAUSKw0LJSYPFg8IEAkEBgQhIB4sITAKEQsGAAAA//8ARAAAAsACjgAmAXcAAAEHBegB6//hAAmxAQK4/+GwNSsA//8AQwAAAw8CjgImAXgAAAEHBegB6//UAAmxAQK4/9SwNSsA////+AAAAdwCjgImAXUAAAEHBegAuv/UAAmxAQK4/9SwNSsA////+AAAAY0CjgImAXYAAAEHBegAo//YAAmxAQK4/9iwNSsA//8AEQAAAfICnwImAXMAAAEHBSMAHACVAAixAgGwlbA1KwAA//8ARAAAAvECnwImAXQAAAEHBSMAnACVAAixAgGwlbA1KwAA////+AAAAdwClQImAXUAAAEHBgj/4ACQAAixAQGwkLA1KwAA////+AAAAY0ClQImAXYAAAEHBgj/4ACQAAixAQGwkLA1KwAA//8AEQAAAfICnwImAXMAAAEHBcEA5f8xAAmxAgK4/zGwNSsA//8ARAAAAvECnwImAXQAAAEHBcEBZf87AAmxAgK4/zuwNSsA////9QAAAdwCpQImAXUAAAEGBgNJuQAJsQECuP+5sDUrAAAA////+AAAAY0CpQImAXYAAAEGBgNTuQAJsQECuP+5sDUrAAAA//8AEf8LAfICnwImAXMAAAEGBS44EQAIsQIDsBGwNSv//wBE/wsC8QKfAiYBdAAAAQcFLgCeABEACLECA7ARsDUrAAD//wBE/zsCwALwAiYBdwAAACcF+gHb/64BBwUqAQz//QASsQEBuP+usDUrsQICuP/9sDUr//8AQ/87Aw8C8QImAXgAAAAnBfsB0v+pAQcFKgEM//0AErEBAbj/qbA1K7ECArj//bA1K////9X/PQHcAvECJgF1AAAAJwX7AIz/qQEGBSoE/wASsQEBuP+psDUrsQICuP//sDUrAAD////n/z0BjQLwAiYBdgAAACcF+gCr/64BBgUqA/8AErEBAbj/rrA1K7ECArj//7A1KwAA//8AEQAAAfICnwImAXMAAAEHBa0A5wCaAAixAgOwmrA1KwAA//8ARAAAAvECnwImAXQAAAEHBa0BZwCkAAixAgOwpLA1KwAA////+AAAAdwC0QImAXUAAAEGBgQAggAJsQEDuP+CsDUrAAAA////+AAAAY0C0QImAXYAAAEGBgQAggAJsQEDuP+CsDUrAAAA//8ARAAAAsAC/wImAXcAAAAnBfoB2f+uAQcGAwF1ABMAEbEBAbj/rrA1K7ECArATsDUrAP//AEMAAAMPAv8CJgF4AAAAJwX7Acn/qQEHBgMBdQATABGxAQG4/6mwNSuxAgKwE7A1KwD////VAAAB3AL/AiYBdQAAACcF+wCM/6kBBgYDUxMAEbEBAbj/qbA1K7ECArATsDUrAAAA////5wAAAY0C/wImAXYAAAAnBfoAq/+uAQYGA1MTABGxAQG4/66wNSuxAgKwE7A1KwAAAP//AET+5gLAAvACJgF3AAAAJwX6Adv/rgEHBScBTQAPABGxAQG4/66wNSuxAgKwD7A1KwD//wBD/uYDDwLxAiYBeAAAACcF+wHJ/6kBBwUnAU0ADwARsQEBuP+psDUrsQICsA+wNSsA////1f7mAdwC8QImAXUAAAAnBfsAjP+pAQYFJ1APABGxAQG4/6mwNSuxAgKwD7A1KwAAAP///+n+5gGNAvACJgF2AAAAJwX6AK3/rgEGBSdQDwARsQEBuP+usDUrsQICsA+wNSsAAAAAAQBE/voCFQKfACcAKEAlHx4bEggHBgFKAAEAAAFZAAEBAGECAQABAFEBAA4MACcBJwMHFisTIiY1NDY2NxcGFRQWMzI2NzY3Jy4DNTQ2NxYWFwcTFhYVFAYHBvlVYAMICCULSUEwZicoEBoKGBYNHyIHFAkXIgoKTURD/vpiWhQiKyEJLhlJUiEdHSKhQJSNbBkiJQcWRB0U/shLdiw9cSMkAAABAET++gJgAp8AMQA5QDYsKCcIBwUDAhIBAQMCTCAfHAMCSgABBAEAAQBlAAICA2EAAwMPA04BACsqJSQODAAxATEFBxYrEyImNTQ2NjcXBhUUFjMyNjc2NycmJicmJjU0NjcWFhcHFhYXFjMyFRUUIyInFxQGBwb5VWADCAglC0lBMGYnKBAaFRwHBwYfIgUUCxcWGAIgKQgIJR8BTURD/vpiWhQiKyEJLhlJUiEdHSKhgLExMkIQIiUHFkIfFMXaGBEIPAgQIT1xIyQAAAAAAf/4AAABBgKfACEAMEAtHxsDAgQCAAFMFhANCgcFAEoBAQAAAmEEAwICAg8CTgAAACEAIR4dGRcVBQcXKzEiNTU0MzI3JiYnJiYnNjY3FhYXFBYXFjMyFRUUIyInBgYICDc2BhAJAwUCDyETAwYCAgEtOwkJSSolOgg8CBo1vokhQyEQGw1F15MiRSIbCDwIJBUPAAEANQAAAnMCnwA3ACtAKB8eAgEAAUwoFxANBABKAAAAAWEDAgIBAQ8BTgAAADcANiIhHBsEBxYrMyc2Njc2Njc2Njc2Njc2NjcWFhceAhcWFhcWMzIVFRQjIiYmJyYmJzQmJicOAgcGBgcGBwYGOAM4WiMjNhUYIAoLCgITIBAEBAEBAgIBCRQLGyUICCE/KgQGCQMEAwIDCw4HChYOK0M3XyMOGRAPKR0hYDo6gUIUGAxLjkdAXE4uBggECQg8CBInHjFWKQo2OxMdVVAYIDQUPgkGBQAAAAH/9gAAAJkCnwAaACZAIwMBAAEBTBANBwMBSgABAQBhAgEAAA8ATgEABgUAGgEaAwcWKyMiNTU0MzI3JiYnJiYnNjY3FhYXFhYXFAYHBgEJCTwzAwYDCRILDiETAwgGBQQBFA8SCDwIGyZKJF62WA8cDTaNWEl/NRE5GiMAAAAAAQCkAEIBwQKwAB4ABrMQAAEyKzcnNjY3NjY1NC4DNTQ2NxYWFwcUHgIVFAYHBgbnQyxOHR0jBgoKBiMfCA8NGwMCAiIcHEdCBSJDIyNNKw84REAyChseBh81IxcFMDsvBTVeKChBAP//AET++gIVA30CJgG9AAABBwWiASgBmwAJsQEBuAGbsDUrAP//AET++gIVA30CJgG9AAABBwWiASgBmwAJsQEBuAGbsDUrAP//AET++gJgA30CJgG+AAABBwWiASgBmwAJsQEBuAGbsDUrAP////YAAAEGA30CJgG/AAABBwWi/9kBmwAJsQEBuAGbsDUrAP///+wAAACeA30CJgHBAAABBwWi/88BmwAJsQEBuAGbsDUrAP//AKQAQgHBA44CJgHCAAABBwWiAO0BrAAJsQEBuAGssDUrAP//AET++gIVA14CJgG9AAABBwUjAOcBkQAJsQEBuAGRsDUrAP//AET++gIVA14CJgG9AAABBwUjAOcBkQAJsQEBuAGRsDUrAP//AET++gJgA14CJgG+AAABBwUjANMBkQAJsQEBuAGRsDUrAP////gAAAEGA14CJgG/AAABBwUj/48BkQAJsQEBuAGRsDUrAP//ADUAAAJzA2ECJgHAAAABBwXbAP4BlAAJsQEBuAGUsDUrAP////YAAACiA14CJgHBAAABBwUj/5kBkQAJsQEBuAGRsDUrAP//AKQAQgHBA3UCJgHCAAABBwUjAK4BqAAJsQEBuAGosDUrAP//AET++gIyA4oCJgG9AAABBwUvANcBfgAJsQEDuAF+sDUrAP//AET++gIyA4oCJgG9AAABBwUvANcBfgAJsQEDuAF+sDUrAP//AET++gJgA4oCJgG+AAABBwUvAMMBfgAJsQEDuAF+sDUrAP///9oAAAEGA4oCJgG/AAABBwUv/38BfgAJsQEDuAF+sDUrAP//ADUAAAJzA48CJgHAAAABBwUvAOcBgwAJsQEDuAGDsDUrAP///+QAAADkA4oCJgHBAAABBwUv/4kBfgAJsQEDuAF+sDUrAP//AKQAQgH2A48CJgHCAAABBwUvAJsBgwAJsQEDuAGDsDUrAP//AET+GQIVAp8CJgG9AAABBwUuADz/HwAJsQEDuP8fsDUrAP//AET+GQJgAp8CJgG+AAABBwUuADz/HwAJsQEDuP8fsDUrAP////j/CwEHAp8CJgG/AAABBgUurhEACLEBA7ARsDUr//8ANf8LAnMCnwImAcAAAAEHBd0A3AARAAixAQOwEbA1KwAAAAT/+P8LARsCnwAZACMALgA3ADJALwMCAgEAAUwPDAkDAEo2NDEtKiciHx0JAUkAAAABYQIBAQEPAU4AAAAZABgVAwcXKzEiNTU0MzI2NjcmAic2NjceAhUUBw4DFxYWFwYHJiYnNicWFhcGBgcmJic2FxYXBgYHJic2CAgoW1IZCRkQDiITBAwKIgU3T0/CFB4KGCIFHBYSbBIdDQkdFA4bDhNqGx0PGwsTJBsIPAgHDQh5AQCGDxwNUb28TiNBCg0IBC0QGwwWJwYbFRwVDRwOCB4XDhwNHjMRIQ4bDhkeIQAA//8ApP8LAgMCsAImAcIAAAEHBd0AqgARAAixAQOwEbA1KwAA//8ARP76AlcCnwImAb0AAAAHBhABLQAA//8ARP76AmACnwImAb4AAAEHBhABLQAKAAixAQKwCrA1KwAA////+AAAAQYCnwImAb8AAAEGBg4A9gAJsQECuP/2sDUrAAAA////9gAAAP0CnwImAcEAAAEGBg4A9gAJsQECuP/2sDUrAAAA//8ApABCAfsCsAImAcIAAAAHBiUBKgAAAAEARP7PAcIBXQAmADBALRoSEQMAAQFMJB4AAwNJAAIAAQACAWkAAAMDAFkAAAADYQADAANRFSUiGgQHGisTNCYnJjU0Njc2NjMnJiMiBgcnNjYzMhcWFhcHIgYHFhYXFhYXBgZXBQUJEw03c1gVPy0YKxogGkwqKjEgNRomapA4BQoFCg4EDiH+zyZeN2UWEzAaERAfaCIoEERJQS07FGYTEx46HTFVJQoTAAACAET+zwIjAVQANAA9ADhANTw4LSsGBQEAJAECAQJMMgACAkkAAAACYQMBAgIPTQABAQJhAwECAg8CTiYlIyEeHBQSBAcWKxMmJicmJjU0Njc2Nz4DNzY2MzIXFhYXFhYXFjMyFRUUIyInByImJjU0NwYHFhcWFhcGBgEmJicOAgcWYwQJBwUGAwIGGAIcKCkQLEMVExIKFg0MGAoXFwsLMSoMRXVIEy4mBxoFDAYPIQE0FiMFHDcrCkH+zzl1PTBSIQYPCRIzBRccHAoaHB8QMiMiMhAgCjcLKyseNSIgJR8dU5snTicKEwFuI2o2DiswFTcAAv/4AAABrgFUACcAMAA1QDIvKw0DAgMnCAcCBAACAkwAAwMAYQUBAgAAD00EAQICAGEFAQIAAA8ATiMpKRUTEAYHHCshIicGBiMiNTU0MzI2NyY1NDY2NzYzMhYXFhYXFhYXFjMyFRUUIyInJyYmJw4CBxYBO15QMD0gCAgRIBAHHC8dLCcRGg0HDAYMHBARFAsLMikTFiIGGTgtC0MeEgwIPAgGBA4RHEZDFyMnHg4gESI6FBQKNwssGSBoOgsrMhY3AAL/+AAAAZkBVAAdACkAMEAtKAECBAoJBAMAAgJMHgECAUsAAwAEAgMEaQACAgBhAQEAAA8ATisoFRUQBQcbKyEiJiYnBgYjIjU1NDMyNjc+Ajc2NjMyFhcWFRQGJyYmJyYjIgcGBgcWAWYoWFIfFzgmCAgmNBMSFRIOGTQdHToSEhYzARoQCwwlIAgXDkgPGxIgHAg8CCUlISsiEiEdSDg3NR8yNS5ZFxEtDCcdJgD//wBE/s8BwgJWAiYB4gAAAQYFLxJKAAixAQOwSrA1K///AET++gIXAXUCJgHvAAABBgUjS6gACbEBAbj/qLA1KwAAAP//AET++gJbAXUCJgHyAAABBgUjS6gACbEBAbj/qLA1KwAAAP////gAAAE1AcwCJgAhAAABBgUj1P8ACbEBAbj//7A1KwAAAP////gAAADhAkQCJgAiAAABBgUjwHcACLEBAbB3sDUr//8ARP4/AhcBdQImAe8AAAAmBSNLqAEHBSUAVv8BABKxAQG4/6iwNSuxAgG4/wGwNSsAAP//AET+PwJbAXUCJgHyAAAAJgUjS6gBBwUlAFb/AQASsQEBuP+osDUrsQIBuP8BsDUrAAD////4/zEBNQHMAiYAIQAAACYFI9T/AQYFJerzABKxAQG4//+wNSuxAgG4//OwNSv////4/zEA4QJEAiYAIgAAACYFI8B3AQYFJdbzABGxAQGwd7A1K7ECAbj/87A1KwAAAQBE/voCFwDiACEAJkAjFxEIBwQBSgABAAABWQABAQBhAgEAAQBRAQAODAAhASEDBxYrEyImNTQ2NjcXBhUUFjMyNjY3JiY1NDY3FhYXFhYVFAYHBvlVYAMICCULSUEyZVAQISAdJAIMCQcHTkRD/vpiWhQiKyEJLhlJUiE5I0dgGBopFhFBLyQ5FT1xIyT////4AAABNQDkAgYAIQAA////+AAAAOEBdgIGACIAAAABAET++gJbAOIAKAA2QDMjHwgHBAMCEQEBAwJMFwECSgABBAEAAQBlAAICA2EAAwMPA04BACIhHRsODAAoASgFBxYrEyImNTQ2NjcXBhUUFjMyNjY3JiY1NDY3FhYXFjMyFRUUIyInFxQGBwb5VWADCAglC0lBMmVQECEgHSQCCwokJQkJIRsBTkRD/vpiWhQiKyEJLhlJUiE5I0dgGBopFhBCMxEIPAgMHT1xIyQA//8ARP76AhcBrQImAe8AAAEGBdZ3gQAJsQECuP+BsDUrAAAA//8ARP76AlsBrQImAfIAAAEGBdZ3gQAJsQECuP+BsDUrAAAA//8ARP6eAhcBdQImAe8AAAAmBSNLqAEHBegAwPzZABKxAQG4/6iwNSuxAgK4/NmwNSsAAP//AET+ngJbAXUCJgHyAAAAJgUjS6gBBwXoAMD82QASsQEBuP+osDUrsQICuPzZsDUrAAD////4/7ABNQHMAiYAIQAAACYFI9T/AQcF6AAc/esAErEBAbj//7A1K7ECArj967A1KwAA////+P+5AOECRAImACIAAAAmBSPAdwEHBegAP/30ABGxAQGwd7A1K7ECArj99LA1KwAAAP//AET++gIXAaECJgHvAAABBgUvO5UACbEBA7j/lbA1KwAAAP//AET++gJbAaECJgHyAAABBgUvO5UACbEBA7j/lbA1KwAAAP////j/CwE1AOQCJgAhAAABBgUu0BEACLEBA7ARsDUr////+P8LARUBdgImACIAAAEGBS68EQAIsQEDsBGwNSsAAgBEAAABZQGYABcAIgAmQCMdGwwGBAFKAwEBAQBhAgEAAA8AThkYAQAYIhkiABcBFwQHFiszIiY1NDY3JiY1NDY3FhYXFhUUDgIHBicyNjcmJwYGFRQWvjdDLzQQCQ8LLVYdHQ0VFwspLxtAEBVcJSwuQjwvXjEJGAsQHAQQSjMyKQ8sLiQHHFAdD2NHIkwdIygAAAIARAAAAdUBqwAiAC0AMkAvHAQCAQIhBQIAAQJMKCYZAwJKAwECAQKFAAEBAGEAAAAPAE4jIyMtIywfHRAEBxcrISImJycHLgInLgI1NDY2NzY3NjY3JzY3FhYXFjMyFRUUJyYmJwYHFhYXFhYBzDxLBgQkJCkaDxIrIA0RBggRGlArBCMmAgYBLzkJnAYNB3czGz8jESQuKRVFBAUIBgcXGwwJJSUKDhEZMxEbHxVRolEbCDwIeCdTKShBGBgFAgIAAAAD//j/NAGOAacALQA3AEMASUBGGgkIAwECQTskBAMFBQECTDUxEAMCSggBBQcBAAUAZQMBAgIBYQYEAgEBDwFOOTgBAEA+OEM5QyMhHhsOCwYFAC0BLQkHFisXIiYnBycHIjU1NDM2Njc2NxcWFhcWFhUUBgc2NjMyFRUUIyIHFhYVFAYGBwYGAzY2NycmJicGBhMyNjcuAiMiBxYW/01iDiYKEggIESAQFackAggHBwYyIClFKAsLKR0YIQ4QBQ0umiRHHAgKCwI3LYYRJgsIMUIkCR4QRMxcVRAsAQg8CAEBAbqeChY7JSIvDixEEwQDCjcLBwopFAkpJwYNEgEXBB8YMy9FFk5y/vcODR0xHQZAQAAD//j/7gHJAbEAGwAoADQALkArMicNAwECKQgHAgQAAQJMLhQCAkoAAgEChQABAQBhAAAADwBOIiAVFAMHGCsFJicGBiMiNTU0MzI2NyY1NDY2NzcWFhcWFRQGJzY1NCYjIgYHBgYHFgU0JicmJxYVFAcWFgGVglsoWDgICCM6Fj4XKBkXVYUlJh3vMyUWECUQCAsFLwEMIh4dIhYyIk4SDy8XFQg9BwYHNkAfQTUMQRhhQEFDFkppJ0MdMRoXCxsOMmcsYSspHSgpSjoPFAD//wBEAAABZQJ5AiYB/QAAAQ4FeXeLRmYACbECAbj/i7A1KwD//wBEAAAB1QJ5AiYB/gAAAQ8FeQCk/4tGZgAJsQIBuP+LsDUrAAAA//8ARAAAAWUBmAIGAf0AAAABAEQAAAG6ALsAGwAjQCABAQIBAUwAAAADAQADaQABAQJhAAICDwJOJiMjJQQHGiszJzY2NzYzMhYXFjMyFRUUIyImJyYmJyYjIgcGXhofORssJRQkFiovCwsiMhQKGhEZExkgIRctQhQhGh04CjcLEA0HGRMbGRkAAAAAAf/4/yUBnACHACcALkArEgEBAgkIAwMAAQJMJxcCAEkAAgEChQMBAQEAYQQBAAAPAE4jKyMVFQUHGysXJiYnBgYjIjU1NDMyNjc2MzIXBhUUFhc3NjY3NjMyFRUUIyIHBgYVwDxHBQwhCwgIIxkGDxgTDQsqKAULNSgkMAsLUyEPB9sueUMGCQg8CA0OIBIcHjBhKh44UBQSCjcLDiBRLgAAAAAC//j/LADhAXYAGQAmADdANAMCAgEAJCAfAwIBAkwQCAIASgQBAgEChgAAAAFhAwEBAQ8BThoaAAAaJhomABkAGRUFBxcrMSI1NTQzMjY3JiYnJjU0NjcWFhcWFRQGBwYXJiY1NDcXBhUUFwYGCAguXyQHEgsTIx4CDAkPExFbIBQaPBsgKQkgCDwIFRAXNyAyERYuEA45LEIkDzcgN9QCHBM8ThEsJRkNEiH//wBEAAABZQJ5AiYB/QAAAQ4FeXeLRmYACbECAbj/i7A1KwD//wBEAAABugHFAiYCBAAAAQ8FeQCi/tdGZgAJsQEBuP7XsDUrAAAA////+P8lAZwBkwImAgUAAAEPBXkAg/6lRmYACbEBAbj+pbA1KwAAAP////j/LADhAlsCJgIGAAABDwV5AEL/bUZmAAmxAgG4/22wNSsAAAAAAwAR/+4CRAGxABkAJwAxAC5AKzAsJQwHBQECKAYCAwABAkwTAQJKAAIBAoUAAQEAYQAAAA8ATiAeJCMDBxgrBSYnBiMiJzcWFjMyNyY1NDY2NzcWFhcWFRQlNjU0JiMiBgcGBgcWFhc0JiYnFhUUBxYCEYFdSlIrWwkVOCIyMz4XKBkXU4clJv70MyUWECUQCAsFFjPyIDgmFjRJEg8vLDIiBAQONEEfQTUMQRdhQUFDMDkpQR0xGhcLGw4ZKVcsXFUgJStLOh0AAAAAAwARAAACiwG/ADEAPQBRALZLsBJQWEAUTEc7CwUFAQUvKQQDAAECTBQBBUobQBRMRzsLBQUBBS8pBAMEAQJMFAEFSllLsBJQWEAWAAUBBYUIBgIDAQEAYQQDBwMAAA8AThtLsCRQWEAiAAUBBYUIBgIDAQEEYQAEBA9NCAYCAwEBAGEDBwIAAA8AThtAGwAFAQWFAAQAAQRZCAYCAwEBAGEDBwIAAA8ATllZQBk/PgEAPlE/UTk3LSsnJSIgCQcAMQExCQcWKzMiJiYnNxYWMzI2NycmJjU0NjY3NxYWFxYWFRQGFRQXFjMyFRUUIyImJwYGIyImJwYGNzY2NTQmIyIGBxYWFzI2NTQmJyYmJxYWFRQHFhYXFhaVEDExEgkjKhgdSBQQGSAWKBwXP2woKC0FDQ0iDAwtMAsLKA4gViIgWXYaFB8ZGToLFTPWEBsZFBU0HA8KLAUYDQ8oEBgKIgQECQsNFz0fGz43DkEQOCUlVCwXHQwRBwgJOgoPFQsLGBAZHY4VMx4bNDYtFipVDhMnQxsbKQ0bIh1FNwEMBQYJAAAAA//4AAACEgG/AC8AOwBPALhLsBJQWEAUSkU5CQQABSwmAwIEAgACTBEBBUobQBdKRTkJBAAFLCYDAwMAAgECAwNMEQEFSllLsBJQWEAWAAUABYUIBgEDAAACYQcEAwMCAg8CThtLsCRQWEAiAAUABYUIBgEDAAADYQADAw9NCAYBAwAAAmEHBAICAg8CThtAGwAFAAWFAAMCAANZCAYBAwAAAmEHBAICAg8CTllZQBg9PAAAPE89Tzc1AC8ALiooJSIfHSUJBxcrMSI1NTQzMzI2NyYmNTQ2Njc3FhYXFhYVFAYVFBcWMzIVFRQjIyInBgYjIiYnBgYjNzY2NTQmIyIGBxYWFzI2NTQmJyYmJxcWFRQHFhYXFhYICBUYQhAfKRYoGxc/bCgoLQUNDSMLCwxKEgwoDh9YISRaMq4WGCAZGTkLFDTVEBsZFBQ1HAsOLAUZDBAnCD0IDAcYRSMbPjcOQRA4JSVULBcdDBAICAk6CiQLCxgQGR2OFDIgGzQ2LRYqVQ4TJ0MbGykNEx0qRzUBDAUGCQD////4/+4ByQGxAgYCAAAA//8AEf/uAkQCoQImAgsAAAEHBaMAwACpAAixAwGwqbA1KwAA//8AEQAAAosCoQImAgwAAAEHBaMAwACpAAixAwGwqbA1KwAA////+AAAAhICoQImAg0AAAEHBaMARwCpAAixAwGwqbA1KwAA////+P/uAckCoQImAgAAAAEHBaMARwCpAAixAwGwqbA1KwAA////+P/uAckDcQImAgAAAAEHBXcApADjAAixAwGw47A1KwAA//8ALgAAAWUCVQImAf0AAAEHBSn/5QCCAAixAgKwgrA1KwAA//8AUwAAAeQCVQAmAf4PAAEHBSkAIQCCAAixAgKwgrA1KwAA//8APQAAAXQCVQAmAf0PAAEHBSn/9ACCAAixAgKwgrA1KwAA//8ARAAAAboBoQImAgQAAAEGBSkQzgAJsQECuP/OsDUrAAAAAAIAIv81AaEBVAAgADAASEBFJAEEBQsBAgQFAQECBAEAAQRMAAMABQQDBWkAAQYBAAEAZQcBBAQCYQACAg8CTiIhAQAqKCEwIjAWFA4MCAYAIAEgCAcWKxciJiYnNxYzMjY3JwYjIiY1NDY3NjMyFhcWFhUUBgYHBhMyNjcmJicmIyIGBwYVFBbBFT07Eg41RUtrHAE8SjpDKSEhIidHFwwLHTMiNBIZKSMMLBoVFRUoDg06yxMeDyETWVUIOUc6M2MfHkVAID8fL15PGiYBHQgOL0gRDx8YGBMkJwACACL/NQHlAVwAJgA1AEtASB8eAgIEBQEBAgQBAAEDTAADAAYEAwZpAAEIAQABAGUJBwIEBAJhBQECAg8CTicnAQAnNSc0LCoiIRwbFhQLCggGACYBJgoHFisXIiYmJzcWMzI2NyYmJyYmNTQ2NzYzMhYXFhYVMzIVFRQjIwYGBwYTLgIjIgYHBhUUFhcWFsEVPTsSDjVFPmAhMUkYNSopISEiKUgWCwo8CAhGD0AsK4gIKDYdFSkNDTQtFzvLEx4PIRM9QQIIBw41NDNkHh5KQSFBIwg8CDlfGhkBGDRUMiAXFxMjJQgEAwAA//8AIv81AaEBVAImAhgAAAEHBegAcP2UAAmxAgK4/ZSwNSsA//8AIv81AeUBXAImAhkAAAEHBegAcP2VAAmxAgK4/ZWwNSsA//8AIv81AaECDgImAhgAAAEGBSklOwAIsQICsDuwNSv//wAi/zUB5QIOAiYCGQAAAQYFKSU7AAixAgKwO7A1K///ACL/NQGhAhsCJgIYAAABBgUjNU4ACLECAbBOsDUr//8AIv81AeUCGwImAhkAAAEGBSM1TgAIsQIBsE6wNSv//wAi/zUBoQIyAiYCGAAAAQ8FeQC3/0RGZgAJsQIBuP9EsDUrAAAA//8AIv81AeUCMgImAhkAAAEPBXkAt/9ERmYACbECAbj/RLA1KwAAAP//////NQGhAfMCJgIYAAABBwYA/9b/HAAJsQIBuP8csDUrAP//////NQHlAf0CJgIZAAABBwYA/9b/JgAJsQIBuP8msDUrAP//AAD/NQGhAfMCJgIYAAABBwYB/+r/HAAJsQIBuP8csDUrAP//AAD/NQHlAf0CJgIZAAABBwYB/+r/JgAJsQIBuP8msDUrAP//ABT/EwGhAVQCJgYmAAAABgYqAAAAAP//ABT/EwHlAVwCJgYnAAAABgYqAAAAAAACACL/NQGhAVQAMAA+AFtAWDQBBgcbAQQGGBACAgMPCQUDAQIEAQABBUwABQAHBgUHaQADAAIBAwJpAAEIAQABAGUJAQYGBGEABAQPBE4yMQEAODYxPjI+JiQeHBQSDQsIBgAwATAKBxYrFyImJic3FjMyNyYmIyIGByc2NjMyFxYWFzY3JwYjIiY1NDY3NjMyFhcWFhUUBgYHBhMyNjcmJiMiBgcGFRQWwRU9OxIONUU/MSEfDwkYDhAaFw8WJREdCycVATxKOkMpISEiJ0cXDAsdMyI0EhkpIxFGJRUoDg06yxMeDyETIA0JBwYZFAsPBwkCK0MIOUc6M2MfHkVAID8fL15PGiYBHQgOQlUfGBgTJCcAAAAAAgAi/zUB5QFcADYARQBeQFsvLgIEBhgQAgIDDwkFAwECBAEAAQRMAAUACAYFCGkAAwACAQMCaQABCgEAAQBlCwkCBgYEYQcBBAQPBE43NwEAN0U3RDw6MjEsKyYkGxoUEg0LCAYANgE2DAcWKxciJiYnNxYzMjcmJiMiBgcnNjYzMhcWFhc2NyYmJyYmNTQ2NzYzMhYXFhYVMzIVFRQjIwYGBwYTLgIjIgYHBhUUFhcWFsEVPTsSDjVFPzEhHw8JGA4QGhcPFiURHQsVFDFJGDUqKSEhIilIFgsKPAgIRg9ALCuICCg2HRUpDQ04NQ44yxMeDyETIA0JBwYZFAsPBwkCGCYCCAcONTQzZB4eSkEhQSMIPAg5XxoZARg0VDIgFxcTIyUKAgMAAAD//wAi/zUBoQI/AiYCGAAAAQYF82QFAAixAgGwBbA1K///ACL/NQHlAj8CJgIZAAABBgXzZAUACLECAbAFsDUr//8AIv81AaECLgImAhgAAAEHBdgAu//QAAmxAgK4/9CwNSsA//8AIv81AeUCLgImAhkAAAEHBdgAu//QAAmxAgK4/9CwNSsA//8AIv81AaECggImAhgAAAEHBXcA5//0AAmxAgG4//SwNSsA//8AIv81AeUCggImAhkAAAEHBXcA5//0AAmxAgG4//SwNSsA//8AIv81AaECKgImAhgAAAEHBaMAmAAyAAixAgGwMrA1KwAA//8AIv81AeUCKgImAhkAAAEHBaMAmAAyAAixAgGwMrA1KwAA//8AIv81AaECRwImAhgAAAEGBS8lOwAIsQIDsDuwNSv//wAi/zUB5QJHAiYCGQAAAQYFLyU7AAixAgOwO7A1KwABAET++gJIAVQAOwA2QDMuJSQTCAcGAQMBTAACAAMBAgNpAAEAAAFZAAEBAGEEAQABAFEBACknIiAODAA7ATsFBxYrEyImNTQ2NjcXBhUUFjMyNjY3NjcmJicmJicmNTQ2Njc2MzIWFwcmJiMiBgYHBgcWFxYWFxYWFRQGBgcG+VVgAwgIJQtJQSRUUh8wDw0ZDkxKFBMdNCAxJyIiCSQHDQwRMTEVHwcbUyQyDxENM1U1Tv76YloUIishCS4ZSVIRHxQeIQIEAQYZHBwtH0pJHCsuKQoKBRkrGiojGgwFDQoMIRceRD4WIQAAAAEARP77AsAAXQAyADFALicIBwMDAhMBAQMCTAABBAEAAQBlAAICA2EAAwMPA04BACUjIB4ODAAyATIFBxYrASImNTQ2NjcXBhUUFjMyNjc2NjcmJicmNTQ2NjMWFjMyFRUUIyImJxYWFxYWFRQGBgcGARtncAMICCULXVBbci0WKBI+SRQXDREGI3hACwsiVyEWLRYPCDJZO1n++2JaFCIrIQkuGUlSFBMJGBAhKw4SBwkhGgcKCzcKCggNGAwJHhQZNjASGgD//wBE/voCSAHLAiYCNAAAAQcFdwDB/z0ACbEBAbj/PbA1KwD//wBE/vsCwAHLAiYCNQAAAQcFdwE5/z0ACbEBAbj/PbA1KwD////XAAAAsgKoAiYF6gAAAQYGAK7RAAmxAQG4/9GwNSsAAAD////PAAABDgKoAiYF6wAAAQYGAKbRAAmxAQG4/9GwNSsAAAD///+5AAAB0gKoAiYABgAAAQYGAJDRAAmxAQG4/9GwNSsAAAD////WAAAAsgKpAiYF6gAAAQYGAcDSAAmxAQG4/9KwNSsAAAD////OAAABDgKpAiYF6wAAAQYGAbjSAAmxAQG4/9KwNSsAAAD///+5AAAB0gKpAiYABgAAAQYGAaPSAAmxAQG4/9KwNSsAAAD//wBE/k4CSAFUAiYCNAAAAQcFKgBc/xAACbEBArj/ELA1KwD//wBE/k4CwABdAiYCNQAAAQcFKgBw/xAACbEBArj/ELA1KwD////4/zsBeQDkAiYF7gAAAQYFKgj9AAmxAQK4//2wNSsAAAD////4/zsBNAF2AiYF7QAAAQYFKvH9AAmxAQK4//2wNSsAAAD//wBD/voCSAFUAiYCNAAAAQ8FeQBL/jRGZgAJsQEBuP40sDUrAAAA//8AQ/77AsABIgImAjUAAAEPBXkAS/40RmYACbEBAbj+NLA1KwAAAP////gAAAE1AeMCJgAhAAABDwV5AFb+9UZmAAmxAQG4/vWwNSsAAAD////4AAAA4QJbAiYAIgAAAQ8FeQBC/21GZgAJsQEBuP9tsDUrAAAA//8ARP76AkgBpQImAjQAAAEHBfMAPf9rAAmxAQG4/2uwNSsA//8ARP77AsABWgImAjUAAAEHBfMAq/8gAAmxAQG4/yCwNSsA////+P87AXkB8AImBe4AAAAmBfMrtgEGBSoI/QASsQEBuP+2sDUrsQICuP/9sDUr////+P87ATQCaAImBe0AAAAmBSrx/QEGBfMnLgARsQECuP/9sDUrsQMBsC6wNSsA//8AQ/5OAkgBVAImAjQAAAAnBSoAXP8QAQ8FeQBL/jRGZgASsQECuP8QsDUrsQMBuP40sDUrAAD//wBD/k4CwAEiAiYCNQAAACcFKgBw/xABDwV5AEv+NEZmABKxAQK4/xCwNSuxAwG4/jSwNSsAAP////j/OwF5AeMCJgXuAAAAJgUqCP0BDwV5AH7+9UZmABKxAQK4//2wNSuxAwG4/vWwNSv////4/zsBNAJbAiYF7QAAACYFKvH9AQ8FeQB6/21GZgASsQECuP/9sDUrsQMBuP9tsDUr//8ARP5OAkgBVAImAjQAAAAnBSoAXP8QAQcFI//J/z4AErEBArj/ELA1K7EDAbj/PrA1K///AET++gJIAZwCJgI0AAABDwXzAWIDIsAAAAmxAQG4AyKwNSsAAAD//wBE/vsCwAFMAiYCNQAAAQ8F8wG9AtLAAAAJsQEBuALSsDUrAAAA////+P89AXkB2wImBe4AAAAmBSr8/wEGBaNf4wASsQECuP//sDUrsQMBuP/jsDUr////+P87ATQCUwImBe0AAAAmBSrx/QEGBaNbWwARsQECuP/9sDUrsQMBsFuwNSsA//8ARP76AkgBVAIGAjQAAP//AET++wLAAF0CBgI1AAD////4/zsBeQDkAiYF7gAAAQYFKgj9AAmxAQK4//2wNSsAAAD////4/zsBNAF2AiYF7QAAAQYFKvH9AAmxAQK4//2wNSsAAAD//wAC/voCSAFUAiYCNAAAAQcFKf+5/ysACbEBArj/K7A1KwD//wA//vsCwAEGAiYCNQAAAQcF4//2/zMACbEBArj/M7A1KwD////4/zsBeQG/AiYF7gAAACYFKezsAQYFKgj9ABKxAQK4/+ywNSuxAwK4//2wNSv////4/zsBGQI3AiYAIgAAACYFKbBkAQYFKtb9ABGxAQKwZLA1K7EDArj//bA1KwD//wAU/voCSAFUAiYCNAAAAQcFL/+5/ysACbEBA7j/K7A1KwD//wA9/vsCwAE3AiYCNQAAAQcF4f/i/ysACbEBA7j/K7A1KwD////4/zsBeQH4AiYF7gAAACYFL+zsAQYFKgj9ABKxAQO4/+ywNSuxBAK4//2wNSv////4/zsBGQJwAiYAIgAAACYFL7BkAQYFKtb9ABGxAQOwZLA1K7EEArj//bA1KwD//wBE/voCSAGmAiYCNAAAAQcGAABA/s8ACbEBAbj+z7A1KwD//wBE/vsCwAEdAiYCNQAAAQcGAAB+/kYACbEBAbj+RrA1KwD////4/zsBeQHJAiYF7gAAACYFKgj9AQcGAAAA/vIAErEBArj//bA1K7EDAbj+8rA1KwAA////+P87ARkCIwImACIAAAAmBSrW/QEHBgD/3/9MABKxAQK4//2wNSuxAwG4/0ywNSsAAP//AET++gJIAaUCJgI0AAABBwYBAFP+zgAJsQEBuP7OsDUrAP//AET++wLAAScCJgI1AAABBwYBAJH+UAAJsQEBuP5QsDUrAP////j/OwF5AckCJgXuAAAAJgUqCP0BBwYBAAD+8gASsQECuP/9sDUrsQMBuP7ysDUrAAD////4/zsBGQIjAiYAIgAAACYFKtb9AQcGAf/s/0wAErEBArj//bA1K7EDAbj/TLA1KwAA//8ARP2eAkgBVAImAjQAAAEHBgcAlvvkAAmxAQG4++SwNSsA//8ARP2fAsAAXQImAjUAAAEHBgcAn/vlAAmxAQG4++WwNSsA////+P63AXkA5AImBe4AAAEHBgcAQfz9AAmxAQG4/P2wNSsA////+P63AOEBdgAmACIAAAEHBgcABPz9AAmxAQG4/P2wNSsAAAEARP76As8BVABIADlANjsxMB8TEgsJBgkBAwFMAAIAAwECA2kAAQAAAVkAAQEAYQQBAAEAUQEANTMuLBoYAEgBSAUHFisBIiY1NDY3BgYHJjU2Njc+AjcXBgYVFBYzMjY2NzY3JiYnJiYnJjU0NjY3NjMyFhcHJiYjIgYGBwYGBxYXFhYXFhYVFAYGBwYBgFVgAwQhQiIJAgMHBzJDHyEJC0lCJFRRHy8RDRoOS0sUEx00IDEnIiIJIwcNDREwMhQQEgUbUyQxDxEOM1U1Tv76YloYKhMSJBIPEA0NCAgdIxAQHDMVSVIRHxQeIQIEAQYZHBwtH0pJHCsuKQoKBRkrGhQmExoMBQ0KCyEYHkQ+FiEAAAEAQ/77A0cAXQA7ADZAMzIGAgMCHQkCAQMCTBEQAgJKAAEEAQABAGUAAgIDYQADAw8DTgEAMC4rKRgWADsBOwUHFisBIiY1NDY3BgYHJjY3PgI3FwYGFRQWMzI2NzY2NyYmJyY1NDY3NjMWFjMyFRUUIyImJxYXFhYVFAYHBgGiZnEDBCFCIgoDCgczQx4hCgpdUFtyLRYoEj5KExcMBwoHI3hACwsiVyFXCAsHbllZ/vtiWhgqEhIkEhEjDQgeJA8RJzILSVEVEwkWECIqDxIHCSALDwcJDDcJCQgvBgkcESdQGhoAAP//AET+bgJIAVQCJgI0AAABBwXoAOz8qQAJsQECuPypsDUrAP//AET+bgLAAF0CJgI1AAABBwXoAQD8qQAJsQECuPypsDUrAP////j/fgE1AOQCJgAhAAABBwXoABz9uQAJsQECuP25sDUrAP////j/hwDhAXYCJgAiAAABBwXoAD/9wgAJsQECuP3CsDUrAP//AET9+QJIAVQCJgI0AAABBwUnAJ3/IgAJsQECuP8isDUrAP//AET9+QLAAF0CJgI1AAABBwUnALH/IgAJsQECuP8isDUrAP////j+5gF5AOQCJgXuAAABBgUnSQ8ACLEBArAPsDUr////+P7mASQBdgImBe0AAAEGBScyDwAIsQECsA+wNSv//wBE/h4CSAFUAiYCNAAAAQcFLgBC/yQACbEBA7j/JLA1KwD//wBE/h4CwABdAiYCNQAAAQcFLgBW/yQACbEBA7j/JLA1KwAAAQBEAAAC8wE6ACsAIUAeIxwTEgQBSgABAQBhAgEAAA8ATgIAIB4AKwIqAwcWKyEiJiYnJjU0NjY3Njc2Njc2NjcXBgYHBgYHBgYHFhYzMjY2NxYVFAcOAwGWL2xjJDALEAYeNAcsJSMgCh8PKCAiLhEVJBg4kWgzcW8vDiMUSFhdBw0JDAsOMS0GIBgDDw0LHxMRLC0KCg8ICRURFhQLFQ4PEhoSChEMBgAAAAEAH/80AkYATAAkADBALR0ZAgMCBQEAAwJMAAMEAQADAGUAAQECYQACAg8CTgEAHBoUEQ4MACQBIwUHFisFIiYnJjU0NjY3PgIzMhUVFCMiBgYHBgYHFjMyNxYVFAYHBgYBLEB7KSkKDAM0isCFCwtDi3orID0caHyCgg8lGihtzBMPDwsIJiIDMjwbCzcKBxIPCx0THRgSERQbBQcGAP//AEQAAALzAh4CJgJ3AAABDwV5AKn/MEZmAAmxAQG4/zCwNSsAAAD//wAQ/zQCNwFDACYCePEAAQ8FeQCL/lVGZgAJsQEBuP5VsDUrAAAA//8AUgAAAwECGgAmAncOAAEHBgAAKv9DAAmxAQG4/0OwNSsA//8ADv80AjUBUQAmAnjvAAEHBgAACf56AAmxAQG4/nqwNSsA////+P87AXkByQImBe4AAAAnBgAAAP7yAQYFKgj9ABKxAQG4/vKwNSuxAgK4//2wNSsAAP////j/OwEZAiMCJgAiAAAAJwYA/9//TAEGBSrW/QASsQEBuP9MsDUrsQICuP/9sDUrAAD//wBSAAADAQIaACYCdw4AAQcGAQA+/0MACbEBAbj/Q7A1KwD//wAO/zQCNQFRACYCeO8AAQcGAQAc/noACbEBAbj+erA1KwD////4/zsBeQHJAiYF7gAAACcGAQAf/vIBBgUqCP0AErEBAbj+8rA1K7ECArj//bA1KwAA////7P87ARkCFwImACIAAAAnBgH/1v9AAQYFKtb9ABKxAQG4/0CwNSuxAgK4//2wNSsAAP//AD//NAJJAYcCJgYoAAABBwUqAK7/9gAJsQICuP/2sDUrAP//ACL/NQHVAisCJgIYAAABBwV5AVz/gQAJsQIBuP+BsDUrAP//ACL/NQHlAh4CJgIZAAABBwV5AVz/dAAJsQIBuP90sDUrAP//ACL/NQGjAisCJgIYAAAAJgWMb70BBwV5ASr/gQASsQICuP+9sDUrsQQBuP+BsDUrAAD//wAi/zUB5QI1AiYCGQAAACYFjG/HAQcFeQEq/4sAErECArj/x7A1K7EEAbj/i7A1KwAA//8ARP76AsYCRwAmAjQAAAEPBXkCQf9ZRmYACbEBAbj/WbA1KwAAAP//AET++wLAAV0CJgI1AAABDwV5AP3+b0ZmAAmxAQG4/m+wNSsAAAD////4AAABXAHvACYAIQAAACcF7ADA/6QBBwYCASwAAAAJsQEBuP+ksDUrAP////gAAAFXAmkAJgAiAAABBwXsAMEAHgAIsQEBsB6wNSsAAP//AEQAAAFlAZgCBgH9AAD//wBE/voDwgH5AiYA9wAAAQcFMgIUAAwACLEBBLAMsDUrAAD//wBE/voECQH5AiYA+AAAAQcFMgIUAAwACLEBBLAMsDUrAAD////4AAACqAH5AiYA+QAAAQcFMgCsAAwACLEBBLAMsDUrAAD////4AAACXwH5AiYA+gAAAQcFMgCsAAwACLEBBLAMsDUrAAD//wBJ/uICdQKWAiYBMwAAAQcFKQAKAMMACLEBArDDsDUrAAD//wBE/uICHgI2AiYBNAAAAQYFKSBjAAixAgKwY7A1K/////gAAAGvAjYCJgE1AAABBgUp/GMACLECArBjsDUr////+AAAAd8CNgImATYAAAEGBSkAYwAIsQECsGOwNSv//wBJ/uICdQLGAiYBMwAAAQcFLQAAAKwACLEBA7CssDUrAAD//wBE/uICHgJmAiYBNAAAAQYFLRZMAAixAgOwTLA1K///AAAAAAG3AmYAJgE1CAABBgUt+kwACLECA7BMsDUr////+AAAAd8CZgImATYAAAEGBS32TAAIsQEDsEywNSv//wBJ/uICdQL0AiYBMwAAAQcFJgBpAO8ACLEBArDvsDUrAAD//wBE/uICHgKUAiYBNAAAAQcFJgB/AI8ACLECArCPsDUrAAD//wAAAAABtwKhACYBNQgAAQcFJgBeAJwACLECArCcsDUrAAD////4AAAB3wKUAiYBNgAAAQcFJgBfAI8ACLEBArCPsDUrAAD//wBE/08DHQHhAiYF6QAAAQcFKgHSABEACLECArARsDUrAAD//wBE/zsDQgGHAiYBTAAAAQcFKgG+//0ACbECArj//bA1KwD////4/zsBlAF2AiYBZAAAAQYFKg/9AAmxAgK4//2wNSsAAAD////4/zsBcQHQAiYBZQAAAQYFKi39AAmxAgK4//2wNSsAAAD//wBE/xcDHQHhAiYF6QAAAQcFMQGzABEACLECA7ARsDUrAAD//wBE/wMDQgGHAiYBTAAAAQcFMQGf//0ACbECA7j//bA1KwD////4/wMBlAF2AiYBZAAAAQYFMfD9AAmxAgO4//2wNSsAAAD////4/wMBcQHQAiYBZQAAAQYFMQ79AAmxAgO4//2wNSsAAAD//wBE/s8BwgIqAiYB4gAAAQYFIyJdAAixAQGwXbA1K///AET+zwIjAioCJgHjAAABBgUjTV0ACLECAbBdsDUr////+AAAAa4CKgImAeQAAAEGBSMHXQAIsQIBsF2wNSv////4AAABmQIqAiYB5QAAAQYFIzZdAAixAgGwXbA1K///AET+zwHCAV0CJgHiAAABBgUlb/MACbEBAbj/87A1KwAAAP//AET+zwIjAVQCJgHjAAABBwUlAJ3/8wAJsQIBuP/zsDUrAP////j/MQGuAVQCJgHkAAABBgUlHvMACbECAbj/87A1KwAAAP////j/MQGZAVQCJgHlAAABBgUlMvMACbECAbj/87A1KwAAAP//AET+SQIXAXUCJgHvAAAAJgUjS6gBBwUqAFb/CwASsQEBuP+osDUrsQICuP8LsDUrAAD//wBE/kkCWwF1AiYB8gAAACcFKgBW/wsBBgUjS6gAErEBArj/C7A1K7EDAbj/qLA1KwAA////+P87AXkBzAImBe4AAAAmBSP8/wEGBSoI/QASsQEBuP//sDUrsQICuP/9sDUr////+P87ATQCRAImBe0AAAAmBSrx/QEGBSP4dwARsQECuP/9sDUrsQMBsHewNSsA//8ARP76AhcCYgImAe8AAAEGBddLqAAJsQEDuP+osDUrAAAA//8ARP76AlsCYgImAfIAAAEGBddLqAAJsQEDuP+osDUrAAAA////8gAAAS8CuQAmACH6AAEGBdfO/wAJsQEDuP//sDUrAAAA////+AAAAQkC1AImBf0AAAEGBdfWGgAIsQEDsBqwNSv//wBE/voCFwJMAiYB7wAAAQYFJEuoAAmxAQK4/6iwNSsAAAD//wBE/voCWwJMAiYB8gAAAQYFJEuoAAmxAQK4/6iwNSsAAAD////4AAABNQKjAiYAIQAAAQYFJNT/AAmxAQK4//+wNSsAAAD////4AAAA8AK+AiYF/QAAAQYFJNYaAAixAQKwGrA1K///AET++gJHAp8CJgG9AAABBwYPAS3/7AAJsQEBuP/ssDUrAP//AET++gJgAp8CJgG+AAABBwYPASP/7wAJsQEBuP/vsDUrAP////gAAAEGAp8CJgG/AAABBgYNANgACbEBAbj/2LA1KwAAAP////YAAADiAp8CJgHBAAABBgYNANgACbEBAbj/2LA1KwAAAP//AKQAQgHzArACJgHCAAABBwYNAREATgAIsQEBsE6wNSsAAP//AET++gPCAiUCJgD3AAABBwUmAlMAIAAIsQECsCCwNSsAAP//AET++gQJAiUAJgD4AAABBwUmAlMAIAAIsQECsCCwNSsAAP////gAAAKoAiUCJgD5AAABBwUmAOsAIAAIsQECsCCwNSsAAP////gAAAJfAiUCJgD6AAABBwUmAOsAIAAIsQECsCCwNSsAAP//AET++gPCAmICJgD3AAABBwY6AkP/9QAJsQEEuP/1sDUrAP//AET++gQJAmICJgD4AAABBwY6AkP/9QAJsQEEuP/1sDUrAP////gAAAKoAmICJgD5AAABBwY6ANv/9QAJsQEEuP/1sDUrAP////gAAAJfAmICJgD6AAABBwY6ANv/9QAJsQEEuP/1sDUrAP//AET+4QJrAn8CJgCJAAABBwXWALcAUwAIsQECsFOwNSsAAP//AET+4gKrAmEAJgCKAAABBwXWALcANQAIsQECsDWwNSsAAP////gAAAKrAmYCJgCLAAABBwXWAIsAOgAIsQECsDqwNSsAAP////gAAAJaAmYCJgCMAAABBwXWAIsAOgAIsQECsDqwNSsAAP//AET++gPCAlICJgD3AAABBwYHAj//cQAJsQEBuP9xsDUrAP//AET++gQJAlICJgD4AAABBwYHAlT/cQAJsQEBuP9xsDUrAP////gAAAKoAlICJgD5AAABBwYHAOL/cQAJsQEBuP9xsDUrAP////gAAAJfAlICJgD6AAABBwYHAOL/cQAJsQEBuP9xsDUrAP//AET++gPCAeMCJgD3AAABBwWjAmf/6wAJsQEBuP/rsDUrAP//AET++gQJAeMCJgD4AAABBwWjAmf/6wAJsQEBuP/rsDUrAP////gAAAKoAeMCJgD5AAABBwWjAP//6wAJsQEBuP/rsDUrAP////gAAAJfAeMCJgD6AAABBwWjAP//6wAJsQEBuP/rsDUrAP//AB//NASEAp8AJwG/Ah4AAAAmAngAAAAnAWUDEwAAAQcFKQMAALwACLEEArC8sDUrAAD////4AAAA2wJuAiYC1gAAAAYFfB0AAAAAAQBEAAABagEYABcAKEAlFRQCAAIBTAABAgGFAAICAGEDAQAADwBOAQASEAgHABcBFwQHFishIicmNTQ2NzMGBhUUFhcWFjMzMhUVFCMBOHA/RQoEJAEDERcXVkYVCAgnK1gpOA0FJAggORUUGQg8CAAAAP////j/VADbAEwCJgLWAAABBgWQGMEACbEBAbj/wbA1KwAAAAAB//gAAADbAEwACwAhQB4JAwIDAQABTAAAAAFfAgEBAQ8BTgAAAAsACyUDBxcrMSI1NTQzMzIVFRQjCAjSCQkIPAgIPAgAAP////gAAADbAiwCJgLWAAAABgWIGAAAAP////gAAADbAl4CJgLWAAAABgWMLAAAAP////gAAADbAkcCJgLWAAAABgWTIgAAAP////gAAADbAqkCJgLWAAABBgWKMAoACLEBArAKsDUr////+AAAANsC5gImAtYAAAEGBY4xBQAIsQEDsAWwNSv////4AAAA2wKzAiYC1gAAAQYFkiJdAAixAQKwXbA1K/////gAAADbAkUCJgLWAAAABgWcIgAAAAACACMAAAHSArAAOABAADdAND47LRsLBQABAgECAAJMKSglFAQBSgABAAGFAwEAAAJhBAECAg8CTgAAOjkAOAA3FyUFBxgrMyI1NDY2MzM+AjcmJicHJjU0NjceAhcWFhc2NTQmJyYmNTQ3FhYXBxcXFAceAxUUBgYHBgY3NjcmJicGBokDBQgCFS04KBQtcDYlMAwWHENGHhcwGRMLCwUGQwcTChsBBiYKGRgQCw4FL5UJWFoLIBUVOgsJHxkjMCsaNF0fAi0pDx0LEzY9HhY0HSwqFVxIIigEMwwfOx0XFY9MSRAsLSQICh8ZAQgJTgMPFC8aHTgAAAEAIgAAAnMCnwA6ADdANDAjFQUEAQAoJwICAQJMMyAaAwBKAAABAIUAAQECYQQDAgICDwJOAAAAOgA6KyolJBgFBxcrMyc2NzY3JiYnIyYmNTQ3HgIXFhYXNjY3Njc2Njc2NjcWEhcWMzIVFRQjIiYnJiYnJiYnBwYGBwYHBjgDczAvISVbNSUUGCYoUEYWBhMNJioNDQUFDAcKFQwFBwMtOwgIO04FBAYDAgcFAwcRCC5uXSMcExMYU4YzFSgZJRIZW244EDQkJ3ZTUpEGCwYIEQiC/uedGwg8CC8oHTodHWVIG0JgHaEOCwD//wAiAAACcwLqAiYC3wAAAQcFeQDBAEAACLEBAbBAsDUrAAD///+5AAACcwKoAiYC3wAAAQYGAJDRAAmxAQG4/9GwNSsAAAD///+5AAACcwKpAiYC3wAAAQYGAaPSAAmxAQG4/9KwNSsAAAD//wAiAAACcwNhAiYC3wAAAQcF2wD+AZQACbEBAbgBlLA1KwD//wAjAAAB9gOPAiYC3gAAAQcFLwCbAYMACbECA7gBg7A1KwD//wAiAAACcwOPAiYC3wAAAQcFLwDnAYMACbEBA7gBg7A1KwD//wAj/wsCAwKwAiYC3gAAAQcF3QCqABEACLECA7ARsDUrAAD//wAi/wsCcwKfAiYC3wAAAQcF3QDcABEACLEBA7ARsDUrAAD////3AAAB0gL2AiYC3gAAAQYFef5MAAixAgGwTLA1K///ABAAAAJzAuoCJgLfAAABBgV5F0AACLEBAbBAsDUr//8AI/8KAdICsAImAt4AAAEGBXtYnwAJsQIBuP+fsDUrAAAA//8AIv8KAnMCnwImAt8AAAEGBXtwnwAJsQEBuP+fsDUrAAAA////1wAAAdICsAImAt4AAAEGBZ4bKAAIsQIBsCiwNSv////cAAACcwKfAiYC3wAAAQYFniAcAAixAQGwHLA1K////+EAAAHSAskCJgLeAAABBgXvFD4ACLECArA+sDUr////+gAAAnMCvgImAt8AAAEGBe8tMwAIsQECsDOwNSv//wBE/uIDewF2ACcAJgKaAAAABgByAAD////4/zEDewF2ACcAJgKaAAAABgBzAAD//wBE/uIDewF2ACcAJgKaAAAABgCKAAD////4/zEDewF2ACcAJgKaAAAABgCLAAD//wBE/k4GfgFUACcAJQVJAAAAJwCLAq8AAAAGAj8AAP//AET+4gN7AikAJwAmApoAAAAGAKIAAP////j/MQN7Ai4AJwAmApoAAAAGAKMAAP//AET+TgZ+Ai4AJwAlBUkAAAAnAKMCrwAAAAYCPwAA////7P8xAskBVQAnACUBlAAAAAYA1AAA////7P8xAskCGgAnACUBlAAAAAYA1gAA//8ARP7PAvEBdgAnACYCEAAAAAYB4wAA//8ARP7PA0UBVAAnACUCEAAAAAYB4wAA////+P8xAtIBVAAnACUBnQAAAAYB5AAA////+P8xAn4BdgAnACYBnQAAAAYB5AAA//8ARP76A34BdQAnACUCSQAAAAYB6AAA////+P8xArIBpwAnACUBfQAAAAYB/wAA////+P8xAl4BpwAnACYBfQAAAAYB/wAA//8ARP77A5ABdgAnACYCrwAAAAYCNQAA//8ARP77A+QA5AAnACUCrwAAAAYCNQAA//8ARP5OA5ABdgAnACYCrwAAAAYCPwAA//8ARP5OA+QA5AAnACUCrwAAAAYCPwAA//8ARP7iA6gCNwAnADgCmgAAAAYAcgAA////+P9IA6gCNwAnADgCmgAAAAYAcwAA////+P9IBUUCNwAnADgENwAAACcAcwGdAAAABgHkAAD//wBE/vsGwgG/ACcANwVJAAAAJwBzAq8AAAAGAjUAAP//AET+TgbCAb8AJwA3BUkAAAAnAHMCrwAAAAYCPwAA//8ARP7iA6gCNwAnADgCmgAAAAYAigAA////+AAAA6gCNwAnADgCmgAAAAYAiwAA//8ARP7iBq0BvwAnADcFNAAAACcAiwKaAAAABgByAAD////4/0gGQgI3ACcAOAU0AAAAJwCLApoAAAAGAHMAAP////gAAAVFAjcAJwA4BDcAAAAnAIsBnQAAAAYB5AAA//8ARP7iA6gCNwAnADgCmgAAAAYAogAA////+AAAA6gCNwAnADgCmgAAAAYAowAA////+AAABUUCNwAnADgENwAAACcAowGdAAAABgHkAAD//wBE/vsGwgIuACcANwVJAAAAJwCjAq8AAAAGAjUAAP//AET+TgbCAi4AJwA3BUkAAAAnAKMCrwAAAAYCPwAA////7P80Aw0BvwAnADcBlAAAAAYA1AAA////7P80Aw0CGgAnADcBlAAAAAYA1gAA//8ARP7PAx4CNwAnADgCEAAAAAYB4wAA//8ARP7PA4kBvwAnADcCEAAAAAYB4wAA////+AAAAxYBvwAnADcBnQAAAAYB5AAA////+AAAAqsCNwAnADgBnQAAAAYB5AAA////+P9IBUUCNwAnADgENwAAACcB5AKaAAAABgBzAAD////4AAAFRQI3ACcAOAQ3AAAAJwHkApoAAAAGAIsAAP////gAAAVFAjcAJwA4BDcAAAAnAeQCmgAAAAYAowAA//8ARP77BcUBvwAnADcETAAAACcB5AKvAAAABgI1AAD//wBE/k4FxQG/ACcANwRMAAAAJwHkAq8AAAAGAj8AAP//AET++gPCAb8AJwA3AkkAAAAGAegAAP////j/NAL2Ab8AJwA3AX0AAAAGAf8AAP////j/NAKLAjcAJwA4AX0AAAAGAf8AAP//AET++wO9AjcAJwA4Aq8AAAAGAjUAAP//AET++wQoAb8AJwA3Aq8AAAAGAjUAAP//AET+TgO9AjcAJwA4Aq8AAAAGAj8AAP//AET+TgQoAb8AJwA3Aq8AAAAGAj8AAP//AET+4gOlAnAAJwBEApoAAAAGAHIAAP///+z/NAMNAfgAJwBDAZQAAAAGANQAAP///+z/NAMNAhoAJwBDAZQAAAAGANYAAP//AET+zwMbAnAAJwBEAhAAAAAGAeMAAP//AET+zwOJAfgAJwBDAhAAAAAGAeMAAP////gAAAMWAfgAJwBDAZ0AAAAGAeQAAP////gAAAKoAnAAJwBEAZ0AAAAGAeQAAP//AET++gPCAfgAJwBDAkkAAAAGAegAAP////j/NAL2AfgAJwBDAX0AAAAGAf8AAP//AET++wO6AnAAJwBEAq8AAAAGAjUAAP//AET++wQoAfgAJwBDAq8AAAAGAjUAAP//AET+TgO6AnAAJwBEAq8AAAAGAj8AAP//AET+TgQoAfgAJwBDAq8AAAAGAj8AAP//AET+4gT0AWcAJwB0ApoAAAAGAIoAAP////j/SAT0AVQAJwB0ApoAAAAGAIsAAP//AET++wejAVQAJwB0BUkAAAAnAIsCrwAAAAYCNQAA//8ARP5OB/QBVAAnAHMFSQAAACcAiwKvAAAABgI/AAD//wBE/s8EagFUACcAdAIQAAAABgHjAAD////4/0gD9wFUACcAdAGdAAAABgHkAAD//wBE/uIG4gFnACcAcwQ3AAAAJwHkApoAAAAGAIoAAP////j/SAaRAVQAJwB0BDcAAAAnAeQCmgAAAAYAiwAA//8ARP77BvcBVAAnAHMETAAAACcB5AKvAAAABgI1AAD//wBE/k4G9wFUACcAcwRMAAAAJwHkAq8AAAAGAj8AAP//AET++wUJAVQAJwB0Aq8AAAAGAjUAAP//AET++wVaAVQAJwBzAq8AAAAGAjUAAP//AET+TgUJAVQAJwB0Aq8AAAAGAj8AAP//AET+TgVaAVQAJwBzAq8AAAAGAj8AAP//AET+4gT0AWcAJwCMApoAAAAGAHIAAP////j/SAT0AVQAJwCMApoAAAAGAHMAAP//AET+Tgf0AVQAJwCLBUkAAAAnAHMCrwAAAAYCPwAA//8ARP7PBGoBVAAnAIwCEAAAAAYB4wAA////+AAAA/cBVAAnAIwBnQAAAAYB5AAA//8ARP77BvcBVAAnAIsETAAAACcB5AKvAAAABgI1AAD//wBE/k4G9wFUACcAiwRMAAAAJwHkAq8AAAAGAj8AAP//AET++wUJAVQAJwCMAq8AAAAGAjUAAP//AET++wVaAVQAJwCLAq8AAAAGAjUAAP//AET+TgUJAVQAJwCMAq8AAAAGAj8AAP//AET+TgVaAVQAJwCLAq8AAAAGAj8AAP////gAAAKrAn8CJgCLAAABBwUmAK4AegAIsQECsHqwNSsAAP////j/SAKrAiECJgCLAAAAJgUpT04BBwUlAKsACgAQsQECsE6wNSuxAwGwCrA1K///AET+4gT0Ai4AJwCkApoAAAAGAHIAAP////j/SAT0Ai4AJwCkApoAAAAGAHMAAP//AET+4gT0Ai4AJwCkApoAAAAGAIoAAP//AET+zwRqAi4AJwCkAhAAAAAGAeMAAP////gAAAP3Ai4AJwCkAZ0AAAAGAeQAAP//AET++wUJAi4AJwCkAq8AAAAGAjUAAP//AET++wVaAi4AJwCjAq8AAAAGAjUAAP//AET+TgUJAi4AJwCkAq8AAAAGAj8AAP//AET+TgVaAi4AJwCjAq8AAAAGAj8AAP//AET+4gT5AWcAJwD6ApoAAAAGAHIAAP////j/SAVCAVQAJwD5ApoAAAAGAHMAAP////j/SAT5AVQAJwD6ApoAAAAGAHMAAP////j/SAeTAVQAJwD6BTQAAAAnAHMCmgAAAAYAiwAA//8ARP77B/EBVAAnAPkFSQAAACcAcwKvAAAABgI1AAD//wBE/uIE+QFnACcA+gKaAAAABgCKAAD////4AAAFQgFUACcA+QKaAAAABgCLAAD////4AAAE+QFUACcA+gKaAAAABgCLAAD////4/0gHkwFUACcA+gU0AAAAJwCLApoAAAAGAHMAAP//AET+4gT5AikAJwD6ApoAAAAGAKIAAP////gAAAVCAi4AJwD5ApoAAAAGAKMAAP////gAAAT5Ai4AJwD6ApoAAAAGAKMAAP//AET++wfxAi4AJwD5BUkAAAAnAKMCrwAAAAYCNQAA//8ARP5OB/ECLgAnAPkFSQAAACcAowKvAAAABgI/AAD////s/zQD8wFVACcA+gGUAAAABgDUAAD////s/zQEPAFVACcA+QGUAAAABgDUAAD//wBE/s8EbwFUACcA+gIQAAAABgHjAAD////4AAAERQFUACcA+QGdAAAABgHkAAD////4AAAD/AFUACcA+gGdAAAABgHkAAD////4/0gGlgFUACcA+gQ3AAAAJwHkApoAAAAGAHMAAP//AET+4gbfAWcAJwD5BDcAAAAnAeQCmgAAAAYAigAA////+AAABpYBVAAnAPoENwAAACcB5AKaAAAABgCLAAD//wBE/s8GVQFUACcA+QOtAAAAJwHkAhAAAAAGAeMAAP////gAAAWZAVQAJwD6AzoAAAAnAeQBnQAAAAYB5AAA////+P80BCUBpwAnAPkBfQAAAAYB/wAA////+P80A9wBpwAnAPoBfQAAAAYB/wAA//8ARP77BQ4BVAAnAPoCrwAAAAYCNQAA//8ARP77BVcBBAAnAPkCrwAAAAYCNQAA//8ARP5OBQ4BVAAnAPoCrwAAAAYCPwAA//8ARP5OBVcBBAAnAPkCrwAAAAYCPwAA////+P8DAl8CAAImAPoAAAAnBS8AjP/0AQYFMXf9ABKxAQO4//SwNSuxBAO4//2wNSsAAP//AET+4gT5AgAAJwEGApoAAAAGAHIAAP//AET+4gVCAgAAJwEFApoAAAAGAHIAAP////j/SAVCAgAAJwEFApoAAAAGAHMAAP////j/SAT5AgAAJwEGApoAAAAGAHMAAP//AET+TgfxAgAAJwEFBUkAAAAnAHMCrwAAAAYCPwAA//8ARP7iBPkCAAAnAQYCmgAAAAYAigAA//8ARP7iBUICAAAnAQUCmgAAAAYAigAA////+AAABUICAAAnAQUCmgAAAAYAiwAA////+AAABPkCAAAnAQYCmgAAAAYAiwAA//8ARP7PB1ICAAAnAQUEqgAAACcAiwIQAAAABgHjAAD////4AAAGlgIAACcBBgQ3AAAAJwCLAZ0AAAAGAeQAAP//AET+TgfxAgAAJwEFBUkAAAAnAIsCrwAAAAYCPwAA//8ARP7iBPkCKQAnAQYCmgAAAAYAogAA//8ARP7iBUICKQAnAQUCmgAAAAYAogAA////+AAABUICLgAnAQUCmgAAAAYAowAA////+AAABPkCLgAnAQYCmgAAAAYAowAA////7P80A/MCAAAnAQYBlAAAAAYA1AAA////7P80BDwCAAAnAQUBlAAAAAYA1AAA//8ARP7PBG8CAAAnAQYCEAAAAAYB4wAA//8ARP7PBLgCAAAnAQUCEAAAAAYB4wAA////+AAABEUCAAAnAQUBnQAAAAYB5AAA////+AAAA/wCAAAnAQYBnQAAAAYB5AAA//8ARP7iBt8CKQAnAQUENwAAACcB5AKaAAAABgCiAAD////4AAAGlgIuACcBBgQ3AAAAJwHkApoAAAAGAKMAAP//AET+zwZVAgAAJwEFA60AAAAnAeQCEAAAAAYB4wAA////+AAABZkCAAAnAQYDOgAAACcB5AGdAAAABgHkAAD////4/zQEJQIAACcBBQF9AAAABgH/AAD////4/zQD3AIAACcBBgF9AAAABgH/AAD//wBE/vsFDgIAACcBBgKvAAAABgI1AAD//wBE/vsFVwIAACcBBQKvAAAABgI1AAD//wBE/k4FDgIAACcBBgKvAAAABgI/AAD//wBE/k4FVwIAACcBBQKvAAAABgI/AAD//wBE/uIFaAFnACcBEgKaAAAABgCKAAD////4AAAFaAFlACcBEgKaAAAABgCLAAD//wBE/uIITwFnACcBEQU0AAAAJwCLApoAAAAGAIoAAP////gAAAgCAWUAJwESBTQAAAAnAIsCmgAAAAYAiwAA//8ARP5OCGQBZQAnAREFSQAAACcAiwKvAAAABgI/AAD////4AAAFaAIuACcBEgKaAAAABgCjAAD////s/zQEYgFlACcBEgGUAAAABgDUAAD////s/zQErwFlACcBEQGUAAAABgDUAAD//wAf/zQF4QKfACcBvwIeAAAAJwESAxMAAAAGAngAAP//AET+zwTeAWUAJwESAhAAAAAGAeMAAP////gAAARrAWUAJwESAZ0AAAAGAeQAAP//AET+zwbIAWUAJwERA60AAAAnAeQCEAAAAAYB4wAA////+AAABggBZQAnARIDOgAAACcB5AGdAAAABgHkAAD//wBE/vsFfQFlACcBEgKvAAAABgI1AAD//wBE/vsFygFlACcBEQKvAAAABgI1AAD//wBE/k4FfQFlACcBEgKvAAAABgI/AAD//wBE/k4FygFlACcBEQKvAAAABgI/AAD//wBE/uIFaAJMACcBHgKaAAAABgByAAD////4/0gFaAJMACcBHgKaAAAABgBzAAD//wBE/uIFaAJMACcBHgKaAAAABgCKAAD////4AAAFaAJMACcBHgKaAAAABgCLAAD//wBE/vsIZAJMACcBHQVJAAAAJwCLAq8AAAAGAjUAAP//AET+TghkAkwAJwEdBUkAAAAnAIsCrwAAAAYCPwAA//8ARP7iBWgCTAAnAR4CmgAAAAYAogAA////+AAABWgCTAAnAR4CmgAAAAYAowAA//8ARP7PB8UCTAAnAR0EqgAAACcAowIQAAAABgHjAAD////4AAAHBQJMACcBHgQ3AAAAJwCjAZ0AAAAGAeQAAP///+z/NARiAkwAJwEeAZQAAAAGANQAAP///+z/NASvAkwAJwEdAZQAAAAGANQAAP//AET+zwTeAkwAJwEeAhAAAAAGAeMAAP////gAAARrAkwAJwEeAZ0AAAAGAeQAAP//AET++wV9AkwAJwEeAq8AAAAGAjUAAP//AET++wXKAkwAJwEdAq8AAAAGAjUAAP//AET+TgV9AkwAJwEeAq8AAAAGAj8AAP//AET+TgXKAkwAJwEdAq8AAAAGAj8AAP//AET+4gT7Ap8AJwEmApoAAAAGAIoAAP////gAAAT7Ap8AJwEmApoAAAAGAIsAAP//AET+zwRxAp8AJwEmAhAAAAAGAeMAAP////gAAAQjAp8AJwElAZ0AAAAGAeQAAP////gAAAP+Ap8AJwEmAZ0AAAAGAeQAAP//AET+4ga9Ap8AJwElBDcAAAAnAeQCmgAAAAYAigAA////+AAABpgCnwAnASYENwAAACcB5AKaAAAABgCLAAD////4AAAFmwKfACcBJgM6AAAAJwHkAZ0AAAAGAeQAAP//AET+TgbSAp8AJwElBEwAAAAnAeQCrwAAAAYCPwAA//8ARP77BRACnwAnASYCrwAAAAYCNQAA//8ARP77BTUCnwAnASUCrwAAAAYCNQAA//8ARP5OBRACnwAnASYCrwAAAAYCPwAA//8ARP5OBTUCnwAnASUCrwAAAAYCPwAA////+AAAAoYCnwImASUAAAEHBSkAzgB2AAixAgKwdrA1KwAA//8ARP7PBHECnwAnAS4CEAAAAAYB4wAA////+AAABCMCnwAnAS0BnQAAAAYB5AAA////+AAAA/4CnwAnAS4BnQAAAAYB5AAA//8ARP7iBHkBdgAnATYCmgAAAAYAcgAA////+P9IBHkBdgAnATYCmgAAAAYAcwAA//8ARP7PBlkBdgAnATUEqgAAACcAcwIQAAAABgHjAAD////4/0gGFgF2ACcBNgQ3AAAAJwBzAZ0AAAAGAeQAAP//AET+zwPvAXYAJwE2AhAAAAAGAeMAAP////gAAAN8AXYAJwE2AZ0AAAAGAeQAAP//AET+zwVcAXYAJwE1A60AAAAnAeQCEAAAAAYB4wAA////+AAABRkBdgAnATYDOgAAACcB5AGdAAAABgHkAAD//wBE/vsF+wF2ACcBNQRMAAAAJwHkAq8AAAAGAjUAAP//AET+TgX7AXYAJwE1BEwAAAAnAeQCrwAAAAYCPwAA//8ARP77BI4BdgAnATYCrwAAAAYCNQAA//8ARP77BF4BdgAnATUCrwAAAAYCNQAA//8ARP5OBI4BdgAnATYCrwAAAAYCPwAA//8ARP5OBF4BdgAnATUCrwAAAAYCPwAA//8ARP7iBHkCQwAnAT4CmgAAAAYAcgAA////+P9IBHkCQwAnAT4CmgAAAAYAcwAA//8ARP7PA+8CQwAnAT4CEAAAAAYB4wAA////+AAAA3wCQwAnAT4BnQAAAAYB5AAA//8ARP7PBVwCQwAnAT0DrQAAACcB5AIQAAAABgHjAAD//wBE/vsF+wJDACcBPQRMAAAAJwHkAq8AAAAGAjUAAP//AET+TgX7AkMAJwE9BEwAAAAnAeQCrwAAAAYCPwAA//8ARP77BI4CQwAnAT4CrwAAAAYCNQAA//8ARP77BF4CQwAnAT0CrwAAAAYCNQAA//8ARP5OBI4CQwAnAT4CrwAAAAYCPwAA//8ARP5OBF4CQwAnAT0CrwAAAAYCPwAA//8ARP7iBAsClQAnAUYCmgAAAAYAcgAA////+P9IBAsClQAnAUYCmgAAAAYAcwAA//8ARP7iBAsClQAnAUYCmgAAAAYAigAA////+AAABAsClQAnAUYCmgAAAAYAiwAA//8ARP7iBAsClQAnAUYCmgAAAAYAogAA////+AAABAsClQAnAUYCmgAAAAYAowAA//8ARP7PBj4COwAnAUUEqgAAACcAowIQAAAABgHjAAD////4AAAFqAKVACcBRgQ3AAAAJwCjAZ0AAAAGAeQAAP//AET+zwOBApUAJwFGAhAAAAAGAeMAAP////gAAAMOApUAJwFGAZ0AAAAGAeQAAP//AET+TgXgAjsAJwFFBEwAAAAnAeQCrwAAAAYCPwAA//8ARP77BCAClQAnAUYCrwAAAAYCNQAA//8ARP77BEMCOwAnAUUCrwAAAAYCNQAA//8ARP5OBCAClQAnAUYCrwAAAAYCPwAA//8ARP5OBEMCOwAnAUUCrwAAAAYCPwAA//8ARP7iBAsCiAAnAWoCmgAAAAYAigAA////+AAABAsCiAAnAWoCmgAAAAYAiwAA//8ARP7PA4ECiAAnAWoCEAAAAAYB4wAA////+AAAAw4CiAAnAWoBnQAAAAYB5AAA//8ARP7iBcsCLgAnAWkENwAAACcB5AKaAAAABgCKAAD////4AAAFqAKIACcBagQ3AAAAJwHkApoAAAAGAIsAAP//AET+zwVBAi4AJwFpA60AAAAnAeQCEAAAAAYB4wAA//8ARP5OBeACLgAnAWkETAAAACcB5AKvAAAABgI/AAD//wBE/vsEIAKIACcBagKvAAAABgI1AAD//wBE/vsEQwIuACcBaQKvAAAABgI1AAD//wBE/k4EIAKIACcBagKvAAAABgI/AAD//wBE/k4EQwIuACcBaQKvAAAABgI/AAD//wBMAAACigKfACcBdgD9AAAABgAEAAD//wBMAAAC2QKfACcBdQD9AAAABgAEAAD//wBE/uIEJwKOACcBdgKaAAAABgByAAD////4/0gEJwKOACcBdgKaAAAABgBzAAD//wBE/uIEJwKOACcBdgKaAAAABgCKAAD////4AAAEJwKOACcBdgKaAAAABgCLAAD//wBE/uIEJwKOACcBdgKaAAAABgCiAAD////4AAAEJwKOACcBdgKaAAAABgCjAAD//wBE/voD3AKfACcBdgJPAAAABgG+AAD//wBE/voEKwKfACcBdQJPAAAABgG+AAD////4AAAC0QKfACcBdQD1AAAABgG/AAD////4AAACggKfACcBdgD1AAAABgG/AAD//wBE/s8DnQKOACcBdgIQAAAABgHjAAD//wBE/s8D7AKOACcBdQIQAAAABgHjAAD////4AAADeQKOACcBdQGdAAAABgHkAAD////4AAADKgKOACcBdgGdAAAABgHkAAD//wBE/s8FiQKOACcBdQOtAAAAJwHkAhAAAAAGAeMAAP////gAAATHAo4AJwF2AzoAAAAnAeQBnQAAAAYB5AAA//8ARP5OBigCjgAnAXUETAAAACcB5AKvAAAABgI/AAD//wBE/vsEPAKOACcBdgKvAAAABgI1AAD//wBE/vsEiwKOACcBdQKvAAAABgI1AAD//wBE/k4EPAKOACcBdgKvAAAABgI/AAD//wBE/k4EiwKOACcBdQKvAAAABgI/AAD//wBE/uIDMwKfACcBwQKaAAAABgByAAD////4/0gDMwKfACcBwQKaAAAABgBzAAD//wBE/uIGOgKfACcBvwU0AAAAJwBzApoAAAAGAHIAAP////j/SAXNAp8AJwHBBTQAAAAnAHMCmgAAAAYAcwAA//8ARP7PBbACnwAnAb8EqgAAACcAcwIQAAAABgHjAAD////4/0gE0AKfACcBwQQ3AAAAJwBzAZ0AAAAGAeQAAP//AET+TgZPAp8AJwG/BUkAAAAnAHMCrwAAAAYCPwAA//8ARP7iAzMCnwAnAcECmgAAAAYAigAA////+AAAAzMCnwAnAcECmgAAAAYAiwAA//8ARP7PBbACnwAnAb8EqgAAACcAiwIQAAAABgHjAAD////4AAAE0AKfACcBwQQ3AAAAJwCLAZ0AAAAGAeQAAP//AET++wZPAp8AJwG/BUkAAAAnAIsCrwAAAAYCNQAA//8ARP5OBk8CnwAnAb8FSQAAACcAiwKvAAAABgI/AAD//wBE/uIDMwKfACcBwQKaAAAABgCiAAD////4AAADMwKfACcBwQKaAAAABgCjAAD//wBE/s8FsAKfACcBvwSqAAAAJwCjAhAAAAAGAeMAAP////gAAATQAp8AJwHBBDcAAAAnAKMBnQAAAAYB5AAAAAIARAAAA2QCnwBAAEkARUBCKCQdGQQFAQQ+BQIAAQJMR0QxLiMgGAcESgYBBAEEhQIBAQEAYQMFAgAADwBOQUEBAEFJQUk9OyclHBoAQAFABwcWKyEiJicnBy4CJy4CNTQ2Nz4CNyc2NjcTFjMyNyYmJzY2NxMWMzI3JiYnJiYnNjY3HgQVFAcGBiMiJwYGJyYmJwYGBxYWAcw8SgcGIiQpGg8SKyARDQ46SCEFDyYXCiw+NjYQEwUOIRULLzs7MwMGBAcUCw8iEwcJBQMCIwo9MkgpJju/Bw4GOFgXJ2EuKSRAAwYIBQcXGw0LLhkcMikOKg4cDv6XGxqKsSUOHA7+aRsbJEolWbZeEBsNcqJvSzMXIUMSESQVD4wpUSgTNx8iFv//AET+zwKpAp8AJwHBAhAAAAAGAeMAAP//AET+zwMWAp8AJwG/AhAAAAAGAeMAAP////gAAAKjAp8AJwG/AZ0AAAAGAeQAAP////gAAAI2Ap8AJwHBAZ0AAAAGAeQAAP//AET+4gU9Ap8AJwG/BDcAAAAnAeQCmgAAAAYAigAA////+AAABNACnwAnAcEENwAAACcB5AKaAAAABgCLAAD//wBE/k4FUgKfACcBvwRMAAAAJwHkAq8AAAAGAj8AAP////j/NAIWAp8AJwHBAX0AAAAGAf8AAP//AET++wNIAp8AJwHBAq8AAAAGAjUAAP//AET++wO1Ap8AJwG/Aq8AAAAGAjUAAP//AET+TgNIAp8AJwHBAq8AAAAGAj8AAP//AET+TgO1Ap8AJwG/Aq8AAAAGAj8AAP//ACIAAAJzA30CJgLfAAAAJwWiATsBmwEHBXkAwQBAABGxAQG4AZuwNSuxAgGwQLA1KwD///+5AAACcwN9AiYC3wAAACcFogE7AZsBBgYAkNEAErEBAbgBm7A1K7ECAbj/0bA1KwAA////uQAAAnMDfQImAt8AAAAnBaIBOwGbAQYGAaPSABKxAQG4AZuwNSuxAgG4/9KwNSsAAP//ACMAAAHSA3UCJgLeAAABBwUjAK4BqAAJsQIBuAGosDUrAP//ACIAAAJzA2ECJgLfAAAAJwXbAP4BlAEHBXkAwQBAABGxAQG4AZSwNSuxAgGwQLA1KwD///+5AAACcwNhAiYC3wAAACcF2wD+AZQBBgYAkNEAErEBAbgBlLA1K7ECAbj/0bA1KwAA////uQAAAnMDYQImAt8AAAAnBdsA/gGUAQYGAaPSABKxAQG4AZSwNSuxAgG4/9KwNSsAAP//ACIAAAJzA48CJgLfAAAAJwUvAOcBgwEHBXkAwQBAABGxAQO4AYOwNSuxBAGwQLA1KwD///+5AAACcwOPAiYC3wAAACcFLwDnAYMBBgYAkNEAErEBA7gBg7A1K7EEAbj/0bA1KwAA////uQAAAnMDjwImAt8AAAAnBS8A5wGDAQYGAaPSABKxAQO4AYOwNSuxBAG4/9KwNSsAAP//ACL/CwJzAuoCJgLfAAAAJwXdANwAEQEHBXkAwQBAABCxAQOwEbA1K7EEAbBAsDUrAAD///+5/wsCcwKoAiYC3wAAACcF3QDcABEBBgYAkNEAEbEBA7ARsDUrsQQBuP/RsDUrAAAA////uf8LAnMCqQImAt8AAAAnBd0A3AARAQYGAaPSABGxAQOwEbA1K7EEAbj/0rA1KwAAAP//AEwAAAKrAp8AJwHkAP0AAAAGAAQAAP//AET+4gQzAWcAJwHlApoAAAAGAHIAAP////j/SAQzAVQAJwHlApoAAAAGAHMAAP////j/SAbNAVQAJwHlBTQAAAAnAHMCmgAAAAYAiwAA////+P9IBs0CLgAnAeUFNAAAACcAcwKaAAAABgCjAAD////4/0gF0AFUACcB5QQ3AAAAJwBzAZ0AAAAGAeQAAP//AET+Tgb3AVQAJwHkBUkAAAAnAHMCrwAAAAYCPwAA//8ARP7iBDMBZwAnAeUCmgAAAAYAigAA////+AAABDMBVAAnAeUCmgAAAAYAiwAA////+P9IBs0BVAAnAeUFNAAAACcAiwKaAAAABgBzAAD////4AAAF0AFUACcB5QQ3AAAAJwCLAZ0AAAAGAeQAAP//AET+Tgb3AVQAJwHkBUkAAAAnAIsCrwAAAAYCPwAA//8ARP7iBDMCKQAnAeUCmgAAAAYAogAA////+AAABDMCLgAnAeUCmgAAAAYAowAA////+P9IBs0CLgAnAeUFNAAAACcAowKaAAAABgBzAAD////4AAAF0AIuACcB5QQ3AAAAJwCjAZ0AAAAGAeQAAP//AET+Tgb3Ai4AJwHkBUkAAAAnAKMCrwAAAAYCPwAA//8ARP7PA6kBVAAnAeUCEAAAAAYB4wAA//8ARP7PA74BVAAnAeQCEAAAAAYB4wAA////+AAAAzYBVAAnAeUBnQAAAAYB5AAA//8ARP5OBfoBVAAnAeQETAAAACcB5AKvAAAABgI/AAD//wBE/vsESAFUACcB5QKvAAAABgI1AAD//wBE/k4ESAFUACcB5QKvAAAABgI/AAD//wBE/uIDewJEACcB6gKaAAAABgByAAD////4/0gDewJEACcB6gKaAAAABgBzAAD//wBE/uIGaQHMACcB6QU0AAAAJwBzApoAAAAGAIoAAP////j/SAYVAkQAJwHqBTQAAAAnAHMCmgAAAAYAiwAA//8ARP7PBd8BzAAnAekEqgAAACcAcwIQAAAABgHjAAD////4/0gFGAJEACcB6gQ3AAAAJwBzAZ0AAAAGAeQAAP//AET++wZ+AcwAJwHpBUkAAAAnAHMCrwAAAAYCNQAA//8ARP5OBn4BzAAnAekFSQAAACcAcwKvAAAABgI/AAD//wBE/uIDewJEACcB6gKaAAAABgCKAAD////4AAADewJEACcB6gKaAAAABgCLAAD////4AAAFGAJEACcB6gQ3AAAAJwCLAZ0AAAAGAeQAAP//AET++wZ+AcwAJwHpBUkAAAAnAIsCrwAAAAYCNQAA//8ARP5OBn4BzAAnAekFSQAAACcAiwKvAAAABgI/AAD//wBE/uIDewJEACcB6gKaAAAABgCiAAD////4AAADewJEACcB6gKaAAAABgCjAAD////s/zQCyQHMACcB6QGUAAAABgDUAAD////s/zQCyQIaACcB6QGUAAAABgDWAAD//wBE/s8C8QJEACcB6gIQAAAABgHjAAD//wBE/s8DRQHMACcB6QIQAAAABgHjAAD////4AAAC0gHMACcB6QGdAAAABgHkAAD////4AAACfgJEACcB6gGdAAAABgHkAAD//wBE/vsFgQHMACcB6QRMAAAAJwHkAq8AAAAGAjUAAP//AET+TgWBAcwAJwHpBEwAAAAnAeQCrwAAAAYCPwAA//8ARP76A34BzAAnAekCSQAAAAYB6AAA////+P80ArIBzAAnAekBfQAAAAYB/wAA////+P80Al4CRAAnAeoBfQAAAAYB/wAA//8ARP77A5ACRAAnAeoCrwAAAAYCNQAA//8ARP77A+QBzAAnAekCrwAAAAYCNQAA//8ARP5OA5ACRAAnAeoCrwAAAAYCPwAA//8ARP5OA+QBzAAnAekCrwAAAAYCPwAA//8ARP7iBGMBsQAnAgACmgAAAAYAcgAA////+P9IBGMBsQAnAgACmgAAAAYAcwAA//8ARP7PA9kBsQAnAgACEAAAAAYB4wAA////+P/uA2YBsQAnAgABnQAAAAYB5AAA////+P9IBgABsQAnAgAENwAAACcB5AKaAAAABgBzAAD////4/+4FAwGxACcCAAM6AAAAJwHkAZ0AAAAGAeQAAP//AET++wR4AbEAJwIAAq8AAAAGAjUAAP//AET+TgR4AbEAJwIAAq8AAAAGAj8AAP//AET+4gPOAXYAJwJBApoAAAAGAHIAAP////j/OwPOAXYAJwJBApoAAAAGAHMAAP//AET+TgbCAVQAJwJABUkAAAAnAHMCrwAAAAYCPwAA//8ARP7iA84BdgAnAkECmgAAAAYAigAA////+P87A84BdgAnAkECmgAAAAYAiwAA//8ARP5OBsIBVAAnAkAFSQAAACcAiwKvAAAABgI/AAD//wBE/uIDzgIpACcCQQKaAAAABgCiAAD////4/zsDzgIuACcCQQKaAAAABgCjAAD////s/zQDDQFVACcCQAGUAAAABgDUAAD////s/zQDDQIaACcCQAGUAAAABgDWAAD//wBE/s8DRAF2ACcCQQIQAAAABgHjAAD//wBE/s8DiQFUACcCQAIQAAAABgHjAAD////4/zsDFgFUACcCQAGdAAAABgHkAAD////4/zsC0QF2ACcCQQGdAAAABgHkAAD//wBE/s8FJgFUACcCQAOtAAAAJwHkAhAAAAAGAeMAAP////j/OwRuAXYAJwJBAzoAAAAnAeQBnQAAAAYB5AAA//8ARP5OBcUBVAAnAkAETAAAACcB5AKvAAAABgI/AAD//wBE/voDwgF1ACcCQAJJAAAABgHoAAD////4/zQC9gGnACcCQAF9AAAABgH/AAD////4/zQCsQGnACcCQQF9AAAABgH/AAD//wBE/vsD4wF2ACcCQQKvAAAABgI1AAD//wBE/vsEKADkACcCQAKvAAAABgI1AAD//wBE/k4D4wF2ACcCQQKvAAAABgI/AAD//wBE/k4EKADkACcCQAKvAAAABgI/AAD//wBMAAAB3gKfACcCRQD9AAAABgAEAAD//wBMAAACMgKfACcCRAD9AAAABgAEAAD//wBE/uIDewJbACcCRQKaAAAABgByAAD////4/0gDewJbACcCRQKaAAAABgBzAAD//wBE/uIDewJbACcCRQKaAAAABgCKAAD////4AAADewJbACcCRQKaAAAABgCLAAD////4AAADewJbACcCRQKaAAAABgCjAAD////s/zQCyQHjACcCRAGUAAAABgDUAAD////s/zQCyQIaACcCRAGUAAAABgDWAAD//wBE/s8C8QJbACcCRQIQAAAABgHjAAD//wBE/s8DRQHjACcCRAIQAAAABgHjAAD////4AAAC0gHjACcCRAGdAAAABgHkAAD////4AAACfgJbACcCRQGdAAAABgHkAAD//wBE/voDfgHjACcCRAJJAAAABgHoAAD////4/zQCsgHjACcCRAF9AAAABgH/AAD////4/zQCXgJbACcCRQF9AAAABgH/AAD//wAi/zUCtQJbACcCRQHUAAAABgIZAAD//wAi/zUDCQHjACcCRAHUAAAABgIZAAD//wAi/zUCtQJbACcCRQHUAAAABgIrAAD//wAi/zUDCQI/ACcCRAHUAAAABgIrAAD//wAi/zUCtQJbACcCRQHUAAAABgItAAD//wAi/zUDCQIuACcCRAHUAAAABgItAAD//wAi/zUCtQKCACcCRQHUAAAABgIvAAD//wAi/zUDCQKCACcCRAHUAAAABgIvAAD//wBE/vsDkAJbACcCRQKvAAAABgI1AAD//wBE/vsD5AHjACcCRAKvAAAABgI1AAD//wBE/k4DkAJbACcCRQKvAAAABgI/AAD//wBE/k4D5AHjACcCRAKvAAAABgI/AAD//wBE/fkDkAJbACcCRQKvAAAABgJyAAD//wBE/fkD5AHjACcCRAKvAAAABgJyAAD////4/uYCSQJbACcCRQFoAAAABgJzAAD//wBEAAACigJbACcCRQGpAAAABgIEAAD//wBEAAAC3gHjACcCRAGpAAAABgIEAAD//wBE/vsDkAJbACcCRQKvAAAABgI1AAD//wBE/vsD5AHjACcCRAKvAAAABgI1AAD////4AAACBQJbACcCRQEkAAAABgAhAAD//wBEAAAEUgNkACYEJgAAACcAAwOgAAABBgYRBzQACLEDArA0sDUr//8ARAAABFICnwAmBCYAAAAHAAMDoAAA////7P8xBQkDdgAnAAcEVwAAACcBdgK4AAAAJwAlAZQAAAAGANQAAP//ACIAAAeqAvoAJwHlBhEAAAAnAIsDdwAAACcFjAa6AG8AJwWTAlQAXwAnBYgEVABZACcFiAI5AM4AJwHkAdoAAAEGALYAAAAgsQMCsG+wNSuxBQGwX7A1K7EGAbBZsDUrsQcBsM6wNSsAAP//AET+zwdxAp8AJwESBKMAAAAnAb8DrgAAACcBNQIQAAAABgHjAAD//wBE/voHcgKfACcA0wYjAAAAJwD6A/UAAAAnAhkCIQAAAAYBvQAA//8ARP87BgACnwAnATYEIQAAACcBvwMsAAAAJwJAAcQAAAAGAf4AAP//AET+zwbgAp8AJwD6AxIAAAAnAhgFPwAAACcBvwIZAAAABgHjAAD//wBE/vsGWwKfACcBvwKYAAAAJwESA40AAAAGAjUAAAALADH/TgSQA3EAJwHYAe8B9gICAgkCFQI1Aj8CSwJXFPdLsBJQWEGyACYAFQACAAEAJwEAANoAAgAAAAEB5wHmARQAzADKALoABgAVAAAB8QABACoAFQC9AAEAFAAqAfoB9AEeAAMAKAAUAgABIgEfARcBCgD7ANcA0gDGAMAAqAALAA8AKAE0AKkAoQCMAIgABQALAA8B7wHZATYA7ACcAGcAXAAHABgACwFJAGQAAgAZABgCBAFQAUoAAwAJABEAeAABAAwACQIHAVYAgQB7AFIABQArAAwBawFTAT4BOQBgAEsABgAHACsBZQABAAYABwI7AhABpAADAAQABQI+AhMCDQADACAABAI4AZgAAgAxACACTwABAB4AMQJTAkMB0wHQAc0BmwGIAAcAHwAeAa8BdgBDAEAALgAFAB0AHwIiAcgBuAADACIAHQHCAAEAJAAiAh0AAQAtACQCHAABACwALQAZAEwA/QABABUAAQBLAQMA3QAQAAUABAAFACcASgHWACgAAgAsAEkbS7ATUFhBtQAmABUAAgABACcBAADaAAIAAAABAecB5gEUAMwAygC6AAYAFQAAAfEAAQAqABUAvQABABQAKgH6AfQBHgADACgAFAIAASIBHwEXAQoA+wDXANIAxgDAAKgACwAPACgBNACpAKEAjACIAAUACwAPAe8B2QE2AOwAnABnAFwABwAYAAsBSQBkAAIAGQAYAgQBUAFKAAMACQARAHgAAQAMAAkCBwFWAIEAewBSAAUAKwAMAVMBPgE5AEsABAAXACsBawBgAAIABwAXAWUAAQAGAAcCOwIQAaQAAwAEAAUCPgITAg0AAwAgAAQCOAGYAAIAMQAgAk8AAQAeADECUwJDAdMB0AHNAZsBiAAHAB8AHgGvAXYAQwBAAC4ABQAdAB8CIgHIAbgAAwAiAB0BwgABACQAIgIdAAEALQAkAhwAAQAsAC0AGgBMAP0AAQAVAAEASwEDAN0AEAAFAAQABQAnAEoB1gAoAAIALABJG0uwFlBYQbUAJgABAAEAAgEAANoAAgADAAEB5wHmARQAzADKALoABgAVAAAB8QABACoAFQC9AAEAFAAqAfoB9AEeAAMAKAAUAgABIgEfARcBCgD7ANcA0gDGAMAAqAALAA8AKAE0AKkAoQCMAIgABQALAA8B7wHZATYA7ACcAGcAXAAHABgACwFJAGQAAgAZABgCBAFQAUoAAwAJABEAeAABAAwACQIHAVYAgQB7AFIABQArAAwBUwE+ATkASwAEABcAKwFrAGAAAgAHABcBZQABAAYABwI7AhABpAADAAQABQI+AhMCDQADACAABAI4AZgAAgAxACACTwABAB4AMQJTAkMB0wHQAc0BmwGIAAcAHwAeAa8BdgBDAEAALgAFAB0AHwIiAcgBuAADACIAHQHCAAEAJAAiAh0AAQAtACQCHAABACwALQAaAEwA/QABABUAAQBLAQMA3QAVABAABQAEAAYAAgBKAdYAKAACACwASRtLsBhQWEG1ACYAAQABAAIBAADaAAIAAwABAecB5gEUAMwAygC6AAYAFQAAAfEAAQAqABUAvQABABQAKgH6AfQBHgADACgAFAIAASIBHwEXAQoA+wDXANIAxgDAAKgACwAPACgBNACpAKEAjACIAAUACwAPAe8B2QE2AOwAnABnAFwABwAYAAsBSQBkAAIAGQAYAgQBUAFKAAMACQARAHgAAQAMAAkCBwFWAIEAewBSAAUAKwAMAVMBPgE5AEsABAAKACsBawBgAAIABwAKAWUAAQAGAAcCOwIQAaQAAwAEAAUCPgITAg0AAwAgAAQCOAGYAAIAMQAgAk8AAQAeADECUwJDAdMB0AHNAZsBiAAHAB8AHgGvAXYAQwBAAC4ABQAdAB8CIgHIAbgAAwAiAB0BwgABACQAIgIdAAEALQAkAhwAAQAsAC0AGgBMAP0AAQAVAAEASwEDAN0AFQAQAAUABAAGAAIASgHWACgAAgAsAEkbS7AbUFhBuAAmAAEAAQACAQAA2gACAAMAAQHnAeYBFADMAMoAugAGABUAAAHxAAEAKgAVAL0AAQAUACoB+gH0AR4AAwAoABQCAAEiAR8BFwEKAPsA1wDSAMYAwACoAAsADwAoATQAqQChAIwAiAAFAAsADwHvAdkBNgDsAJwAZwBcAAcAGAALAUkAZAACABkAGAIEAVABSgADAAkAEQB4AAEADAAJAgcBVgCBAHsAUgAFACsADAFTAT4BOQBLAAQACgArAWsAYAACAAcACgFlAAEABgAHAhABpAACAC8ABQI7AAEABAAvAj4CEwINAAMAIAAEAjgBmAACADEAIAJPAAEAHgAxAlMCQwHTAdABzQGbAYgABwAfAB4BrwF2AEMAQAAuAAUAHQAfAiIByAG4AAMAIgAdAcIAAQAkACICHQABAC0AJAIcAAEALAAtABsATAD9AAEAFQABAEsBAwDdABUAEAAFAAQABgACAEoB1gAoAAIALABJG0uwHlBYQbsAFQABAAIAJwAmAAEAAQACAQAA2gACAAMAAQHnAeYBFADMAMoAugAGABUAAAHxAAEAKgAVAL0AAQAUACoB+gH0AR4AAwAoABQCAAEiAR8BFwEKAPsA1wDSAMYAwACoAAsADwAoATQAqQChAIwAiAAFAAsADwHvAdkBNgDsAJwAZwBcAAcAGAALAUkAZAACABkAGAIEAVABSgADAAkAEQB4AAEADAAJAgcBVgCBAHsAUgAFACsADAFTAT4BOQBLAAQACgArAWsAYAACAAcACgFlAAEABgAHAhABpAACAC8ABQI7AAEABAAvAj4CEwINAAMAIAAEAjgBmAACADEAIAJPAAEAHgAxAlMCQwHTAdABzQGbAYgABwAfAB4BrwF2AEMAQAAuAAUAHQAfAiIByAG4AAMAIgAdAcIAAQAkACICHQABAC0AJAIcAAEALAAtABwATAD9AAEAFQABAEsBAwDdABAABQAEAAUAJwBKAdYAKAACACwASRtLsC5QWEG7ABUAAQACACcAJgABAAEAAgEAANoAAgADAAEB5wHmARQAzADKALoABgAVAAAB8QABACoAFQC9AAEAFAAqAfoB9AEeAAMAKAAUAgABIgEfARcBCgD7ANcA0gDGAMAAqAALAA8AKAE0AKkAoQCMAIgABQALAA8B7wHZATYA7ACcAGcAXAAHABgACwFJAGQAAgAZABgCBAFQAUoAAwAJABEAeAABAAwACQIHAVYAgQB7AFIABQArAAwBUwE+ATkASwAEAAoAKwFrAGAAAgAHAAoBZQABAAYABwIQAaQAAgAvAAUCOwABAAQALwI+AhMCDQADACAABAI4AZgAAgAxACACTwABAB4AMQJTAkMB0wHQAc0BmwGIAAcAHwAeAa8BdgBDAEAALgAFAB0AHwIiAcgBuAADACIAHQHCAAEALgAiAh0AAQAtACQCHAABACwALQAcAEwA/QABABUAAQBLAQMA3QAQAAUABAAFACcASgHWACgAAgAsAEkbQcIAFQABAAIAJwAmAAEAAQACAQAA2gACAAMAAQHnAeYBFADMAMoAugAGABUAAAHxAAEAKgAVAL0AAQAUACoB+gH0AR4AAwAoABQCAAEiAR8BCgD7ANcA0gDGAMAAqAAKABIAKAE0AKkAoQCMAIgABQALAA8B7wE2AJwAAwANAAsA7ABnAFwAAwAYAA0BSQBkAAIAGQAYAgQBUAFKAAMACQARAHgAAQAMAAkCBwFWAIEAewBSAAUAKwAMAVMBPgE5AEsABAAKACsBawBgAAIABwAKAWUAAQAGAAcCEAGkAAIALwAFAjsAAQAEAC8CPgITAg0AAwAgAAQCOAGYAAIAMQAgAk8AAQAeADECUwJDAdMB0AHNAZsBiAAHAB8AHgGvAXYAQwBAAC4ABQAdADACIgHIAbgAAwAiAB0BwgABAC4AIgIdAAEALQAkAhwAAQAsAC0AHQBMAP0AAQAVARcAAQASAdkAAQANAAMASwEDAN0AEAAFAAQABQAnAEoB1gAoAAIALABJWVlZWVlZWUuwEVBYQKwAJwEnhQAUKigqFCiANAEoDyooD34ACREMEQkMgAAgBDEEIDGAAB4xHzEeH4ACAQEDMwIAFQEAaTUpExIQBQ8WDg0DCxgPC2kAGAAZERgZaQARAAwrEQxqNgErBwYrWRsaAgYFBwZaFwoIAwccAQUEBwVqLwEEADEeBDFpAC03ASwtLGUAKioVYQAVFQ5NIQEfHyJhIwEiIg9NOTI4MAQdHSRhLiYlAyQkDyROG0uwElBYQKwAJwEnhQAUKigqFCiANAEoDyooD34ACREMEQkMgAAgBDEEIDGAAB4xHzEeH4ACAQEDMwIAFQEAaTUpExIQBQ8WDg0DCxgPC2kAGAAZERgZaQARAAwrEQxqNgErBwYrWRsaAgYFBwZaFwoIAwccAQUEBwVqLwEEADEeBDFpAC03ASwtLGUAKioVYQAVFQ5NODAhAx8fImEjASIiD005MgIdHSRhLiYlAyQkDyROG0uwE1BYQK0AJwEnhQAUKigqFCiANAEoDyooD34ACREMEQkMgAAgBDEEIDGAAB4xHzEeH4ACAQEDMwIAFQEAaTUpExIQBQ8WDg0DCxgPC2kAGAAZERgZaQARAAwrEQxqNgErFwYrWQAXGxoCBgUXBmoKCAIHHAEFBAcFai8BBAAxHgQxaQAtNwEsLSxlACoqFWEAFRUOTTgwIQMfHyJhIwEiIg9NOTICHR0kYS4mJQMkJA8kThtLsBZQWEC0ABQqKCoUKIA0ASgPKigPfgAJEQwRCQyAACAEMQQgMYAAHjEfMR4fgCcBAgADAAIDaQABMwEAFQEAaTUpExIQBQ8WDg0DCxgPC2kAGAAZERgZaQARAAwrEQxqABcHGhdZNgErAAYaKwZqGwEaBQcaWgoIAgccAQUEBwVqLwEEADEeBDFpAC03ASwtLGUAKioVYQAVFQ5NODAhAx8fImEjASIiD005MgIdHSRhLiYlAyQkDyROG0uwGFBYQLAAFCooKhQogDQBKA8qKA9+AAkRDBEJDIAAIAQxBCAxgAAeMR8xHh+AJwECAAMAAgNpAAEzAQAVAQBpNSkTEhAFDxYODQMLGA8LaQAYABkRGBlpABEADCsRDGo2ASsABhorBmoXAQobARoFChpqCAEHHAEFBAcFaS8BBAAxHgQxaQAtNwEsLSxlACoqFWEAFRUOTTgwIQMfHyJhLiMCIiIPTTkyAh0dJGEmJQIkJA8kThtLsBpQWEC2ABQqKCoUKIA0ASgPKigPfgAJEQwRCQyAAAQvIC8EIIAAIDEvIDF+AB4xHzEeH4AnAQIAAwACA2kAATMBABUBAGk1KRMSEAUPFg4NAwsYDwtpABgAGREYGWkAEQAMKxEMajYBKwAGGisGahcBChsBGgUKGmoIAQccAQUvBwVpAC8AMR4vMWkALTcBLC0sZQAqKhVhABUVDk04MCEDHx8iYS4jAiIiD005MgIdHSRhJiUCJCQPJE4bS7AbUFhAtAAUKigqFCiANAEoDyooD34ACREMEQkMgAAELyAvBCCAACAxLyAxfgAeMR8xHh+AJwECAAMAAgNpAAEzAQAVAQBpNSkTEhAFDxYODQMLGA8LaQAYABkRGBlpABEADCsRDGo2ASsABhorBmoXAQobARoFChpqCAEHHAEFLwcFaQAvADEeLzFpODAhAx8uIwIiJB8iaQAtNwEsLSxlACoqFWEAFRUOTTkyAh0dJGEmJQIkJA8kThtLsB5QWEC4ACcCJ4UAFCooKhQogDQBKA8qKA9+AAkRDBEJDIAABC8gLwQggAAgMS8gMX4AHjEfMR4fgAACAAMAAgNpAAEzAQAVAQBpNSkTEhAFDxYODQMLGA8LaQAYABkRGBlpABEADCsRDGo2ASsABhorBmoXAQobARoFChpqCAEHHAEFLwcFaQAvADEeLzFpODAhAx8uIwIiJB8iaQAtNwEsLSxlACoqFWEAFRUOTTkyAh0dJGEmJQIkJA8kThtLsC1QWEC/ACcCJ4UAFCooKhQogDQBKA8qKA9+AAkRDBEJDIAABC8gLwQggAAgMS8gMX4AHjEfMR4fgAACAAMAAgNpAAEzAQAVAQBpNSkTEhAFDxYODQMLGA8LaQAYABkRGBlpABEADCsRDGo2ASsABhorBmoXAQobARoFChpqCAEHHAEFLwcFaQAvADEeLzFpIwEiLh8iWQAtNwEsLSxlACoqFWEAFRUOTTgwIQMfHy5hAC4uD005MgIdHSRhJiUCJCQPJE4bS7AuUFhAvQAnAieFABQqKCoUKIA0ASgPKigPfgAJEQwRCQyAAAQvIC8EIIAAIDEvIDF+AB4xHzEeH4AAAgADAAIDaQABMwEAFQEAaTUpExIQBQ8WDg0DCxgPC2kAGAAZERgZaQARAAwrEQxqNgErAAYaKwZqFwEKGwEaBQoaaggBBxwBBS8HBWkALwAxHi8xaSMBIi4fIlk4MCEDHwAuJB8uaQAtNwEsLSxlACoqFWEAFRUOTTkyAh0dJGEmJQIkJA8kThtAxAAnAieFABQqKCoUKIA0ASgSKigSfgAJEQwRCQyAAAQvIC8EIIAAIDEvIDF+AB4xHzEeH4AAAgADAAIDaQABMwEAFQEAaQASAAsNEgtpNSkTEAQPFg4CDRgPDWkAGAAZERgZaQARAAwrEQxqNgErAAYaKwZqFwEKGwEaBQoaaggBBxwBBS8HBWkALwAxHi8xaSEBHyMBIi4fImk4ATAALiQwLmkALTcBLC0sZQAqKhVhABUVDk05MgIdHSRhJiUCJCQPJE5ZWVlZWVlZWVlZQYUCTAJMAkECQAIXAhYCAwIDAfgB9wHwAfAAAQAAAkwCVwJMAlcCRwJFAkACSwJBAksCLgIsAicCJQIgAh4CFgI1AhcCNQIDAgkCAwIJAf4B/AH3AgIB+AICAfAB9gHwAfYB4wHiAcoByQHGAcQBwAG+AbYBtAGtAawBngGdAZYBlQGMAYoBhAGCAXoBeAFxAXABaQFnAWMBYQFOAUwBRgFEAT0BOgExAS8BJgEkARwBGwEMAQsA+QD3AOkA5wDWANQAxADCAKUAowCfAJ0AlwCVAIsAiQB+AH0AdwB1AG4AbABjAGEATQBMAEcARgA5ADgAIwAiABQAEgALAAkAAAAnAAEAJwA6AAcAFisBIjU0NxcGFRQWMzI1NCY1NxYWMzI3JiYnJjU0NjcWFhUUBiMiJicGAyYmJyYmNTQ3NjY3NjY3NjYzMhYXHgIXNjY3JiYnIiYmJycHJiYnJiY1NDY3NjY3JzY2Nx4CFxYWMyc2NjcUHgIXMjY2NzY2NzY2MzIXBgYHFhYzMjY1LgInJiYnBiMiJx4CFRQGBwYGIyImNTQ2NwYjIiYnBgYjIiYnJwcmJicuAjU0Njc2NjcnNjY3FBYXFhQVFhYzMjY3LgInNjcWFhUUFhcWFjMyNyYmJzY2NxYWFRQGBwYVFBYzMjY2NycuAjU0NjYzFhYzMjY3JicmJic2NjcWFhceAhUWFzQuAjE2NjcWFhc2NzY2MzIXBxYWFzY2MzIWFxYVFAYHBgYjIiYmJwYHFhYXFjMyMjcmNTQ2NzYzMhcWFwcmJiMiBgcUFhc2NjcWFhUUBgcGBgcGBiMiJicGBiMiJicGBgcGBiMUHgIVFhYzMjY2NzY2NzY2MzIWFwYHFhYzMjY3NjY3NjY3NjMyFwYGBxYWMyYmNTQ2NxYWFxYVFAYHIiYnBgYHBgYjIiYnBgYHDgIjIiYnBgYjIiYnByImJjU0NjcGBgcWFhcGBgE0JicuAjU0NjcWFhcHMB4CFRQGByUnBgYHFhYFMjY3JiYjIgYHFhYFJwYGBxYWBSYmJzY2NxYWFwYGBSImJyYmJzcWMzI2NycGBiMiJjU0NjYzMhYXFhUUBgYlJic2NjcWFhcGBTI2NyYmIyIGFRQWBSYmJw4CBxYXFhYBAxsHCgMJBRABDgIJDQoCAQQEAgsGBwMRDwQMBAzKAwUDAgMFBgYJBSIOESwOCA8HCBEWDgsbEAQHAxpAMAMDEhMgEwofDgQNMhwDCRAPAQICAQo+LggKEAkCAgIBMGRaIRcVCgUODQkGBg0GCxoQCBADBQQEDhoJFhUVExgaCiQdHEQhMUADAw0ZESALDRsRFjAEAxIRGA8HGhYPBAc3HAIKEwsCAQILGhINEwgECQYCEhICAgEBDBsQIRkFDQgKEAkHBw4FBjYlIUU5DhwFIBwFCQUSGg0LFAwHBQIFAwsQCQIBAQEBARMXAgMCChAJAgMBEg8FEAsGBxAGDQcyVSUVMQ0KJR0WORkWODIODxkBAQIWIREfDiAXERUXExERFQkSFQ0QIhEtLhM0EwUDCg0YMhscPSMRIQsGDhEVJg4DBAJJmV8BAgEMLygeST8PCxMHAg0IBQcDDQkKGgsOEgoCBAIEBwYIDQYGBwsFFSsdBxcWDQIGBAgOBiYxGQECAgQrEQ0mDQMJBhZGSRomNg4SIREQFgsHGz8tBAYLFwoEEgcKEQHABAIBBQMTEAQKBA4BAQEBAv6GDhgwEBcvA1smRiAQLhQZRycVJ/zrDRgxEBY0AZ0JEggHDg0FFgUDEAHECxwNChAGCBslKzYOAQ0iGB8jFSIUHScIBh01/b0UDwgTBwkQBwYCLhAYDQclFRMcH/zUDRICCh8aBBIbDh0DFx0NDgEMBAoHEwUNAgkQEwcFCwkEAQYLAQsYBQ0YBQUX/DcpPxoaLhgIEQ4PBwUXCgsVFQ8SMCQBAQUHIVEuBhMTEyIBBggFEgsGIAsSIgwWCgoKGU1HEgcEjwsMBwwtNTAPAwgIBgsTCyAEDxsOCBEJCh0wQzgBCAkQBwwODxATJA4OEjA0DxUNAwgKBwsTGxMiAgMEAgwRCAcgCg8nCxYKDQcmOBcWIxEHBwgGH0tJGxILJT0cGzUdCAYOQYlECwwGepAQCiYGExMrJgsZFBADEhQFAxEPAwMEB0BGI0MiCwwHJ0olEzc3FQwBBSQtIQsMBypHIgYVChoFJAMFAkFEIxINCx8uDAoMCA0JFQczYjcNARQhFy4PEgoIFBAHBQoJIi8EBBEIBAoECRADBwsFBQYJCgcLFQwEBQMYDgksNCwLBwcICwQEDg4FEwICIg8EAgMCBAoFDh4KDgMPJBAKCho4Cw4WCA8fDh4YCisMCwwCBQIHBwkHAwUDBgsHBQ0KCAwLFw0bFQgTCQcQCDRsNgoOAo8sYR4QOjgLFBQBExwPDCQxKgUbRDBGVQkaFBMJIhkUHh0xLwUD8FUJGhQRDaAIEAkFEA4GFAYHEeoJBgUKBBEKNCkEDBIlHxsyITkhGBYiRS/XDxEJFAcKDggLWQYGHTIlDhUTEhU3GgUXGgwQCQQGABAAMf+5BBkCuAAaACYA2wDnAREBLAE4AUMBXgFtAXkBpwGvAcEBzQHZDERLsBFQWEFyAF8AIAACAA8AAQDcAHQAXAAiAAUABQACAA8BAgD7AO8AeAB3AAIABgAAABUBPwE9ASEA6wBWAAUAEQAAAS0BFwACABAAEQEUAAEAFAAQAWgAbQACABsACQFuAKUAWQA/AAQABwAIAYQBSQBkAAMAGgAHAUYAnACZAAMABgAaAaoBlAC9ALcAjQCHAAYABQAGAb8BuQG3AAMADQAFAa0AxAACAB4ADQHIAaEBkQDRAGYASgBGAAcABAAeAdQBywHFAX8ALQAFAB0ADgHXAdEBjgADABwAHQAQAEwBEAABABEAAQBLG0uwFVBYQXUAXwAgAAIADwABANwAdABcACIABQAFAAIADwECAPsA7wB4AHcAAgAGAAAAFQE/AT0BIQDrAFYABQARABIBLQABABMAEQEXAAEAEAATARQAAQAUABABaABtAAIAGwAJAW4ApQBZAD8ABAAHAAgBhAFJAGQAAwAaAAcBRgCcAJkAAwAGABoBqgGUAL0AtwCNAIcABgAFAAYBvwG5AbcAAwANAAUBrQDEAAIAHgANAcgBoQGRANEAZgBKAEYABwAEAB4B1AHLAcUBfwAtAAUAHQAOAdcB0QGOAAMAHAAdABEATAEQAAEAEQABAEsbS7AWUFhBeABfACAAAgAPAAEA3AB0AFwAIgAFAAUAAgAPAPsA7wB4AHcAAgAFABcAFQECAAEAAAAXAT8BPQEhAOsAVgAFABEAEgEtAAEAEwARARcAAQAQABMBFAABABQAEAFoAG0AAgAbAAkBbgClAFkAPwAEAAcACAGEAUkAZAADABoABwFGAJwAmQADAAYAGgGqAZQAvQC3AI0AhwAGAAUABgG/AbkBtwADAA0ABQGtAMQAAgAeAA0ByAGhAZEA0QBmAEoARgAHAAQAHgHUAcsBxQF/AC0ABQAdAA4B1wHRAY4AAwAcAB0AEgBMARAAAQARAAEASxtLsCJQWEF4AF8AIAACAA8AAQDcAHQAXAAiAAUABQACAA8A+wDvAHgAdwACAAUAFwAVAQIAAQAAABcBPwE9ASEA6wBWAAUAEQASAS0AAQATABEBFwABABAAEwEUAAEAFAAQAWgAbQACABsACQFuAKUAWQA/AAQABwAIAYQBSQBkAAMAGgAHAUYAnACZAAMABgAKAaoBlAC9ALcAjQCHAAYABQAGAb8BuQG3AAMADQAFAa0AxAACAB4ADQHIAaEBkQDRAGYASgBGAAcABAAeAdQBywHFAX8ALQAFAB0ADgHXAdEBjgADABwAHQASAEwBEAABABEAAQBLG0uwJ1BYQXgAXwAgAAIADwABANwAdABcACIABQAFAAIADwD7AO8AeAB3AAIABQAXABUBAgABAAAAFwE/AT0BIQDrAFYABQARABIBLQABABMAEQEXAAEAEAATARQAAQAUABABaABtAAIAGwAZAW4ApQBZAD8ABAAHAAgBhAFJAGQAAwAaAAcBRgCcAJkAAwAGAAoBqgGUAL0AtwCNAIcABgAFAAYBvwG5AbcAAwANAAUBrQDEAAIAHgANAcgBoQGRANEAZgBKAEYABwAEAB4B1AHLAcUBfwAtAAUAHQAOAdcB0QGOAAMAHAAdABIATAEQAAEAEQABAEsbS7AuUFhBeABfACAAAgAPAAEA3AB0AFwAIgAFAAUAAgAPAPsA7wB4AHcAAgAFABcAFQECAAEAAAAXAT8BPQEhAOsAVgAFABEAEgEtAAEAEwARARcAAQAQABMBFAABABQAFgFoAG0AAgAbABkBbgClAFkAPwAEAAcACAGEAUkAZAADABoABwFGAJwAmQADAAYACgGqAZQAvQC3AI0AhwAGAAUABgG/AbkBtwADAA0ABQGtAMQAAgAeAA0ByAGhAZEA0QBmAEoARgAHAAQAHgHUAcsBxQF/AC0ABQAdAA4B1wHRAY4AAwAcAB0AEgBMARAAAQARAAEASxtBeABfACAAAgAPAAEA3AB0AFwAIgAFAAUAAgAPAPsA7wB4AHcAAgAFABcAFQECAAEAAAAXAT8BPQEhAOsAVgAFABEAEgEtAAEAEwARARcAAQAQABMBFAABABQAFgFoAG0AAgAbABkBbgClAFkAPwAEAAcACAGEAUkAZAADABoABwFGAJwAmQADAAYACgGqAZQAvQC3AI0AhwAGAAUABgG/AbkBtwADAA0ABQGtAMQAAgAgAA0ByAGhAZEA0QBmAEoARgAHAAQAHgHUAcsBxQF/AC0ABQAdAA4B1wHRAY4AAwAcAB0AEgBMARAAAQARAAEAS1lZWVlZWUuwEVBYQHwAAg8VDwIVgCQBFBAJEBQJgAAaBwYHGgaAAAEADwIBD2kAFRchAgARFQBpEgERFhMjAxAUERBpABsICRtZGQEJAAgHCQhpCgEHJRgLAwYFBwZpDAEFIAENHgUNaSgfJwMeAA4dHg5pAAQAHRwEHWkiAQMDHGEmARwcEANOG0uwFVBYQIIAAg8VDwIVgCQBFBAJEBQJgAAaBwYHGgaAAAEADwIBD2kAFRchAgASFQBpABIAExASE2kAERYjAhAUERBpABsICRtZGQEJAAgHCQhpCgEHJRgLAwYFBwZpDAEFIAENHgUNaSgfJwMeAA4dHg5pAAQAHRwEHWkiAQMDHGEmARwcEANOG0uwFlBYQIkAAg8VDwIVgCEBABcSFwASgCQBFBAJEBQJgAAaBwYHGgaAAAEADwIBD2kAFQAXABUXaQASABMQEhNpABEWIwIQFBEQaQAbCAkbWRkBCQAIBwkIaQoBByUYCwMGBQcGaQwBBSABDR4FDWkoHycDHgAOHR4OaQAEAB0cBB1pIgEDAxxhJgEcHBADThtLsBdQWECPAAIPFQ8CFYAhAQAXEhcAEoAkARQQCRAUCYAAGgcKBxoKgAAKBgcKBn4AAQAPAgEPaQAVABcAFRdpABIAExASE2kAERYjAhAUERBpABsICRtZGQEJAAgHCQhpAAclGAsDBgUHBmkMAQUgAQ0eBQ1pKB8nAx4ADh0eDmkABAAdHAQdaSIBAwMcYSYBHBwQA04bS7AiUFhAlQACDxUPAhWAIQEAFxIXABKAJAEUEAkQFAmAABoHCgcaCoAACgYHCgZ+AAEADwIBD2kAFQAXABUXaQASABMQEhNpABEWIwIQFBEQaQAbCAkbWRkBCQAIBwkIaQAHJRgLAwYFBwZpDAEFIAENHgUNaSgfJwMeAA4dHg5pAAQAHRwEHWkmARwDAxxZJgEcHANiIgEDHANSG0uwJ1BYQJYAAg8VDwIVgCEBABcSFwASgCQBFBAJEBQJgAAaBwoHGgqAAAoGBwoGfgABAA8CAQ9pABUAFwAVF2kAEgATEBITaQARFiMCEBQREGkAGQAbCBkbaQAJAAgHCQhpAAclGAsDBgUHBmkMAQUgAQ0eBQ1pKB8nAx4ADh0eDmkABAAdHAQdaSYBHAMDHFkmARwcA2IiAQMcA1IbS7AuUFhAnAACDxUPAhWAIQEAFxIXABKAABYQFBAWFIAkARQJEBQJfgAaBwoHGgqAAAoGBwoGfgABAA8CAQ9pABUAFwAVF2kAEgATEBITaQARIwEQFhEQaQAZABsIGRtpAAkACAcJCGkAByUYCwMGBQcGaQwBBSABDR4FDWkoHycDHgAOHR4OaQAEAB0cBB1pJgEcAwMcWSYBHBwDYiIBAxwDUhtAowACDxUPAhWAIQEAFxIXABKAABYQFBAWFIAkARQJEBQJfgAaBwoHGgqAAAoGBwoGfgANBSAFDSCAAAEADwIBD2kAFQAXABUXaQASABMQEhNpABEjARAWERBpABkAGwgZG2kACQAIBwkIaQAHJRgLAwYFBwZpDAEFACAeBSBpKB8nAx4ADh0eDmkABAAdHAQdaSYBHAMDHFkmARwcA2IiAQMcA1JZWVlZWVlZQWUBsQGwAagBqAF7AXoBRQFEARMBEgDpAOgAKAAnAAEAAAG9AbsBsAHBAbEBwQGoAa8BqAGvAZ4BnAF6AacBewGnAXUBcwFcAVsBUAFOAUQBXgFFAV4BNAEyASoBKQEeARwBEgEsARMBLAEPAQ0BAQD/APQA8wDoAREA6QERAOMA4QDPAM0AyADGAMMAwQC7ALkAtgC0AKwAqgCjAKEAkwCSAIsAiQCFAIMASQBHACcA2wAoANsAGAAXAAsACgAAABoAAQAaACkABwAWKxMiNTY2NyYmNTQ2MzIXFhUUBgcWFjMGBiMGBiU3PgI3FhUUBwYGASImNTQ2NyYmJy4CNTQ2Njc2NjcnNjY3HgIVFBYXFjM2NyYmJyYmJwcmJjU0NxYWFyYmJzY2Nx4DFxYXNjY3PgI3LgI1NDY3FhYXBx4DFR4CFxYWMzI2NyYmIyIGByYmNTQ2MzIWFxYWFzY2NyYmJyYmIyIGByYmNTQ2MzIWFxYWFxYXFjIzByYmIyIGBxYWFxYyMwcmJiMiBgcOAiMiJicVFgYHDgIHBgYDNjY1NCYjIgYVFBYFIiY1NDY3FwYGFRQzMjY1NCY1NxYWFxYzMjcmJjU0NjcWFhUUBiMiJwYFIjU2NjcmJjU0NjMyFxYVFAYHFhYzBgYjBgY3NjY1NCYjIgYVFBYFNzY2NxYVFAcGBgUiNTY2NyYmNTQ2MzIXFhUUBgcWFjMGBiMGBgUnJicmJjU0NjcWFhUUBiU2NjU0JiMiBhUUFhMyPgI3LgMnBgYHBgYHBgYHJzY2NyYmJxYWFxQGBwYGByYmJycHBgYVFBYnJicGBgcWFiEyNjY3NjY3JicGBiMmJicXFhcmJic2NjcWFhcGBgcmJic2NjcWFhcGBtEVFiULCxEcEAgKCAECBgsGARURDSkBwgMUMjAPAREcPv5hQz4PCwoXCQodFgwRBw4wGgMLEhABAQIBARomIxgEBwQlSygWExIQOF4oBQsFCxIMAgUFBgE4NgUJBBMUCwIFCwYZEQUNBA8CBQYFAQYFAQseETBEGgsXDBYZEwIGKCEIEggRHw4IEQgNJg4QIBEWHQ8EBS4aEyQQGy4kExkIEwsXBhoFFhoTEioUCBMLFwkaAhwqJiIzNiUOFAsBBgoTT14sGzkmAgEICAYKEAHQDBEEBAoBARAICgIQAgUFBQsKBAEMDQYHBRQQDQsL/sUVFiULCxIcEAoJBwECBwsGARYRDSkzAwEICAcKEAF8BB9GHwERHD38vhUWJQsLERwQCQkIAQIGCwYBFRENKQGdBAIEAgINDQQBBv6LAgEICAYLEa0hV11VIAIJCgsEBhQQCBILIls+DThgGxcsFwIBAQ0IBh8bKCkFBBQHED0TCAcdNRMYOgHzIi4uIBAeDBcPK1A0Dx0MAxTzCxcKDBgHCxIKBxGhCxYKCxcIChMLBRQCNhMFEAcFEwwQHwoLGwQLBQEBDgYNGyMfBhASBwQFEggNFf1XOzYjOR8BBQMEEBIGCR4cBw4fChkLDAsWMy4NEzMPDwINJlMuJTscBA4WEg8NJUQoPHc1Cw8IHVNZUBpDYQULBhlJVCkiRTkPFwsEESUSDAwvNSoIDjQtBQcHEAoICQwLAwsJFxoEAwcUCwUFAgcaCQsQDAsCDwYbFg0KECINBwIBLAECBwkKDAIBLAECDRQSFAkFBA4QHA0bOTMRCg8CwAMIAgsPCAQKDnUREAUTBwEHBgYSDQcBEgQKCxEFBQgFGQQICgEMGgcOHAsaLxQEEQcFEwwQHgoKHAQKBQEBDwUNHEMDCAILDwgECg5AIAgbCwMGEAkOFJ4UBBAHBhINEB4KCR0ECwQBAQ8FDRwTASUZDRMHDgwDGSELEB5GAwgCCw8IBAoO/loeMTkbETtEQRhBWiEQHAslLhoRLEwcKD8aHDUbCiYLCQsBAhgbFCUTLRw3J+ooOQofFhYKBQ4NBgwCCwwWGAEFByAKaggWCgoZCA0RCQgWRQkVCgoYCQoVCAcYAAAA//8AI/8WAdICsAImAt4AAAEGBXZdpwAJsQIBuP+nsDUrAAAA//8AIv8PAnMCnwImAt8AAAEGBXZhoAAJsQEBuP+gsDUrAAAA//8AIwAAAdIC4gImAt4AAAEHBXkAlAA4AAixAgGwOLA1KwAA//8AIwAAAdIDjgImAt4AAAEHBaIA7QGsAAmxAgG4AaywNSsA//8AIgAAAnMDfQImAt8AAAEHBaIBOwGbAAmxAQG4AZuwNSsA//8ANQAAAnMDfQImAcAAAAEHBaIBOwGbAAmxAQG4AZuwNSsAAAIAFf/rAKYAuAAXACAAc0ASGgEEBQgBAgQEAQECAwEAAQRMS7AYUFhAHQADAAUEAwVpBwEEAAIBBAJpAAEBAGEGAQAADwBOG0AdAAMABQQDBWkHAQQAAgEEAmkGAQAAAWEAAQEQAE5ZQBcZGAEAHRsYIBkgEhANCwcFABcBFwgHFisXIiYnNxYzMjc1BgYjIjU0NjMyFhUUBgY1MjcmIyIGFRRRDCULBhYXOhUJIA4xIRYbJRcmGBMQJA0VFQ8JDQhCBAkNMCEvPCYcMR5rCTkaCx0AAQArAAABLwB3ACEAIUAeGxYPDgQBSgABAQBhAgEAAA8ATgIAGRcAIQIhAwcWKzMiJicmNTQ2NzY2Nzc2NxcGBgcOAgcWMzI2NxYVFAYHBqoWKxUpCAQHDwkiFgcMBQ8OFRUNCyZMHEUZBiAbJQMEBgcHIQIGCwQMCBAHDhQEBggICA8KBwYHCgwFBgACADH/9gILAtUACwAbADFALgABAAMCAQNpBQECAAACWQUBAgIAYQQBAAIAUQ0MAQAVEwwbDRsHBQALAQsGBhYrBSImNTQ2MzIWFRQGJzI3NjU0JyYjIgcGFRQXFgEddHhzeXR6dng8IjdAIDU+IzM3IQq7tbu0vLO6tkspQbvIPR8sQLi+QCcAAQBZAAABYwLKAA4AH0AcCwoEAwEAAUwAAAEAhQIBAQF2AAAADgAOHAMGFyshETQ2NwYGBwYGByc3MxEBDQICBw4IGTcdLsFJAfMhPhwHDgYVLRg7lv02AAEAMAAAAggC1AAgADNAMA8OAgIAAQEDAgJMAAEAAAIBAGkAAgMDAlcAAgIDXwQBAwIDTwAAACAAICskKgUGGSszNTc2Njc2NjU0JiMiBgcnNjMyFhcWFRQGBwYGBwcVIRUwuyo4Dh4YQzssTjIvZXZLaRYOIiIROSmVAWlJvSs9EilAKDhAHiY7VDo1JSouWC0WPyeTBFAAAQAt//YCAwLUACcATkBLFwEDBCEBAgMEAQECAwEAAQRMGAEEAUsABQAEAwUEaQADAAIBAwJpAAEAAAFZAAEBAGEGAQABAFEBABwaFRMPDQwKCAYAJwEnBwYWKxciJic1FhYzMjU0IyM1MzI2NTQmIyIGByc2NjMyFhUUBgcVFhYVFAbxOGAsLWgwtchFRlRfRT0vUTQsLHNAaHVNR1ZUjQoRFlIWGY1+S0dAMzoaIjwjJ19SRFgPBApYR2ZvAAAAAgAVAAACKALOAAoAFAA6QDcQAQIBAwEAAgJMAAECAYUGAQQABIYFAQIAAAJXBQECAgBfAwEAAgBPAAAMCwAKAAoRERIRBwYaKyE1ITUBMxEzFSMVJSE1NDY3IwYGBwFr/qoBUFtoaP6rAQADAgQFGwuiSwHh/iNPovHhIU8uDTMOAAAAAQA///YCAwLKABsAR0BEEw4CAgUNAwIBAgIBAAEDTAADAAQFAwRnAAUAAgEFAmkAAQAAAVkAAQEAYQYBAAEAUQEAFxUSERAPDAoHBQAbARsHBhYrFyInNRYWMzI2NTQjIgcnEyEVIQc2NjMyFhUUBvh3QiRmMFRcszlBLBsBZv7lERw3G3CAjgonUxYZUEqRDxwBUVDPBQZwYm9/AAAAAAIAN//2Ag0C1AAZACkAS0BICQECAQoBAwIRAQQFA0wAAQACAwECaQADAAUEAwVpBwEEAAAEWQcBBAQAYQYBAAQAURsaAQAhHxopGykVEw0LCAYAGQEZCAYWKwUiJicmNRAhMhcVJiMiBgcGBzM2MzIWFRQGJzI2NTQmIyIGBwYVFBYXFgEqSW4eHgFHOR8lMUpnHR8EBjZzYW56akRJR0MrRhYVKCIjCktHSGEBowlLDDo+Q2ZUdWRvgUpXT0dNIx8fIjFVGRgAAQAsAAACCwLKAAYAKkAnBQEAAQFMAwECAAKGAAEAAAFXAAEBAF8AAAEATwAAAAYABhERBAYYKzMBITUhFQGIASX+fwHf/t4CelBE/XoAAAADADH/9gIKAtQAFgAkADAAOUA2KxEFAwMCAUwAAQACAwECaQUBAwAAA1kFAQMDAGEEAQADAFEmJQEAJTAmMB4cDAoAFgEWBgYWKwUiJjU0NyYmNTQ2MzIWFRQGBxYWFRQGAzY2NTQmIyIGFRQWFxYTMjY1NCYnBgYVFBYBIHJ9lEM6dGFidEFHVkl+a0M4RTs6RBoZGTRGUExaRkFRCmRcekUmWDdOXFtQNVUlKlw7WGsBphxDKzE2NzAfLxIT/olCOi9JIR5KNDhBAAACADL/9gIIAtQAGwArAEpARwsBBAUDAQECAgEAAQNMAAMABQQDBWkHAQQAAgEEAmkAAQAAAVkAAQEAYQYBAAEAUR0cAQAnJRwrHSsWFA8NBgQAGwEbCAYWKxciJzUWMzI2NzY2NyMGIyImNTQ2NjMyFhcWFRADMjY3NjU0JicmIyIGFRQWwDsfJDRQZxsOEAEGOHJebzhnRUluHR72LEgVFCYjIi5DSkYKCksNQD0fUTNTc2VKbDtLR0hh/l0BWyMfHyExVBkZV09HTAACAAsAfgC+AZsACAAYADFALgABAAMCAQNpBQECAAACWQUBAgIAYQQBAAIAUQoJAQASEAkYChgGBAAIAQgGBxYrNyImNTQzMhUUJzI3NjU0JyYjIgcGFRQXFmQsLVlaWhoLDAwKGxoKCwsLfktFjY2QKBoaNDQYGRkZNDIbGgAAAAEAFgCEAIoBlgAOACZAIwsKBAMBAAFMAAABAQBXAAAAAV8CAQEAAU8AAAAOAA4cAwcXKzc1NDY3BgYHBgYHJzczEWIBAQQIBAMVEhRKKoSxCh0RBAgDAhIPIDv+7gAAAAEADgCEALEBnAAaADZAMw0BAAEMAQIAAQEDAgNMAAEAAAIBAGkAAgMDAlcAAgIDXwQBAwIDTwAAABoAGhgkKAUHGSs3NTc2NzY1NCYjIgYHJzYzMhYVFAYHBgYHMxUOQiQKChUREhwPFCYrJSoSEAUnI3CEIkgnEREXFBQPDB0kKSMWJhMHKyQnAAAAAQAKAH4AuQGcACIATUBKFQEEBRQBAwQdAQIDAwEBAgIBAAEFTAAFAAQDBQRpAAMAAgEDAmkAAQAAAVkAAQEAYQYBAAEAUQEAGRcTEQ0LCggGBAAiASIHBxYrNyInNRYzMjU0IyM1MzI2NTQmIyIHJzY2MzIWFRQHFRYVFAZVKCMmJjk/ISEdGxYSHiIVFCoWJi0sMzN+EygWMSwjGRQSFRkeEQ8mIS8SAQ00Jy0AAAACAAYAhADFAZcACgASADdANBABAgEDAQACAkwAAQIEAVcFAQIDAQAEAgBnAAEBBF8GAQQBBE8AAAwLAAoAChEREhEHBxorNzUjNTczFTMVIxUnMzU0NjcGB3ZwcSsjI3RIAQECGoQ+ILWwJT5jPBEgEAUtAAEAEQB+ALgBlgAfAEZAQxcBAgUPBAIBAgMBAAEDTAADAAQFAwRnAAUAAgEFAmkAAQAAAVkAAQEAYQYBAAEAUQEAGxkWFRQTDQsIBgAfAR8HBxYrNyImJzUWFjMyNjU0IyIGByYmJzczFSMHNjYzMhYVFAZZFSgLECcPGh86DxUJBAkFCotqBQkRBykxMH4KBysLDBkcMwUCBAYDhSRFAgIvJiwyAAAAAAIADAB+AL8BmwAYACMAS0BICAECAQkBAwIPAQUDA0wAAQACAwECaQADAAUEAwVpBwEEAAAEWQcBBAQAYQYBAAQAURoZAQAfHRkjGiMUEg0LBgQAGAEYCAcWKzciJjU0MzIWFxUmJiMiBgczNjYzMhYVFAYnMjY1NCMiBhUUFmgrMXwLEAcHEgsoKQMBCCAYIykuKhYbLxYgHn5DOaECAyYEAi4zDRMuKCw1JR0eMxoTHCUAAAEAEACEALoBlgAGACpAJwUBAAEBTAMBAgAChgABAAABVwABAQBfAAABAE8AAAAGAAYREQQHGCs3NyM1MxUHLmF/qmKE6ycg8gAAAAADAA4AfwC7AZwAFwAgACsAOUA2JxAGAwMCAUwAAQACAwECaQUBAwAAA1kFAQMDAGEEAQADAFEiIQEAISsiKx4cDAoAFwEXBgcWKzciJjU0NjcmNTQ2MzIWFRQHFhcWFhUUBic2NjU0IyIVFBcyNjU0JicGFRQWZSkuFRckLSIkKikgDAMCLyYRFScnJhkWGh0mFn8qIxYkDhkoICclISkXDxwGDwgjLKgIFREkJByXFxIUGAsTJBIXAAAAAAIACgB+AL4BnAAWACIASkBHCAECBAMBAQICAQABA0wAAwAFBAMFaQcBBAACAQQCaQABAAABWQABAQBhBgEAAQBRGBcBAB4cFyIYIhIQDAoHBQAWARYIBxYrNyInNRYWMzI3IwYjIiY1NDYzMhYVFAYnMjY1NCYjIgYVFBZCExEHEw1NBgIWJyUqMSYrMj8fFx8dGBYaF34FJwMEYx8uKSo0QTlUUIwbEh4iHhoZHAAAAQAD/7sAmwDkABsAQkALAwEBAAFMEwQCAEpLsB1QWEAMAgEBAQBhAAAAEAFOG0ARAAABAQBZAAAAAWECAQEAAVFZQAoAAAAbABsmAwcXKxcmJic3FhYzMjY3NiYnJiY3NjY3FhYXFhUUBgY8CiQLBRAVBhweCAIVCAUGAQINDAYPBBcbK0QBGA0VBAQqKhAkEw4WChARCw0bDD0fKkYp//8AAwGGAJsCrwMHBOAAAAHLAAmxAAG4AcuwNSsAAAAAAQAC/9EAjQDqABUAHUAaAQACAEkAAQAAAVkAAQEAYQAAAQBRJCgCBxgrFzU+AycmBiMiJjU0NjMyFhUUBgYCFCcdCggHBg0QHiAXHiQgPS8kBio1LQkICRIXGBorIChVQwAAAAEAAv+jAMcA6gADABdAFAAAAQCFAgEBAXYAAAADAAMRAwcXKxcTMwMCfkeMXQFH/rkAAAABAEwAjADuAToADgAGswgAATIrNyYmJzY3NjY3FhYXFhUUtxs1GxENBw8JGycMF4wUIw4bGw0ZDRAZCxQICwD//wCaAIwBPAE6AAYE5E4AAAEATQAAAPACOAASAAazDAABMiszJy4CJyYmJyY1NDceAhUUBuImAQcQDwohFQg2ITEbBwlFZVcvJEwpDRApIDuHgjUxXwAA//8AhAAAAScCOAAGBOY3AAABAE0AAAGEAjgAHQAoQCUXAQEAAUwTEgwDAEoBAQFJAAABAQBZAAAAAWEAAQABUSMvAgcYKzMnLgInJiYnJjU0NxYXFjMyNxcUIyInFhYVFAYG4iYBCBEQCx4UCDYeFhcqVw0oixIRDQ0CBglKaVguI0glDRArHjc6EoIIzgcxVCQfNT4AAAAAAQBNAAABxAI4ADAARUBCGBUCAQIoAQMALQEEAwNMFgwCAkoBAQRJAAIBAoUAAAMEAFkAAQADBAEDaQAAAARhAAQABFEsKiclIB4bGRIQBQcWKzMnNCYmJyYmJyY1NDcWFhcWMzI1NCc3FhUWMzI1NCYnFxYWFRQGIyInBgYjIicWFRTiJgcQDgohFwg2DhsLGCA4Ax4OEQwsAQEpAgIxJRgYCiQYEBQaCUFfVjAkUi0NECofGzkdEksWGAoxLQZGCQ8EBBAaCzdFDhQaB2BJYgAAAP//ACYAAAGdAjgABgTp2QAAAQBE//4BfgJUAC4AJ0AkKQEAAQFMKB0ZBgQBSgABAQBhAgEAAA8ATgEAJyUALgEuAwcWKwUiJjU0NjcmJjU0Njc2NjcWFRQHBgYHBgYHFhYXFwYGBwYGFRQWMzI3FwYGBwYGARhPVSceNEEZERJHPAMeFTYOBw0GFjsoGxIrEQsORzE1JQoJEQsMHwIzMh1gLhYxFRpKGRk1HwkMLhkRKQsFDQUYGQk7EzobEx8KFB8KEgwZCwsPAAIALwA8AU8CAwATACIAMUAuAAEAAwIBA2kFAQIAAAJZBQECAgBhBAEAAgBRFRQBAB0bFCIVIgoIABMBEwYHFis3IiY1NDY2NzYzMhYXFhUUBgYHBicyNjY1NCYmIyIGBwYVFK1DOxYoGiYiJTsQEBQlGiUfHDEgHDEfFC8QDzxGSitfWiEyPDk3PzxJLREZVBsqFjRTMT8wMCJS//8AVgA8AXYCAwAGBOwnAAABABn//wF+AkMAGQApQCYTDAcDAAEBTBINAgFKGQEASQABAAABWQABAQBhAAABAFE0KAIHGCsFJiY1NDY2NwYjIiYnNxYzMjY3FwYGFRQWFwFDJSQGCQIoPihMGBc8OSFJKAwJCCclAVyeVBY7MwkGDwtVGQUFECI7GnWwSgAAAAEACQAAAasCNgAbAAazCgABMiszLgMnJjU0NjceAxc+AzcXDgIHBgbMGiopMCAGGxYXNDIoCgsiKiwVKik9LhUQG1d8ZGE8CwsTKRAjcIOBMzyBfWslSFOTkFANEwAA//8AEQAAAbMCNgAGBO8IAAABAAgAAAGrAjYAGQAGswYAATIrMyc+Ajc3HgMXFhUUBy4CJyYmJw4CMyssOy4VNhgpKjIhBTEVMS4UDRQGDzE7R1eOjlQoUntmZD0MCigkIGV3PChLH1OrlgAA//8AEgAAAbUCNgAGBPEKAAACACr//wFyAkkAEgAfADRAMQQBAAIBTBIBAEkAAQADAgEDaQQBAgAAAlkEAQICAGEAAAIAURUTGhgTHxUfJSUFBxgrBS4CJwYjIjU0Njc2MzIWFRQXAzI2NyYmIyIGFRQWFgE3GB8OARsbkSgfHx06PE+dBwwFCConGSUkOwE/aGAyDH4hTBkZZG6cnQFCAQE4OCEPEx4R//8ATACMAO4BOgIGBOQAAP//AE0AAADwAjgCBgTmAAD//wBNAAABhAI4AgYE6AAA//8ATQAAAcQCOAIGBOkAAAABAE0AAAG3AkAAMQA1QDIaDAIBAC0lIiAbEQYCAQJMAQECSQACAQKGAAABAQBZAAAAAWEAAQABUSwqHhwZFwMHFiszJyYmJy4CJyY1NDcWFxYWFyY1NDY3NjMyFwcmIyIGBxQXNjY3FhUUBgYjIicWFRQG4iYBDA0HEx0WCDYmGAwdEBgcFhkZKS4XIB4NGBNSFCgTBilDKCUmFwYJWXQxITtDLA8PKh5FRwcIAhIfGjgREUEXFgkKQAgGEAoSDhklEwxeQC5g//8ANQAAAZ8CQAAGBPjoAAACAC8ABwGXAkkAGwAvADFALhoBAAIBTCkfDQMCSgMFAgICAGEBBAIAAA8ATh0cAQAjIRwvHS8ZFwAbARsGBxYrNyImNTQ3PgI3JjU0Nx4CFxYVFAYHBiMiJwYnMjc3FBYzMjY1NCYmJw4CFRQWnzU7NQYMFhUHFENUMA4UGxkZICsdHCowBiQcHxokKUkyITUfKAdaUVpnDRQkIw0XKCI9XEwnPEMxVRkYJSVVSRMrMSseI1tjLiZaWSYjNgAAAP//AC8ABwGXAkkABgT6AAAAAQA8AAABngI3ACgAL0AsEQEBAB8YEgMCAQJMBgECAkkAAgEChgAAAQEAWQAAAAFhAAEAAVEUJi0DBxkrMyc+AzcmJjU0Njc2MzIWFwcuAiMiBxQXFhc2NjcWFhUUBw4DWR0UHRshFjdEJSAhKCI+MBIeIRYPQjFAJDQaQCgEBzotRzs7MCY1KysbDkUrL1kaGyApIwwKAiNLHBABEyISBR4JHCYeOD9TAAAA//8APAAAAZ4CNwAGBPwAAP//AAkAAAGrAjYCBgTvAAD//wAIAAABqwI2AgYE8QAA//8AK///AXMCSQAGBPMBAAABABsA0ACCAT8ADAAGswcAATIrNyYnNjY3NjcWFxYVFF8iIgUKBAcNJA4O0BoSCREIERAUDQsHCAAAAAABACYAkgB9AbQAEgAGsw0AATIrNycuAicmJicmJjU0NxYWFRQGdBgBAgYGCRQFAgMjGRsEkgYhLSYVICgJBgkEGxQsZi4cJgABACUAkgDGAbYAIgAjQCAVFAIASgEBAUkAAAEBAFkAAAABYQABAAFRHBoTEQIHFis3JzQmJyYmJyY1NDY3FhYXFhYzMjcXFAYHBgYjIicWFhUUBnQZAgcGERAGHwUIDQYHDQcoBhkECAccFwkEBQIBkgYiRBkXKhwMCg8ZAw8fDwQCRAUVKBERFgIWJg4YMQABACUAkgDoAbUAOABvQBccAQECLgEDADMBBAMDTBkMAgJKAQEESUuwGlBYQB4AAgEBAnAAAAMEAFkAAQADBAEDagAAAARhAAQABFEbQB0AAgEChQAAAwQAWQABAAMEAQNqAAAABGEABAAEUVlADTEwKyolIyAeExEFBxYrNycmJicmJicmNTQ2NxYWFxYWMzI2NTQmNTcWFhUWFjMyNTQ0JxcWFhUUBiMiJicGBiMiJxYWFRQGdBkBAgYGEw4GHgYIDQYFCwYLDgETBAUDBgMSARsBARkUBQ0HBhELBggFAwKSBiM/HBcuGQwLDhkDDx8QBAEQFgULBgcMGQwBASQECAMDCBAHHSoEBAoOAhYmDhouAAABACAAkQDEAcMALQAtQCooAQABAUwnHhoRBgUBSgABAAABWQABAQBhAgEAAQBRAQAmJAAtAS0DBxYrNyImNTQ2NyYmJyY1NDY3NjY3FhQVFAYHBgYHFhYXFw4CFRQWMzI3FwYGBwYGhyItFg0QGQgKFxULHBIBEhkGDAcJGxQRChkSKhAcEgYJCwcFD5EaGg4xEgkQBwoLEjESCREJAgcEECEPBAcFCg8FJgkeGwcKCgYLEREFBAQAAgAYALEAqwGaABIAHgAxQC4AAQADAgEDaQUBAgAAAlkFAQICAGEEAQACAFEUEwEAGhgTHhQeCQcAEgESBgcWKzciJjU0Njc2MzIWFxYWFRQGBwYnMjY1NCYjIgYGFRRZIh8cExIRFB0IBQMRFRYPEx4eFwkUDbEkJSRNFxggHA4eECgqEA81FAsmMhwnESMAAQAKAJEAxAG6ABcAKUAmEQoGAwABAUwQCwIBShcBAEkAAQAAAVkAAQEAYQAAAQBRMycCBxgrNyYmNTQ2NwYjIic3FjMyNjcXBgYVFBYXnhISBQQSHS8bDyAbFCQTBwQEExORLE4pDi4KAw02DAIDCg0jFjVQJAAAAQAFAJIA2gG1AB4ABrMLAAEyKzcmJicmJicmNTQ2NxYWFxYWFzM2Njc2NjcXBgYHBgZlCBYLCRYUBBYJDhkKChQEAQUPCgkWChscJhEHEZIeRBgVLSgHCw8XBxc6Ghs/GBw9HR41Ey43aDwGDQAAAAEABACSANsBtAAgAAazCgABMis3JzY2NzY2NzY2NxYWFxYWFxYVFAYHJiYnJiYnIwYGBwYgHAsVCgsVCggRCQoXCAodDQQWCg0bCQkUBQEEEAkVki4VLBkaPycHDAcfRRUYORkHCBAYCBM+GxhDFhw9HT4AAgAWAJIAvwG+ABQAHwA3QDQYAQMCAwEAAwJMFAEASQABAAIDAQJpBAEDAAADWQQBAwMAYQAAAwBRFRUVHxUfLiQkBQcZKzcmJicGIyI1NDY2MzIWFhUUFhcWFycyMjcmJiMiFRQWmhAUAQ0MRhQfEBkZCgQFCRhWAgQDAxMTFySSJ0gkBEsSJhofMhsRIhIpKp0BEhwREA4AAAAAAQAlAJIA4wG6ADYAOkA3GwEBACYhHAMCATEBAwIDTAEBA0kAAAABAgABaQACAwMCWQACAgNhAAMCA1EuLSMiIB4ZFwQHFis3JyYmJyYmJyY1NDY2NxYWFxYWFyY1NDYzMhYXByYmIyIHFBc2NjcWFhUUBgcGIyImJxYWFRQGdhgBBQgHDhAGDxIDCxAGBgsECxwWDBUNDwsLBwsLKAgSCgMCCgwcGgcRCQUDA5IGJj8cGCccDgkJEg4CEyYTAgIBCBETKRERDwUEBCIBAgcGBw4FBxEGDAMBFCQNGTAAAAACABgAlgDRAb4AGAAoADhANRcBAAIBTCQcCwMCSgMFAgIAAAJZAwUCAgIAYQEEAgACAFEaGQEAHx0ZKBooFhQAGAEYBgcWKzciJjU0NjcmNSY2NxYWFxYWFRQGBiMiJwYnMjc3FDMyNjU0JicGBhUUUSAZIRUCAQkGHiwPDhANGBMWDw8XGAQWHQsPMSAVH5Y3IShHIggGDRoKHDAXFzIeFSwdExM3IQwtDQ8cQxwZPh4iAAABAB4AkgDWAbUAJQAvQCwOAQEAGhQPAwIBAkwEAQICSQACAQKGAAABAQBZAAAAAWEAAQABURMlKgMHGSs3JzY2NyYmNTQ2NjMyFhcHJiYjIgcUFhc2NjcWFhUUBgcGBgcGBjIUEh0UGiUSIRYRIhgMDRcMIBQnLQsaDwMFHQ4THA4OHZIfISwWBiQWFi0eEhQVAwQNFhkBBw0GBBIGDRcJDRcODScAAAABAFL//gG2Ak0AMQAfQBwxHhADAUkAAAEBAFkAAAABYQABAAFRIyEsAgcXKxcmJyYmJyY1NDY3NjYzMhYXPgI3FhQVFAcOAgcHLgIjIgYVFBcWFhceAhcWFhXoDCkeKAkSGxwLHREdPxYPIS4iAh4VGhYPLQokLRgQEwkHGBAcHQ8FBAQCTU84UhkzHiVVGQsMSDgiMCoZBxEKPx4THigkEi9KKxMNEB4VNiE3SDUdFSgUAAABABsAAAFxAi4AIgAfQBweGQIBSgABAQBhAgEAAA8ATgEAHRsAIgEiAwcWKzMiJiYnJiY1NDY3Njc2NjcWFRQHBgYHBgYVFhYzMjcUBgcGshk6MAsGAxkYMo0NGA0UGEdZESIkEVMtUE4dFzEFBwMWGxYnUStVog4gEBYaHx5XbhctOxQIDQwmLAUJAAEAQwAAAZcCOAAeAChAJRkBAQABTBMSDAMASgEBAUkAAAEBAFkAAAABYQABAAFRJS8CBxgrMycmJicmJicmNTQ2NxYXFjMyNxcUBwYjIicWFRQGBtgmARAQDSIXCB0ZHhQnOlcMKCsiPScbGwIGCV6JOC1SKxANFCgNNDcYggh1MicMbEIfNT4AAAAAAQBT//4BrAJUACsAI0AgJiUbFw8EBgFKAAEBAGECAQAADwBOAQAjIQArASsDBxYrBSI1NDcmJjU0Njc2NjcWFQYGBw4DBxYWFxcOAhUUFjMyNjcXBgYHBgYBK6hKM0cZExNGOgQBCQoDISsmCRRMNxwkPSVIMR0zJwsMIgwQIgJlPmsUNxMaTBkaNB0JDw4eDQQbIh4IFiALOxY6OBUUHwgMEgoiCg0RAAAAAAEAG///AZ0CQwAaAClAJhQMBwMAAQFMEw0CAUoaAQBJAAEAAAFZAAEBAGEAAAEAUTUoAgcYKwUmJjU0NjY3BiMiJic3FhYzMjY3FwYGFRQWFwFiJSMFBwQqSytXGBcYTB8sUCUMCgcnJQFcnlQSNDQTBhAKVQoPBQUQKzEbdK9MAAAAAAIAQ///AZQCSQATACAAOEA1FwECAwQBAAICTBMBAEkAAQADAgEDaQQBAgAAAlkEAQICAGEAAAIAURUUGxkUIBUgJiUFBxgrBS4CJwYjIiY1NDY3NjMyFhUUFwMyNjc3JiMiBgYVFBYBWRgdDwInIjpNKyAgIjk8T64JFQwBDE4RIhZDAT9iWC4OTTwlUhkYZG6cnQEuAwIBgBQfERwmAAABAC7/vgCnAEAACwAGswgAATIrFyYmJzY2NzY3FhUUfRgoDwUMBAoNTUIRGgkLEwoSFC0OCwAAAAEAKf9/AMAAdAAIAB5AGwAAAQEAVwAAAAFfAgEBAAFPAAAACAAIEwMGFysXNjY3MxcGBgcpEBkJXgcNLB2BPXs9CzR1QQAAAAIALv++AKcBoQAMABkACLUWDQkAAjIrEyYmJzY2NzY2NxYVFAMmJic2Njc2NjcWFRR9GCgPBgsFBQsGTSoVKBIGCwUFCwZNASARGQoKEwoKEwksDwr+Yg8aCwsTCgoTCS0OCwACAC7/vgCnAmYACwAaAAi1FQwFAAIyKzcmJic2Nx4EFQcmJic2Njc2NjcWFxYVFFUHDwwaKgMEAgEBBRYiFwYLBQULBisQEnKF1WIaHjVUTlVrScgRFg0LEwoKEwkYDg8GCwAAAgBK/38A7QImAAsAFQAwQC0AAQQBAAIBAGkAAgMDAlcAAgIDXwUBAwIDTwwMAQAMFQwVERAHBQALAQsGBhYrEyImNTQ2MzIWFRQGEyYmJzczFhYXF4kaJSUaGiMjCBgwDgdeBxAODgGeICQlHx8lJCD94TZ/NQssUjs8AAEATP9/AOMAdAAJAB5AGwAAAQEAVwAAAAFfAgEBAAFPAAAACQAJFAMGFysXJiYnNzMWFhcXoh4rDQdeBhQKDoFDdDMLKGUsPAAAAAABAEj/8gDEAHkACQA2S7AkUFhADAABAQBhAgEAAA8AThtAEQABAAABWQABAQBhAgEAAQBRWUALAQAGBAAJAQkDBxYrFyImNTQzMhUUBoUdID0/IQ4kIENDHyUAAAIASP/yAMQCygADAA4AT0uwJFBYQBUAAAQBAQMAAWcAAwMCYQUBAgIPAk4bQBoAAAQBAQMAAWcAAwICA1kAAwMCYQUBAgMCUVlAEgUEAAAKCAQOBQ4AAwADEQYHFys3AzMDByImNTQzMhYVFAZqGWsZHR4gPh4gIckCAf3/1yQgQyMgHyUAAAABACgAAAEaAE4AAwAeQBsAAAEBAFcAAAABXwIBAQABTwAAAAMAAxEDBhcrMzUzFSjyTk4AAQAoAAABGgBOAAMAHkAbAAABAQBXAAAAAV8CAQEAAU8AAAADAAMRAwYXKzM1MxUo8k5OAAIAOwBbAYEBmQANABsAJkAjGhYVFBMMCAcGBQoASgMBAgMAAHYPDgEADhsPGwANAQ0EBhYrJSInJiYnNxcHFxYWFwYjIicmJic3FwcXFhYXBgFQFQoWLBd3IVgwDhwPFrkWChYrFnYgVzARHQsWWw4gQiGtH4c7ESIRGQ4gQiGtH4c7FiIMGQACAB0AUAFjAY8ADQAbABhAFRsUEA8NBgIBCABJAQEAAHYsJwIGGCs3JzcnJiYnNjMyFxYWFwUnNycmJic2MzIXFhYX7SFYMA0dDxcZFwoWKxb+7CFYGBYoExcaFgkXKxZQIIc7ESIRGQ8gQiCuIIcdHDAWGQ8gQiAAAAAAAgA/AFsBgAGZABAAIQAiQB8fGhkOCQgGAEoDAQIDAAB2EhEBABEhEiEAEAEQBAcWKyUiJicmNTQ2NxcGFRQWFwYGIyImJyY1NDY3FwYVFBYXBgYBURwwERdBMSBUNy4HG6odMhAWQTEgVDguBxtbIBojLy9hIh8/RShGFAsOIRskLC9hIh8/RShGFAsOAAAAAgAdAFEBXwGPABMAJgAVQBIcFQkBBABJAQEAAHYgHisCBxcrNyc2NTQmJyYmJzY2MzIWFxYVFAYHJzY1NCcmJic2NjMyFhcWFRQG7CBUIB0JFAsIGg0hNw8OP9EhVDkKFgwIGg0eMhAVPlEfPkYgNxUGDAQLDisiHiItYCQfPkY+KwcNBQsOIxwiLC1gAAAAAgA1/wUBE//JABkAIwBRsQZkREBGEAEDASIdBQQDBQIDAgEAAgNMDgoHAwFKAAEAAwIBA2kFAQIAAAJZBQECAgBhBAEAAgBRGxoBACEfGiMbIxQSABkBGQYHFiuxBgBEFyInNxc3JjU0NjcWFhcHFRc2NjMyFhYVFAYnMjY3JiYjIgcWr0Y0Bi4SGBgRAwYCCAMPIwwOHxY6KBciExAYCx0pF/sgGAcUVgkIFAQKEQYGDjMPFhQaCiApKwoHDggkAwAA//8AgQFDAQkBzQIGBdsAAP//AFcBQwEaAqQCJgXbAAABBgXzKmoACLEBAbBqsDUr//8AcP8+APj/yQIGBecAAP//ADUBFQCzAgUCBgXfAAD//wA6/tcAt//IAgYF5QAA//8AOv7XALf/yAIGBeUAAP//AEkBRAFeAdMCBgXjAAD//wAu/z4BQ//NAgYF3AAA//8ALv8+AUP/zQIGBdwAAP//AC7/vgE8AL8AJgUUAAAAJgUUSH8BBwUUAJUAAAAIsQEBsH+wNSv//wBeAVUBXwIaAgYF4gAA//8AWf76AVn/wgIGBd0AAP//AFsBRAFbAgwCBgXhAAD//wBbAUQBWwIMAgYF4QAA//8AXP8GAV3/ywEPBeYBsP7WwAAACbEAA7j+1rA1KwD//wBOAS0BFwHtAgYF3gAA//8ATv8EARf/xAIGBeQAAAACAGL/FwDX/8wADAAZAAi1FQ0IAAIyKxcnNCcmJjU0NxYWFRQHJzQnJiY1NDcWFhUUwwUNAQEcBgZhBQ0BARwGBukBRDMFCAQgDB0xJSYcAUQzBQgEIAwdMSUm//8ATv8EARf/xAIGBeQAAP//AE7/BAEX/8QCBgXkAAAAAgBAAXMApwHtAAsAFgBksQZkREuwD1BYQB4AAQMCAXAAAwIDhQUBAgAAAlkFAQICAGIEAQACAFIbQB0AAQMBhQADAgOFBQECAAACWQUBAgIAYgQBAAIAUllAEw0MAQASEQwWDRYHBQALAQsGBxYrsQYARBMiJjU0NjMyFhUUBicyNjU0JiMiBhUUchYcIBMUICETDBMXDgcSAXMgGRwlJR0XISQNBwkWDQwaAAwAFf9DA4UCswAPABIAFQAdACAALAAvADIAPgBBAEQARwCDQIAeHBkTCwUGCgUyLgIMCjEvDAQECwwwLQIJC0RBHRgNAwYGCQVMEggCAUoEAgIBCAcCBQoBBWcACgAMCwoMaREBCxABCQYLCWkODQIGAAAGVw4NAgYGAF8PAwIABgBPNDMiIUdGQ0JAPzo4Mz40PigmISwiLBQTERMRFBIUERIHHysFJyM1Jzc1MzcXMxUXBxUjAzMnBTcjEzM3NScjBxUBNSMDIiY1NDYzMhYVFAYlNQcFNycFMjY1NCYjIgYVFBYHMycFMzUFNyMBzH+4gIC4f4K3gIC316tW/ul4eKTopKTopAIweKEkNTUkJjU1/qNUAsRTU/7HGiIiGhgiIv94eAG4eP7nVqu9gbaBgbaBgbaBgbYCblbsd/3Ro+ikpOgBFXf+jjUlJjQ0JiU1BapVVVVVkSMZGSMjGRkj23d3d+xVAAEAKQAAAMsATAALABlAFgAAAAFfAgEBAQ8BTgAAAAsACTMDBxcrMyI1NTQzMzIVFRQjLAMDmgUFA0YDA0YDAAABACj/xACpAM0AEQAXQBQPCAcDAEoBAQAAdgEAABEBEQIHFisXIiY1NDY2NxcOAhUUFhcGBmYcIhsvHBsQGQ8eGgQmPCkZGUZKHhcRMzISDRIIHSYAAAAAAgAs/74ArQG4ABEAHQAcQBkPCAcDAEoaFQIASQEBAAB2AQAAEQERAgcWKzciJjU0NjY3Fw4CFRQWFwYGByYmJzY2NzY3FhUUahwiGy8cGxAZDh4ZBCUHGCgPBQwECg1NsCkZGUVKHhcRMjISDRIIHiXyERoJCxMKEhQtDgsAAAACADf/uQF1AjIAIwAyACBAHSwnHBYBBQFJAAABAQBZAAAAAWEAAQABUScvAgcYKzcnNjY1NCcmJicmNTQ2NzYzMhcWFRQHJiYjIgYHHgIXFhUUByYmJzY2NzY3FhYXFhUU7yYEAxAHHhdNNSssMEkrDh4eLiMuThEWODcTHjgZIhQGCwQKDRQeCRFiGBQbBxQPCBQNKUYvXxwdLhARHRocFDwyDRodFB8nK+wSFwsLEwoSFAwSCA8GCwAAAAEAIv/kBFQA5AAkADKxBmREQCcfDwUEBAFKAgEBAAABWQIBAQEAXwMBAAEATwEAHRsJBgAkASMEBxYrsQYARAUgJyY1NxYWMyE2Njc2NjcWFhUUBwYGBwYVFBYzMjY3BgYHBiEB3f7PSz8BRuymAXQBBwoPQj0DAhcZKwsZKhQdQSMEDgs//l8cDAg8CggJFyUQGzEfBw4HIxAPGwgTCg0JCAgWJhEP//8AIv/kBFQA5AIGBT0AAAABABMAOwF+AZQADgAYQBUGAQBKDgwLAgQASQEBAAB2EyMCBxgrNyc3JzcXNzMXNxcHFwcnWgMvcwKKJgUniwJ1LwNuOwKDTgQEhoYEA0+DAlYAAAABACD/GgZMACsAVQDdsQZkREAQBQEEAi8GAgoELSwCAQgDTEuwHlBYQCwGAQIACgMCCmkABAAIAQQIaQABBwABWQUBAwkBBwADB2kAAQEAYQsBAAEAURtLsCdQWEAxBgECAAoFAgppAAUDBwVZAAQACAEECGkAAQcAAVkAAwkBBwADB2kAAQEAYQsBAAEAURtAMgYBAgAKBQIKaQAEAAgBBAhpAAEHAAFZAAUABwkFB2kAAwAJAAMJaQABAQBhCwEAAQBRWVlAHQEAUE5JR0NBPTsoJiEfHRsYFhMRDAoAVQFVDAcWK7EGAEQXIiY1NDcXBhUUFjMyNzY2NzYzMhYXFjMyPgIzMhcWMzI+Ajc2MzIWFRQHJzY1NCMiBgcOAwcGBiMiJicmJiMiBw4CIyInJiYnJiMiBwYGBwZtJSgjIBcQEAwPCBgRNjcQGxAgKQwUEhEIBA2T8jd3kbh5GBISFlEbPAMFEQ0gYHN6OzpqMW24UwYJAgcMChcXCyUnDBEFCQoqMg0XCRHmQTo0RhExHBctEgknHl4bIDwYIBgFQwkXKB8HGQ89UBg8FgQEBAobGxkJCAkfIgMCDgsdFjYSFwQJWhkjChQAAAD//wAg/xoGTAArAgYFQAAA//8AIP8aBkwAKwIGBUAAAP//ACD/GgZMACsCBgVAAAD//wAg/xoGTAArAgYFQAAAAAIAM//iBLMA1gAfAC0AcLEGZES1FwECAwFMS7AnUFhAHQABAAMCAQNpBgQCAgAAAlkGBAICAgBfBQEAAgBPG0AhBgEEAgACBHIAAQADAgEDaQACBAACVwACAgBfBQEAAgBPWUAVICABACAtIC0lIxYOCQcAHwEcBwcWK7EGAEQXIiYmNTQ2NjMyFhUUBgc6AzYzNjY3FAYHDgMjJTY1NCMiBgYVFBYXFhbTKEkvJT0jJzkNCUmjnIFVCFSjVBggE1NlXiD9rQk7ECkdJRkRJx4TMC0iPCY7KREkDgEBCwsfKwgFCAQCTRQNQwoWExMUBQIDAP//ADP/4gSzANYCBgVFAAAAAgAY/z4EGgBXADgARQBVsQZkREBKRDwlIyARBgECOAcEAwABAkwBAQBJAAIGAQYCAYAAAwAGAgMGaQcFAgEAAAFZBwUCAQEAYQQBAAEAUTo5QD45RTpFKSwVODkIBxsrsQYARAUnNjY3JiYnBgYjISImJyYmNR4CMyEyNzY2NzYzMhYXBgYHFhc2Njc2MzIWFxYVFAYHBgYjIiYnNzI2NyYmIyIHBgYHFgLnHgUMBgoVCh9LOP7zLlIxIh00TEYqASVOIBERCQoVBg0FCA0HERUaMBYvIB0tDQ0jGho8GRwzGnAePyIUKxcWGg0eESPCEgsUCgUJBRkVBwoHJyIIDAUcDhoOEgQFEx4LCQYmPBUqHhgYHCc3EhIQCQg7ExofGRkMIxcGAP//ABj/PgQaAFcCBgVHAAAAAQAp/2kA5ACbAAwABrMJAAEyKxcmJjU0Njc2NjcXBgZFDQ82JBQoExI9TJcEGA0gXy8aLhM8PHYAAAAAAgBEAAADRQE3AB8ALwA1QDIYFwICAwFMAAEAAwIBA2kGBAICAgBhBQEAAA8ATiAgAQAgLyAvKCYSEAsJAB8BHwcHFishIicmNTQ2NzY2MzIWFRQGBgc2Njc2NjcXBgYHBgcGBic2NTQmJyYjIgYHBhUUFxYBVps7PC4kFzYcO0UFBQJ+tzkaJAwaChULMGo1lmgIHBgaGyA9FBQ8Oh8gTCtKFw8RV0cLIx4FAhoZCxoPFRQgCzAXCwtIIg0jMw4OFhESECkXFgACACb+1wI5AhQAQABOAEpARxYBAQAiHxwXCQUCAUw+MAQEBAUBAQMEBEwAAAABAgABaQYBBAADBANlAAICBWEABQUPBU5CQUtJQU5CTjw6NDIbGRQSBwcWKxMnNjY3JjU0NjcmJicmNTQ2NzYzMhYXByYmIyIHFBYXNjY3FhUUBgcGBgcGBwYVFBc2NjMyFhUUBgcGIyImJwYGNzI2NzY1NCYmIyIHFhZDHRgtFhEyKidAERAuJSQpIUArEigjFjtFXE8hUjELEQ4HFhGCPjYNRoBCRko/Njc7P18dEir/Ml0WFx0vGnh8HlP+1xYmRR8wQD5yKgkvHx0dK10dHSImHw8IJD9PBxUkEBAPEBkJBQsHM0lBSCkeWFRSTTdbGho0LxY5OR8bGiIVIxWDHyH//wA2/t0BowEVAiYAAQAAAQYFNFLGAAmxAQK4/8awNSsAAAAACABH/vEEpQQgAEIATgBiAHYAsADEANkA5QDGQMNWAQQHkwEGDDIWAgEGOA0CCwCuAQoQyAERFQZMJAEFSgAFBwWFCQEHBAeFFgEEDASFDQEMBgyFAgEBBgAGAQCAAwEACwYAC34OAQsQBgsQfg8ZAgoQFRAKFYAAFREQFRF+EhoCERQQERR+GwEUFIQYCBcDBgYQYRMBEBAPEE7b2rGxeHdkY1BPREPh39rl2+XT0cbFscSxxLu5rKqlpJeVkY+Af3eweLBtbGN2ZHZbWk9iUGJKSENORE46OTEwGyocBxgrASYmJyYmJyYmJzYzMhcmJicmJjU0NjcGIyInJic2NzY2NzY2NxYWFxYXFhcGBgcGIyInFhYVFAYHNjMyFhcGBgcGBgMyNjU0JiMiBhUUFgUyNzY2NzY1NCcmJiMGBgcGFRQWBTI3NjU0JyYmJyIGFRQWFxYWFxYDMjc2NzY2NxYzNjY3NjUmJicGBiMmJicmIyIGByYmIyIHBgYHIiYnBgYVFBYXMjY3FhcWMzI2NxYWFzY2NzY1NCcmIyIHBgYHBgYVFBYFMjY3NjU0JicmJicmIyIGFRQXFhYXMjY1NCYjIgYVFBYCdhE0IEx9MzRSIAYMCgkQHQoKDCsiCQoHAwYCQW42e0IfMxMRNiWHb2xEAgUCAwcJDCQhISQLCQYKBELSkiI5ERkfHxkXHR3+2g0KI1koFwEDFw00XCgPEwKJDAkIDipcNA4aDgoqViQL4AsLRzEaLRULEyM1ESMBREcKDwUpcToQDxslDQohGg8RQHApBBMJRkZGRgkSBU2MCwgaKBARLVg1XCgPBwoMCwwjWCkKDxr+5Q0XBAEOCidbIgsMDBIOKV3IGR8gGBcdHf7xFzIEDjcnJmM6CAghRCMjSSdLikIEAQIBd1EoNQkENhkZMggdTk12AQEBAQRCiUxOikMIBQN0mSIHLwRQHhkXHR0XGR5NCxcrDAoWBAMOCRIvHQcUDBIDCwoLEggdLRMJEQsVAwwrGAv8vwQXGw8kGAYgSypUXWOjRgUDKUILAw8MCxADDT8qAwVGo2Ngpz8DA1ojAxYNDhZXEi8cCxAKCQwMFiwLBA8LEg4CCg4DBAoPBAwuFgsVDRAIHi03HhcXHx4YFx4AAAAIAEf+8QSlBCAAQgBOAGIAdgCwAMQA2QDlAcVLsAlQWEAgVgEEB5MBBgwyFgIBBjgNAgsArgEKEMgBERUGTCQBBUobS7AKUFhAIFYBBAmTAQYMMhYCAQg4DQILAK4BChDIAREVBkwkAQVKG0AgVgEEB5MBBgwyFgIBBjgNAgsArgEKEMgBERUGTCQBBUpZWUuwCVBYQE8ABQcFhQkBBwQHhRYBBAwEhQ0BDAYMhRgIFwMGAQaFAgEBAAGFAwEACwCFDgELEAuFEwEQChCFDxkCChUKhQAVERWFEhoCERQRhRsBFBR2G0uwClBYQFcABQcFhQAHCQeFAAkECYUWAQQMBIUNAQwGDIUXAQYIBoUYAQgBCIUCAQEAAYUDAQALAIUOAQsQC4UTARAKEIUPGQIKFQqFABURFYUSGgIRFBGFGwEUFHYbQE8ABQcFhQkBBwQHhRYBBAwEhQ0BDAYMhRgIFwMGAQaFAgEBAAGFAwEACwCFDgELEAuFEwEQChCFDxkCChUKhQAVERWFEhoCERQRhRsBFBR2WVlARdvasbF4d2RjUE9EQ+Hf2uXb5dPRxsWxxLHEu7msqqWkl5WRj4B/d7B4sG1sY3ZkdltaT2JQYkpIQ05ETjo5MTAbKhwGGCsBJiYnJiYnJiYnNjMyFyYmJyYmNTQ2NwYjIicmJzY3NjY3NjY3FhYXFhcWFwYGBwYjIicWFhUUBgc2MzIWFwYGBwYGAzI2NTQmIyIGFRQWBTI3NjY3NjU0JyYmIwYGBwYVFBYFMjc2NTQnJiYnIgYVFBYXFhYXFgMyNzY3NjY3FjM2Njc2NSYmJwYGIyYmJyYjIgYHJiYjIgcGBgciJicGBhUUFhcyNjcWFxYzMjY3FhYXNjY3NjU0JyYjIgcGBgcGBhUUFgUyNjc2NTQmJyYmJyYjIgYVFBcWFhcyNjU0JiMiBhUUFgJ2ETQgTH0zNFIgBgwKCRAdCgoMKyIJCgcDBgJBbjZ7Qh8zExE2JYdvbEQCBQIDBwkMJCEhJAsJBgoEQtKSIjkRGR8fGRcdHf7aDQojWSgXAQMXDTRcKA8TAokMCQgOKlw0DhoOCipWJAvgCwtHMRotFQsTIzURIwFERwoPBSlxOhAPGyUNCiEaDxFAcCkEEwlGRkZGCRIFTYwLCBooEBEtWDVcKA8HCgwLDCNYKQoPGv7lDRcEAQ4KJ1siCwwMEg4pXcgZHyAYFx0d/vEXMgQONycmYzoICCFEIyNJJ0uKQgQBAgF3USg1CQQ2GRkyCB1OTXYBAQEBBEKJTE6KQwgFA3SZIgcvBFAeGRcdHRcZHk0LFysMChYEAw4JEi8dBxQMEgMLCgsSCB0tEwkRCxUDDCsYC/y/BBcbDyQYBiBLKlRdY6NGBQMpQgsDDwwLEAMNPyoDBUajY2CnPwMDWiMDFg0OFlcSLxwLEAoJDAwWLAsEDwsSDgIKDgMECg8EDC4WCxUNEAgeLTceFxcfHhgXHgAAAP//AEf+8QSlBCACBgVNAAD//wBH/vEEpQQgAgYFTQAA//8AR/7xBKUEIAIGBU0AAAAFAAj+4gHKAtUACgAmADAAQwBOACZAIzouKykYEg0JBgMKAEpMSUZEOQUASQEBAAB2DAsLJgwmAgcWKwEmJic2NjcWFhcGAyInNjY3Njc2Njc2NzcXBwYGBwYGBwYGBwcGBgcmJzY3FhYXBgYBJyYnJicmJic3FhcWFhcWFhcXByYnNjY3FhYXBgYBrAsQBQUPDAcQBw/oIRUYIQ4VFwEDBAlWNQwYHhgFBg0GARAPChgtjSsdJx8PIRMMIQFINVYJCAEBBAIfHA0DCwcGGhsYDBULBQ8MBg8JBw8ClgoPBQQQDQgOCBP+FS4GCgUJDWJvH1kdESEJCyotMmUyByAaEhIWGyIgJSQRIBEQJf58Eh1ZM24dLA4TMToQTT8uKgkJbhAOBBANBw8HChEAAAAABQAq/uIB7ALVAAkAHQAnAEMATQAfQBwkHRQHBQIGAEpLSUY+OCYhBwBJAAAAdjc1AQcWKxMmJzY2NxYXBgYTJicmJicmJicnNxcWFxYWFxYWFxcmJic2NjcWFwYBJzc2Njc2Njc2Nzc2NjMyFwYGBwYGBxQGBwYHByYnNjY3FhcGBkgMEggPBxULBBB6HQwECQgHGRsZDDVUDAMEAQIDA7kPIRIMIBQqHif+gwwZHhcGBgsIAR8KGC0KIRQYIQ4KFQ0EAwpWNQ0RCA8HFwkEEAKWDw8KEAcQDgQQ/lQvPA9OPi8qCQgiEhxaGVA3HiwOehEgEBElEyAhJf57IgkLKiwzZTILNhIRFy8FCQcECwdhbx9ZHV4RDQoQBxILBBEAAf/u/9MAEgKnAAMAF0AUAAABAIUCAQEBdgAAAAMAAxEDBhcrBxEzERIkLQLU/SwAAf+V/9MAawKnAA4AIUAeDQwLCgkIBwYFBAMCAQ0ASgEBAAB2AAAADgAOAgYWKwcRByc3JzcXNxcHFwcnERJBGFRUGFNTGFRUGEEtAkJCGFJSGFNTGFJSGEL9vgAB/+7/0wDSAqcACgA0QDEFAQEABwYCAgECTAQDAgBKAwECAQKGAAABAQBXAAAAAV8AAQABTwAAAAoAChYRBAYYKwcRMyc3FwcnNyMREqVDGGpqGEOCLQJ9PxhpaBg//aYAAAAAAf8u/9MAEQKnAAoANEAxBQEAAQQDAgIAAkwHBgIBSgMBAgAChgABAAABVwABAQBfAAABAE8AAAAKAAoWEQQGGCsHESMXByc3FwczERKBQhdqahdCpC0CWj8YaGkYP/2DAAAAAAIARP76A0QBVAA8AEoAS0BIPzAtEAQEBTcdEQYEAgQbCQIBAgNMAAMABQQDBWkAAQYBAAEAZQcBBAQCYQACAg8CTj49AQBEQj1KPkooJiAeFxUAPAE8CAcWKwEiJjU0NjcGBgcmNTQ3NjY3FwYVFBYzMjY3NjcmJwYjIiY1NDY3NjMyFhcWFhc2NjcWFRQHBgYHFRQHBgYTMjcmJyYjIgYHBhUUFgGAVWADBCFCIgkXFEg0IRRJQjRvLCwZAQVLQzlEKSEhIiNBGRAZBxgzHB0PEDkiYS9/YC84HjMZFBUoDQ05/vpiWhkrERIkEg8QHA8OKBkQUBRJUiAdHCcrJBxHOjNjHx4tKRg9Iw4iFA8ZEhITJBENf1UpMgFYFmElESAXFhUlJgD////4/scBcQKVAiYBZQAAACcFI//yAMgBBgYtAOIAEbECAbDIsDUrsQMBuP/isDUrAAAAAAMAZv8jA6cC6wB3AIIAjQB2QHM7AQYESjQCBQZwbVoxIhsaEgcJAwVRTQgDBwOMiYaBfnsRBwEIBUw4NS4tKgUESgAFBgMGBQOAAAQABgUEBmkAAwACCAMCaQABCQEAAQBlAAcHCGEACAgPCE4BAGdlUE5IRkRDPz0gHhYUDgwAdwF3CgcWKwUiJjU0NjY3FwYVFBYzMjY2NycGBiMiJjU0NxcGFRQzMjY3LgI1NDc2NxYWFwcWFhc2NjcnNjY3FhYXNjYzMhYVFAYjJiYjIgYHFhYXFjMyNyYmJyYmNTQ2NxcWFhcWFhUUBwYGIyImJicmJicGBgcWFhUUBgcGJSYmJzY2NxYWFwYHJiYnNjY3FhYXBgEbVWAFCwolEklCMmVPEB0oWCw3PiokH0IjRjgPGA4dDxUIEgoYBw8IFy8XDw4iEwIDAh06ISIvEgoJIBkbMh0CAwIlMDMjBRMNCQojHwsICwMCAiIiLh0bMSEDBQwIFSwXCQhPREMCCBUhDA4gEw4fEBO/FiELDCEUDR8QFN1iWhotNSYJQiZJUiE5I7U/RVlWVGELTjl6PE1donMWLBQLBCI5HBU3gEsgORn0Dh0NH3lcFxQqIhEiFhAQFVWeSxocFDgiGiIHFS0SOSY1Dw4bDSNDHBIYKBctjV4YPydJbiM9cSMjJxEfDQ0hFRAeDh80Eh8NCyIWEB4PHgAAAAAQADD/qgIiAZwABwAPABcAHwAnAC8ANwA/AEcATwBXAF8AZwBvAHcAfwD0QPEAASABAAcBAGkFAQMiBCEDAgYDAmkJAQckCCMDBgsHBmkNAQsmDCUDCg8LCmkRAQ8oECcDDhMPDmkVARMqFCkDEhcTEmkZARcsGCsDFh8XFmkAHxoeH1kdARsuHC0DGh4bGmkAHx8eYS8BHh8eUXl4cXBpaGFgWVhRUElIQUA5ODEwKSghIBkYERAJCAEAfXt4f3l/dXNwd3F3bWtob2lvZWNgZ2FnXVtYX1lfVVNQV1FXTUtIT0lPRUNAR0FHPTs4Pzk/NTMwNzE3LSsoLykvJSMgJyEnHRsYHxkfFRMQFxEXDQsIDwkPBQMABwEHMAYWKwEiNTQzMhUUByI1NDMyFRQzIjU0MzIVFAUiNTQzMhUUISI1NDMyFRQFIjU0MzIVFCEiNTQzMhUUBSI1NDMyFRQhIjU0MzIVFAUiNTQzMhUUISI1NDMyFRQFIjU0MzIVFCEiNTQzMhUUBSI1NDMyFRQzIjU0MzIVFAciNTQzMhUUASkaGhpvGhobkBwcGf70GhobASAbGxn+fRwbGgGAGhsa/jobGxoBpBsbGf46GxsaAYAaGhv+exoaGwEgGxsZ/vQaGhuQHBwZbxoaGgFnGxoaGxEaGxsaGhsbGjAZHBsaGhsbGkkbGhobGxoaG1QaGhoaGhoaGlUaGxsaGhsbGkkaGxsaGhsbGi8ZGxkbGRsbGRIbGhobAAAABgAV/8MB6gJ7AAgAEQAYAB8AJgAtALxAKhELBwEEAwQBTC0sKyopJSQjIiEfHh0cGxcWFRQTEA8ODQwGBQQDAh4FSkuwF1BYQCAABQQFhQgBBAMCBHAAAQYBAAEAZAADAxBNBwECAg8CThtLsB1QWEAfAAUEBYUIAQQDBIUAAQYBAAEAZAADAxBNBwECAg8CThtAHwAFBAWFCAEEAwSFAAMCA4UAAQYBAAEAZAcBAgIPAk5ZWUAbICASEgAAKCcgJiAmGhkSGBIYCgkACAAICQcWKxc3NSc3FwcVFyUhJxE3JwcXERcRJzcXBxEnMxE3JwcXExEnNxcHESczETcnBxcVX1/q62Bg/mABakJYy8pXEz6enj6smDSAgDQWJVtbJVxLIEVEHj1s9Wvs7Gv1bBdNAQZizMxi/votAU9Gn59G/rEXAUU5gIA5/twBNShdXSj+yxIBKSJHRyL//wBE/s8BwgFdAiYB4gAAAQcFNACKABYACLEBArAWsDUrAAAAAwAi/8QBbwH6ABAAHgAuAAq3KB8ZEQcAAzIrFyYnPgM3FhYVFAcOAxMmJzY2NzY2NxYXFhUUEyYmJzY2NzY2NxYWFxYVFFgXDBQ+SksgDBADIElGPBg5JwgMBwYOBzQTFYkZMBcHDQYHDgcaIwoVPBIfKX6TkDsIHg4JCDyNjHYBUCoVDRgMCxgMHhETCA7+cRMgDQ0YCwwYDA8XChIIDAAAAP////b/+wIQAqsCJgYSAAABBwYBAKv/1AAJsQEBuP/UsDUrAP////b/+wIQArwCJgYSAAAABgYTAAAAAAAEACL/xAIfAfoADgAcACwAPAANQAo2LSYdFw8HAAQyKxcmJz4DNxYWFRQHBgYDJic2Njc2NjcWFxYVFBMmJic2Njc2NjcWFhcWFRQXJiYnNjY3NjY3FhYXFhUUWBcMG0JHRR4MEANPfwU5JwgMBwYOBzQTFYkZMBcHDQYHDgcaIwoVfhkwGAgNBgYPBxkkChU8Eh85hYyENwgeDgkIlvkBFCoVDRgMCxgMHhETCAz+bxMgDQ0YDAsYDA8XChIIDEoTIA0NGAwLGAwPFwoSCAwABQAi/8QCyQH6AA4AHAAsADwATAAPQAxGPTYtJh0XDwcABTIrFyYnPgM3FhYVFAcGBgMmJzY2NzY2NxYXFhUUEyYmJzY2NzY2NxYWFxYVFBcmJic2Njc2NjcWFhcWFRQXJiYnNjY3NjY3FhYXFhUUWBcMG0JHRR4MEANPfwU5JwgMBwYOBzQTFYkZMBcHDQYHDgcaIwoVfhkwGAgNBgYPBxkkChV4GTAYCA0GBg8HGSQKFTwSHzmFjIQ3CB4OCQiW+QEUKhUNGAwLGAweERMIDP5vEyANDRgMCxgMDxcKEggMShMgDQ0YDAsYDA8XChIIDEoTIA0NGAwLGAwPFwoSCAwAAP//AC3/AhfMAdkALwY/EZQAACZmAC8Eug6CAAAmZgAvBkAIwQAAJmYALgZBLQAmZgAvBZAXZf9xJmYALwWQDtf/0yZmAC8FkAkM/0MmZgAvBZAR8f/TJmYALwWQAIb/0yZmAC8FsQHR/70mZgAvBbELcf//JmYALwWxFo3/vCZmAC8Fdwnq/+cmZgAvBYoPpgBEJmYALwWKDL4ARCZmAC8FigbiAEYmZgEPBZAFfv/TJmYAcrEbAbj/cbA1K7EcAbj/07A1K7EdAbj/Q7A1K7EeAbj/07A1K7EfAbj/07A1K7EgAbj/vbA1K7EhAbj//7A1K7EiAbj/vLA1K7EjAbj/57A1K7EkArBEsDUrsSYCsESwNSuxKAKwRrA1K7EqAbj/07A1KwAA//8ANv+CB2ABwgIGBWgAAAABADb/ggdgAcIAbwC5sQZkREuwG1BYQBo5AQMFPAEJA1YBBwQcAQACBExIAQVKEAEASRtAGjkBAwU8AQkDVgEHBBwBAQIETEgBBUoQAQBJWUuwG1BYQCgABQMFhQADAAkEAwlpBgEECAEHAgQHaQACAAACVwACAgBhAQEAAgBRG0AsAAUDBYUAAAEAhgADAAkEAwlpBgEECAEHAgQHaQACAQECVwACAgFfAAECAU9ZQBBkYl5cUVAWKSUrSGEVCgcdK7EGAEQFJiYnJiYnJiYjISIOAwc0Njc+AjMhMhYWFyYmNTQ2Njc2MzIWFx4CMzI2Nz4DNzY2MzIXBgYHFhYzJiYnJiY1NDY3HgIVFAYGByImJyYmJwYGBw4CIyImJyYmIyIGBhUUFhYXFhUUBgURBxwUQYA9PW8x/tM6U0VDUDccIjaJk0UBLSpzejILDRwuGg4MEiERKTY5KxkkEgMKDAsCChYRDQsNFAooUTcJFAkJCywWBRMPDREIKD0ZGSkVAQUCBSc0Gi1YIR0rExMgFBw2JwsJfgMPCgsPBAUFAQQGDAgiJwcMDQQLEQkULh0iS0APCQ4KGBwMBAYFHyYhBhwZCBtDIBMVIDgWFyELGywNIk9KHQwwMA8GBQYQCgQIBQkLBhUUERgjMBInLBsQBRQOJAAAAAAC/70B8gDAAsQAFwAvAFuxBmREQFAfGgIFAgoEAgEFAwEAAQNMEQEDSgADAAIFAwJpAAQABQEEBWcAAQAAAVkAAQEAYQgGBwMAAQBRGBgBABgvGC8rKiknJSMeHAgGABcBFwkHFiuxBgBEEyImJzcWFjMyNjcmJicmNTQ3FxYWFRQGJzQ3JyYjIgcmNTQ2MzIXFjM3BwcjBgYHawkmCwcPFQYdJQkIDQULHAsJCDDEThQLDhkOCRsUFyQjGR0OGAE1RwEB8hMJEAQEIhsRGAkTDhIMJBghCC4/BFMlBQQOBgwPEgwNAiUBCTslAAT/yQHJANwCvwAJACEAOQBBAG2xBmREQGIoAQYCQDwrJhIFBQY2AQMFDgEBAw0BAAEFTBoEAgEEAkoAAgAGBQIGaQkBBQADAQUDaQABAAABWQABAQBhCAQHAwABAFE7OiIiCwo/PTpBO0EiOSI5NDIuLBEPCiELIQoHFiuxBgBEESc3FhcGBgcGBhciJic3FjMyNyYmJyY1NDY3FhYXFhUUBgc2NTQnNjcWFhc2MzIWFRQGIyImJxQHBjcyNyYjIgcWHxcMEwIFAwMGgwokCwUgCzgTBQ0ICg4NAgUEETDSAxQJDAUMBTEuFCE6LAoVCQYLOj0VCiIqIA0ChxsdCBQDBwUEB7wTCQ8GPAsaDRMOCw4GCRMJLRQtQAQhECExHQ0PGQpSIhIpMAICJQsTYx8hOwX//wAJ/5MAv//7AAYFkBQAAAL/9gHKALUClwAXACAAQ7EGZERAOAcBAwAeGgkFBAIDFAEBAgNMAAAAAwIAA2kEAQIBAQJZBAECAgFhAAECAVEZGB0bGCAZICQqBQcYK7EGAEQTNjY1NCc2NxYXNjMyFhUUBiMiJicUBwY3MjcmIyIHFhYHAgIVDAoNCDIuFCA5LAsUCgYKOj4TCiEqIQcSAcoLGQ0kLh8LIhBSIhIoMgMCIg4SYh4iOwMCAAAB//QBygC+ApUAKwBJsQZkREA+FwcCAgEhHxwYBQUAAioBAwADTAUBBAMEhgABAAIAAQJpAAADAwBZAAAAA2EAAwADUQAAACsAKywkJSwGBxorsQYARBM2NjU0JzY3FhYXFhYzMycmNTQ2MzIWFwcmIyIHFBcXNjcWFRQGBwYjIicUBQICFQgNCw4EBxAIAgUIJRYLFhMJFREUEyYLIhUDIh4bHBMOAcoOGQokLhcSGR4FCgUGCw8aMAsRDQgLJgwDCgoGBxMUBgcDPgAB/34B/wDXAl4AGgBzsQZkREuwG1BYQA0KAgIAAgFMGRgSAwJKG0ANCgICAQIBTBkYEgMCSllLsBtQWEAUAwECAAACWQMBAgIAYQEEAgACAFEbQBUDAQIAAQACAWkDAQICAGEEAQACAFFZQA8BABcVDQwHBAAaARoFBxYrsQYARBMiJwYGIyImJyY1NRcyNjc2NxcVFBYzMjcXBokeEhU6JBknDhonPkkTExcOFAoYHgwfAf8kEQsCAwYTBgEHCQccBQkNFzMKVQAAAAAC/7ABtQCTAoAAGAAgAEqxBmREQD8cEQcGBQQGAgMBTA4NDAMBSgABAAMCAQNpBQECAAACWQUBAgIAYQQBAAIAURoZAQAfHRkgGiAUEgAYARgGBxYrsQYARBMiJiYnNxc3JyY1NDcXBxYWFzYzMhYVFAYnMjY3JiMiBycmKRkPBDgSCwkXDQoDAwEsIxInOS8ZKBcSHyUuAbUHDQoMCxU8MBAWBSkIGSwSMB8PICUbCw4iOgAC/z4B6QCvAyAACgA1AGaxBmREQBArJR4bGRYKBwYFAwIADQFKS7AxUFhAGAACAQABAnIAAQIAAVkAAQEAYQMBAAEAURtAGQACAQABAgCAAAECAAFZAAEBAGEDAQABAFFZQA0NCzIxKSYLNQ0zBAcWK7EGAEQTNScnNDcXBxcVBwciJicmNTQ2NzY2NyYmJzY3FxYVFAYHBwYGBxYzMjY2NxYVFAcGBzYOAhUDBxgOCgIDQiI6GTEHAxhGPgIJBgwMBgMIBUYVKhImcBpNSBQGDApCFwosOgI5I0dhFwUtCEUVVVMDBAgIAxgEFR4KI1c0DgdkKyQHGQkSBg8LCwUHBAYHCQYGBgEBAwIAAv+4AaIAPgK5AAcAHgA2sQZkREArCwEAAQFMGBEMBgMCBgFKAAEAAAFZAAEBAGECAQABAFEJCBAOCB4JHgMHFiuxBgBEEyYnNxYWFwYHIiYnNxYWMzI3JyYmNTQ2NxcWFhUUBg8QCxsHDQUKNAkeCwUKEgg6FQsKEAwLEQYFLgKFDgocBw0ED/ANCA0CAkIYESAMCgsFMRAZCS0/AAAAAAH/yQHOAEkCIgAKAAazBAABMisDNzY2NxcUBwYGBzcFJTwZARcXJxEBzioLFQoKGgoKDgUAAAAAAv/KAcAARgJbABMAHQB6sQZkREAKBAEDAgIBAAMCTEuwEFBYQCMAAgQDBAIDgAADAAQDcAUBAACEAAEEBAFZAAEBBGEABAEEURtAJAACBAMEAgOAAAMABAMAfgUBAACEAAEEBAFZAAEBBGEABAEEUVlAEQEAGRgSEQ8OCggAEwETBgcWK7EGAEQDIjU2NyY1NDYzMhYVFAcXBgYjBjc3NjU0IyIGFRQZHSUSISQTDBQGFQESEyEaAgMTBwwBwBkJDhIcFyYYHREPARAMKUoECQUfCgUWAAAAAf/J/00ASf+iAAoABrMEAAEyKwc3NjY3FxQHBgYHNwUlPBkBFxcnEbMrCxUKCxoJCg4G//8AUv7yAVP/twEGBeb/5wAJsQADuP/nsDUrAAAB/9b/bwCZABoALADisQZkREuwEVBYQA8hHBcWBAQDIgsBAwUAAkwbS7AbUFhAEiEcFxYEBAMiCwIGAAEBBQYDTBtAFSEcFxYEBAMLAQEAIgEGAQEBBQYETFlZS7ARUFhAIwAEAwADBACAAAIAAwQCA2kBAQAFBQBZAQEAAAVhBgEFAAVRG0uwG1BYQCQABAMAAwQAgAACAAMEAgNpAQEAAAYFAAZpAQEAAAVhAAUABVEbQCgABAMAAwQAgAACAAMEAgNpAAEGBQFZAAAABgUABmkAAQEFYQAFAQVRWVlACiUUJCUnJBMHBx0rsQYARAcnNjYzMhYXFjMyNyY1NDY3NjMyFRQHJzQjIgYHFhYzMjcHBgYjIiYnJiYjIhwOCh0KBQgHDAwJByUWEhEOJAYRDAsaCAsgDRIWAhw5EAUMBwYLBQ6RCxEZAwIEAxMdESULCyMMDwUQFA8KDwklEBUFAwMEAAH/8QGjACUCjgANAAazCQABMisTJzQmJyY1NDY3FhYVFAsGCQgDFRAIBwGjAixNIgwKFB4GJUAxMAAAAAH/8f8EACX/7wAMAAazCAABMisXJzQnJjU0NjcWFhUUCwYRAxUQCAf8AlRGDgkUHQcmQDAwAAD////5AfYAeQKqAgYF2QAAAAP/8gHyAHIDCgALABYANACGQBQhAQUEMCwlJAQGBQJMMRkYFwQGSUuwFlBYQCUAAQMCAXAAAwIDhQAGBQaGCAECBwEABAIAagAFBQRhAAQEDgVOG0AkAAEDAYUAAwIDhQAGBQaGCAECBwEABAIAagAFBQRhAAQEDgVOWUAZDQwBAC8tKiggHhIRDBYNFgcFAAsBCwkHFisTIiY1NDYzMhYVFAYnMjY1NCYjIgYVFAcnNyY1NDY2MzIVFAYHJzc2JiMiBgcWMzI3Bw4CFhARFA0OFhcMCgsOCwUMDgInJBYiESgEAxMBAQkFEBkGDh8WHgIeKCACuRgOEhkYExAWGAsDBQ8ICBLfJxYIIBQkFyUGDQgFBAUIFg8UBicHDg8AAAEAF/9rAJcAHgAbADaxBmREQCsXEw0MBAIBAUwYAgEABAJJAAIBAoYAAAEBAFkAAAABYQABAAFRIyYnAwcZK7EGAEQXJzcmNTQ2NjMyFRQHJzc0IyIGBxYzMjcHDgIZAickFyIRJwcTAQwQGgUPHhgbAh0mIpUnFgggFCMXJA4OBgQMFg8UBycHDRAAAv/rAcEAqwJuAAkAEwAItQ4KBAACMisDNzY2NxcUBwYGBzc2NjcXFAcGBhUDLFgsARYfTyYEJFk0ARcdUAIJKA0fEQsbCg0aVisLHhQKGQsNHQAAA//iAfIAcgMZAAkAEwAxADdANB4BAQAtKSIhBAIBAkwPDgUEBABKLhYVFAQCSQACAQKGAAEBAGEAAAAOAU4sKiclHRsDBxYrAzc2NjcXFAcGBgc3NjY3FxQGBwYHJzcmNTQ2NjMyFRQGByc3NCYjIgYHFjMyNwcOAh4DIDobAQ8XNRcDFTsmAREUHC4CJyQWIhEoBAMTAQgFEBkGDh8WHgIeKCAC1hoKFQoIEAgJEjgdBhQPCQwLCAzGJxYIIBQkFyUGDQgFBAUIFg8UBicHDg8AAAAD/9YByQCSAsgACgAXADwAPUA6OgEBAUszLycmHh0PBAEACgJKAAEDAAFZAAIAAwACA2kAAQEAYQQBAAEAURkYOTcqKSQiGDwZPAUHFisDNzY2NxUUBgcGBgc3NjY3FxQHBgYHBgYXIiY1NDcXBgYVFDMyNSc3FhYzMjY3JicnNDY3FhUUBiMiJwYGKgIXVz4SIxU9IAMXVj4BBwQbGQszBxEYCw4BAxYYAhUDDxMGCwEDDAEPCRAcFBULBxcCaSEHHhkGEA4PCRVCIQYfGQcLCQQOCgQUfBYVEhgDBwwFGBsfDRkbBwQPFQcJEgIfHRchDhERAAAA//8AAgHLAMICeAEGBXwXCgAIsQACsAqwNSsAAAAC/8kBtACgAl4AJQAxANCxBmRES7AbUFhAERwBAwERAQADDwkIBAQFAANMG0AUHAEDAREBAAMJAQQADwgEAwUEBExZS7AVUFhAJgcBBQAABXEAAgAGAQIGaQABAwABWQADAAADWQADAwBhBAEAAwBRG0uwG1BYQCUHAQUABYYAAgAGAQIGaQABAwABWQADAAADWQADAwBhBAEAAwBRG0AmBwEFBAWGAAIABgECBmkAAwAEA1kAAQAABAEAaQADAwRhAAQDBFFZWUAQAAAuLAAlACUTFikkFQgHGyuxBgBEEyImNTc0IyIHJzYzMhUUBzY3JjU0NjMyFxYVFAcWFjMGBiMnBgY3NzY0NTQmIyIGFRQRAxYECgcaCB4VGgQkGCUmFAwMCgQIEAcBExQOEzhJAgIKCwkNAbQRAhoPDQonJA4MCRIUHBcmDw4jEgcBARAMARMhVgUEBgMOFQsFFgAAA//MAfIAcgMeACMALgBMAQhLsCdQWEAgCQgEAwIAOQEJCEhEPTwECgkDTBINAgABS0kxMC8ECkkbQCMJAQUACAQCAgU5AQkISEQ9PAQKCQRMEg0CAAFLSTEwLwQKSVlLsBVQWEArAAoJCoYAAwAHAQMHaQQBAQUBAAIBAGkAAgsBBggCBmkACQkIYQAICA4JThtLsCdQWEAwAAoJCoYAAwAHAQMHaQABBAABWQAEBQEAAgQAaQACCwEGCAIGaQAJCQhhAAgIDglOG0AxAAoJCoYAAwAHAQMHaQABAAAFAQBpAAQABQIEBWkAAgsBBggCBmkACQkIYQAICA4JTllZQBcAAEdFQkA4NiooACMAIyEVJSMjJQwHHCsDIiY1NzQjIgcnNjMyFRQGBzY3JjU0NjMyFxYVFAcXBiMjBgY3NjU0JiMiBhUUFwcnNyY1NDY2MzIVFAYHJzc0JiMiBgcWMzI3Bw4CBAIPAgUFEQYSEBEBARMVGRkNCgYIAxQBGgkNJDADBwgFCRU9AickFiIRKAQDEwEIBRAZBg4fFh4CHiggAq0LARIKCQYbGAQJBQQODxEQGQoKFwkIARINFTkGBgkOBwMPCPYnFgggFCQXJQYNCAUEBQgWDxQGJwcODwAAA/+rAckAkQLhACoANgBbARJLsBtQWEAjLwEBBSQBAAFSFQ0MBAUEAE9GRT08BQgEWQEHCAVMFwEAAUsbQCYvAQEFJAEAAQ0BAwBSFQwEBAQDT0ZFPTwFCARZAQcIBkwXAQABS1lLsBVQWEAsCgEEAAgABHIAAgAFAQIFaQAIAAkGCAlpAAcLAQYHBmUDAQAAAWEAAQEOAE4bS7AbUFhALQoBBAAIAAQIgAACAAUBAgVpAAgACQYICWkABwsBBgcGZQMBAAABYQABAQ4AThtAMwADAAQAAwSACgEECAAECH4AAgAFAQIFaQAIAAkGCAlpAAcLAQYHBmUAAAABYQABAQ4ATllZQBs4NwAAWFZJSEJAN1s4WzMxACoAKhkpKCYMBxorAyImNTc3NCMiBwYGByc2NzYzMhUUBzY3JjU0NjMyFxYVFAYHFwYGIycGBjc3NjY1NCYjIgYVFAciJjU0NxcGFRQzMjY1JzcWFjMyNjcmJyc0NjcWFRQGIyInBgYNAxYCAQgGBgMLCAgPDg0JGgQhGyUmFA4KCwICHgETFA0TOUkCAQEKCwkNDhIXCw0DFQoOAhUEEREFCwEBDQIPCREdFBQLBxkCNhIBCxAPAwEGBAoUCgokDwsHExUbFycPECEJDAUCEAsBEyJXBAQHAw0VCwUV0xYVEhgDDAwZDw0eDh0XBwQHHQgIEQIgHBYiDxER////9QHMAMwCdgEGBYAsGAAIsQACsBiwNSsAAAACAAD/egC2AAgACgAUAAi1DgsEAAIyKxU3NjY3FxQGBwYGBzc2NxUUBgcGBgMvVyUBFBcNPDQDXk4dIBYxUiEOHA8GEA4JBhRHIRseBhMRCwgQAAADABf/CwCbAB4AGwAlAC8ATkAUFxMNDAQCAQFMKyohIBgCAQAIAklLsCpQWEAQAAIBAoYAAAABYQABAQ8BThtAFQACAQKGAAABAQBZAAAAAWEAAQABUVm1IyYnAwcZKxcnNyY1NDY2MzIVFAcnNzQjIgYHFjMyNwcOAgc3NjY3FxQHBgYHNzY2NxcUBwYGGQInJBciEScHEwEMEBoFDx4YGwIdJiIVAiM6GAEPFjUXAxQ8JQEPFzaVJxYIIBQjFyQODgYEDBYPFAcnBw0QPBsLEwoIEAgJEjgcBxQOCA8IChMAAAAD//IBPACwAkYAIwAuADoAQkA/IQEBAigBAAMCTBoYDw4HBgYCSjMBAEkAAQMAAVkAAgADAAIDaQABAQBhBAEAAQBRAQAgHhIRDAoAIwEjBQcWKxMiJjU0NjcXBhUUMzI1JzcWFjMyNjcmJyc0NxYVFAYjIicGBgc3NjY3FxQGBwYGBzc2NjcVFAcGBwYGGhEXBQYNAxUYARUDDhQFCwEBDQIYER0UFAsHGC4DF1Y+ARUhFTwhAxdWPwcGMRAvAb4WFgoVCgMPCBkbHw4ZGwcEDBgHEwkfHRYiDxESTSAHHxgGEA4OCRZCIQcfGAYOBwcUBxL//wAF/3oAuwAIAAYFhAUAAAH/9QHEAKsCLAAJAAazBAABMisDNzY2NxcUBwYGCwU/WBkBFxtPAcQrFR4KCxkKDR0AAAAC/+oB8gByAuwACgAoADVAMhUBAQAkIBkYBAIBAkwFBAIASiUNDAsEAkkAAgEChgABAQBhAAAADgFOIyEeHBQSAwcWKwM3NjY3FxQGBwYGByc3JjU0NjYzMhUUBgcnNzQmIyIGBxYzMjcHDgIWAyw6EAEREw0rFAInJBYiESgEAxMBCAUQGQYOHxYeAh4oIAKmHQ4UBwgMDAgFD74nFgggFCQXJQYNCAUEBQgWDxQGJwcODwAC/9oBxACSAp8ACAAtAD1AOisBAQIBTCQhHRgXDw4FBAkCSgABAwABWQACAAMAAgNpAAEBAGEEAQABAFEKCSooHBoUEgktCi0FBxYrAzc2NjcXFAcGByImNTQ3FwYVFDMyNjUnNxYWMzI3JiYnJzQ2NxYVFAYjIicGBiYENVkjARczMBEYDA0DFQoOARQFEBAOBAEIBQIPCRAcFBQLBxgCNysRHg4LGAsYlRYWExYDDAsZDg0fDh8WDAYSDAgIEQIfHRYiDxES//8ACQHJAL8CMQEGBYgUBQAIsQABsAWwNSsAAAAC/+QBsgCTAl4AFgAhAESxBmREQDkFAQMCAgEAAwJMBQEAAwCGAAEABAIBBGkAAgMDAlcAAgIDYQADAgNRAQAeHBQTERALCQAWARYGBxYrsQYARBEiNTY2NyY1NDYzMhcWFRQHFwYGIycGNzc2NTQmIyIGFRQcGi4VJScTDQgOBR8BExQMNjICAwsLCQ0BshkFEw4RHhgmCxEkDA0CEAwBNlgFCAUOFAoFFgAAA//mAfIAcgMeABMAHQA7AGNAYA0BAgQEAQMCAgEAAygBBgU3MywrBAcGBUw4IB8eBAdJCAEAAwUDAAWAAAcGB4YAAQAEAgEEaQACAAMAAgNpAAYGBWEABQUOBk4BADY0MS8nJRgXERAPDgoIABMBEwkHFisDIjU2NyY1NDYzMhcWFQcXBiMjBjc2NTQjIgYVFBcHJzcmNTQ2NjMyFRQGByc3NCYjIgYHFjMyNwcOAgcTIR0ZGg0JBwcDFQEbByQhAw4GCRY9AickFiIRKAQDEwEIBRAZBg4fFh4CHiggAqsRBhMPERAZCgwVEQESJDsGBhcHAw8I9icWCCAUJBclBg0IBQQFCBYPFAYnBw4PAAAD/9kByQCSAuEAGQAlAEoAckBvHhACAgQFAQMCQQICAAM0MywrBAcASAEGBwVMPgEAAUsJAQADBwMAB4AAAQAEAgEEaQACAAMAAgNpAAYIBQZZAAcACAUHCGkABgYFYQoBBQYFUScmAQBHRTc2MS8mSidKISAXFhQTDQsAGQEZCwcWKwMiNTY2NzcmJjU0NjMyFxYVFAYHFwYGIycGNzc2NjU0JiMiBhUUByImNTQ3FwYVFDMyNSc3FhYzMjY3JiYnJzQ2NxYVFAYjIicGBgscGDISAQ8WJxMOCgsDAh8BExQMNjICAgELCwgOExEYDA0DFRgCFQQREAYLAQEIBQIOChAcFBULBxcCNRoFEwwBBxkQFyYPECEHDAcBEAwBNlEEBQYDDRULBBfLFhUUFgMMDBkcHw0eFgcEBxILCAkRAh8dFyEOERH//wALAcoAugJ2AQYFjCcYAAixAAKwGLA1KwAAAAH/9f+TAKv/+wAJAAazBAABMisHNzY2NxcUBwYGCwU/WBkBFxtPbSsVHgoKGgoNHQAAAAACABf/OQCXAB4AGwAlAExAEhcTDQwEAgEBTCEgGAIBAAYCSUuwKlBYQBAAAgEChgAAAAFhAAEBDwFOG0AVAAIBAoYAAAEBAFkAAAABYQABAAFRWbUjJicDBxkrFyc3JjU0NjYzMhUUByc3NCMiBgcWMzI3Bw4CBzc2NjcXFAYHBhkCJyQXIhEnBxMBDBAaBQ8eGBsCHSYiFQQrOhABDhQalScWCCAUIxckDg4GBAwWDxQHJwcNED4cDhUGCAoNCAwAAAAAAv/4AWsArgJWACUALgBDQEAjAQECKgEAAwJMGxgPDg0GBQcCSisBAEkAAQMAAVkAAgADAAIDaQABAQBhBAEAAQBRAQAiIBIRCwkAJQElBQcWKxMiJjU0NxcGFRQzMjYnJzcWFjMyNjcmJyc0NjcWFhUUBiMiJwYGBzc2NjcXFAcGIxEXCw0DFQsOAQEVBBEQBQsBAQ0BDgoJBxwUFQsHGDkFMFgoARcyAc4XFREYAgoOGQ4NHw4fFggDCRwICBECEiIIFiIPERJjKw8fDwoaCRgAAAAB/+8BwACUAkcAIwBCsQZkREA3IgEBAgFMGhcODQYFBgJKAAEDAAFZAAIAAwACA2kAAQEAYQQBAAEAUQEAIR8REAsJACMBIwUHFiuxBgBEEyImNTQ3FwYVFDMyNSc3FhYzMjY3JicnNDY3FhYVFAYjIicGGBIXDA0DFRgCFQUQEAYLAQENAg8JCQccFBQMDgHAFhYTFgMKDhkcHw0fFQcEBx0ICBECESIIFyEOIgAAAAL/7wHKAJQDDgANADEAPUA6MAEBAgFMKCUhHBsUEwgACQJKAAEDAAFZAAIAAwACA2kAAQEAYQQBAAEAUQ8OLy0gHhkXDjEPMQUHFisTJzQnJjU0NjceAhUUByImNTQ3FwYVFDMyNSc3FhYzMjcmJicnNDY3FhYVFAYjIicGLwUPAxIOBwUBLRIXDA0DFRgCFQQREA4EAQcGAg8JCQccFBQMDgJFAkc8DAcRGwUhKiEUKpoWFhQVAgoOGRweDR4VCgUTDQgIEQISIggWIg8jAP//ABEBygC2AlEBBgWTIgoACLEAAbAKsDUrAAD////hAckAxwLhAAYFgjYA//8AAQHKAL8C1AEHBYYADwCOAAixAAOwjrA1K///AAYByQC+AqQBBgWKLAUACLEAArAFsDUrAAD//wAFAckAvgLhAAYFjiwA//8ADAHIAMICswEGBZIUXQAIsQACsF2wNSsAAP//ABEBygC2Aw4ABgWUIgAAAgAQAcsAdgJFAAsAFgBksQZkREuwD1BYQB4AAQMCAXAAAwIDhQUBAgAAAlkFAQICAGIEAQACAFIbQB0AAQMBhQADAgOFBQECAAACWQUBAgIAYgQBAAIAUllAEw0MAQASEQwWDRYHBQALAQsGBxYrsQYARBMiJjU0NjMyFhUUBicyNjU0JiMiBhUUQRYbHxMUICETDBQXDggRAcsgGRwlJR0XISQNBwkVDAwa//8AMgHLAJgCRQAGBZwiAAAB/7wCMAC7AnYAEAAssQZkREAhEAsBAwEAAUwKAQBKAAABAQBZAAAAAWEAAQABUSUmAgcYK7EGAEQDJzYzFxYWMzI2NxcGBiMiJzgMOwcRHSoNETEOCBQ2Ei45AjAVLwIHBwsHFhQaFQAAAAABAAoCAgCwAn0ADQA4sQZkREAtDAECAQFMCwEBSgABAgGFAAIAAAJZAAICAGEDAQACAFEBAAoIBwYADQENBAcWK7EGAEQTIicmNTQ2MxYzMjcXBlwwFwsPCQk0MRAQEQICMBkWCxFRUQN4AAAC//0B0gDOAp4AGgAlAFGxBmREQEYRAQIDEwEBAiEFAgQBA0wAAwIDhQACAQKFAAEEAYUGAQQAAARZBgEEBABhBQEABABRHBsBABslHCUPDQoJBwYAGgEaBwcWK7EGAEQTIiY1NDcmIzY2Mxc2NjMyFhUGBwcWFhUUBgYnMjY1NCYnBhUUFkYSFwUTEgIXFw4dPhUQE0AuAhQZFSACCw8dEQUMAdIpIxENAxIOAR4iEA4MHwIJHRISIhUmDQUQFwQGDhAZAAAAAAEAGgHOAOIB+wAKACqxBmREQB8EAQBKAAABAQBZAAAAAWECAQEAAVEAAAAKAAghAwcXK7EGAEQTNzI2NxQGBw4CGggaYEYKCBpHQgHOIwUFCxUBAwUEAAABAB0BPwDPAeIAGAAcsQZkREAREQEASg0DAgBJAAAAdiYBBxcrsQYARBMmJic3NjYzMhYXFhYXPgI3FhYVFAYHBmwSJBkPBQkEDhUOCA0HAxATCAgOJA4TAT8mSCIEAQEUGAwcERAuKgoBDQ4JQS0LAAAA//8AHQFUAM8B+AAGBfTiAAAB/7f+sQBJ/0QACwAnsQZkREAcAAEAAAFZAAEBAGECAQABAFEBAAcFAAsBCwMHFiuxBgBEESImNTQ2MzIWFRQGHisrHh4rK/6xKx4fKysfHisAAAACAAwB5ADdArAAGgAkAE+xBmREQEQdCgICBBYBAQIYAQABA0wAAgQBBAIBgAABAAQBAH4FAQAAhAADBAQDWQADAwRhAAQDBFEBACEfEQ8JBwUEABoBGgYHFiuxBgBEEyImJwciJicyNjcmJjU0NjMyFhUUBgcWFxQGJzY1NCYjIgYVFLoWPR0OFxcCCBMKAgMXERksExkvQRN7Lg8LDQwB5CMeAQ4SAQEJDwckJy0bERoOIAwOEWkQGwUOGRAJAAADABgB8QDSApkABwAWAB8ACrcaFw8IAwADMisTJzY3FhYXBgc3NjY3NjY3FhUUBgcGBhcnNzcWFhcGBk4kAxYJEgoQQAoXRC4KEAcGFxUoR3AlDQ0HEgwHDgJYHwUdBhAKGmIoCCwkCAsEDw0QDA4bKBofEhAEEAsNEAAAAv8HAekA8QMzAEIATABtsQZkREBiSkUgHxQFAQIzLgIEAT04AgYEA0wRDgIDSgACCAEIAgGAAAMACAIDCGkKBwIBBQEEBgEEaQAGAAAGWQAGBgBhCQEABgBRREMBAElHQ0xETDw6MjAsKiMhHRsXFQBCAUELBxYrsQYARAMiJicmNTQ2NzY2NyYmJzY2NxYVFxYzMjY3NzYzMhcHFzYzMhYWFRQGBwYjIiYnBwYjIicHDgIHFhYzMjcWFRQGBjcyNyYmIyIHFhZ7FSoVKgcDGUJAAggFBwwGAwIUExITBQsFBwUECxNHMwwgFyEcGiIfNxAGDR0dDhcTNTAPDz8gTCsFLTvhMzEMIQ0oOQ4bAekEBAgHAxgEFhwLFmBMBgsEJ0RsCQwJEwcEGghgExkKGiIKCgwJBw4OBQQOEgoEBwkHBg4KAWogFhRFAgMAAAAABP9RAekAtgMzAD4ARQBMAFYAdbEGZERAak8OAgYHJhsCAgYaFQIBAi4BBAE4MwIFBAVMSklIRENCFBEIA0oAAwAHBgMHaQkBBgACAQYCaQABAAQFAQRpAAUAAAVZAAUFAGEIAQAFAFFOTQEAU1FNVk5WNzUtKyQiHhwYFgA+AT0KBxYrsQYARAMiJicmNTQ2NzY2NyYmJyYmJzY2NxcWMzI2NycGIyImNTQ2MzIWFRQGBwYGIyInBgYHBgcWFjMyNxYVFAcGBhMmJic3FwYHJic3FwYGFzI3JiYjIgYVFDIVKxQpBwIYQkECBAIBBAIGDQcEFBQmPBYFFiAQFRwTHyQGCBo8Jh0PHScJOBkQPyBLKwYjEC+aBgwHGRcHSBEIGBcCCxwTFQYYDgwPAekEBAcIBRcDFRwMFy8ZGjAZBQsF1wkKCh0UGRIeK0EsBRMOEQoOBwkDEBAEBwkHBhIDAgIBEwULBxkWDBQNCxkXBA1zChQaFA4WAAL/PgGNAEYCQQAXACIAT7EGZERARCABBQQVAQIFAwEDAANMAAEABAUBBGkHAQUAAgAFAmkAAAMDAFkAAAADYQYBAwADURgYAAAYIhgiGxoAFwAXFiUUCAcZK7EGAEQDIiYnNTI2Nz4CMzIWFhUUBgciJicGBjcmJiMiBwYGBxYWoAcRCjhKHAkQFA8NFA0LCBUuERVAoQEPCQgJBAoGEBwBjQMBCjUuDh8WHSkTDxIGDAofK1EULQwGEQsJCQAAAAH/wQF+AEMCaAAkADixBmREQC0aFBMDAAEBTCIeAAMDSQACAAEAAgFpAAADAwBZAAAAA2EAAwADURQjJRoEBxorsQYARAM0JyYmNTQ3PgI3JyYmJyYjIgcnNjMyFhYXByIGBxcWFhcGBjoDAQEMFRUVFAgHCgMHBg0SDBUfDBMWDg4eLBoGAwQBBA0BficbDRUHCB4GBQIBCAsMBAYXBTQYIAwnBQkkEBsNBAsAAAAAAv+ZAdEAPALVACoAMwA/sQZkREA0Mi8iFRIHBgABAUwfHhsZDQUBSgABAAGFAAACAgBZAAAAAmEDAQIAAlEAAAAqACgVIwQHGCuxBgBEAyc0NjMzNjcmJicjJic0NxYWFzU2NTQmJyc0NxYWFwcXFAcWFRQGBwYGIzc3NjY3JiYnBkABBAEJLBAUKBQNEQINHj8RBwQEBRoCBgUKAw8dCQMQLx0XCw4bDQQLCBMB0QUGEiMWFiIKEA8QBhY6FQIOEQgkGR0TBgYWEQk+HhoxBwUSAgQDHgEBAwEIEQoYAAAC/5QBgQBkAoUAJwAwAFWxBmREQEoYAQMCDAYCBAEvLCslBAUEJgEABQRMAAIAAQQCAWkAAwAEBQMEaQAFAAAFWQAFBQBhBgEABQBRAQAjIRsaFxUSEAoIACcBJwcHFiuxBgBEEyImNTQ2NycmIyIGByY1NDYzMhcWFjMzNwcHBgYHBhUUFjMyNjcXBicmJic3FhYXBgIuMzIlHwsLDxEKBRoUEi0VIQsMDQoTIUgUFDElFy4XBTgrCA0FGggMBQoBgTIuJ0EOCgQJBgwDDhIOBgcBHQIBIBcZGh0jBQUNHUQHDAYbBw0EDQAAAAP/oQEcAF0BugAHAA4AFwAKtxUPDAgEAAMyKxEmJzY3FhcGFyYnNzcXBgcmJic2NjcXBhYUGBAMGx8wHw0XFCgTfQ0VCgQVEigSAW4MGBUTERcgThsOFxUoGxoJFQoDFhMnGwAB/2ABpgCwApAAQQBWsQZkREBLMyQTAwIDOQYFAwUCDgEBBQNMLgEDSgADAgOFBAECBgEFAQIFaQABAAABWQABAQBhBwEAAQBRAQA9Ozc2KCcgHhkXDAoAQQFBCAcWK7EGAEQDIiY1NDcXBhUUFjMyNjcmNTQ2NxYWFxYzMjc3Njc2MzIXBgYHBxYWMyYnJjU0NxYXFhYVFAYHIiYnBwYjIxUUBgZcIyEHDgUeFx01ChgOCgIFAgsNFwkFCgUGCAUDAwcFARQaFAYIBxkDBgIDCgQWJRUDBiwNHzEBpigfGBkEEggdHRoUNRQMDwYNFgsDBAohCgkDBxUQBAkFFxIPChEPFhULEwkIHggGCQUKDhgqGgAC/8oBxAA2AjcACwAVADmxBmREQC4AAQADAgEDaQUBAgAAAlkFAQICAGEEAQACAFENDAEAExEMFQ0VBwUACwELBgcWK7EGAEQRIiY1NDYzMhYVFAYnMjY1NCYjIhUUFiAgFhgeHhgMEhIMHgHEGx8eGx0cHB4aDxERDyAgAAAAAAL/ygGdADYCNwAHABEAObEGZERALgABAAMCAQNpBQECAAACWQUBAgIAYQQBAAIAUQkIAQAODAgRCREFAwAHAQcGBxYrsQYARBEiNTQzMhUUJzI1NCYjIgYVFDY2NjYeDhAQDgGdTU1NTRozHBgYHDMAAAAAAf+vAcQAVAJEABUAVLEGZERADBIBAAEJBAEDAgACTEuwDlBYQBYAAgAAAnEAAQAAAVkAAQEAYQAAAQBRG0AVAAIAAoYAAQAAAVkAAQEAYQAAAQBRWbUUJjUDBxkrsQYARAMnNjY3JiYjIgcmJjU0NjMyFhcGBiM1HDI6FA0eDyEPAQQbHxopFw8+OAHEEgQdJgMCEQMHBg0WCAgxPwAAAAH/YP7+ALD/6ABEAFixBmREQE00JiMSBAIDQDoGBQQFAg4BAQUDTC8BA0oAAwIDhQQBAgYBBQECBWkAAQAAAVkAAQEAYQcBAAEAUQEAPz44NygnISAZFwwKAEQBRAgHFiuxBgBEAyImNTQ3FwYVFBYzMjY3JjU0NxYWFxYWMzI2Nzc2Njc2MzIXBgYHFjMmJicmNTQ3FhcWFhUUBgciJicHBgcGIycVFAYGXB8lBw4FHhcdNgkYGAIFAgcMBQsQBQUFBwIGCQYCBAgEICIDBwQHGQMGAgMLAxYiGAMDDg0UDR8y/v4lIhUbAxIIHh0dEzMVFA0NFgsBAwMBCxEVBAkCChgODw0VCA8KEQ8WFQoUCQcgBwYKBgQDAwEPGCoaAAH/kgHVAGUCFgAQACyxBmREQCEQCwEDAQABTAoBAEoAAAEBAFkAAAABYQABAAFRJSYCBxgrsQYARAMnNjMXFhYzMjY3FwYGIyInYwszBA4YJAoOKAwGEC0PKC0B1RQrAgYHCwYUExgTAAAAAAH/egHAAH0CNgAeAC6xBmREQCMZFA0MBAFKAAEAAAFZAAEBAGECAQABAFEBABcVAB4BHgMHFiuxBgBEAyInJjU0Njc2Nzc2NxcGBgcOAgcWMzI2NxYVFAYGBi0qKQgEDRMrDAcMBQ8OFBIODiNQG0UZBSw9AcAGCAYHIQINCBAGDQYOFAQGBwgJDwoHBgcLDwcAAAAC/6QBpABUAnkACAAmAC+xBmREQCQdGRAPBAUBSgABAAABWQABAQBhAgEAAQBRCgkXFQkmCiYDBxYrsQYARAMmJic3FhYXBgciJjU0NjcXBgYVFBYzMjY3JjU0NxYWFxYWFRQGBhAJDQQZCw0CCBkjIQIFDgICHRcdNwgXFwIFAwIDHjECRQgMBBwKDQIMsCcfCxAWAwoNBB0dGxQ2EhMPDRkLCxULGCoaAAAAAv+//qoAQf8tAAMABwAItQYEAgACMisRJzcXBzcnB0FBQUEpKSn+qkJBQSkpKSkAAv+/Ac4AQQJRAAMABwAItQYEAgACMisRJzcXBzcnB0FBQUEpKSkBzkJBQSkpKCgAAf+3AbcASQJKAAsAJ7EGZERAHAABAAABWQABAQBhAgEAAQBRAQAHBQALAQsDBxYrsQYARBEiJjU0NjMyFhUUBh4rKx4eKysBtysfHisrHh8rAAAAAf/B/qQAQ/+OACMAPLEGZERAMRkTEgMAAQFMIR0AAwNJAAABAwEAA4AAAwOEAAIBAQJZAAICAWEAAQIBURQjIxsEBxorsQYARAM0JyYmNTQ3NjY3NjYzJyYjIgcnNjMyFhYXByIGBxcWFhcGBjoDAQEMEhoICQ8HEg0KDRIMFR8MExYODhkyGQYDBAEEDf6kJxsNFQcIHgUGAQEBFxIXBTQYIAwnBAokEBsNBAsAAAAB/6QB7ACMAsgAHwApsQZkREAeDg0CAEofAQFJAAABAQBZAAAAAWEAAQABUS8XAgcYK7EGAEQDJiY1NDY3NhY3NjU0JzcWFhUUBgcGBgciBgcGFRQWF0wHCRgfFy4cMQ4RDg4kIw0bDA0WCiIFBAHsDhwOGCgHBQIIDxYLFA4KGg4cLQkEBAEBAQMeCRIIAAAC/+sBsgCgAl4ALAA4AEqxBmREQD8hCwoDAQQUAQIBEQEDAgNMBQEDAgOGAAAABAEABGkAAQICAVkAAQECYQACAQJRAAA0MgAsACwnJiMiGxkGBxYrsQYARBMiJiY1JiY1NDY3FwYUFRQWFzY2NyYmNTQ2MzIXFhUUBgcWFjMGBiMiJiMGBjc2NjU0JiMiBhUUFgoBDgwBAwkFDAEEAxchDBAWJxMMDAoBAwgQBwETFAMGBRU6TQMBCgsJDRcBsgkKAQoSCg8fDQQJDwUNGQ0FEQgHGBEXJg8OIwQNCAEBEAwBFCJYBggEDhULBQ0SAP///6T+ewCM/1cDBwW6AAD8jwAJsQABuPyPsDUrAAAA////cAHsAIwDNgImBboAAAEGBbrMbgAIsQEBsG6wNSv///+yAbIAoAL3AiYFuwAAAQ8FuwBSBKnAAAAJsQICuASpsDUrAAAA////f/5rAJv/tAAnBboAD/x/AQcFuv/b/OwAErEAAbj8f7A1K7EBAbj87LA1KwAAAAH/zgJ3ADUC4AAKAAazBQABMisTJic2NjcWFhcGBgQkEhMaBwwaDQkZAnccFRQbCQ0ZCw0cAAD///+hAnoAYgLeAQcF8f8uASAACbEAArgBILA1KwAAAAAC/7kCGwB2AqAAHAAmAE+xBmREQEQUCgIDASUgGhYEAgMXAQACA0wRAQFKAAEAAwIBA2kFAQIAAAJZBQECAgBhBAEAAgBRHh0BACMhHSYeJgkHABwBHAYHFiuxBgBEAyImNTQ2NzYzMhc2Njc2NjcXBgYHFhcHJiYnBgYnMjY3JiMiBgcWHBIZBQMbIx0eBQcEBQYEDAYNCBcVDgsaDRIpEA4ZDRYKExIRCwIbIhkKEAcUEwULBQUKBAgKFQsSGgsHDgYWISYPCwgEBhgAAAD////O/xEANf96AwcFwAAA/JoACbEAAbj8mrA1KwAAAP///6H/EgBi/3YBBwXx/y79uAAJsQACuP24sDUrAAAA////uf7wAHb/dQMHBcIAAPzVAAmxAAK4/NWwNSsAAAD////1AcQBAgKgAiYFiAAAAQYFiFd0AAixAQGwdLA1K////+QBsgElAl4AJwWMAJIAAAIGBYwAAP////X/HwEC//sCJgWQAIwBBgWQVwAACbEAAbj/jLA1KwAAAP///7MBigBEAlcBBwTK/54BnwAJsQACuAGfsDUrAAAAAAP/qQHgAGACcwAaACUALwBQsQZkREBFEwECAS4rKSQeGxUUEgkDAgUCAgADA0wAAQACAwECaQUBAwAAA1kFAQMDAGEEAQADAFEnJgEAJi8nLyIgEA4AGgEaBgcWK7EGAEQTIicGBgc0Njc3NCY1NDYzMhYXNwcHFhYVFAYnNjY3JiYjIgYVFBcyNjU0JwYGBxYFIxAFGgoFCRMBJxYOHQolAhoBASdBECEQBg8LDhQqEhkEESIQCgHgIAMOBQsQBQwCBgQlLBETER4MBAsGHiNCCBEIBQgXDwUhEwoGBggRCAgA////mgHCAGACaAImBisAAAAGBikAAAAA////mv7tAGD/nQInBisAAP1tAQcGKQAc/MIAErEAAbj9bbA1K7EBAbj8wrA1KwAAAAH/rAGnADQCVgAYABqxBmREQA8TDwMABABJAAAAdiwBBxcrsQYARBMmJic0Njc2NzY2NzYzMhcOAgcWFhcWFjQmOSkCBCQWCxIHBwoLBgcjJw4ZIBAOCgGnEBUIERYNGhQLDgMEEAgdHAcEDQsKIgAB/68BuwBRAmsAGQAgsQZkREAVEg8HAwQASgEBAAB2AQAAGQEZAgcWK7EGAEQDIiYnPgI3JiYnJjU0NjcWFhcGBgcGBgcGOAoOAQopLhAWGwwLBAIdNCIECQgVJA8fAbsQCQYQDgIMFhAODgwQBxkoFBAUCwcOCA8AAAD///+s/wgANP+3AwcFzQAA/WEACbEAAbj9YbA1KwAAAP///6//BgBR/7YDBwXOAAD9SwAJsQABuP1LsDUrAAAA////gAGrAJ4CawAmBc7RAAEGBc5N8AAJsQEBuP/wsDUrAAAA////VwGrAJ4CawAmBc7RAAAmBc5N8AEGBimH1QASsQEBuP/wsDUrsQIBuP/VsDUr////igG7AFECawImBc4AAAEGBim61QAJsQEBuP/VsDUrAAAA////sQGyAIQCXgAmBYzxAAEGBinh1QAJsQIBuP/VsDUrAAAA//8ANQFnARMCLAIGBeAAAP//AFUBQwEzAroCJgXbAAABBwXwAJ8AnQAIsQECsJ2wNSsAAAAC/+QBsgCTAl4AFgAhADxAOQUBAwICAQADAkwFAQADAIYAAQAEAgEEaQACAwMCVwACAgNhAAMCA1EBAB4cFBMREAsJABYBFgYGFisRIjU2NjcmNTQ2MzIXFhUUBxcGBiMnBjc3NjU0JiMiBhUUHBouFSUnEw0IDgUfARMUDDYyAgMLCwkNAbIZBRMOER4YJgsRJAwNAhAMATZYBQgFDhQKBRYAAAH/+QH2AHkCqgAbADWxBmREQCoYFA0MBAIBAUwZAgEDAkkAAgEChgAAAQEAWQAAAAFhAAEAAVEjJycDBxkrsQYARAMnNyY1NDY2MzIVFAcnNzQmIyIGBxYzMjcHBgYEAyglFiMRJwcTAQgFEBkGER8SHwIpNQH2JxYGIhMkGCUODQUEBQgXDxQGJgoWAAAAAAEAHQE/AM8B4gAYABRAEREBAEoNAwIASQAAAHYmAQYXKxMmJic3NjYzMhYXFhYXPgI3FhYVFAYHBmwSJBkPBQkEDhUOCA0HAxATCAgOJA4TAT8mSCIEAQEUGAwcERAuKgoBDQ4JQS0LAAAAAAEAgQFDAQkBzQAIAAazBAABMisTJic2NxYWFwbJLBwzExAhERcBQyEgMxYRIQ8hAAIALv8+AUP/zQAJABQACLUPCgQAAjIrBSYnNjcWFhcGBgcmJzY2NxYWFwYGAQYqFxYqBR4bCh6qLRYOIBMEHxoJHrMhHBUuBB4aDiIjJBkMIhYFHhkPIgADAFn++gFZ/8IACQAUAB0ACrcaFRAKBQADMisFFhYXBgcmJic2JxYWFwYGByYmJzYXFhcGBgcmJzYBHRQeChgiBRwWEmwSHQ0JHRQOGw4TahsdDxsLEyQbPhAbDBYnBhsVHBUNHA4IHhcOHA0eMxEhDhsOGR4hAAAEAE4BLQEXAe0ACQATAB4AKQANQAokHxoUDwoFAAQyKxMmJzY2NxYWFwYHJic2NjcWFhcGByYmJzY2NxYXBgYXJic2NjcWFhcGBuMeEgMYFAUXEBCAHhIKFw4GFhAPEhAYCAgYDxQYCxVaGxUIFw8LFgsIFQGRFxUDGBUGFg8YGRcVCRgPBhYPGH0MFgoIGRAXFQ8ZCRQYCBkQDBULDRkAAgA1ARUAswIFAAsAFQAItRIMBgACMisTJiYnNjY3FhYXBgYHJiYnNjY3FhcGcRUdCgoeEwUcFgoaBhMdCwgdFQwrEQGSERsLCR4VBhsVDh6ODhwNCB8WDSoaAAACADUBZwETAiwAGwAkAIOxBmREQBURAQIBEgEEAh8HBgMEAwQCAQADBExLsAlQWEAhAAECAwFwAAIABAMCBGkGAQMAAANZBgEDAwBiBQEAAwBSG0AgAAECAYUAAgAEAwIEaQYBAwAAA1kGAQMDAGIFAQADAFJZQBUdHAEAIyEcJB0kFhQNDAAbARsHBxYrsQYARBMiJzcWFhc3JiY1NDY3FhYXBxc2NjMyFhYVFAYnMjY3JiYjIgevQzcGCRcOEgoOFhMDBgIIAxEfDg4fFjooFyITEBgLHSoBZyAZAgQCFCIzCg4SAQsRBgZBDhcTGgogKiwJBw4IJAAAAAADAFsBRAFbAgwACAATAB0ACrcZFA4JBQADMisTJic2NjcWFwYXJic2NjcWFhcGBgcmJic2NxYXBgbeGh4PGwsSJR0wIRoJHhMOGw4KGqATHgsYIgotCRoBoxAhDxwNGCAiYRodCB8WDhwNDh8dDxwMFicMKw0fAAADAF4BVQFfAhoACgAVAB8ACrcbFhALBgADMisBJiYnNjY3FhcGBgcmJic2NxYWFwYGFyYnNjY3FhcGBgEpER4NCB4VDygJG6ETHgsYIgkcEwkbOCEYDhsNGB4QGQGmDhsOCB8WECcNHx4QGwwVKAoaEg4fVRYbDRwOHhkRGAAAAAIASQFEAV4B0wAKABUACLUQCwYAAjIrASYmJzY2NxYWFwYHJiYnNjcWFhcGBgEiFiELBSAbBh8ZE78WIgsaJwYeGQkeAVIRHw4FIhwGHhgeNRIfDBkrBx4XDyIABABO/wQBF//EAAkAEwAeACkADUAKJB8ZFA8KBQAEMisXJic2NjcWFwYGByYnNjY3FhYXBgcmJzY2NxYWFwYGFyYnNjY3FhYXBgbjHhIMGAsLIQgVcx4SCBgPBhYQEBEeEgwYCwYWEAoVWR0TDBcLCxYLBxWYFxULGQwMIAsYDRcVCBgQBhYQGHwXFQwYDAYWDw0ZCxYWDBgMCxULCxkAAAIAOv7XALf/yAAJABUACLUQCgYAAjIrFyYmJzY2NxYXBgcmJic2NjcWFhcGBnUUHgkMHREMKxEYFR4JCB0VCBsUCRqsERwLCx4TDSoaoBEcCwgeFgcbFA0gAAAAAwBT/wsBVP/QAAoAFgAgAAq3HBcRCwYAAzIrBSYmJzY2NxYWFwYHJiYnNjY3FhYXBgYXJiYnNjcWFhcGAR4THgsOHRANGw8RtBIeDAgdFgUcFQgbOg0dDh4XCBwTHqMOHA4LHhIOGg4aMA4cDQceGAcbFQ0eVwkZEBsdChwSIgABAHD/PgD4/8kACAAGswUAATIrFyYmJzY3FhcGuBYkDi0ZIx8YwhEhECseJR0hAAAC//4BxQCCAkIACwAXADFALgABAAMCAQNpBQECAAACWQUBAgIAYQQBAAIAUQ0MAQATEQwXDRcHBQALAQsGBhYrEyImNTQ2MzIWFRQGJzI2NTQmIyIGFRQWPxwlJR0cJiYdDg8QDA0QDwHFIh0dISEcHiIiEA0ODg4ODRAAAgBEAAADHQHhACwAOABLQEgwAQUBFwEDBQJMAAEGBQYBBYAABAAGAQQGaQgBBQADAgUDaQACAgBiBwEAAA8ATi4tAQA0Mi04LjgiIBoYEA4GBQAsASwJBxYrISI1NDY3MwYGFRQWFxYWMzI+AjcmJicGIyImNTQ2NzYzMhYXFhYVFAcOAjcyNjcmJiMiBhUUFgFA/AcHJAICLTIaQikycnFiIQMGBUNLKzghGh0kMlEYCwwjLI6p0xg1HBM9Jx4pKaogNhgMGQw9PREHCQsTGg8UJxM3QjMyVxofSEQiSykkQhwoFfYNDD1ANyIgHQAAAQBNAAAAsgILABIABrMIAAEyKzMuBDU0NxYWFwcWFBcUBgdlAgYHBQRBBhMLGwEBAgQ1dG9fPwk/DR47HRclSCQZc1kAAAAAAQBOAAABDgIOABIAH0AcEQEAAQFMDAgGAwFKAAEBAGEAAAAPAE4sEAIHGCshIiYnJicnNjcWFhcXFjMyFRUUAQU6TgUOChIXLgIFAQIvOQkvKGF1xRELLqt9URsIPAgAAQAIAYQAlgJLACAALkArHBgREAQCAQFMHQQBAAQCSQACAQKGAAABAQBZAAAAAWEAAQABUSMoKgMGGSsTJzY2NyY1NDY3NjMyFhUUByc3NCYjIgYHFjMyNwcOAgsDCxYLKRkSERYYFAgVAQgHEB0GESIZHwMaKSkBhCsHDAYIJBUpDA0YEQsSBQUFCRkRFQYrBg0TAAH/+AAAASQBdgAbACVAIgMCAgEAAUwSCQIASgAAAAFhAgEBAQ8BTgAAABsAGxUDBxcrMSI1NTQzMjY2NyYmJyYmNTQ2Nx4CFRQHDgIICDFfThcFEg4JCiQeDRAIIhtVYgg8CAkRCxQ3IxohCBYuEDpLNh4jQxAZDgAAAAAB//gAAAF5AOQAIQBSQA4PAQABHRkYAwIFAwACTEuwCVBYQBQAAQAAAXACAQAAA2IFBAIDAw8DThtAEwABAAGFAgEAAANiBQQCAwMPA05ZQA0AAAAhACEVFyYVBgcaKzEiNTU0MzI2NzY2NzYzMhcOAgcWFjMyFRUUIyInBw4CCAg+VhILHh0JFgwNBhUSBCI6IggIXUEQBz9VCDwICw0HMTcRCA8vKAQVEQg8CDoVBxIMAAAC/80CHACSAosAFQAfAEBAPR4ZBgMDABQBAgIDAkwAAAQDBAADgAABAAQAAQRpBQEDAgIDWQUBAwMCYQACAwJRFxYcGhYfFx8mJBMGBhkrAyc2NjMyFzY2MzIWFRQHBgYjIiYnBjcyNjcmIyIGBxYlDhEaEAUMFiMVERoIERsSEyIaE2IVEBALFgsWFBUCHAoZGgYgGCEZDRUMBwoNCBkDBxcKEQYAAAAC/7YBWACUAh0AGgAjAEhARREBAwEeBwYDBAIDAgEAAgNMEAwJAwFKAAEAAwIBA2kFAQIAAAJZBQECAgBhBAEAAgBRHBsBACIgGyMcIxUTABoBGgYGFisTIic3FhYXNyY1NDY3FhYXBxc2NjMyFhYVFAYnMjY3JiYjIgcwQzcGCRcPERgYEQMGAggDER8ODh8WOigXIhMQGAscKwFYIBkBBAIUVgkIFAQKEQYHQA4XFBoKICosCQgOByMAAAIAcwFaATQBvgAHABAACLUNCAQAAjIrASYnNjcWFwYHJic2NjcXBgYBCiAODCETGBh7IQ0LFwsqChQBYxoRDCQWFR8aGRIKGA0qDhgAAAAAAwAl/zYBzP/RAAkAFQAfAAq3GhYQCgUAAzIrBSYnNjY3FhYXBgcmJic2NjcWFhcGBgcmJzY3FhYXBgYBkC4UDSAUDh0SE78WIQsNIBMFHhoJHqcrFxwkBh4ZCR2vJRgLIhYPHBEfNBIfDAwiFgUeGQ8iICIbHCcGHhgNIwAAAAABAC0BhgDwAjoAGAAVQBIQAQBKEwwDAwBJAAAAdiUBBhcrEyYmJzc2MzIWFxYWFz4CNxYWFRQHBgcGgxMoGxALCREeDQgMBQQTFQcJDg8WERQBhilQJQUDIhsNGg0VNSwJAQ8PBx0tMwwAAAABADsBVADtAfgAGAAasQZkREAPFQwJAwQASgAAAHYuAQcXK7EGAEQTJiY1NDY2NzY3FhYXBgYjIiYnJiYnDgJRCA4RGAoVGxUnEwcTBxAVCwQPCgMQEwFUAQ4OBSIzHA0ELkgaAgQYEQgfFhIuKAAAAf+2AfoAeQKmACwASkBHIx4XFgQEAwsBAQAkAQYBAQEFBgRMAAQDAAMEAIAAAgADBAIDaQABBgUBWQAAAAYFAAZpAAEBBWEABQEFUSIkJCcnJBMHBh0rAyc2NjMyFhcWMzI3JjU0Njc2MzIVFAcnNzQmIyIGBxYWMzI3BwYGIyInJiMiPA4JHQwIDgQOAwcJJBYREg4iBBIBBwULGggLIAwSFgIdOQ4JEA0IEAH6CxAaBAIDAxQcESQMDCMQCwQEBAgUDgsPCSQRFQgHAAAA////xgAAAdIDAQImAt4AAAEGBfUQWwAIsQIBsFuwNSv////jAAACcwL4AiYC3wAAAQYF9S1SAAixAQGwUrA1K///ACMAAAHzArACJgLeAAABBwYNAREATgAIsQIBsE6wNSsAAP//ACIAAAJzAp8CJgLfAAABBwYNAVUAfQAIsQEBsH2wNSsAAAAB/zwCUgDTA0IAFQAGswsAATIrAyYmNTQzNjY3NjY3FhUGBgcGBgcGBrwEBAIcOBxGikUQAgoQDyIbNJcCUhEXBwUOHQ4gQiEMEg8RBwcPDBdGAAAB/0kCSwEPA0gAEwAGswsAATIrAyYmNTQ3NjY3NjY3FhUGBgcHBgavBAQCK2U6QXU0EAMLD3tGkwJLERYHBQIVLxkdNhgMEhARBjYdQgAAAAL/+AAAAWoBkQAfAC4AQUA+IwEEBQoBAgQCTAADAAUEAwVpBwEEAAIBBAJpAAEBAGEGAQAADwBOISABACgmIC4hLhYUDgwHBAAfAR4IBxYrMyI1NTQzMjY2NycGBiMiJjU0Njc2MzIWFxYVFAcOAjcyNjcuAiMiBgcGFRQWAQkJWHNNHgQcMRpJQScfICArShYXIylUcnccLxAEIzIaFSQLCzoJOgkECgoXDAg+KzBbGxtDOTpBJUwPEgijCAUjQCkZFhYUGScAAf/4AAAA4gEXABYAJUAiAwICAQABTA8IAgBKAAAAAWECAQEBDwFOAAAAFgAWFQMHFysxIjU1NDMyNjcmJicmNTQ3FhYVFAYHBggILGIkBAsICj4KCRIRagg8CAsJDyITGAwsIzRGEQ83ICb//wAw//4BlAJNAAYFDt4AAAEALwAAAaYCRwAkAB9AHCAZAgFKAAEBAGECAQAADwBOAQAeGwAkASQDBxYrMyImJicmJjU0NjY3NjY3FhUUBwYGBw4CFRYWMzI2NjcUBgcG1BpBNgwFAzJVNCM+GRQXOVcdFSseE1g1GEBBFx0XMAUHAxYbFjludEEqSSIUHh8dQ2omGz01EAkMBAYEJisFDAAAAAABACkBuwDEAtcAGgAoQCUWAQEAAUwSEQoDAEoBAQFJAAABAQBZAAAAAWEAAQABUSMuAgYYKxMnJicmJicmNTQ3FhYXFjMyNxcUIyInFhUUBnQTAQwGEw4EGwgMBg4SLAYURQkJDgQBuwVNLh0xGwcHFg8OGxAIQQRnBDIjFTAAAAEAFgG7ANIC1wAtAEFAPhkVEgMBAiUBAwApAQQDA0wTCgICSgEBBEkAAgEChQAAAwQAWQABAAMEAQNpAAAABGEABAAEUSIlIyYuBQYbKxMnJiYnJicmNTQ3FhYXFjMyNSc3FhUWMzI1NCY1FxYWFQYGIyInBiMiJxYWFRRhEwEGBgwbBBsHDQYLERwBDwcIBhYBFQEBARgSCw0PFAkJBgcBuwUpQBYwNQcIFQ8MHBEIJRcFGRYDIwQHAwIHDAcdIQcXBBcqFDcAAf/4AAAAMABMAAsAIkAfCQgDAgQBAAFMAAAAAV8CAQEBDwFOAAAACwALFQMHFysxIjU1NDMzMhUVFCMICCgICAg8CAg8CAAAAv+sAmUAWQLsAAcAEAAItQ4IBQACMisTJic2NjcXBgcmJic2NjcXBi8fDwYXECsYaBAWBwsWCysVApEZEgYXEysfPQ0WCAsXDSoeAAMAAAKbALIDTwAIABIAGwAKtxkTDwkFAAMyKxMmJic2NxYXBhcmJic2NjcXBgYHJiYnNjY3FwZHChoIGBIMHhcxEBcIBxcQKgwVYxAXBwsXCyoSAv4HFgoVFRAaG0cNFQkFGBIqEBgwDBYICxcNKhsAAAAAAwAAAj8AsQLbAAoAEwAeAAq3GhQQCwUAAzIrEyYnNjY3FhYXBgYXJic2NjcWFwYHJiYnNjY3FhcGBlIXIggLAgobEwcLSyAdBg0EGR8DchAeDgIMCBseAQoCkgMPExwIBAsGFhknBw4PHg8NCBVYBAkHBR8ZDAkLHgAAAAQAGwFMAOkCbQAeACgAMwA8AFFAThQBAwEmIgcGAwUCAwIBAAIDTBMPAgFKOzg2Mi8sBgBJAAEAAwIBA2kFAQIAAAJZBQECAgBhBAEAAgBRIB8BACUjHyggJxkXAB4BHgYGFisTIic3FhYXNyYmNTQ3NjY3FhYXBxc2NzYzMhYWFRQGJzI2NyYjIgcWFhcmJic2NjcWFhcGByYnNjcWFhcGiT4wBQgUDREKDAoFDQkDBQMIAhgMCwkNHBQ0JRUeExkWGiYKFUAQFwcGFxALFQoXfCANFhYLFgoVAbseFgIEAhMnKgQHCAUGBAoPBQY6FQYGEhcJHSYoCAcTIAEBjQ0VCQUYEgoVCx8bGhEUGwoVCx4AAAEAIQG6ANMC4QAmACRAIQ8BAEomHhkMBAFJAAABAQBZAAAAAWEAAQABUR0bKAIGFysTJicmNTQ2NzYzMhYXNjY3FhQVFAYHBgYHByYmIyIVFBYXFhcWFhVsBhUwDQ0OEA4fDAkbHAEHCBAPCxcHIRERAgIIEBoQAbonJ1khEykNDCMdFx8UBAkFDBkIDhcZCSMvEQMLBxchND0fAAEANQGSAKcCBQALAAazBgABMisTJiYnNjY3FhYXBgZxFR0KCh4TBRwWChoBkhEbCwkeFQYbFQ4eAAAAAAH/+AAAAaQBdgAbACVAIgMCAgEAAUwRCgIASgAAAAFhAgEBAQ8BTgAAABsAGiUDBxcrMSI1NTQzMjY3NjcmJyY1NDY3HgIVFAcGBwYGCAhWmDY6FgsZEyIgBREPIzWyK04IPAgKCAgLKUUyERUtEh9QTR0iRCAPBAQAAAAB//gAAAH5AOQAKQAyQC8WEwIAASIdHAMCBQMAAkwAAQABhQIBAAADYQUEAgMDDwNOAAAAKQApFRcpJQYHGisxIjU1NDMyPgI3Njc2Njc2MzIXBgYHHgIzMhUVFCMiJicGBgcOAwgIHU1NOgkMCgUWFgkXDAsLGA0YLDorCAhHYiAECAQEP1pbCDwIAgYJBgkPCCgoEQgcNRoOEQYIPAgdHgYLBQUMDAgAAAAAAQA0//wBGQF3ACgAI0AgEQEASigaDgMBSQAAAQEAWQAAAAFhAAEAAVEgHioCBxcrFyYnJiY1NDY3NjYzMhYXNjY3FBYVFAcGBgcHJiYnJiMiFRQXFhcWFhWSCRkWJhITCBEKEigPCyIlAhgTEQwjBRALFRUKBwgTHBgENS8qVh0ZNxAFBywlGygcBw4HMhgSGh8PFyoQHQoOFBUkOVkrAAEAEv/9AO4BYwAjAB9AHB4NAgFKAAEBAGECAQAADwBOAQAcGgAjASMDBxYrFyImJyY0NTQ2Njc2NjcWFRQGBwYGBwYGBxYWMzI2NxQGBwYGchc5DAQeRDcIEAgQGRcrLgsKCgIJNRwaMxYJBw06AwYFDRMNI0lbPwkUCxIVGB8aMTUODRMFBwsGBBUbCBAFAAEAAAFpAOIBmQADAB5AGwAAAQEAVwAAAAFfAgEBAAFPAAAAAwADEQMGFysRNTMV4gFpMDAAAAAAAgAAAQcA/QGZAAMABwAvQCwAAAQBAQIAAWcAAgMDAlcAAgIDXwUBAwIDTwQEAAAEBwQHBgUAAwADEQYGFysRNTMVBzUzFe/g7gFoMTFhMTEAAAAAAQAAAT4BGgFuAAMAHkAbAAABAQBXAAAAAV8CAQEAAU8AAAADAAMRAwYXKxE1IRUBGgE+MDAAAAACAAAA3QEqAW4AAwAHAC9ALAAABAEBAgABZwACAwMCVwACAgNfBQEDAgNPBAQAAAQHBAcGBQADAAMRBgcXKxE1IRUFNSEVARr+9gEaAT4wMGEwMAACAcsCCgKGAzAACwAzADxAOTEBAQIBTCglGxoSEQgFCAJKAAEDAAFZAAIAAwACA2kAAQEAYQQBAAEAUQ0MLi0eHRcVDDMNMwUGFisBJyYmJyc0NjcWFRQHIiY1NDcXBhUUMzI2NSc3FhYzMjY3JiYnJzQ2NxYWFRQGIyImJwYGAhcFAQQFAhAOBjEUGg0PBBgLEAEXBBAXBgwBAQgHAhEKCwggFwgVBwcdApUCEzEfExAPBCcrLacaGBgWAw0OHBINIw8bHwgEBRUPCQoSAxUmChkmCAkTFAAB//b/+wIQAp0ABwAgQB0BAQBKAwEBSQAAAQEAVwAAAAFfAAEAAU8RFAIHGCsFATcTEzMVIwEP/ucr5IWGUgUCixf96gEKNQABALwBkQFZArwAKgAuQCsmAQABAUwlGxcRDwUGAUoAAQAAAVkAAQEAYQIBAAEAUQEAIyEAKgEqAwcWKwEiNTQ2NyYnJjU0Njc2NjcWFRQGBwYGBxYWFxcOAhUUFjMyNjcXBgYHBgEoVBgLGxAQFBULGxECFxYHDQgLHRQOCRkSJhYOFgkFAQQCEAGRMw82EAwMDQkTLhIIEggGBhMWDwULCAwMBB4JIh8ICw8DAgkBBQMcAP//ACL/5ARUAOQCBgU9AAD//wAi/+QEVADkAgYFPQAA//8AM//iBLMA1gIGBUUAAP//ADP/4gSzANYCBgVFAAD//wAY/z4EGgBXAgYFRwAA//8AGP8+BBoAVwIGBUcAAP//ADb/ggdgAcICBgVoAAD//wA2/4IHYAHCAgYFaAAA//8ANv+CB2ABwgIGBWgAAAABACIAZAFNAhYAFwATQBANAQBKFwEASQAAAHYWAQcXKyUmJicmJicjJiY1NDY3FhYXFhYXHgIXATQWMBwVMBolFxUXDyA9HBwvEAUUEwVkM1onIDgZFycZFRoHFEAmJlMnDTY2DwAA//8ANQAAAnMCnwImAcAAAAEHBg0BVQB9AAixAQGwfbA1KwAA////3ABkAU0CkgImBh0AAAEGBZ4gHAAIsQEBsBywNSv////fAGQBTQL1AiYGHQAAAQYF9SlPAAixAQGwT7A1K///ACL/DwFNAhYCJgYdAAABBgV2YaAACbEBAbj/oLA1KwAAAP////oAZAFNAr4CJgYdAAABBgXvLTMACLEBArAzsDUr//8AEABkAU0C6gImBh0AAAEGBXkXQAAIsQEBsECwNSv//wAi/woBTQIWAiYGHQAAAQYFe3CfAAmxAQG4/5+wNSsAAAAAAgAAAWkA0QH6AAMABwAvQCwAAAQBAQIAAWcAAgMDAlcAAgIDXwUBAwIDTwQEAAAEBwQHBgUAAwADEQYHFysRNTMVBzUzFcO0wgHKMDBhMDAAAAAAAgAU/xMBoQFUAC8APwBIQEUzAQQFDgECBAUBAQIEAQABBEwAAwAFBAMFaQABBgEAAQBlBwEEBAJhAAICDwJOMTABADo4MD8xPxsZEhAIBgAvAS8IBxYrFyImJic3FjMyNjc2NjcnBgYjIiY1NDY3NjYzMhYXFhYXFhYXFhYVFAcGBgcGBgcGEzI2NyYmJyYmIyIGBhUUFrIVPzoQDjZDKkcdHSoMARlBLDpDGRMUMxoXKRIIEAcOFAQDAiAHEwsWNB8dNx0vGQgbERElEhYpGTrtFB0PIBIlHR1LJggXIkg5KEwdHiQXFAoWDRo6Hg8cDkVIESIQIDIODgE/Cg0hNxQUFiEtFCUmAAAAAAIAFP8TAeUBXAAuAD4AS0BIJiUCAgQFAQECBAEAAQNMAAMABgQDBmkAAQgBAAEAZQkHAgQEAmEFAQICDwJOLy8BAC8+Lz02NCkoIyIaGA4NCAYALgEuCgcWKxciJiYnNxYzMjY3NjY3JiYnJiY1NDY3NjYzMhYXFhYXFhYVMzIVFRQjIwYGBwYGEyYmJyYmIyIGBhUUFhcWFrIVPzoQDjZDIz4aGioPHTscPz4ZExQzGhcqEhsiBwMCPAgIRwssHR5JnQUaEhIqFRYpGTg1ETztFB0PIBIaFhc6HwIFBAo7OClLHR4kFxQfWisQHxIIPAgrVSIiKQE6J0MZGh0gLRQjJQoDAwAAAAIAP/9RAkkBhwAxAD4ANEAxPDQdCgQBAygjIgMCAQJMMQECSQAAAAMBAANpAAEBAmEAAgIPAk45NyYlIB8VEwQHFisXJiY1NDY2NzY2NyYmNTQ2Njc2NjMyFhcWFhUUBgcWFjMyFRUUIyImJwYGBwYGFRQWFxM2Ny4CIyIGBgcWFoMZKyU8IhUpExMZBQkFGVMyFSYODhEdHRw9IggIN2osIEkyLiwWDso/LQIbIgwUMS0PFDCvGUsiJjAaBQIIBBMxGQwqKQkuOiAaGUEgIkAaBQYIPAgSFg4TBwcoHRkoEQETGCgZOiohNB0ZJAAAAAH/0AIrAA0CaAALACexBmREQBwAAQAAAVkAAQEAYQIBAAEAUQEABwUACwELAwcWK7EGAEQDIiY1NDYzMhYVFAYSDRERDQ0SEgIrEg0NERENDRIAAAEATv+FAMb//wAIAAazBQABMisXJiYnNjcWFwaNEx8NLBIWJBV7Dh4NLBUZIR4AAAH/mgHCAGACMAAIAAazBAABMisDNzY2NxcUBwZmBUZgGgEXRwHCKxciCgsYCyIAAAEArv7lAPb/7QAOAAazCQABMisTJzQmJyY1NDY3FhYVFAbRCAsLBR4WDAgP/uUDL1YoDwoXIQcyPSMmNgAAAAABAB3+5gDP/4oAFQAUQBEOAQBKCgMCAEkAAAB2FQEGFysTJiYnNzcyFhcWFz4CNxYWFRQGBwZsEiQZEA8PFg8ODQQQEwcIDiQOEv7mJkgiBAMTGhkgEy8pCAEODQlCLQoAAAD//wAd/wQAz/+oAQcF9P/i/bAACbEAAbj9sLA1KwAAAP//ACD/GgZMACsCBgVAAAAAAv/qAfIAcgLsAAgAJQAxQC4lFBADAQABTAYBAkobGhkVBAFJAAEAAYYAAgAAAlkAAgIAYQAAAgBRLSMsAwYZKxMXFAcGBgc3Nhc3NCYjIgYHFjMyNwcOAgcnNyY1NDY2MzIVFAdjAQQGRSsDVgkBCAUQGQYPHhYeAh8nIRUCJyQXIREoBwLsCAgGBxsOHRx0BAQJFg8UBicHDhALJxYHIRMlFyUMDwAAA//iAfIAcgMZAAcAEAAtADdANC0cGAMBAAFMDgkFAQQCSiMiIR0EAUkAAQABhgACAAACWQACAgBhAAACAFEqKBsZFhQDBhYrExcUBwYHNzYXFxQHBgYHNzYXNzQmIyIGBxYzMjcHDgIHJzcmNTQ2NjMyFRQHWgEPLjwDPz4BBQZFKgMpNwEIBRAZBg8eFh4CHychFQInJBchESgHAxkIEQcTEBoTFwkJBQcbDR0MZAQECRYPFAYnBw4QCycWByETJRclDA8AAP////j/LADhAmkAJgIGAAABBgXsIR4ACLECAbAesDUrAAL/tgFYAJQCHQAaACMASEBFEQEDAR4HBgMEAgMCAQACA0wQDAkDAUoAAQADAgEDaQUBAgAAAlkFAQICAGEEAQACAFEcGwEAIiAbIxwjFRMAGgEaBgYWKxMiJzcWFhc3JjU0NjcWFhcHFzY2MzIWFhUUBicyNjcmJiMiBzBDNwYJFw8RGBgRAwYCCAMRHw4OHxY6KBciExAYCxwrAVggGQEEAhRWCQgUBAoRBgdADhcUGgogKiwJCA4HIwD//wBw/z4A+P/JAgYF5wAAAAIAOAFYARYCHQAaACMASEBFEQEDAR4HBgMEAgMCAQACA0wQDAkDAUoAAQADAgEDaQUBAgAAAlkFAQICAGEEAQACAFEcGwEAIiAbIxwjFRMAGgEaBgYWKxMiJzcWFhc3JjU0NjcWFhcHFzY2MzIWFhUUBicyNjcmJiMiB7JDNwYJFw8RGBgRAwYCCAMRHw4OHxY6KBciExAYCxwrAVggGQEEAhRWCQgUBAoRBgdADhcUGgogKiwJCA4HIwAAAgAa/x4A+P/jABoAIwBIQEURAQMBHgcGAwQCAwIBAAIDTBAMCQMBSgABAAMCAQNpBQECAAACWQUBAgIAYQQBAAIAURwbAQAiIBsjHCMVEwAaARoGBhYrFyInNxYWFzcmNTQ2NxYWFwcXNjYzMhYWFRQGJzI2NyYmIyIHlEM3BgkXDxEYGBEDBgIIAxEfDg4fFjooFyITEBgLHCviIBkBBAIUVgkIFAQKEQYHQA4XFBoKICosCQgOByMAAAAEABsBTADpAm0AHgAoADMAPABRQE4UAQMBJiIHBgMFAgMCAQACA0wTDwIBSjs4NjIvLAYASQABAAMCAQNpBQECAAACWQUBAgIAYQQBAAIAUSAfAQAlIx8oICcZFwAeAR4GBhYrEyInNxYWFzcmJjU0NzY2NxYWFwcXNjc2MzIWFhUUBicyNjcmIyIHFhYXJiYnNjY3FhYXBgcmJzY3FhYXBok+MAUIFA0RCgwKBQ0JAwUDCAIYDAsJDRwUNCUVHhMZFhomChVAEBcHBhcQCxUKF3wgDRYWCxYKFQG7HhYCBAITJyoEBwgFBgQKDwUGOhUGBhIXCR0mKAgHEyABAY0NFQkFGBIKFQsfGxoRFBsKFQseAAAEABv+/wDpACAAHgAoADMAPABRQE4UAQMBJiIHBgMFAgMCAQACA0wTDwIBSjs4NjIvLAYASQABAAMCAQNpBQECAAACWQUBAgIAYQQBAAIAUSAfAQAlIx8oICcZFwAeAR4GBhYrFyInNxYWFzcmJjU0NzY2NxYWFwcXNjc2MzIWFhUUBicyNjcmIyIHFhYXJiYnNjY3FhYXBgcmJzY3FhYXBok+MAUIFA0RCgwKBQ0JAwUDCAIYDAsJDRwUNCUVHhMZFhomChVAEBcHBhcQCxUKF3wgDRYWCxYKFZIeFgIEAhMnKgQFCgUGBAoPBQY6FQYGEhcJHSYoCAcTIAEBjQ0VCQUYEgoVCx8bGhEUGwoVCx4AAP//ABsBTADpAm0CBgYGAAD//wAiAGQBTQLqAiYGHQAAAQcFeQDBAEAACLEBAbBAsDUrAAD///+5AGQBTQKoAiYGHQAAAQYGAJDRAAmxAQG4/9GwNSsAAAD///+5AGQBTQKpAiYGHQAAAQYGAaPSAAmxAQG4/9KwNSsAAAD//wA1AAACcwKfAiYBwAAAAQcGDgFRAHgACLEBArB4sDUrAAAABgAA/tQKWQF7ADQAPQCPAJsAtQC+ATdLsB5QWEAxdjgCCQBkAQcJpHlnVVI8LSsGCQEHn56Zk4yEf0dGQSQLAgEETKwBAEq9u7kyAAUCSRtANDgBCwB2AQkLZAEHCaR5Z1VSPC0rBgkBB5+emZOMhH9HRkEkCwIBBUysAQBKvbu5MgAFAklZS7AeUFhAPwsBCQAHAAkHgAAHAQAHAX4AAAACYRUSFA8ODQUTBAMKAgIPTREQDAoIBgYBAQJiFRIUDw4NBRMEAwoCAg8CThtARQALAAkACwmAAAkHAAkHfgAHAQAHAX4AAAACYRUSFA8ODQUTBAMKAgIPTREQDAoIBgYBAQJiFRIUDw4NBRMEAwoCAg8CTllANZyckZA/Ppy1nLWioZeVkJuRm4uJgoF9e3Vza2ljYVlXUU9KSURDPo8/jyYlIyEeHBQSFgcWKxMmJicmJjU0Njc2Nz4DNzY2MzIXFhYXFhYXFjMyFRUUIyInByImJjU0NwYHFhcWFhcGBgEmJicOAgcWBSImJwYGIyI1NTQzMjc2Njc2MzIXBgYHFhYzMjY3PgI3NjYzMhcGBgcWFjMyNzY2NzY2NzYzMhcGBgcWFjMyFRUUIyImJwYGBwYGIyInDgIFIjU1NDMhMhUVFCMhIjU1NDMyNjcmJicmNTQ2NxYWFxYVFAYHBhcmJic2NxYXBh8ECQcFBgMCBhgCHCgpECxDFRMSChYNDBgKFxcLCzEqDEV1SBMuJgcaBQwGDyEBNBYjBRw3KwpBBtkcShQOOCIICDMeERETCRYOCg8VBQgxIBcnEAcREgYJEA8OCgcZFA8WEBodBQ4KCBIJCBcOCgcZEiMyIgkJJ00aAwYFBzYbLiIHKjP6GjAwBOE1NQKuCAguXyQHEgsTIx4CDAkPExFbLBYkDi0ZIx8Y/tQ5dT0wUiEGDwkSMwUXHBwKGhwfEDIjIjIQIAo3CysrHjUiICUfHVObJ04nChMBbiNqNg4rMBU3UxINDhEIPAgXDhsfEAcmLQgECQYECzY4DRQMBxBIQgcFCgYkHhksDw8HEEc3ExAIPAgeGAcNBwwPGQkLBAEIPAgIPAgIPAgVEBc3IDIRFi4QDjksQiQPNyA3zxEhECseJR0hAP//AAD++gkMAp8AJwCMA7oAAAAnAeQCEQAAACcAAwhaAAAAJwHBB8cAAAAnANQGJwAAAAYB6LwAAAoAAP7PDYMCnwAsADsAVgCAALUAvgDKAOwA9gEBAR1ASLkBAQLaEwIDE72urAsEBAOHQyEDAATo5OPOzcjCpX15eGI/OwMCEAUA/PTx71wtBgkF//lbAwgJB0xrTEk3NjMGAkqzgQIISUuwCVBYQEIAEwEDABNyAAMEAQMEfgAEAAEEAH4MAQIAARMCAWkACRkBCAkIZRQSEQ0KBwYAAAViGxYVGhAPDgsYBhcLBQUPBU4bQEMAEwEDARMDgAADBAEDBH4ABAABBAB+DAECAAETAgFpAAkZAQgJCGUUEhENCgcGAAAFYhsWFRoQDw4LGAYXCwUFDwVOWUBAy8vAv1hXPTwAAMvsy+zn5uHg2dfR0MbEv8rAyqempKKfnZWTfHt2dWBeV4BYgEJBPFY9VgAsACsjFiYpFRwHGyshIjU1NDMyNjc2NjcmJyYmIyIGByY1NDYzMhYXFhYXFjMHJiYjIgYGBw4DISYCJyY2NxYWFwcWFgYHBSI1NTQzMjcmJicmJic2NjcWFhcWFhcUBgcGBSImJic3FhYzMjY3JiYnJiY1NDY3FhYXFhYXFhYXFjMyFRUUIyInDgIFJiYnJiY1NDY3Njc+Azc2NjMyFxYWFxYWFxYzMhUVFCMiJwciJiY1NDcGBxYXFhYXBgYBJiYnDgIHFgUiNTU0MyEyFRUUIyEiNTU0MzI2NzY2NzYzMhcOAgcWFjMyFRUUIyInBw4CFyYnNjcWFhcGBgcmJzY2NxYWFwYGCC8ICEpzKkNBGCcmKDMaHyolDkI2IEErNUklJkgnECkXFyk3LSU/RVgEzAMNCwQfJwQTDBsEAwMF/tYJCTwzAwYDCRILDiETAwgGBQQBFA8S/mwQMTIQDRwvFEtrHAcPCSIgHhwECQcQFQUGBwEiLAgILyMLPlz02gQJBwUGAwIGGAIcKCkQLEMVExIKFg0MGAoXFwsLMSoMRXVIEy4mBxoFDAYPIQE0FiMFHDcrCkECYzAwBOE1NfmyCAg+VhILHh0JFgwNBhUSBCI6IggIXUEQBz9V5ioXFioFHhsKHqotFg4gEwQfGgkeCDwIDQ0XFQQTHyAYEBcQFyYuGyAmJwgJSQIDCBgXExkOBrABI3MpKQcTQSMXSpy0bwgIPAgbJkokXrZYDxwNNo1YSX81ETkaI8wSGgwhBQZaVRAfED1JGRofDQ0fESk9ExMfDRQIPAgYRWc4ZTl1PTBSIQYPCRIzBRccHAoaHB8QMiMiMhAgCjcLKyseNSIgJR8dU5snTicKEwFuI2o2DiswFTdTCDwICDwICDwICw0HMTcRCA8vKAQVEQg8CDoVBxIMtiEcFS4EHhoOIiMkGQwiFgUeGQ8iAAAAAAAADwC6AAMAAQQJAAAAaAAAAAMAAQQJAAEAIgBoAAMAAQQJAAIADgCKAAMAAQQJAAMARACYAAMAAQQJAAQAMgDcAAMAAQQJAAUAoAEOAAMAAQQJAAYALgGuAAMAAQQJAAcARAHcAAMAAQQJAAgAKgIgAAMAAQQJAAkAZgJKAAMAAQQJAAoAhAKwAAMAAQQJAAsAPgM0AAMAAQQJAAwAPANyAAMAAQQJAA0BIAOuAAMAAQQJAA4ANATOAEMAbwBwAHkAcgBpAGcAaAB0ACAAMgAwADEAOQAtADIAMAAyADAAIABHAG8AbwBnAGwAZQAgAEwATABDAC4AIABBAGwAbAAgAFIAaQBnAGgAdABzACAAUgBlAHMAZQByAHYAZQBkAC4ATgBvAHQAbwAgAE4AYQBzAGsAaAAgAEEAcgBhAGIAaQBjAFIAZQBnAHUAbABhAHIAMgAuADAAMAA0ADsARwBPAE8ARwA7AE4AbwB0AG8ATgBhAHMAawBoAEEAcgBhAGIAaQBjAC0AUgBlAGcAdQBsAGEAcgBOAG8AdABvACAATgBhAHMAawBoACAAQQByAGEAYgBpAGMAIABSAGUAZwB1AGwAYQByADIALgAwADAANAA7ACAAdAB0AGYAYQB1AHQAbwBoAGkAbgB0ACAAKAB2ADEALgA4AC4AMwApACAALQBsACAAOAAgAC0AcgAgADUAMAAgAC0ARwAgADIAMAAwACAALQB4ACAAMQA0ACAALQBEACAAYQByAGEAYgAgAC0AZgAgAG4AbwBuAGUAIAAtAGEAIABxAHMAcQAgAC0AWAAgACIAIgBOAG8AdABvAE4AYQBzAGsAaABBAHIAYQBiAGkAYwAtAFIAZQBnAHUAbABhAHIATgBvAHQAbwAgAGkAcwAgAGEAIAB0AHIAYQBkAGUAbQBhAHIAawAgAG8AZgAgAEcAbwBvAGcAbABlACAATABMAEMALgBNAG8AbgBvAHQAeQBwAGUAIABJAG0AYQBnAGkAbgBnACAASQBuAGMALgBNAG8AbgBvAHQAeQBwAGUAIABEAGUAcwBpAGcAbgAgAFQAZQBhAG0ALAAgAEQAYQB2AGkAZAAgAFcAaQBsAGwAaQBhAG0AcwAsACAATQBvAGgAYQBtAGEAZAAgAEQAYQBrAGEAawBEAGUAcwBpAGcAbgBlAGQAIABiAHkAIABNAG8AbgBvAHQAeQBwAGUAIABEAGUAcwBpAGcAbgAgAFQAZQBhAG0ALAAgAEQAYQB2AGkAZAAgAFcAaQBsAGwAaQBhAG0AcwAgAGEAbgBkACAATQBvAGgAYQBtAGEAZAAgAEQAYQBrAGEAawBoAHQAdABwADoALwAvAHcAdwB3AC4AZwBvAG8AZwBsAGUALgBjAG8AbQAvAGcAZQB0AC8AbgBvAHQAbwAvAGgAdAB0AHAAOgAvAC8AdwB3AHcALgBtAG8AbgBvAHQAeQBwAGUALgBjAG8AbQAvAHMAdAB1AGQAaQBvAFQAaABpAHMAIABGAG8AbgB0ACAAUwBvAGYAdAB3AGEAcgBlACAAaQBzACAAbABpAGMAZQBuAHMAZQBkACAAdQBuAGQAZQByACAAdABoAGUAIABTAEkATAAgAE8AcABlAG4AIABGAG8AbgB0ACAATABpAGMAZQBuAHMAZQAsACAAVgBlAHIAcwBpAG8AbgAgADEALgAxAC4AIABUAGgAaQBzACAAbABpAGMAZQBuAHMAZQAgAGkAcwAgAGEAdgBhAGkAbABhAGIAbABlACAAdwBpAHQAaAAgAGEAIABGAEEAUQAgAGEAdAA6ACAAaAB0AHQAcAA6AC8ALwBzAGMAcgBpAHAAdABzAC4AcwBpAGwALgBvAHIAZwAvAE8ARgBMAGgAdAB0AHAAOgAvAC8AcwBjAHIAaQBwAHQAcwAuAHMAaQBsAC4AbwByAGcALwBPAEYATAACAAAAAAAA/7UAMgAAAAAAAAAAAAAAAAAAAAAAAAAABkIAAAECAQMBBAEFAQYBBwEIAQkBCgELAQwBDQEOAQ8BEAERARIBEwEUARUBFgEXARgBGQEaARsBHAEdAR4BHwEgASEBIgEjASQBJQEmAScBKAEpASoBKwEsAS0BLgEvATABMQEyATMBNAE1ATYBNwE4ATkBOgE7ATwBPQE+AT8BQAFBAUIBQwFEAUUBRgFHAUgBSQFKAUsBTAFNAU4BTwFQAVEBUgFTAVQBVQFWAVcBWAFZAVoBWwFcAV0BXgFfAWABYQFiAWMBZAFlAWYBZwFoAWkBagFrAWwBbQFuAW8BcAFxAXIBcwF0AXUBdgF3AXgBeQF6AXsBfAF9AX4BfwGAAYEBggGDAYQBhQGGAYcBiAGJAYoBiwGMAY0BjgGPAZABkQGSAZMBlAGVAZYBlwGYAZkBmgGbAZwBnQGeAZ8BoAGhAaIBowGkAaUBpgGnAagBqQGqAasBrAGtAa4BrwGwAbEBsgGzAbQBtQG2AbcBuAG5AboBuwG8Ab0BvgG/AcABwQHCAcMBxAHFAcYBxwHIAckBygHLAcwBzQHOAc8B0AHRAdIB0wHUAdUB1gHXAdgB2QHaAdsB3AHdAd4B3wHgAeEB4gHjAeQB5QHmAecB6AHpAeoB6wHsAe0B7gHvAfAB8QHyAfMB9AH1AfYB9wH4AfkB+gH7AfwB/QH+Af8CAAIBAgICAwIEAgUCBgIHAggCCQIKAgsCDAINAg4CDwIQAhECEgITAhQCFQIWAhcCGAIZAhoCGwIcAh0CHgIfAiACIQIiAiMCJAIlAiYCJwIoAikCKgIrAiwCLQIuAi8CMAIxAjICMwI0AjUCNgI3AjgCOQI6AjsCPAI9Aj4CPwJAAkECQgJDAkQCRQJGAkcCSAJJAkoCSwJMAk0CTgJPAlACUQJSAlMCVAJVAlYCVwJYAlkCWgJbAlwCXQJeAl8CYAJhAmICYwJkAmUCZgJnAmgCaQJqAmsCbAJtAm4CbwJwAnECcgJzAnQCdQJ2AncCeAJ5AnoCewJ8An0CfgJ/AoACgQKCAoMChAKFAoYChwKIAokCigKLAowCjQKOAo8CkAKRApICkwKUApUClgKXApgCmQKaApsCnAKdAp4CnwKgAqECogKjAqQCpQKmAqcCqAKpAqoCqwKsAq0CrgKvArACsQKyArMCtAK1ArYCtwK4ArkCugK7ArwCvQK+Ar8CwALBAsICwwLEAsUCxgLHAsgCyQLKAssCzALNAs4CzwLQAtEC0gLTAtQC1QLWAtcC2ALZAtoC2wLcAt0C3gLfAuAC4QLiAuMC5ALlAuYC5wLoAukC6gLrAuwC7QLuAu8C8ALxAvIC8wL0AvUC9gL3AvgC+QL6AvsC/AL9Av4C/wMAAwEDAgMDAwQDBQMGAwcDCAMJAwoDCwMMAw0DDgMPAxADEQMSAxMDFAMVAxYDFwMYAxkDGgMbAxwDHQMeAx8DIAMhAyIDIwMkAyUDJgMnAygDKQMqAysDLAMtAy4DLwMwAzEDMgMzAzQDNQM2AzcDOAM5AzoDOwM8Az0DPgM/A0ADQQNCA0MDRANFA0YDRwNIA0kDSgNLA0wDTQNOA08DUANRA1IDUwNUA1UDVgNXA1gDWQNaA1sDXANdA14DXwNgA2EDYgNjA2QDZQNmA2cDaANpA2oDawNsA20DbgNvA3ADcQNyA3MDdAN1A3YDdwN4A3kDegN7A3wDfQN+A38DgAOBA4IDgwOEA4UDhgOHA4gDiQOKA4sDjAONA44DjwOQA5EDkgOTA5QDlQOWA5cDmAOZA5oDmwOcA50DngOfA6ADoQOiA6MDpAOlA6YDpwOoA6kDqgOrA6wDrQOuA68DsAOxA7IDswO0A7UDtgO3A7gDuQO6A7sDvAO9A74DvwPAA8EDwgPDA8QDxQPGA8cDyAPJA8oDywPMA80DzgPPA9AD0QPSA9MD1APVA9YD1wPYA9kD2gPbA9wD3QPeA98D4APhA+ID4wPkA+UD5gPnA+gD6QPqA+sD7APtA+4D7wPwA/ED8gPzA/QD9QP2A/cD+AP5A/oD+wP8A/0D/gP/BAAEAQQCBAMEBAQFBAYEBwQIBAkECgQLBAwEDQQOBA8EEAQRBBIEEwQUBBUEFgQXBBgEGQQaBBsEHAQdBB4EHwQgBCEEIgQjBCQEJQQmBCcEKAQpBCoEKwQsBC0ELgQvBDAEMQQyBDMENAQ1BDYENwQ4BDkEOgQ7BDwEPQQ+BD8EQARBBEIEQwREBEUERgRHBEgESQRKBEsETARNBE4ETwRQBFEEUgRTBFQEVQRWBFcEWARZBFoEWwRcBF0EXgRfBGAEYQRiBGMEZARlBGYEZwRoBGkEagRrBGwEbQRuBG8EcARxBHIEcwR0BHUEdgR3BHgEeQR6BHsEfAR9BH4EfwSABIEEggSDBIQEhQSGBIcEiASJBIoEiwSMBI0EjgSPBJAEkQSSBJMElASVBJYElwSYBJkEmgSbBJwEnQSeBJ8EoAShBKIEowSkBKUEpgSnBKgEqQSqBKsErAStBK4ErwSwBLEEsgSzBLQEtQS2BLcEuAS5BLoEuwS8BL0EvgS/BMAEwQTCBMMExATFBMYExwTIBMkEygTLBMwEzQTOBM8E0ATRBNIE0wTUBNUE1gTXBNgE2QTaBNsE3ATdBN4E3wTgBOEE4gTjBOQE5QTmBOcE6ATpBOoE6wTsBO0E7gTvBPAE8QTyBPME9AT1BPYE9wT4BPkE+gT7BPwE/QT+BP8FAAUBBQIFAwUEBQUFBgUHBQgFCQUKBQsFDAUNBQ4FDwUQBREFEgUTBRQFFQUWBRcFGAUZBRoFGwUcBR0FHgUfBSAFIQUiBSMFJAUlBSYFJwUoBSkFKgUrBSwFLQUuBS8FMAUxBTIFMwU0BTUFNgU3BTgFOQU6BTsFPAU9BT4FPwVABUEFQgVDBUQFRQVGBUcFSAVJBUoFSwVMBU0FTgVPBVAFUQVSBVMFVAVVBVYFVwVYBVkFWgVbBVwFXQVeBV8FYAVhBWIFYwVkBWUFZgVnBWgFaQVqBWsFbAVtBW4FbwVwBXEFcgVzBXQFdQV2BXcFeAV5BXoFewV8BX0FfgV/BYAFgQWCBYMFhAWFBYYFhwWIBYkFigWLBYwFjQWOBY8FkAWRBZIFkwWUBZUFlgWXBZgFmQWaBZsFnAWdBZ4FnwWgBaEFogWjBaQFpQWmBacFqAWpBaoFqwWsBa0FrgWvBbAFsQWyBbMFtAW1BbYFtwW4BbkFugW7BbwFvQW+Bb8FwAXBBcIFwwXEBcUFxgXHBcgFyQXKBcsFzAXNBc4FzwXQBdEF0gXTBdQF1QXWBdcF2AXZBdoF2wXcBd0F3gXfBeAF4QXiBeMF5AXlBeYF5wXoBekF6gXrBewF7QXuBe8F8AXxBfIF8wX0BfUF9gX3BfgF+QX6BfsF/AX9Bf4F/wYABgEGAgYDBgQGBQYGBgcGCAYJBgoGCwYMBg0GDgYPBhAGEQYSBhMGFAYVBhYGFwYYBhkGGgYbBhwGHQYeBh8GIAYhBiIGIwYkBiUGJgYnBigGKQYqBisGLAYtBi4GLwYwBjEGMgYzBjQGNQY2BjcGOAY5BjoGOwY8Bj0GPgY/BkAGQQZCBkMGRAZFBkYGRwZIBkkGSgZLBkwGTQZOBk8GUAZRBlIGUwZUBlUGVgZXBlgGWQZaBlsGXAZdBl4GXwZgBmEGYgZjBmQGZQZmBmcGaAZpBmoGawZsBm0GbgZvBnAGcQZyBnMGdAZ1BnYGdwZ4BnkGegZ7BnwGfQZ+Bn8GgAaBBoIGgwaEBoUGhgaHBogGiQaKBosGjAaNBo4GjwaQBpEGkgaTBpQGlQaWBpcGmAaZBpoGmwacBp0GngafBqAGoQaiBqMGpAalBqYGpwaoBqkGqgarBqwGrQauBq8GsAaxBrIGswa0BrUGtga3BrgGuQa6BrsGvAa9Br4GvwbABsEGwgbDBsQGxQbGBscGyAbJBsoGywbMBs0GzgbPBtAG0QbSBtMG1AbVBtYG1wbYBtkG2gbbBtwG3QbeBt8G4AbhBuIG4wbkBuUG5gbnBugG6QbqBusG7AbtBu4G7wbwBvEG8gbzBvQG9Qb2BvcG+Ab5BvoG+wb8Bv0G/gb/BwAHAQcCBwMHBAcFBwYHBwcIBwkHCgcLBwwHDQcOBw8HEAcRBxIHEwcUBxUHFgcXBxgHGQcaBxsHHAcdBx4HHwcgByEHIgcjByQHJQcmBycHKAcpByoHKwcsBy0HLgcvBzAHMQcyBzMHNAc1BzYHNwc4BzkHOgc7BzwHPQc+Bz8HQAdBB0IHdW5pMDYyMQd1bmkwNjc0B3VuaTA2MjcHdW5pRkU4RQt1bmlGRThFLjAwMQx1bmlGRThFLnJsaWcHdW5pMDYyMwd1bmlGRTg0DHVuaUZFODQucmxpZwd1bmkwNjI1B3VuaUZFODgMdW5pRkU4OC5ybGlnB3VuaTA2NzIMdW5pMDY3Mi5maW5hEXVuaTA2NzIuZmluYS5ybGlnB3VuaTA2NzMMdW5pMDY3My5maW5hEXVuaTA2NzMuZmluYS5ybGlnB3VuaTA2MjIHdW5pRkU4Mgx1bmlGRTgyLnJsaWcHdW5pMDY3MQx1bmkwNjcxLnJsaWcHdW5pRkI1MQd1bmlGRDNEDHVuaUZEM0QucmxpZwd1bmlGRDNDB3VuaTA2NzUMdW5pMDY3NS5maW5hEXVuaTA2NzUuZmluYS5ybGlnB3VuaTA2NkUMdW5pMDY2RS5maW5hDHVuaTA2NkUubWVkaQx1bmkwNjZFLmluaXQHdW5pMDYyOAd1bmlGRTkwB3VuaUZFOTIHdW5pRkU5MQd1bmkwNjdFB3VuaUZCNTcHdW5pRkI1OQd1bmlGQjU4B3VuaTA2N0IHdW5pRkI1Mwd1bmlGQjU1B3VuaUZCNTQHdW5pMDY4MAd1bmlGQjVCB3VuaUZCNUQHdW5pRkI1Qwd1bmlGQkU5B3VuaUZCRTgHdW5pMDYyQQd1bmlGRTk2B3VuaUZFOTgHdW5pRkU5Nwd1bmkwNjdDDHVuaTA2N0MuZmluYQx1bmkwNjdDLm1lZGkMdW5pMDY3Qy5pbml0B3VuaTA2N0QMdW5pMDY3RC5maW5hDHVuaTA2N0QubWVkaQx1bmkwNjdELmluaXQHdW5pMDYyQgd1bmlGRTlBB3VuaUZFOUMHdW5pRkU5Qgd1bmkwNjc5B3VuaUZCNjcHdW5pRkI2OQd1bmlGQjY4B3VuaTA2N0EHdW5pRkI1Rgd1bmlGQjYxB3VuaUZCNjAHdW5pMDY3Rgd1bmlGQjYzB3VuaUZCNjUHdW5pRkI2NAd1bmkwNzUwDHVuaTA3NTAuZmluYQx1bmkwNzUwLm1lZGkMdW5pMDc1MC5pbml0B3VuaTA3NTEMdW5pMDc1MS5maW5hDHVuaTA3NTEubWVkaQx1bmkwNzUxLmluaXQHdW5pMDc1Mgx1bmkwNzUyLmZpbmEMdW5pMDc1Mi5tZWRpDHVuaTA3NTIuaW5pdAd1bmkwNzUzDHVuaTA3NTMuZmluYQx1bmkwNzUzLm1lZGkMdW5pMDc1My5pbml0B3VuaTA3NTQMdW5pMDc1NC5maW5hDHVuaTA3NTQubWVkaQx1bmkwNzU0LmluaXQHdW5pMDc1NQx1bmkwNzU1LmZpbmEMdW5pMDc1NS5tZWRpDHVuaTA3NTUuaW5pdAd1bmkwNzU2DHVuaTA3NTYuZmluYQx1bmkwNzU2Lm1lZGkMdW5pMDc1Ni5pbml0B3VuaTA4QTAMdW5pMDhBMC5maW5hDHVuaTA4QTAubWVkaQx1bmkwOEEwLmluaXQHdW5pMDYyQwd1bmlGRTlFB3VuaUZFQTAHdW5pRkU5Rgd1bmkwNjg2B3VuaUZCN0IHdW5pRkI3RAd1bmlGQjdDB3VuaTA2ODcHdW5pRkI3Rgd1bmlGQjgxB3VuaUZCODAHdW5pMDZCRgx1bmkwNkJGLmZpbmEMdW5pMDZCRi5tZWRpDHVuaTA2QkYuaW5pdAd1bmkwNjgzB3VuaUZCNzcHdW5pRkI3OQd1bmlGQjc4B3VuaTA2ODQHdW5pRkI3Mwd1bmlGQjc1B3VuaUZCNzQHdW5pMDYyRAd1bmlGRUEyB3VuaUZFQTQHdW5pRkVBMwd1bmkwNjgxDHVuaTA2ODEuZmluYQx1bmkwNjgxLm1lZGkMdW5pMDY4MS5pbml0B3VuaTA2ODIMdW5pMDY4Mi5maW5hDHVuaTA2ODIubWVkaQx1bmkwNjgyLmludGkHdW5pMDc1Nwx1bmkwNzU3LmZpbmEMdW5pMDc1Ny5tZWRpDHVuaTA3NTcuaW5pdAd1bmkwNjg1DHVuaTA2ODUuZmluYQx1bmkwNjg1Lm1lZGkMdW5pMDY4NS5pbml0B3VuaTA3NTgMdW5pMDc1OC5maW5hDHVuaTA3NTgubWVkaQx1bmkwNzU4LmluaXQHdW5pMDYyRQd1bmlGRUE2B3VuaUZFQTgHdW5pRkVBNwd1bmkwNzZFDHVuaTA3NkUuZmluYQx1bmkwNzZFLm1lZGkMdW5pMDc2RS5pbml0B3VuaTA3NkYMdW5pMDc2Ri5maW5hDHVuaTA3NkYubWVkaQx1bmkwNzZGLmluaXQHdW5pMDc3Qwx1bmkwNzdDLmZpbmEMdW5pMDc3Qy5tZWRpDHVuaTA3N0MuaW5pdAd1bmkwOEEyDHVuaTA4QTIuZmluYQx1bmkwOEEyLm1lZGkMdW5pMDhBMi5pbml0B3VuaTA2MkYHdW5pRkVBQQd1bmkwNjMwB3VuaUZFQUMHdW5pRkM1Qgx1bmlGQzVCLmZpbmEHdW5pMDY4OAd1bmlGQjg5B3VuaTA2OEMHdW5pRkI4NQd1bmkwNjhEB3VuaUZCODMHdW5pMDY4OQx1bmkwNjg5LmZpbmEHdW5pMDY4QQx1bmkwNjhBLmZpbmEHdW5pMDY4Qgx1bmkwNjhCLmZpbmEHdW5pMDY4RQd1bmlGQjg3B3VuaTA2OEYMdW5pMDY4Ri5maW5hB3VuaTA2OTAMdW5pMDY5MC5maW5hB3VuaTA2RUUMdW5pMDZFRS5maW5hB3VuaTA3NTkMdW5pMDc1OS5maW5hB3VuaTA3NUEMdW5pMDc1QS5maW5hB3VuaTA2MzEHdW5pRkVBRQd1bmkwNjMyB3VuaUZFQjAHdW5pMDY5MQd1bmlGQjhEB3VuaTA2OTIMdW5pMDY5Mi5maW5hB3VuaTA2OTMMdW5pMDY5My5maW5hB3VuaTA2OTQMdW5pMDY5NC5maW5hB3VuaTA2OTUMdW5pMDY5NS5maW5hB3VuaTA2OTYMdW5pMDY5Ni5maW5hB3VuaTA2OTcMdW5pMDY5Ny5maW5hB3VuaTA2OTgHdW5pRkI4Qgd1bmkwNjk5DHVuaTA2OTkuZmluYQd1bmkwNkVGDHVuaTA2RUYuZmluYQd1bmkwNzVCDHVuaTA3NUIuZmluYQd1bmkwNzZCDHVuaTA3NkIuZmluYQd1bmkwNzZDDHVuaTA3NkMuZmluYQd1bmkwNzcxDHVuaTA3NzEuZmluYQd1bmlGQzVDDHVuaUZDNUMuZmluYQd1bmkwOEFBDHVuaTA4QUEuZmluYQd1bmkwNjMzB3VuaUZFQjIHdW5pRkVCNAd1bmlGRUIzB3VuaTA2OUEMdW5pMDY5QS5maW5hDHVuaTA2OUEubWVkaQx1bmkwNjlBLmluaXQHdW5pMDY5Qgx1bmkwNjlCLmZpbmEMdW5pMDY5Qi5tZWRpDHVuaTA2OUIuaW5pdAd1bmkwNjM0B3VuaUZFQjYHdW5pRkVCOAd1bmlGRUI3B3VuaTA2RkEMdW5pMDZGQS5maW5hDHVuaTA2RkEubWVkaQx1bmkwNkZBLmluaXQHdW5pMDY5Qwx1bmkwNjlDLmZpbmEMdW5pMDY5Qy5tZWRpDHVuaTA2OUMuaW5pdAd1bmkwNjM1B3VuaUZFQkEHdW5pRkVCQwd1bmlGRUJCB3VuaTA2OUQMdW5pMDY5RC5maW5hDHVuaTA2OUQubWVkaQx1bmkwNjlELmluaXQHdW5pMDY5RQx1bmkwNjlFLmZpbmEMdW5pMDY5RS5tZWRpDHVuaTA2OUUuaW5pdAd1bmkwNjM2B3VuaUZFQkUHdW5pRkVDMAd1bmlGRUJGB3VuaTA2RkIMdW5pMDZGQi5maW5hDHVuaTA2RkIubWVkaQx1bmkwNkZCLmluaXQHdW5pMDYzNwd1bmlGRUMyB3VuaUZFQzQHdW5pRkVDMwd1bmkwNjlGDHVuaTA2OUYuZmluYQx1bmkwNjlGLm1lZGkMdW5pMDY5Ri5pbml0B3VuaTA2MzgHdW5pRkVDNgd1bmlGRUM4B3VuaUZFQzcHdW5pMDhBMwx1bmkwOEEzLmZpbmEMdW5pMDhBMy5tZWRpDHVuaTA4QTMuaW5pdAd1bmkwNjM5B3VuaUZFQ0EHdW5pRkVDQwd1bmlGRUNCB3VuaTA2QTAMdW5pMDZBMC5maW5hDHVuaTA2QTAubWVkaQx1bmkwNkEwLmluaXQHdW5pMDYzQQd1bmlGRUNFB3VuaUZFRDAHdW5pRkVDRgd1bmkwNkZDDHVuaTA2RkMuZmluYQx1bmkwNkZDLm1lZGkMdW5pMDZGQy5pbml0B3VuaTA2NDEHdW5pRkVEMgd1bmlGRUQ0B3VuaUZFRDMHdW5pMDZBNAd1bmlGQjZCB3VuaUZCNkQHdW5pRkI2Qwd1bmkwNkExDHVuaTA2QTEuZmluYQx1bmkwNkExLm1lZGkMdW5pMDZBMS5pbml0B3VuaTA2QTIMdW5pMDZBMi5maW5hDHVuaTA2QTIubWVkaQx1bmkwNkEyLmluaXQHdW5pMDZBMwx1bmkwNkEzLmZpbmEMdW5pMDZBMy5tZWRpDHVuaTA2QTMuaW5pdAd1bmkwNkE1DHVuaTA2QTUuZmluYQx1bmkwNkE1Lm1lZGkMdW5pMDZBNS5pbml0B3VuaTA4QTQMdW5pMDhBNC5maW5hDHVuaTA4QTQubWVkaQx1bmkwOEE0LmluaXQHdW5pMDZBNgd1bmlGQjZGB3VuaUZCNzEHdW5pRkI3MAd1bmkwNjZGDHVuaTA2NkYubWVkaQx1bmkwNjZGLmluaXQMdW5pMDY2Ri5maW5hB3VuaTA2NDIHdW5pRkVENgd1bmlGRUQ4B3VuaUZFRDcHdW5pMDZBNwx1bmkwNkE3LmZpbmEHdW5pMDZBOAx1bmkwNkE4LmZpbmEHdW5pMDhBNQx1bmkwOEE1LmZpbmEMdW5pMDhBNS5tZWRpDHVuaTA4QTUuaW5pdAd1bmkwNjQzB3VuaUZFREEHdW5pRkVEQwd1bmlGRURCB3VuaTA2QTkHdW5pRkI4Rgd1bmlGQjkxB3VuaUZCOTAHdW5pMDYzQgx1bmkwNjNCLmZpbmEMdW5pMDYzQi5tZWRpDHVuaTA2M0IuaW5pdAd1bmkwNjNDDHVuaTA2M0MuZmluYQx1bmkwNjNDLm1lZGkMdW5pMDYzQy5pbml0B3VuaTA3NjIMdW5pMDc2Mi5maW5hDHVuaTA3NjIubWVkaQx1bmkwNzYyLmluaXQHdW5pMDc2Mwx1bmkwNzYzLmZpbmEMdW5pMDc2My5tZWRpDHVuaTA3NjMuaW5pdAd1bmkwNzY0DHVuaTA3NjQuZmluYQx1bmkwNzY0Lm1lZGkMdW5pMDc2NC5pbml0B3VuaTA2QUYHdW5pRkI5Mwd1bmlGQjk1B3VuaUZCOTQHdW5pMDZCMAx1bmkwNkIwLmZpbmEMdW5pMDZCMC5tZWRpDHVuaTA2QjAuaW5pdAd1bmkwNkI0DHVuaTA2QjQuZmluYQx1bmkwNkI0Lm1lZGkMdW5pMDZCNC5pbml0B3VuaTA2QUEMdW5pMDZBQS5maW5hDHVuaTA2QUEubWVkaQx1bmkwNkFBLmluaXQHdW5pMDZBQgx1bmkwNkFCLmZpbmEMdW5pMDZBQi5tZWRpDHVuaTA2QUIuaW5pdAd1bmkwNkFDDHVuaTA2QUMuZmluYQx1bmkwNkFDLm1lZGkMdW5pMDZBQy5pbml0B3VuaTA3N0YMdW5pMDc3Ri5maW5hDHVuaTA3N0YubWVkaQx1bmkwNzdGLmluaXQHdW5pMDZBRQx1bmkwNkFFLmZpbmEHdW5pMDZCMgx1bmkwNkIyLmZpbmEMdW5pMDZCMi5tZWRpDHVuaTA2QjIuaW5pdAd1bmkwNkFEB3VuaUZCRDQHdW5pRkJENgd1bmlGQkQ1B3VuaTA2QjEHdW5pRkI5Qgd1bmlGQjlEB3VuaUZCOUMHdW5pMDZCMwd1bmlGQjk3B3VuaUZCOTkHdW5pRkI5OAd1bmkwNjQ0B3VuaUZFREUHdW5pRkVFMAx1bmlGRUUwLnJsaWcHdW5pRkVERgx1bmlGRURGLnJsaWcHdW5pMDZCNQx1bmkwNkI1Lmlzb2wMdW5pMDZCNS5maW5hDHVuaTA2QjUubWVkaQx1bmkwNkI1LmluaXQMdW5pMDZCNS5ybGlnB3VuaTA2QjYMdW5pMDZCNi5pc29sDHVuaTA2QjYuZmluYQx1bmkwNkI2Lm1lZGkRdW5pMDZCNi5tZWRpLnJsaWcMdW5pMDZCNi5pbml0DHVuaTA2QjYucmxpZwd1bmkwNkI3DHVuaTA2QjcuaXNvbAx1bmkwNkI3LmZpbmEMdW5pMDZCNy5tZWRpEXVuaTA2QjcubWVkaS5ybGlnDHVuaTA2QjcuaW5pdAx1bmkwNkI3LnJsaWcHdW5pMDZCOAx1bmkwNkI4LmZpbmEMdW5pMDZCOC5tZWRpEXVuaTA2QjgubWVkaS5ybGlnDHVuaTA2QjguaW5pdAx1bmkwNkI4LnJsaWcHdW5pMDhBNgx1bmkwOEE2LmZpbmEMdW5pMDhBNi5tZWRpDHVuaTA4QTYuaW5pdAx1bmkwOEE2LnJsaWcHdW5pMDY0NQd1bmlGRUUyB3VuaUZFRTQHdW5pRkVFMwd1bmkwOEE3B3VuaTA2NDYHdW5pRkVFNgd1bmlGRUU4B3VuaUZFRTcHdW5pMDZCOQx1bmkwNkI5LmZpbmEMdW5pMDZCOS5tZWRpDHVuaTA2QjkuaW5pdAd1bmkwNkJBDHVuaTA2QkEubWVkaQx1bmkwNkJBLmluaXQHdW5pRkI5Rgd1bmkwNkJCB3VuaUZCQTEHdW5pMDZCQwx1bmkwNkJDLmZpbmEMdW5pMDZCQy5tZWRpDHVuaTA2QkMuaW5pdAd1bmkwNkJEDHVuaTA2QkQuZmluYQx1bmkwNkJELm1lZGkMdW5pMDZCRC5pbml0B3VuaTA2NDcHdW5pRkVFQQd1bmlGRUVDB3VuaUZFRUIHdW5pMDZDMAd1bmlGQkE1B3VuaTA2QzEHdW5pRkJBNwd1bmlGQkE5B3VuaUZCQTgHdW5pMDZDMgx1bmkwNkMyLmZpbmEMdW5pMDZDMi5tZWRpDHVuaTA2QzIuaW5pdAd1bmkwNkJFB3VuaUZCQUIHdW5pRkJBRAd1bmlGQkFDB3VuaTA2RkYMdW5pMDZGRi5maW5hDHVuaTA2RkYubWVkaQx1bmkwNkZGLmluaXQHdW5pRkNEOQd1bmkwNjI5B3VuaUZFOTQHdW5pMDZDMwx1bmkwNkMzLmZpbmEHdW5pMDY0OAd1bmlGRUVFB3VuaTA2QzQMdW5pMDZDNC5maW5hB3VuaTA2Q0EMdW5pMDZDQS5maW5hB3VuaTA2Q0YMdW5pMDZDRi5maW5hB3VuaTA2MjQHdW5pRkU4Ngd1bmkwNzc4DHVuaTA3NzguZmluYQd1bmkwNzc5DHVuaTA3NzkuZmluYQd1bmkwOEFCDHVuaTA4QUIuZmluYQd1bmkwNkM1B3VuaUZCRTEHdW5pMDZDNgd1bmlGQkRBB3VuaTA2QzcHdW5pRkJEOAd1bmkwNkM4B3VuaUZCREMHdW5pMDZDOQd1bmlGQkUzB3VuaTA2Q0IHdW5pRkJERgd1bmkwNjQ5B3VuaUZFRjAHdW5pRkM1RAd1bmlGQzkwB3VuaTA3NzMMdW5pMDc3My5maW5hEXVuaTA3NzMuZmluYS5ybGlnB3VuaTA3NzQMdW5pMDc3NC5maW5hEXVuaTA3NzQuZmluYS5ybGlnB3VuaTA2NEEHdW5pRkVGMgd1bmlGRUY0B3VuaUZFRjMHdW5pMDYyNgd1bmlGRThBB3VuaUZFOEMHdW5pRkU4Qgd1bmkwNkNFDHVuaTA2Q0UuZmluYQx1bmkwNkNFLm1lZGkMdW5pMDZDRS5pbml0B3VuaTA4QTgMdW5pMDhBOC5maW5hDHVuaTA4QTgubWVkaQx1bmkwOEE4LmluaXQHdW5pMDhBOQd1bmkwNjNEDHVuaTA2M0QuZmluYQx1bmkwNjNELm1lZGkMdW5pMDYzRC5pbml0B3VuaTA2Q0MHdW5pRkJGRAd1bmlGQkZGB3VuaUZCRkUHdW5pMDYzRQx1bmkwNjNFLmZpbmEMdW5pMDYzRS5tZWRpDHVuaTA2M0UuaW5pdAd1bmkwNjNGDHVuaTA2M0YuZmluYQx1bmkwNjNGLm1lZGkMdW5pMDYzRi5pbml0B3VuaTA3NzUMdW5pMDc3NS5maW5hDHVuaTA3NzUubWVkaQx1bmkwNzc1LmluaXQHdW5pMDc3Ngx1bmkwNzc2LmZpbmEMdW5pMDc3Ni5tZWRpDHVuaTA3NzYuaW5pdAd1bmkwNzc3DHVuaTA3NzcuZmluYQx1bmkwNzc3Lm1lZGkMdW5pMDc3Ny5pbml0B3VuaTA2Q0QMdW5pMDZDRC5maW5hB3VuaTA2MjAMdW5pMDYyMC5maW5hDHVuaTA2MjAubWVkaQx1bmkwNjIwLmluaXQHdW5pMDZEMAd1bmlGQkU1B3VuaUZCRTcHdW5pRkJFNgd1bmkwNkQxDHVuaTA2RDEuZmluYQd1bmkwNkQyB3VuaUZCQUYHdW5pMDZEMwd1bmlGQkIxB3VuaTA3N0EMdW5pMDc3QS5maW5hDHVuaTA3N0EubWVkaQx1bmkwNzdBLmluaXQHdW5pMDc3Qgx1bmkwNzdCLmZpbmEMdW5pMDc3Qi5tZWRpDHVuaTA3N0IuaW5pdAd1bmkwOEFDB3VuaTA2NzYMdW5pMDY3Ni5maW5hB3VuaTA2NzcMdW5pMDY3Ny5maW5hB3VuaTA2NzgMdW5pMDY3OC5maW5hDHVuaTA2NzgubWVkaQx1bmkwNjc4LmluaXQHdW5pMDZENQd1bmkwNzVDDHVuaTA3NUMuZmluYQx1bmkwNzVDLm1lZGkMdW5pMDc1Qy5pbml0B3VuaTA3NUQMdW5pMDc1RC5maW5hDHVuaTA3NUQubWVkaQx1bmkwNzVELmluaXQHdW5pMDc1RQx1bmkwNzVFLmZpbmEMdW5pMDc1RS5tZWRpDHVuaTA3NUUuaW5pdAd1bmkwNzVGDHVuaTA3NUYuZmluYQx1bmkwNzVGLm1lZGkMdW5pMDc1Ri5pbml0B3VuaTA3NjAMdW5pMDc2MC5maW5hDHVuaTA3NjAubWVkaQx1bmkwNzYwLmluaXQHdW5pMDc2MQx1bmkwNzYxLmZpbmEMdW5pMDc2MS5tZWRpDHVuaTA3NjEuaW5pdAd1bmkwNzY1DHVuaTA3NjUuZmluYQx1bmkwNzY1Lm1lZGkMdW5pMDc2NS5pbml0B3VuaTA3NjYMdW5pMDc2Ni5maW5hDHVuaTA3NjYubWVkaQx1bmkwNzY2LmluaXQHdW5pMDc2Nwx1bmkwNzY3LmZpbmEMdW5pMDc2Ny5tZWRpDHVuaTA3NjcuaW5pdAd1bmkwNzY4DHVuaTA3NjguZmluYQx1bmkwNzY4Lm1lZGkMdW5pMDc2OC5pbml0B3VuaTA3NjkMdW5pMDc2OS5maW5hDHVuaTA3NjkubWVkaQx1bmkwNzY5LmluaXQHdW5pMDc2QQx1bmkwNzZBLmZpbmEMdW5pMDc2QS5tZWRpDHVuaTA3NkEuaW5pdAx1bmkwNzZBLnJsaWcHdW5pMDc2RAx1bmkwNzZELmZpbmEMdW5pMDc2RC5tZWRpDHVuaTA3NkQuaW5pdAd1bmkwNzcwDHVuaTA3NzAuZmluYQx1bmkwNzcwLm1lZGkMdW5pMDc3MC5pbml0B3VuaTA3NzIMdW5pMDc3Mi5maW5hDHVuaTA3NzIubWVkaQx1bmkwNzcyLmluaXQHdW5pMDc3RAx1bmkwNzdELmZpbmEMdW5pMDc3RC5tZWRpDHVuaTA3N0QuaW5pdAd1bmkwNzdFDHVuaTA3N0UuZmluYQx1bmkwNzdFLm1lZGkMdW5pMDc3RS5pbml0B3VuaUZERjEHdW5pRkU3MQd1bmlGRTczB3VuaUZFN0IHdW5pMDY0MAd1bmlGRTc3B3VuaUZFNzkHdW5pRkU3RAd1bmlGQ0YyB3VuaUZDRjMHdW5pRkNGNAd1bmlGRTdGB3VuaUZFRkIHdW5pRkVGQw91bmlGRUZDX3VuaTA2NzUPdW5pRkVGQ191bmkwNzczD3VuaUZFRkNfdW5pMDc3NA91bmlGRUZDX3VuaUZCQjIPdW5pRkVGQl91bmlGQkI2D3VuaUZFRkNfdW5pRkJCNg91bmlGRUZCX3VuaUZCQjcPdW5pRkVGQ191bmlGQkI3B3VuaUZFRjcHdW5pRkVGOAd1bmlGRUY5B3VuaUZFRkEHdW5pRkVGNQd1bmlGRUY2D3VuaUZFREZfdW5pRkI1MQ91bmlGRUUwX3VuaUZCNTEHdW5pRkMwNQd1bmlGQzlDB3VuaUZDMDYHdW5pRkM5RAd1bmlGREMyB3VuaUZDMDcHdW5pRkM5RQd1bmlGRDlFB3VuaUZDNkEHdW5pRkM2Qgd1bmlGQzA4B3VuaUZDNkMHdW5pRkNFMQd1bmlGQzlGB3VuaUZDNkQHdW5pRkNFMgd1bmlGQ0EwB3VuaUZDMDkHdW5pRkM2RQd1bmlGQzBBB3VuaUZDNkYHdW5pRkMwQgd1bmlGQ0ExB3VuaUZENTAHdW5pRkRBMAd1bmlGRDlGB3VuaUZDMEMHdW5pRkNBMgd1bmlGRDUxB3VuaUZENTIHdW5pRkQ1Mwd1bmlGQzBEB3VuaUZDQTMHdW5pRkQ1NAd1bmlGREEyB3VuaUZEQTEHdW5pRkM3MAd1bmlGQzcxB3VuaUZDMEUHdW5pRkM3Mgd1bmlGQ0UzB3VuaUZDQTQHdW5pRkQ1NQd1bmlGRDU2B3VuaUZENTcHdW5pRkRBNAd1bmlGREEzB3VuaUZDNzMHdW5pRkNFNAd1bmlGQ0E1B3VuaUZDMEYHdW5pRkM3NAd1bmlGQzEwB3VuaUZDNzUHdW5pRkMxMQd1bmlGQzc2B3VuaUZDNzcHdW5pRkMxMgd1bmlGQzc4B3VuaUZDRTUHdW5pRkNBNgd1bmlGQzc5B3VuaUZDRTYHdW5pRkMxMwd1bmlGQzdBB3VuaUZDMTQHdW5pRkM3Qgd1bmlGQzE1B3VuaUZDQTcHdW5pRkRBNgd1bmlGREJFB3VuaUZDMTYHdW5pRkNBOAd1bmlGRDU4B3VuaUZENTkHdW5pRkRBNwd1bmlGREE1B3VuaUZEMDEHdW5pRkQxRAd1bmlGRDAyB3VuaUZEMUUHdW5pRkMxNwd1bmlGQ0E5B3VuaUZEQkYHdW5pRkMxOAd1bmlGQ0FBB3VuaUZENUIHdW5pRkQ1QQd1bmlGQ0ZGB3VuaUZEMUIHdW5pRkQwMAd1bmlGRDFDD3VuaUZFQTRfdW5pRkJCRBd1bmlGRUE0X3VuaUZCQjRfdW5pRkJCMwd1bmlGQzE5B3VuaUZDQUIHdW5pRkMxQQd1bmlGQzFCB3VuaUZDQUMHdW5pRkQwMwd1bmlGRDFGB3VuaUZEMDQHdW5pRkQyMAd1bmlGQzFDB3VuaUZEMzQHdW5pRkNBRAd1bmlGRDVEB3VuaUZENUUHdW5pRkMxRAd1bmlGRDM1B3VuaUZDQUUHdW5pRkQ1Qwd1bmlGQzFFB3VuaUZEMzYHdW5pRkNBRgd1bmlGREE4B3VuaUZEQzYHdW5pRkQwRQd1bmlGRDJBB3VuaUZDMUYHdW5pRkNFNwd1bmlGQ0IwB3VuaUZENjEHdW5pRkQ1Rgd1bmlGRDYwB3VuaUZENjIHdW5pRkQ2Mwd1bmlGQ0U4B3VuaUZEMzEHdW5pRkNGQgd1bmlGRDE3B3VuaUZDRkMHdW5pRkQxOBd1bmlGRUIzX3VuaUZCQjZfdW5pRkJCNwd1bmlGRDA5B3VuaUZEMjUHdW5pRkQzNwd1bmlGRDJEB3VuaUZENjkHdW5pRkQwQQd1bmlGRDI2B3VuaUZEMzgHdW5pRkQyRQd1bmlGRDY3B3VuaUZENjgHdW5pRkRBQQd1bmlGRDBCB3VuaUZEMjcHdW5pRkQzOQd1bmlGRDJGB3VuaUZEMEQHdW5pRkQyOQd1bmlGRDBDB3VuaUZEMjgHdW5pRkNFOQd1bmlGRDMwB3VuaUZENkEHdW5pRkQ2Qgd1bmlGRDZDB3VuaUZENkQHdW5pRkNFQQd1bmlGRDMyB3VuaUZDRkQHdW5pRkQxOQd1bmlGQ0ZFB3VuaUZEMUEHdW5pRkMyMAd1bmlGQ0IxB3VuaUZENjQHdW5pRkQ2NQd1bmlGREE5B3VuaUZDQjIHdW5pRkQwRgd1bmlGRDJCB3VuaUZERjAHdW5pRkMyMQd1bmlGQ0IzB3VuaUZENjYHdW5pRkRDNQd1bmlGRDA1B3VuaUZEMjEHdW5pRkQwNgd1bmlGRDIyB3VuaUZDMjIHdW5pRkNCNAd1bmlGQzIzB3VuaUZDQjUHdW5pRkQ2RQd1bmlGREFCB3VuaUZDMjQHdW5pRkNCNgd1bmlGRDZGB3VuaUZENzAHdW5pRkQxMAd1bmlGRDJDB3VuaUZDMjUHdW5pRkNCNwd1bmlGRDA3B3VuaUZEMjMHdW5pRkQwOAd1bmlGRDI0B3VuaUZDMjYHdW5pRkNCOAd1bmlGQzI3B3VuaUZEM0EHdW5pRkQzMwd1bmlGRDcxB3VuaUZENzIHdW5pRkQ3Mwd1bmlGRDc0B3VuaUZDRjUHdW5pRkQxMQd1bmlGQ0Y2B3VuaUZEMTIPdW5pRkVDNF91bmlGQkI0B3VuaUZDMjgHdW5pRkQzQgd1bmlGQ0I5B3VuaUZDMjkHdW5pRkNCQQd1bmlGRDc1B3VuaUZEQzQHdW5pRkMyQQd1bmlGQ0JCB3VuaUZENzYHdW5pRkQ3Nwd1bmlGRDc4B3VuaUZEQjYHdW5pRkNGNwd1bmlGRDEzB3VuaUZDRjgHdW5pRkQxNAd1bmlGQzJCB3VuaUZDQkMHdW5pRkMyQwd1bmlGQ0JEB3VuaUZENzkHdW5pRkQ3Qgd1bmlGRDdBB3VuaUZDRjkHdW5pRkQxNQd1bmlGQ0ZBB3VuaUZEMTYHdW5pRkMyRAd1bmlGQ0JFB3VuaUZDMkUHdW5pRkNCRgd1bmlGQzJGB3VuaUZDQzAHdW5pRkQ3Qwd1bmlGRDdEB3VuaUZDMzAHdW5pRkNDMQd1bmlGREMxB3VuaUZDMzEHdW5pRkM3Qwd1bmlGQzMyB3VuaUZDN0QHdW5pRkMzMwd1bmlGQ0MyB3VuaUZDMzQHdW5pRkNDMwd1bmlGRDdFB3VuaUZEQjQHdW5pRkQ3Rgd1bmlGREIyB3VuaUZDMzUHdW5pRkM3RQd1bmlGQzM2B3VuaUZDN0YHdW5pRkMzNwd1bmlGQzgwB3VuaUZDMzgHdW5pRkNDNAd1bmlGQzM5B3VuaUZDQzUHdW5pRkMzQQd1bmlGQ0M2B3VuaUZDM0IHdW5pRkM4MQd1bmlGQ0VCB3VuaUZDQzcHdW5pRkMzQwd1bmlGQzgyB3VuaUZDRUMHdW5pRkNDOAd1bmlGREJCB3VuaUZEQzMHdW5pRkRCNwd1bmlGQzNEB3VuaUZDODMHdW5pRkMzRQd1bmlGQzg0B3VuaUZDM0YHdW5pRkNDOQd1bmlGRDg0B3VuaUZEODMHdW5pRkRCQwd1bmlGREJBB3VuaUZEQUMHdW5pRkM0MAd1bmlGQ0NBB3VuaUZEODAHdW5pRkRCNQd1bmlGRDgyB3VuaUZEODEHdW5pRkM0MQd1bmlGQ0NCB3VuaUZEODUHdW5pRkQ4Nhd1bmlGRURGX3VuaUZFRTBfdW5pRkVFQQd1bmlGQzQyB3VuaUZDODUHdW5pRkNFRAd1bmlGQ0NDB3VuaUZEODcHdW5pRkQ4OAd1bmlGREFEB3VuaUZDQ0QHdW5pRkM0Mwd1bmlGQzg2B3VuaUZDNDQHdW5pRkM4Nw91bmkwNkI1X3VuaTA2NzUPdW5pMDZCNV91bmkwNzczD3VuaTA2QjVfdW5pMDc3NA91bmkwNkI2X3VuaTA2MjcPdW5pMDZCNl91bmkwNjc1D3VuaTA2QjZfdW5pMDc3Mw91bmkwNkI2X3VuaTA3NzQPdW5pMDZCN191bmkwNjc1D3VuaTA2QjdfdW5pMDc3Mw91bmkwNkI3X3VuaTA3NzQPdW5pMDZCOF91bmkwNjc1D3VuaTA2QjhfdW5pMDc3Mw91bmkwNkI4X3VuaTA3NzQHdW5pRkM4OAd1bmlGQzQ1B3VuaUZDQ0UHdW5pRkQ4Qwd1bmlGRDkyB3VuaUZEOEQHdW5pRkRDMAd1bmlGQzQ2B3VuaUZDQ0YHdW5pRkQ4OQd1bmlGRDhBB3VuaUZEOEIHdW5pRkM0Nwd1bmlGQ0QwB3VuaUZEOEUHdW5pRkQ4Rgd1bmlGREI5B3VuaUZDNDgHdW5pRkM4OQd1bmlGQ0QxB3VuaUZEQjEHdW5pRkM0OQd1bmlGQzRBB3VuaUZDNEIHdW5pRkNEMgd1bmlGREJEB3VuaUZEQjgHdW5pRkQ5Nwd1bmlGRDk4B3VuaUZEOTkHdW5pRkRDNwd1bmlGQzRDB3VuaUZDRDMHdW5pRkQ5NQd1bmlGRDk2B3VuaUZEQjMHdW5pRkM0RAd1bmlGQ0Q0B3VuaUZDOEEHdW5pRkM4Qgd1bmlGQzRFB3VuaUZDOEMHdW5pRkNFRQd1bmlGQ0Q1B3VuaUZEOUIHdW5pRkQ5QQd1bmlGQzhEB3VuaUZDRUYHdW5pRkNENgd1bmlGQzRGB3VuaUZDOEUHdW5pRkM1MAd1bmlGQzhGB3VuaUZDNTEHdW5pRkNENwd1bmlGQzUyB3VuaUZDRDgHdW5pRkQ5Mwd1bmlGRDk0B3VuaUZDNTMHdW5pRkM1NAd1bmlGQzU1B3VuaUZDREEHdW5pRkRBRgd1bmlGQzU2B3VuaUZDREIHdW5pRkRBRQd1bmlGQzU3B3VuaUZDREMHdW5pRkM5MQd1bmlGQzkyB3VuaUZDNTgHdW5pRkM5Mwd1bmlGQ0YwB3VuaUZDREQHdW5pRkQ5Qwd1bmlGRDlEB3VuaUZEQjAHdW5pRkM5NAd1bmlGQ0YxB3VuaUZDREUHdW5pRkM1OQd1bmlGQzk1B3VuaUZDNUEHdW5pRkM5Ngd1bmlGQkVBB3VuaUZCRUIHdW5pRkMwMAd1bmlGQzk3B3VuaUZDMDEHdW5pRkM5OAd1bmlGQzk5B3VuaUZDNjQHdW5pRkM2NQd1bmlGQzAyB3VuaUZDNjYHdW5pRkNERgd1bmlGQzlBB3VuaUZDNjcHdW5pRkNFMAd1bmlGQzlCB3VuaUZCRUUHdW5pRkJFRgd1bmlGQkYyB3VuaUZCRjMHdW5pRkJGMAd1bmlGQkYxB3VuaUZCRjQHdW5pRkJGNQd1bmlGQzAzB3VuaUZDNjgHdW5pRkMwNAd1bmlGQzY5B3VuaUZCRjYHdW5pRkJGNwd1bmlGQkY4B3VuaUZCRUMHdW5pRkJFRAd1bmlGQkY5B3VuaUZCRkEHdW5pRkJGQgd1bmlGREYyC3VuaUZERjIuMDAxB3VuaUZERjMHdW5pRkRGNAd1bmlGREY1B3VuaUZERjYHdW5pRkRGNwd1bmlGREY4B3VuaUZERjkHdW5pRkRGQQd1bmlGREZCD3VuaUZFRkJfdW5pMDY1Rg91bmlGRUZDX3VuaTA2NUYPdW5pRkVGQl91bmkwNjU0D3VuaUZFRkJfdW5pMDY1QQ91bmlGRUZDX3VuaTA2NUEUdW5pRkVFMF91bmkwNjVBLnJsaWcHdW5pMDZFNQd1bmkwNkU2B3VuaTAwMzAHdW5pMDAzMQd1bmkwMDMyB3VuaTAwMzMHdW5pMDAzNAd1bmkwMDM1B3VuaTAwMzYHdW5pMDAzNwd1bmkwMDM4B3VuaTAwMzkMdW5pMDAzMC5jY21wDHVuaTAwMzEuY2NtcAx1bmkwMDMyLmNjbXAMdW5pMDAzMy5jY21wDHVuaTAwMzQuY2NtcAx1bmkwMDM1LmNjbXAMdW5pMDAzNi5jY21wDHVuaTAwMzcuY2NtcAx1bmkwMDM4LmNjbXAMdW5pMDAzOS5jY21wB3VuaTA2NkIJdW5pMDY2Qi4xB3VuaTA2NkMMdW5pMDY2Qi5sb2NsB3VuaTA2NjAJdW5pMDY2MC40B3VuaTA2NjEJdW5pMDY2MS41B3VuaTA2NjIHdW5pMDY2Mwl1bmkwNjYzLjYHdW5pMDY2NAd1bmkwNjY1CXVuaTA2NjUuNwd1bmkwNjY2B3VuaTA2NjcJdW5pMDY2Ny44B3VuaTA2NjgJdW5pMDY2OC45B3VuaTA2NjkHdW5pMDZGMAd1bmkwNkYxB3VuaTA2RjIHdW5pMDZGMwd1bmkwNkY0CnVuaTA2RjQuMTAHdW5pMDZGNQp1bmkwNkY1LjExB3VuaTA2RjYKdW5pMDZGNi4xMgd1bmkwNkY3B3VuaTA2RjgHdW5pMDZGOQx1bmkwNjYwLmNjbXAMdW5pMDY2MS5jY21wDHVuaTA2NjIuY2NtcAx1bmkwNjYzLmNjbXAMdW5pMDY2NC5jY21wDHVuaTA2NjUuY2NtcAx1bmkwNjY2LmNjbXAMdW5pMDY2Ny5jY21wDHVuaTA2NjguY2NtcAx1bmkwNjY5LmNjbXAMdW5pMDZGNC5jY21wDHVuaTA2RjUuY2NtcAx1bmkwNkY2LmNjbXAMdW5pMDZGNC5sb2NsDHVuaTA2RjcubG9jbAx1bmkwNjYyLnRudW0MdW5pMDY2NC50bnVtDHVuaTA2NjYudG51bQx1bmkwNjY5LnRudW0HdW5pMDAyRQd1bmkwMDJDB3VuaTAwM0EHdW5pMDAyMQd1bmkyMDRGB3VuaTJFNDEMdW5pMDAyRS5sb2NsDHVuaTAwMjEubG9jbAd1bmkyMDEwB3VuaTIwMTEHdW5pMDBBQgd1bmkwMEJCDHVuaTAwQUIubG9jbAx1bmkwMEJCLmxvY2wHdW5pRkJDMQd1bmlGQkIyC3VuaUZCQjIuMDAxB3VuaUZCQjMHdW5pRkJCRAd1bmlGQkJFC3VuaUZCQkUuMDAxB3VuaUZCQjQHdW5pRkJCNQt1bmlGQkI1LjAwMQd1bmkwNjFFB3VuaUZCQjgHdW5pRkJCOQd1bmlGQkI2C3VuaUZCQjYuMDAxB3VuaUZCQjcHdW5pRkJCQQd1bmlGQkJCB3VuaUZCQkMLdW5pRkJCQi4wMDELdW5pRkJCQi4wMDIHdW5pRkJCRgd1bmkwNkRFB3VuaTA2RDQHdW5pMDYwQwd1bmkwNjFCB3VuaTA2MUYHdW5pMDYwMAx1bmkwNjAwLmNjbXAHdW5pMDY2RAd1bmkwNjAxCnVuaTA2MDEuMTMKdW5pMDYwMS4xNAp1bmkwNjAxLjE1CnVuaTA2MDEuMTYHdW5pMDYwMgx1bmkwNjAyLmNjbXAHdW5pMDYwMwx1bmkwNjAzLmNjbXAHdW5pMDYwRAd1bmkwNjBFB3VuaTA2MEYHdW5pMDZGRAd1bmkwNkREC3VuaTA2REQuMDAxCnVuaTA2REQuMTcKdW5pMDZERC4xOAp1bmkwNkRELjE5B3VuaUZEM0UHdW5pRkQzRgd1bmkwMDIwB3VuaTIwMEIHdW5pMjAwQwd1bmkyMDBEB3VuaTIwMEUHdW5pMjAwRgJDUgd1bmkwNjA4B3VuaTA2MEIHdW5pRkRGQwd1bmkyNUNDB3VuaTA2RTkHdW5pMDZGRQd1bmkwNjZBB3VuaTA2MDYHdW5pMDYwNwd1bmkwNjA5B3VuaTA2MEEHdW5pRkRGRAx1bmkwNjA0LmNjbXAHdW5pMDYwNAd1bmkwNjEyB3VuaTA2MTMHdW5pRkU3QQd1bmkwNjEwB3VuaTA2MTEHdW5pMDYxNAd1bmkwNjE1B3VuaTA2MTYHdW5pMDYxNwd1bmkwNjE4B3VuaTA2MTkHdW5pMDYxQRV0aHJlZWRvdHNkb3duY2VudGVyYXIHdW5pMDY1Rgd1bmkwNjcwB3VuaTA2NTYHdW5pMDY1NA91bmkwNjU0X3VuaTA2NTIHdW5pMDY1NQd1bmkwNjRCD3VuaTA2NEJfdW5pMDY1NA91bmkwNjRCX3VuaTA2NTEHdW5pRkU3MAd1bmkwNjRDD3VuaTA2NENfdW5pMDY1NA91bmkwNjRDX3VuaTA2NTEHdW5pRkU3Mgd1bmkwNjRED3VuaTA2NERfdW5pMDY1NQ91bmkwNjREX3VuaTA2NTEHdW5pRkU3NAd1bmkwNjRFD3VuaTA2NEVfdW5pMDY1NA91bmkwNjRFX3VuaTA2NTEHdW5pRkU3Ngd1bmkwNjRGD3VuaTA2NEZfdW5pMDY1NA91bmkwNjRGX3VuaTA2NTEHdW5pRkU3OAd1bmkwNjUwD3VuaTA2NTBfdW5pMDY1NQ91bmkwNjUwX3VuaTA2NTEHdW5pMDY1MQ91bmkwNjUxX3VuaTA2NzAHdW5pRkU3Qwd1bmlGQzVFB3VuaUZDNUYHdW5pRkM2MAd1bmlGQzYxB3VuaUZDNjIHdW5pRkM2Mwd1bmkwNjUyB3VuaUZFN0UHdW5pMDY1Mwd1bmkwNjU4B3VuaTA2NTcHdW5pMDY1OQd1bmkwNjVBB3VuaTA2NUIHdW5pMDY1Qwd1bmkwNjVEB3VuaTA2NUUHdW5pMDZENgd1bmkwNkQ3B3VuaTA2RDgHdW5pMDZFMgd1bmkwNkQ5B3VuaTA2REEHdW5pMDZEQgd1bmkwNkRDB3VuaTA2REYHdW5pMDZFMAd1bmkwNkUxB3VuaTA2RTMHdW5pMDZFNAd1bmkwNkU3B3VuaTA2RTgHdW5pMDZFQQd1bmkwNkVCB3VuaTA2RUMHdW5pMDZFRAd1bmkwOEU0B3VuaTA4RTUHdW5pMDhFNgd1bmkwOEU3B3VuaTA4RTgHdW5pMDhFOQd1bmkwOEVBB3VuaTA4RUIHdW5pMDhFQwd1bmkwOEVEB3VuaTA4RUUHdW5pMDhFRgd1bmkwOEYwB3VuaTA4RjEHdW5pMDhGMgd1bmkwOEYzB3VuaTA4RjQHdW5pMDhGNQd1bmkwOEY2B3VuaTA4RjcHdW5pMDhGOAd1bmkwOEY5B3VuaTA4RkEHdW5pMDhGQgd1bmkwOEZDB3VuaTA4RkQHdW5pMDhGRQd1bmkwMDAwB3VuaUZCQzAMX3VuaUZCQjIuMDAxCF91bmkwNjRGBl9oYW16YQdfdmFib3ZlCF91bmlGQkIyCF91bmlGQkI1CF91bmlGQkI5CF91bmlGQkJBCF91bmlGQkJECF91bmlGQkMwCF91bmlGQkI2CF91bmlGQkI4CF91bmlGQkI0CF91bmlGQkJCCF91bmlGQkJFCF91bmlGQkI3CF91bmlGQkIzBF8wODIEXzUxNQRfNTE3BF81MTgEXzUxOQRfNTIyBF81MjQEXzUyOARfNTMxBF81MzQEXzUzNQRfNTM2BF81MzcEXzUzOARfNjcwBF82NzEEXzc3OARfNzc5BF83ODAEXzc4MQRfNzgyBF83ODMEXzg2MwRfODY2BF84NjcEXzg2OARfODcxBF84NzIEXzg3MwRfODc0BF84NzUEXzg3NgRfODc3BF84ODUEXzg4NgRfOTcxBF85NzIEXzk3MwRfOTc0BF85NzUEXzk3NgRfOTc4BF85NzkEXzk4MARfOTgyC18xNDg1X18xNDg1BF85OTAEXzk5MQRfOTk0BF85OTUEXzk5OARfOTk5EV8xNDg1X18xNDg1X18xNDg1BV8xMDA4BV8xMDI0BV8xMDI1BV8xMDI2BV8xMDI3BV8xMDI4BV8xMDI5BV8xMDMxBV8xMDMyBV8xMDM0BV8xMDM1BV8xMDM2BV8xMDM3BV8xMDM4BV8xMDM5BV8xNDg1BV8xNTA5BV8xNTE3BV8xNTE4BF85ODYEXzU0MARfNTM5BF83ODYIXzUzMS4wMDEMX3VuaUZCQjMuMDAxCF81MzEuMDAyCF81MzEuMDAzCF84NzUuMDAxCF84NzUuMDAyC190d29kb3RzdGFoBV8xMDMwBV8xMDMzBV8xMDQwBV8xMDIzA2JzbQdhbHJobWFuB2FscmFoZW0AAQAB//8ADwABAAIADgAAAAAAAAFQAAIANQABAt0AAQLeBLkAAgS6BLoAAQS7BMEAAgTEBMgAAgTJBMsAAQUjBScAAQUpBSoAAQUtBS8AAQUxBTMAAQU1BTUAAQU+BT4AAQVBBUQAAQVGBUYAAQVIBUgAAQVLBUwAAQVPBVEAAQVcBVwAAQVeBV4AAQVgBWAAAQVmBWYAAgVnBWcAAQVpBWoAAwVrBWsAAQVsBXQAAwV2BX4AAwV/BX8AAQWABYIAAwWDBYMAAQWEBYYAAwWHBYcAAQWIBYoAAwWLBYsAAQWMBY4AAwWPBY8AAQWQBZQAAwWVBZUAAQWWBZwAAwWdBZ0AAQWeBdQAAwXWBdgAAQXpBekAAQXtBe8AAQXyBfMAAQX1BfkAAQX8Bf0AAQYJBgoAAQYUBiQAAQYmBicAAQYuBi8AAQYzBjMAAQY3BjcAAQY5BkEAAQABAAIAAAAMAAAAMAABABAFdAV2BXgFewWEBYUFkAWRBaQFsgW2BbkFvAXIBc8F0AACABcFaQVqAAAFbAVzAAIFdwV3AAoFeQV6AAsFfAV+AA0FgAWCABAFhgWGABMFiAWKABQFjAWOABcFkgWUABoFmgWaAB0FnAWcAB4FngWjAB8FpQWxACUFswW1ADIFtwW4ADUFuwW7ADcFvgW+ADgFwAXDADkFxQXHAD0FygXKAEAFzQXOAEEF0QXUAEMAAAABAAAACgBIAGYAA0RGTFQAMGFyYWIAFGxhdG4AMAAgAARGQVIgACBLU0ggACBTTkQgACBVUkQgACAABAAAAAD//wACAAAAAQACbWFyawAObWttawAWAAAAAgAAAAEAAAACAAIAAwAEAAozfjYCNwwABAAAAAEACAABAAwAcAAFARYCuAACABAFaQVqAAAFbAV0AAIFdgV+AAsFgAWCABQFhAWGABcFiAWKABoFjAWOAB0FkAWUACAFlgWcACUFngW5ACwFuwW8AEgFvgW+AEoFwAXDAEsFxQXIAE8FygXKAFMFzQXUAFQAAgAbAAEC3QAABLoEugLdBMkEywLeBSMFJwLhBSkFKgLmBS0FLwLoBTEFMwLrBTUFNQLuBUsFTALvBVwFXALxBV4FXgLyBWAFYALzBWsFawL0BYsFiwL1BdYF2AL2BekF6QL5Be0F7wL6BfIF8wL9BfUF+QL/BfwF/QMEBgkGCgMGBh0GJAMIBiYGJwMQBi4GLwMSBjMGMwMUBjcGNwMVBjkGQQMWAFwAADe+AAA3xAAAN8oAADfQAAA31gAAN9wAADiQAAA34gAAN+gAADfuAAE1XAACAXIAADf0AAE1YgADAXgAADf6AAI1dAAAOAAAADgGAAA4DAAAOBIAADg2AAA4GAABNWgAATV0AAA4HgAAOQgAADgkAAA4KgAAODAAADg2AAA4PAABNW4AATV0AAA4QgAAOEgAADhOAAQBfgAEAYQABAGKAAQBkAAAOFQABAGWAAA4WgAAOGAAADhmAAA4bAAAOHIAADh4AAMBnAABNXoAADh+AAA4hAAAOIoAADiQAAA4lgAAOJwAADiiAAA4qAAAOK4AADi0AAA4wAAAOLoAADjAAAE1gAAAOMYAADjMAAA40gABNYYAADjYAAA43gABNYwAADjkAAE1kgAAOOQAADjqAAA48AAAOPYAADj8AAA5AgAAOQgAADkOAAE1mAAAORQAADkaAAA5JgABNZ4AATWkAAA5IAAAOSAAADkmAAA5LAABAEYALQABAEIB5QABAHcBuwABAGcBygABAG4BvwABAG0BvgABAGoBwAABAGgBKAMfLowukgAAAAAAAB84AAAAAAAAAAAfbijOH8IAAAAAH7wvyh/OAAAAAB+8L8ofzgAAAAAvZCjmAAAAAAAAHz4ozh/CAAAAAB9EL8ofzgAAAAAfSijmAAAAAAAAH24fUB/CAAAAAB+8H1YfzgAAAAAfgB9cAAAAAAAAH2Iozh/CAAAAAB9oL8ofzgAAAAAvWCjmAAAAAAAAH24fdB/CAAAAAB+8H3ofzgAAAAAfgB+GAAAAAAAAH4wozh/CAAAAAB+SL8ofzgAAAAAfmCjmAAAAAAAAH54ozh/CAAAAAB+kKOYAAAAAAAAfqi/KH84AAAAAH8gozh/CAAAAAB+wH7YAAAAAAAAfvC/KH84AAAAAH8gozh/CAAAAAB/IL8ofzgAAAAAf1CjmAAAAAAAAITAhHiE8IUIAACFIISohVCFaAAAp4izWLNws4gAAKe4viC+OMDAAACEwIKYhPCFCAAAhSCCyIVQhWgAAKeInbCzcLOIAACnuJ3IvjjAwAAAhMB/aITwhQgAAIUgf4CFUIVoAAC8iH+YvLi80AAAvCh/sLxYvHAAAITAf8iE8IUIAACFIH/ghVCFaAAAp4h/+LNws4gAAKe4gBC+OMDAAACEwIAohPCFCAAAhSCAQIVQhWgAALyIgFi8uLzQAAC8KIBwvFi8cAAAp4izWLNws4gAAKe4viC+OMDAAACC+IR4hPCFCAAAgyiEqIVQhWgAAKV4vKC8uLzQAAClkL4gvjjAwAAAgviAiITwhQgAAIMogIiFUIVoAACleICgvLi80AAApZCn0L44wMAAAIC4hHiE8IUIAACA0ISohVCFaAAAgOi8oLy4vNAAAIEAvEC8WLxwAACCgIR4hPCFCAAAgrCEqIVQhWgAAKXYvKC8uLzQAACl8L4gvjjAwAAAgRiEeITwhQgAAIEwhKiFUIVoAAC5oLygvLi80AAAgUi+IL44wMAAAIFghHiE8IUIAACBeISohVCFaAAAgZC8oLy4vNAAAIGovEC8WLxwAACBwIR4hPCFCAAAgdiEqIVQhWgAAIHwvKC8uLzQAACCCLxAvFi8cAAAhMCCIITwhQgAAIUggjiFUIVoAAC+yIJQvvi/EAAAvmiCaL6YvrAAAIKAgpiE8IUIAACCsILIhVCFaAAAguCdsLNws4gAAKXwnci+OMDAAACEwIMQhPCFCAAAhSCDQIVQhWgAALyIg1i8uLzQAAC8KIOIvFi8cAAAgviDEITwhQgAAIMog0CFUIVoAACleINYvLi80AAAg3CDiLxYvHAAAIOgg7iE8IUIAACD0IPohVCFaAAAsiCyOLy4vNAAALJQsmi8WLxwAACEwIQAhPCFCAAAhSCEGIVQhWgAALyIhDC8uLzQAAC8KIRIvFi8cAAAhGCEeITwhQgAAISQhKiFUIVoAACkELygvLi80AAApCi8QLxYvHAAAITAhNiE8IUIAACFIIU4hVCFaAAAvIiFgLy4vNAAAIWYhbCFyIXgAACHwLWYAAC1sAAAh9i14AAAtfgAAIgIiIC2QLZYAACICIiAtkC2WAAAh8C1mAAAtbAAAIfYteAAALX4AACICIYQtkC2WAAAiAiGELZAtlgAAIfAtZgAALWwAACH2LXgAAC1+AAAiAiF+LZAtlgAAIgIhfi2QLZYAACHSLWYAAC1sAAAh2C14AAAtfgAAId4hhC2QLZYAACHeIYQtkC2WAAAh8C1mAAAtbAAAIfYteAAALX4AACICIYotkC2WAAAiAiGKLZAtlgAAIfAtZgAALWwAACH2LXgAAC1+AAAiAiGQLZAtlgAAIgIhkC2QLZYAACHwLWYAAC1sAAAh9i14AAAtfgAAIgItii2QLZYAACICLYotkC2WAAAhli1mAAAtbAAAIZwteAAALX4AACGiLYotkC2WAAAhoi2KLZAtlgAAIagtZgAALWwAACGuLXgAAC1+AAAhtC2KLZAtlgAAIbQtii2QLZYAACIOLWYAAC1sAAAiFC14AAAtfgAAIhotii2QLZYAACIaLYotkC2WAAAhui1mAAAtbAAAIcAteAAALX4AACHGLYotkC2WAAAhxi2KLZAtlgAAIfAtZgAALWwAACH2LXgAAC1+AAAiAiHMLZAtlgAAIgIhzC2QLZYAACHSLWYAAC1sAAAh2C14AAAtfgAAId4tii2QLZYAACHeLYotkC2WAAAh8C1mAAAtbAAAIfYteAAALX4AACICIeQtkC2WAAAiAiHkLZAtlgAAIfAtZgAALWwAACH2LXgAAC1+AAAiAiHqLZAtlgAAIgIh6i2QLZYAACHwLWYAAC1sAAAh9i14AAAtfgAAIgIh/C2QLZYAACICIggtkC2WAAAiDi1mAAAtbAAAIhQteAAALX4AACIaIiAtkC2WAAAiGiIgLZAtlgAAIsgikiLUItoAACLgIp4i7CLyAAAiPiKSItQi2gAAIiYiniLsIvIAACIsIpIi1CLaAAAiMiKeIuwi8gAAIqQvcCKwIrYAACK8Ip4i7CLyAAAiOCKSItQi2gAAIj4iniLsIvIAACLIIkQi1CLaAAAi4CJKIuwi8gAAIsgiUCLUItoAACLgIlAi7CLyAAAiyCJcItQi2gAAIuAiYiLsIvIAACJWIlwi1CLaAAAivCJiIuwi8gAAImgikiLUItoAACJuIp4i7CLyAAAidCKSItQi2gAAInoiniLsIvIAACKAIpIi1CLaAAAihiKeIuwi8gAAIowikiLUItoAACKYIp4i7CLyAAAipCKqIrAitgAAIrwiwiLsIvIAACLIIs4i1CLaAAAi4CLmIuwi8gAAIzojWCNeI2QAACM6I1gjXiNkAAAjFiNYI14jZAAAIxYjWCNeI2QAACL4I1gjXiNkAAAi+CNYI14jZAAAIv4jWCNeI2QAACL+I1gjXiNkAAAjOiMEI14jZAAAIzojBCNeI2QAACM6IxwjXiNkAAAjOiMcI14jZAAAIzojCiNeI2QAACM6IxAjXiNkAAAjFiMcI14jZAAAIxYjHCNeI2QAACMiI1gjXiNkAAAjIiNYI14jZAAAIygjWCNeI2QAACMoI1gjXiNkAAAjLiNYI14jZAAAIy4jWCNeI2QAACM0I1gjXiNkAAAjNCNYI14jZAAAIzojWCNeI2QAACM6I1gAAAAAAAAjQCNYI14jZAAAI0AjWCNeI2QAACNGI1gjXiNkAAAjRiNYI14jZAAAI0wjWCNeI2QAACNMI1gjXiNkAAAjUiNYI14jZAAAI1IjWCNeI2QAACNqI3AAAAAAAAAjaiNwAAAAAAAAI4Itri20LboAACOCLa4ttC26AAAjiC3GLcwt0gAAI4gtxi3MLdIAACN2I44ttC26AAAjdiOOLbQtugAAI3wjlC3MLdIAACN8I5QtzC3SAAAjgiOgLbQtugAAI4IjoC20LboAACOII6wtzC3SAAAjiCOsLcwt0gAAI5otri20LboAACOaLa4ttC26AAAjpi3GLcwt0gAAI6Ytxi3MLdIAACOaI44ttC26AAAjmiOOLbQtugAAI6YjlC3MLdIAACOmI5QtzC3SAAAjmiOgLbQtugAAI5ojoC20LboAACOmI6wtzC3SAAAjpiOsLcwt0gAAI7Ij1iPuI/QAACOyI9Yj7iP0AAAjviPcJAYkDAAAI74j3CQGJAwAACOyI7gj7iP0AAAjsiO4I+4j9AAAI74jxCQGJAwAACO+I8QkBiQMAAAjyiPWI+4j9AAAI8oj1iPuI/QAACPQI9wkBiQMAAAj0CPcJAYkDAAAI+Ij1iPuI/QAACPiI9Yj7iP0AAAj+iPcJAYkDAAAI/oj3CQGJAwAACPiI+gj7iP0AAAj4iPoI+4j9AAAI/okACQGJAwAACP6JAAkBiQMAAAkEiROJFQkWgAAJBIkTiRUJFoAACQYJHIkeCRmAAAkHiRyJHgkfgAAJCQkTiRUJFoAACQkJE4kVCRaAAAkKiRyJHgkZgAAJDAkciR4JH4AACQ2JE4kVCRaAAAkNiROJFQkWgAAJDwkciR4JGYAACRCJHIkeCR+AAAkSCROJFQkWgAAJEgkTiRUJFoAACRgJHIkeCRmAAAkbCRyJHgkfgAAJIQuhgAAK2IAACSKK24AACt0AAAkkCssKzIrOAAAJJYrmCueK6QAACScLoYAACtiAAAkoituAAArdAAAJKgrLCsyKzgAACSuK5grniukAAAktC6GAAArYgAAJLorbgAAK3QAACTAKywrMis4AAAkzCuYK54rpAAAJLQuhgAAK2IAACS6K24AACt0AAAkwCTGKzIrOAAAJMwk0iueK6QAACTYLvgu/i8EAAAk3iUyK9Qr2gAAJOQoJivsK/IAAC6YLp4upC6qAAAk/C74Lv4vBAAAJQglMivUK9oAACUUKCYr7CvyAAAlGi9wL3YvfAAALvIu+C7+LwQAACvIJTIr1CvaAAAr4CgmK+wr8gAAK/guni6kLqoAAC7yJQIu/i8EAAAryCUOK9Qr2gAAK+AlgCvsK/IAACv4JYwupC6qAAAk2CUCLv4vBAAAJN4lDivUK9oAACTkJYAr7CvyAAAumCWMLqQuqgAALvIk6i7+LwQAACvIJPAr1CvaAAAr4CWkK+wr8gAAK/gk9i6kLqoAACT8JQIu/i8EAAAlCCUOK9Qr2gAAJRQlgCvsK/IAACUaJSAvdi98AAAlJi74Lv4vBAAAJSwlMivUK9oAACU4KCYr7CvyAAAlPi9wL3YvfAAAJUQs+i0AJWgAACvgKCYr7CvyAAAr+C6eLqQuqgAAJUQs+i0AJXQAACViLPotACVoAAAlbiz6LQAldAAAJXooJivsK/IAACWGLp4upC6qAAAlSiz6LQAlaAAAJVAs+i0AJXQAACVWLPotACVoAAAlXCz6LQAldAAAJWInZi0AJWgAACVuJ2YtACV0AAAleiWAK+wr8gAAJYYljC6kLqoAACZMJnwmgiaIAAAmWCaUJpomoAAAJc4oJivsAAAAACXUKCYr7AAAAAAmHCayJsoAAAAAJcImsibKAAAAACXOKCYr7AAAAAAl1CgmK+wAAAAAJaomsibKAAAAACWqJrImygAAAAAlkigmK+wAAAAAJZgoJivsAAAAACYcJZ4mygAAAAAlwiWeJsoAAAAAJc4lpCvsAAAAACXUJaQr7AAAAAAlqiayJsoAAAAAJaomsibKAAAAACWwKCYr7AAAAAAlsCgmK+wAAAAAJbYmsibKAAAAACW2JrImygAAAAAlvCgmK+wAAAAAJbwoJivsAAAAACYcJcgmygAAAAAlwiXIJsoAAAAAJc4r5ivsAAAAACXUK+Yr7AAAAAAmviayJsoAAAAAJr4msibKAAAAACXaKCYr7AAAAAAl2igmK+wAAAAAJr4msibKAAAAACa+JrImygAAAAAl2igmK+wAAAAAJdooJivsAAAAACXgJrImygAAAAAl4CayJsoAAAAAJeYoJivsAAAAACXmKCYr7AAAAAAl7CXyAAAAAAAAJfgl/gAAAAAAACYEJgoAAAAAAAAmECYWAAAAAAAAJhwmsibKAAAAACYcJrImygAAAAAmIigmK+wAAAAAJiIoJivsAAAAACYoJnwmgiaIAAAmLiaUJpomoAAAJkYoJivsAAAAACZGKCYr7AAAAAAmNCZ8JoImiAAAJjomlCaaJqAAACZAKCYr7AAAAAAmRigmK+wAAAAAJkwmUiaCJogAACZYJl4mmiagAAAmviZkJsoAAAAAJr4mZCbKAAAAACbQJmor7AAAAAAm0CZwK+wAAAAAJnYmfCaCJogAACaOJpQmmiagAAAmpigmK+wAAAAAJqYoJivsAAAAACasJrImygAAAAAmrCayJsoAAAAAJrgoJivsAAAAACa4KCYr7AAAAAAmvibEJsoAAAAAJr4mxCbKAAAAACbQJtYr7AAAAAAm0CbWK+wAAAAALPQs+i0ALO4AACz0LPotAC0GAAAtDC0SLRgtHgAAMGAwZgAAAAAAAC0kLSotMC02AAAtPC1CAAAAAAAAJtws+i0ALO4AACbcLPotACzuAAAm3Cz6LQAtBgAAJuItEi0YLR4AACboLSotMC02AAAm7i1CAAAAAAAAJvQs+i0ALO4AACb0LPotACzuAAAm+iz6LQAtBgAAJwAtEi0YLR4AACcGMGYAAAAAAAAnDC0qLTAtNgAAJxItQgAAAAAAACcYLPotACzuAAAnGCz6LQAs7gAAJx4s+i0ALQYAACckLRItGC0eAAAnKjBmAAAAAAAAJzAtKi0wLTYAACc2LUIAAAAAAAAs9Cc8LQAs7gAALPQnPC0ALQYAAC0MJ0ItGC0eAAAwYCdIAAAAAAAAJ04nVAAAAAAAAC08J1oAAAAAAAAs9Cz6LQAs7gAALPQs+i0ALQYAAC0MLRItGC0eAAAtJC0qLTAtNgAALTwtQgAAAAAAAC68LsIuyC7OAAAsNCwQLEAsRgAALEwsHCxYLF4AACxkLCgscCx2AAAnYC7CLsguzgAALHws+i0ALMoAACx8LPotACzKAAAniizWLNws4gAAJ5AviC+OMDAAACx8J2YtACzKAAAsfCdmLQAsygAAJ4onbCzcLOIAACeQJ3IvjjAwAAAneCz6LQAsygAAKeIs1izcLOIAACnuL4gvjjAwAAAneCz6LQAsygAAJ34s+i0ALMoAACd+LPotACzKAAAsfCeELQAsygAALHwnhC0ALMoAACeKKegs3CziAAAnkCn0L44wMAAAJ5Ys+i0ALMoAACeWLPotACzKAAAp4iecLNws4gAAKe4noi+OMDAAACr8KwIAACsIAAAyuCe6AAAnwAAAJ6gnrgAAAAAAACgCKCYAACgsAAAn2CsCAAArCAAAJ7QnugAAJ8AAACr8KwIAACsIAAAnxihiAAAoaAAAJ8wn6gAAJ/AAACfSJ/YAADAwAAAn2CsCAAArCAAAJ94oYgAAKGgAACfkJ+oAACfwAAAo8if2AAAwMAAAJ/woDgAAKBQAACf8KA4AACgUAAAoAigmAAAoLAAAKAIoJgAAKCwAACgIKA4AACgUAAAoCCgOAAAoFAAAKBooJgAAKCwAACgaKCYAACgsAAAoICgmAAAoLAAAKDIrAgAAKwgAACg4KD4AAChEAAAoSihQAAAoVgAAKFwoYgAAKGgAACiYKsAAADAAAAAomCrAAAAwAAAAKJgqwAAAMAAAACiYKsAAADAAAAAobirAAAAwAAAAKG4qwAAAMAAAACh0KsAAADAAAAAodCrAAAAwAAAAKHoqwAAAMAAAACh6KsAAADAAAAAogCrAAAAwAAAAKIYqwAAAMAAAACiMKsAAADAAAAAokirAAAAwAAAAL/owDAAAMAAAADAGMAwAADASAAAomCrAAAAwAAAAKJgqwAAAMAAAACieKsAAADAAAAAonirAAAAwAAAAKKQqwAAAMAAAACikKsAAADAAAAAoqirAAAAwAAAAKKoqwAAAMAAAACiwKsAAADAAAAAosCrAAAAwAAAAKLYqwAAAMAAAACi2KsAAADAAAAAqxirMKtIq6gAAKhgq3irkKuoAACi8Kswq0irqAAAowireKuQq6gAAKNQozgAAAAAAACjUKNoAAAAAAAAoyCjmAAAAAAAAKNQozgAAAAAAACjUKNoAAAAAAAAo4CjmAAAAAAAAKsYpLirSKuoAACoYKRYq5CrqAAAvIiyOLy4vNAAALwosmi8WLxwAACkQKswq0irqAAApECreKuQq6gAAKOws1izcLOIAACjyL4gvjjAwAAAo+CrMKtIq6gAAKP4q3irkKuoAACkELI4vLi80AAApCiyaLxYvHAAAKRApLirSKuoAACkQKRYq5CrqAAApHCyOLy4vNAAAKSIsmi8WLxwAACkoKS4q0irqAAApNCrMKtIq6gAAKToq3irkKuoAAClAKUYvLi80AAApTCyaLxYvHAAAKsYqzCrSKuoAACoYKt4q5CrqAAAvIiyOLy4vNAAALwosmi8WLxwAAClSKswq0irqAAApWCreKuQq6gAAKV4sji8uLzQAAClkKqgvjjAwAAApairMKtIq6gAAKXAq3irkKuoAACl2LI4vLi80AAApfCqoL44wMAAAKYIqzCrSKuoAACmIKt4q5CrqAAAqbCyOLy4vNAAAKnIqqC+OMDAAACmOKswq0irqAAAplCreKuQq6gAAKZosji8uLzQAACmgKqgvjjAwAAAqximmKtIq6gAAKhgprCrkKuoAAC8iKbIvLi80AAAp7im4L44wMAAAKb4pxAAAAAAAACnKKdAAAAAAAAAqxinWKtIq6gAAKhgp3CrkKuoAACniKegs3CziAAAp7in0L44wMAAAKsYp+irSKuoAACoYKgAq5CrqAAAvIioGLy4vNAAALwoqDC8WLxwAACrGKhIq0irqAAAqGCoeKuQq6gAAKiQqQgAAKkgAACoqKjAAACo2AAAqPCpCAAAqSAAAKk4qVAAAKloAACpgKn4AACqEAAAqZiqQAAAqlgAAKmwsji8uLzQAACpyKqgvjjAwAAAqeCp+AAAqhAAAKooqkAAAKpYAACqcLI4vLi80AAAqoiqoL44wMAAAAAAqrgAAAAAAACq0KsAAADAAAAAqtCrAAAAwAAAAKrQqwAAAMAAAACq6KsAAADAAAAAqxirMKtIq6gAAKtgq3irkKuoAACrwLNYs3CziAAAq9i+IL44wMAAAKvwrAgAAKwgAACsOLa4ttC26AAArDi2uLbQtugAAKxQtxi3MLdIAACsULcYtzC3SAAArGi6GAAArYgAAKyArbgAAK3QAACsmKywrMis4AAArPiuYK54rpAAAK0QuhgAAK2IAACtKK24AACt0AAArUCuAK4YrjAAAK1YrmCueK6QAACtcLoYAACtiAAAraCtuAAArdAAAK3orgCuGK4wAACuSK5grniukAAAu8iuqLv4vBAAAK8grsCvUK9oAACvgK7Yr7CvyAAAr+Cu8LqQuqgAALvIrwi7+LwQAACvIK84r1CvaAAAr4CvmK+wr8gAAK/gr/i6kLqoAACwELsIuyC7OAAAsCiwQLEAsRgAALBYsHCxYLF4AACwiLCgscCx2AAAuvCwuLsguzgAALDQsOixALEYAACxMLFIsWCxeAAAsZCxqLHAsdgAALHwsgi0ALMoAACx8LIItACzKAAAsiCyOLy4vNAAALJQsmi8WLxwAACygLPotACzKAAAsoCz6LQAsygAALKYsrCyyLLgAACy+L4gvji+UAAAsxCz6LQAsygAALMQs+i0ALMoAACzQLNYs3CziAAAs6C+IL44vlAAALPQs+i0ALO4AACz0LPotAC0GAAAtDC0SLRgtHgAALSQtKi0wLTYAAC08LUIAAAAAAAAtSC2uLbQtugAALUgtri20LboAAC1OLcYtzC3SAAAtTi3GLcwt0gAALVQtri20LboAAC1ULa4ttC26AAAtWi3GLcwt0gAALVotxi3MLdIAAC1gLWYAAC1sAAAtci14AAAtfgAALYQtii2QLZYAAC2ELYotkC2WAAAtnC2uLbQtugAALZwtri20LboAAC2iLcYtzC3SAAAtoi3GLcwt0gAALagtri20LboAAC2oLa4ttC26AAAtwC3GLcwt0gAALcAtxi3MLdIAAC3YLd4t5C3qAAAuCC4OAAAAAAAALfAt9gAAAAAAAC4ILfwAAAAAAAAuCC4OAAAAAAAALgIuDgAAAAAAAC4ILg4AAAAAAAAuCC4OAAAAAAAALgguDgAAAAAAAC4ILg4AAAAAAAAuCC4OAAAAAAAALgguDgAAAAAAADLWMtwuFAAAAAAuGjBmAAAAAAAALiAuJgAAAAAAAC4sLjIAAAAAAAAuOAAAAAAAAAAALj4AAAAAAAAAAAAALlYAAAAAAAAuRAAAAAAAAAAAAAAuSgAAAAAAAC5QAAAAAAAAAAAAAC5WAAAAAAAALlwAAAAAAAAAAAAALmIAAAAAAAAuaAAAAAAAAAAAAAAubgAAAAAAAC50AAAAAAAAAAAAAC56AAAAAAAAAAAuegAAAAAAAC6ALoYAAAAAAAAujC6SAAAAAAAALpguni6kLqoAAC6wLrYAAAAAAAAuvC7CLsguzgAAAAAu1AAAAAAAAC7aAAAAAAAAAAAu4AAAAAAAAAAALuYAAAAAAAAAAC7sAAAAAAAAAAAu8i74Lv4vBAAALwovEC8WLxwAAC8iLygvLi80AAAvOgAAAAAAAAAAAAAvQAAAAAAAAC9GAAAAAAAAAAAvTAAAAAAAAAAAL1IAAAAAAAAvWC9eAAAAAAAAL+4AAAAAAAAAAC9kAAAAAAAAAAAv7i9qL3Avdi98AAAvgi+IL44vlAAAL5ovoC+mL6wAAC+yL7gvvi/EAAAv7i/KAAAAAAAAMGAwZgAAAAAAAC/QMFoAAAAAAAAv1jBaAAAAAAAAL+4v3AAAAAAAAC/iMFoAAAAAAAAv6DBaAAAAAAAAL+4v9AAAAAAAAC/6MAwAADAAAAAwBjAMAAAwEgAAAAAwGAAAAAAAAAAAMB4AAAAAAAAwJDAqAAAwMAAAAAAwNgAAAAAAAAAAMDwAAAAAAAAwQgAAAAAAAAAAMEgwWgAAAAAAADBOMFoAAAAAAAAwVDBaAAAAAAAAMGAwZgAAAAAAADBsMHIweDB+AAAwhDCKMJAwlgAAMJwwojCoMK4AAAABADkC6wABAHEDjgABAG0DjgABADUDAAABAIv/HAABAK7/FgABAK3/CwABAHYDlAABAGEDmQABAHoC0AABAIP/DQABAKr/DQABAE0CXwABAJn/AQABAHsDMAABAHcDMAABAGACuAABAHoDVAABAFEC0QABAHYDVAABAEEC/AABALr/nAABAHYC0AABAH//0AABAHsC7gABAKb/0AABAMsC7AABAYX+9QABAaP+9QABAMn+9QABALL+9QABAYf+0QABAaX+0QABAK3+0QABAJn+0QABAYf+7wABAaX+7wABAMv+7wABALT+7wABAXb/iAABAIf/iAABAWMB/QABAW0B/QABAL0B8wABALkCawABAX4CKQABAYgCKQABAJwClwABAWECSAABAWsCSAABALsCPgABALcCtgABAV8CFgABAWkCFgABALkCDAABALUChAABAXr/FgABAZj/BwABAQ7/BwABAOz/BwABAWoCFQABAXv/EQABAXQCFQABAZn/EQABAJwCCwABAWMCAQABAYT+5gABAW0CAQABAaL+5gABAMj+5gABALkCbwABALH+5gABAWoCAQABAXv/GwABAXQCAQABAZn/GwABAXz/BAABAZr/BAABAMD/BAABAKn/BAABAWsCJQABAXr/nAABAXUCJQABAZj/nAABAW0B1AABAXz+8gABAXr/4QABAW0BFQABAXcB1AABAZr+8gABAZj/4QABAXcBFQABAMD+8gABAMQCBgABAKr+8gABAKj/4QABAMQBgwABAW7/BgABAWz/DAABAWL/MgABAW7+6AABAUwCeQABAUwCWwABASACYAABAUoCuQABAUoCmwABAR4CoAABAVMChgABAVMCaAABAScCbQABAWv+/QABAVMCcgABAVMCVAABAScCWQABAWf/OAABAVr+0gABAVYCTwABAVYCMQABAWH+oQABASsCQAABAaX+0wABAUwCcgABAUwCVAABASACWQABAWL/KAABAOACggABANcDcgABAN4DcgABANICggABANkCggABANL/JQABANj/JQABANf/iAABAO0CqgABANL/GwABANj/GwABANkClgABAOAClgABANICfgABANkCfgABAM4ClwABANUClwABANsCfgABAND/nAABAOICfgABANb/nAABAPACqgABAOH+2wABANT/6wABAN8BlgABAPQCqgABAOT+2wABANsCSwABANP/DgABANH/6wABANwBlgABAOICSwABANn/DgABANf/6wABAOMBlgABAN4CbQABAMsCaQABALn+zgABAHb+PQABAGz+LAABAMoCRQABAHX+XAABAMMCRQABAMoCWQABAL8CWgABAMwCQQABAMkCGgABAMECjAABAMQCNgABAMkCqAABAMcDCgABAHT+/gABAHT/LAABAM0BWQABAQsCGgABALb+/gABAswB/wABAWQB/wABAscB0gABAWYB0gABAtP/EQABAU3/EQABAswCEwABAt3+9QABAWQCEwABAVf+9QABAtECSgABAwX/GwABAYMCSgABAbf/GwABAtYCiwABAYgCiwABAwT/nAABAbb/nAABAtYCdwABAwX/EQABAwT/4QABAtkBiwABAYgCdwABAbf/EQABAbb/4QABAYsBiwABAf8CSgABAaECVAABAZcCVAABAgQCiwABAaYClQABAZwClQABAgQCdwABAaYCgQABAZwCgQABAf0CdwABAYj/nAABAYj/4QABAgcBiwABAZ8CgQABAakBlQABAZUCgQABAWD/nAABAWD/4QABAZ8BlQABAOUClwABAPsCNwABANcCNwABANsCNwABAOIC4gABAPgCggABANQCggABANgCggABAOICzgABAPgCbgABANQCbgABAN//EQABANgCbgABARH/EQABAmgC0QABAngCbwABAMkCZgABApP/CQABAn/+9QABAO7+9QABAmgC5QABAon/JQABAngCgwABAnX/EQABAMkCegABALoCmAABANT/EQABAl0C5gABAm0ChAABAVL/nAABAL4CewABAK8CmQABAZcCAwABAZUCRAABAZUCTgABAZUCWAABAZUCYgABAY4CRAABAZgBWAABAY4CTgABAZgBYgABAMICZgABAMb/EQABALMCwAABAOT/EQABAEYCnAABAFACnAABAc3+9QABAND+9QABAWgCnAABADwCnAABAW0C2AABAEsC2AABAY0CQAABAcz+5gABAG4CQQABAG0CSgABAFMCwQABAW0DQQABAEsDQQABAYMCLgABAwb/nAABAYUCNwABAxj/zAABAQgCNwABAqj/zAABAQMCNwABApH/zAABAWMCPwABAHMCWgABAOQCjQABAWQCjQABAOYCMAABAWYCOgABAEYCsAABAFACsAABAOcBtgABARP+9QABAWcBwAABAXn+9QABAcP/GwABALv/HQABALr/HQABAOYCbAABAQj/nAABAQj/4QABAOcBoQABAWYCdgABAW7/nAABAW7/4QABAWcBoQABAEsC4gABAWgDCgABAcL/nAABAFADCgABAWsCrQABAc/+0QABAcL/4QABAF0CwQABANL+0QABAaUDlgABAFYDlgABAEwDlgABAW0DnwABAa8DiQABAZsDiQABAFcDiQABAccDcQABAGEDiQABAXcDhQABAa8DnQABAZsDnQABAFcDnQABAcMDmQABAGEDnQABAXcDjwABARf+AwABAIn+9QABAbT++QABAM4C0AABAJr++QABAYL++QABAOoCaQABAQ3+HwABAKH/EQABAI3/EQABARMBYQABAScByAABAQP+gAABAJwB9wABAIgCbwABARMBtAABAKv+9QABAJf+9QABAOUCYAABAOn/AwABAOMClAABAPv/nAABAO0BoQABAOsBrAABAMwBWAABAIECQgABALYClAABAOEB4AABAMIBrgABAL3+9AABAMwAuwABAI7++wABASgCkAABAK8CkAABAScCuQABASv/nAABASgB0QABAK4CuQABAK0DgAABAMb/nAABAK8B0QABALYCjQABAPICjQABAQr/nAABAPwBoQABAMUCjQABAM7/nAABAM8BoQABAOEB2QABAOr/nAABAOsA7QABAPYCRgABAP0CRgABAPYCTQABAPQCEgABAPQCHAABAP4CEgABAQgCHAABAQACGQABAP4CagABAPUCRgABAPACeQABAP8CQgABAP0CWgABAMoB2gABAUIB2gABABICwgABAH//nAABAHYCzAABAIj/nAABABsCwgABAJj/nAABAJUB/gABAIECdgABAN0BuQABAUgBcwABAMUCGwABAMECkwABAIoBPQABASf+LgABAL0B/gABALkCdgABAJEBNgABARP+LgABAMkBuQABASgBcwABAMYB8wABALP/HQABAMICawABAIoBNgABAMkBLQABAL0B9wABAIECbwABAJEBSgABALwBTQABAMQCCwABAIgCgwABALgBxwABAPoBPAABANIBxwABAQ4BSwABAIUB7wABAG4CSQABAeH+3QABAgn+3QABAK/+nQABAHP+nQABAXIBOgABAYD+ygABAcsA6gABAaL+yQABAS/+eAABAUP+eAABAJ8BygABAF//iAABAJUCOAABAIL/kQABAer/DQABAgj++QABAMv+0QABALT+0QABAgT/DQABAUQA6gABAhL++QABAPMByQABAOQA7gABAU7/AwABAOMAawABAOgCOQABAaz/nAABAPIBRgABAMoBXgABAT//AwABANQAawABAKgCQAABAIUBcwABAHsB5QABAFsCSQABAL0CSAABAbr/nAABAQABRgABAJkBcwABAT3/AwABANIAawABAJkB5QABAF4CSQABAI3/GwABAWX/FAABAQICRAABAQICTwABAMH/BAABAMwA6gABARL+xQABARL+9AABATwBeAABASb+xQABASb+9AABAJQASgABAQcCDwABAP0ChgABAMACYAABAL//nAABAMABoQABAsECFAABAVkCFAABANsCzgABAPECbgABAM0CbgABAN7/nAABAN7/4QABANcBggABANECbgABANsCygABAPECagABANUCagABANECagABANkDFQABAOUB4gABAO8CtQABAR3+qgABAPsBggABAM4CwgABAOb/nAABAOb/4QABAN8BggABAM8CtQABARD/nAABARD/4QABANsBggABAon/LwABAnX/GwABAMb/GwABAOT/GwABApL++gABAnsCTAABAn7+5gABAnT/4QABAnsBgwABAMwCTAABAM/+5gABAMX/4QABAMwBegABALcCYQABAO3+5gABAOoCVQABARUCVQABAVT/nAABAM8CVQABANX/nAABAP4CVQABAOn/nAABASb/EQABARcCKAABAVT/EQABAVP/4QABARgBaQABANECKAABANX/EQABANT/4QABANIBaQABAQACKAABAOn/EQABAOj/4QABAQEBaQABARMBoAABAQ3+KQABAMQB9wABAL//GwABAMACbwABAKj/GwABARMCcgABAJYCyQABAJr/nAABAJr/4QABAJkBCwABAJ4C5AABARMCaAABARYAtAABAJwCvwABAKD/nAABAKD/4QABAJ8BCwABAJ4C2gABAbICnQABAbIC0AABAQz+uAABAQz+7wABAZ4CnQABAGMC0AABAH7/nAABAH7/4QABAFoCnQABAFkC0AABAHP/nAABAHP/4QABAGQCnQABAXcC4QABAYX/nAABAsMCRgABAVsCRgABAssCYgABAWMCYgABAWcCmgABAVr+sAABAVYBhgABAWcCfAABAUH+sgABAVYBaAABATsCgQABATn/nAABAWH/+AABASoBbQABAsUCbgABAVMCbgABAs4B+wABAtL/nAABAtL/4QABAs8BEwABAWYB+wABAUz/nAABAUz/4QABAWcBEwABANoA7gABAUT/AwABA/b/4QABA9AB1AABAGcBYAABAG7/nAABAG7/VAABAGUCQAABAGMBxAABAGP/nAABBB//0AABAboDjwABAGUBuAABAF3/ogABAK0BuAABAKr/ogABAMgB+AABAMgCwAABAHACJgABAIL+wgABANECCwABALf/HgABANsCHgABANv+5AABANgCHwABAN/+6QABAK0CCAABALr+4AABAOUCpwABAW3+qgABAM4BRgABAO3/qgABALoCwAABAOP/nAABAOP/4QABAL0B1AABAS4BzgABAS7/cAABAOwCKAABASb/nAABASX/4QABAO0BaQABAGr/kwABAGECRQABALACRwABAMgCygABADoCdgABAXACOwABAYT/nAABAoj/9QABAmsB5QABAMMCBgABAKf/nAABAKf/4QABAMMBgwABAMcBjgABAL7/nAABAL7/4QABAMcBCwABADkCmQABAPj/EgABAJoCZQABACACvgABADADGQABAEYDDgABAE0DEAABAEACUAABALcCBwABANP/nAABANP/4QABAL0BmAABAKEBqQABAIz/nAABAIz/4QABAKEBJgABARcCBgABAOz/nAABAOz/4QABARcBgwABAQ0BjgABAQ7/nAABAQ7/4QABAQ0BCwABAKb/nAABAGUCrAABAEAC+gABAKT/FwABAGsCxwABAEAC8AABAFkCRAABAMz/FwABAQAB3QABAQABWgABAQAB5wABAMH+6gABAQABZAABAIP+wQABAHD+8QABAHsClAABAHj/CgABAIsBgwABAJT/BgABAIf+5gABAIgCbQABAOoC5gABABICzAABABsCzAABAKL/nAABAcUC0AABAfH/nAABBmQByQABBmT/oQABCCT/5gABCD8BFgABAM8BoAABAMj+uAABAMj+7wABANIAtAABBZECSwABBZEAIwABApUAaAABAp4BkgAFAAAAAQAIAAEADACIAAIAkgHgAAIAFAVpBWoAAAVsBXQAAgV3BXgACwV6BXoADQV8BX4ADgWABYIAEQWEBYYAFAWIBYoAFwWMBY4AGgWQBZQAHQWaBZoAIgWcBZwAIwWeBaIAJAWkBbkAKQW7BbwAPwW+Bb4AQQXABcMAQgXFBcgARgXKBcoASgXNBdQASwABAAMEJgS5BWYAUwAABM4AAATUAAAE2gAABOAAAATmAAAE7AAABaAAAATyAAAE+AAABP4AAQJsAAAFBAABAnIAAAUKAAAFEAAABRYAAAUcAAAFIgAABUYAAAUoAAECeAABAoQAAAUuAAAGGAAABTQAAAU6AAAFQAAABUYAAAVMAAECfgABAoQAAAVSAAAFWAAABV4AAAVkAAAFagAABXAAAAV2AAAFfAAABYIAAAWIAAECigAABY4AAAWUAAAFmgAABaAAAAWmAAAFrAAABbIAAAW4AAAFvgAABcQAAAXQAAAFygAABdAAAQKQAAAF1gAABdwAAAXiAAEClgAABegAAAXuAAECnAAABfQAAQKiAAAF9AAABfoAAAYAAAAGBgAABgwAAAYSAAAGGAAABh4AAQKoAAAGJAAABioAAAY2AAECrgABArQAAAYwAAAGMAAABjYAAAY8AAMACAAoAGoAAwA+AEQADgBQABQAGgABAjMCLwABAO0CYAABAPz/nAAEABIAGAAeACQAKgAwADYAPAABBBoC0AABBB//nAABAxoCzQABA0L/nAABAjMDgwABAkf/nAABAO0CIwABAPz/ngADAA4AFAAaACAAJgAsAAEQXgGuAAEQdv/EAAEP1AFPAAEP4P/EAAEPEAFtAAEPGf/EAAYAEAABAAoAAAABAAwALAABAD4AxgABAA4FdAV4BYQFhQWQBZEFpAWyBbYFuQW8BcgFzwXQAAEABwV2BXsFkAW8BcgFzwXQAA4AAAA6AAAAQAAAAEYAAABSAAAATAAAAFIAAABYAAAAXgAAAGQAAABqAAAAcAAAAHYAAAB8AAAAggABAAD/pwABAAn//gABAE8ABgABAE7/+gABAFQALQABAAD/UgABAAD/1QABAAD/PAABAAD/nAABABz/RgABAE7/hgABAAH/yAABAAH/wQAHABAAFgAcACIAKAAuADQAAQBK/2oAAQBc/3MAAQBW/5MAAQAd/o8AAQCt/5MAAQAG/wYAAQAG/v8ABgAQAAEACgABAAEADACaAAEAugNEAAEARQVpBWoFbAVtBW4FbwVwBXEFcgVzBXcFegV8BX0FfgWABYEFggWGBYgFiQWKBYwFjQWOBZIFkwWUBZoFnAWeBZ8FoAWhBaIFpQWmBacFqAWpBaoFqwWsBa0FrgWvBbAFsQWzBbQFtQW3BbgFuwW+BcAFwQXCBcMFxQXGBccFygXNBc4F0QXSBdMF1AABAA4FdwV5BYgFngWiBaMFrQXBBcYFzQXOBdEF0gXTAEUAAAEWAAABHAAAASIAAAEoAAABLgAAATQAAAHoAAABOgAAAUAAAAFGAAABTAAAAVIAAAFYAAABXgAAAWQAAAFqAAABjgAAAXAAAAF2AAACYAAAAXwAAAGCAAABiAAAAY4AAAGUAAABmgAAAaAAAAGmAAABrAAAAbIAAAG4AAABvgAAAcQAAAHKAAAB0AAAAdYAAAHcAAAB4gAAAegAAAHuAAAB9AAAAfoAAAIAAAACBgAAAgwAAAIYAAACEgAAAhgAAAIeAAACJAAAAioAAAIwAAACNgAAAjwAAAI8AAACQgAAAkgAAAJOAAACVAAAAloAAAJgAAACZgAAAmwAAAJyAAACfgAAAngAAAJ4AAACfgAAAoQAAQAcAegAAQAnAcQAAQBjAdUAAQBbAdUAAQAsAfUAAQAwAawAAQAAAZgAAQAAAb8AAQAAAbYAAQALAa0AAQA8AgUAAQBVAcIAAQA4AfwAAQBCAb4AAQBaAbQAAQBBAbsAAQBYATwAAQA5AfwAAQBCAboAAQBKAbIAAQA7AfwAAQBBAb4AAQBFAWwAAQBFAbYAAQBIAcAAAQBZAckAAQBCAcEAAQBDAigAAQBcAfgAAQBLAckAAQB9AcUAAQCKATUAAQBPAdsAAQBoAecAAQAJAd8AAQAAAd8AAf//AbcAAQAPAZMAAQAAAcgAAQAAAXcAAQAAARwAAQBLAaYAAQAAAZMAAQAAAboAAQAAAdEAAf/6AbYAAf/3AZoAAQAAAcUAAQAAAa4AAQBTAagAAQAEAm4AAQACAoUAAf/8AhEAAQAE/wgAAf/8/uYAAQBaAboAAQDcAbIAAQAFAdYAAf/oAZ4AAf/RAbQAAQAAAbQAAQA7AbIADgAeACQAKgAwADYAPABCAEgATgBUAGAAWgBaAGAAAQAJAp0AAQA5AsIAAQBNAkAAAQBEAogAAQB9AfsAAQBnAhAAAf//AdIAAQABAv8AAQCkArQAAf/oAm0AAQBNAnMAAQAAAoMAAAABAAAACgDOAboAA0RGTFQAFGFyYWIAKmxhdG4ArgAEAAAAAP//AAYAAAACAAMABwALAA8AHAAERkFSIAA2S1NIIABqU05EIABQVVJEIABqAAD//wAKAAAAAQADAAQABQAGAAsADQAOABAAAP//AAoAAAABAAMABAAFAAYACAANAA4AEAAA//8ACgAAAAEAAwAEAAUABgAKAA0ADgAQAAD//wAKAAAAAQADAAQABQAGAAkADQAOABAABAAAAAD//wAGAAAAAgADAAcADAAPABFhYWx0AGhjY21wAHBjY21wAH5kbGlnAIRmaW5hAIppbml0AJBsaWdhAJZsaWdhAKBsb2NsAKZsb2NsAK5sb2NsALZsb2NsAL5sb2NsAMRtZWRpAMxybGlnANJybGlnAN50bnVtAOYAAAACAAAAAQAAAAUABwAIAAkACgAHAAAAAQAHAAAAAQAXAAAAAQALAAAAAQAMAAAAAwAQABEAEAAAAAEAEAAAAAIAEgAVAAAAAgASABIAAAACABIAFAAAAAEAEgAAAAIAEgATAAAAAQANAAAABAAOAA8ADwAOAAAAAgAOAA8AAAABABYAGAAyAawIeAjOCSQJPgmMCmILDAsyDHwNshD8EiQUahTcFmoWjhbCFuAW+hcUFzYXoAABAAAAAQAIAAIAugBaAAQACAALAA4AEQAUABgAGwAdALYAuAC8AL4AwADCAMQAxgDIAMoAzADOANAA0gDUANYA2ADaANwA3gDgAOIA5ADmAOgA6gDsAO4A8ADyAPYCAgIVAhcCGQIbAh0CHwIhAiMCJQInAikCKwItAi8CMQIzAjkCPAJsAngCegKFAocB/gTjBOEE5QTnBRAE6gURBO0FEgTwBPIFEwTlBOcFEATqBPsE8gUTBf4F/wUaBRsFIAUhAAEAWgADAAcACgANABAAEwAWABkAHAC1ALcAuwC9AL8AwQDDAMUAxwDJAMsAzQDPANEA0wDVANcA2QDbAN0A3wDhAOMA5QDnAOkA6wDtAO8A8QD1AgECFAIWAhgCGgIcAh4CIAIiAiQCJgIoAioCLAIuAjACMgI4AjsCawJ3AnkChAKGAowE4ATiBOQE5gToBOkE6wTsBO4E7wTxBPME9AT1BPYE9wT6BP8FAAUOBQ8FFAUXBR4FHwADAAAAAQAIAAEFoACQASYBLgE2AT4BRgFOAVYBXgFmAW4BdgF+AYYBjgGWAZ4BpgGuAbYBvgHGAc4B1gHeAeYB7gH2Af4CBgIOAhYCHgImAi4CNgI+AkYCTgJWAl4CZgJuAnYCfgKGAo4ClgKeAqYCrgK2Ar4CxgLOAtYC3gLmAu4C9gL+AwYDDgMWAx4DJgMuAzYDPgNGA04DVgNeA2YDbgN2A34DhgOOA5YDngOmA64DtgO+A8YDzgPWA94D5gPuA/YD/gQGBA4EFgQeBCYELgQ2BD4ERgROBFYEXgRmBG4EdgR+BIYEjgSWBJ4EpgSuBLYEvgTGBM4E1gTeBOYE7gT2BP4FBgUOBRYFHgUmBS4FNgU+BUYFTgVWBV4FZgVuBXYFfgWGBY4FlAWaAAMAIAAhACIAAwAkACUAJgADACgAKQAqAAMALAAtAC4AAwAwADEAMgADADYANwA4AAMAOgA7ADwAAwA+AD8AQAADAEIAQwBEAAMARgBHAEgAAwBKAEsATAADAE4ATwBQAAMAUgBTAFQAAwBWAFcAWAADAFoAWwBcAAMAXgBfAGAAAwBiAGMAZAADAGYAZwBoAAMAagBrAGwAAwBuAG8AcAADAHIAcwB0AAMAdgB3AHgAAwB6AHsAfAADAH4AfwCAAAMAggCDAIQAAwCGAIcAiAADAIoAiwCMAAMAjgCPAJAAAwCSAJMAlAADAJYAlwCYAAMAmgCbAJwAAwCeAJ8AoAADAKIAowCkAAMApgCnAKgAAwCqAKsArAADAK4ArwCwAAMAsgCzALQAAwD4APkA+gADAPwA/QD+AAMBAAEBAQIAAwEEAQUBBgADAQgBCQEKAAMBDAENAQ4AAwEQAREBEgADARQBFQEWAAMBGAEZARoAAwEcAR0BHgADASABIQEiAAMBJAElASYAAwEoASkBKgADASwBLQEuAAMBMAExATIAAwE0ATUBNgADATgBOQE6AAMBPAE9AT4AAwFAAUEBQgADAUQBRQFGAAMBSAFJAUoAAwFMAWQBZQADAVABUQFSAAMBVAFVAVYAAwFYAVkBWgADAVwBXQFeAAMBYAFhAWIAAwFkAWUBZgADAWgBaQFqAAMBRQFGAWwAAwFJAUoBbgADAXABcQFyAAMBdAF1AXYAAwF4AXkBegADAXwBfQF+AAMBgAGBAYIAAwGEAYUBhgADAYgBiQGKAAMBjAGNAY4AAwGQAZEBkgADAZQBlQGWAAMBmAGZAZoAAwGcAZ0BngADAaABoQGiAAMBpAGlAaYAAwGoAakBqgADAYEBggGsAAMBrgGvAbAAAwGyAbMBtAADAbYBtwG4AAMBugG7AbwAAwG+Ab8BwQADAcUBxgHHAAMBywHMAc4AAwHSAdMB1QADAdgB2QHbAAMB3gHfAeAAAwHjAeQB5QADAegB6QHqAAMB7AHtAe4AAwAzADQB8gADAEcASAH0AAMB9gH3AfgAAwApACoB+gADAf4B/wIAAAMCBAIFAgYAAwIIAgkCCgADAgwCDQIOAAMCEAIRAhIAAwAzADQCNQADAj8CQAJBAAMCQwJEAkUAAwJHAkgCSQADAksCTAJNAAMCUAJRAlIAAwJUAlUCVgADAlgCWQJaAAMCXAJdAl4AAwJgAmECYgADAmQCZQJmAAMCaAJpAmoAAwJuAm8CcAADAnICcwJ0AAMAKQAqAnYAAwJ8An0CfgADAoACgQKCAAMCiQKKAosAAwKOAo8CkAADApICkwKUAAMClgKXApgAAwKaApsCnAADAp4CnwKgAAMCogKjAqQAAwKmAqcCqAADAqoCqwKsAAMCrgKvArAAAwKyArMCtAADArYCtwK4AAMCugK7ArwAAwK/AsACwQADAsMCxALFAAMCxwLIAskAAwLLAswCzQADAs8C0ALRAAIE+QUOAAIE7gT9AAIE8AUPAAEAkAAfACMAJwArAC8ANQA5AD0AQQBFAEkATQBRAFUAWQBdAGEAZQBpAG0AcQB1AHkAfQCBAIUAiQCNAJEAlQCZAJ0AoQClAKkArQCxAPcA+wD/AQMBBwELAQ8BEwEXARsBHwEjAScBKwEvATMBNwE7AT8BQwFHAUsBTwFTAVcBWwFfAWMBZwFrAW0BbwFzAXcBewF/AYMBhwGLAY8BkwGXAZsBnwGjAacBqwGtAbEBtQG5Ab0BwwHJAdAB1wHdAeIB5wHrAe8B8wH1AfkB/QIDAgcCCwIPAjQCPgJCAkYCSgJPAlMCVwJbAl8CYwJnAm0CcQJ1AnsCfwKIAo0CkQKVApkCnQKhAqUCqQKtArECtQK5Ar4CwgLGAsoCzgT4BPwE/gABAAAAAQAIAAIAKAARAAYACQAMAA8AEgAVABcAHgHCAcgBzwHWAdwB4QI6Aj0CvQABABEABAAIAAsADgARABQAGAAdAcEBxwHOAdUB2wHgAjkCPAK8AAEAAAABAAgAAgAoABEGHQYjBiQGIAYhBh8GIgY7AcAEyQHNAdQB2gY+BjwGPQYeAAEAEQAEAAgACwAOABEAFAAYAB0BvwHGAcwB0wHZAd8COQI8ArsAAQAAAAEACAACA94ABgU+BUEFRgVIBU8FZwABAAAAAQAIAAID1AAgBNYE1wTYBNkE2gTbBNwE3QTeBN8FAQUCBQMFBAUFBQYFBwUIBQkFCgUBBQIFAwUEBQsFDAUNBQgFCQUKBgsGDAACAAAAAQAIAAEECAAZADgAPgBEAEoAUABWAFwAYgBoAG4AdAB6AIAAhgCMAJIAmACeAKQAqgCwALYAvADCAMgAAgYsBNYAAgYsBNcAAgYsBNgAAgYsBNkAAgYsBNoAAgYsBNsAAgYsBNwAAgYsBN0AAgYsBN4AAgYsBN8AAgYsBQEAAgYsBQIAAgYsBQMAAgYsBQQAAgYsBQUAAgYsBQYAAgYsBQcAAgYsBQgAAgYsBQkAAgYsBQoAAgYsBQsAAgYsBQwAAgYsBQ0AAgYsBgsAAgYsBgwABAAIAAEACAABAwgABgASACYARgBaAG4AggACAAYADgYVAAMGLAYsBhQAAgYsAAMACAASABoFRAAEBiwGLAYsBUMAAwYsBiwFQgACBiwAAgAGAA4GFwADBiwGLAYWAAIGLAACAAYADgYZAAMGLAYsBhgAAgYsAAIABgAOBVEAAwYsBiwFUAACBiwAAwAIABIAGgYcAAQGLAYsBiwGGwADBiwGLAYaAAIGLAACAAAAAQAIAAEACgACABIAGAABAAIABwAKAAIAAwV5AAIAAwV7AAQAAAABAAgAAQEmAAwAHgAwADoAZAB2AIgAmgCsAL4A0ADiARwAAgAGAAwABwACBXkACgACBXsAAQAEBZQAAgWTAAUADAASABgAHgAkBX0AAgV8BYEAAgWABYkAAgWIBY0AAgWMBXoAAgWcAAIABgAMBYUAAgWEBZEAAgWQAAIABgAMBX4AAgWTBX0AAgV5AAIABgAMBYIAAgWTBYEAAgV5AAIABgAMBYYAAgWTBYUAAgV7AAIABgAMBYoAAgWTBYkAAgV5AAIABgAMBY4AAgWTBY0AAgV5AAIABgAMBZIAAgWTBZEAAgV7AAcAEAAWABwAIgAoAC4ANAV+AAIFfAWCAAIFgAWGAAIFhAWKAAIFiAWOAAIFjAWSAAIFkAWUAAIFdwABAAQFegACBXkAAQAMAAMFdwV5BXsFfAWABYQFiAWMBZAFkwWcAAYAAAAGABIAOgBcAHgA4gEGAAMAAAAFAHwAjACMAIwAjAAAAAUAAAAEAAEABQACAAUAAwAFAAQABQADAAAABABUAGQAZABkAAAABAAAAAQAAQAFAAIABQADAAUAAwAAAAMAMgBCAEIAAAADAAAABAABAAUAAgAFAAMAAAACABYAJgAAAAIAAAAEAAEABQABAAYFPQVABUUFRwVNBWgAAQAgBMwEzQTOBM8E0ATRBNIE0wTUBNUE5ATmBOgE6QTrBOwE7gTvBPEE8wT0BPUE9gT3BPgE+gT8BP4E/wUABQ4FDwADAAIAPgAUAAEAPgAAAAEAAAAGAAEABgU+BUEFRgVIBU8FZwADAAIAGgAUAAEAGgAAAAEAAAAGAAEAAQYsAAIAAwTWBN8AAAUBBQ0ACgYLBgwAFwABAAAAAQAIAAIBogDOAAQACAALAA4AEQAUABgAGwAdACAAJAAoACwAMAA2ADoAPgBCAEYASgBOAFIAVgBaAF4AYgBmAGoAbgByAHYAegB+AIIAhgCKAI4AkgCWAJoAngCiAKYAqgCuALIAtgC4ALwAvgDAAMIAxADGAMgAygDMAM4A0ADSANQA1gDYANoA3ADeAOAA4gDkAOYA6ADqAOwA7gDwAPIA9gD4APwBAAEEAQgBDAEQARQBGAEcASABJAEoASwBMAE0ATgBPAFAAUQBSAFMAVABVAFYAVwBYAFmAWgBbAFuAXABdAF4AXwBgAGEAYgBjAGQAZQBmAGcAaABpAGoAawBrgGyAbYBugG+AcUBywHSAdgB3gHjAegB7AHyAfQB9gH6Af4CAgIEAggCDAIQAhUCFwIZAhsCHQIfAiECIwIlAicCKQIrAi0CLwIxAjMCNQI5AjwCPwJDAkcCSwJQAlQCWAJcAmACZAJoAmwCbgJyAnYCeAJ6AnwCgAKFAocCiQH+Ao4CkgKWApoCngKiAqYCqgKuArICtgK6Ar8CwwLHAssCzwABAM4AAwAHAAoADQAQABMAFgAZABwAHwAjACcAKwAvADUAOQA9AEEARQBJAE0AUQBVAFkAXQBhAGUAaQBtAHEAdQB5AH0AgQCFAIkAjQCRAJUAmQCdAKEApQCpAK0AsQC1ALcAuwC9AL8AwQDDAMUAxwDJAMsAzQDPANEA0wDVANcA2QDbAN0A3wDhAOMA5QDnAOkA6wDtAO8A8QD1APcA+wD/AQMBBwELAQ8BEwEXARsBHwEjAScBKwEvATMBNwE7AT8BQwFHAUsBTwFTAVcBWwFfAWMBZwFrAW0BbwFzAXcBewF/AYMBhwGLAY8BkwGXAZsBnwGjAacBqwGtAbEBtQG5Ab0BwwHJAdAB1wHdAeIB5wHrAe8B8wH1AfkB/QIBAgMCBwILAg8CFAIWAhgCGgIcAh4CIAIiAiQCJgIoAioCLAIuAjACMgI0AjgCOwI+AkICRgJKAk8CUwJXAlsCXwJjAmcCawJtAnECdQJ3AnkCewJ/AoQChgKIAowCjQKRApUCmQKdAqECpQKpAq0CsQK1ArkCvgLCAsYCygLOAAEAAAABAAgAAgJIAI0AIgAmACoALgAyADgAPABAAEQASABMAFAAVABYAFwAYABkAGgAbABwAHQAeAB8AIAAhACIAIwAkACUAJgAnACgAKQAqACsALAAtAD6AP4BAgEGAQoBDgESARYBGgEeASIBJgEqAS4BMgE2AToBPgFCAUYBSgFlAVIBVgFaAV4BYgFlAWoBRgFKAXIBdgF6AX4BggGGAYoBjgGSAZYBmgGeAaIBpgGqAYIBsAG0AbgBvAHBAccBzgHVAdsB4AHlAeoB7gA0AEgB+AAqAgACBgIKAg4CEgA0AkECRQJJAk0CUgJWAloCXgJiAmYCagJwAnQAKgJ+AoICiwKQApQCmAKcAqACpAKoAqwCsAK0ArgCvALBAsUCyQLNAtEAAQAAAAEACAACASAAjQAhACUAKQAtADEANwA7AD8AQwBHAEsATwBTAFcAWwBfAGMAZwBrAG8AcwB3AHsAfwCDAIcAiwCPAJMAlwCbAJ8AowCnAKsArwCzAPkA/QEBAQUBCQENAREBFQEZAR0BIQElASkBLQExATUBOQE9AUEBRQFJAWQBUQFVAVkBXQFhAWQBaQFFAUkBcQF1AXkBfQGBAYUBiQGNAZEBlQGZAZ0BoQGlAakBgQGvAbMBtwG7Ab8BxgHMAdMB2QHfAeQB6QHtADMARwH3ACkB/wIFAgkCDQIRADMCQAJEAkgCTAJRAlUCWQJdAmECZQJpAm8CcwApAn0CgQKKAo8CkwKXApsCnwKjAqcCqwKvArMCtwK7AsACxALIAswC0AABAI0AHwAjACcAKwAvADUAOQA9AEEARQBJAE0AUQBVAFkAXQBhAGUAaQBtAHEAdQB5AH0AgQCFAIkAjQCRAJUAmQCdAKEApQCpAK0AsQD3APsA/wEDAQcBCwEPARMBFwEbAR8BIwEnASsBLwEzATcBOwE/AUMBRwFLAU8BUwFXAVsBXwFjAWcBawFtAW8BcwF3AXsBfwGDAYcBiwGPAZMBlwGbAZ8BowGnAasBrQGxAbUBuQG9AcMByQHQAdcB3QHiAecB6wHvAfMB9QH5Af0CAwIHAgsCDwI0Aj4CQgJGAkoCTwJTAlcCWwJfAmMCZwJtAnECdQJ7An8CiAKNApEClQKZAp0CoQKlAqkCrQKxArUCuQK+AsICxgLKAs4ABgAIAAIACgAyAAMAAAACABYAUAAAAAIAAAACAAEAAgABAAcBwQHHAc4B1QHbAeACvAADAAAAAgAWACgAAAACAAAAAwABAAMAAQAHAb8BxgHMAdMB2QHfArsAAQAKAAQACAALAA4AEQAUABgAHQI5AjwABAAIAAEACAABAWoADAAeAGoApgDIANIA9AD+ASABKgFMAVYBYAAKABYAHAAiACgALgXuADQAOgBAAEYF9wACAA4ExQACABEC4AACAB0C4QACAjkC4gACAjwC7QACABQC6QACAAgC6wACAAsC3wACAAQACAASABgAHgXMACQAKgAwADYF9gACAA4ExAACABEExgACAB0C7AACABQC6AACAAgC6gACAAsC3gACAAQABAAKABAAFgAcBDMAAgAdBDQAAgI5BDUAAgI8BMgAAgAEAAEABATHAAIABAAEAAoAEAAWABwENwACAB0EOAACAjkEOQACAjwC4wACAAQAAQAEBDYAAgAEAAQACgAQABYAHAQ6AAIAHQQ7AAICOQQ8AAICPALlAAIABAABAAQC5AACAAQABAAKABAAFgAcBD0AAgAdBD4AAgI5BD8AAgI8AucAAgAEAAEABALmAAIABAABAAQF+QACAAQAAQAEBfgAAgAEAAEADAG/AcEBxgHHAcwBzgHTAdUB2QHbArsCvAAEAAgAAQAIAAEAFgABAAgAAgAGBJgEJgADAb8CBAABAAEBwQAEAAAAAQAIAAEAJgABAAgAAgAGABIEuQAFAcEBvwWUAgQEuQAFAcEBvwWUAf4AAQABAAMAAQAAAAEACAACAAwAAwUOBO4FDwABAAME+AT8BP4AAQAAAAEACAACAAoAAgUaBRsAAQACBRQFFwABAAAAAQAIAAIACgACBO4FDwABAAIE/AT+AAEAAAABAAgAAgAOAAQE4wThBSAFIQABAAQE4ATiBR4FHwABAAAAAQAIAAIAMgAWBOUE5wUQBOoFEQTtBRIE8ATyBRME5QTnBRAE6gT5BPsE/QTwBPIFEwX+Bf8AAQAWBOQE5gToBOkE6wTsBO4E7wTxBPME9AT1BPYE9wT4BPoE/AT+BP8FAAUOBQ8ABAAJAAEACAABBEIAIQBIAH4AtADqAQQBJgFeAZYBzgHqAgYCIgI0AkYCWAKEArAC3ALuAwADEgMkAz4DYAOEA5YDqAO6A9gD4gP0BA4EKAAGAA4AFgAeACQAKgAwALIAAwUpBSUAfgADBTEFIwCWAAIFKQCaAAIFLwB2AAIFLgCSAAIFJgAGAA4AFgAeACQAKgAwAH8AAwUjBTEAswADBSkFJQCXAAIFKQCbAAIFLwB3AAIFLgCTAAIFJgAGAA4AFgAeACQAKgAwAIAAAwUjBTEAtAADBSkFJQCYAAIFKQCYAAIFLwB4AAIFLgCUAAIFJgADAAgADgAUAMQAAgUlAMoAAgUtAMwAAgUyAAQACgAQABYAHADeAAIFJQDkAAIFKQDoAAIFMgDuAAIFJgAGAA4AFgAeACYALAAyAPwAAwUjBSUBCAADBS8FJQEMAAMFLwUxAQAAAgUxAo4AAgUyAr8AAgUmAAYADgAWAB4AJgAsADIA/QADBSMFJQEJAAMFLwUlAQ0AAwUvBTEBAQACBTECjwACBTICwAACBSYABgAOABYAHgAmACwAMgD+AAMFIwUlAQoAAwUvBSUBDgADBS8FMQECAAIFMQKQAAIFMgLBAAIFJgADAAgAEAAWASAAAwUlBSMBFAACBSoBGAACBS8AAwAIABAAFgEhAAMFJQUjARUAAgUqARkAAgUvAAMACAAQABYBIgADBSMFJQEWAAIFKgEaAAIFLwACAAYADAEwAAIFKQEoAAIFLwACAAYADAExAAIFKQEpAAIFLwACAAYADAEyAAIFKQEqAAIFLwAFAAwAFAAaACAAJgFAAAMFIwUlApIAAgUpATgAAgUvApYAAgUtApoAAgUmAAUADAAUABoAIAAmAUEAAwUjBSUCkwACBSkBOQACBS8ClwACBS0CmwACBSYABQAMABQAGgAgACYBQgADBSMFJQKUAAIFKQE6AAIFLwKYAAIFLQKcAAIFJgACAAYADAGkAAIFIwGsAAIFMQACAAYADAGBAAIFMQGBAAIFLgACAAYADAGOAAIFMQGCAAIFLgACAAYADAGMAAIFMQGAAAIFLgADAAgADgAUAcsAAgUjAdIAAgUvAdgAAgUxAAQACgAQABYAHALvAAIAGAHMAAIFIwHTAAIFLwHZAAIFMQAEAAoAEgAYAB4EJgADAb8B/gLuAAIAGAHOAAIFIwHVAAIFLwACAAYADAKmAAIFIwKqAAIFJQACAAYADAKnAAIFIwKrAAIFJQACAAYADAKoAAIFIwKsAAIFJQADAAgAEAAYAewAAwUlBSMCrgADBSoFIwH6AAIFLwABAAQCFwACBSkAAgAGAAwCHwACBSMCHQACBSkAAwAIAA4AFAJYAAIFKQJcAAIFLwJ2AAIFMQADAAgADgAUBDYAAgUjAuQAAgUvAuYAAgUxAAMACAAOABQC4wACBSMC5QACBS8C5wACBTEAAQAhAIoAiwCMALYA1AD4APkA+gEQAREBEgEkASUBJgE0ATUBNgF0AXUBdgF4Ab4BvwHBAeMB5AHlAfICBAIZAjUC3gLfAAAAAAABAAAAAA=="

def ensure_arabic_font():
    """ينشئ خط عربي محليًا من داخل التطبيق؛ لا يحتاج رفع ملف خط منفصل."""
    import tempfile
    font_path = os.path.join(tempfile.gettempdir(), "NotoNaskhArabic-Regular.ttf")
    if not os.path.exists(font_path) or os.path.getsize(font_path) < 50000:
        with open(font_path, "wb") as f:
            f.write(base64.b64decode(ARABIC_FONT_B64))
    return font_path

def pdf_bytes(pdf):
    """إرجاع PDF كـ bytes بدون latin-1 حتى تعمل العربية مع fpdf2."""
    return bytes(pdf.output())

# ====== دالة التصدير للاكسل RTL صح 100% ======
def to_excel(df):
    df = df.fillna('-')
    df = df.astype(str)

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='التقرير')
        worksheet = writer.sheets['التقرير']
        worksheet.sheet_view.rightToLeft = True

        for col in worksheet.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            worksheet.column_dimensions[column].width = max_length + 2

        from openpyxl.styles import Alignment
        for row in worksheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    return output.getvalue()

# ====== دالة التصدير للورد ======
def to_word(df, title, region):
    doc = Document()
    doc.add_heading(fix_arabic('الهيئة القومية للتأمين الاجتماعى'), 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(fix_arabic('الإدارة المركزية للإدارات القانونية'), 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(fix_arabic('الإدارة العامة للقضايا'), 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(fix_arabic(f'ديوان عام {region}'), 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(fix_arabic(title), 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = fix_arabic(str(col_name))
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = fix_arabic(str(val))
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(fix_arabic('تفضلوا بقبول وافر الاحترام\n'))
    p.add_run(fix_arabic('عضو الادارة.................. مدير الإدارة..................\n'))
    p.add_run(fix_arabic(f'تحر في {datetime.now().strftime("%Y-%m-%d")}'))

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# ====== دالة التصدير للـ PDF ======
def to_pdf(df, title, region):
    pdf = FPDF(orientation='L', unit='mm', format='A4')

    font_path = ensure_arabic_font()
    pdf.add_font('Cairo', '', font_path)

    pdf.add_page()
    pdf.set_font('Cairo', '', 16)
    pdf.cell(0, 10, fix_arabic('الهيئة القومية للتأمين الاجتماعى'), 0, 1, 'C')
    pdf.set_font('Cairo', '', 12)
    pdf.cell(0, 8, fix_arabic('الإدارة المركزية للإدارات القانونية'), 0, 1, 'C')
    pdf.cell(0, 8, fix_arabic('الإدارة العامة للقضايا'), 0, 1, 'C')
    pdf.cell(0, 8, fix_arabic(f'ديوان عام {region}'), 0, 1, 'C')
    pdf.cell(0, 8, fix_arabic(title), 0, 1, 'C')
    pdf.ln(5)

    pdf.set_font('Cairo', '', 8)
    col_width = 280 / len(df.columns)
    row_height = 8

    for col in df.columns:
        pdf.cell(col_width, row_height, fix_arabic(str(col)), 1, 0, 'C')
    pdf.ln()

    for _, row in df.iterrows():
        for item in row:
            pdf.cell(col_width, row_height, fix_arabic(str(item)), 1, 0, 'C')
        pdf.ln()

    pdf.ln(8)
    pdf.set_font('Cairo', '', 11)
    pdf.cell(0, 8, fix_arabic('تفضلوا بقبول وافر الاحترام'), 0, 1, 'R')
    pdf.ln(5)

    pdf.set_font('Cairo', '', 10)
    cell_w = 90
    pdf.cell(cell_w, 8, fix_arabic('العضو القانوني'), 0, 0, 'C')
    pdf.cell(cell_w, 8, fix_arabic('مدير إدارة القضايا'), 0, 0, 'C')
    pdf.cell(cell_w, 8, fix_arabic('مدير عام الإدارات القانونية'), 0, 1, 'C')

    pdf.ln(10)
    pdf.cell(cell_w, 8, '..................', 0, 0, 'C')
    pdf.cell(cell_w, 8, '..................', 0, 0, 'C')
    pdf.cell(cell_w, 8, '..................', 0, 1, 'C')

    pdf.ln(5)
    pdf.set_font('Cairo', '', 10)
    pdf.cell(0, 8, fix_arabic(f'تحر في {datetime.now().strftime("%d-%m-%Y")}'), 0, 1, 'L')

    return pdf_bytes(pdf)

# ====== دالة حفظ صحيفة الدعوى ======
def create_paper_pdf(case_data):
    pdf = FPDF(orientation='P', unit='mm', format='A4')
    font_path = ensure_arabic_font()
    pdf.add_font('Cairo', '', font_path)
    pdf.add_page()
    pdf.set_font('Cairo', '', 14); pdf.set_right_margin(15)
    pdf.cell(0,10,fix_arabic(f"صحيفة {case_data.get('مسندة_ل','')}"),ln=1,align='R')
    pdf.ln(5)
    pdf.cell(0,10,fix_arabic(f"محكمة: {case_data.get('محكمة_اسم','')}"),ln=1,align='R')
    pdf.cell(0,10,fix_arabic(f"رقم: {case_data.get('رقم','')} لسنة {case_data.get('سنة','')}"),ln=1,align='R')
    pdf.cell(0,10,fix_arabic(f"المدعي: {case_data.get('مدعي','')}"),ln=1,align='R')
    pdf.cell(0,10,fix_arabic(f"ضد: {case_data.get('مدعي_عليه','')}"),ln=1,align='R')
    pdf.multi_cell(0,10,fix_arabic(f"الموضوع: {case_data.get('موضوع','')}"),align='R')

    # <--- التعديل: بدل ما نحفظ في فولدر هنرجع bytes
    return pdf_bytes(pdf)

def print_case_report(case):
    نوع = case.get('نوع', '').lower()
    if 'استئناف' in نوع:
        طرف1_عنوان = "المستأنف"
        طرف2_عنوان = "المستأنف ضده"
    elif 'طعن' in نوع:
        طرف1_عنوان = "الطاعن"
        طرف2_عنوان = "المطعون ضده"
    else:
        طرف1_عنوان = "المدعي"
        طرف2_عنوان = "المدعى عليه"

    html = f"""
    <html dir="rtl" lang="ar">
    <head>
    <meta charset="UTF-8">
    <style>
        @page {{ size: A4; margin: 1.5cm; }}
        body {{ font-family: 'Arial'; direction: rtl; text-align: right; color: #000; background: #f8f9fa; }}
       .header {{ text-align: center; padding: 25px; margin-bottom: 25px; background: linear-gradient(135deg, #1E2A47 0%, #D4AF37 100%); color: #FFF; border-radius: 15px; box-shadow: 0 5px 15px rgba(0,0,0,0.2); }}
       .logo {{ font-size: 22px; font-weight: 900; color: #FFF; text-shadow: 2px 2px 4px rgba(0,0,0,0.3); }}
       .sub {{ font-size: 16px; color: #FFF9E6; margin: 8px 0; }}
       .title {{ text-align: center; font-size: 26px; font-weight: 900; color: #1E2A47; margin: 25px 0; border: 3px solid #D4AF37; padding: 15px; border-radius: 15px; background: linear-gradient(90deg, #FFF9E6, #FFF); box-shadow: 0 3px 10px rgba(212,175,55,0.3); }}
       .section {{ padding: 20px; border-radius: 15px; margin-bottom: 25px; box-shadow: 0 4px 15px rgba(0,0,0,0.1); border: 2px solid transparent; }}
       .section-title {{ font-weight: 900; font-size: 20px; color: #FFF; margin-bottom: 20px; text-align: center; padding: 12px; border-radius: 10px; }}
       .sec1 {{ background: linear-gradient(135deg, #1E2A47, #3498db); border-color: #1E2A47; }}
       .sec2 {{ background: linear-gradient(135deg, #27ae60, #2ecc71); border-color: #27ae60; }}
       .sec3 {{ background: linear-gradient(135deg, #8e44ad, #9b59b6); border-color: #8e44ad; }}
       .sec4 {{ background: linear-gradient(135deg, #c0392b, #e74c3c); border-color: #c0392b; }}
       .row {{ display: flex; justify-content: space-between; margin-bottom: 12px; background: linear-gradient(90deg, #fff, #f8f9fa); padding: 12px; border-radius: 8px; border-right: 4px solid #D4AF37; }}
       .label {{ font-weight: 900; color: #1E2A47; width: 35%; font-size: 15px; }}
       .value {{ width: 65%; color: #000; font-weight: 700; font-size: 15px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 15px; border-radius: 10px; overflow: hidden; box-shadow: 0 3px 10px rgba(0,0,0,0.1); }}
        th {{ background: linear-gradient(135deg, #1E2A47, #34495e); color: #D4AF37; padding: 12px; border: none; text-align: center; font-size: 16px; font-weight: 900; }}
        td {{ padding: 12px; border-bottom: 1px solid #ddd; text-align: center; background: #fff; }}
        tr:nth-child(even) td {{ background: #f8f9fa; }}
        tr:hover td {{ background: #FFF9E6; }}
    </style>
    </head>
    <body>

    <div class="header">
        <div class="logo">الهيئة القومية للتأمين الاجتماعي</div>
        <div class="sub">الإدارة المركزية للإدارات القانونية</div>
        <div class="sub">الإدارة العامة للشئون القانونية منطقة: _____________</div>
    </div>

    <div class="title">📄 تقرير تفاصيل القضية رقم {case.get('رقم')} لسنة {case.get('سنة')}</div>

    <div class="section sec1">
        <div class="section-title sec1">1- بيانات القضية</div>
        <div class="row"><div class="label">رقم القضية:</div><div class="value">{case.get('رقم')}</div></div>
        <div class="row"><div class="label">السنة:</div><div class="value">{case.get('سنة')}</div></div>
        <div class="row"><div class="label">النوع:</div><div class="value">{case.get('نوع')}</div></div>
        <div class="row"><div class="label">المحكمة:</div><div class="value">{case.get('محكمة_اسم')} {f'- مأمورية {case.get("مأمورية")}' if case.get('مأمورية') else ''}</div></div>
        <div class="row"><div class="label">الدائرة:</div><div class="value">{case.get('دائرة')}</div></div>
        <div class="row"><div class="label">الحالة:</div><div class="value">{case.get('حالة')}</div></div>
        <div class="row"><div class="label">الموضوع:</div><div class="value">{case.get('موضوع')}</div></div>
    </div>

    <div class="section sec2">
        <div class="section-title sec2">2- بيانات الخصوم</div>
        <div class="row"><div class="label">{طرف1_عنوان}:</div><div class="value">{case.get('مدعي')}</div></div>
        <div class="row"><div class="label">{طرف2_عنوان}:</div><div class="value">{case.get('مدعي_عليه')}</div></div>
    </div>
    """

    if case.get("جلسات"):
        html += """
        <div class="section sec3">
            <div class="section-title sec3">3- الجلسات والإجراءات</div>
            <table>
                <tr><th>م</th><th>الرول</th><th>الجلسات</th><th>الإجراءات</th><th>ملاحظات</th></tr>
        """
        for i, ج in enumerate(case["جلسات"], 1):
            html += f"<tr><td>{i}</td><td>{ج.get('الرول')}</td><td>{ج.get('تاريخ')}</td><td>{ج.get('الاجراء')}</td><td>{ج.get('ملاحظات')}</td></tr>"
        html += "</table></div>"

    if case.get('حالة') == 'منتهية':
        html += f"""
        <div class="section sec4">
            <div class="section-title sec4">4- منطوق الحكم</div>
            <div class="row"><div class="label">تاريخ الحكم:</div><div class="value">{case.get('تاريخ_الحكم')}</div></div>
            <div class="row"><div class="label">مسندة لـ:</div><div class="value">{case.get('مسندة_ل_الحكم')}</div></div>
            <div class="row"><div class="label">المنطوق:</div><div class="value">{case.get('منطوق_الحكم')}</div></div>
        </div>
        """

    html += "</body></html>"
    return html
# ====== دالة التحميل والحفظ الوحيدة ==
# ==== دالة التحميل والحفظ الوحيدة ======
DATA_FILE = "cases_data.json" # هنسيبه كـ باك اب بس
TOKENS_FILE = "tokens.json"

def load_data():
    """تحميل القضايا والمكتبة. المكتبة تستخدم جدول library إن وجد، وإلا تستخدم cases كبديل سحابي."""
    local = {"cases": [], "library": []}
    if os.path.exists(DATA_FILE):
        try:
            with open(DATA_FILE, "r", encoding="utf-8") as f:
                local = json.load(f)
            if not isinstance(local, dict):
                local = {"cases": [], "library": []}
            local.setdefault("cases", [])
            local.setdefault("library", [])
        except Exception:
            local = {"cases": [], "library": []}

    cloud_cases = None
    cloud_library = None

    try:
        response = supabase.table("cases").select("*").order("created_at", desc=True).execute()
        cloud_cases, fallback_library = [], []
        for row in (response.data or []):
            item = row.get("data", {}) or {}
            if item.get("_record_type") == "library":
                item = dict(item)
                item.pop("_record_type", None)
                item["id"] = row.get("id")
                fallback_library.append(item)
            else:
                item = dict(item)
                item["id"] = row.get("id")
                cloud_cases.append(item)
        cloud_library = fallback_library
    except Exception:
        cloud_cases = None

    # نحاول استخدام جدول library القديم لو كان موجودًا، لكن عدم وجوده لا يوقف التطبيق.
    if cloud_cases is not None:
        try:
            library_res = supabase.table("library").select("*").order("created_at", desc=True).execute()
            cloud_library = []
            for row in (library_res.data or []):
                item = dict(row.get("data", {}) or {})
                item["id"] = row.get("id")
                cloud_library.append(item)
        except Exception:
            pass

    return {
        "cases": cloud_cases if cloud_cases is not None else local.get("cases", []),
        "library": cloud_library if cloud_library is not None else local.get("library", [])
    }

def save_data(data):
    data.setdefault("cases", [])
    data.setdefault("library", [])

    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

    # حفظ القضايا في السحابة
    try:
        for case in data.get("cases", []):
            case_id = case.get("id")
            payload = {"data": case, "updated_at": datetime.now().isoformat()}
            if case_id:
                supabase.table("cases").update(payload).eq("id", case_id).execute()
            else:
                payload["created_at"] = datetime.now().isoformat()
                result = supabase.table("cases").insert(payload).execute()
                if result.data:
                    case["id"] = result.data[0]["id"]
    except Exception:
        pass

    # حفظ المكتبة في جدول library إن كان موجودًا، وإلا في جدول cases كنوع سجل مستقل.
    try:
        for item in data.get("library", []):
            item_id = item.get("id")
            payload = {"data": item, "updated_at": datetime.now().isoformat()}
            if item_id:
                try:
                    supabase.table("library").update(payload).eq("id", item_id).execute()
                except Exception:
                    # لو جدول library غير موجود لا نحاول تحويل معرف المكتبة إلى معرف قضية.
                    pass
            else:
                try:
                    payload["created_at"] = datetime.now().isoformat()
                    result = supabase.table("library").insert(payload).execute()
                    if result.data:
                        item["id"] = result.data[0]["id"]
                except Exception:
                    try:
                        record = dict(item)
                        record["_record_type"] = "library"
                        payload2 = {"data": record, "updated_at": datetime.now().isoformat(), "created_at": datetime.now().isoformat()}
                        result = supabase.table("cases").insert(payload2).execute()
                        if result.data:
                            item["id"] = result.data[0]["id"]
                    except Exception:
                        pass
    except Exception:
        pass

def save_library_item(item_data):
    data = load_data()
    data.setdefault("library", []).append(item_data)
    save_data(data)

def delete_library_item(item_id):
    data = load_data()
    data["library"] = [x for x in data.get("library", []) if x.get("id") != item_id]
    save_data(data)

def load_tokens():
    if os.path.exists(TOKENS_FILE):
        try:
            with open(TOKENS_FILE,"r",encoding="utf-8") as f: return json.load(f)
        except: pass
    return {"tokens":[]}

def save_tokens(tokens_data):
    with open(TOKENS_FILE,"w",encoding="utf-8") as f:
        json.dump(tokens_data, f, ensure_ascii=False, indent=4)

# ===== دوال التنبيهات ======
from datetime import datetime, timedelta

def get_alert_cases():
    data = load_data()
    today = datetime.now().date()
    all_cases = data.get("cases", [])
    alerts = {"sessions": [], "appeals": []}

    for case in all_cases:
        if case.get("حالة") == "متداولة" and case.get("تاريخ_جلسة"):
            try:
                session_date = datetime.strptime(case["تاريخ_جلسة"], "%Y-%m-%d").date()
                days_left = (session_date - today).days
                if 0 <= days_left <= 7:
                    case_copy = case.copy()
                    case_copy["days_left"] = days_left
                    alerts["sessions"].append(case_copy)
            except Exception:
                pass

        if case.get("حالة") == "منتهية" and case.get("مسندة_ل_الحكم") == "الضد" and case.get("تاريخ_الحكم"):
            try:
                judgment_date = datetime.strptime(case["تاريخ_الحكم"], "%Y-%m-%d").date()
                appeal_days = 40 if case.get("نوع") == "دعوى" else 60
                last_appeal_day = judgment_date + timedelta(days=appeal_days)
                notify_start = last_appeal_day - timedelta(days=15)
                days_left_appeal = (last_appeal_day - today).days
                if notify_start <= today <= last_appeal_day and days_left_appeal >= 0:
                    case_copy = case.copy()
                    case_copy["days_left_appeal"] = days_left_appeal
                    case_copy["deadline"] = last_appeal_day.strftime("%Y-%m-%d")
                    alerts["appeals"].append(case_copy)
            except Exception:
                pass

    return alerts

def send_alert_email(to_email, alerts):
    subject = f"🔔 تنبيهات ادارة القضايا - {datetime.now().strftime('%Y-%m-%d')}"
    body = "<div style='direction:rtl; text-align:right; font-family:Arial;'>"
    body += "<h2 style='color:#C9A961; text-align:center;'>مركز التنبيهات</h2>"

    if alerts["sessions"]:
        body += "<h3 style='color:#FFD700;'>⚖️ جلسات خلال 7 ايام</h3>"
        for case in alerts["sessions"]:
            body += f"<p style='border:1px solid #C9A961; padding:10px; border-radius:8px;'>"
            body += f"<b>رقم القضية:</b> {case.get('رقم_كامل','')}<br>"
            body += f"<b>الموضوع:</b> {case.get('موضوع_الدعوى','')}<br>"
            body += f"<b>الجلسة:</b> {case.get('تاريخ_جلسة')} - <b style='color:red;'>فاضل {case['days_left']} يوم</b></p>"
    else:
        body += "<p>✅ مفيش جلسات خلال 7 ايام</p>"

    if alerts["appeals"]:
        body += "<h3 style='color:#FF4500;'>📄 طعون خلال 15 يوم</h3>"
        for case in alerts["appeals"]:
            body += f"<p style='border:1px solid #FF4500; padding:10px; border-radius:8px;'>"
            body += f"<b>رقم القضية:</b> {case.get('رقم_كامل','')}<br>"
            body += f"<b>الموضوع:</b> {case.get('موضوع_الدعوى','')}<br>"
            body += f"<b style='color:red;'>اخر ميعاد للطعن: {case['deadline']} - فاضل {case['days_left_appeal']} يوم</b></p>"
    else:
        body += "<p>✅ مفيش طعون قريبة</p>"

    body += "</div>"

    try:
        send_email(to_email, subject, body)
        return True
    except Exception as e:
        st.error(f"فشل الارسال: {e}")
        return False

LIBRARY_SECTIONS = {
    "القوانين": "#FF4500", "القرارات الوزارية": "#FF8C00", "قرارات الهيئة": "#FFD700",
    "المنشورات الوزارية": "#ADFF2F", "منشورات الهيئة": "#32CD32", "الكتب الدورية": "#20B2AA",
    "تعليمات الهيئة": "#00CED1", "رسائل الهيئة": "#1E90FF", "المرصد الفنى": "#4169E1",
    "فتاوى لجنة الشئون القانونية بالوزارة": "#8A2BE2", "فتاوى الادارة المركزية للشئون القانونية": "#9400D3",
    "احكام المحكمة الدستورية العليا": "#DC143C", "احكام محكمة النقض": "#B22222", "احكام المحكمة الإدارية العليا": "#8B0000",
    "احكام المحاكم الاستئنافية": "#A0522D", "احكام محاكم القضاء الإدارى": "#D2691E", "احكام المحاكم الابتدائية": "#CD853F",
    "احكام المحكمة الإدارية": "#DEB887", "منشورات القضاء العادى": "#5F9EA0", "منشورات مجلس الدولة": "#4682B4",
    "فتاوى الجمعية العمومية": "#7B68EE", "صحف طعون": "#6A5ACD", "صحف استئنافات": "#483D8B",
    "صحف دعاوى": "#E6E6FA", "مذكرات دفاع": "#FFF0F5", "أخرى": "#808080"
}

SENDER_EMAIL=st.secrets.get("SENDER_EMAIL", ""); SENDER_PASSWORD=st.secrets.get("SENDER_PASSWORD", ""); APP_URL="https://qpyqpsmkqcvdou4imbfunp.streamlit.app/"
ANWA3_MOSTANDAT = ["صحيفة دعوى","صحيفة استئناف","صحيفة طعن","مذكرة دفاع","حافظة مستندات","تقرير خبير","تقرير طب شرعى","تقرير لجنة طبية","صحيفة تجديد من الشطب","صحيفة تعجيل من الوقف","صورة حكم تمهيدى","أخرى"]

if "page" not in st.session_state: st.session_state.page="الرئيسية"
if "selected_case_id" not in st.session_state: st.session_state.selected_case_id=None

# ============= التصميم النهائي =============
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;700;900&display=swap');
*{font-family:'Cairo',sans-serif!important;}
html,body{direction:rtl;color:#FFF!important;}
.stApp{background:linear-gradient(180deg,#0A1428 0%,#1E2A47 100%);}
.marquee{background:linear-gradient(90deg,#D4AF37,#FFD700,#D4AF37);color:#0A1428;padding:12px;font-weight:900;font-size:16px;white-space:nowrap;overflow:hidden;border-radius:0 0 15px 15px;}
.marquee span{display:inline-block;animation:marquee 15s linear infinite;}
@keyframes marquee{0%{transform:translateX(-100%);}100%{transform:translateX(100%);}}
.main-title{color:#D4AF37;text-align:center;font-size:36px;font-weight:900;padding:15px 0;}
h1,h2,h3{color:#D4AF37!important;text-align:center!important;}
div[data-testid="column"]{display:flex;justify-content:center;}
[data-testid="stForm"] label,.stMarkdown{color:#FFF!important;font-weight:700;}
.stButton>button{width:100%!important;max-width:400px!important;border:none!important;border-radius:15px!important;font-size:18px!important;font-weight:900!important;padding:16px!important;color:#000!important;}
</style>
""", unsafe_allow_html=True)

st.markdown("""<div class="marquee"><span>مع تحيات وليد حماد - الإدارة العامة للشئون القانونية بديوان عام منطقة البحيرة بالهيئة القومية للتأمين الاجتماعي</span></div>""", unsafe_allow_html=True)
st.markdown('<div class="main-title">⚖️ إدارة القضايا ⚖️</div>', unsafe_allow_html=True)

# =========================================
# =======================================
# ==================================================
# الصفحة الرئيسية
# ==================================================

if st.session_state.page == "الرئيسية":

    st.markdown("<h2>الأقسام</h2>", unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        st.markdown('<div class="btn-add">', unsafe_allow_html=True)
        if st.button(" تسجيل القضايا", use_container_width=True):
            st.session_state.page = "تسجيل"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="btn-list">', unsafe_allow_html=True)
        if st.button("📋 الحصر العام", use_container_width=True):
            st.session_state.page = "الحصر"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown('<div class="btn-alert">', unsafe_allow_html=True)
        if st.button("🔴 مركز التنبيهات", use_container_width=True):
            st.session_state.page = "التنبيهات"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        st.markdown('<div class="btn-report">', unsafe_allow_html=True)
        if st.button("📊 التقارير", use_container_width=True):
            st.session_state.page = "تقارير"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="btn-lib">', unsafe_allow_html=True)
        if st.button("📚 المكتبة القانونية", use_container_width=True):
            st.session_state.page = "مكتبة"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown('<div class="btn-arch">', unsafe_allow_html=True)
        if st.button("🗄️ الأرشيف", use_container_width=True):
            st.session_state.page = "الأرشيف"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown('<div class="btn-search">', unsafe_allow_html=True)
        if st.button("🔍 البحث عن دعوى", use_container_width=True):
            st.session_state.page = "بحث"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
        # =============================
# ====== الجزء الثاني: تسجيل القضية ============
elif st.session_state.page == "تسجيل":
    data = load_data()
    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    st.markdown("<h2 style='color:#D4AF37; text-align:center'> تسجيل القضية</h2>", unsafe_allow_html=True)
    if st.button("⬅️ العودة للرئيسية", key="back_add", use_container_width=True):
        st.session_state.page = "الرئيسية"
        st.rerun()

    st.markdown("<label style='color:#FFF; font-weight:700; text-align:right; width:100%; display:block;'>نوع القضية</label>", unsafe_allow_html=True)
    نوع = st.selectbox("", ["دعوى", "استئناف", "طعن"], key="case_type_add")

    with st.form("form_case_add", clear_on_submit=True):
        # 1- بيانات المحكمة
        st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
        st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>1- بيانات المحكمة</div>", unsafe_allow_html=True)
        محكمة_اسم = st.text_input("اسم المحكمة", key="court_name_add")
        مأمورية = st.text_input("المأمورية", key="mamoria_add") if نوع == "استئناف" else ""
        st.markdown("</div>", unsafe_allow_html=True)

        # 2- بيانات القضية
        st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
        st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>2- بيانات القضية</div>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1: رقم = st.text_input("رقم القضية / الاستئناف / الطعن", key="case_num_add")
        with col2: سنة = st.text_input("السنة القضائية", key="case_year_add")
        دائرة = st.text_input("الدائرة", key="circle_add")
        st.markdown("</div>", unsafe_allow_html=True)

        # 3- بيانات الخصوم
        st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
        st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>3- بيانات الخصوم</div>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1: مدعي = st.text_input("اسم المدعى / المستأنف / الطاعن", key="plaintiff_add")
        with col2: مدعي_عليه = st.text_input("اسم المدعى عليه / المستأنف ضده / المطعون ضده", key="defendant_add")
        موضوع = st.text_area("موضوع القضية", height=100, key="subject_add")
        st.markdown("</div>", unsafe_allow_html=True)

        # 4- بيانات الجلسة
        st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
        st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>4- بيانات الجلسة</div>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1: تاريخ_جلسة = st.date_input("تاريخ أول جلسة", value=datetime.now().date(), key="session_date_add")
        with col2: الرول = st.text_input("الرول", key="roll_add")
        الاجراء = st.text_input("الاجراء", key="reason_add")
        ملاحظات = st.text_area("ملاحظات", height=100, key="notes_add")

        # رفع مستندات
        uploaded_files = st.file_uploader("ارفع المستندات", type=["pdf", "jpg", "png", "docx"], accept_multiple_files=True)
        st.markdown("</div>", unsafe_allow_html=True)

        if st.form_submit_button("💾 حفظ القضية", use_container_width=True, type="primary"):
            if not رقم or not سنة:
                st.error("❌ من فضلك ادخل رقم القضية والسنة")
            else:
                # 1. نعمل صحيفة pdf
                case_for_pdf = {"نوع":نوع,"رقم":رقم,"سنة":سنة,"دائرة":دائرة,"محكمة_اسم":محكمة_اسم,"مدعي":مدعي,"مدعي_عليه":مدعي_عليه,"موضوع":موضوع,"تاريخ_جلسة":str(تاريخ_جلسة)}
                paper_bytes = create_paper_pdf(case_for_pdf)

                # 2. نحول الملفات ل base64
                مستندات = []
                if paper_bytes:
                    مستندات.append({
                        "name": f"صحيفة_{رقم}_{سنة}.pdf",
                        "data": base64.b64encode(paper_bytes).decode()
                    })

                for f in uploaded_files:
                    مستندات.append({
                        "name": f.name,
                        "data": base64.b64encode(f.getvalue()).decode()
                    })

                # 3. نجهز القضية
                username = st.session_state.user["username"] if st.session_state.user else "غير معروف"
                new_case = {
                    "نوع": نوع, "محكمة_اسم": محكمة_اسم, "مأمورية": مأمورية,
                    "رقم": رقم, "سنة": سنة, "دائرة": دائرة, "مدعي": مدعي, "مدعي_عليه": مدعي_عليه,
                    "موضوع": موضوع, "تاريخ_جلسة": str(تاريخ_جلسة), "الرول": الرول, "الاجراء": الاجراء,
                    "ملاحظات": ملاحظات, "جلسات": [], "مستندات": مستندات, "حالة": "متداولة",
                    "assigned_to": username,
                    "created_at": datetime.now().isoformat()
                }
                if الرول or الاجراء:
                    new_case["جلسات"].append({"تاريخ":str(تاريخ_جلسة),"الرول":الرول,"الاجراء":الاجراء,"ملاحظات":ملاحظات})

                data["cases"].append(new_case)
                save_data(data) # دي بتحفظ في السحابة والمحلي

                st.success(f"✅ تم الحفظ بنجاح - جاهز لتسجيل قضية جديدة")
                st.rerun()

# ===============================================
# ====== الجزء الثالث: الحصر العام ============
# ================================================
elif st.session_state.page == "الحصر":
    data = load_data()
    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    st.markdown("<h2 style='color:#FFFFFF; text-align:center'>📊 الحصر العام الخارجي</h2>", unsafe_allow_html=True)
    if st.button("⬅️ العودة للرئيسية", use_container_width=True): st.session_state.page = "الرئيسية"; st.rerun()

    if not data["cases"]:
        st.info("لا توجد قضايا مسجلة")
    else:
        # ======= تحديث اخر جلسة والاجراء من الجلسات =======
        for case in data["cases"]:
            if "جلسات" in case and case["جلسات"]:
                جلسات_مرتبة = sorted(case["جلسات"], key=lambda x: x.get("تاريخ","9999-12-31"), reverse=True)
                اخر_جلسة = جلسات_مرتبة[0]
                case["تاريخ_جلسة"] = اخر_جلسة.get("تاريخ","")
                case["الاجراء"] = اخر_جلسة.get("الاجراء","")
                case["الحالة"] = اخر_جلسة.get("الحالة", case.get("الحالة","متداولة"))

        # ======= نجيب قضايا العضو ده بس لو مش ادمن =======
        username = st.session_state.user["username"] if st.session_state.user else ""
        user_role = st.session_state.user["role"] if st.session_state.user else "member"

        if user_role == "admin":
            active_cases = [c for c in data["cases"] if c.get('حالة') == 'متداولة']
        else:
            active_cases = [c for c in data["cases"] if c.get('حالة') == 'متداولة' and c.get('assigned_to') == username]

        sorted_cases = sorted(active_cases, key=lambda x: x.get("تاريخ_جلسة","9999-12-31"))
        total = len(active_cases)
        today = datetime.now().date()
        start_week = today - timedelta(days=(today.weekday() + 2) % 7) # السبت
        end_week = start_week + timedelta(days=5) # الخميس

        this_week = len([c for c in active_cases if c.get('تاريخ_جلسة') and start_week <= datetime.strptime(c['تاريخ_جلسة'],'%Y-%m-%d').date() <= end_week])
        reserved = len([c for c in active_cases if any(k in str(c.get('الاجراء','')) for k in ['حكم', 'للحكم', 'الحكم'])])

        st.markdown(f"<div style='background:#1E2A47; padding:20px; border-radius:15px; border:2px solid #D4AF37; text-align:center; margin-bottom:20px'>", unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)
        with col1: st.markdown(f"<div style='font-size:28px; font-weight:900; color:#D4AF37'>📊 {total}</div><div style='font-size:18px; color:#FFF; font-weight:700'>اجمالي القضايا</div>", unsafe_allow_html=True)
        with col2: st.markdown(f"<div style='font-size:28px; font-weight:900; color:#4DA8DA'>📅 {this_week}</div><div style='font-size:18px; color:#FFF; font-weight:700'>جلسات هذا الاسبوع</div>", unsafe_allow_html=True)
        with col3: st.markdown(f"<div style='font-size:28px; font-weight:900; color:#FF5252'>⚖️ {reserved}</div><div style='font-size:18px; color:#FFF; font-weight:700'>المحجوز للحكم</div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("""
        <style>
.case-table {width:100%; border-collapse: collapse; font-size:11px; color:white; text-align:center; margin-bottom:5px;}
.case-table th {background:#D4AF37; color:#0B1426; padding:6px; font-weight:900;}
.case-table td {background:#1E2A47; padding:6px; border:1px solid #D4AF37; vertical-align:top;}
.plaintiff {background:#FFF3CD; color:#000; font-weight:700; border-radius:6px; padding:6px; font-size:11px;}
.plaintiff-hey2a {background:#DC3545!important; color:#FFF!important; font-weight:900; border-radius:6px; padding:6px; font-size:11px;}
.defendant {background:#CFF4FC; color:#000; font-weight:700; border-radius:6px; padding:6px; font-size:11px;}
.date-gold {color:#FFD700; font-weight:900;}
.status-green {color:#4CAF50; font-weight:900;}
        </style>
        """, unsafe_allow_html=True)

        for idx, case in enumerate(sorted_cases, 1):
            رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
            محكمة_كاملة = f"{case.get('نوع','')} {case.get('محكمة_اسم','')}"
            if case.get('مأمورية',''): محكمة_كاملة += f"<br>مأمورية {case.get('مأمورية','')}"
            دائرة_كاملة = f"{case.get('دائرة', '' )}" if case.get('دائرة', '' ) else ""
            if دائرة_كاملة: محكمة_كاملة += f"<br>دائرة {دائرة_كاملة}"

            نوع = case.get('نوع','')
            if نوع == "استئناف": لقب1, لقب2 = "المستأنف:", "المستأنف ضده:"
            elif نوع == "طعن": لقب1, لقب2 = "الطاعن:", "المطعون ضده:"
            else: لقب1, لقب2 = "المدعى:", "المدعى عليه:"

            if "الهيئة" in str(case.get('مدعي','')):
                طرف1_html = f"<div class='plaintiff-hey2a'><b>{لقب1}</b><br>{case.get('مدعي','')}</div>"
            else:
                طرف1_html = f"<div class='plaintiff'><b>{لقب1}</b><br>{case.get('مدعي','')}</div>"
            طرف2_html = f"<div class='defendant'><b>{لقب2}</b><br>{case.get('مدعي_عليه','')}</div>"
            خصوم = طرف1_html + "<div style='height:4px'></div>" + طرف2_html

            table_html = "<table class='case-table'><tr>"
            headers = ["م", "الرقم والسنة", "المحكمة والدائرة", "الخصوم", "الموضوع", "اخر جلسة", "الاجراء", "الحالة"]
            for h in headers: table_html += f"<th>{h}</th>"
            table_html += "</tr>"
            table_html += f"<tr><td>{idx}</td><td>{رقم_كامل}</td><td>{محكمة_كاملة}</td><td>{خصوم}</td><td>{case.get('موضوع','')}</td><td class='date-gold'>{case.get('تاريخ_جلسة','')}</td><td>{case.get('الاجراء','')}</td><td class='status-green'>{case.get('حالة','متداولة')}</td></tr></table>"
            st.markdown(table_html, unsafe_allow_html=True)

            c1, c2, c3 = st.columns([4,1,4])
            with c2:
                if st.button("فتح", key=f"open_{case.get('id', idx)}", use_container_width=True):
                    st.session_state.selected_case_id = case.get('id'); st.session_state.page = "تفاصيل"; st.rerun()
# ============================================
# ============ الجزء الرابع: تفاصيل القضية ============
elif st.session_state.page == "تفاصيل":
    data = load_data()
    case = next((c for c in data["cases"] if c.get("id") == st.session_state.selected_case_id), None)
    if not case: st.error("القضية غير موجودة"); st.session_state.page = "الحصر"; st.rerun()
    if 'جلسات' not in case: case['جلسات'] = []
    if 'مستندات' not in case: case['مستندات'] = []

    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    st.markdown(f"<h2 style='color:#D4AF37; text-align:center'>📄 تفاصيل القضية رقم {case.get('رقم')} لسنة {case.get('سنة')}</h2>", unsafe_allow_html=True)

    if st.button("⬅️ العودة للحصر", use_container_width=True): st.session_state.page = "الحصر"; st.rerun()

    # 1- بيانات القضية
    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:15px'>1- بيانات القضية</div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1: st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>رقم القضية</div><div style='color:#FFF; font-weight:900; font-size:22px'>{case.get('رقم')}</div></div>", unsafe_allow_html=True)
    with col2: st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>السنة</div><div style='color:#FFF; font-weight:900; font-size:22px'>{case.get('سنة')}</div></div>", unsafe_allow_html=True)
    with col3: دائرة_نص = f"{case.get('دائرة')}" if case.get('دائرة') else ""; st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>الدائرة</div><div style='color:#FFF; font-weight:900; font-size:18px'>{دائرة_نص}</div></div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1: st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>النوع</div><div style='color:#FFF; font-weight:900; font-size:18px'>{case.get('نوع')}</div></div>", unsafe_allow_html=True)
    with col2: محكمة_كاملة = f"{case.get('محكمة_اسم')}";
    if case.get('مأمورية'): محكمة_كاملة += f" - مأمورية {case.get('مأمورية')}"; st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>المحكمة</div><div style='color:#FFF; font-weight:700; font-size:14px'>{محكمة_كاملة}</div></div>", unsafe_allow_html=True)
    with col3: st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>الحالة</div><div style='color:#4CAF50; font-weight:900; font-size:18px'>{case.get('حالة')}</div></div>", unsafe_allow_html=True)
    st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>الموضوع</div><div style='color:#FFF; font-weight:700; font-size:16px'>{case.get('موضوع')}</div></div>", unsafe_allow_html=True)
    st.markdown("<style>div[data-testid='stExpander'] summary p{color:#D4AF37!important; font-weight:900!important;}</style>", unsafe_allow_html=True)
    with st.expander("✏️ تعديل بيانات القضية"):
        with st.form("edit_case_form"):
            col1, col2, col3 = st.columns(3)
            with col1: رقم_جديد = st.text_input("رقم القضية", value=case.get('رقم','')); سنة_جديد = st.text_input("السنة", value=case.get('سنة','')); نوع_جديد = st.selectbox("النوع", ["دعوى", "استئناف", "طعن"], index=["دعوى", "استئناف", "طعن"].index(case.get('نوع','دعوى')))
            with col2: محكمة_جديد = st.text_input("اسم المحكمة", value=case.get('محكمة_اسم','')); مأمورية_جديد = st.text_input("المأمورية", value=case.get('مأمورية','')); دائرة_جديد = st.text_input("الدائرة", value=case.get('دائرة',''))
            with col3: مدعي_جديد = st.text_input("المدعي", value=case.get('مدعي','')); مدعي_عليه_جديد = st.text_input("المدعي عليه", value=case.get('مدعي_عليه','')); حالة_جديد = st.selectbox("الحالة", ["متداولة", "مؤجلة", "منتهية", "شطب"], index=["متداولة", "مؤجلة", "منتهية", "شطب"].index(case.get('حالة','متداولة')) if case.get('حالة') in ["متداولة", "مؤجلة", "منتهية", "شطب"] else 0)
            موضوع_جديد = st.text_area("الموضوع", value=case.get('موضوع',''), height=100)
            if st.form_submit_button("💾 حفظ التعديلات", use_container_width=True, type="primary"):
                case['رقم']=رقم_جديد; case['سنة']=سنة_جديد; case['نوع']=نوع_جديد; case['محكمة_اسم']=محكمة_جديد; case['مأمورية']=مأمورية_جديد; case['دائرة']=دائرة_جديد; case['مدعي']=مدعي_جديد; case['مدعي_عليه']=مدعي_عليه_جديد; case['حالة']=حالة_جديد; case['موضوع']=موضوع_جديد
                save_data(data); st.success("✅ تم حفظ التعديلات"); st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

    # 2- بيانات الخصوم
    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>2- بيانات الخصوم</div>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1: st.markdown(f"<div style='background:#FFF3CD; padding:10px; border-radius:10px; color:#000; text-align:center'><b>المدعى:</b><br>{case.get('مدعي')}</div>", unsafe_allow_html=True)
    with col2: st.markdown(f"<div style='background:#CFF4FC; padding:10px; border-radius:10px; color:#000; text-align:center'><b>المدعى عليه:</b><br>{case.get('مدعي_عليه')}</div>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # 3- الجلسات والإجراءات
    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>3- الجلسات والإجراءات</div>", unsafe_allow_html=True)
    if case.get("جلسات"):
        for i, ج in enumerate(case["جلسات"]):
            st.markdown(f"<div style='background:#142038; padding:15px; border-radius:12px; border:2px solid #D4AF37; margin-bottom:10px; text-align:right; direction:rtl'><div style='display:flex; justify-content:flex-end; margin-bottom:10px'><div style='background:#D4AF37; color:#000; padding:5px 15px; border-radius:8px; font-weight:900; font-size:16px'>جلسة {i+1}</div></div><div style='margin-bottom:8px'><span style='color:#D4AF37; font-weight:900'>التاريخ:</span> <span style='color:#FFF'>{ج.get('تاريخ')}</span></div><div style='margin-bottom:8px'><span style='color:#D4AF37; font-weight:900'>الرول:</span> <span style='color:#FFF'>{ج.get('الرول')}</span></div><div style='margin-bottom:8px'><span style='color:#D4AF37; font-weight:900'>الاجراء:</span> <span style='color:#FFF'>{ج.get('الاجراء')}</span></div><div><span style='color:#D4AF37; font-weight:900'>ملاحظات:</span> <span style='color:#FFF'>{ج.get('ملاحظات')}</span></div></div>", unsafe_allow_html=True)
            if st.button("✏️ تعديل الجلسة", key=f"edit_session_{i}", use_container_width=True):
                st.session_state.edit_session_index = i; st.rerun()
        if 'edit_session_index' in st.session_state and st.session_state.edit_session_index is not None:
            idx = st.session_state.edit_session_index; جلسة = case["جلسات"][idx]
            with st.form("edit_session_form"):
                st.warning(f"تعديل الجلسة رقم {idx+1}")
                تاريخ_تعديل = st.date_input("التاريخ", value=datetime.strptime(جلسة.get('تاريخ'),'%Y-%m-%d').date())
                رول_تعديل = st.text_input("الرول", value=جلسة.get('الرول','')); اجراء_تعديل = st.text_input("الاجراء", value=جلسة.get('الاجراء','')); ملاحظات_تعديل = st.text_area("الملاحظات", value=جلسة.get('ملاحظات',''))
                c1,c2 = st.columns(2)
                with c1:
                    if st.form_submit_button("💾 حفظ تعديل الجلسة", use_container_width=True):
                        case["جلسات"][idx] = {"تاريخ":str(تاريخ_تعديل),"الرول":رول_تعديل,"الاجراء":اجراء_تعديل,"ملاحظات":ملاحظات_تعديل}
                        جلسات_مرتبة = sorted(case["جلسات"], key=lambda x: x.get("تاريخ","9999-12-31"), reverse=True)
                        case["تاريخ_جلسة"] = جلسات_مرتبة[0].get("تاريخ",""); case["الاجراء"] = جلسات_مرتبة[0].get("الاجراء","")
                        save_data(data); st.session_state.edit_session_index = None; st.success("تم التعديل"); st.rerun()
                with c2:
                    if st.form_submit_button("❌ الغاء", use_container_width=True): st.session_state.edit_session_index = None; st.rerun()
    else: st.info("لا توجد جلسات مسجلة")
    st.markdown("<style>div[data-testid='stExpander'] summary p{color:white!important; font-weight:900!important;}</style>", unsafe_allow_html=True)
    with st.expander("اضافة جلسة جديدة"):
        with st.form("add_session"):
            تاريخ_جديد = st.date_input("تاريخ الجلسة", value=datetime.now().date()); رول_جديد = st.text_input("الرول"); الاجراء_جديد = st.text_input("الاجراء"); ملاحظات_جديدة = st.text_area("ملاحظات")
            if st.form_submit_button("حفظ الجلسة"):
                case["جلسات"].append({"تاريخ":str(تاريخ_جديد),"الرول":رول_جديد,"الاجراء":الاجراء_جديد,"ملاحظات":ملاحظات_جديدة})
                case["تاريخ_جلسة"] = str(تاريخ_جديد); case["الاجراء"] = الاجراء_جديد; save_data(data); st.success("تم اضافة الجلسة"); st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

    # 4- المستندات
    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>4- المستندات</div>", unsafe_allow_html=True)

    ANWA3_MOSTANDAT = [
        "صحيفة دعوى", "صحيفة استئناف", "صحيفة طعن", "مذكرة دفاع",
        "حافظة مستندات", "تقرير خبير", "تقرير طب شرعى", "تقرير لجنة طبية",
        "صحيفة تجديد من الشطب", "صحيفة تعجيل من الوقف", "صورة حكم تمهيدى", "أخرى"
    ]

    نوع_المستند = st.selectbox("نوع المستند", ANWA3_MOSTANDAT, key="select_doc_type")

    اسم_نهائي = نوع_المستند
    if نوع_المستند == "أخرى":
        اسم_نهائي = st.text_input("✍️ اكتب اسم المستند", placeholder="مثال: طلب / انذار / الخ")

    with st.form("upload_form"):
        uploaded_file = st.file_uploader("اختر الملف", type=['pdf', 'jpg', 'jpeg', 'png', 'doc', 'docx'])
        if st.form_submit_button("رفع المستند"):
            if uploaded_file and اسم_نهائي and اسم_نهائي.strip()!= "":
                file_name = f"{اسم_نهائي}_{uploaded_file.name}"
                file_base64 = base64.b64encode(uploaded_file.getvalue()).decode('utf-8')
                case['مستندات'].append({"name": file_name, "data": file_base64}) # <--- وحدت الاسم
                save_data(data); st.success("✅ تم رفع المستند"); st.rerun()
            else:
                st.error("❌ لازم تختار ملف وتكتب اسم المستند")
    st.markdown("</div>", unsafe_allow_html=True)

    # عرض المستندات
    if case.get('مستندات'):
        st.markdown("<div style='background:#142038; padding:15px; border-radius:12px; margin-top:10px'>", unsafe_allow_html=True)
        st.markdown("<div style='color:#D4AF37; font-weight:900; margin-bottom:10px'>المستندات المرفوعة:</div>", unsafe_allow_html=True)
        for i, مستند in enumerate(case['مستندات']):
            اسم_المستند = مستند.get('name', f'ملف رقم {i+1}')
            محتوى_المستند = مستند.get('data', '')
            col1, col2, col3 = st.columns([4,1,1])
            with col1: st.write(f"📄 {اسم_المستند}")
            with col2:
                if محتوى_المستند:
                    try: file_data = base64.b64decode(محتوى_المستند); st.download_button("📥 تحميل", data=file_data, file_name=اسم_المستند, mime="application/octet-stream", key=f"dl_{i}_{case['id']}", use_container_width=True)
                    except: st.write("❌")
            with col3:
                if st.button("🗑️ حذف", key=f"del_{i}_{case['id']}", use_container_width=True): case['مستندات'].pop(i); save_data(data); st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    # 5- جلسة الحكم
    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #FF5252; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FF5252; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>5- جلسة الحكم</div>", unsafe_allow_html=True)
    if case.get('حالة') == 'منتهية':
        لون = "#4CAF50" if case.get('مسندة_ل_الحكم') == "الصالح" else "#FF5252"
        st.markdown(f"<div style='background:#142038; padding:15px; border-radius:12px; border:2px solid {لون}; margin-bottom:10px'><b style='color:{لون}'>تاريخ جلسة الحكم:</b> {case.get('تاريخ_الحكم')}<br><b style='color:{لون}'>مسندة لـ:</b> {case.get('مسندة_ل_الحكم')}<br><b style='color:{لون}'>منطوق الحكم:</b> {case.get('منطوق_الحكم')}</div>", unsafe_allow_html=True)
        st.success("✅ القضية منتهية")
        with st.expander("✏️ تعديل بيانات الحكم"):
            with st.form("edit_judgment_form"):
                تاريخ_حكم_تعديل = st.date_input("تاريخ الحكم", value=datetime.strptime(case.get('تاريخ_الحكم'),'%Y-%m-%d').date())
                منطوق_الحكم_تعديل = st.text_area("منطوق الحكم", value=case.get('منطوق_الحكم',''), height=150)
                مسندة_ل_تعديل = st.selectbox("مسندة لـ", ["الصالح", "الضد"], index=["الصالح", "الضد"].index(case.get('مسندة_ل_الحكم','الصالح')))
                if st.form_submit_button("💾 حفظ تعديل الحكم", use_container_width=True, type="primary"):
                    case['تاريخ_الحكم'] = str(تاريخ_حكم_تعديل); case['منطوق_الحكم'] = منطوق_الحكم_تعديل; case['مسندة_ل_الحكم'] = مسندة_ل_تعديل
                    for ج in reversed(case['جلسات']):
                        if 'الحكم' in ج.get('الاجراء',''): ج['تاريخ'] = str(تاريخ_حكم_تعديل); ج['الاجراء'] = f'الحكم - مسندة لـ {مسندة_ل_تعديل}'; ج['ملاحظات'] = منطوق_الحكم_تعديل; break
                    case['تاريخ_جلسة'] = str(تاريخ_حكم_تعديل); case['الاجراء'] = f'الحكم - مسندة لـ {مسندة_ل_تعديل}'
                    save_data(data); st.success("✅ تم تعديل الحكم"); st.rerun()
    else:
        with st.form("judgment_form"):
            st.markdown("<div style='background:#142038; padding:10px; border-radius:10px; margin-bottom:10px'>", unsafe_allow_html=True)
            st.markdown("<label style='color:#FFD700; font-weight:900; font-size:16px'>1- تاريخ الجلسة</label>", unsafe_allow_html=True); تاريخ_حكم = st.date_input("تاريخ الجلسة", value=datetime.now().date(), label_visibility="collapsed"); st.markdown("</div>", unsafe_allow_html=True)
            st.markdown("<div style='background:#142038; padding:10px; border-radius:10px; margin-bottom:10px'>", unsafe_allow_html=True)
            st.markdown("<label style='color:#FFD700; font-weight:900; font-size:16px'>2- منطوق الحكم</label>", unsafe_allow_html=True); منطوق_الحكم = st.text_area("منطوق الحكم", height=150, placeholder="اكتب منطوق الحكم هنا...", label_visibility="collapsed"); st.markdown("</div>", unsafe_allow_html=True)
            st.markdown("<div style='background:#142038; padding:10px; border-radius:10px; margin-bottom:10px'>", unsafe_allow_html=True)
            st.markdown("<label style='color:#FFD700; font-weight:900; font-size:16px'>3- مسندة لـ</label>", unsafe_allow_html=True); مسندة_ل = st.selectbox("مسندة لـ", ["الصالح", "الضد"], label_visibility="collapsed"); st.markdown("</div>", unsafe_allow_html=True)
            if st.form_submit_button("💾 حفظ الحكم", use_container_width=True, type="primary"):
                if not منطوق_الحكم: st.error("❌ لازم تكتب منطوق الحكم")
                else: case['حالة'] = 'منتهية'; case['تاريخ_الحكم'] = str(تاريخ_حكم); case['منطوق_الحكم'] = منطوق_الحكم; case['مسندة_ل_الحكم'] = مسندة_ل
                case['جلسات'].append({'تاريخ':str(تاريخ_حكم),'الرول':'-','الاجراء':f'الحكم - مسندة لـ {مسندة_ل}','ملاحظات':منطوق_الحكم}); case['تاريخ_جلسة'] = str(تاريخ_حكم); case['الاجراء'] = f'الحكم - مسندة لـ {مسندة_ل}'; save_data(data); st.success(f"✅ حفظت ونقلت للارشيف"); st.session_state.page = "الأرشيف"; st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

    # 6- الطباعة والتحميل
    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px; text-align:center'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; margin-bottom:10px'>🖨️ الطباعة والتقرير</div>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🖨️ معاينة للطباعة", use_container_width=True, type="primary"):
            html_report = print_case_report(case)
            st.components.v1.html(html_report, height=800, scrolling=True)
            st.success("✅ اضغط Ctrl+P للطباعة")
    with col2:
        html_report = print_case_report(case)
        st.download_button(label="📥 تحميل التقرير",data=html_report.encode('utf-8'),file_name=f"تقرير_قضية_{case.get('رقم')}_{case.get('سنة')}.html",mime="text/html",use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # 7- حذف نهائى - منطقة خطر
    st.markdown("<div style='background:#2A0A0A; padding:20px; border-radius:15px; border:3px solid #FF0000; margin-bottom:15px; text-align:center'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FF0000; font-size:22px; font-weight:900; margin-bottom:10px'>⚠️ منطقة خطر</div>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FFF; font-size:14px; margin-bottom:15px'>تحذير: حذف القضية نهائي ولا يمكن التراجع عنه</div>", unsafe_allow_html=True)

    if st.button("🗑️ حذف القضية نهائي", use_container_width=True, type="secondary"):
        st.session_state.confirm_delete = True
        st.rerun()

    if st.session_state.get('confirm_delete', False):
        st.warning("⚠️ هل انت متأكد 100% انك عايز تحذف القضية دي؟")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("نعم احذفها", use_container_width=True, type="primary"):
                # حذف من السحابة والمحلي
                supabase.table("cases").delete().eq("id", case["id"]).execute()
                data["cases"] = [c for c in data["cases"] if c["id"]!= case["id"]]
                save_data(data); st.session_state.confirm_delete = False
                st.success("✅ تم حذف القضية بنجاح"); st.session_state.page = "الحصر"; st.rerun()
        with col2:
            if st.button("الغاء", use_container_width=True):
                st.session_state.confirm_delete = False; st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)
# ========================================
# ============ الجزء الخامس: الأرشيف ============
elif st.session_state.page == "الأرشيف":
    data = load_data()

    st.markdown("""
    <style>
        label { color: #FFD700 !important; font-weight: 900 !important; font-size: 15px !important; }
        input::placeholder, textarea::placeholder {
            color: #FFD700 !important;
            opacity: 1 !important;
            font-weight: 600;
        }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    st.markdown("<h2 style='color:#D4AF37; text-align:center'>📁 الأرشيف</h2>", unsafe_allow_html=True)
    
    if st.button("⬅️ العودة للرئيسية", use_container_width=True): 
        st.session_state.page = "الرئيسية"; 
        st.rerun()

    # 1- شريط البحث
    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FFF; font-size:18px; font-weight:900; text-align:center; margin-bottom:10px'>🔍 البحث عن قضية صدر فيها الحكم</div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([3,3,1])
    with col1: بحث_مدعي = st.text_input("بحث بالاسم", placeholder="اكتب اي اسم")
    with col2: بحث_رقم = st.text_input("بحث برقم وسنة", placeholder="مثال: 123 لسنة 2024")
    with col3: st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True); بحث_زر = st.button("🔍 بحث", use_container_width=True, type="primary")
    st.markdown("</div>", unsafe_allow_html=True)

    # فلترة القضايا المنتهية فقط + بتاعت العضو لو مش ادمن
    username = st.session_state.user["username"] if st.session_state.user else ""
    user_role = st.session_state.user["role"] if st.session_state.user else "member"

    if user_role == "admin":
        قضايا_منتهية = [c for c in data["cases"] if c.get("حالة") == "منتهية"]
    else:
        قضايا_منتهية = [c for c in data["cases"] if c.get("حالة") == "منتهية" and c.get('assigned_to') == username]
    
    # فلترة البحث
    if بحث_زر:
        if بحث_مدعي: 
            بحث_مدعي = بحث_مدعي.lower()
            قضايا_منتهية = [c for c in قضايا_منتهية if any(
                بحث_مدعي in str(قيمة).lower() 
                for قيمة in c.values() 
                if isinstance(قيمة, str)
            )]
        if بحث_رقم: 
            قضايا_منتهية = [c for c in قضايا_منتهية if بحث_رقم in f"{c.get('رقم')} لسنة {c.get('سنة')}"]

    # نقسمهم 2
    قضايا_جاري = [c for c in قضايا_منتهية if not c.get("تم_الحفظ_النهائي")]
    قضايا_محفوظة = [c for c in قضايا_منتهية if c.get("تم_الحفظ_النهائي")]

    # تبويب 1: احكام صادرة وجاري اتخاذ الاجراء
    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #FFD700; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FFD700; font-size:20px; font-weight:900; text-align:center; margin-bottom:15px'>1- احكام صادرة وجاري اتخاذ الاجراء اللازم بشأنها</div>", unsafe_allow_html=True)
    
    if قضايا_جاري:
        for case in قضايا_جاري:
            لون = "#4CAF50" if case.get('مسندة_ل_الحكم') == "الصالح" else "#FF5252"
            
            نوع = case.get('نوع', '').lower()
            if 'استئناف' in نوع:
                طرف1_عنوان = "المستأنف"
                طرف2_عنوان = "المستأنف ضده"
            elif 'طعن' in نوع:
                طرف1_عنوان = "الطاعن"
                طرف2_عنوان = "المطعون ضده"
            else:
                طرف1_عنوان = "المدعي"
                طرف2_عنوان = "المدعي عليه"

            st.markdown(f"<div style='background:#142038; padding:15px; border-radius:12px; border:2px solid {لون}; margin-bottom:10px'>", unsafe_allow_html=True)
            
            st.markdown(f"""
            <table style='width:100%; border-collapse:collapse; margin-bottom:10px;'>
                <tr><th colspan='2' style='background:{لون}; color:#FFF; padding:10px; text-align:center; font-size:16px; border-radius:8px 8px 0 0;'>رقم {case.get('رقم')} لسنة {case.get('سنة')} - {case.get('نوع')}</th></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; width:35%; font-weight:900;'>المحكمة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('محكمة_اسم')} {f'- مأمورية {case.get("مأمورية")}' if case.get('مأمورية') else ''}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الدائرة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('دائرة')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الموضوع</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('موضوع')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{طرف1_عنوان}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{طرف2_عنوان}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي_عليه')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>تاريخ الحكم</td><td style='background:#FFF; color:{لون}; padding:8px; font-weight:900;'>{case.get('تاريخ_الحكم')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>المنطوق</td><td style='background:#FFF; color:{لون}; padding:8px; font-weight:900;'>{case.get('منطوق_الحكم')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900; border-radius:0 0 0 8px;'>مسندة لـ</td><td style='background:#FFF; color:{لون}; padding:8px; font-weight:900; border-radius:0 0 8px 0;'>{case.get('مسندة_ل_الحكم')}</td></tr>
            </table>
            """, unsafe_allow_html=True)
            
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("📄 فتح", key=f"open_{case['id']}", use_container_width=True):
                    st.session_state.selected_case_id = case["id"]; st.session_state.page = "تفاصيل"; st.rerun()
            with col2:
                if st.button("💾 حفظ نهائي", key=f"save_{case['id']}", use_container_width=True):
                    st.session_state.save_case_id = case["id"]; st.rerun()
            with col3:
                if st.button("🗑️ حذف", key=f"del_arch_{case['id']}", use_container_width=True):
                    st.session_state.del_arch_id = case["id"]; st.rerun()
            
            # فورم الحفظ النهائي
            if st.session_state.get('save_case_id') == case['id']:
                with st.form(f"save_form_{case['id']}"):
                    st.warning("حفظ القضية نهائي")
                    سبب_الحفظ = st.text_area("سبب الحفظ", placeholder="مثال: تم الطعن / حكم نهائي / عدم جدوى")
                    مستندات_الحفظ = st.file_uploader("ارفع مستندات الحفظ", type=['pdf','jpg','png','doc','docx'], accept_multiple_files=True)
                    
                    if st.form_submit_button("💾 تأكيد الحفظ النهائي", use_container_width=True, type="primary"):
                        case['سبب_الحفظ'] = سبب_الحفظ
                        case['مستندات_الحفظ'] = []
                        for f in مستندات_الحفظ:
                            file_base64 = base64.b64encode(f.getvalue()).decode('utf-8')
                            case['مستندات_الحفظ'].append({"name": f.name, "data": file_base64}) # <--- وحدنا الاسم
                        case['تم_الحفظ_النهائي'] = True
                        case['تاريخ_الحفظ'] = str(datetime.now().date())
                        save_data(data); st.session_state.save_case_id = None
                        st.success("✅ تم حفظ القضية نهائي"); st.rerun()

            # منطقة خطر الحذف
            if st.session_state.get('del_arch_id') == case['id']:
                st.error("⚠️ هل انت متأكد 100% من حذف القضية نهائي من الارشيف؟")
                c1,c2 = st.columns(2)
                with c1:
                    if st.button("نعم احذف", key=f"confirm_del_{case['id']}"):
                        supabase.table("cases").delete().eq("id", case["id"]).execute() # <--- حذف من السحابة
                        data["cases"] = [c for c in data["cases"] if c["id"]!= case["id"]]
                        save_data(data); st.session_state.del_arch_id = None; st.success("تم الحذف"); st.rerun()
                with c2:
                    if st.button("الغاء", key=f"cancel_del_{case['id']}"):
                        st.session_state.del_arch_id = None; st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)
    else: st.info("لا توجد احكام")
    st.markdown("</div>", unsafe_allow_html=True)

    # تبويب 2: احكام صادرة وتم حفظها
    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #4CAF50; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#4CAF50; font-size:20px; font-weight:900; text-align:center; margin-bottom:15px'>2- احكام صادرة وتم اتخاذ الاجراء اللازم بشأنها وحفظت</div>", unsafe_allow_html=True)

    if قضايا_محفوظة:
        for case in قضايا_محفوظة:
            نوع = case.get('نوع', '').lower()
            if 'استئناف' in نوع:
                طرف1_عنوان = "المستأنف"
                طرف2_عنوان = "المستأنف ضده"
            elif 'طعن' in نوع:
                طرف1_عنوان = "الطاعن"
                طرف2_عنوان = "المطعون ضده"
            else:
                طرف1_عنوان = "المدعي"
                طرف2_عنوان = "المدعي عليه"

            st.markdown(f"<div style='background:#142038; padding:15px; border-radius:12px; border:2px solid #4CAF50; margin-bottom:10px'>", unsafe_allow_html=True)
            st.markdown(f"""
            <table style='width:100%; border-collapse:collapse; margin-bottom:10px;'>
                <tr><th colspan='2' style='background:#4CAF50; color:#FFF; padding:10px; text-align:center; font-size:16px; border-radius:8px 8px 0 0;'>رقم {case.get('رقم')} لسنة {case.get('سنة')} - {case.get('نوع')}</th></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; width:35%; font-weight:900;'>المحكمة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('محكمة_اسم')} {f'- مأمورية {case.get("مأمورية")}' if case.get('مأمورية') else ''}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الدائرة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('دائرة')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الموضوع</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('موضوع')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{طرف1_عنوان}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{طرف2_عنوان}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي_عليه')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>تاريخ الحفظ</td><td style='background:#FFF; color:#4CAF50; padding:8px; font-weight:900;'>{case.get('تاريخ_الحفظ')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900; border-radius:0 0 0 8px;'>سبب الحفظ</td><td style='background:#FFF; color:#FFD700; padding:8px; font-weight:900; border-radius:0 0 8px 0;'>{case.get('سبب_الحفظ')}</td></tr>
            </table>
            """, unsafe_allow_html=True)
            
            if case.get('مستندات_الحفظ'):
                st.markdown("<div style='color:#D4AF37; margin-top:10px'>مستندات الحفظ:</div>", unsafe_allow_html=True)
                for i, مستند in enumerate(case['مستندات_الحفظ']):
                    file_data = base64.b64decode(مستند['data']) # <--- وحدنا الاسم
                    st.download_button(f"📥 {مستند['name']}", data=file_data, file_name=مستند['name'], key=f"dl_save_{case['id']}_{i}")

            col1, col2 = st.columns(2)
            with col1:
                if st.button("📄 فتح", key=f"open_saved_{case['id']}", use_container_width=True):
                    st.session_state.selected_case_id = case["id"]; st.session_state.page = "تفاصيل"; st.rerun()
            with col2:
                if st.button("🗑️ حذف نهائي", key=f"del_saved_{case['id']}", use_container_width=True):
                    st.session_state.del_saved_id = case["id"]; st.rerun()

            if st.session_state.get('del_saved_id') == case['id']:
                st.error("⚠️ هل انت متأكد 100% من حذف القضية نهائي؟")
                c1,c2 = st.columns(2)
                with c1:
                    if st.button("نعم احذف", key=f"confirm_del_saved_{case['id']}"):
                        supabase.table("cases").delete().eq("id", case["id"]).execute() # <--- حذف من السحابة
                        data["cases"] = [c for c in data["cases"] if c["id"]!= case["id"]]
                        save_data(data); st.session_state.del_saved_id = None; st.success("تم الحذف"); st.rerun()
                with c2:
                    if st.button("الغاء", key=f"cancel_del_saved_{case['id']}"):
                        st.session_state.del_saved_id = None; st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)
    else: st.info("لا توجد قضايا محفوظة نهائي")
    st.markdown("</div>", unsafe_allow_html=True)
# ==========================================
# =========== الجزء السادس: البحث ============
elif st.session_state.page == "بحث":
    import base64
    data = load_data()
    
    st.markdown("""
    <style>
        label { color: #FFD700 !important; font-weight: 900 !important; font-size: 15px !important; }
        input::placeholder { color: #FFD700 !important; opacity: 1 !important; font-weight: 600; }
        .case-table {width:100%; border-collapse: collapse; font-size:12px; color:#FFF; text-align:center; margin-bottom:10px;}
        .case-table th {background:#D4AF37; color:#0B1426; padding:8px; font-weight:900;}
        .case-table td {background:#1E2A47; padding:8px; border:1px solid #D4AF37; vertical-align:top; color:#FFF;}
        .plaintiff {background:#FFF3CD; color:#000; font-weight:700; border-radius:6px; padding:6px; font-size:12px;}
        .plaintiff-hey2a {background:#DC3545!important; color:#FFF!important; font-weight:900; border-radius:6px; padding:6px; font-size:12px;}
        .defendant {background:#CFF4FC; color:#000; font-weight:700; border-radius:6px; padding:6px; font-size:12px;}
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    st.markdown("<h2 style='color:#4DA8DA; text-align:center'>🔍 البحث عن دعوى</h2>", unsafe_allow_html=True)

    if st.button("⬅️ العودة للرئيسية", use_container_width=True, key="back_from_search"):
        st.session_state.page = "الرئيسية"
        st.rerun()

    st.markdown("<div style='background:#1E2A47; padding:20px; border-radius:15px; border:2px solid #4DA8DA; margin-bottom:20px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FFF; font-size:18px; font-weight:900; text-align:center; margin-bottom:15px'>ابحث بالاسم او برقم وسنة الدعوى</div>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1: 
        بحث_اسم = st.text_input("بحث بالاسم", placeholder="اكتب اسم المدعي او المدعى عليه")
    with col2: 
        بحث_رقم = st.text_input("بحث برقم وسنة", placeholder="مثال: 123 لسنة 2024")
    
    بحث_زر = st.button("🔍 بحث", use_container_width=True, type="primary")
    st.markdown("</div>", unsafe_allow_html=True)

    if بحث_زر:
        if not بحث_اسم.strip() and not بحث_رقم.strip():
            st.error("اكتب اسم او رقم للبحث")
        else:
            # فلترة حسب المستخدم
            username = st.session_state.user["username"] if st.session_state.user else ""
            user_role = st.session_state.user["role"] if st.session_state.user else "member"

            if user_role == "admin":
                all_cases = data["cases"]
            else:
                all_cases = [c for c in data["cases"] if c.get('assigned_to') == username]

            results = []
            بحث_اسم = بحث_اسم.lower()
            
            for case in all_cases:
                match = False
                if بحث_اسم:
                    if any(بحث_اسم in str(قيمة).lower() for قيمة in case.values() if isinstance(قيمة, str)):
                        match = True
                if بحث_رقم:
                    رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
                    if بحث_رقم in رقم_كامل:
                        match = True
                if match:
                    results.append(case)

            if not results:
                st.warning("⚠️ لا يوجد بحث مطابق")
            else:
                st.success(f"✅ تم العثور على {len(results)} نتيجة")
                for idx, case in enumerate(results, 1):
                    
                    نوع = case.get('نوع', '').lower()
                    if 'استئناف' in نوع:
                        لقب1, لقب2 = "المستأنف:", "المستأنف ضده:"
                    elif 'طعن' in نوع:
                        لقب1, لقب2 = "الطاعن:", "المطعون ضده:"
                    else:
                        لقب1, لقب2 = "المدعي:", "المدعي عليه:"

                    if "الهيئة" in str(case.get('مدعي','')):
                        طرف1_html = f"<div class='plaintiff-hey2a'><b>{لقب1}</b><br>{case.get('مدعي','')}</div>"
                    else:
                        طرف1_html = f"<div class='plaintiff'><b>{لقب1}</b><br>{case.get('مدعي','')}</div>"
                    طرف2_html = f"<div class='defendant'><b>{لقب2}</b><br>{case.get('مدعي_عليه','')}</div>"
                    خصوم = طرف1_html + "<div style='height:4px'></div>" + طرف2_html

                    رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
                    محكمة_كاملة = f"{case.get('نوع','')} {case.get('محكمة_اسم','')}"
                    if case.get('مأمورية',''): محكمة_كاملة += f"<br>مأمورية {case.get('مأمورية','')}"
                    if case.get('دائرة',''): محكمة_كاملة += f"<br>دائرة {case.get('دائرة','')}"

                    if case.get('حالة') == 'منتهية':
                        حالة_لون = "#FF5252"
                        مكان = "📁 الأرشيف"
                    else:
                        حالة_لون = "#4CAF50"
                        مكان = "📋 الحصر العام"

                    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
                    st.markdown(f"""
                    <table class='case-table'>
                    <tr><th>م</th><th>الرقم والسنة</th><th>المحكمة والدائرة</th><th>الخصوم</th><th>الموضوع</th><th>اخر جلسة</th><th>الحالة</th><th>المكان</th></tr>
                    <tr>
                        <td>{idx}</td>
                        <td>{رقم_كامل}</td>
                        <td>{محكمة_كاملة}</td>
                        <td>{خصوم}</td>
                        <td>{case.get('موضوع','')}</td>
                        <td style='color:#FFD700; font-weight:900'>{case.get('تاريخ_جلسة','-')}</td>
                        <td style='color:{حالة_لون}; font-weight:900'>{case.get('حالة','متداولة')}</td>
                        <td style='color:#4DA8DA; font-weight:900'>{مكان}</td>
                    </tr>
                    </table>
                    """, unsafe_allow_html=True)

                    # التفاصيل
                    st.markdown("<div style='background:#142038; padding:20px; border-radius:12px; border:2px solid #4DA8DA; margin-top:10px'>", unsafe_allow_html=True)
                    st.markdown(f"<div style='color:#4DA8DA; font-size:20px; font-weight:900; text-align:center; margin-bottom:15px'>📄 تفاصيل كاملة - {رقم_كامل}</div>", unsafe_allow_html=True)
                    
                    st.markdown("<div style='color:#FFD700; font-size:16px; font-weight:900; margin-bottom:10px'>1- البيانات الاساسية</div>", unsafe_allow_html=True)
                    st.markdown(f"""
                    <table style='width:100%; border-collapse:collapse; margin-bottom:15px;'>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; width:30%; font-weight:900;'>نوع الدعوى</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('نوع')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>المحكمة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('محكمة_اسم')} {f'- مأمورية {case.get("مأمورية")}' if case.get('مأمورية') else ''}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الدائرة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('دائرة')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الموضوع</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('موضوع')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{لقب1}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{لقب2}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي_عليه')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الحالة الحالية</td><td style='background:#FFF; color:{حالة_لون}; padding:8px; font-weight:900;'>{case.get('حالة','متداولة')} - {مكان}</td></tr>
                    </table>
                    """, unsafe_allow_html=True)

                    st.markdown("<div style='color:#FFD700; font-size:16px; font-weight:900; margin-bottom:10px'>2- اخر اجراء</div>", unsafe_allow_html=True)
                    st.markdown(f"""
                    <table style='width:100%; border-collapse:collapse; margin-bottom:15px;'>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; width:30%; font-weight:900;'>تاريخ اخر جلسة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('تاريخ_جلسة','-')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الاجراء</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('الاجراء','-')}</td></tr>
                    </table>
                    """, unsafe_allow_html=True)

                    if case.get('جلسات'):
                        st.markdown("<div style='color:#FFD700; font-size:16px; font-weight:900; margin-bottom:10px'>3- سجل الجلسات</div>", unsafe_allow_html=True)
                        جلسات_مرتبة = sorted(case['جلسات'], key=lambda x: x.get("تاريخ",""), reverse=True)
                        for ج in جلسات_مرتبة:
                            st.markdown(f"<div style='background:#1E2A47; padding:10px; border-radius:8px; margin-bottom:5px; border:1px solid #D4AF37'><b style='color:#FFD700'>تاريخ:</b> <span style='color:#FFF'>{ج.get('تاريخ')}</span> | <b style='color:#FFD700'>الاجراء:</b> <span style='color:#FFF'>{ج.get('الاجراء')}</span></div>", unsafe_allow_html=True)
                    else:
                        st.info("لا يوجد سجل جلسات مسجل")

                    if case.get('مستندات'):
                        st.markdown("<div style='color:#FFD700; font-size:16px; font-weight:900; margin:15px 0 10px 0'>4- المستندات المرفقة</div>", unsafe_allow_html=True)
                        for i, مستند in enumerate(case['مستندات']):
                            # <--- التعديل هنا عشان الاسم الجديد
                            اسم_الملف = مستند.get('name', مستند.get('نوع', f'ملف_{i}'))
                            محتوى_الملف = مستند.get('data', مستند.get('محتوى', ''))
                            if محتوى_الملف:
                                file_data = base64.b64decode(محتوى_الملف)
                                st.download_button(f"📥 تحميل {اسم_الملف}", data=file_data, file_name=اسم_الملف, key=f"dl_search_{case['id']}_{i}", use_container_width=True)
                    else:
                        st.info("لا يوجد مستندات مرفقة")
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                    st.markdown("</div>", unsafe_allow_html=True)
# ============ مركز التنبيهات =============
# =========== مركز التنبيهات ====================
elif st.session_state.page == "التنبيهات":
    st.markdown("<h1 style='text-align:center; color:#C9A961;'>مركز التنبيهات</h1>", unsafe_allow_html=True)
    
    if st.button("⬅️ العودة للرئيسية", use_container_width=True):
        st.session_state.page = "الرئيسية"
        st.rerun()

    st.divider()
    
    with st.container(border=True):
        st.markdown("<h2 style='text-align:center; color:#C9A961;'>ارسال التنبيهات بالايميل</h2>", unsafe_allow_html=True)
        
        # نجيب الايميل المحفوظ من التوكنز
        tokens = load_tokens()
        username = st.session_state.user["username"] if st.session_state.user else ""
        user_token = next((t for t in tokens["tokens"] if t["username"] == username), {})
        user_email = user_token.get("email", "")

        user_email_input = st.text_input("سجل ايميلك عشان يجيلك التنبيهات", value=user_email, key="alert_email_input")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("1. حفظ الايميل", use_container_width=True):
                if user_email_input and "@" in user_email_input:
                    # نحفظ الايميل في التوكنز
                    found = False
                    for t in tokens["tokens"]:
                        if t["username"] == username:
                            t["email"] = user_email_input
                            found = True
                            break
                    if not found:
                        tokens["tokens"].append({"username": username, "email": user_email_input})
                    save_tokens(tokens)
                    st.success(f"✅ تم حفظ الايميل: {user_email_input}")
                    st.rerun()
                else:
                    st.warning("دخل ايميل صحيح")

        with col2:
            if st.button("2. 📧 ارسل التنبيهات دلوقتي", use_container_width=True):
                if user_email_input:
                    alerts = get_alert_cases()
                    body = "<div style='direction:rtl; text-align:right; font-family:Arial;'>"
                    body += "<h2 style='color:#C9A961; text-align:center;'>تنبيهات القضايا</h2>"
                    
                    body += "<h3 style='color:#FFD700;'>1. جلسات خلال 7 ايام</h3>"
                    if alerts["sessions"]:
                        for case in alerts["sessions"]:
                            رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
                            body += f"<p style='border:1px solid #C9A961; padding:10px; border-radius:8px;'>"
                            body += f"<b>رقم:</b> {رقم_كامل} <br> "
                            body += f"<b>الموضوع:</b> {case.get('موضوع','')} <br> "
                            body += f"<b>الجلسة:</b> {case.get('تاريخ_جلسة','')} - <b style='color:red;'>فاضل {case.get('days_left',0)} يوم</b>"
                            body += f"</p>"
                    else:
                        body += "<p>✅ لا توجد جلسات قريبة</p>"

                    body += "<h3 style='color:#FF4500;'>2. طعون خلال 15 يوم</h3>"
                    if alerts["appeals"]:
                        for case in alerts["appeals"]:
                            رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
                            body += f"<p style='border:1px solid #FF4500; padding:10px; border-radius:8px;'>"
                            body += f"<b>رقم:</b> {رقم_كامل} <br> "
                            body += f"<b>الموضوع:</b> {case.get('موضوع','')} <br> "
                            body += f"<b style='color:red;'>اخر ميعاد للطعن: {case.get('deadline','')} - فاضل {case.get('days_left_appeal',0)} يوم</b>"
                            body += f"</p>"
                    else:
                        body += "<p>✅ لا توجد طعون قريبة</p>"
                    body += "</div>"

                    if send_email(user_email_input, "تنبيهات القضايا من النظام", body):
                        st.success("✅ تم ارسال التنبيهات بنجاح للايميل")
                else:
                    st.error("❌ سجل الايميل الاول من الزرار اللي جنبه")

    st.divider()
    
    # نجيب التنبيهات بتاعت المستخدم بس
    alerts = get_alert_cases()
    st.markdown(f"<h3 style='text-align:center; color:#C9A961;'>التنبيهات الموجوده حاليا</h3>", unsafe_allow_html=True)

    # الجلسات
    st.markdown("<h2 style='text-align:center; color:#C9A961;'>الجلسات خلال 7 ايام</h2>", unsafe_allow_html=True)
    if alerts["sessions"]:
        for case in alerts["sessions"]:
            رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
            with st.container(border=True):
                st.write(f"**رقم القضية:** {رقم_كامل}")
                st.write(f"**الموضوع:** {case.get('موضوع','')}")
                st.write(f"**تاريخ الجلسة:** {case.get('تاريخ_جلسة','')}")
                st.write(f"**فاضل:** {case.get('days_left',0)} يوم")
                if st.button("📄 فتح القضية", key=f"open_alert_s_{case['id']}"):
                    st.session_state.selected_case_id = case['id']
                    st.session_state.page = "تفاصيل"
                    st.rerun()
    else:
        st.info("لا توجد جلسات خلال 7 ايام")

    # الطعون
    st.markdown("<h2 style='text-align:center; color:#C9A961;'>الطعون خلال 15 يوم</h2>", unsafe_allow_html=True)
    if alerts["appeals"]:
        for case in alerts["appeals"]:
            رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
            with st.container(border=True):
                st.write(f"**رقم القضية:** {رقم_كامل}")
                st.write(f"**الموضوع:** {case.get('موضوع','')}")
                st.write(f"**اخر ميعاد للطعن:** {case.get('deadline','')}")
                st.write(f"**فاضل:** {case.get('days_left_appeal',0)} يوم")
                if st.button("📄 فتح القضية", key=f"open_alert_a_{case['id']}"):
                    st.session_state.selected_case_id = case['id']
                    st.session_state.page = "تفاصيل"
                    st.rerun()
    else:
        st.info("لا توجد طعون خلال 15 يوم")
        
    st.divider()
        # ================================================
# ============ صفحة المكتبة القانونية ============
elif st.session_state.page == "مكتبة":
    data = load_data()
    st.markdown('<h1 style="text-align: center; color: #FFD700;">المكتبة 📚<br>القانونية</h1>', unsafe_allow_html=True)
    
    if st.button("⬅️ العودة للصفحة الرئيسية", use_container_width=True):
        st.session_state.page = "الرئيسية"
        for k in ["selected_section", "search_filters", "show_upload"]:
            st.session_state.pop(k, None)
        st.rerun()

    # فلترة: كل واحد يشوف بتاعه بس. الادمن يشوف الكل
    library_data = data.get("library", [])
    if st.session_state.user["role"] == "admin":
        my_library = library_data
    else:
        my_library = [f for f in library_data if f.get("user_id") == st.session_state.user["id"]]

    # 1. البحث
    st.markdown("### 🔍 بحث سريع")
    col1, col2 = st.columns([4,1])
    with col1: 
        search_query = st.text_input("اكتب اسم او رقم او سنة", key="search_q_lib", placeholder="مثال: 148 او تأمينات او 2019")
    with col2: 
        st.write("")
        st.write("")
        if st.button("🔍 بحث", use_container_width=True, key="btn_search_lib"):
            if search_query.strip():
                st.session_state.search_filters = {"q": search_query.strip()}
                st.session_state.pop("selected_section", None)
                st.rerun()
            else:
                st.warning("اكتب كلمة للبحث")

    st.divider()

    # 2. الاقسام
    st.markdown("### 📁 الاقسام")
    cols = st.columns(4)
    for i, (section, color) in enumerate(LIBRARY_SECTIONS.items()):
        with cols[i % 4]:
            count = len([f for f in my_library if str(f.get("section")).strip() == str(section).strip()])
            
            if st.button(f"{section}", key=f"btn_section_{i}", use_container_width=True, help=f"عدد الملفات: {count}"):
                st.session_state.selected_section = section
                st.session_state.pop("search_filters", None)
                st.rerun()
            
            st.markdown(f"<p style='text-align:center; color:{color}; font-size:12px; margin-top:-10px;'>({count} ملف)</p>", unsafe_allow_html=True)

    st.divider()
    
    # 3. تحديد ايه اللي هنعرضه
    files_to_show = []
    title = ""
    
    if "selected_section" in st.session_state:
        sec = st.session_state.selected_section
        title = f"📂 القسم: {sec}"
        files_to_show = [f for f in my_library if str(f.get("section")).strip() == str(sec).strip()]
        
    elif "search_filters" in st.session_state:
        q = st.session_state.search_filters["q"].lower()
        title = f"🔍 نتائج البحث عن: '{q}'"
        files_to_show = [
            item for item in my_library 
            if q in item.get("name","").lower() 
            or q in item.get("number","").lower() 
            or q in item.get("year","").lower()
            or q in item.get("section","").lower()
        ]
    else:
        title = "📂 اختار قسم من فوق"
        files_to_show = []

    if title: st.subheader(title)

    # 4. زر الاضافة
    if st.button("➕ اضافة مادة قانونية لمكتبتك", key="btn_add_lib", type="primary", use_container_width=True):
        st.session_state.show_upload = True

    if st.session_state.get("show_upload", False):
        with st.form("form_add_doc_lib"):
            section_select = st.selectbox("1- اختر القسم", list(LIBRARY_SECTIONS.keys()), key="sel_section")
            doc_name = st.text_input("2- اسم المستند", placeholder="مثال: قانون التأمينات 148 لسنة 2019", key="inp_name")
            col1, col2 = st.columns(2)
            with col1: doc_number = st.text_input("3- الرقم", placeholder="148", key="inp_num")
            with col2: doc_year = st.text_input("4- السنة", placeholder="2019", key="inp_year")
            uploaded_file = st.file_uploader("5- ارفع الملف", type=['pdf', 'doc', 'docx', 'jpg', 'png'], key="upl_file")
            
            if st.form_submit_button("💾 حفظ في القسم", use_container_width=True):
                if uploaded_file and doc_name.strip() and section_select:
                    file_base64 = base64.b64encode(uploaded_file.getvalue()).decode('utf-8')
                    new_doc = {
                        "id": secrets.token_hex(8),
                        "user_id": st.session_state.user["id"],
                        "name": doc_name.strip(), 
                        "section": section_select.strip(),
                        "number": doc_number.strip(),
                        "year": doc_year.strip(),
                        "file_type": uploaded_file.name.split('.')[-1],
                        "content": file_base64
                    }
                    data.setdefault("library", []).append(new_doc)
                    save_data(data) # <--- بيحفظ في السحابة والمحلي
                    st.success(f"✅ تم حفظ '{doc_name}' في قسم '{section_select}'")
                    st.session_state.show_upload = False
                    st.session_state.selected_section = section_select
                    st.rerun()
                else:
                    st.error("❌ لازم تختار قسم + اسم + ملف")

    # 5. عرض الملفات
    if files_to_show:
        st.write(f"عدد الملفات: {len(files_to_show)}")
        for doc in files_to_show:
            color = LIBRARY_SECTIONS.get(doc.get("section"), "#7F8C8D")
            st.markdown(f"<div style='border-right:5px solid {color}; padding:12px; margin:8px 0; background:#1e1e1e; border-radius:8px;'>", unsafe_allow_html=True)
            st.markdown(f"**{doc.get('name')}**")
            st.caption(f"رقم: {doc.get('number','-')} | سنة: {doc.get('year','-')}")
            
            col1, col2 = st.columns([3,1])
            with col1:
                if doc.get("content"):
                    file_data = base64.b64decode(doc["content"])
                    file_name = f"{doc.get('name')}.{doc.get('file_type')}"
                    mime_type = {
                        "pdf":"application/pdf",
                        "doc":"application/msword",
                        "docx":"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "jpg":"image/jpeg",
                        "png":"image/png"
                    }.get(doc.get('file_type'),"application/octet-stream")
                    st.download_button("⬇️ تحميل", data=file_data, file_name=file_name, mime=mime_type, key=f"dl_{doc['id']}", use_container_width=True)
            
            with col2:
                if doc.get("user_id") == st.session_state.user["id"] or st.session_state.user["role"] == "admin":
                    if st.button("🗑️ حذف", key=f"del_{doc['id']}", use_container_width=True):
                        data["library"] = [d for d in data["library"] if d["id"] != doc["id"]]
                        save_data(data) # <--- حذف من السحابة
                        st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
    
    elif "selected_section" in st.session_state:
        st.warning(f"القسم '{st.session_state.selected_section}' فاضي. ارفع اول ملف من الزرار اللي فوق")
    elif "search_filters" in st.session_state:
        st.info("مفيش نتائج للبحث ده")
    else:
        st.info("اختار قسم من الازرار اللي فوق عشان تشوف الملفات")
        # ============================================
# ================================================
# =========== الجزء الثامن: التقارير ============
# ================= القسم 1 من 2 =================

def get_export_html(content_html, title):
    return f"""<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
    @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;700;900&display=swap');
    body {{ font-family: "Cairo", "Times New Roman", serif; direction: rtl; background: white; color: black; padding: 20px; }}
.case-table {{width:100%; border-collapse:collapse; font-size:12px; margin-top:12px;}}
.case-table th {{ background: #B8860B; color:#000; padding:8px; border:1px solid #8B6914; font-weight:900; text-align:center; }}
.case-table td {{padding:6px; border:1px solid #B8860B; text-align:center; background:#fff; color:#000;}}
.table-container {{overflow-x:auto}}
    h1, h2, h3 {{ color: #8B6914!important; text-align: center; }}
</style>
</head>
<body>
{content_html}
</body>
</html>"""

if st.session_state.page == "تقارير":
    import io
    import pandas as pd
    from datetime import datetime
    data = load_data()

    # فلترة حسب المستخدم
    username = st.session_state.user["username"] if st.session_state.user else ""
    user_role = st.session_state.user["role"] if st.session_state.user else "member"
    if user_role == "admin":
        all_cases = data["cases"]
    else:
        all_cases = [c for c in data["cases"] if c.get('assigned_to') == username]

    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;700;900&display=swap');
.case-table {width:100%; border-collapse:collapse; font-size:10px; margin-top:12px; direction:rtl; font-family: "Cairo", sans-serif;}
.case-table th {background: linear-gradient(135deg, #8B6914 0%, #B8860B 100%); color:#000; padding:5px; border:1px solid #8B6914; font-weight:900; text-align:center; white-space: nowrap; font-size:10px;}
.case-table td {padding:4px; border:1px solid #B8860B; text-align:center; background:#1E2A47; color:#fff; font-size:10px;}
.table-container {overflow-x:auto}
    h1, h2, h3, h4, label {color: #B8860B!important; font-family: "Cairo", sans-serif;}
.stButton>button {border: 2px solid #B8860B!important; color: #B8860B!important; font-family: "Cairo", sans-serif; font-size:13px;}
.stButton>button:hover {background: #B8860B!important; color: #000!important;}
    [data-testid="stTab"] button {color: #B8860B!important; font-family: "Cairo", sans-serif; font-size:13px;}
    [data-testid="stTab"] button[aria-selected="true"] {border-bottom: 3px solid #B8860B!important;}
    </style>
    """, unsafe_allow_html=True)

    st.markdown("<h2 style='color:#B8860B; text-align:center; font-family: Cairo; font-size:18px;'>📑 مركز التقارير الحكومية</h2>", unsafe_allow_html=True)
    if st.button("⬅️ العودة للرئيسية", use_container_width=True): st.session_state.page = "الرئيسية"; st.rerun()

    if 'last_report_html' not in st.session_state: st.session_state.last_report_html = ""; st.session_state.last_report_title = "تقرير"
    if 'last_report_df' not in st.session_state: st.session_state.last_report_df = pd.DataFrame()

    tab1, tab2, tab3, tab4 = st.tabs(["📊 بيان بجميع الدعاوى المتداولة", "⚖️ بيان الاحكام", "📈 الإحصائيات", "📊 بيان عددي"])

    def report_header(region, title, مدير_عام, مدير_ادارة, عضو_قانوني):
        return f"""<div style="text-align:right; color:#B8860B; border:2px double #B8860B; padding:12px 10px; background: #0A1428; border-radius:5px; margin-bottom:12px;">
        <h2 style="margin:2px 0; font-size:15px; font-weight:900;">الهيئة القومية للتأمين الاجتماعي</h2>
        <h3 style="margin:1px 0; font-size:12px; font-weight:700;">الإدارة المركزية للإدارات القانونية</h3>
        <h3 style="margin:4px 0; font-size:12px; font-weight:700;">ديوان عام منطقة {region}</h3>
        <hr style="border:1px solid #B8860B; margin:8px 0;">
        <h3 style="margin:6px 0; font-size:13px; font-weight:900; text-align:center; text-decoration: underline;"> {title} </h3>
        </div>"""

    # ========= تبويب 1: الدعاوى المتداولة =========
    with tab1:
        st.markdown("<div style='background:#1E2A47; padding:12px; border-radius:8px; border:2px solid #B8860B; margin-bottom:12px'>", unsafe_allow_html=True)
        نوع_تقرير_متداولة = st.selectbox("نوع البيان", ["بيان بجميع الدعاوى المتداولة", "بيان بالدعاوى المتداولة حسب موضوع الدعوى"], key="no1_tقرير")
        colA, colB, colC = st.columns(3)
        with colA: region = st.text_input("ديوان عام منطقة", key="region1")
        with colB: مدير_عام1 = st.text_input("اسم مدير عام الإدارات القانونية", key="modir1")
        with colC: مدير_ادارة1 = st.text_input("اسم مدير إدارة القضايا", key="modir_idara1")
        عضو_قانوني1 = st.text_input("اسم العضو القانوني", key="odo1")
        col1, col2, col3 = st.columns(3)
        with col1: from_date = st.date_input("من الفترة", key="from1")
        with col2: to_date = st.date_input("حتى الفترة", key="to1")
        with col3: lawyer = st.text_input("طرف الاستاذ/ المحامي", key="lawyer1")
        topic = ""
        if "حسب موضوع" in نوع_تقرير_متداولة: topic = st.text_input("موضوع الدعوى للفلترة", key="topic1")
        st.markdown("</div>", unsafe_allow_html=True)

        if st.button("🔍 عرض بيان الدعاوى المتداولة", use_container_width=True, type="primary", key="show1"):
            cases = [c for c in all_cases if str(c.get('حالة','')).strip() == 'متداولة']
            فلترة_بالمدة = []
            for c in cases:
                if c.get('تاريخ_جلسة'):
                    try:
                        ت_جلسة = datetime.strptime(c['تاريخ_جلسة'], '%Y-%m-%d').date()
                        if from_date <= ت_جلسة <= to_date: فلترة_بالمدة.append(c)
                    except: pass
            cases = فلترة_بالمدة
            if "حسب موضوع" in نوع_تقرير_متداولة and topic: cases = [c for c in cases if topic.lower() in str(c.get('موضوع','')).lower()]
            cases = sorted(cases, key=lambda x: x.get("تاريخ_جلسة","0000-00-00"), reverse=True)
            title = f"{نوع_تقرير_متداولة} خلال الفترة من {from_date.strftime('%d-%m-%Y')} إلى {to_date.strftime('%d-%m-%Y')} طرف الاستاذ/ {lawyer} المحامي"
            header_html = report_header(region, title, مدير_عام1, مدير_ادارة1, عضو_قانوني1)
            if not cases: st.warning(f"⚠️ لا توجد دعاوى متداولة في الفترة")
            else:
                st.success(f"✅ تم العثور على {len(cases)} دعوى متداولة")
                html = "<table class='case-table'><tr><th>م</th><th>رقم</th><th>سنة</th><th>المحكمة والدائرة</th><th>الخصوم</th><th>الموضوع</th><th>اخر جلسة</th><th>الاجراء</th><th>ملاحظات</th></tr>"
                df_data = []
                for idx, c in enumerate(cases, 1):
                    محكمة = f"{c.get('نوع','')} {c.get('محكمة_اسم','')}"
                    if c.get('مأمورية'): محكمة += f" - مأمورية {c.get('مأمورية')}"
                    if c.get('دائرة'): محكمة += f" - دائرة {c.get('دائرة')}"
                    مدعي = c.get('مدعي',''); مدعي_عليه = c.get('مدعي_عليه','')
                    خصوم_html = f"<span style='color:#dc3545; font-weight:900'>{مدعي} ضد {مدعي_عليه}</span>" if "الهيئة" in str(مدعي) else f"{مدعي} ضد {مدعي_عليه}"
                    html += f"<tr><td>{idx}</td><td>{c.get('رقم','')}</td><td>{c.get('سنة','')}</td><td>{محكمة}</td><td>{خصوم_html}</td><td>{c.get('موضوع','')}</td><td><b style='color:#B8860B'>{c.get('تاريخ_جلسة','')}</b></td><td>{c.get('الاجراء','')}</td><td>{c.get('ملاحظات','')}</td></tr>"
                    df_data.append({'م': idx, 'رقم': c.get('رقم',''), 'سنة': c.get('سنة',''), 'المحكمة': محكمة, 'الخصوم': f"{مدعي} ضد {مدعي_عليه}", 'الموضوع': c.get('موضوع',''), 'اخر جلسة': c.get('تاريخ_جلسة',''), 'الاجراء': c.get('الاجراء',''), 'ملاحظات': c.get('ملاحظات','')})
                html += "</table>"
                footer = f"""<div style="margin-top:25px; color:#B8860B; font-size:12px;"><p style="text-align:center; margin-bottom:20px; font-size:13px; font-weight:700;">تفضلوا بقبول وافر الاحترام والتقدير،</p><table style="width:100%;"><tr><td style="width:50%; text-align:right; vertical-align:top;"><div style="font-weight:900;">العضو القانوني</div><div>{عضو_قانوني1}</div><div style="margin-top:12px;">....................</div><div style="margin-top:20px;">تحر في: {datetime.now().strftime('%d-%m-%Y')}</div></td><td style="width:50%; text-align:left; vertical-align:top;"><div style="font-weight:900;">مدير إدارة القضايا</div><div>{مدير_ادارة1}</div><div style="margin-top:12px;">....................</div></td></tr></table><div style="text-align:center; margin-top:20px;"><div style="font-weight:900; color:#dc3545;">مدير عام الإدارات القانونية</div><div>{مدير_عام1}</div><div style="margin-top:12px;">....................</div></div></div>"""
                full_html = header_html + f"<div class='table-container'>{html}</div>" + footer
                st.markdown(full_html, unsafe_allow_html=True)
                st.session_state.last_report_html = full_html; st.session_state.last_report_title = f"بيان_الدعاوى_المتداولة_{region}"; st.session_state.last_report_df = pd.DataFrame(df_data)
                html_export = get_export_html(full_html, st.session_state.last_report_title)
                c1,c2,c3 = st.columns(3)
                with c1: st.download_button("⬇️ PDF", data=html_export.encode('utf-8'), file_name=f"بيان_الدعاوى_{region}.html", mime="text/html", use_container_width=True, key="dl1")
                with c2: st.download_button("⬇️ Word", data=html_export.encode('utf-8'), file_name=f"بيان_الدعاوى_{region}.doc", mime="application/msword", use_container_width=True, key="dl2")
                with c3:
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                        df = pd.DataFrame(df_data)
                        df.to_excel(writer, index=False, sheet_name='Sheet1')
                        worksheet = writer.sheets['Sheet1']
                        for col in worksheet.columns: worksheet.column_dimensions[col[0].column_letter].width = 25
                        for cell in worksheet[1]: cell.alignment = cell.alignment.copy(wrap_text=True, horizontal='right')
                    st.download_button("⬇️ Excel", data=excel_buffer.getvalue(), file_name=f"بيان_الدعاوى_{region}.xlsx", use_container_width=True, key="dlx1")

    # ========= تبويب 2: الاحكام =========
    with tab2:
        st.markdown("<div style='background:#1E2A47; padding:12px; border-radius:8px; border:2px solid #B8860B; margin-bottom:12px'>", unsafe_allow_html=True)
        نوع_التقرير = st.selectbox("نوع البيان", ["بيان بجميع الاحكام الصادرة للصالح والضد", "بيان بالاحكام الصادرة للصالح", "بيان بالاحكام الصادرة للضد", "بيان بالاحكام الصادرة حسب موضوع الدعوى"], key="no3_tقرير")
        colA, colB, colC = st.columns(3)
        with colA: region2 = st.text_input("ديوان عام منطقة", key="region2")
        with colB: مدير_عام2 = st.text_input("اسم مدير عام الإدارات القانونية", key="modir2")
        with colC: مدير_ادارة2 = st.text_input("اسم مدير إدارة القضايا", key="modir_idara2")
        عضو_قانوني2 = st.text_input("اسم العضو القانوني", key="odo2")
        col1, col2, col3 = st.columns(3)
        with col1: from_date2 = st.date_input("من الفترة", key="from2")
        with col2: to_date2 = st.date_input("حتى الفترة", key="to2")
        with col3: lawyer2 = st.text_input("طرف الاستاذ/ المحامي", key="lawyer2")
        topic2 = ""
        if "حسب موضوع" in نوع_التقرير: topic2 = st.text_input("موضوع الدعوى للفلترة", key="topic2")
        st.markdown("</div>", unsafe_allow_html=True)

        if st.button("🔍 عرض بيان الاحكام", use_container_width=True, type="primary", key="show2"):
            archive = [c for c in all_cases if c.get("حالة") == "منتهية" and not c.get("تم_الحفظ_النهائي")]
            cases = []
            for c in archive:
                if c.get('تاريخ_الحكم'):
                    try:
                        ت_حكم = datetime.strptime(c['تاريخ_الحكم'], '%Y-%m-%d').date()
                        if from_date2 <= ت_حكم <= to_date2: cases.append(c)
                    except: pass
            if نوع_التقرير == "بيان بالاحكام الصادرة للصالح": cases = [c for c in cases if c.get('مسندة_ل_الحكم') == 'الصالح']
            elif نوع_التقرير == "بيان بالاحكام الصادرة للضد": cases = [c for c in cases if c.get('مسندة_ل_الحكم') == 'الضد']
            elif "حسب موضوع" in نوع_التقرير and topic2: cases = [c for c in cases if topic2.lower() in str(c.get('موضوع','')).lower()]
            cases = sorted(cases, key=lambda x: x.get("تاريخ_الحكم","0000-00-00"), reverse=True)
            title = f"{نوع_التقرير} خلال الفترة من {from_date2.strftime('%d-%m-%Y')} إلى {to_date2.strftime('%d-%m-%Y')} طرف الاستاذ/ {lawyer2} المحامي"
            header_html = report_header(region2, title, مدير_عام2, مدير_ادارة2, عضو_قانوني2)
            if not cases: st.warning(f"⚠️ لا توجد احكام في الفترة")
            else:
                st.success(f"✅ تم العثور على {len(cases)} حكم")
                html = "<table class='case-table'><tr><th>م</th><th>رقم</th><th>سنة</th><th>المحكمة والدائرة</th><th>الخصوم</th><th>الموضوع</th><th>تاريخ الحكم</th><th>منطوق الحكم</th><th>مسندة ل</th><th>ملاحظات</th></tr>"
                df_data = []
                for idx, c in enumerate(cases, 1):
                    محكمة = f"{c.get('نوع','')} {c.get('محكمة_اسم','')}"
                    if c.get('مأمورية'): محكمة += f" - مأمورية {c.get('مأمورية')}"
                    if c.get('دائرة'): محكمة += f" - دائرة {c.get('دائرة')}"
                    مدعي = c.get('مدعي',''); مدعي_عليه = c.get('مدعي_عليه','')
                    خصوم_html = f"<span style='color:#dc3545; font-weight:900'>{مدعي} ضد {مدعي_عليه}</span>" if "الهيئة" in str(مدعي) else f"{مدعي} ضد {مدعي_عليه}"
                    لون_مسندة = "#28a745" if c.get('مسندة_ل_الحكم') == 'الصالح' else "#dc3545"
                    html += f"<tr><td>{idx}</td><td>{c.get('رقم','')}</td><td>{c.get('سنة','')}</td><td>{محكمة}</td><td>{خصوم_html}</td><td>{c.get('موضوع','')}</td><td><b style='color:#B8860B'>{c.get('تاريخ_الحكم','')}</b></td><td>{c.get('منطوق_الحكم','')}</td><td><b style='color:{لون_مسندة}'>{c.get('مسندة_ل_الحكم')}</b></td><td>{c.get('ملاحظات','')}</td></tr>"
                    df_data.append({'م': idx, 'رقم': c.get('رقم',''), 'سنة': c.get('سنة',''), 'المحكمة': محكمة, 'الخصوم': f"{مدعي} ضد {مدعي_عليه}", 'الموضوع': c.get('موضوع',''), 'تاريخ الحكم': c.get('تاريخ_الحكم',''), 'منطوق الحكم': c.get('منطوق_الحكم',''), 'مسندة ل': c.get('مسندة_ل_الحكم'), 'ملاحظات': c.get('ملاحظات','')})
                html += "</table>"
                footer = f"""<div style="margin-top:25px; color:#B8860B; font-size:12px;"><p style="text-align:center; margin-bottom:20px; font-size:13px; font-weight:700;">تفضلوا بقبول وافر الاحترام والتقدير،</p><table style="width:100%;"><tr><td style="width:50%; text-align:right; vertical-align:top;"><div style="font-weight:900;">العضو القانوني</div><div>{عضو_قانوني2}</div><div style="margin-top:12px;">....................</div><div style="margin-top:20px;">تحرر في: {datetime.now().strftime('%d-%m-%Y')}</div></td><td style="width:50%; text-align:left; vertical-align:top;"><div style="font-weight:900;">مدير إدارة القضايا</div><div>{مدير_ادارة2}</div><div style="margin-top:12px;">....................</div></td></tr></table><div style="text-align:center; margin-top:20px;"><div style="font-weight:900; color:#dc3545;">مدير عام الإدارات القانونية</div><div>{مدير_عام2}</div><div style="margin-top:12px;">....................</div></div></div>"""
                full_html = header_html + f"<div class='table-container'>{html}</div>" + footer
                st.markdown(full_html, unsafe_allow_html=True)
                st.session_state.last_report_html = full_html; st.session_state.last_report_title = f"بيان_الاحكام_{region2}"; st.session_state.last_report_df = pd.DataFrame(df_data)
                html_export = get_export_html(full_html, st.session_state.last_report_title)
                c1,c2,c3 = st.columns(3)
                with c1: st.download_button("⬇️ PDF", data=html_export.encode('utf-8'), file_name=f"بيان_الاحكام_{region2}.html", mime="text/html", use_container_width=True, key="dl21")
                with c2: st.download_button("⬇️ Word", data=html_export.encode('utf-8'), file_name=f"بيان_الاحكام_{region2}.doc", mime="application/msword", use_container_width=True, key="dl22")
                with c3:
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                        df = pd.DataFrame(df_data)
                        df.to_excel(writer, index=False, sheet_name='Sheet1')
                        worksheet = writer.sheets['Sheet1']
                        for col in worksheet.columns: worksheet.column_dimensions[col[0].column_letter].width = 25
                        for cell in worksheet[1]: cell.alignment = cell.alignment.copy(wrap_text=True, horizontal='right')
                    st.download_button("⬇️ Excel", data=excel_buffer.getvalue(), file_name=f"بيان_الاحكام_{region2}.xlsx", use_container_width=True, key="dlx2")
                    # ================= القسم 2 من 2 =================
    # ========= تبويب 3: الاحصائيات =========
    with tab3:
        st.markdown("<h3 style='color:#B8860B; text-align:center; font-size:16px;'>📈 الإحصائيات العددية</h3>", unsafe_allow_html=True)
        st.markdown("<div style='background:#1E2A47; padding:12px; border-radius:8px; border:2px solid #B8860B; margin-bottom:12px'>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1: stat_from = st.date_input("من تاريخ", key="s1")
        with col2: stat_to = st.date_input("حتى تاريخ", key="s2")
        st.markdown("</div>", unsafe_allow_html=True)
        if st.button("استخراج الإحصائيات", use_container_width=True, type="primary"):
            متداولة = [c for c in all_cases if c.get('حالة') == 'متداولة' and c.get('تاريخ_جلسة') and stat_from <= datetime.strptime(c['تاريخ_جلسة'], '%Y-%m-%d').date() <= stat_to]
            احكام = [c for c in all_cases if c.get('حالة') == 'منتهية' and not c.get('تم_الحفظ_النهائي') and c.get('تاريخ_الحكم') and stat_from <= datetime.strptime(c['تاريخ_الحكم'], '%Y-%m-%d').date() <= stat_to]
            للصالح = [c for c in احكام if c.get('مسندة_ل_الحكم') == 'الصالح']
            للضد = [c for c in احكام if c.get('مسندة_ل_الحكم') == 'الضد']
            c1,c2,c3,c4 = st.columns(4)
            c1.metric("عدد القضايا المتداولة", len(متداولة))
            c2.metric("عدد الاحكام الصادرة", len(احكام))
            c3.metric("عدد الاحكام للصالح", len(للصالح))
            c4.metric("عدد الاحكام للضد", len(للضد))

    # ========= تبويب 4: البيان العددي =========
    with tab4:
        st.markdown("<h3 style='color:#B8860B; text-align:center; font-size:16px;'>📊 البيان العددي</h3>", unsafe_allow_html=True)
        st.markdown("<div style='background:#1E2A47; padding:12px; border-radius:8px; border:2px solid #B8860B; margin-bottom:12px'>", unsafe_allow_html=True)
        نوع_البيان_العددي = st.selectbox("نوع البيان العددي", ["جميع الدعاوى المتداولة", "الدعاوى المتداولة حسب موضوع الدعوى", "جميع الاحكام للصالح والضد", "الاحكام الصادرة للصالح", "الاحكام الصادرة للضد", "الاحكام حسب موضوع الدعوى"], key="no4_عددي")
        colA, colB, colC = st.columns(3)
        with colA: region_stat = st.text_input("ديوان عام منطقة", key="region_stat")
        with colB: مدير_عام_stat = st.text_input("اسم مدير عام الإدارات القانونية", key="modir_stat")
        with colC: مدير_ادارة_stat = st.text_input("اسم مدير إدارة القضايا", key="modir_idara_stat")
        عضو_قانوني_stat = st.text_input("اسم العضو القانوني", key="odo_stat")
        lawyer_stat = st.text_input("طرف الاستاذ/ المحامي", key="lawyer_stat")
        col1, col2 = st.columns(2)
        with col1: stat_from2 = st.date_input("من الفترة", key="s_from2")
        with col2: stat_to2 = st.date_input("حتى الفترة", key="s_to2")
        topic_stat = ""
        if "حسب موضوع" in نوع_البيان_العددي: topic_stat = st.text_input("موضوع الدعوى للفلترة", key="topic_stat")
        st.markdown("</div>", unsafe_allow_html=True)

        if st.button("🔍 عرض البيان العددي", use_container_width=True, type="primary", key="show_stat"):
            العدد = 0; اسم_البيان = نوع_البيان_العددي
            if "الدعاوى المتداولة" in نوع_البيان_العددي:
                cases = [c for c in all_cases if str(c.get('حالة','')).strip() == 'متداولة']
                for c in cases:
                    if c.get('تاريخ_جلسة'):
                        try:
                            ت_جلسة = datetime.strptime(c['تاريخ_جلسة'], '%Y-%m-%d').date()
                            if stat_from2 <= ت_جلسة <= stat_to2:
                                if "حسب موضوع" in نوع_البيان_العددي and topic_stat:
                                    if topic_stat.lower() in str(c.get('موضوع','')).lower(): العدد += 1
                                elif "جميع" in نوع_البيان_العددي: العدد += 1
                        except: pass
            elif "الاحكام" in نوع_البيان_العددي:
                archive = [c for c in all_cases if c.get("حالة") == "منتهية" and not c.get("تم_الحفظ_النهائي")]
                for c in archive:
                    if c.get('تاريخ_الحكم'):
                        try:
                            ت_حكم = datetime.strptime(c['تاريخ_الحكم'], '%Y-%m-%d').date()
                            if stat_from2 <= ت_حكم <= stat_to2:
                                مسندة = c.get('مسندة_ل_الحكم')
                                if "حسب موضوع" in نوع_البيان_العددي and topic_stat:
                                    if topic_stat.lower() in str(c.get('موضوع','')).lower(): العدد += 1
                                elif "للصالح" in نوع_البيان_العددي and مسندة == 'الصالح': العدد += 1
                                elif "للضد" in نوع_البيان_العددي and مسندة == 'الضد': العدد += 1
                                elif "للصالح والضد" in نوع_البيان_العددي: العدد += 1
                        except: pass
            st.success(f"✅ العدد الإجمالي: {العدد}")
            title_stat = f"بيان عددي ب{اسم_البيان} خلال الفترة من {stat_from2.strftime('%d-%m-%Y')} إلى {stat_to2.strftime('%d-%m-%Y')} طرف الاستاذ/ {lawyer_stat} المحامي"
            header_html = report_header(region_stat, title_stat, مدير_عام_stat, مدير_ادارة_stat, عضو_قانوني_stat)
            ملاحظات = f"عن موضوع: {topic_stat}" if topic_stat else ""
            html = f"<table class='case-table'><tr><th>م</th><th>البيان</th><th>العدد</th><th>ملاحظات</th></tr>"
            html += f"<tr><td>1</td><td>{اسم_البيان}</td><td><b style='color:#dc3545; font-size:14px'>{العدد}</b></td><td>{ملاحظات}</td></tr>"
            html += "</table>"
            footer = f"""<div style="margin-top:25px; color:#B8860B; font-size:12px;"><p style="text-align:center; margin-bottom:20px; font-size:13px; font-weight:700;">تفضلوا بقبول وافر الاحترام والتقدير،</p><table style="width:100%;"><tr><td style="width:50%; text-align:right; vertical-align:top;"><div style="font-weight:900;">العضو القانوني</div><div>{عضو_قانوني_stat}</div><div style="margin-top:12px;">....................</div><div style="margin-top:20px;">تحرر في: {datetime.now().strftime('%d-%m-%Y')}</div></td><td style="width:50%; text-align:left; vertical-align:top;"><div style="font-weight:900;">مدير إدارة القضايا</div><div>{مدير_ادارة_stat}</div><div style="margin-top:12px;">....................</div></td></tr></table><div style="text-align:center; margin-top:20px;"><div style="font-weight:900; color:#dc3545;">مدير عام الإدارات القانونية</div><div>{مدير_عام_stat}</div><div style="margin-top:12px;">....................</div></div></div>"""
            full_html = header_html + f"<div class='table-container'>{html}</div>" + footer
            st.markdown(full_html, unsafe_allow_html=True)
            df_data = pd.DataFrame([{'م': 1, 'البيان': اسم_البيان, 'العدد': العدد, 'ملاحظات': ملاحظات}])
            st.session_state.last_report_html = full_html; st.session_state.last_report_title = f"بيان_عددي_{region_stat}"; st.session_state.last_report_df = df_data
            html_export = get_export_html(full_html, st.session_state.last_report_title)
            c1,c2,c3 = st.columns(3)
            with c1: st.download_button("⬇️ PDF", data=html_export.encode('utf-8'), file_name=f"بيان_عددي_{region_stat}.html", mime="text/html", use_container_width=True, key="dl_stat1")
            with c2: st.download_button("⬇️ Word", data=html_export.encode('utf-8'), file_name=f"بيان_عددي_{region_stat}.doc", mime="application/msword", use_container_width=True, key="dl_stat2")
            with c3:
                excel_buffer = io.BytesIO()
                with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                    df = df_data
                    df.to_excel(writer, index=False, sheet_name='Sheet1')
                    worksheet = writer.sheets['Sheet1']
                    for col in worksheet.columns: worksheet.column_dimensions[col[0].column_letter].width = 25
                    for cell in worksheet[1]: cell.alignment = cell.alignment.copy(wrap_text=True, horizontal='right')
                st.download_button("⬇️ Excel", data=excel_buffer.getvalue(), file_name=f"بيان_عددي_{region_stat}.xlsx", use_container_width=True, key="dl_stat3")
# ========================= نهاية التقارير =====================
import json
import os
from datetime import datetime, date

BANNERS_FILE = "banners_v2.json"

def load_banners(): 
    try:
        with open(BANNERS_FILE, "r", encoding="utf-8") as f: return json.load(f)
    except: return []

def save_banners(banners):
    with open(BANNERS_FILE, "w", encoding="utf-8") as f: json.dump(banners, f, ensure_ascii=False, indent=2)
    try: supabase.table("banners").upsert({"id": 1, "data": banners}).execute()
    except: pass

# =========================================
# ============ صفحة إدارة اليافطات ============
# =========================================
if st.session_state.page == "إدارة اليافطات":

    try:
        res = supabase.table("banners").select("data").eq("id", 1).single().execute()
        st.session_state.banners = res.data["data"] if res.data and res.data["data"] else load_banners()
    except: st.session_state.banners = load_banners()
    
    users = load_users() # نجيب كل اليوزرز

    st.title("⚙️ إدارة اليافطات")
    
    if st.session_state.user["role"] != "admin":
        st.warning("⚠️ هذه الصفحة متاحة للادمن فقط")
        st.stop()

    # ========== 1. اضافة يافطة جديدة ==========
    with st.form("add_banner_form"):
        st.subheader("➕ اضافة يافطة جديدة")
        banner_text = st.text_area("نص اليافطة", placeholder="اكتب هنا نص الاعلان...", height=100)
        col1, col2 = st.columns(2)
        with col1: banner_color = st.color_picker("لون اليافطة", "#FFC107")
        with col2: banner_expire = st.date_input("تاريخ الانتهاء", value=date.today())

        st.markdown("### 👥 مين يشوف اليافطة دي؟")
        audience_type = st.radio("اختر الجمهور", ["الكل", "اعضاء محددين"], horizontal=True, key="audience_type")
        
        visible_to = []
        if audience_type == "اعضاء محددين":
            all_usernames = [u["username"] for u in users]
            visible_to = st.multiselect("حدد الاعضاء", all_usernames, key="visible_users")
            if not visible_to: st.warning("لازم تختار عضو واحد على الاقل")
        
        if st.form_submit_button("🚀 نشر اليافطة", use_container_width=True, type="primary"):
            if banner_text.strip() and (audience_type == "الكل" or visible_to):
                st.session_state.banners.append({
                    "text": banner_text, "color": banner_color, "expire": str(banner_expire),
                    "created_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "audience": audience_type, "visible_to": visible_to # <--- الجديد
                })
                save_banners(st.session_state.banners)
                st.success("✅ تم نشر اليافطة بنجاح!"); st.rerun()
            else: st.error("❌ املى كل الحقول")

    st.markdown("---")

    # ========== 2. حذف اليافطات ==========
    st.subheader(f"🗑️ اليافطات الحالية - العدد: {len(st.session_state.banners)}")
    for i, banner in enumerate(st.session_state.banners):
        with st.container(border=True):
            col1, col2 = st.columns([4,1])
            with col1:
                st.markdown(f'<div style="background:{banner["color"]};padding:15px;border-radius:10px;color:black;font-weight:bold;">📢 {banner["text"]}</div>', unsafe_allow_html=True)
                audience = "الكل" if banner.get("audience") == "الكل" else f"محدد: {', '.join(banner.get('visible_to', []))}"
                st.caption(f"📅 تنتهي: {banner.get('expire')} | 👥 تظهر لـ: {audience}")
            with col2:
                if st.button("🗑️ حذف", key=f"del_{i}", type="primary", use_container_width=True):
                    st.session_state.banners.pop(i); save_banners(st.session_state.banners); st.rerun()

    if st.button("⬅️ الرجوع للرئيسية", use_container_width=True):
        st.session_state.page = "الرئيسية"; st.rerun()
