
import json, os, bcrypt, smtplib, random, io, requests, base64, html, secrets
from datetime import datetime
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

import streamlit as st
import pandas as pd

from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import arabic_reshaper
from bidi.algorithm import get_display
from openpyxl.styles import Font, Alignment, PatternFill # <-- زودت PatternFill هنا


# ===== التصميم القضائي الجديد =====
st.markdown("""
<style>
/* صفحة الدخول */
.login-shell{
    max-width:760px;margin:45px auto;padding:28px;
    background:linear-gradient(145deg,#111c33,#1f3154);
    border:2px solid #D4AF37;border-radius:24px;
    box-shadow:0 20px 55px rgba(0,0,0,.45);
}
.login-shell h1,.login-shell h2,.login-shell p,.login-shell label{color:#fff!important;}
div[data-testid="stTabs"] button{
    color:#fff!important;font-weight:900!important;font-size:18px!important;
}
div[data-testid="stTabs"] button[aria-selected="true"]{
    color:#D4AF37!important;border-bottom:3px solid #D4AF37!important;
}
.login-shell input,.login-shell textarea{
    background:#fff!important;color:#111!important;
}
.admin-control-top{
    background:linear-gradient(135deg,#101b31,#1e2f50);
    border:2px solid #D4AF37;border-radius:20px;padding:14px;
    margin:8px 0 18px;box-shadow:0 10px 30px rgba(0,0,0,.28);
}
.judicial-card{
    min-height:118px;padding:20px 14px;border-radius:22px;
    border:2px solid #D4AF37;
    background:linear-gradient(145deg,#13213a,#213a61);
    box-shadow:0 12px 28px rgba(0,0,0,.30);
    text-align:center;margin-bottom:18px;
}
.judicial-card-title{color:#D4AF37;font-size:20px;font-weight:900;margin-top:8px;}
.judicial-card-icon{font-size:38px;line-height:1;}
.banner-admin-card{
    border:2px solid #D4AF37;border-radius:18px;padding:14px;
    background:#142038;margin-bottom:12px;
}
.status-active{color:#42d77d;font-weight:900}
.status-temp{color:#ffca28;font-weight:900}
.status-final{color:#ff5252;font-weight:900}
.incoming-messages-title{color:#D4AF37;font-size:24px;font-weight:900;text-align:center;margin:8px 0 14px;padding:10px;border-bottom:2px solid #D4AF37;}

/* إبراز الأيقونات والأزرار الملتصقة بخلفية الموقع */
button, [data-testid="stButton"] button, [data-testid="stDownloadButton"] button {
    color:#fff !important;
    font-weight:900 !important;
}
button svg, [data-testid="stButton"] svg, [data-testid="stDownloadButton"] svg, [data-testid="stFileUploader"] svg, [data-testid="stTabs"] svg, [data-testid="stExpander"] svg {
    color:#fff !important; fill:#fff !important; stroke:#fff !important; opacity:1 !important; visibility:visible !important;
}
button [data-testid="stIconMaterial"], button span, button p { opacity:1 !important; visibility:visible !important; }
[data-testid="stButton"] button, [data-testid="stDownloadButton"] button, [data-testid="stFormSubmitButton"] button {
    color:#fff !important; background:#142038 !important; border:1px solid #D4AF37 !important;
}
[data-testid="stTabs"] button { color:#fff !important; opacity:1 !important; visibility:visible !important; }
[data-testid="stFileUploader"] section * { color:#fff !important; opacity:1 !important; visibility:visible !important; }
input, textarea, [data-baseweb="select"], [data-testid="stFileUploader"] {color:#111!important;}
label, [data-testid="stCheckbox"] label, [data-testid="stRadio"] label, [data-testid="stMultiSelect"] label, [data-testid="stSelectbox"] label {color:#fff!important;font-weight:900!important;}
[data-testid="stCheckbox"] p, [data-testid="stRadio"] p {color:#fff!important;font-weight:800!important;}
[data-testid="stFileUploader"] section {background:#142038!important;border:1px solid #D4AF37!important;}
[data-testid="stButton"] button svg, [data-testid="stDownloadButton"] button svg, [data-testid="stTabs"] button svg, [data-testid="stExpander"] svg, [data-testid="stFileUploader"] svg { display:inline-block !important; visibility:visible !important; opacity:1 !important; color:#fff !important; fill:#fff !important; stroke:#fff !important; }
[data-testid="stButton"] button > div, [data-testid="stDownloadButton"] button > div { display:flex !important; align-items:center !important; justify-content:center !important; gap:.45rem !important; }

[data-testid="stFileUploader"] small, [data-testid="stFileUploader"] span {color:#fff!important;}
</style>
""", unsafe_allow_html=True)

def fix_arabic(text):
    """ بتظبط العربي عشان ميطلعش متقطع """
    if not text: return ""
    reshaped_text = arabic_reshaper.reshape(str(text))
    bidi_text = get_display(reshaped_text)
    return bidi_text
    # ===== نظام اليافطة ====
# ====== اليافطة نظام ====
from datetime import timedelta

LOCAL_DATA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")
os.makedirs(LOCAL_DATA_DIR, exist_ok=True)

def _secret(name, default=""):
    try:
        value = st.secrets.get(name, default)
        if value:
            return str(value)
    except Exception:
        pass
    return os.getenv(name, default)

SUPABASE_URL = _secret("SUPABASE_URL").rstrip("/")
SUPABASE_KEY = _secret("SUPABASE_KEY")
ADMIN_USERNAME = _secret("ADMIN_USERNAME", "admin")
ADMIN_DEFAULT_PASS = _secret("ADMIN_DEFAULT_PASS", "admin123")

SENDER_EMAIL = _secret("SENDER_EMAIL", "")
SENDER_PASSWORD = _secret("SENDER_PASSWORD", "")
APP_URL = _secret("APP_URL", "https://qpyqpsmkqcvdou4imbfunp.streamlit.app/")

CLOUD_ENABLED = bool(SUPABASE_URL and SUPABASE_KEY)

def _cloud_headers():
    return {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}", "Content-Type": "application/json", "Prefer": "resolution=merge-duplicates,return=minimal"}

def _cloud_get(key):
    if not CLOUD_ENABLED: return None
    try:
        r=requests.get(f"{SUPABASE_URL}/rest/v1/app_state", params={"key":f"eq.{key}","select":"key,payload,updated_at"}, headers=_cloud_headers(), timeout=12)
        r.raise_for_status(); rows=r.json(); return rows[0] if rows else None
    except Exception:
        return None

def _cloud_put(key,payload):
    if not CLOUD_ENABLED: return False
    try:
        r=requests.post(f"{SUPABASE_URL}/rest/v1/app_state", params={"on_conflict":"key"}, headers=_cloud_headers(), json={"key":key,"payload":payload,"updated_at":datetime.now().astimezone().isoformat()}, timeout=15)
        r.raise_for_status(); return True
    except Exception:
        return False

def load_json_persistent(key, local_path, default):
    local_obj=None
    try:
        with open(local_path,"r",encoding="utf-8") as f: local_obj=json.load(f)
    except Exception: pass
    cloud=_cloud_get(key)
    if cloud and cloud.get("payload") is not None:
        try:
            cloud_obj=cloud["payload"]
            local_obj = cloud_obj if local_obj is None else local_obj
            if local_obj is not None and os.path.exists(local_path):
                # cloud is authoritative when available; every save is mirrored locally and remotely
                local_obj=cloud_obj
            os.makedirs(os.path.dirname(local_path) or ".", exist_ok=True)
            with open(local_path,"w",encoding="utf-8") as f: json.dump(local_obj,f,ensure_ascii=False,indent=4)
        except Exception: pass
    if local_obj is None:
        local_obj=default
        os.makedirs(os.path.dirname(local_path) or ".", exist_ok=True)
        with open(local_path,"w",encoding="utf-8") as f: json.dump(local_obj,f,ensure_ascii=False,indent=4)
        _cloud_put(key,local_obj)
    return local_obj

def save_json_persistent(key, local_path, payload):
    os.makedirs(os.path.dirname(local_path) or ".", exist_ok=True)
    tmp=local_path+".tmp"
    with open(tmp,"w",encoding="utf-8") as f: json.dump(payload,f,ensure_ascii=False,indent=4)
    os.replace(tmp,local_path)
    _cloud_put(key,payload)
    return True

BANNERS_FILE = os.path.join(LOCAL_DATA_DIR, "banners_v2.json")

def load_banners():
    return load_json_persistent("banners", BANNERS_FILE, [])

def save_banners(banners):
    save_json_persistent("banners", BANNERS_FILE, banners)


if "RESET_CODES" not in st.session_state:
    st.session_state.RESET_CODES = {}

def send_email(to_email, subject, body):
    try:
        msg = MIMEText(body, "plain", "utf-8")
        msg["Subject"] = subject
        msg["From"] = SENDER_EMAIL
        msg["To"] = to_email
        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.sendmail(SENDER_EMAIL, to_email, msg.as_string())
        server.quit()
        return True
    except Exception as e:
        st.error(f"خطأ في الارسال: {e}")
        return False

import os
import json

USERS_FILE = os.path.join(LOCAL_DATA_DIR, "users.json")  # <<< ده اهم سطر . ده اللي بيخلي الحفظ يشتغل

import os
import json

USERS_FILE = os.path.join(LOCAL_DATA_DIR, "users.json")  # <<< السطر ده هو السر كله

def load_users():
    admin_pass = bcrypt.hashpw(ADMIN_DEFAULT_PASS.encode(), bcrypt.gensalt()).decode()
    default_users = [{"id":1,"username":ADMIN_USERNAME,"display_name":ADMIN_USERNAME,"password":admin_pass,"email":SENDER_EMAIL,"recovery_email":"","role":"admin","status":"active","password_set":True,"job_title":"","insurance_area":""}]
    return load_json_persistent("users", USERS_FILE, default_users)

def save_users(users):
    save_json_persistent("users", USERS_FILE, users)

def check_login(username, password):
    users = load_users()
    changed = False
    for u in users:
        if not u.get("display_name"):
            u["display_name"] = u.get("username", "")
            changed = True
    if changed:
        save_users(users)
    for user in users:
        typed = str(username or "").strip().casefold()
        login_username = str(user.get("username", "")).strip().casefold()
        display_name = str(user.get("display_name", "")).strip().casefold()
        if typed not in (login_username, display_name):
            continue
        if user.get("status", "active") != "active":
            return None
        stored = user.get("password") or ""
        if not stored:
            return None
        try:
            if bcrypt.checkpw(password.encode(), stored.encode()):
                return user
        except Exception:
            return None
    return None

def is_admin_email(email):
    users = load_users()
    admin = next((u for u in users if u["role"] == "admin"), None)
    if not admin: return False
    return email == admin["email"] or email == admin.get("recovery_email","")

def normalize_person_name(value):
    import re
    text = str(value or "").strip().casefold().replace("ـ", "")
    return re.sub(r"\s+", " ", text)

def find_user_for_activation(users, typed_name):
    key = normalize_person_name(typed_name)
    if not key:
        return None
    for u in users:
        if normalize_person_name(u.get("username")) == key or normalize_person_name(u.get("display_name")) == key:
            return u
    compact = key.replace(" ", "")
    for u in users:
        for field in ("username", "display_name"):
            candidate = normalize_person_name(u.get(field)).replace(" ", "")
            if candidate and candidate == compact:
                return u
    return None

def login_page():
    st.markdown("""
    <div class="login-shell">
      <h1 style="text-align:center;color:#D4AF37!important;">⚖️ إدارة القضايا دخول السادة الأعضاء ⚖️</h1>
    </div>
    """, unsafe_allow_html=True)

    tab1, tab2 = st.tabs(["🔐 تسجيل الدخول", "🪪 تفعيل حساب جديد"])

    with tab1:
        st.markdown("<div class='login-shell' style='margin-top:0;'>", unsafe_allow_html=True)
        st.markdown("<p style='color:white!important;font-weight:900;'>اسم المستخدم</p>", unsafe_allow_html=True)
        username = st.text_input("اسم المستخدم", key="login_user", label_visibility="collapsed")
        st.markdown("<p style='color:white!important;font-weight:900;'>كلمة السر</p>", unsafe_allow_html=True)
        password = st.text_input("كلمة المرور", type="password", key="login_pass", label_visibility="collapsed")

        if st.button("دخول", type="primary", use_container_width=True, key="login_btn"):
            user = check_login(username.strip(), password)
            if user:
                st.session_state.user = user
                st.session_state.role = user["role"]
                if user["role"] == "member" and not user.get("password_set", False):
                    st.session_state.page = "set_password"
                    st.session_state.temp_user_id = user.get("id")
                    st.session_state.temp_user = user.get("username")
                else:
                    st.session_state.page = "الرئيسية"
                st.rerun()
            else:
                st.error("اسم المستخدم أو كلمة السر غير صحيحة أو العضوية موقوفة")

        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown("<p style='color:white!important;font-weight:900;'>نسيت بياناتك؟ استرجعها بالإيميل</p>", unsafe_allow_html=True)

        admin_recover_email = st.text_input("إيميل الادمن", key="admin_recover")
        if st.button("إرسال كود للادمن", key="admin_send", use_container_width=True):
            if is_admin_email(admin_recover_email):
                code = str(random.randint(100000, 999999))
                st.session_state.RESET_CODES[admin_recover_email] = {"code": code, "role": "admin"}
                body = f"كود إعادة تعيين كلمة سر الادمن: {code}"
                if send_email(admin_recover_email, "كود استرجاع الادمن", body):
                    st.success("تم إرسال الكود على الإيميل")
                    st.session_state.show_reset_admin = True
            else:
                st.error("هذا الإيميل غير مسجل كادمن")

        member_recover_email = st.text_input("إيميل العضو", key="member_recover")
        if st.button("إرسال كود للعضو", key="member_send", use_container_width=True):
            users = load_users()
            found = [u for u in users if u.get("email") == member_recover_email]
            if found:
                user = found[0]
                code = str(random.randint(100000, 999999))
                st.session_state.RESET_CODES[member_recover_email] = {"code": code, "user_id": user["id"]}
                body = f"مرحبا {user['username']}\nاسم المستخدم: {user['username']}\nكود إعادة التعيين: {code}"
                if send_email(member_recover_email, "استرجاع بيانات الدخول", body):
                    st.success("تم إرسال الكود على إيميلك")
                    st.session_state.show_reset_member = True
            else:
                st.error("الإيميل ده مش متسجل")

        if st.session_state.get("show_reset_admin") or st.session_state.get("show_reset_member"):
            email_to_reset = admin_recover_email if st.session_state.get("show_reset_admin") else member_recover_email
            code_input = st.text_input("ادخل الكود")
            new_pass = st.text_input("كلمة السر الجديدة", type="password")
            if st.button("تأكيد وتغيير كلمة السر", key="confirm_reset"):
                saved = st.session_state.RESET_CODES.get(email_to_reset, {})
                if saved.get("code") == code_input:
                    users = load_users()
                    logged_user = None
                    if saved.get("role") == "admin":
                        admin = next((u for u in users if u["role"] == "admin"), None)
                        if admin:
                            hashed = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt()).decode()
                            admin["password"] = hashed
                            save_users(users)
                            logged_user = admin
                    else:
                        uid = saved.get("user_id")
                        for u in users:
                            if u["id"] == uid:
                                u["password"] = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt()).decode()
                                u["status"] = "active"
                                u["password_set"] = True
                                logged_user = u
                                break
                        save_users(users)

                    if logged_user:
                        st.session_state.RESET_CODES.clear()
                        st.session_state.show_reset_admin = False
                        st.session_state.show_reset_member = False
                        st.session_state.user = logged_user
                        st.session_state.role = logged_user["role"]
                        st.session_state.page = "الرئيسية"
                        st.success("تم تسجيل الدخول بنجاح")
                        st.rerun()
                else:
                    st.error("الكود غلط")
        st.markdown("</div>", unsafe_allow_html=True)

    with tab2:
        st.markdown("<div class='login-shell' style='margin-top:0;'>", unsafe_allow_html=True)
        st.markdown("<h3 style='color:#D4AF37!important;text-align:center;'>تفعيل حساب جديد</h3>", unsafe_allow_html=True)
        member_name = st.text_input("اكتب الاسم الذي استخرجه لك الادمن", key="new_user")
        if st.button("🪪 تفعيل الحساب", use_container_width=True, type="primary", key="activate_account_v2"):
            users = load_users()
            found_user = find_user_for_activation(users, member_name)
            if not member_name.strip():
                st.error("اكتب اسم العضو أولاً")
            elif not found_user:
                st.error("الاسم غير موجود ضمن العضويات المستخرجة. اكتب الاسم كما سجله الادمن عند استخراج العضوية.")
            elif found_user.get("role") == "admin":
                st.error("هذا حساب الادمن ولا يتم تفعيله من هنا.")
            elif found_user.get("status") in ("suspended_final", "deleted", "banned"):
                st.error("هذه العضوية موقوفة نهائياً ولا يمكن تفعيلها.")
            elif found_user.get("password_set") and found_user.get("password"):
                st.error("هذه العضوية مفعلة بالفعل. استخدم تسجيل الدخول أو استرجاع كلمة السر.")
            else:
                st.session_state.activation_user_id = found_user.get("id")
                st.session_state.activation_user_name = found_user.get("display_name") or found_user.get("username")
                st.session_state.show_activation_password = True
                st.rerun()

        if st.session_state.get("show_activation_password"):
            users = load_users()
            activation_id = st.session_state.get("activation_user_id")
            target = next((u for u in users if u.get("id") == activation_id), None)
            if target:
                st.markdown(f"<div style='background:#14233d;border:2px solid #D4AF37;border-radius:14px;padding:14px;margin-top:14px;color:#fff;text-align:center;font-weight:900;'>تم العثور على العضوية: {html.escape(str(target.get('display_name') or target.get('username')))}<br>أدخل كلمة السر التي تريد استخدامها.</div>", unsafe_allow_html=True)
                first = st.text_input("كلمة السر", type="password", key="activation_pass_1")
                second = st.text_input("تأكيد كلمة السر", type="password", key="activation_pass_2")
                activation_photo = st.file_uploader("الصورة الشخصية", type=["jpg","jpeg","png"], key="activation_profile_photo")
                if st.button("💾 حفظ كلمة السر وتفعيل العضوية", type="primary", use_container_width=True, key="confirm_activation_v2"):
                    if not first:
                        st.error("اكتب كلمة السر أولاً")
                    elif first != second:
                        st.error("كلمتا السر غير متطابقتين")
                    else:
                        for u in users:
                            if u.get("id") == activation_id:
                                u["password"] = bcrypt.hashpw(first.encode(), bcrypt.gensalt()).decode()
                                u["password_set"] = True
                                if activation_photo:
                                    u["profile_photo"] = base64.b64encode(activation_photo.getvalue()).decode("utf-8")
                                    u["profile_photo_type"] = activation_photo.type
                                if u.get("status") in (None, "suspended_temp"):
                                    u["status"] = "active"
                                target = u
                                break
                        save_users(users)
                        st.session_state.user = target
                        st.session_state.role = target.get("role")
                        st.session_state.page = "الرئيسية"
                        st.session_state.pop("activation_user_id", None)
                        st.session_state.pop("activation_user_name", None)
                        st.session_state.show_activation_password = False
                        st.success("تم تفعيل العضوية وتسجيل الدخول بنجاح")
                        st.rerun()
            else:
                st.session_state.show_activation_password = False
                st.session_state.pop("activation_user_id", None)
                st.error("تعذر العثور على العضوية، اكتب الاسم مرة أخرى.")

        st.markdown("</div>", unsafe_allow_html=True)

def calculate_age(birth_date, on_date=None):
    """حساب العمر بدقة بالسنوات الكاملة."""
    if not birth_date:
        return None
    if isinstance(birth_date, str):
        try:
            birth_date = datetime.strptime(birth_date, "%Y-%m-%d").date()
        except Exception:
            return None
    on_date = on_date or datetime.now().date()
    return on_date.year - birth_date.year - ((on_date.month, on_date.day) < (birth_date.month, birth_date.day))


def member_age_status(user):
    """يعيد حالة السن، مع مراعاة تاريخ مد العضوية إن وجد."""
    birth = user.get("birth_date")
    age = calculate_age(birth)
    extension = user.get("membership_extended_until")
    extended_active = False
    if extension:
        try:
            extended_active = datetime.fromisoformat(str(extension)).date() >= datetime.now().date()
        except Exception:
            try:
                extended_active = datetime.strptime(str(extension), "%Y-%m-%d").date() >= datetime.now().date()
            except Exception:
                extended_active = False
    return age, extended_active


def extract_member_page():
    st.markdown("<h2 style='text-align:center; color:#C9A961'>🪪 استخراج عضوية جديدة</h2>", unsafe_allow_html=True)
    st.markdown("""
    <style>
    div[data-testid="stTextInput"] label, div[data-testid="stDateInput"] label,
    div[data-testid="stSelectbox"] label {color:white!important;font-weight:900!important;}
    </style>
    """, unsafe_allow_html=True)

    if st.button("⬅️ العودة للوحة التحكم", use_container_width=True):
        st.session_state.page = "لوحة التحكم"
        st.rerun()

    with st.container(border=True):
        st.markdown("<h3 style='color:#D4AF37;text-align:center;'>بيانات العضوية</h3>", unsafe_allow_html=True)
        new_username = st.text_input("الاسم", key="new_username")
        birth_date = st.date_input(
            "تاريخ الميلاد",
            value=datetime.now().date().replace(year=max(1900, datetime.now().year - 30)),
            min_value=datetime(1900,1,1).date(),
            max_value=datetime.now().date(),
            key="new_member_birth_date"
        )
        insurance_area = st.text_input("المنطقة التأمينية التابع لها العضو", key="new_member_insurance_area")
        job_role = st.selectbox("الوظيفة", ["عضو قانونى", "مدير إدارة", "مدير عام", "قيادي مركز رئيسى 1", "قيادي مركز رئيسى 2", "قيادي مركز رئيسى 3"], key="new_member_job_role")
        profile_photo = st.file_uploader("الصورة الشخصية", type=["jpg","jpeg","png"], key="new_member_profile_photo")

        if st.button("🪪 استخراج العضوية", use_container_width=True, type="primary", key="extract_member_new_v3"):
            name = new_username.strip()
            area = insurance_area.strip()
            if not name:
                st.error("لازم تكتب اسم العضو")
            elif not area:
                st.error("لازم تكتب المنطقة التأمينية")
            else:
                users = load_users()
                existing_user = next((u for u in users if u.get("username") == name), None)
                if existing_user:
                    existing_user["birth_date"] = birth_date.isoformat()
                    existing_user["insurance_area"] = area
                    existing_user["job_title"] = job_role
                    if profile_photo:
                        existing_user["profile_photo"] = base64.b64encode(profile_photo.getvalue()).decode("utf-8")
                        existing_user["profile_photo_type"] = profile_photo.type
                    existing_user["role"] = {"عضو قانونى":"member", "مدير إدارة":"manager", "مدير عام":"general_manager", "قيادي مركز رئيسى 1":"leader1", "قيادي مركز رئيسى 2":"leader2", "قيادي مركز رئيسى 3":"leader3"}.get(job_role, "member")
                    existing_user.setdefault("membership_extended_until", None)
                    if existing_user.get("status") in ("banned", "deleted"):
                        existing_user["status"] = "active"
                    if not existing_user.get("password_set"):
                        existing_user["password"] = ""
                    save_users(users)
                    st.success(f"تم حفظ بيانات العضوية للعضو: {name}")
                else:
                    new_id = max([int(u.get('id', 0)) for u in users if str(u.get('id','')).isdigit()] or [0]) + 1
                    users.append({
                        "id": new_id,
                        "username": name,
                        "password": "",
                        "email": "",
                        "recovery_email": "",
                        "role": {"عضو قانونى":"member", "مدير إدارة":"manager", "مدير عام":"general_manager", "قيادي مركز رئيسى 1":"leader1", "قيادي مركز رئيسى 2":"leader2", "قيادي مركز رئيسى 3":"leader3"}.get(job_role, "member"),
                        "job_title": job_role,
                        "status": "active",
                        "password_set": False,
                        "birth_date": birth_date.isoformat(),
                        "insurance_area": area,
                        "membership_extended_until": None,
                        "age_action": None,
                        "age_action_at": None,
                        "transferred_to": None,
                        "profile_photo": base64.b64encode(profile_photo.getvalue()).decode("utf-8") if profile_photo else "",
                        "profile_photo_type": profile_photo.type if profile_photo else ""
                    })
                    save_users(users)
                    st.success(f"تم استخراج العضوية بنجاح: {name}")
                st.rerun()


def _transfer_member_data(source_user, target_user):
    """نقل القضايا والمكتبة للعضو الجديد مع الإبقاء على السجلات دون حذف."""
    data = load_data()
    source_name = source_user.get("username", "")
    target_name = target_user.get("username", "")
    target_id = target_user.get("id")

    changed = False
    for case in data.get("cases", []):
        if case.get("assigned_to") == source_name:
            case["assigned_to"] = target_name
            changed = True
        if case.get("user_id") == source_user.get("id"):
            case["user_id"] = target_id
            changed = True

    for item in data.get("library", []):
        if item.get("user_id") == source_user.get("id"):
            item["user_id"] = target_id
            changed = True

    if changed:
        save_data(data)
    return changed


def _age_action_dialog(user_id):
    """حوار إجراءات بلوغ الستين للعضو."""
    users = load_users()
    user = next((u for u in users if u.get("id") == user_id), None)
    if not user:
        return

    age, extended_active = member_age_status(user)
    if age is None or age < 60 or extended_active:
        return

    @st.dialog("إجراءات العضوية لبلوغ سن الستين")
    def age_dialog():
        st.warning(f"⚠️ العضو **{user.get('username','')}** بلغ سن {age} سنة.")
        st.write(f"تاريخ الميلاد: {user.get('birth_date','غير مسجل')}")
        st.write(f"المنطقة التأمينية: {user.get('insurance_area','غير مسجلة')}")

        action = st.radio(
            "اختر الإجراء",
            [
                "إيقاف عضوية لبلوغ السن",
                "إيقاف عضوية لبلوغ السن مع نقل البيانات لعضو آخر",
                "مد العضوية لفترة"
            ],
            key=f"age_action_{user_id}"
        )

        if action == "إيقاف عضوية لبلوغ السن":
            if st.button("⛔ تنفيذ إيقاف العضوية", type="primary", use_container_width=True, key=f"age_stop_{user_id}"):
                user["status"] = "suspended_age"
                user["age_action"] = "إيقاف عضوية لبلوغ السن"
                user["age_action_at"] = datetime.now().isoformat()
                save_users(users)
                st.success("تم إيقاف العضوية، ولم يتم حذف أي قضية أو مستند.")
                st.rerun()

        elif action == "إيقاف عضوية لبلوغ السن مع نقل البيانات لعضو آخر":
            targets = [u for u in users if u.get("role") != "admin" and u.get("id") != user.get("id") and u.get("status") == "active"]
            if not targets:
                st.error("لا يوجد عضو آخر نشط يمكن نقل البيانات إليه.")
            else:
                target_map = {u.get("username", ""): u for u in targets}
                target_name = st.selectbox("ادخل اسم العضو الآخر", list(target_map.keys()), key=f"age_target_{user_id}")
                st.info("سيتم نقل القضايا والمكتبة الخاصة بالعضو للعضو المختار، مع عدم حذف السجلات.")
                if st.button("🔄 نقل البيانات ثم إيقاف العضوية", type="primary", use_container_width=True, key=f"age_transfer_{user_id}"):
                    target = target_map[target_name]
                    _transfer_member_data(user, target)
                    user["status"] = "suspended_age"
                    user["age_action"] = "إيقاف عضوية لبلوغ السن مع نقل البيانات"
                    user["age_action_at"] = datetime.now().isoformat()
                    user["transferred_to"] = target.get("username")
                    save_users(users)
                    st.success(f"تم نقل البيانات إلى {target.get('username')} ثم إيقاف العضوية.")
                    st.rerun()

        else:
            extension_date = st.date_input(
                "مد العضوية حتى تاريخ",
                min_value=datetime.now().date(),
                value=(datetime.now() + timedelta(days=365)).date(),
                key=f"age_extension_{user_id}"
            )
            if st.button("📅 حفظ مد العضوية", type="primary", use_container_width=True, key=f"age_extend_{user_id}"):
                user["status"] = "active"
                user["membership_extended_until"] = extension_date.isoformat()
                user["age_action"] = "مد العضوية"
                user["age_action_at"] = datetime.now().isoformat()
                save_users(users)
                st.success(f"تم مد العضوية حتى {extension_date.strftime('%d-%m-%Y')}")
                st.rerun()

    age_dialog()

def manage_users_page():
    if not st.session_state.user or st.session_state.user.get("role") != "admin":
        st.error("هذه الصفحة متاحة للادمن فقط")
        return

    st.markdown("<h2 style='color:#D4AF37;text-align:center;'>👥 إدارة الأعضاء</h2>", unsafe_allow_html=True)
    users = load_users()
    members = [u for u in users if u.get("role") != "admin"]

    st.info("إيقاف العضو أو حذف حساب الدخول لا يحذف القضايا أو المستندات. بيانات العضو تظل محفوظة.")

    if not members:
        st.info("لا يوجد أعضاء حالياً.")
        return

    for user in members:
        birth = user.get("birth_date")
        age, extended_active = member_age_status(user)
        status = user.get("status", "active")
        if status == "active":
            status_text, cls = "نشط", "status-active"
        elif status == "suspended_temp":
            status_text, cls = "إيقاف مؤقت", "status-temp"
        elif status == "suspended_final":
            status_text, cls = "إيقاف نهائي", "status-final"
        elif status == "suspended_age":
            status_text, cls = "إيقاف لبلوغ السن", "status-final"
        else:
            status_text, cls = status, "status-final"

        with st.container(border=True):
            c1, c2, c3 = st.columns([4, 2, 4])
            with c1:
                if user.get("profile_photo"):
                    st.markdown(_photo_html(user.get("profile_photo"), user.get("profile_photo_type") or "image/png", 85), unsafe_allow_html=True)
                st.markdown(f"### 👤 {user.get('username', '')}")
                st.markdown(f"**الوظيفة:** {user.get('job_title') or role_label(user.get("role"))}")
                st.markdown(f"**تاريخ الميلاد:** {birth or 'غير مسجل'}")
                st.markdown(f"**المنطقة التأمينية:** {user.get('insurance_area') or 'غير مسجلة'}")
                if age is not None:
                    age_line = f"**العمر:** {age} سنة"
                    if extended_active:
                        age_line += f" — **العضوية ممددة حتى:** {user.get('membership_extended_until')}"
                    st.markdown(age_line)
                st.markdown(f"الإيميل: {user.get('email') or 'غير مسجل'}")
                st.markdown(f"الحالة: <span class='{cls}'>{status_text}</span>", unsafe_allow_html=True)

                if age is not None and age >= 60 and not extended_active and status not in ("suspended_age", "suspended_final"):
                    st.warning("⚠️ بلغ سن الستين — يلزم اتخاذ إجراء بشأن العضوية.")
                    if st.button("⚖️ إجراءات بلوغ سن الستين", key=f"age_dialog_btn_{user['id']}", use_container_width=True):
                        _age_action_dialog(user["id"])
                elif status == "suspended_age":
                    st.warning("⛔ العضوية موقوفة لبلوغ السن.")
                    if user.get("transferred_to"):
                        st.caption(f"تم نقل البيانات إلى: {user.get('transferred_to')}")

            with c2:
                if status != "active":
                    if st.button("✅ إعادة تفعيل", key=f"reactivate_{user['id']}", use_container_width=True):
                        user["status"] = "active"
                        save_users(users)
                        st.rerun()

                if age is not None and age >= 60 and not extended_active and status not in ("suspended_age", "suspended_final"):
                    if st.button("⚖️ إجراء السن", key=f"age_action_btn_{user['id']}", use_container_width=True):
                        _age_action_dialog(user["id"])

            with c3:
                a,b,c = st.columns(3)
                with a:
                    if status == "active":
                        if st.button("⏸️ إيقاف مؤقت", key=f"temp_{user['id']}", use_container_width=True):
                            user["status"] = "suspended_temp"
                            save_users(users)
                            st.rerun()
                with b:
                    if status != "suspended_final":
                        if st.button("⛔ إيقاف نهائي", key=f"final_{user['id']}", use_container_width=True):
                            user["status"] = "suspended_final"
                            user["password"] = ""
                            user["password_set"] = False
                            save_users(users)
                            st.rerun()
                with c:
                    if st.button("🗑️ حذف الحساب", key=f"delete_member_{user['id']}", use_container_width=True):
                        st.session_state[f"confirm_member_delete_{user['id']}"] = True
                        st.rerun()

            if st.session_state.get(f"confirm_member_delete_{user['id']}"):
                st.warning("تأكيد حذف حساب الدخول فقط — القضايا والمستندات ستظل محفوظة.")
                y,n = st.columns(2)
                with y:
                    if st.button("نعم، احذف الحساب", key=f"yes_delete_member_{user['id']}", type="primary", use_container_width=True):
                        users = [u for u in users if u.get("id") != user.get("id")]
                        save_users(users)
                        st.session_state.pop(f"confirm_member_delete_{user['id']}", None)
                        st.rerun()
                with n:
                    if st.button("إلغاء", key=f"no_delete_member_{user['id']}", use_container_width=True):
                        st.session_state.pop(f"confirm_member_delete_{user['id']}", None)
                        st.rerun()


# ===== الرسائل والبيانات الداخلية =====
MESSAGES_FILE = os.path.join(LOCAL_DATA_DIR, "messages_v1.json")

def load_messages():
    data = load_json_persistent("messages", MESSAGES_FILE, [])
    return data if isinstance(data, list) else []

def save_messages(items):
    save_json_persistent("messages", MESSAGES_FILE, items)

def role_label(role):
    return {
        "admin":"الادمن", "member":"عضو قانونى", "manager":"مدير إدارة",
        "general_manager":"مدير عام (قيادى مناطق)",
        "leader1":"قيادي مركز رئيسى 1", "leader2":"قيادي مركز رئيسى 2",
        "leader3":"قيادي مركز رئيسى 3"
    }.get(role, role or "")

def can_message(sender_role, target_role):
    # جميع العضويات تستطيع مراسلة الادمن مباشرة.
    if target_role == "admin" and sender_role != "admin":
        return True
    if sender_role == "admin":
        return target_role in ("admin", "member", "manager", "general_manager", "leader1", "leader2", "leader3")
    if sender_role == "member":
        return target_role in ("admin", "member", "manager")
    if sender_role == "manager":
        return target_role in ("admin", "member", "manager", "general_manager")
    if sender_role == "leader1":
        return target_role in ("admin", "member", "manager", "general_manager", "leader1", "leader2", "leader3")
    if sender_role == "leader2":
        return target_role in ("admin", "member", "manager", "general_manager", "leader1", "leader2", "leader3")
    if sender_role == "leader3":
        return target_role in ("admin", "member", "manager", "general_manager", "leader2", "leader3")
    if sender_role == "general_manager":
        return target_role in ("admin", "member", "manager", "general_manager", "leader2", "leader3")
    return False

def person_display(u, include_role=True):
    name = u.get("display_name") or u.get("username") or "بدون اسم"
    role = role_label(u.get("role"))
    if u.get("role") == "admin":
        role = "الادمن"
    area = u.get("insurance_area") or "غير محددة"
    if include_role:
        return f"{name} — {role} — المنطقة التأمينية: {area}"
    return name

def _recipient_groups(candidates, key_prefix, preselected=None):
    preselected = set(preselected or [])
    selected_ids = []
    grouped = {}
    for u in candidates:
        area = (u.get("insurance_area") or "غير محددة").strip() or "غير محددة"
        grouped.setdefault(area, []).append(u)
    for area in sorted(grouped, key=lambda x: x.casefold()):
        st.markdown(f"<div style='color:#D4AF37;font-size:17px;font-weight:900;margin:10px 0 6px;'>المنطقة: {html.escape(area)}</div>", unsafe_allow_html=True)
        cols = st.columns(3)
        for i, u in enumerate(sorted(grouped[area], key=lambda x: (role_order(x.get("role")), str(x.get("display_name") or x.get("username") or "").casefold()))):
            uid = u.get("id")
            label = person_display(u)
            with cols[i % 3]:
                checked = st.checkbox(label, value=(uid in preselected), key=f"{key_prefix}_{uid}")
                if checked:
                    selected_ids.append(uid)
    return selected_ids

def messaging_page():
    user = st.session_state.get("user") or {}
    sr = user.get("role")
    if sr not in ("admin", "member", "manager", "general_manager", "leader1", "leader2", "leader3"):
        return

    users = load_users()
    messages = load_messages()
    st.markdown("<h2 style='color:#D4AF37;text-align:center;'>📢 الرسائل والبيانات</h2>", unsafe_allow_html=True)

    target_options = [
        ("الادمن", "admin"), ("عضو قانونى", "member"), ("مدير إدارة", "manager"),
        ("مدير عام", "general_manager"), ("قيادي مركز رئيسى 1", "leader1"),
        ("قيادي مركز رئيسى 2", "leader2"), ("قيادي مركز رئيسى 3", "leader3")
    ]
    allowed_roles = [r for _, r in target_options if can_message(sr, r)]
    preselected = st.session_state.pop("message_target_ids", [])
    candidates = [u for u in users if u.get("status") == "active" and u.get("id") != user.get("id") and can_message(sr, u.get("role"))]

    if not candidates:
        st.info("لا يوجد مستلمون متاحون وفق صلاحيات الحساب")
        return

    target_kind = st.selectbox(
        "نوع المستلمين",
        ["الجميع"] + [label for label, role in target_options if role in allowed_roles] + ["محددون"],
        key=f"msg_target_kind_{sr}"
    )
    text = st.text_area("نص الرسالة أو البيان", height=130, key=f"msg_text_{sr}")

    attach_count_key = f"msg_attach_count_{sr}"
    if attach_count_key not in st.session_state:
        st.session_state[attach_count_key] = 1
    attachment_files = []
    for ai in range(st.session_state[attach_count_key]):
        af = st.file_uploader(
            "📎 المرفق" if ai == 0 else f"📎 مرفق إضافي {ai + 1}",
            type=None,
            key=f"msg_attachment_{sr}_{ai}",
            help="يمكن رفع PDF أو Word أو Excel أو صورة أو أي ملف آخر"
        )
        attachment_files.append(af)
        if af is not None and st.button("➕ إضافة مرفق آخر", key=f"msg_add_attachment_{sr}_{ai}", use_container_width=True):
            st.session_state[attach_count_key] += 1
            st.rerun()

    image_file = st.file_uploader("🖼️ الصورة", type=["jpg","jpeg","png"], key=f"msg_image_{sr}")

    if target_kind == "الجميع":
        selected_ids = [u.get("id") for u in candidates]
    elif target_kind == "محددون":
        selected_ids = _recipient_groups(candidates, f"msg_chk_{sr}", preselected)
    else:
        rr = dict(target_options).get(target_kind)
        role_candidates = [u for u in candidates if u.get("role") == rr]
        selected_ids = _recipient_groups(role_candidates, f"msg_chk_{sr}_{rr}", preselected)

    if st.button("📤 إرسال الرسالة / البيان", type="primary", use_container_width=True, key=f"send_message_{sr}"):
        if not text.strip():
            st.error("اكتب نص الرسالة أو البيان")
            return
        if not selected_ids:
            st.error("حدد مستلمًا واحدًا على الأقل")
            return
        messages.append({
            "id": secrets.token_hex(10),
            "sender_id": user.get("id"),
            "sender_name": user.get("display_name") or user.get("username"),
            "sender_role": sr,
            "sender_job_title": user.get("job_title") or role_label(sr),
            "sender_insurance_area": user.get("insurance_area") or "",
            "sender_photo": user.get("profile_photo") or "",
            "sender_photo_type": user.get("profile_photo_type") or "",
            "text": text.strip(),
            "recipient_ids": selected_ids,
            "created_at": datetime.now().astimezone().isoformat(),
            "image_data": base64.b64encode(image_file.getvalue()).decode("utf-8") if image_file else "",
            "image_type": image_file.type if image_file else "",
            "attachments": [{"data": base64.b64encode(af.getvalue()).decode("utf-8"), "name": af.name, "type": af.type or "application/octet-stream"} for af in attachment_files if af is not None],
            "attachment_data": base64.b64encode(attachment_files[0].getvalue()).decode("utf-8") if attachment_files and attachment_files[0] else "",
            "attachment_name": attachment_files[0].name if attachment_files and attachment_files[0] else "",
            "attachment_type": attachment_files[0].type if attachment_files and attachment_files[0] else "application/octet-stream",
            "in_reply_to": ""
        })
        save_messages(messages)
        st.session_state.pop("message_target_ids", None)
        st.success("تم إرسال الرسالة وحفظها")
        st.rerun()

def _photo_html(photo_data, photo_type="image/png", size=72):
    if not photo_data:
        return ""
    return f"<img src='data:{html.escape(photo_type)};base64,{photo_data}' style='width:{size}px;height:{size}px;border-radius:50%;object-fit:cover;border:2px solid #D4AF37;margin:0 auto 8px;display:block;'>"

def _visible_banners_for_user(user):
    username = user.get("username", "")
    now = datetime.now()
    out = []
    for b in load_banners():
        try:
            expire = datetime.fromisoformat(str(b.get("expire_at") or b.get("expire")))
        except Exception:
            continue
        if expire <= now:
            continue
        audience = b.get("audience", "الكل")
        if audience not in ("الكل", "كل الأعضاء") and username not in b.get("visible_to", []):
            continue
        out.append(b)
    return out

def show_top_notices():
    user = st.session_state.get("user") or {}
    uid = user.get("id")
    if not uid:
        return
    users = load_users()
    by_id = {u.get("id"): u for u in users}
    notices = []
    for m in load_messages():
        if uid in m.get("recipient_ids", []):
            sender = by_id.get(m.get("sender_id"), {})
            notices.append({"kind":"message", "created_at":m.get("created_at", ""), "item":m, "sender":sender})
    for b in _visible_banners_for_user(user):
        sender = by_id.get(b.get("sender_id"), {})
        notices.append({"kind":"banner", "created_at":b.get("created_at", ""), "item":b, "sender":sender})
    if not notices:
        return
    notices.sort(key=lambda x: str(x.get("created_at", "")), reverse=True)
    total = len(notices)
    idx = int(st.session_state.get("top_notice_index", 0))
    idx = max(0, min(idx, total - 1))
    st.session_state.top_notice_index = idx
    n = notices[idx]
    item = n["item"]
    sender = n["sender"]
    sender_name = item.get("sender_name") or sender.get("display_name") or sender.get("username") or "الادمن"
    sender_role = item.get("sender_role") or sender.get("role") or "admin"
    sender_job = item.get("sender_job_title") or sender.get("job_title") or role_label(sender_role)
    sender_area = item.get("sender_insurance_area") or sender.get("insurance_area") or "غير محددة"
    photo = item.get("sender_photo") or sender.get("profile_photo") or ""
    photo_type = item.get("sender_photo_type") or sender.get("profile_photo_type") or "image/png"
    text = item.get("text", "")
    image_data = item.get("image_data", "")
    image_type = item.get("image_type", "image/png")
    title = "رسالة واردة" if n["kind"] == "message" else "بيان وارد"
    created = html.escape(str(item.get("created_at", "")))
    sender_color = "#00E676" if sender_role == "admin" else "#D4AF37"
    card = f"<div style='background:linear-gradient(135deg,#101b31,#203b63);border:2px solid {sender_color};border-radius:18px;padding:16px;margin:0 0 10px;box-shadow:0 10px 28px rgba(0,0,0,.3);text-align:center;'>"
    card += _photo_html(photo, photo_type, 70)
    card += f"<div style='color:#D4AF37;font-size:22px;font-weight:900;'>📢 {title}</div>"
    admin_prefix = "<span style=\"color:#ff3b30;font-weight:1000;margin-left:8px;\">الادمن</span>" if sender_role == "admin" else ""
    card += f"<div style=\'color:{sender_color};font-size:20px;font-weight:900;margin-top:5px;\'>المرسل: {admin_prefix} {html.escape(str(sender_name))}</div>"
    card += f"<div style='color:#fff;font-size:15px;font-weight:800;margin-top:4px;'>الوظيفة: {html.escape(str(sender_job))} — المنطقة التأمينية: {html.escape(str(sender_area))}</div>"
    card += f"<div style='color:#fff;font-size:18px;font-weight:700;margin-top:12px;white-space:pre-wrap;text-align:right;direction:rtl;'>{html.escape(str(text))}</div>"
    if image_data:
        card += f"<img src='data:{html.escape(image_type)};base64,{image_data}' style='max-width:100%;max-height:360px;border-radius:14px;margin-top:12px;border:1px solid #D4AF37;'>"
    attachments = item.get("attachments") or []
    if not attachments and item.get("attachment_data"):
        attachments = [{"data": item.get("attachment_data"), "name": item.get("attachment_name") or "مرفق", "type": item.get("attachment_type") or "application/octet-stream"}]
    for att in attachments:
        card += f"<div style='margin-top:10px;padding:9px;border:1px solid #D4AF37;border-radius:12px;color:#fff;font-weight:800;'>📎 {html.escape(str(att.get('name') or 'مرفق'))}</div>"
    card += f"<div style='color:#cbd5e1;font-size:12px;margin-top:10px;'>{created}</div></div>"
    st.markdown(card, unsafe_allow_html=True)
    attachments = item.get("attachments") or []
    if not attachments and item.get("attachment_data"):
        attachments = [{"data": item.get("attachment_data"), "name": item.get("attachment_name") or "مرفق", "type": item.get("attachment_type") or "application/octet-stream"}]
    for ai, att in enumerate(attachments):
        try:
            attachment_bytes = base64.b64decode(att.get("data") or "")
            st.download_button(
                f"📎 فتح / تنزيل: {att.get('name') or 'مرفق'}",
                data=attachment_bytes, file_name=att.get("name") or "مرفق",
                mime=att.get("type") or "application/octet-stream", use_container_width=True,
                key=f"download_notice_attachment_{item.get('id','x')}_{ai}"
            )
        except Exception:
            pass
    c1,c2,c3 = st.columns([1,1,1])
    with c1:
        if st.button("⬅️ السابقة", disabled=(idx >= total-1), use_container_width=True, key="top_notice_prev_v4"):
            st.session_state.top_notice_index = min(total-1, idx+1)
            st.rerun()
    with c2:
        st.markdown(f"<div style='text-align:center;color:#fff;font-weight:900;padding-top:8px;'>الرسالة {idx+1} من {total}</div>", unsafe_allow_html=True)
    with c3:
        if st.button("التالية ➡️", disabled=(idx <= 0), use_container_width=True, key="top_notice_next_v4"):
            st.session_state.top_notice_index = max(0, idx-1)
            st.rerun()

    if n["kind"] == "message":
        st.markdown("<div style='color:#D4AF37;font-size:18px;font-weight:900;text-align:center;margin-top:10px;'>الرد على الرسالة</div>", unsafe_allow_html=True)
        reply_text = st.text_area("نص الرد", height=110, key=f"top_reply_text_{item.get('id','x')}")
        reply_count_key = f"reply_attach_count_{item.get('id','x')}"
        if reply_count_key not in st.session_state:
            st.session_state[reply_count_key] = 1
        reply_attachments = []
        for ri in range(st.session_state[reply_count_key]):
            rf = st.file_uploader("📎 مرفق الرد" if ri == 0 else f"📎 مرفق رد إضافي {ri + 1}", type=None, key=f"top_reply_attachment_{item.get('id','x')}_{ri}", help="PDF أو Word أو Excel أو صورة أو أي ملف")
            reply_attachments.append(rf)
            if rf is not None and st.button("➕ إضافة مرفق آخر للرد", key=f"reply_add_attachment_{item.get('id','x')}_{ri}", use_container_width=True):
                st.session_state[reply_count_key] += 1
                st.rerun()
        if st.button("↩️ إرسال الرد", type="primary", use_container_width=True, key=f"top_reply_send_{item.get('id','x')}"):
            if not reply_text.strip() and not any(reply_attachments):
                st.error("اكتب الرد أو أرفق ملفًا")
            else:
                target_id = item.get("sender_id")
                if target_id is not None:
                    msgs = load_messages()
                    msgs.append({
                        "id": secrets.token_hex(10),
                        "sender_id": user.get("id"),
                        "sender_name": user.get("display_name") or user.get("username"),
                        "sender_role": user.get("role"),
                        "sender_job_title": user.get("job_title") or role_label(user.get("role")),
                        "sender_insurance_area": user.get("insurance_area") or "",
                        "sender_photo": user.get("profile_photo") or "",
                        "sender_photo_type": user.get("profile_photo_type") or "",
                        "text": reply_text.strip(),
                        "recipient_ids": [target_id],
                        "created_at": datetime.now().astimezone().isoformat(),
                        "image_data": "",
                        "image_type": "",
                        "attachments": [{"data": base64.b64encode(rf.getvalue()).decode("utf-8"), "name": rf.name, "type": rf.type or "application/octet-stream"} for rf in reply_attachments if rf is not None],
                        "attachment_data": base64.b64encode(reply_attachments[0].getvalue()).decode("utf-8") if reply_attachments and reply_attachments[0] else "",
                        "attachment_name": reply_attachments[0].name if reply_attachments and reply_attachments[0] else "",
                        "attachment_type": reply_attachments[0].type if reply_attachments and reply_attachments[0] else "application/octet-stream",
                        "in_reply_to": item.get("id", "")
                    })
                    save_messages(msgs)
                    st.success("تم إرسال الرد")
                    st.rerun()

def role_order(role):
    return {"leader1":0,"leader2":1,"leader3":2,"general_manager":3,"manager":4,"member":5,"admin":-1}.get(role,99)

def show_member_directory():
    users = [u for u in load_users() if u.get("status") != "deleted"]
    users.sort(key=lambda u: (role_order(u.get("role")), str(u.get("insurance_area") or "").casefold(), str(u.get("display_name") or u.get("username") or "").casefold()))
    current = st.session_state.get("user") or {}
    st.markdown("<div style='color:#D4AF37;font-size:21px;font-weight:900;text-align:center;margin:14px 0 8px;'>العضويات</div>", unsafe_allow_html=True)
    for area in sorted(set((u.get("insurance_area") or "غير محددة").strip() or "غير محددة" for u in users), key=lambda x:x.casefold()):
        st.markdown(f"<div style='color:#D4AF37;font-size:14px;font-weight:900;margin:7px 0 5px;'>المنطقة: {html.escape(area)}</div>", unsafe_allow_html=True)
        area_users = [u for u in users if ((u.get("insurance_area") or "غير محددة").strip() or "غير محددة") == area]
        cols = st.columns(6)
        for i,u in enumerate(area_users):
            with cols[i % 6]:
                name = html.escape(str(u.get("display_name") or u.get("username") or ""))
                raw_job = str(u.get("job_title") or "").strip()
                job_value = "مدير عام (قيادى مناطق)" if (u.get("role") == "general_manager" or raw_job in ("مدير عام", "مدير عام (قيادى مناطق)", "مدير عام (قيادى مناطق )")) else (raw_job or role_label(u.get("role")))
                job = html.escape(job_value)
                photo = _photo_html(u.get("profile_photo"), u.get("profile_photo_type") or "image/png", 42)
                status = html.escape(str(u.get("status") or ""))
                is_admin = u.get("role") == "admin"
                border = "#00E676" if is_admin else "#D4AF37"
                name_color = "#00E676" if is_admin else "#fff"
                admin_badge = "<div style='color:#00E676;font-size:10px;font-weight:900;'>ADMIN</div>" if is_admin else ""
                card = f"<div style='background:#142038;border:2px solid {border};border-radius:12px;padding:6px;text-align:center;min-height:112px;margin-bottom:5px;'>{photo}<div style='color:{name_color};font-size:12px;font-weight:900;line-height:1.2;'>{name}</div>{admin_badge}<div style='color:#D4AF37;font-size:10px;font-weight:800;line-height:1.2;'>{job}</div><div style='color:#cbd5e1;font-size:9px;line-height:1.2;'>{html.escape(area)}</div></div>"
                st.markdown(card, unsafe_allow_html=True)
                if u.get("id") != current.get("id") and can_message(current.get("role"), u.get("role")):
                    if st.button("✉️ إرسال", key=f"dir_msg_v5_{current.get('id')}_{u.get('id')}", use_container_width=True):
                        st.session_state.message_target_ids = [u.get("id")]
                        st.session_state.page = "الرسائل"
                        st.rerun()

def global_people_search():
    user=st.session_state.get("user") or {}
    users=load_users()
    q=st.text_input("🔎 بحث سريع بالاسم", placeholder="اكتب اسم العضو أو مدير الإدارة أو المدير العام", key="global_people_search")
    if q.strip():
        matches=[u for u in users if q.strip().lower() in str(u.get("display_name") or u.get("username","")).lower()]
        if matches:
            for u in matches:
                st.markdown(f"**👤 {u.get('display_name') or u.get('username')}** — {role_label(u.get('role'))} — المنطقة التأمينية: {u.get('insurance_area') or 'لا توجد منطقة'} — الحالة: {u.get('status','')}")
        else: st.info("لا توجد نتيجة بهذا الاسم")

def admin_control_page():
    if not st.session_state.user or st.session_state.user.get("role") != "admin":
        st.error("هذه اللوحة للادمن فقط")
        return

    st.markdown("<h1 style='color:#D4AF37;text-align:center;'>⚖️ لوحة التحكم</h1>", unsafe_allow_html=True)
    st.markdown("<p style='color:white;text-align:center;'>جميع أدوات إدارة الموقع في مكان واحد</p>", unsafe_allow_html=True)

    tab_users, tab_messages, tab_search, tab_account = st.tabs(["👥 الأعضاء", "📢 الرسائل والبيانات", "🔎 البحث والملف الشخصي", "🔐 حساب الادمن والإيميل"])

    with tab_users:
        manage_users_page()
        st.markdown("---")
        if st.button("➕ استخراج عضوية جديدة", use_container_width=True, key="admin_extract_from_panel"):
            st.session_state.page = "extract_member"
            st.rerun()

    with tab_messages:
        messaging_page()
        st.markdown("---")
        st.markdown("<h3 style='color:#D4AF37;'>📜 سجل البيانات والرسائل</h3>", unsafe_allow_html=True)
        users = load_users()
        banners = load_banners()
        now = datetime.now()

        with st.form("admin_banner_add_v2"):
            text = st.text_area("نص الرسالة أو البيان", height=90)
            color = st.color_picker("لون الرسالة أو البيان", "#D4AF37")
            banner_image = st.file_uploader("صورة البيان", type=["jpg","jpeg","png"], key="banner_image_add_v3")
            col1,col2 = st.columns(2)
            with col1:
                duration_value = st.number_input("مدة الإبقاء", min_value=1, max_value=3650, value=7, step=1)
            with col2:
                duration_unit = st.selectbox("وحدة المدة", ["أيام","ساعات","دقائق"])
            audience = st.radio("يظهر لـ", ["كل الأعضاء","أعضاء محددين"], horizontal=True, key="banner_audience_v2")
            visible_to = []
            if audience == "أعضاء محددين":
                banner_candidates = [u for u in users if u.get("status") == "active"]
                selected_banner_ids = _recipient_groups(banner_candidates, "banner_recipient_chk_v3")
                visible_to = [u.get("username") for u in banner_candidates if u.get("id") in selected_banner_ids]
            if st.form_submit_button("🚀 نشر الرسالة / البيان", use_container_width=True, type="primary"):
                if not text.strip():
                    st.error("اكتب نص الرسالة أو البيان")
                elif audience == "أعضاء محددين" and not visible_to:
                    st.error("حدد عضواً واحداً على الأقل")
                else:
                    minutes = duration_value * (1440 if duration_unit == "أيام" else 60 if duration_unit == "ساعات" else 1)
                    created = datetime.now()
                    banner = {
                        "id": secrets.token_hex(8),
                        "text": text.strip(),
                        "color": color,
                        "created_at": created.isoformat(),
                        "expire_at": (created + timedelta(minutes=minutes)).isoformat(),
                        "audience": "الكل" if audience == "كل الأعضاء" else "اعضاء محددين",
                        "visible_to": visible_to,
                        "sender_id": st.session_state.user.get("id"),
                        "sender_name": st.session_state.user.get("display_name") or st.session_state.user.get("username"),
                        "sender_role": st.session_state.user.get("role"),
                        "sender_job_title": st.session_state.user.get("job_title") or role_label(st.session_state.user.get("role")),
                        "sender_insurance_area": st.session_state.user.get("insurance_area") or "",
                        "sender_photo": st.session_state.user.get("profile_photo") or "",
                        "sender_photo_type": st.session_state.user.get("profile_photo_type") or "",
                        "image_data": base64.b64encode(banner_image.getvalue()).decode("utf-8") if banner_image else "",
                        "image_type": banner_image.type if banner_image else ""
                    }
                    banners.append(banner)
                    save_banners(banners)
                    st.success("تم نشر الرسالة / البيان")
                    st.rerun()

        st.markdown("---")
        st.markdown(f"<h3 style='color:#D4AF37;'>الرسائل والبيانات الظاهرة أعلى الصفحة: {len(banners)}</h3>", unsafe_allow_html=True)

        for i,b in enumerate(banners):
            expire_raw = b.get("expire_at") or b.get("expire")
            try:
                expire_dt = datetime.fromisoformat(str(expire_raw))
            except Exception:
                expire_dt = None
            expired = bool(expire_dt and expire_dt <= now)
            audience_text = "الكل" if b.get("audience") in ("الكل","كل الأعضاء") else "، ".join(b.get("visible_to", []))
            with st.container(border=True):
                st.markdown(
                    f"<div class='banner-admin-card'><div style='color:#D4AF37;font-size:20px;font-weight:900;'>📢 {b.get('text','')}</div>"
                    f"<div style='color:white;margin-top:8px;'>الجمهور: {audience_text}</div>"
                    f"<div style='color:#ddd;'>تنتهي: {expire_raw or 'غير محدد'} | الحالة: {'منتهية' if expired else 'فعالة'}</div></div>",
                    unsafe_allow_html=True
                )
                e1,e2 = st.columns(2)
                with e1:
                    if st.button("✏️ تعديل", key=f"edit_banner_{b.get('id',i)}", use_container_width=True):
                        st.session_state["edit_banner_id"] = b.get("id")
                        st.rerun()
                with e2:
                    if st.button("🗑️ حذف", key=f"delete_banner_{b.get('id',i)}", use_container_width=True):
                        banners = [x for x in banners if x.get("id") != b.get("id")]
                        save_banners(banners)
                        st.rerun()

                if st.session_state.get("edit_banner_id") == b.get("id"):
                    with st.form(f"edit_banner_form_{b.get('id',i)}"):
                        nt = st.text_area("نص الرسالة أو البيان", value=b.get("text",""), key=f"et_{b.get('id',i)}")
                        nc = st.color_picker("اللون", value=b.get("color","#D4AF37"), key=f"ec_{b.get('id',i)}")
                        nimg = st.file_uploader("الصورة", type=["jpg","jpeg","png"], key=f"eimg_{b.get('id',i)}")
                        na = st.radio("المستلمون", ["كل الأعضاء","أعضاء محددين"], index=0 if b.get("audience") in ("الكل","كل الأعضاء") else 1, horizontal=True, key=f"ea_{b.get('id',i)}")
                        nvis = []
                        if na == "أعضاء محددين":
                            edit_candidates = [u for u in users if u.get("status") == "active"]
                            default_ids = [u.get("id") for u in edit_candidates if u.get("username") in b.get("visible_to", [])]
                            selected_edit_ids = _recipient_groups(edit_candidates, f"banner_edit_recipient_{b.get('id',i)}", default_ids)
                            nvis = [u.get("username") for u in edit_candidates if u.get("id") in selected_edit_ids]
                        new_expire = st.date_input("تاريخ الانتهاء", value=expire_dt.date() if expire_dt else datetime.now().date(), key=f"ed_{b.get('id',i)}")
                        ec1,ec2 = st.columns(2)
                        with ec1:
                            if st.form_submit_button("💾 حفظ التعديل", use_container_width=True):
                                b["text"] = nt.strip()
                                b["color"] = nc
                                b["audience"] = "الكل" if na == "كل الأعضاء" else "اعضاء محددين"
                                b["visible_to"] = nvis
                                if nimg:
                                    b["image_data"] = base64.b64encode(nimg.getvalue()).decode("utf-8")
                                    b["image_type"] = nimg.type
                                old_dt = expire_dt or datetime.now()
                                b["expire_at"] = datetime.combine(new_expire, old_dt.time()).isoformat()
                                save_banners(banners)
                                st.session_state.pop("edit_banner_id", None)
                                st.rerun()
                        with ec2:
                            if st.form_submit_button("إلغاء", use_container_width=True):
                                st.session_state.pop("edit_banner_id", None)
                                st.rerun()

    with tab_search:
        st.markdown("<h3 style='color:#D4AF37;'>🔎 البحث عن عضو أو مدير إدارة أو مدير عام أو قيادي</h3>", unsafe_allow_html=True)
        global_people_search()
        st.markdown("---")
        st.markdown("<h3 style='color:#D4AF37;'>👤 الملف الشخصي</h3>", unsafe_allow_html=True)
        admin_user = next((u for u in load_users() if u.get("role") == "admin"), st.session_state.user)
        st.write(f"**الاسم:** {admin_user.get('username','')}")
        st.write(f"**الوظيفة:** {role_label(admin_user.get('role'))}")
        st.write(f"**البريد الإلكتروني:** {admin_user.get('email') or 'غير مسجل'}")
        st.write(f"**الحالة:** {admin_user.get('status','')}")

    with tab_account:
        st.markdown("<h3 style='color:#D4AF37;'>👤 بيانات حساب الادمن</h3>", unsafe_allow_html=True)
        users = load_users()
        admin_user = next((u for u in users if u.get("role") == "admin"), None)
        if admin_user:
            with st.form("admin_profile_edit_form_v4"):
                display_name = st.text_input("الاسم", value=admin_user.get("display_name") or admin_user.get("username") or "admin")
                job_title = st.text_input("الوظيفة", value=admin_user.get("job_title") or "")
                insurance_area = st.text_input("المنطقة التأمينية", value=admin_user.get("insurance_area") or "")
                login_name = st.text_input("اسم الدخول", value=admin_user.get("username") or "admin")
                admin_photo = st.file_uploader("الصورة الشخصية", type=["jpg","jpeg","png"], key="admin_profile_photo_upload")
                if st.form_submit_button("💾 حفظ بيانات الادمن", use_container_width=True, type="primary"):
                    login_name = login_name.strip() or "admin"
                    if login_name != admin_user.get("username") and any(u.get("username") == login_name and u.get("id") != admin_user.get("id") for u in users):
                        st.error("اسم الدخول مستخدم بالفعل")
                    else:
                        admin_user["display_name"] = display_name.strip() or login_name
                        admin_user["username"] = login_name
                        admin_user["job_title"] = job_title.strip()
                        admin_user["insurance_area"] = insurance_area.strip()
                        if admin_photo:
                            admin_user["profile_photo"] = base64.b64encode(admin_photo.getvalue()).decode("utf-8")
                            admin_user["profile_photo_type"] = admin_photo.type
                        save_users(users)
                        st.session_state.user = admin_user
                        st.success("تم حفظ اسم الادمن والوظيفة والمنطقة التأمينية وبيانات الدخول")
                        st.rerun()
        st.info("تظل صفة الحساب ظاهرة دائمًا باسم: الادمن. ويمكن لمن لديه صلاحية مراسلة الادمن رؤيته بالاسم والوظيفة والمنطقة التأمينية.")
        recovery_settings_page()
        st.markdown("---")
        change_password_page()

    st.markdown("---")
    if st.button("🚪 تسجيل الخروج", use_container_width=True, key="admin_control_logout_v3"):
        st.session_state.user = None
        st.session_state.role = None
        st.session_state.page = "login"
        st.rerun()

    if st.button("⬅️ العودة للرئيسية", use_container_width=True, key="back_admin_panel"):
        st.session_state.page = "الرئيسية"
        st.rerun()

def recovery_settings_page():
    st.markdown("<h2 style='text-align:center; color:#C9A961'>تأكيد البريد الالكتروني</h2>", unsafe_allow_html=True)
    users = load_users()
    user = next((u for u in users if u["id"] == st.session_state.user["id"]), None)
    st.info("هذا البريد سيستخدم لاسترجاع حسابك في حالة الفقد")
    email = st.text_input("البريد الالكتروني", value=user.get("email",""))
    if user["role"] == "admin":
        recovery_email = st.text_input("ايميل استرجاع اضافي للادمن", value=user.get("recovery_email",""))
    else: recovery_email = user.get("recovery_email","")
    if st.button("حفظ البريد", use_container_width=True):
        user["email"] = email
        if user["role"] == "admin": user["recovery_email"] = recovery_email
        save_users(users); st.session_state.user = user; st.success("تم حفظ البريد بنجاح")

def change_password_page():
    st.markdown("<h1 style='text-align:center; color:#C9A961'>تغيير كلمة السر</h1>", unsafe_allow_html=True)
    old_pass = st.text_input("كلمة السر القديمة", type="password")
    new_pass = st.text_input("كلمة السر الجديدة", type="password")
    if st.button("تغيير", use_container_width=True):
        if bcrypt.checkpw(old_pass.encode(), st.session_state.user["password"].encode()):
            users = load_users()
            for user in users:
                if user["id"] == st.session_state.user["id"]:
                    user["password"] = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt()).decode()
                    save_users(users); st.session_state.user = user; st.success("تم تغيير الباسورد"); st.rerun()
        else: st.error("كلمة السر القديمة غلط")

def set_password_page():
    st.markdown("<div class='login-shell'>", unsafe_allow_html=True)
    st.markdown("<h1 style='text-align:center;color:#D4AF37!important;'>🔐 تفعيل العضوية</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align:center;color:#fff!important;font-weight:900;'>أدخل كلمة السر التي تريد استخدامها لحسابك</p>", unsafe_allow_html=True)
    new_pass = st.text_input("كلمة السر الجديدة", type="password", key="first_new_pass")
    confirm_pass = st.text_input("تأكيد كلمة السر", type="password", key="first_confirm_pass")
    profile_photo = st.file_uploader("الصورة الشخصية", type=["jpg","jpeg","png"], key="first_profile_photo")
    if st.button("حفظ كلمة السر والدخول", type="primary", use_container_width=True, key="save_first_password"):
        if not new_pass:
            st.error("اكتب كلمة السر أولاً")
            return
        if new_pass != confirm_pass:
            st.error("كلمتا السر غير متطابقتين")
            return
        users = load_users()
        target_id = st.session_state.get("temp_user_id")
        target_name = str(st.session_state.get("temp_user", "")).strip().casefold()
        target = None
        for user in users:
            if target_id is not None and user.get("id") == target_id:
                target = user
                break
            if normalize_person_name(user.get("username")) == normalize_person_name(target_name) or normalize_person_name(user.get("display_name")) == normalize_person_name(target_name):
                target = user
                break
        if not target:
            st.error("تعذر العثور على العضوية. ارجع لتسجيل الدخول وأعد إدخال اسم العضو كما هو.")
            return
        target["password"] = bcrypt.hashpw(new_pass.encode(), bcrypt.gensalt()).decode()
        target["password_set"] = True
        if profile_photo:
            target["profile_photo"] = base64.b64encode(profile_photo.getvalue()).decode("utf-8")
            target["profile_photo_type"] = profile_photo.type
        if target.get("status") in (None, "suspended_temp"):
            target["status"] = "active"
        save_users(users)
        st.session_state.user = target
        st.session_state.role = target.get("role")
        st.session_state.page = "الرئيسية"
        st.session_state.pop("temp_user", None)
        st.session_state.pop("temp_user_id", None)
        st.success("تم تفعيل العضوية وتسجيل الدخول بنجاح")
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

def personal_control_page():
    user = st.session_state.get("user") or {}
    role = user.get("role")
    if not user or role == "admin":
        return

    st.markdown("<h1 style='color:#D4AF37;text-align:center;'>⚙️ لوحة تحكمى</h1>", unsafe_allow_html=True)
    st.markdown("<p style='color:white;text-align:center;font-weight:800;'>ملفك الشخصي وصلاحياتك وحسابك ورسائلك</p>", unsafe_allow_html=True)

    tab_profile, tab_permissions, tab_messages, tab_account = st.tabs(
        ["👤 الملف الشخصي", "🛡️ صلاحياتى", "📢 الرسائل والبيانات", "🔐 الحساب"]
    )

    with tab_profile:
        st.markdown("<h3 style='color:#D4AF37;'>بيانات العضوية</h3>", unsafe_allow_html=True)
        st.write(f"**الاسم:** {user.get('username','')}")
        st.write(f"**الوظيفة:** {user.get('job_title') or role_label(role)}")
        st.write(f"**تاريخ الميلاد:** {user.get('birth_date') or 'غير مسجل'}")
        st.write(f"**المنطقة التأمينية:** {user.get('insurance_area') or 'غير مسجلة'}")
        st.write(f"**البريد الإلكتروني:** {user.get('email') or 'غير مسجل'}")
        st.write(f"**الحالة:** {user.get('status','')}")
        st.info("إيقاف حساب الدخول لا يحذف القضايا أو المستندات أو الرسائل المحفوظة.")

    with tab_permissions:
        st.markdown("<h3 style='color:#D4AF37;'>الصلاحيات المقررة لحسابك</h3>", unsafe_allow_html=True)
        targets = [
            ("الادمن","admin"), ("عضو قانونى","member"), ("مدير إدارة","manager"),
            ("مدير عام","general_manager"), ("قيادي مركز رئيسى 1","leader1"),
            ("قيادي مركز رئيسى 2","leader2"), ("قيادي مركز رئيسى 3","leader3")
        ]
        allowed = [label for label, target_role in targets if can_message(role, target_role)]
        if allowed:
            st.success("يمكنك إرسال رسالة أو بيان إلى: " + "، ".join(allowed))
        else:
            st.info("لا توجد صلاحية إرسال رسائل لهذا الحساب.")

    with tab_messages:
        messaging_page()

    with tab_account:
        st.markdown("<h3 style='color:#D4AF37;'>👤 بيانات حساب الادمن</h3>", unsafe_allow_html=True)
        users = load_users()
        admin_user = next((u for u in users if u.get("role") == "admin"), None)
        if admin_user:
            with st.form("admin_profile_edit_form_v4"):
                display_name = st.text_input("الاسم", value=admin_user.get("display_name") or admin_user.get("username") or "admin")
                job_title = st.text_input("الوظيفة", value=admin_user.get("job_title") or "")
                insurance_area = st.text_input("المنطقة التأمينية", value=admin_user.get("insurance_area") or "")
                login_name = st.text_input("اسم الدخول", value=admin_user.get("username") or "admin")
                if st.form_submit_button("💾 حفظ بيانات الادمن", use_container_width=True, type="primary"):
                    login_name = login_name.strip() or "admin"
                    if login_name != admin_user.get("username") and any(u.get("username") == login_name and u.get("id") != admin_user.get("id") for u in users):
                        st.error("اسم الدخول مستخدم بالفعل")
                    else:
                        admin_user["display_name"] = display_name.strip() or login_name
                        admin_user["username"] = login_name
                        admin_user["job_title"] = job_title.strip()
                        admin_user["insurance_area"] = insurance_area.strip()
                        save_users(users)
                        st.session_state.user = admin_user
                        st.success("تم حفظ اسم الادمن والوظيفة والمنطقة التأمينية وبيانات الدخول")
                        st.rerun()
        st.info("تظل صفة الحساب ظاهرة دائمًا باسم: الادمن. ويمكن لمن لديه صلاحية مراسلة الادمن رؤيته بالاسم والوظيفة والمنطقة التأمينية.")
        recovery_settings_page()
        st.markdown("---")
        change_password_page()

    st.markdown("---")
    if st.button("⬅️ العودة للرئيسية", use_container_width=True, key="personal_control_back_v3"):
        st.session_state.page = "الرئيسية"
        st.rerun()
    if st.button("🚪 تسجيل الخروج", use_container_width=True, key="personal_control_logout_v3"):
        st.session_state.user = None
        st.session_state.role = None
        st.session_state.page = "login"
        st.rerun()

# ===== عرض الرسائل والبيانات أعلى الصفحة =====
def _early_load_banners():
    try:
        data = load_json_persistent("banners", BANNERS_FILE, [])
        return data if isinstance(data, list) else []
    except Exception:
        return []

def show_banners():
    if not st.session_state.get("user"):
        return
    username = st.session_state.user.get("username", "")
    now = datetime.now()
    for b in _early_load_banners():
        try:
            expire = datetime.fromisoformat(str(b.get("expire_at") or b.get("expire")))
        except Exception:
            continue
        if expire <= now:
            continue
        audience = b.get("audience", "الكل")
        if audience not in ("الكل", "كل الأعضاء") and username not in b.get("visible_to", []):
            continue
        st.markdown(
            f"<div style='background:linear-gradient(90deg,{b.get('color','#D4AF37')},#ffffff22);color:#fff;padding:16px;border-radius:16px;margin:0 0 14px;font-weight:900;font-size:20px;text-align:center;border:2px solid {b.get('color','#D4AF37')};'>📢 {b.get('text','')}</div>",
            unsafe_allow_html=True
        )

# ======= الجزء الاول: الاساسيات ============
# ============================================
import streamlit as st
import pandas as pd
import json
import os
import io
import smtplib
import secrets
import base64
import arabic_reshaper # جديد
from bidi.algorithm import get_display # جديد
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

st.set_page_config(page_title="إدارة القضايا", layout="wide", page_icon="⚖️")

# ===== تثبيت ظهور أسماء الأزرار والأيقونات بوضوح =====
st.markdown("""
<style>
/* كل الأزرار: خلفية قضائية داكنة ونص وأيقونات بيضاء ظاهرة */
.stButton > button,
[data-testid="stButton"] > button,
[data-testid="stDownloadButton"] > button,
[data-testid="stFormSubmitButton"] > button {
    background: linear-gradient(135deg,#16243d,#243b63) !important;
    color: #ffffff !important;
    border: 2px solid #D4AF37 !important;
    border-radius: 15px !important;
    font-weight: 900 !important;
    min-height: 52px !important;
    opacity: 1 !important;
    box-shadow: 0 6px 16px rgba(0,0,0,.22) !important;
}
.stButton > button:hover,
[data-testid="stButton"] > button:hover,
[data-testid="stDownloadButton"] > button:hover {
    background: linear-gradient(135deg,#203455,#2e4c7c) !important;
    color:#ffffff !important;
    border-color:#f0cf62 !important;
}
.stButton > button p,
[data-testid="stButton"] > button p,
[data-testid="stDownloadButton"] > button p,
.stButton > button span,
[data-testid="stButton"] > button span {
    color:#ffffff !important;
    opacity:1 !important;
    font-weight:900 !important;
}
.stButton > button svg,
[data-testid="stButton"] > button svg,
[data-testid="stDownloadButton"] > button svg {
    color:#ffffff !important;
    fill:#ffffff !important;
    stroke:#ffffff !important;
    opacity:1 !important;
}
/* تبويبات الدخول ظاهرة بوضوح */
div[data-testid="stTabs"] button,
div[data-testid="stTabs"] button p {
    color:#ffffff !important;
    font-weight:900 !important;
}
div[data-testid="stTabs"] button[aria-selected="true"],
div[data-testid="stTabs"] button[aria-selected="true"] p {
    color:#D4AF37 !important;
}
/* تسميات الحقول */
div[data-testid="stTextInput"] label,
div[data-testid="stPasswordInput"] label,
div[data-testid="stDateInput"] label,
div[data-testid="stSelectbox"] label,
div[data-testid="stTextArea"] label {
    color:#ffffff !important;
    font-weight:900 !important;
}
</style>
""", unsafe_allow_html=True)
def fix_arabic(text):
    if not text:
        return ""
    text = str(text)
    reshaped_text = arabic_reshaper.reshape(text)
    bidi_text = get_display(reshaped_text)
    return bidi_text
# ====== دالة التصدير للاكسل RTL صح 100% ======
def to_excel(df):
    df = df.fillna('-') # عشان ميضربش لو في خلايا فاضية
    df = df.astype(str) # نحول كله لنص

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='التقرير')
        worksheet = writer.sheets['التقرير']
        worksheet.sheet_view.rightToLeft = True # RTL

        for col in worksheet.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            worksheet.column_dimensions[column].width = max_length + 2

        from openpyxl.styles import Alignment
        for row in worksheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    return output.getvalue()

# ====== دالة التصدير للورد ======
def to_word(df, title, region):
    doc = Document()
    doc.add_heading(fix_arabic('الهيئة القومية للتأمين الاجتماعى'), 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(fix_arabic('الإدارة المركزية للإدارات القانونية'), 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(fix_arabic('الإدارة العامة للقضايا'), 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(fix_arabic(f'ديوان عام {region}'), 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(fix_arabic(title), 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = fix_arabic(str(col_name))
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = fix_arabic(str(val))
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(fix_arabic('تفضلوا بقبول وافر الاحترام\n'))
    p.add_run(fix_arabic('عضو الادارة.................. مدير الإدارة..................\n'))
    p.add_run(fix_arabic(f'تحر في {datetime.now().strftime("%Y-%m-%d")}'))

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# ====== دالة التصدير للـ PDF ======
def to_pdf(df, title, region):
    pdf = FPDF(orientation='L', unit='mm', format='A4') # عرض
    pdf.add_page()
    pdf.add_font('Cairo', '', 'Cairo-Regular.ttf', uni=True)

    pdf.set_font('Cairo', '', 16)
    pdf.cell(0, 10, fix_arabic('الهيئة القومية للتأمين الاجتماعى'), 0, 1, 'C')
    pdf.set_font('Cairo', '', 12)
    pdf.cell(0, 8, fix_arabic('الإدارة المركزية للإدارات القانونية'), 0, 1, 'C')
    pdf.cell(0, 8, fix_arabic('الإدارة العامة للقضايا'), 0, 1, 'C')
    pdf.cell(0, 8, fix_arabic(f'ديوان عام {region}'), 0, 1, 'C')
    pdf.cell(0, 8, fix_arabic(title), 0, 1, 'C')
    pdf.ln(5)

    pdf.set_font('Cairo', '', 8)
    col_width = 280 / len(df.columns)
    row_height = 8

    for col in df.columns:
        pdf.cell(col_width, row_height, fix_arabic(str(col)), 1, 0, 'C')
    pdf.ln()

    for _, row in df.iterrows():
        for item in row:
            pdf.cell(col_width, row_height, fix_arabic(str(item)), 1, 0, 'C')
        pdf.ln()

    pdf.ln(8)
    pdf.set_font('Cairo', '', 11)
    pdf.cell(0, 8, fix_arabic('تفضلوا بقبول وافر الاحترام'), 0, 1, 'R')
    pdf.ln(5)
    
    pdf.ln(8)
    pdf.set_font('Cairo', '', 11)
    pdf.cell(0, 8, fix_arabic('تفضلوا بقبول وافر الاحترام'), 0, 1, 'R')
    pdf.ln(5)
    
    pdf.set_font('Cairo', '', 10)
    cell_w = 90
    pdf.cell(cell_w, 8, fix_arabic('العضو القانوني'), 0, 0, 'C')
    pdf.cell(cell_w, 8, fix_arabic('مدير إدارة القضايا'), 0, 0, 'C')
    pdf.cell(cell_w, 8, fix_arabic('مدير عام الإدارات القانونية'), 0, 1, 'C')
    
    pdf.ln(10)
    pdf.cell(cell_w, 8, '..................', 0, 0, 'C')
    pdf.cell(cell_w, 8, '..................', 0, 0, 'C')
    pdf.cell(cell_w, 8, '..................', 0, 1, 'C')

    pdf.ln(5)
    pdf.set_font('Cairo', '', 10)
    pdf.cell(0, 8, fix_arabic(f'تحر في {datetime.now().strftime("%d-%m-%Y")}'), 0, 1, 'L')

    return pdf.output()  # في fpdf2 الجديدة ده بيرجع bytes اصلا
# ====== دالة حفظ صحيفة الدعوى === 
# ====== دالة حفظ صحيفة الدعوى ======
def create_paper_pdf(case_data):
    if not os.path.exists("papers"): os.makedirs("papers")
    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.add_page()
    pdf.add_font('Cairo', '', 'Cairo-Regular.ttf', uni=True)
    pdf.set_font('Cairo', '', 14); pdf.set_right_margin(15)
    pdf.cell(0,10,fix_arabic(f"صحيفة {case_data.get('مسندة_ل','')}"),ln=1,align='R')
    pdf.ln(5)
    pdf.cell(0,10,fix_arabic(f"محكمة: {case_data.get('محكمة_اسم','')}"),ln=1,align='R')
    pdf.cell(0,10,fix_arabic(f"رقم: {case_data.get('رقم','')} لسنة {case_data.get('سنة','')}"),ln=1,align='R')
    pdf.cell(0,10,fix_arabic(f"المدعي: {case_data.get('مدعي','')}"),ln=1,align='R')
    pdf.cell(0,10,fix_arabic(f"ضد: {case_data.get('مدعي_عليه','')}"),ln=1,align='R')
    pdf.multi_cell(0,10,fix_arabic(f"الموضوع: {case_data.get('موضوع','')}"),align='R')
    name = f"papers/صحيفة_{case_data.get('رقم')}_{case_data.get('سنة')}.pdf"; pdf.output(name); return name

import base64
from io import BytesIO

def print_case_report(case):
    نوع = case.get('نوع', '').lower()
    if 'استئناف' in نوع:
        طرف1_عنوان = "المستأنف"
        طرف2_عنوان = "المستأنف ضده"
    elif 'طعن' in نوع:
        طرف1_عنوان = "الطاعن"
        طرف2_عنوان = "المطعون ضده"
    else:  # دعوى عادية
        طرف1_عنوان = "المدعي"
        طرف2_عنوان = "المدعى عليه"

    html = f"""
    <html dir="rtl" lang="ar">
    <head>
    <meta charset="UTF-8">
    <style>
        @page {{ size: A4; margin: 1.5cm; }}
        body {{ font-family: 'Arial'; direction: rtl; text-align: right; color: #000; background: #f8f9fa; }}
        .header {{ text-align: center; padding: 25px; margin-bottom: 25px; background: linear-gradient(135deg, #1E2A47 0%, #D4AF37 100%); color: #FFF; border-radius: 15px; box-shadow: 0 5px 15px rgba(0,0,0,0.2); }}
        .logo {{ font-size: 22px; font-weight: 900; color: #FFF; text-shadow: 2px 2px 4px rgba(0,0,0,0.3); }}
        .sub {{ font-size: 16px; color: #FFF9E6; margin: 8px 0; }}
        .title {{ text-align: center; font-size: 26px; font-weight: 900; color: #1E2A47; margin: 25px 0; border: 3px solid #D4AF37; padding: 15px; border-radius: 15px; background: linear-gradient(90deg, #FFF9E6, #FFF); box-shadow: 0 3px 10px rgba(212,175,55,0.3); }}
        .section {{ padding: 20px; border-radius: 15px; margin-bottom: 25px; box-shadow: 0 4px 15px rgba(0,0,0,0.1); border: 2px solid transparent; }}
        .section-title {{ font-weight: 900; font-size: 20px; color: #FFF; margin-bottom: 20px; text-align: center; padding: 12px; border-radius: 10px; }}
        .sec1 {{ background: linear-gradient(135deg, #1E2A47, #3498db); border-color: #1E2A47; }}
        .sec2 {{ background: linear-gradient(135deg, #27ae60, #2ecc71); border-color: #27ae60; }}
        .sec3 {{ background: linear-gradient(135deg, #8e44ad, #9b59b6); border-color: #8e44ad; }}
        .sec4 {{ background: linear-gradient(135deg, #c0392b, #e74c3c); border-color: #c0392b; }}
        .row {{ display: flex; justify-content: space-between; margin-bottom: 12px; background: linear-gradient(90deg, #fff, #f8f9fa); padding: 12px; border-radius: 8px; border-right: 4px solid #D4AF37; }}
        .label {{ font-weight: 900; color: #1E2A47; width: 35%; font-size: 15px; }}
        .value {{ width: 65%; color: #000; font-weight: 700; font-size: 15px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 15px; border-radius: 10px; overflow: hidden; box-shadow: 0 3px 10px rgba(0,0,0,0.1); }}
        th {{ background: linear-gradient(135deg, #1E2A47, #34495e); color: #D4AF37; padding: 12px; border: none; text-align: center; font-size: 16px; font-weight: 900; }}
        td {{ padding: 12px; border-bottom: 1px solid #ddd; text-align: center; background: #fff; }}
        tr:nth-child(even) td {{ background: #f8f9fa; }}
        tr:hover td {{ background: #FFF9E6; }}
    </style>
    </head>
    <body>
    
    <div class="header">
        <div class="logo">الهيئة القومية للتأمين الاجتماعي</div>
        <div class="sub">الإدارة المركزية للإدارات القانونية</div>
        <div class="sub">الإدارة العامة للشئون القانونية منطقة: _____________</div>
    </div>

    <div class="title">📄 تقرير تفاصيل القضية رقم {case.get('رقم')} لسنة {case.get('سنة')}</div>

    <div class="section sec1">
        <div class="section-title sec1">1- بيانات القضية</div>
        <div class="row"><div class="label">رقم القضية:</div><div class="value">{case.get('رقم')}</div></div>
        <div class="row"><div class="label">السنة:</div><div class="value">{case.get('سنة')}</div></div>
        <div class="row"><div class="label">النوع:</div><div class="value">{case.get('نوع')}</div></div>
        <div class="row"><div class="label">المحكمة:</div><div class="value">{case.get('محكمة_اسم')} {f'- مأمورية {case.get("مأمورية")}' if case.get('مأمورية') else ''}</div></div>
        <div class="row"><div class="label">الدائرة:</div><div class="value">{case.get('دائرة')}</div></div>
        <div class="row"><div class="label">الحالة:</div><div class="value">{case.get('حالة')}</div></div>
        <div class="row"><div class="label">الموضوع:</div><div class="value">{case.get('موضوع')}</div></div>
    </div>

    <div class="section sec2">
        <div class="section-title sec2">2- بيانات الخصوم</div>
        <div class="row"><div class="label">{طرف1_عنوان}:</div><div class="value">{case.get('مدعي')}</div></div>
        <div class="row"><div class="label">{طرف2_عنوان}:</div><div class="value">{case.get('مدعي_عليه')}</div></div>
    </div>
    """

    if case.get("جلسات"):
        html += """
        <div class="section sec3">
            <div class="section-title sec3">3- الجلسات والإجراءات</div>
            <table>
                <tr><th>م</th><th>الرول</th><th>الجلسات</th><th>الإجراءات</th><th>ملاحظات</th></tr>
        """
        for i, ج in enumerate(case["جلسات"], 1):
            html += f"<tr><td>{i}</td><td>{ج.get('الرول')}</td><td>{ج.get('تاريخ')}</td><td>{ج.get('الاجراء')}</td><td>{ج.get('ملاحظات')}</td></tr>"
        html += "</table></div>"

    if case.get('حالة') == 'منتهية':
        html += f"""
        <div class="section sec4">
            <div class="section-title sec4">4- منطوق الحكم</div>
            <div class="row"><div class="label">تاريخ الحكم:</div><div class="value">{case.get('تاريخ_الحكم')}</div></div>
            <div class="row"><div class="label">مسندة لـ:</div><div class="value">{case.get('مسندة_ل_الحكم')}</div></div>
            <div class="row"><div class="label">المنطوق:</div><div class="value">{case.get('منطوق_الحكم')}</div></div>
        </div>
        """
    
    html += "</body></html>"
    return html
# ====== دالة التحميل والحفظ الوحيدة ======
DATA_FILE = os.path.join(LOCAL_DATA_DIR, "cases_data.json")
TOKENS_FILE = os.path.join(LOCAL_DATA_DIR, "tokens.json")
UPLOAD_FOLDER = os.path.join(LOCAL_DATA_DIR, "uploads")
os.makedirs(UPLOAD_FOLDER,exist_ok=True)

def _current_user_identity():
    user = st.session_state.get("user") or {}
    return user.get("id"), user.get("username", "")

def _case_belongs_to_current_user(case):
    uid, username = _current_user_identity()
    if case.get("owner_user_id") is not None and uid is not None:
        return case.get("owner_user_id") == uid
    if case.get("assigned_to"):
        return case.get("assigned_to") == username
    return False

def get_my_cases(data):
    return [c for c in data.get("cases", []) if _case_belongs_to_current_user(c)]

def load_data():
    data=load_json_persistent("data", DATA_FILE, {"cases":[],"library":[]})
    if not isinstance(data,dict): data={}
    data.setdefault("cases",[]); data.setdefault("library",[])
    changed = False
    admin = next((u for u in load_users() if u.get("role") == "admin"), None)
    admin_id = admin.get("id") if admin else None
    admin_username = admin.get("username") if admin else ""
    for case in data["cases"]:
        if "owner_user_id" not in case and case.get("assigned_to"):
            owner = next((u for u in load_users() if u.get("username") == case.get("assigned_to")), None)
            if owner:
                case["owner_user_id"] = owner.get("id")
                changed = True
        if "assigned_to" not in case and admin_id is not None:
            case["assigned_to"] = admin_username
            case["owner_user_id"] = admin_id
            changed = True
    if changed:
        save_json_persistent("data", DATA_FILE, data)
    return data

def save_data(data):
    data.setdefault("cases",[]); data.setdefault("library",[])
    save_json_persistent("data", DATA_FILE, data)

def load_tokens():
    return load_json_persistent("tokens", TOKENS_FILE, {"tokens":[]})

def save_tokens(tokens_data):
    save_json_persistent("tokens", TOKENS_FILE, tokens_data)


# ===== دوال التنبيهات ======
from datetime import datetime, timedelta

def get_alert_cases():
    data = load_data()
    today = datetime.now().date()
    username = (st.session_state.user or {}).get("username", "")
    all_cases = [c for c in data.get("cases", []) if c.get("assigned_to") == username]
    alerts = {"sessions": [], "appeals": []}
    
    for case in all_cases:
        if case.get("حالة") == "متداولة" and case.get("تاريخ_جلسة"):
            try:
                session_date = datetime.strptime(case["تاريخ_جلسة"], "%Y-%m-%d").date()
                days_left = (session_date - today).days
                if 0 <= days_left <= 7:
                    case_copy = case.copy()
                    case_copy["days_left"] = days_left
                    alerts["sessions"].append(case_copy)
            except Exception: 
                pass
                
        if case.get("حالة") == "منتهية" and case.get("مسندة_ل_الحكم") == "الضد" and case.get("تاريخ_الحكم"):
            try:
                judgment_date = datetime.strptime(case["تاريخ_الحكم"], "%Y-%m-%d").date()
                appeal_days = 40 if case.get("نوع") == "دعوى" else 60
                last_appeal_day = judgment_date + timedelta(days=appeal_days)
                notify_start = last_appeal_day - timedelta(days=15)
                days_left_appeal = (last_appeal_day - today).days
                if notify_start <= today <= last_appeal_day and days_left_appeal >= 0:
                    case_copy = case.copy()
                    case_copy["days_left_appeal"] = days_left_appeal
                    case_copy["deadline"] = last_appeal_day.strftime("%Y-%m-%d")
                    alerts["appeals"].append(case_copy)
            except Exception: 
                pass
                
    return alerts

def send_alert_email(to_email, alerts):
    subject = f"🔔 تنبيهات ادارة القضايا - {datetime.now().strftime('%Y-%m-%d')}"
    body = "<div style='direction:rtl; text-align:right; font-family:Arial;'>"
    body += "<h2 style='color:#C9A961; text-align:center;'>مركز التنبيهات</h2>"
    
    if alerts["sessions"]:
        body += "<h3 style='color:#FFD700;'>⚖️ جلسات خلال 7 ايام</h3>"
        for case in alerts["sessions"]:
            body += f"<p style='border:1px solid #C9A961; padding:10px; border-radius:8px;'>"
            body += f"<b>رقم القضية:</b> {case.get('رقم_كامل','')}<br>"
            body += f"<b>الموضوع:</b> {case.get('موضوع_الدعوى','')}<br>"
            body += f"<b>الجلسة:</b> {case.get('تاريخ_جلسة')} - <b style='color:red;'>فاضل {case['days_left']} يوم</b></p>"
    else:
        body += "<p>✅ مفيش جلسات خلال 7 ايام</p>"
    
    if alerts["appeals"]:
        body += "<h3 style='color:#FF4500;'>📄 طعون خلال 15 يوم</h3>"
        for case in alerts["appeals"]:
            body += f"<p style='border:1px solid #FF4500; padding:10px; border-radius:8px;'>"
            body += f"<b>رقم القضية:</b> {case.get('رقم_كامل','')}<br>"
            body += f"<b>الموضوع:</b> {case.get('موضوع_الدعوى','')}<br>"
            body += f"<b style='color:red;'>اخر ميعاد للطعن: {case['deadline']} - فاضل {case['days_left_appeal']} يوم</b></p>"
    else:
        body += "<p>✅ مفيش طعون قريبة</p>"
    
    body += "</div>"
    
    try:
        send_email(to_email, subject, body)
        return True
    except Exception as e:
        st.error(f"فشل الارسال: {e}")
        return False
        
LIBRARY_SECTIONS = {
    "القوانين": "#FF4500", "القرارات الوزارية": "#FF8C00", "قرارات الهيئة": "#FFD700",
    "المنشورات الوزارية": "#ADFF2F", "منشورات الهيئة": "#32CD32", "الكتب الدورية": "#20B2AA",
    "تعليمات الهيئة": "#00CED1", "رسائل الهيئة": "#1E90FF", "المرصد الفنى": "#4169E1",
    "فتاوى لجنة الشئون القانونية بالوزارة": "#8A2BE2", "فتاوى الادارة المركزية للشئون القانونية": "#9400D3",
    "احكام المحكمة الدستورية العليا": "#DC143C", "احكام محكمة النقض": "#B22222", "احكام المحكمة الإدارية العليا": "#8B0000",
    "احكام المحاكم الاستئنافية": "#A0522D", "احكام محاكم القضاء الإدارى": "#D2691E", "احكام المحاكم الابتدائية": "#CD853F",
    "احكام المحكمة الإدارية": "#DEB887", "منشورات القضاء العادى": "#5F9EA0", "منشورات مجلس الدولة": "#4682B4",
    "فتاوى الجمعية العمومية": "#7B68EE", "صحف طعون": "#6A5ACD", "صحف استئنافات": "#483D8B",
    "صحف دعاوى": "#E6E6FA", "مذكرات دفاع": "#FFF0F5", "أخرى": "#808080"
}

ANWA3_MOSTANDAT = ["صحيفة دعوى","صحيفة استئناف","صحيفة طعن","مذكرة دفاع","حافظة مستندات","تقرير خبير","تقرير طب شرعى","تقرير لجنة طبية","صحيفة تجديد من الشطب","صحيفة تعجيل من الوقف","صورة حكم تمهيدى","أخرى"]

if "page" not in st.session_state: st.session_state.page="الرئيسية"
if "selected_case_id" not in st.session_state: st.session_state.selected_case_id=None

# ============= التصميم النهائي =============
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;700;900&display=swap');
*{font-family:'Cairo',sans-serif!important;}
html,body{direction:rtl;color:#FFF!important;}
.stApp{background:linear-gradient(180deg,#0A1428 0%,#1E2A47 100%);}
.marquee{background:linear-gradient(90deg,#D4AF37,#FFD700,#D4AF37);color:#0A1428;padding:12px;font-weight:900;font-size:16px;white-space:nowrap;overflow:hidden;border-radius:0 0 15px 15px;}
.marquee span{display:inline-block;animation:marquee 15s linear infinite;}
@keyframes marquee{0%{transform:translateX(-100%);}100%{transform:translateX(100%);}}
.main-title{color:#D4AF37;text-align:center;font-size:36px;font-weight:900;padding:15px 0;}
h1,h2,h3{color:#D4AF37!important;text-align:center!important;}
div[data-testid="column"]{display:flex;justify-content:center;}
[data-testid="stForm"] label,.stMarkdown{color:#FFF!important;font-weight:700;}
.stButton>button{width:100%!important;max-width:400px!important;border-radius:15px!important;font-size:18px!important;font-weight:900!important;padding:16px!important;color:#fff!important;background:linear-gradient(135deg,#1a2b4b,#263f69)!important;border:1px solid #D4AF37!important;}
</style>
""", unsafe_allow_html=True)

if not st.session_state.get("user"):
    st.session_state.page = "login"
    login_page()
    st.stop()

st.markdown("""<div class="marquee"><span>مع تحيات وليد حماد - الإدارة العامة للشئون القانونية بديوان عام منطقة البحيرة بالهيئة القومية للتأمين الاجتماعي</span></div>""", unsafe_allow_html=True)
st.markdown('<div class="main-title">⚖️ إدارة القضايا ⚖️</div>', unsafe_allow_html=True)

st.session_state.setdefault("top_notice_index", 0)
user_header = st.session_state.get("user") or {}
header_name = html.escape(str(user_header.get("display_name") or user_header.get("username") or ""))
st.markdown(f"<div style='text-align:center;color:#fff;font-size:20px;font-weight:900;margin:8px 0 14px;'>أهلاً {header_name}</div>", unsafe_allow_html=True)
st.markdown("<div style='color:#D4AF37;font-size:21px;font-weight:900;text-align:center;margin:8px 0;'>الرسائل والبيانات</div>", unsafe_allow_html=True)
show_top_notices()

# =========================================
# =======================================
# ==================================================
# ==================================================

if st.session_state.page == "الرئيسية":
    user = st.session_state.get("user") or {}
    is_admin = user.get("role") == "admin"

    st.markdown("<h2>الأقسام</h2>", unsafe_allow_html=True)
    def judicial_card(icon, title, page, key):
        st.markdown(f"<div class='judicial-card'><div class='judicial-card-icon'>{icon}</div><div class='judicial-card-title'>{title}</div></div>", unsafe_allow_html=True)
        if st.button(f"فتح {title}", use_container_width=True, key=key):
            st.session_state.page = page
            st.rerun()

    a,b,c = st.columns([1,1.2,1])
    with a: judicial_card("⚖️","تسجيل القضايا","تسجيل","card_add_v2")
    with b: judicial_card("📋","الحصر العام","الحصر","card_list_v2")
    with c: judicial_card("🔍","البحث عن دعوى","بحث","card_search_v2")

    a,b = st.columns(2)
    with a: judicial_card("📊","التقارير","تقارير","card_reports_v2")
    with b: judicial_card("📚","المكتبة القانونية","مكتبة","card_library_v2")

    a,b,c = st.columns([1,1.2,1])
    with b: judicial_card("🔴","مركز التنبيهات","التنبيهات","card_alerts_v2")

    a,b,c = st.columns([1,1.2,1])
    with b: judicial_card("🗄️","الأرشيف","الأرشيف","card_archive_v2")

    if is_admin:
        st.markdown("<div class='admin-control-top'><div style='text-align:center;color:#D4AF37;font-size:25px;font-weight:900;'>⚙️ لوحة التحكم</div><div style='text-align:center;color:#fff;'>إدارة الموقع والأعضاء والحسابات والبحث والرسائل</div></div>", unsafe_allow_html=True)
        if st.button("⚙️ لوحة التحكم", use_container_width=True, type="primary", key="open_admin_control_bottom_v5"):
            st.session_state.page = "لوحة التحكم"
            st.rerun()
    else:
        st.markdown("<div class='admin-control-top'><div style='text-align:center;color:#D4AF37;font-size:25px;font-weight:900;'>👤 لوحة تحكمى</div><div style='text-align:center;color:#fff;'>ملفك الشخصي وصلاحياتك وحسابك ورسائلك</div></div>", unsafe_allow_html=True)
        if st.button("👤 فتح لوحة تحكمى", use_container_width=True, type="primary", key="open_personal_control_bottom_v5"):
            st.session_state.page = "لوحة تحكمى"
            st.rerun()

    show_member_directory()

# ====== الجزء الثاني: تسجيل القضية ============
elif st.session_state.page == "تسجيل":
    data = load_data()
    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    st.markdown("<h2 style='color:#D4AF37; text-align:center'> تسجيل القضية</h2>", unsafe_allow_html=True)
    if st.button("⬅️ العودة للرئيسية", key="back_add", use_container_width=True):
        st.session_state.page = "الرئيسية"
        st.rerun()

    st.markdown("<label style='color:#FFF; font-weight:700; text-align:right; width:100%; display:block;'>نوع القضية</label>", unsafe_allow_html=True)
    نوع = st.selectbox("", ["دعوى", "استئناف", "طعن"], key="case_type_add")
    
    with st.form("form_case_add", clear_on_submit=True):  # <-- دي بتفضي الفورم لوحدها
        st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
        st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>1- بيانات المحكمة</div>", unsafe_allow_html=True)
        محكمة_اسم = st.text_input("اسم المحكمة", key="court_name_add")
        مأمورية = st.text_input("المأمورية", key="mamoria_add") if نوع == "استئناف" else ""
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
        st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>2- بيانات القضية</div>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1: رقم = st.text_input("رقم القضية / الاستئناف / الطعن", key="case_num_add")
        with col2: سنة = st.text_input("السنة القضائية", key="case_year_add")
        دائرة = st.text_input("الدائرة", key="circle_add")
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
        st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>3- بيانات الخصوم</div>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1: مدعي = st.text_input("اسم المدعى / المستأنف / الطاعن", key="plaintiff_add")
        with col2: مدعي_عليه = st.text_input("اسم المدعى عليه / المستأنف ضده / المطعون ضده", key="defendant_add")
        موضوع = st.text_area("موضوع القضية", height=100, key="subject_add")
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
        st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>4- بيانات الجلسة</div>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1: تاريخ_جلسة = st.date_input("تاريخ أول جلسة", value=datetime.now().date(), key="session_date_add")
        with col2: الرول = st.text_input("الرول", key="roll_add")
        الاجراء = st.text_input("الاجراء", key="reason_add") # <-- 1. غيرنا الاسم هنا
        ملاحظات = st.text_area("ملاحظات", height=100, key="notes_add")
        st.markdown("</div>", unsafe_allow_html=True)

        if st.form_submit_button("💾 حفظ القضية", use_container_width=True, type="primary"):
            if not رقم or not سنة: 
                st.error("❌ من فضلك ادخل رقم القضية والسنة")
            else:
                case_for_pdf = {"نوع":نوع,"رقم":رقم,"سنة":سنة,"دائرة":دائرة,"محكمة_اسم":محكمة_اسم,"مدعي":مدعي,"مدعي_عليه":مدعي_عليه,"موضوع":موضوع,"تاريخ_جلسة":str(تاريخ_جلسة)}
                paper_path = create_paper_pdf(case_for_pdf)

                new_case = {
                    "id": len(data["cases"])+1, "نوع": نوع, "محكمة_اسم": محكمة_اسم, "مأمورية": مأمورية, 
                    "رقم": رقم, "سنة": سنة, "دائرة": دائرة, "مدعي": مدعي, "مدعي_عليه": مدعي_عليه, 
                    "موضوع": موضوع, "تاريخ_جلسة": str(تاريخ_جلسة), "الرول": الرول, "الاجراء": الاجراء, # <-- 2. وهنا
                    "ملاحظات": ملاحظات, "جلسات": [], "مستندات": [paper_path], "حالة": "متداولة",
                    "owner_user_id": st.session_state.user.get("id"),
                    "assigned_to": st.session_state.user.get("username", "")
                }
                if الرول or الاجراء: # <-- 3. وهنا
                    new_case["جلسات"].append({"تاريخ":str(تاريخ_جلسة),"الرول":الرول,"الاجراء":الاجراء,"ملاحظات":ملاحظات}) # <-- 4. وهنا
                
                data["cases"].append(new_case)
                save_data(data)
                
                st.success(f"✅ تم الحفظ بنجاح -ونقلت للحصر العام- جاهز لتسجيل قضية جديدة")
                # =======================
# ================================================
# ====== الجزء الثالث: الحصر العام ============
# ================================================
elif st.session_state.page == "الحصر":
    data = load_data()
    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    st.markdown("<h2 style='color:#FFFFFF; text-align:center'>📊 الحصر العام الخارجي</h2>", unsafe_allow_html=True)
    if st.button("⬅️ العودة للرئيسية", use_container_width=True): st.session_state.page = "الرئيسية"; st.rerun()

    if st.session_state.get('open_from_search', False):
        st.session_state.open_from_search = False
        st.info("جاري فتح القضية من البحث...")

    if not data["cases"]:
        st.info("لا توجد قضايا مسجلة")
    else:
        for i, case in enumerate(data["cases"]):
            if "id" not in case: case["id"] = i + 1
            if "مستندات" not in case: case["مستندات"] = []

        save_data(data)

        # ======= تحديث اخر جلسة والاجراء من الجلسات =======
        for case in data["cases"]:
            if "جلسات" in case and case["جلسات"]:
                جلسات_مرتبة = sorted(case["جلسات"], key=lambda x: x.get("تاريخ","9999-12-31"), reverse=True)
                اخر_جلسة = جلسات_مرتبة[0]
                case["تاريخ_جلسة"] = اخر_جلسة.get("تاريخ","")
                case["الاجراء"] = اخر_جلسة.get("الاجراء","") # <-- هنا بقت الاجراء
                case["الحالة"] = اخر_جلسة.get("الحالة", case.get("الحالة","متداولة"))
        save_data(data)
        # ============================================

        # ======= التعديل 1: نجيب المتداولة من الحصر العام فقط =======
        active_cases = [c for c in get_my_cases(data) if c.get('حالة') == 'متداولة']
        # ==================================================

        sorted_cases = sorted(active_cases, key=lambda x: x.get("تاريخ_جلسة","9999-12-31"))
        total = len(active_cases) # اجمالي المتداولة فقط
        today = datetime.now().date()
        start_week = today - timedelta(days=(today.weekday() + 2) % 7) # السبت
        end_week = start_week + timedelta(days=5) # الخميس

        # ======= التعديل 2: جلسات الاسبوع من المتداولة فقط =======
        this_week = len([c for c in active_cases if c.get('تاريخ_جلسة') and start_week <= datetime.strptime(c['تاريخ_جلسة'],'%Y-%m-%d').date() <= end_week])
        # =========================================================

        # ====== التعديل 3: المحجوز للحكم من المتداولة فقط =======
        reserved = len([c for c in active_cases if any(k in str(c.get('الاجراء','')) for k in ['حكم', 'للحكم', 'الحكم'])])
        # =======================================================

        st.markdown(f"<div style='background:#1E2A47; padding:20px; border-radius:15px; border:2px solid #D4AF37; text-align:center; margin-bottom:20px'>", unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)
        with col1: st.markdown(f"<div style='font-size:28px; font-weight:900; color:#D4AF37'>📊 {total}</div><div style='font-size:18px; color:#FFF; font-weight:700'>اجمالي القضايا</div>", unsafe_allow_html=True)
        with col2: st.markdown(f"<div style='font-size:28px; font-weight:900; color:#4DA8DA'>📅 {this_week}</div><div style='font-size:18px; color:#FFF; font-weight:700'>جلسات هذا الاسبوع</div>", unsafe_allow_html=True)
        with col3: st.markdown(f"<div style='font-size:28px; font-weight:900; color:#FF5252'>⚖️ {reserved}</div><div style='font-size:18px; color:#FFF; font-weight:700'>المحجوز للحكم</div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("""
        <style>
.case-table {width:100%; border-collapse: collapse; font-size:11px; color:white; text-align:center; margin-bottom:5px;}
.case-table th {background:#D4AF37; color:#0B1426; padding:6px; font-weight:900;}
.case-table td {background:#1E2A47; padding:6px; border:1px solid #D4AF37; vertical-align:top;}
.plaintiff {background:#FFF3CD; color:#000; font-weight:700; border-radius:6px; padding:6px; font-size:11px;}
.plaintiff-hey2a {background:#DC3545!important; color:#FFF!important; font-weight:900; border-radius:6px; padding:6px; font-size:11px;}
.defendant {background:#CFF4FC; color:#000; font-weight:700; border-radius:6px; padding:6px; font-size:11px;}
.date-gold {color:#FFD700; font-weight:900;}
.status-green {color:#4CAF50; font-weight:900;}
        </style>
        """, unsafe_allow_html=True)

        for idx, case in enumerate(sorted_cases, 1):
            رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
            محكمة_كاملة = f"{case.get('نوع','')} {case.get('محكمة_اسم','')}"
            if case.get('مأمورية',''): محكمة_كاملة += f"<br>مأمورية {case.get('مأمورية','')}"
            دائرة_كاملة = f"{case.get('دائرة', '' )}" if case.get('دائرة', '' ) else ""
            if دائرة_كاملة: محكمة_كاملة += f"<br>دائرة {دائرة_كاملة}"

            نوع = case.get('نوع','')
            if نوع == "استئناف": لقب1, لقب2 = "المستأنف:", "المستأنف ضده:"
            elif نوع == "طعن": لقب1, لقب2 = "الطاعن:", "المطعون ضده:"
            else: لقب1, لقب2 = "المدعى:", "المدعى عليه:"

            if "الهيئة" in str(case.get('مدعي','')):
                طرف1_html = f"<div class='plaintiff-hey2a'><b>{لقب1}</b><br>{case.get('مدعي','')}</div>"
            else:
                طرف1_html = f"<div class='plaintiff'><b>{لقب1}</b><br>{case.get('مدعي','')}</div>"
            طرف2_html = f"<div class='defendant'><b>{لقب2}</b><br>{case.get('مدعي_عليه','')}</div>"
            خصوم = طرف1_html + "<div style='height:4px'></div>" + طرف2_html

            table_html = "<table class='case-table'><tr>"
            headers = ["م", "الرقم والسنة", "المحكمة والدائرة", "الخصوم", "الموضوع", "اخر جلسة", "الاجراء", "الحالة"] # <-- هنا بقت الاجراء
            for h in headers: table_html += f"<th>{h}</th>"
            table_html += "</tr>"
            table_html += f"<tr><td>{idx}</td><td>{رقم_كامل}</td><td>{محكمة_كاملة}</td><td>{خصوم}</td><td>{case.get('موضوع','')}</td><td class='date-gold'>{case.get('تاريخ_جلسة','')}</td><td>{case.get('الاجراء','')}</td><td class='status-green'>{case.get('حالة','متداولة')}</td></tr></table>" # <-- وهنا كمان
            st.markdown(table_html, unsafe_allow_html=True)

            c1, c2, c3 = st.columns([4,1,4])
            with c2:
                if st.button("فتح", key=f"open_{case['id']}", use_container_width=True):
                    st.session_state.selected_case_id = case['id']; st.session_state.page = "تفاصيل"; st.rerun()

# ===================================
# ================================================
# ============ الجزء الرابع: تفاصيل القضية ============
# ================================================
elif st.session_state.page == "تفاصيل":
    data = load_data()
    case = next((c for c in get_my_cases(data) if c["id"] == st.session_state.selected_case_id), None)
    if not case: st.error("القضية غير موجودة"); st.session_state.page = "الحصر"; st.rerun()
    if 'جلسات' not in case: case['جلسات'] = []
    if 'مستندات' not in case: case['مستندات'] = []

    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    st.markdown(f"<h2 style='color:#D4AF37; text-align:center'>📄 تفاصيل القضية رقم {case.get('رقم')} لسنة {case.get('سنة')}</h2>", unsafe_allow_html=True)

    if st.button("⬅️ العودة للحصر", use_container_width=True): st.session_state.page = "الحصر"; st.rerun()

    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:15px'>1- بيانات القضية</div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1: st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>رقم القضية</div><div style='color:#FFF; font-weight:900; font-size:22px'>{case.get('رقم')}</div></div>", unsafe_allow_html=True)
    with col2: st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>السنة</div><div style='color:#FFF; font-weight:900; font-size:22px'>{case.get('سنة')}</div></div>", unsafe_allow_html=True)
    with col3: دائرة_نص = f"{case.get('دائرة')}" if case.get('دائرة') else ""; st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>الدائرة</div><div style='color:#FFF; font-weight:900; font-size:18px'>{دائرة_نص}</div></div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1: st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>النوع</div><div style='color:#FFF; font-weight:900; font-size:18px'>{case.get('نوع')}</div></div>", unsafe_allow_html=True)
    with col2: محكمة_كاملة = f"{case.get('محكمة_اسم')}";
    if case.get('مأمورية'): محكمة_كاملة += f" - مأمورية {case.get('مأمورية')}"; st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>المحكمة</div><div style='color:#FFF; font-weight:700; font-size:14px'>{محكمة_كاملة}</div></div>", unsafe_allow_html=True)
    with col3: st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; margin-bottom:10px; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>الحالة</div><div style='color:#4CAF50; font-weight:900; font-size:18px'>{case.get('حالة')}</div></div>", unsafe_allow_html=True)
    st.markdown(f"<div style='background:#142038; padding:12px; border-radius:12px; border:1px solid #D4AF37; text-align:center'><div style='color:#D4AF37; font-weight:900; font-size:14px'>الموضوع</div><div style='color:#FFF; font-weight:700; font-size:16px'>{case.get('موضوع')}</div></div>", unsafe_allow_html=True)
    st.markdown("<style>div[data-testid='stExpander'] summary p{color:#D4AF37!important; font-weight:900!important;}</style>", unsafe_allow_html=True)
    with st.expander("✏️ تعديل بيانات القضية"):
        with st.form("edit_case_form"):
            col1, col2, col3 = st.columns(3)
            with col1: رقم_جديد = st.text_input("رقم القضية", value=case.get('رقم','')); سنة_جديد = st.text_input("السنة", value=case.get('سنة','')); نوع_جديد = st.text_input("النوع", value=case.get('نوع',''))
            with col2: محكمة_جديد = st.text_input("اسم المحكمة", value=case.get('محكمة_اسم','')); مأمورية_جديد = st.text_input("المأمورية", value=case.get('مأمورية','')); دائرة_جديد = st.text_input("الدائرة", value=case.get('دائرة',''))
            with col3: مدعي_جديد = st.text_input("المدعي", value=case.get('مدعي','')); مدعي_عليه_جديد = st.text_input("المدعي عليه", value=case.get('مدعي_عليه','')); حالة_جديد = st.selectbox("الحالة", ["متداولة", "مؤجلة", "منتهية", "شطب"], index=["متداولة", "مؤجلة", "منتهية", "شطب"].index(case.get('حالة','متداولة')) if case.get('حالة') in ["متداولة", "مؤجلة", "منتهية", "شطب"] else 0)
            موضوع_جديد = st.text_area("الموضوع", value=case.get('موضوع',''), height=100)
            if st.form_submit_button("💾 حفظ التعديلات", use_container_width=True, type="primary"):
                case['رقم']=رقم_جديد; case['سنة']=سنة_جديد; case['نوع']=نوع_جديد; case['محكمة_اسم']=محكمة_جديد; case['مأمورية']=مأمورية_جديد; case['دائرة']=دائرة_جديد; case['مدعي']=مدعي_جديد; case['مدعي_عليه']=مدعي_عليه_جديد; case['حالة']=حالة_جديد; case['موضوع']=موضوع_جديد
                save_data(data); st.success("✅ تم حفظ التعديلات"); st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>2- بيانات الخصوم</div>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1: st.markdown(f"<div style='background:#FFF3CD; padding:10px; border-radius:10px; color:#000; text-align:center'><b>المدعى:</b><br>{case.get('مدعي')}</div>", unsafe_allow_html=True)
    with col2: st.markdown(f"<div style='background:#CFF4FC; padding:10px; border-radius:10px; color:#000; text-align:center'><b>المدعى عليه:</b><br>{case.get('مدعي_عليه')}</div>", unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>3- الجلسات والإجراءات</div>", unsafe_allow_html=True)
    if case.get("جلسات"):
        for i, ج in enumerate(case["جلسات"]):
            st.markdown(f"<div style='background:#142038; padding:15px; border-radius:12px; border:2px solid #D4AF37; margin-bottom:10px; text-align:right; direction:rtl'><div style='display:flex; justify-content:flex-end; margin-bottom:10px'><div style='background:#D4AF37; color:#000; padding:5px 15px; border-radius:8px; font-weight:900; font-size:16px'>جلسة {i+1}</div></div><div style='margin-bottom:8px'><span style='color:#D4AF37; font-weight:900'>التاريخ:</span> <span style='color:#FFF'>{ج.get('تاريخ')}</span></div><div style='margin-bottom:8px'><span style='color:#D4AF37; font-weight:900'>الرول:</span> <span style='color:#FFF'>{ج.get('الرول')}</span></div><div style='margin-bottom:8px'><span style='color:#D4AF37; font-weight:900'>الاجراء:</span> <span style='color:#FFF'>{ج.get('الاجراء')}</span></div><div><span style='color:#D4AF37; font-weight:900'>ملاحظات:</span> <span style='color:#FFF'>{ج.get('ملاحظات')}</span></div></div>", unsafe_allow_html=True)
            if st.button("✏️ تعديل الجلسة", key=f"edit_session_{i}", use_container_width=True):
                st.session_state.edit_session_index = i; st.rerun()
        if 'edit_session_index' in st.session_state and st.session_state.edit_session_index is not None:
            idx = st.session_state.edit_session_index; جلسة = case["جلسات"][idx]
            with st.form("edit_session_form"):
                st.warning(f"تعديل الجلسة رقم {idx+1}")
                تاريخ_تعديل = st.date_input("التاريخ", value=datetime.strptime(جلسة.get('تاريخ'),'%Y-%m-%d'))
                رول_تعديل = st.text_input("الرول", value=جلسة.get('الرول','')); اجراء_تعديل = st.text_input("الاجراء", value=جلسة.get('الاجراء','')); ملاحظات_تعديل = st.text_area("الملاحظات", value=جلسة.get('ملاحظات',''))
                c1,c2 = st.columns(2)
                with c1:
                    if st.form_submit_button("💾 حفظ تعديل الجلسة", use_container_width=True):
                        case["جلسات"][idx] = {"تاريخ":str(تاريخ_تعديل),"الرول":رول_تعديل,"الاجراء":اجراء_تعديل,"ملاحظات":ملاحظات_تعديل}
                        جلسات_مرتبة = sorted(case["جلسات"], key=lambda x: x.get("تاريخ","9999-12-31"), reverse=True)
                        case["تاريخ_جلسة"] = جلسات_مرتبة[0].get("تاريخ",""); case["الاجراء"] = جلسات_مرتبة[0].get("الاجراء","")
                        save_data(data); st.session_state.edit_session_index = None; st.success("تم التعديل"); st.rerun()
                with c2:
                    if st.form_submit_button("❌ الغاء", use_container_width=True): st.session_state.edit_session_index = None; st.rerun()
    else: st.info("لا توجد جلسات مسجلة")
    st.markdown("<style>div[data-testid='stExpander'] summary p{color:white!important; font-weight:900!important;}</style>", unsafe_allow_html=True)
    with st.expander("اضافة جلسة جديدة"):
        with st.form("add_session"):
            تاريخ_جديد = st.date_input("تاريخ الجلسة", value=datetime.now()); رول_جديد = st.text_input("الرول"); الاجراء_جديد = st.text_input("الاجراء"); ملاحظات_جديدة = st.text_area("ملاحظات")
            if st.form_submit_button("حفظ الجلسة"):
                case["جلسات"].append({"تاريخ":str(تاريخ_جديد),"الرول":رول_جديد,"الاجراء":الاجراء_جديد,"ملاحظات":ملاحظات_جديدة})
                case["تاريخ_جلسة"] = str(تاريخ_جديد); case["الاجراء"] = الاجراء_جديد; save_data(data); st.success("تم اضافة الجلسة"); st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>4- المستندات</div>", unsafe_allow_html=True)

    ANWA3_MOSTANDAT = [
        "صحيفة دعوى", "صحيفة استئناف", "صحيفة طعن", "مذكرة دفاع",
        "حافظة مستندات", "تقرير خبير", "تقرير طب شرعى", "تقرير لجنة طبية",
        "صحيفة تجديد من الشطب", "صحيفة تعجيل من الوقف", "صورة حكم تمهيدى", "أخرى"
    ]

    نوع_المستند = st.selectbox("نوع المستند", ANWA3_MOSTANDAT, key="select_doc_type")

    اسم_نهائي = نوع_المستند
    if نوع_المستند == "أخرى":
        اسم_نهائي = st.text_input("✍️ اكتب اسم المستند", placeholder="مثال: طلب / انذار / الخ")

    with st.form("upload_form"):
        uploaded_file = st.file_uploader("اختر الملف", type=['pdf', 'jpg', 'jpeg', 'png', 'doc', 'docx'])
        if st.form_submit_button("رفع المستند"):
            if uploaded_file and اسم_نهائي and اسم_نهائي.strip()!= "":
                file_name = f"{اسم_نهائي}_{uploaded_file.name}"
                file_base64 = base64.b64encode(uploaded_file.getvalue()).decode('utf-8')
                case['مستندات'].append({"نوع": file_name, "محتوى": file_base64})
                save_data(data); st.success("✅ تم رفع المستند"); st.rerun()
            else:
                st.error("❌ لازم تختار ملف وتكتب اسم المستند")
    st.markdown("</div>", unsafe_allow_html=True)

    if case.get('مستندات'):
        st.markdown("<div style='background:#142038; padding:15px; border-radius:12px; margin-top:10px'>", unsafe_allow_html=True)
        st.markdown("<div style='color:#D4AF37; font-weight:900; margin-bottom:10px'>المستندات المرفوعة:</div>", unsafe_allow_html=True)
        مستندات_جديدة = []
        for i, مستند in enumerate(case['مستندات']):
            if isinstance(مستند, str): مستند = {"نوع": مستند, "محتوى": ""}
            اسم_المستند = مستند.get('نوع', f'ملف رقم {i+1}')
            محتوى_المستند = مستند.get('محتوى', '')
            مستندات_جديدة.append({"نوع": اسم_المستند, "محتوى": محتوى_المستند})
            col1, col2, col3 = st.columns([4,1,1])
            with col1: st.write(f"📄 {اسم_المستند}")
            with col2:
                if محتوى_المستند:
                    try: file_data = base64.b64decode(محتوى_المستند); st.download_button("📥 تحميل", data=file_data, file_name=اسم_المستند, mime="application/octet-stream", key=f"dl_{i}", use_container_width=True)
                    except: st.write("❌")
            with col3:
                if st.button("🗑️ حذف", key=f"del_{i}", use_container_width=True): case['مستندات'].pop(i); save_data(data); st.rerun()
        case['مستندات'] = مستندات_جديدة; save_data(data)
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #FF5252; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FF5252; font-size:20px; font-weight:900; text-align:center; margin-bottom:10px'>5- جلسة الحكم</div>", unsafe_allow_html=True)
    if case.get('حالة') == 'منتهية':
        لون = "#4CAF50" if case.get('مسندة_ل_الحكم') == "الصالح" else "#FF5252"
        st.markdown(f"<div style='background:#142038; padding:15px; border-radius:12px; border:2px solid {لون}; margin-bottom:10px'><b style='color:{لون}'>تاريخ جلسة الحكم:</b> {case.get('تاريخ_الحكم')}<br><b style='color:{لون}'>مسندة لـ:</b> {case.get('مسندة_ل_الحكم')}<br><b style='color:{لون}'>منطوق الحكم:</b> {case.get('منطوق_الحكم')}</div>", unsafe_allow_html=True)
        st.success("✅ حفظت ونقلت للارشيف للمتابعه يتم الانتقال للارشيف")
        with st.expander("✏️ تعديل بيانات الحكم"):
            with st.form("edit_judgment_form"):
                تاريخ_حكم_تعديل = st.date_input("تاريخ الحكم", value=datetime.strptime(case.get('تاريخ_الحكم'),'%Y-%m-%d'))
                منطوق_الحكم_تعديل = st.text_area("منطوق الحكم", value=case.get('منطوق_الحكم',''), height=150)
                مسندة_ل_تعديل = st.selectbox("مسندة لـ", ["الصالح", "الضد"], index=["الصالح", "الضد"].index(case.get('مسندة_ل_الحكم','الصالح')))
                if st.form_submit_button("💾 حفظ تعديل الحكم", use_container_width=True, type="primary"):
                    case['تاريخ_الحكم'] = str(تاريخ_حكم_تعديل); case['منطوق_الحكم'] = منطوق_الحكم_تعديل; case['مسندة_ل_الحكم'] = مسندة_ل_تعديل
                    for ج in reversed(case['جلسات']):
                        if 'الحكم' in ج.get('الاجراء',''): ج['تاريخ'] = str(تاريخ_حكم_تعديل); ج['الاجراء'] = f'الحكم - مسندة لـ {مسندة_ل_تعديل}'; ج['ملاحظات'] = منطوق_الحكم_تعديل; break
                    case['تاريخ_جلسة'] = str(تاريخ_حكم_تعديل); case['الاجراء'] = f'الحكم - مسندة لـ {مسندة_ل_تعديل}'
                    save_data(data); st.success("✅ تم تعديل الحكم"); st.rerun()
    else:
        with st.form("judgment_form"):
            st.markdown("<div style='background:#142038; padding:10px; border-radius:10px; margin-bottom:10px'>", unsafe_allow_html=True)
            st.markdown("<label style='color:#FFD700; font-weight:900; font-size:16px'>1- تاريخ الجلسة</label>", unsafe_allow_html=True); تاريخ_حكم = st.date_input("تاريخ الجلسة", value=datetime.now().date(), label_visibility="collapsed"); st.markdown("</div>", unsafe_allow_html=True)
            st.markdown("<div style='background:#142038; padding:10px; border-radius:10px; margin-bottom:10px'>", unsafe_allow_html=True)
            st.markdown("<label style='color:#FFD700; font-weight:900; font-size:16px'>2- منطوق الحكم</label>", unsafe_allow_html=True); منطوق_الحكم = st.text_area("منطوق الحكم", height=150, placeholder="اكتب منطوق الحكم هنا...", label_visibility="collapsed"); st.markdown("</div>", unsafe_allow_html=True)
            st.markdown("<div style='background:#142038; padding:10px; border-radius:10px; margin-bottom:10px'>", unsafe_allow_html=True)
            st.markdown("<label style='color:#FFD700; font-weight:900; font-size:16px'>3- مسندة لـ</label>", unsafe_allow_html=True); مسندة_ل = st.selectbox("مسندة لـ", ["الصالح", "الضد"], label_visibility="collapsed"); st.markdown("</div>", unsafe_allow_html=True)
            if st.form_submit_button("💾 حفظ الحكم", use_container_width=True, type="primary"):
                if not منطوق_الحكم: st.error("❌ لازم تكتب منطوق الحكم")
                else: case['حالة'] = 'منتهية'; case['تاريخ_الحكم'] = str(تاريخ_حكم); case['منطوق_الحكم'] = منطوق_الحكم; case['مسندة_ل_الحكم'] = مسندة_ل
                case['جلسات'].append({'تاريخ':str(تاريخ_حكم),'الرول':'-','الاجراء':f'الحكم - مسندة لـ {مسندة_ل}','ملاحظات':منطوق_الحكم}); case['تاريخ_جلسة'] = str(تاريخ_حكم); case['الاجراء'] = f'الحكم - مسندة لـ {مسندة_ل}'; save_data(data); st.success(f"✅ حفظت ونقلت للارشيف للمتابعه يتم الانتقال للارشيف"); st.session_state.page = "الأرشيف"; st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px; text-align:center'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#D4AF37; font-size:20px; font-weight:900; margin-bottom:10px'>🖨️ الطباعة والتقرير</div>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🖨️ معاينة للطباعة", use_container_width=True, type="primary"):
            html_report = print_case_report(case)
            st.components.v1.html(html_report, height=800, scrolling=True)
            st.success("✅ اضغط Ctrl+P للطباعة")
    with col2:
        html_report = print_case_report(case)
        st.download_button(label="📥 تحميل التقرير",data=html_report.encode('utf-8'),file_name=f"تقرير_قضية_{case.get('رقم')}_{case.get('سنة')}.html",mime="text/html",use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)
    st.markdown("<div style='background:#2A0A0A; padding:20px; border-radius:15px; border:3px solid #FF0000; margin-bottom:15px; text-align:center'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FF0000; font-size:22px; font-weight:900; margin-bottom:10px'>⚠️ منطقة خطر</div>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FFF; font-size:14px; margin-bottom:15px'>تحذير: حذف القضية نهائي ولا يمكن التراجع عنه</div>", unsafe_allow_html=True)
    
    if st.button("🗑️ حذف القضية نهائي", use_container_width=True, type="secondary"):
        st.session_state.confirm_delete = True
        st.rerun()

    if st.session_state.get('confirm_delete', False):
        st.warning("⚠️ هل انت متأكد 100% انك عايز تحذف القضية دي؟")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("نعم احذفها", use_container_width=True, type="primary"):
                data["cases"] = [c for c in data["cases"] if c["id"]!= case["id"]]
                save_data(data); st.session_state.confirm_delete = False
                st.success("✅ تم حذف القضية بنجاح"); st.session_state.page = "الحصر"; st.rerun()
        with col2:
            if st.button("الغاء", use_container_width=True):
                st.session_state.confirm_delete = False; st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)
# ======================================
# ==========================================
# ==============================================
# ============ الجزء الخامس: الأرشيف ============
# ==============================================
elif st.session_state.page == "الأرشيف":
    data = load_data()

    st.markdown("""
    <style>
        label { color: #FFD700 !important; font-weight: 900 !important; font-size: 15px !important; }
        input::placeholder, textarea::placeholder {
            color: #FFD700 !important;  /* اصفر دهبي */
            opacity: 1 !important;
            font-weight: 600;
        }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    st.markdown("<h2 style='color:#D4AF37; text-align:center'>📁 الأرشيف</h2>", unsafe_allow_html=True)
    
    if st.button("⬅️ العودة للرئيسية", use_container_width=True): 
        st.session_state.page = "الرئيسية"; 
        st.rerun()

    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FFF; font-size:18px; font-weight:900; text-align:center; margin-bottom:10px'>🔍 البحث عن قضية صدر فيها الحكم</div>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([3,3,1])
    with col1: بحث_مدعي = st.text_input("بحث بالاسم", placeholder="اكتب اي اسم")
    with col2: بحث_رقم = st.text_input("بحث برقم وسنة", placeholder="مثال: 123 لسنة 2024")
    with col3: st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True); بحث_زر = st.button("🔍 بحث", use_container_width=True, type="primary")
    st.markdown("</div>", unsafe_allow_html=True)

    قضايا_منتهية = [c for c in get_my_cases(data) if c.get("حالة") == "منتهية"]
    
    if بحث_زر:
        if بحث_مدعي: 
            بحث_مدعي = بحث_مدعي.lower()
            قضايا_منتهية = [c for c in قضايا_منتهية if any(
                بحث_مدعي in str(قيمة).lower() 
                for قيمة in c.values() 
                if isinstance(قيمة, str)
            )]
        if بحث_رقم: 
            قضايا_منتهية = [c for c in قضايا_منتهية if بحث_رقم in f"{c.get('رقم')} لسنة {c.get('سنة')}"]

    قضايا_جاري = [c for c in قضايا_منتهية if not c.get("تم_الحفظ_النهائي")]
    قضايا_محفوظة = [c for c in قضايا_منتهية if c.get("تم_الحفظ_النهائي")]

    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #FFD700; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FFD700; font-size:20px; font-weight:900; text-align:center; margin-bottom:15px'>1- احكام صادرة وجاري اتخاذ الاجراء اللازم بشأنها</div>", unsafe_allow_html=True)
    
    if قضايا_جاري:
        for case in قضايا_جاري:
            لون = "#4CAF50" if case.get('مسندة_ل_الحكم') == "الصالح" else "#FF5252"
            
            نوع = case.get('نوع', '').lower()
            if 'استئناف' in نوع:
                طرف1_عنوان = "المستأنف"
                طرف2_عنوان = "المستأنف ضده"
            elif 'طعن' in نوع:
                طرف1_عنوان = "الطاعن"
                طرف2_عنوان = "المطعون ضده"
            else:
                طرف1_عنوان = "المدعي"
                طرف2_عنوان = "المدعي عليه"

            st.markdown(f"<div style='background:#142038; padding:15px; border-radius:12px; border:2px solid {لون}; margin-bottom:10px'>", unsafe_allow_html=True)
            
            st.markdown(f"""
            <table style='width:100%; border-collapse:collapse; margin-bottom:10px;'>
                <tr><th colspan='2' style='background:{لون}; color:#FFF; padding:10px; text-align:center; font-size:16px; border-radius:8px 8px 0 0;'>رقم {case.get('رقم')} لسنة {case.get('سنة')} - {case.get('نوع')}</th></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; width:35%; font-weight:900;'>المحكمة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('محكمة_اسم')} {f'- مأمورية {case.get("مأمورية")}' if case.get('مأمورية') else ''}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الدائرة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('دائرة')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الموضوع</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('موضوع')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{طرف1_عنوان}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{طرف2_عنوان}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي_عليه')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>تاريخ الحكم</td><td style='background:#FFF; color:{لون}; padding:8px; font-weight:900;'>{case.get('تاريخ_الحكم')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>المنطوق</td><td style='background:#FFF; color:{لون}; padding:8px; font-weight:900;'>{case.get('منطوق_الحكم')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900; border-radius:0 0 0 8px;'>مسندة لـ</td><td style='background:#FFF; color:{لون}; padding:8px; font-weight:900; border-radius:0 0 8px 0;'>{case.get('مسندة_ل_الحكم')}</td></tr>
            </table>
            """, unsafe_allow_html=True)
            
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("📄 فتح", key=f"open_{case['id']}", use_container_width=True):
                    st.session_state.selected_case_id = case["id"]; st.session_state.page = "تفاصيل"; st.rerun()
            with col2:
                if st.button("💾 حفظ نهائي", key=f"save_{case['id']}", use_container_width=True):
                    st.session_state.save_case_id = case["id"]; st.rerun()
            with col3:
                if st.button("🗑️ حذف", key=f"del_arch_{case['id']}", use_container_width=True):
                    st.session_state.del_arch_id = case["id"]; st.rerun()
            
            if st.session_state.get('save_case_id') == case['id']:
                with st.form(f"save_form_{case['id']}"):
                    st.warning("حفظ القضية نهائي")
                    سبب_الحفظ = st.text_area("سبب الحفظ", placeholder="مثال: تم الطعن / حكم نهائي / عدم جدوى")
                    مستندات_الحفظ = st.file_uploader("ارفع مستندات الحفظ", type=['pdf','jpg','png','doc','docx'], accept_multiple_files=True)
                    
                    if st.form_submit_button("💾 تأكيد الحفظ النهائي", use_container_width=True, type="primary"):
                        case['سبب_الحفظ'] = سبب_الحفظ
                        case['مستندات_الحفظ'] = []
                        for f in مستندات_الحفظ:
                            file_base64 = base64.b64encode(f.getvalue()).decode('utf-8')
                            case['مستندات_الحفظ'].append({"نوع": f.name, "محتوى": file_base64})
                        case['تم_الحفظ_النهائي'] = True
                        case['تاريخ_الحفظ'] = str(datetime.now().date())
                        save_data(data); st.session_state.save_case_id = None
                        st.success("✅ تم حفظ القضية نهائي"); st.rerun()

            if st.session_state.get('del_arch_id') == case['id']:
                st.error("⚠️ هل انت متأكد 100% من حذف القضية نهائي من الارشيف؟")
                c1,c2 = st.columns(2)
                with c1:
                    if st.button("نعم احذف", key=f"confirm_del_{case['id']}"):
                        data["cases"] = [c for c in data["cases"] if c["id"]!= case["id"]]
                        save_data(data); st.session_state.del_arch_id = None; st.success("تم الحذف"); st.rerun()
                with c2:
                    if st.button("الغاء", key=f"cancel_del_{case['id']}"):
                        st.session_state.del_arch_id = None; st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)
    else: st.info("لا توجد احكام")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #4CAF50; margin-bottom:15px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#4CAF50; font-size:20px; font-weight:900; text-align:center; margin-bottom:15px'>2- احكام صادرة وتم اتخاذ الاجراء اللازم بشأنها وحفظت</div>", unsafe_allow_html=True)

    if قضايا_محفوظة:
        for case in قضايا_محفوظة:
            نوع = case.get('نوع', '').lower()
            if 'استئناف' in نوع:
                طرف1_عنوان = "المستأنف"
                طرف2_عنوان = "المستأنف ضده"
            elif 'طعن' in نوع:
                طرف1_عنوان = "الطاعن"
                طرف2_عنوان = "المطعون ضده"
            else:
                طرف1_عنوان = "المدعي"
                طرف2_عنوان = "المدعي عليه"

            st.markdown(f"<div style='background:#142038; padding:15px; border-radius:12px; border:2px solid #4CAF50; margin-bottom:10px'>", unsafe_allow_html=True)
            st.markdown(f"""
            <table style='width:100%; border-collapse:collapse; margin-bottom:10px;'>
                <tr><th colspan='2' style='background:#4CAF50; color:#FFF; padding:10px; text-align:center; font-size:16px; border-radius:8px 8px 0 0;'>رقم {case.get('رقم')} لسنة {case.get('سنة')} - {case.get('نوع')}</th></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; width:35%; font-weight:900;'>المحكمة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('محكمة_اسم')} {f'- مأمورية {case.get("مأمورية")}' if case.get('مأمورية') else ''}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الدائرة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('دائرة')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الموضوع</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('موضوع')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{طرف1_عنوان}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{طرف2_عنوان}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي_عليه')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>تاريخ الحفظ</td><td style='background:#FFF; color:#4CAF50; padding:8px; font-weight:900;'>{case.get('تاريخ_الحفظ')}</td></tr>
                <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900; border-radius:0 0 0 8px;'>سبب الحفظ</td><td style='background:#FFF; color:#FFD700; padding:8px; font-weight:900; border-radius:0 0 8px 0;'>{case.get('سبب_الحفظ')}</td></tr>
            </table>
            """, unsafe_allow_html=True)
            
            if case.get('مستندات_الحفظ'):
                st.markdown("<div style='color:#D4AF37; margin-top:10px'>مستندات الحفظ:</div>", unsafe_allow_html=True)
                for i, مستند in enumerate(case['مستندات_الحفظ']):
                    file_data = base64.b64decode(مستند['محتوى'])
                    st.download_button(f"📥 {مستند['نوع']}", data=file_data, file_name=مستند['نوع'], key=f"dl_save_{case['id']}_{i}")

            col1, col2 = st.columns(2)
            with col1:
                if st.button("📄 فتح", key=f"open_saved_{case['id']}", use_container_width=True):
                    st.session_state.selected_case_id = case["id"]; st.session_state.page = "تفاصيل"; st.rerun()
            with col2:
                if st.button("🗑️ حذف نهائي", key=f"del_saved_{case['id']}", use_container_width=True):
                    st.session_state.del_saved_id = case["id"]; st.rerun()

            if st.session_state.get('del_saved_id') == case['id']:
                st.error("⚠️ هل انت متأكد 100% من حذف القضية نهائي؟")
                c1,c2 = st.columns(2)
                with c1:
                    if st.button("نعم احذف", key=f"confirm_del_saved_{case['id']}"):
                        data["cases"] = [c for c in data["cases"] if c["id"]!= case["id"]]
                        save_data(data); st.session_state.del_saved_id = None; st.success("تم الحذف"); st.rerun()
                with c2:
                    if st.button("الغاء", key=f"cancel_del_saved_{case['id']}"):
                        st.session_state.del_saved_id = None; st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)
    else: st.info("لا توجد قضايا محفوظة نهائي")
    st.markdown("</div>", unsafe_allow_html=True)
# =========================================
# ======================================
# ================================================
# ============ الجزء السادس: البحث ============
# ================================================
elif st.session_state.page == "بحث":
    import base64
    data = load_data()
    
    st.markdown("""
    <style>
        label { color: #FFD700 !important; font-weight: 900 !important; font-size: 15px !important; }
        input::placeholder { color: #FFD700 !important; opacity: 1 !important; font-weight: 600; }
        .case-table {width:100%; border-collapse: collapse; font-size:12px; color:#FFF; text-align:center; margin-bottom:10px;}
        .case-table th {background:#D4AF37; color:#0B1426; padding:8px; font-weight:900;}
        .case-table td {background:#1E2A47; padding:8px; border:1px solid #D4AF37; vertical-align:top; color:#FFF;}
        .plaintiff {background:#FFF3CD; color:#000; font-weight:700; border-radius:6px; padding:6px; font-size:12px;}
        .plaintiff-hey2a {background:#DC3545!important; color:#FFF!important; font-weight:900; border-radius:6px; padding:6px; font-size:12px;}
        .defendant {background:#CFF4FC; color:#000; font-weight:700; border-radius:6px; padding:6px; font-size:12px;}
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
    st.markdown("<h2 style='color:#4DA8DA; text-align:center'>🔍 البحث عن دعوى</h2>", unsafe_allow_html=True)

    if st.button("⬅️ العودة للرئيسية", use_container_width=True, key="back_from_search"):
        st.session_state.page = "الرئيسية"
        st.rerun()

    st.markdown("<div style='background:#1E2A47; padding:20px; border-radius:15px; border:2px solid #4DA8DA; margin-bottom:20px'>", unsafe_allow_html=True)
    st.markdown("<div style='color:#FFF; font-size:18px; font-weight:900; text-align:center; margin-bottom:15px'>ابحث بالاسم او برقم وسنة الدعوى</div>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1: 
        بحث_اسم = st.text_input("بحث بالاسم", placeholder="اكتب اسم المدعي او المدعى عليه")
    with col2: 
        بحث_رقم = st.text_input("بحث برقم وسنة", placeholder="مثال: 123 لسنة 2024")
    
    بحث_زر = st.button("🔍 بحث", use_container_width=True, type="primary")
    st.markdown("</div>", unsafe_allow_html=True)

    if بحث_زر:
        if not بحث_اسم.strip() and not بحث_رقم.strip():
            st.error("اكتب اسم او رقم للبحث")
        else:
            results = []
            بحث_اسم = بحث_اسم.lower()
            
            for case in get_my_cases(data):
                match = False
                if بحث_اسم:
                    if any(بحث_اسم in str(قيمة).lower() for قيمة in case.values() if isinstance(قيمة, str)):
                        match = True
                if بحث_رقم:
                    رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
                    if بحث_رقم in رقم_كامل:
                        match = True
                if match:
                    results.append(case)

            if not results:
                st.warning("⚠️ لا يوجد بحث مطابق")
            else:
                st.success(f"✅ تم العثور على {len(results)} نتيجة")
                for idx, case in enumerate(results, 1):
                    
                    نوع = case.get('نوع', '').lower()
                    if 'استئناف' in نوع:
                        لقب1, لقب2 = "المستأنف:", "المستأنف ضده:"
                    elif 'طعن' in نوع:
                        لقب1, لقب2 = "الطاعن:", "المطعون ضده:"
                    else:
                        لقب1, لقب2 = "المدعي:", "المدعي عليه:"

                    if "الهيئة" in str(case.get('مدعي','')):
                        طرف1_html = f"<div class='plaintiff-hey2a'><b>{لقب1}</b><br>{case.get('مدعي','')}</div>"
                    else:
                        طرف1_html = f"<div class='plaintiff'><b>{لقب1}</b><br>{case.get('مدعي','')}</div>"
                    طرف2_html = f"<div class='defendant'><b>{لقب2}</b><br>{case.get('مدعي_عليه','')}</div>"
                    خصوم = طرف1_html + "<div style='height:4px'></div>" + طرف2_html

                    رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
                    محكمة_كاملة = f"{case.get('نوع','')} {case.get('محكمة_اسم','')}"
                    if case.get('مأمورية',''): محكمة_كاملة += f"<br>مأمورية {case.get('مأمورية','')}"
                    if case.get('دائرة',''): محكمة_كاملة += f"<br>دائرة {case.get('دائرة','')}"

                    if case.get('حالة') == 'منتهية':
                        حالة_لون = "#FF5252"
                        مكان = "📁 الأرشيف"
                    else:
                        حالة_لون = "#4CAF50"
                        مكان = "📋 الحصر العام"

                    st.markdown("<div style='background:#1E2A47; padding:15px; border-radius:15px; border:2px solid #D4AF37; margin-bottom:15px'>", unsafe_allow_html=True)
                    st.markdown(f"""
                    <table class='case-table'>
                    <tr><th>م</th><th>الرقم والسنة</th><th>المحكمة والدائرة</th><th>الخصوم</th><th>الموضوع</th><th>اخر جلسة</th><th>الحالة</th><th>المكان</th></tr>
                    <tr>
                        <td>{idx}</td>
                        <td>{رقم_كامل}</td>
                        <td>{محكمة_كاملة}</td>
                        <td>{خصوم}</td>
                        <td>{case.get('موضوع','')}</td>
                        <td style='color:#FFD700; font-weight:900'>{case.get('تاريخ_جلسة','-')}</td>
                        <td style='color:{حالة_لون}; font-weight:900'>{case.get('حالة','متداولة')}</td>
                        <td style='color:#4DA8DA; font-weight:900'>{مكان}</td>
                    </tr>
                    </table>
                    """, unsafe_allow_html=True)

                    st.markdown("<div style='background:#142038; padding:20px; border-radius:12px; border:2px solid #4DA8DA; margin-top:10px'>", unsafe_allow_html=True)
                    st.markdown(f"<div style='color:#4DA8DA; font-size:20px; font-weight:900; text-align:center; margin-bottom:15px'>📄 تفاصيل كاملة - {رقم_كامل}</div>", unsafe_allow_html=True)
                    
                    st.markdown("<div style='color:#FFD700; font-size:16px; font-weight:900; margin-bottom:10px'>1- البيانات الاساسية</div>", unsafe_allow_html=True)
                    st.markdown(f"""
                    <table style='width:100%; border-collapse:collapse; margin-bottom:15px;'>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; width:30%; font-weight:900;'>نوع الدعوى</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('نوع')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>المحكمة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('محكمة_اسم')} {f'- مأمورية {case.get("مأمورية")}' if case.get('مأمورية') else ''}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الدائرة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('دائرة')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الموضوع</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('موضوع')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{لقب1}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>{لقب2}</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('مدعي_عليه')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الحالة الحالية</td><td style='background:#FFF; color:{حالة_لون}; padding:8px; font-weight:900;'>{case.get('حالة','متداولة')} - {مكان}</td></tr>
                    </table>
                    """, unsafe_allow_html=True)

                    st.markdown("<div style='color:#FFD700; font-size:16px; font-weight:900; margin-bottom:10px'>2- اخر اجراء</div>", unsafe_allow_html=True)
                    st.markdown(f"""
                    <table style='width:100%; border-collapse:collapse; margin-bottom:15px;'>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; width:30%; font-weight:900;'>تاريخ اخر جلسة</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('تاريخ_جلسة','-')}</td></tr>
                        <tr><td style='background:#1E2A47; color:#FFD700; padding:8px; font-weight:900;'>الاجراء</td><td style='background:#FFF; color:#000; padding:8px; font-weight:700;'>{case.get('الاجراء','-')}</td></tr>
                    </table>
                    """, unsafe_allow_html=True)

                    if case.get('جلسات'):
                        st.markdown("<div style='color:#FFD700; font-size:16px; font-weight:900; margin-bottom:10px'>3- سجل الجلسات</div>", unsafe_allow_html=True)
                        جلسات_مرتبة = sorted(case['جلسات'], key=lambda x: x.get("تاريخ",""), reverse=True)
                        for ج in جلسات_مرتبة:
                            st.markdown(f"<div style='background:#1E2A47; padding:10px; border-radius:8px; margin-bottom:5px; border:1px solid #D4AF37'><b style='color:#FFD700'>تاريخ:</b> <span style='color:#FFF'>{ج.get('تاريخ')}</span> | <b style='color:#FFD700'>الاجراء:</b> <span style='color:#FFF'>{ج.get('الاجراء')}</span> | <b style='color:#FFD700'>الحالة:</b> <span style='color:#FFF'>{ج.get('الحالة')}</span></div>", unsafe_allow_html=True)
                    else:
                        st.info("لا يوجد سجل جلسات مسجل")

                    if case.get('مستندات'):
                        st.markdown("<div style='color:#FFD700; font-size:16px; font-weight:900; margin:15px 0 10px 0'>4- المستندات المرفقة</div>", unsafe_allow_html=True)
                        for i, مستند in enumerate(case['مستندات']):
                            file_data = base64.b64decode(مستند['محتوى'])
                            st.download_button(f"📥 تحميل {مستند['نوع']}", data=file_data, file_name=مستند['نوع'], key=f"dl_search_{case['id']}_{i}", use_container_width=True)
                    else:
                        st.info("لا يوجد مستندات مرفقة")
                    
                    st.markdown("</div>", unsafe_allow_html=True)
                    st.markdown("</div>", unsafe_allow_html=True)
                    # ======
# ================================================
# ============ مركز التنبيهات ====================
# ================================================
elif st.session_state.page == "التنبيهات":
    st.markdown("<h1 style='text-align:center; color:#C9A961;'>مركز التنبيهات</h1>", unsafe_allow_html=True)
    
    # === زر العودة للرئيسية ===
    if st.button("⬅️ العودة للرئيسية", use_container_width=True):
        st.session_state.page = "الرئيسية"
        st.rerun()
    # ===========================

    st.divider() # فاصل فوق
    
    with st.container(border=True):
        st.markdown("<h2 style='text-align:center; color:#C9A961;'>ارسال التنبيهات بالايميل</h2>", unsafe_allow_html=True)
        
        user_email = st.text_input("سجل ايميلك عشان يجيلك التنبيهات", key="alert_email_input")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("1. حفظ الايميل", use_container_width=True):
                if user_email and "@" in user_email:
                    st.session_state['saved_email'] = user_email
                    st.success(f"✅ تم حفظ الايميل: {user_email}")
                else:
                    st.warning("دخل ايميل صحيح")

        with col2:
            if st.button("2. 📧 ارسل التنبيهات دلوقتي", use_container_width=True):
                if 'saved_email' in st.session_state:
                    alerts = get_alert_cases()
                    body = "<div style='direction:rtl; text-align:right; font-family:Arial;'>"
                    body += "<h2 style='color:#C9A961; text-align:center;'>تنبيهات القضايا</h2>"
                    
                    body += "<h3 style='color:#FFD700;'>1. جلسات خلال 7 ايام</h3>"
                    if alerts["sessions"]:
                        for case in alerts["sessions"]:
                            رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
                            body += f"<p style='border:1px solid #C9A961; padding:10px; border-radius:8px;'>"
                            body += f"<b>رقم:</b> {رقم_كامل} <br> "
                            body += f"<b>الموضوع:</b> {case.get('موضوع','')} <br> "
                            body += f"<b>الجلسة:</b> {case.get('تاريخ_جلسة','')} - <b style='color:red;'>فاضل {case.get('days_left',0)} يوم</b>"
                            body += f"</p>"
                    else:
                        body += "<p>✅ لا توجد جلسات قريبة</p>"

                    body += "<h3 style='color:#FF4500;'>2. طعون خلال 15 يوم</h3>"
                    if alerts["appeals"]:
                        for case in alerts["appeals"]:
                            رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
                            body += f"<p style='border:1px solid #FF4500; padding:10px; border-radius:8px;'>"
                            body += f"<b>رقم:</b> {رقم_كامل} <br> "
                            body += f"<b>الموضوع:</b> {case.get('موضوع','')} <br> "
                            body += f"<b style='color:red;'>اخر ميعاد للطعن: {case.get('deadline','')} - فاضل {case.get('days_left_appeal',0)} يوم</b>"
                            body += f"</p>"
                    else:
                        body += "<p>✅ لا توجد طعون قريبة</p>"
                    body += "</div>"

                    if send_email(st.session_state['saved_email'], "تنبيهات القضايا من النظام", body):
                        st.success("✅ تم ارسال التنبيهات بنجاح للايميل")
                else:
                    st.error("❌ سجل الايميل الاول من الزرار اللي جنبه")

    st.divider()
    alerts = get_alert_cases()
    st.markdown(f"<h3 style='text-align:center; color:#C9A961;'>التنبيهات الموجوده حاليا</h3>", unsafe_allow_html=True)

    st.markdown("<h2 style='text-align:center; color:#C9A961;'>الجلسات خلال 7 ايام</h2>", unsafe_allow_html=True)
    if alerts["sessions"]:
        for case in alerts["sessions"]:
            رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
            with st.container(border=True):
                st.write(f"**رقم القضية:** {رقم_كامل}")
                st.write(f"**الموضوع:** {case.get('موضوع','')}")
                st.write(f"**تاريخ الجلسة:** {case.get('تاريخ_جلسة','')}")
                st.write(f"**فاضل:** {case.get('days_left',0)} يوم")
    else:
        st.info("لا توجد جلسات خلال 7 ايام")

    st.markdown("<h2 style='text-align:center; color:#C9A961;'>الطعون خلال 15 يوم</h2>", unsafe_allow_html=True)
    if alerts["appeals"]:
        for case in alerts["appeals"]:
            رقم_كامل = f"{case.get('رقم','')} لسنة {case.get('سنة','')}"
            with st.container(border=True):
                st.write(f"**رقم القضية:** {رقم_كامل}")
                st.write(f"**الموضوع:** {case.get('موضوع','')}")
                st.write(f"**اخر ميعاد للطعن:** {case.get('deadline','')}")
                st.write(f"**فاضل:** {case.get('days_left_appeal',0)} يوم")
    else:
        st.info("لا توجد طعون خلال 15 يوم")
        
    st.divider() # فاصل تحت
# ================================================
# ============ صفحة المكتبة القانونية ============
        # ================================================
# ============ صفحة المكتبة القانونية ============
# ==================================================
elif st.session_state.page == "مكتبة":
    data = load_data()
    st.markdown('<h1 style="text-align: center; color: #FFD700;">المكتبة 📚<br>القانونية</h1>', unsafe_allow_html=True)
    
    if st.button("⬅️ العودة للصفحة الرئيسية", use_container_width=True):
        st.session_state.page = "الرئيسية"
        for k in ["selected_section", "search_filters", "show_upload"]:
            st.session_state.pop(k, None)
        st.rerun()

    library_data = data.get("library", [])
    my_library = [f for f in library_data if f.get("user_id") == st.session_state.user["id"]]

    st.markdown("### 🔍 بحث سريع")
    col1, col2 = st.columns([4,1])
    with col1: 
        search_query = st.text_input("اكتب اسم او رقم او سنة", key="search_q_lib", placeholder="مثال: 148 او تأمينات او 2019")
    with col2: 
        st.write("")
        st.write("")
        if st.button("🔍 بحث", use_container_width=True, key="btn_search_lib"):
            if search_query.strip():
                st.session_state.search_filters = {"q": search_query.strip()}
                st.session_state.pop("selected_section", None)
                st.rerun()
            else:
                st.warning("اكتب كلمة للبحث")

    st.divider()

    st.markdown("### 📁 الاقسام")
    cols = st.columns(4)
    for i, (section, color) in enumerate(LIBRARY_SECTIONS.items()):
        with cols[i % 4]:
            count = len([f for f in my_library if str(f.get("section")).strip() == str(section).strip()])
            
            if st.button(f"{section}", key=f"btn_section_{i}", use_container_width=True, help=f"عدد الملفات: {count}"):
                st.session_state.selected_section = section
                st.session_state.pop("search_filters", None)
                st.rerun()
            
            st.markdown(f"<p style='text-align:center; color:{color}; font-size:12px; margin-top:-10px;'>({count} ملف)</p>", unsafe_allow_html=True)

    st.divider()
    
    files_to_show = []
    title = ""
    
    if "selected_section" in st.session_state:
        sec = st.session_state.selected_section
        title = f"📂 القسم: {sec}"
        files_to_show = [f for f in my_library if str(f.get("section")).strip() == str(sec).strip()]
        
    elif "search_filters" in st.session_state:
        q = st.session_state.search_filters["q"].lower()
        title = f"🔍 نتائج البحث عن: '{q}'"
        files_to_show = [
            item for item in my_library 
            if q in item.get("name","").lower() 
            or q in item.get("number","").lower() 
            or q in item.get("year","").lower()
            or q in item.get("section","").lower()
        ]
    else:
        title = "📂 اختار قسم من فوق"
        files_to_show = []

    if title: st.subheader(title)

    if st.button("➕ اضافة مادة قانونية لمكتبتك", key="btn_add_lib", type="primary", use_container_width=True):
        st.session_state.show_upload = True

    if st.session_state.get("show_upload", False):
        with st.form("form_add_doc_lib"):
            section_select = st.selectbox("1- اختر القسم", list(LIBRARY_SECTIONS.keys()), key="sel_section")
            doc_name = st.text_input("2- اسم المستند", placeholder="مثال: قانون التأمينات 148 لسنة 2019", key="inp_name")
            col1, col2 = st.columns(2)
            with col1: doc_number = st.text_input("3- الرقم", placeholder="148", key="inp_num")
            with col2: doc_year = st.text_input("4- السنة", placeholder="2019", key="inp_year")
            uploaded_file = st.file_uploader("5- ارفع الملف", type=['pdf', 'doc', 'docx', 'jpg', 'png'], key="upl_file")
            
            if st.form_submit_button("💾 حفظ في القسم", use_container_width=True):
                if uploaded_file and doc_name.strip() and section_select:
                    file_base64 = base64.b64encode(uploaded_file.getvalue()).decode('utf-8')
                    new_doc = {
                        "id": secrets.token_hex(8),
                        "user_id": st.session_state.user["id"],
                        "name": doc_name.strip(), 
                        "section": section_select.strip(),
                        "number": doc_number.strip(),
                        "year": doc_year.strip(),
                        "file_type": uploaded_file.name.split('.')[-1],
                        "content": file_base64
                    }
                    data.setdefault("library", []).append(new_doc)
                    save_data(data)
                    st.success(f"✅ تم حفظ '{doc_name}' في قسم '{section_select}'")
                    st.session_state.show_upload = False
                    st.session_state.selected_section = section_select # يفتح القسم بعد الحفظ
                    st.rerun()
                else:
                    st.error("❌ لازم تختار قسم + اسم + ملف")

    if files_to_show:
        st.write(f"عدد الملفات: {len(files_to_show)}")
        for doc in files_to_show:
            color = LIBRARY_SECTIONS.get(doc.get("section"), "#7F8C8D")
            st.markdown(f"<div style='border-right:5px solid {color}; padding:12px; margin:8px 0; background:#1e1e1e; border-radius:8px;'>", unsafe_allow_html=True)
            st.markdown(f"**{doc.get('name')}**")
            st.caption(f"رقم: {doc.get('number','-')} | سنة: {doc.get('year','-')}")
            
            col1, col2 = st.columns([3,1])
            with col1:
                if doc.get("content"):
                    file_data = base64.b64decode(doc["content"])
                    file_name = f"{doc.get('name')}.{doc.get('file_type')}"
                    mime_type = {
                        "pdf":"application/pdf",
                        "doc":"application/msword",
                        "docx":"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "jpg":"image/jpeg",
                        "png":"image/png"
                    }.get(doc.get('file_type'),"application/octet-stream")
                    st.download_button("⬇️ تحميل", data=file_data, file_name=file_name, mime=mime_type, key=f"dl_{doc['id']}", use_container_width=True)
            
            with col2:
                if doc.get("user_id") == st.session_state.user["id"]:
                    if st.button("🗑️ حذف", key=f"del_{doc['id']}", use_container_width=True):
                        data["library"] = [d for d in data["library"] if d["id"] != doc["id"]]
                        save_data(data)
                        st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
    
    elif "selected_section" in st.session_state:
        st.warning(f"القسم '{st.session_state.selected_section}' فاضي. ارفع اول ملف من الزرار اللي فوق")
    elif "search_filters" in st.session_state:
        st.info("مفيش نتائج للبحث ده")
    else:
        st.info("اختار قسم من الازرار اللي فوق عشان تشوف الملفات")
        # ================================================
        # =================================================
        # =========== الجزء الثامن: التقارير ============
# ================= القسم 1 من 2 =================

def get_export_html(content_html, title):
    return f"""<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
    @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;700;900&display=swap');
    body {{ font-family: "Cairo", "Times New Roman", serif; direction: rtl; background: white; color: black; padding: 20px; }}
  .case-table {{width:100%; border-collapse:collapse; font-size:12px; margin-top:12px;}}
  .case-table th {{ background: #B8860B; color:#000; padding:8px; border:1px solid #8B6914; font-weight:900; text-align:center; }}
  .case-table td {{padding:6px; border:1px solid #B8860B; text-align:center; background:#fff; color:#000;}}
  .table-container {{overflow-x:auto}}
    h1, h2, h3 {{ color: #8B6914!important; text-align: center; }}
</style>
</head>
<body>
{content_html}
</body>
</html>"""

if st.session_state.page == "تقارير":
    import io
    import pandas as pd
    from datetime import datetime
    data = load_data()
    my_cases = get_my_cases(data)

    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;700;900&display=swap');
  .case-table {width:100%; border-collapse:collapse; font-size:10px; margin-top:12px; direction:rtl; font-family: "Cairo", sans-serif;}
  .case-table th {background: linear-gradient(135deg, #8B6914 0%, #B8860B 100%); color:#000; padding:5px; border:1px solid #8B6914; font-weight:900; text-align:center; white-space: nowrap; font-size:10px;}
  .case-table td {padding:4px; border:1px solid #B8860B; text-align:center; background:#1E2A47; color:#fff; font-size:10px;}
  .table-container {overflow-x:auto}
    h1, h2, h3, h4, label {color: #B8860B!important; font-family: "Cairo", sans-serif;}
  .stButton>button {border: 2px solid #B8860B!important; color: #B8860B!important; font-family: "Cairo", sans-serif; font-size:13px;}
  .stButton>button:hover {background: #B8860B!important; color: #000!important;}
    [data-testid="stTab"] button {color: #B8860B!important; font-family: "Cairo", sans-serif; font-size:13px;}
    [data-testid="stTab"] button[aria-selected="true"] {border-bottom: 3px solid #B8860B!important;}
    </style>
    """, unsafe_allow_html=True)

    st.markdown("<h2 style='color:#B8860B; text-align:center; font-family: Cairo; font-size:18px;'>📑 مركز التقارير الحكومية</h2>", unsafe_allow_html=True)
    if st.button("⬅️ العودة للرئيسية", use_container_width=True): st.session_state.page = "الرئيسية"; st.rerun()

    if 'last_report_html' not in st.session_state: st.session_state.last_report_html = ""; st.session_state.last_report_title = "تقرير"
    if 'last_report_df' not in st.session_state: st.session_state.last_report_df = pd.DataFrame()

    tab1, tab2, tab3, tab4 = st.tabs(["📊 بيان بجميع الدعاوى المتداولة", "⚖️ بيان الاحكام", "📈 الإحصائيات", "📊 بيان عددي"])

    def report_header(region, title, مدير_عام, مدير_ادارة, عضو_قانوني):
        return f"""<div style="text-align:right; color:#B8860B; border:2px double #B8860B; padding:12px 10px; background: #0A1428; border-radius:5px; margin-bottom:12px;">
        <h2 style="margin:2px 0; font-size:15px; font-weight:900;">الهيئة القومية للتأمين الاجتماعي</h2>
        <h3 style="margin:1px 0; font-size:12px; font-weight:700;">الإدارة المركزية للإدارات القانونية</h3>
        <h3 style="margin:4px 0; font-size:12px; font-weight:700;">ديوان عام منطقة {region}</h3>
        <hr style="border:1px solid #B8860B; margin:8px 0;">
        <h3 style="margin:6px 0; font-size:13px; font-weight:900; text-align:center; text-decoration: underline;"> {title} </h3>
        </div>"""
        # ========= تبويب 1: الدعاوى المتداولة =========
    with tab1:
        st.markdown("<div style='background:#1E2A47; padding:12px; border-radius:8px; border:2px solid #B8860B; margin-bottom:12px'>", unsafe_allow_html=True)
        نوع_تقرير_متداولة = st.selectbox("نوع البيان", ["بيان بجميع الدعاوى المتداولة", "بيان بالدعاوى المتداولة حسب موضوع الدعوى"], key="no1_tقرير")
        colA, colB, colC = st.columns(3)
        with colA: region = st.text_input("ديوان عام منطقة", key="region1")
        with colB: مدير_عام1 = st.text_input("اسم مدير عام الإدارات القانونية", key="modir1")
        with colC: مدير_ادارة1 = st.text_input("اسم مدير إدارة القضايا", key="modir_idara1")
        عضو_قانوني1 = st.text_input("اسم العضو القانوني", key="odo1")
        col1, col2, col3 = st.columns(3)
        with col1: from_date = st.date_input("من الفترة", key="from1")
        with col2: to_date = st.date_input("حتى الفترة", key="to1")
        with col3: lawyer = st.text_input("طرف الاستاذ/ المحامي", key="lawyer1")
        topic = ""
        if "حسب موضوع" in نوع_تقرير_متداولة: topic = st.text_input("موضوع الدعوى للفلترة", key="topic1")
        st.markdown("</div>", unsafe_allow_html=True)

        if st.button("🔍 عرض بيان الدعاوى المتداولة", use_container_width=True, type="primary", key="show1"):
            cases = [c for c in my_cases if str(c.get('حالة','')).strip() == 'متداولة']
            فلترة_بالمدة = []
            for c in cases:
                if c.get('تاريخ_جلسة'):
                    try:
                        ت_جلسة = datetime.strptime(c['تاريخ_جلسة'], '%Y-%m-%d').date()
                        if from_date <= ت_جلسة <= to_date: فلترة_بالمدة.append(c)
                    except: pass
            cases = فلترة_بالمدة
            if "حسب موضوع" in نوع_تقرير_متداولة and topic: cases = [c for c in cases if topic.lower() in str(c.get('موضوع','')).lower()]
            cases = sorted(cases, key=lambda x: x.get("تاريخ_جلسة","0000-00-00"), reverse=True)
            title = f"{نوع_تقرير_متداولة} خلال الفترة من {from_date.strftime('%d-%m-%Y')} إلى {to_date.strftime('%d-%m-%Y')} طرف الاستاذ/ {lawyer} المحامي"
            header_html = report_header(region, title, مدير_عام1, مدير_ادارة1, عضو_قانوني1)
            if not cases: st.warning(f"⚠️ لا توجد دعاوى متداولة في الفترة")
            else:
                st.success(f"✅ تم العثور على {len(cases)} دعوى متداولة")
                html = "<table class='case-table'><tr><th>م</th><th>رقم</th><th>سنة</th><th>المحكمة والدائرة</th><th>الخصوم</th><th>الموضوع</th><th>اخر جلسة</th><th>الاجراء</th><th>ملاحظات</th></tr>"
                df_data = []
                for idx, c in enumerate(cases, 1):
                    محكمة = f"{c.get('نوع','')} {c.get('محكمة_اسم','')}"
                    if c.get('مأمورية'): محكمة += f" - مأمورية {c.get('مأمورية')}"
                    if c.get('دائرة'): محكمة += f" - دائرة {c.get('دائرة')}"
                    مدعي = c.get('مدعي',''); مدعي_عليه = c.get('مدعي_عليه','')
                    خصوم_html = f"<span style='color:#dc3545; font-weight:900'>{مدعي} ضد {مدعي_عليه}</span>" if "الهيئة" in str(مدعي) else f"{مدعي} ضد {مدعي_عليه}"
                    html += f"<tr><td>{idx}</td><td>{c.get('رقم','')}</td><td>{c.get('سنة','')}</td><td>{محكمة}</td><td>{خصوم_html}</td><td>{c.get('موضوع','')}</td><td><b style='color:#B8860B'>{c.get('تاريخ_جلسة','')}</b></td><td>{c.get('الاجراء','')}</td><td>{c.get('ملاحظات','')}</td></tr>"
                    df_data.append({'م': idx, 'رقم': c.get('رقم',''), 'سنة': c.get('سنة',''), 'المحكمة': محكمة, 'الخصوم': f"{مدعي} ضد {مدعي_عليه}", 'الموضوع': c.get('موضوع',''), 'اخر جلسة': c.get('تاريخ_جلسة',''), 'الاجراء': c.get('الاجراء',''), 'ملاحظات': c.get('ملاحظات','')})
                html += "</table>"
                # ===== الفوتر المعدل: تحرر في بعد سطرين =====
                footer = f"""<div style="margin-top:25px; color:#B8860B; font-size:12px;"><p style="text-align:center; margin-bottom:20px; font-size:13px; font-weight:700;">تفضلوا بقبول وافر الاحترام والتقدير،</p><table style="width:100%;"><tr><td style="width:50%; text-align:right; vertical-align:top;"><div style="font-weight:900;">العضو القانوني</div><div>{عضو_قانوني1}</div><div style="margin-top:12px;">....................</div><div style="margin-top:20px;">تحرر في: {datetime.now().strftime('%d-%m-%Y')}</div></td><td style="width:50%; text-align:left; vertical-align:top;"><div style="font-weight:900;">مدير إدارة القضايا</div><div>{مدير_ادارة1}</div><div style="margin-top:12px;">....................</div></td></tr></table><div style="text-align:center; margin-top:20px;"><div style="font-weight:900; color:#dc3545;">مدير عام الإدارات القانونية</div><div>{مدير_عام1}</div><div style="margin-top:12px;">....................</div></div></div>"""
                full_html = header_html + f"<div class='table-container'>{html}</div>" + footer
                st.markdown(full_html, unsafe_allow_html=True)
                st.session_state.last_report_html = full_html; st.session_state.last_report_title = f"بيان_الدعاوى_المتداولة_{region}"; st.session_state.last_report_df = pd.DataFrame(df_data)
                html_export = get_export_html(full_html, st.session_state.last_report_title)
                c1,c2,c3 = st.columns(3)
                with c1: st.download_button("⬇️ PDF", data=html_export.encode('utf-8'), file_name=f"بيان_الدعاوى_{region}.html", mime="text/html", use_container_width=True, key="dl1")
                with c2: st.download_button("⬇️ Word", data=html_export.encode('utf-8'), file_name=f"بيان_الدعاوى_{region}.doc", mime="application/msword", use_container_width=True, key="dl2")
                with c3:
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                        df = pd.DataFrame(df_data)
                        df.to_excel(writer, index=False, sheet_name='Sheet1')
                        worksheet = writer.sheets['Sheet1']
                        for col in worksheet.columns: worksheet.column_dimensions[col[0].column_letter].width = 25
                        for cell in worksheet[1]: cell.alignment = cell.alignment.copy(wrap_text=True, horizontal='right')
                    st.download_button("⬇️ Excel", data=excel_buffer.getvalue(), file_name=f"بيان_الدعاوى_{region}.xlsx", use_container_width=True, key="dlx1")
                    # ========= تبويب 2: الاحكام =========
    with tab2:
        st.markdown("<div style='background:#1E2A47; padding:12px; border-radius:8px; border:2px solid #B8860B; margin-bottom:12px'>", unsafe_allow_html=True)
        نوع_التقرير = st.selectbox("نوع البيان", ["بيان بجميع الاحكام الصادرة للصالح والضد", "بيان بالاحكام الصادرة للصالح", "بيان بالاحكام الصادرة للضد", "بيان بالاحكام الصادرة حسب موضوع الدعوى"], key="no3_tقرير")
        colA, colB, colC = st.columns(3)
        with colA: region2 = st.text_input("ديوان عام منطقة", key="region2")
        with colB: مدير_عام2 = st.text_input("اسم مدير عام الإدارات القانونية", key="modir2")
        with colC: مدير_ادارة2 = st.text_input("اسم مدير إدارة القضايا", key="modir_idara2")
        عضو_قانوني2 = st.text_input("اسم العضو القانوني", key="odo2")
        col1, col2, col3 = st.columns(3)
        with col1: from_date2 = st.date_input("من الفترة", key="from2")
        with col2: to_date2 = st.date_input("حتى الفترة", key="to2")
        with col3: lawyer2 = st.text_input("طرف الاستاذ/ المحامي", key="lawyer2")
        topic2 = ""
        if "حسب موضوع" in نوع_التقرير: topic2 = st.text_input("موضوع الدعوى للفلترة", key="topic2")
        st.markdown("</div>", unsafe_allow_html=True)

        if st.button("🔍 عرض بيان الاحكام", use_container_width=True, type="primary", key="show2"):
            archive = [c for c in my_cases if c.get("حالة") == "منتهية" and not c.get("تم_الحفظ_النهائي")]
            cases = []
            for c in archive:
                if c.get('تاريخ_الحكم'):
                    try:
                        ت_حكم = datetime.strptime(c['تاريخ_الحكم'], '%Y-%m-%d').date()
                        if from_date2 <= ت_حكم <= to_date2: cases.append(c)
                    except: pass
            if نوع_التقرير == "بيان بالاحكام الصادرة للصالح": cases = [c for c in cases if c.get('مسندة_ل_الحكم') == 'الصالح']
            elif نوع_التقرير == "بيان بالاحكام الصادرة للضد": cases = [c for c in cases if c.get('مسندة_ل_الحكم') == 'الضد']
            elif "حسب موضوع" in نوع_التقرير and topic2: cases = [c for c in cases if topic2.lower() in str(c.get('موضوع','')).lower()]
            cases = sorted(cases, key=lambda x: x.get("تاريخ_الحكم","0000-00-00"), reverse=True)
            title = f"{نوع_التقرير} خلال الفترة من {from_date2.strftime('%d-%m-%Y')} إلى {to_date2.strftime('%d-%m-%Y')} طرف الاستاذ/ {lawyer2} المحامي"
            header_html = report_header(region2, title, مدير_عام2, مدير_ادارة2, عضو_قانوني2)
            if not cases: st.warning(f"⚠️ لا توجد احكام في الفترة")
            else:
                st.success(f"✅ تم العثور على {len(cases)} حكم")
                html = "<table class='case-table'><tr><th>م</th><th>رقم</th><th>سنة</th><th>المحكمة والدائرة</th><th>الخصوم</th><th>الموضوع</th><th>تاريخ الحكم</th><th>منطوق الحكم</th><th>مسندة ل</th><th>ملاحظات</th></tr>"
                df_data = []
                for idx, c in enumerate(cases, 1):
                    محكمة = f"{c.get('نوع','')} {c.get('محكمة_اسم','')}"
                    if c.get('مأمورية'): محكمة += f" - مأمورية {c.get('مأمورية')}"
                    if c.get('دائرة'): محكمة += f" - دائرة {c.get('دائرة')}"
                    مدعي = c.get('مدعي',''); مدعي_عليه = c.get('مدعي_عليه','')
                    خصوم_html = f"<span style='color:#dc3545; font-weight:900'>{مدعي} ضد {مدعي_عليه}</span>" if "الهيئة" in str(مدعي) else f"{مدعي} ضد {مدعي_عليه}"
                    لون_مسندة = "#28a745" if c.get('مسندة_ل_الحكم') == 'الصالح' else "#dc3545"
                    html += f"<tr><td>{idx}</td><td>{c.get('رقم','')}</td><td>{c.get('سنة','')}</td><td>{محكمة}</td><td>{خصوم_html}</td><td>{c.get('موضوع','')}</td><td><b style='color:#B8860B'>{c.get('تاريخ_الحكم','')}</b></td><td>{c.get('منطوق_الحكم','')}</td><td><b style='color:{لون_مسندة}'>{c.get('مسندة_ل_الحكم')}</b></td><td>{c.get('ملاحظات','')}</td></tr>"
                    df_data.append({'م': idx, 'رقم': c.get('رقم',''), 'سنة': c.get('سنة',''), 'المحكمة': محكمة, 'الخصوم': f"{مدعي} ضد {مدعي_عليه}", 'الموضوع': c.get('موضوع',''), 'تاريخ الحكم': c.get('تاريخ_الحكم',''), 'منطوق الحكم': c.get('منطوق_الحكم',''), 'مسندة ل': c.get('مسندة_ل_الحكم'), 'ملاحظات': c.get('ملاحظات','')})
                html += "</table>"
                # ===== الفوتر المعدل: تحرر في بعد سطرين =====
                footer = f"""<div style="margin-top:25px; color:#B8860B; font-size:12px;"><p style="text-align:center; margin-bottom:20px; font-size:13px; font-weight:700;">تفضلوا بقبول وافر الاحترام والتقدير،</p><table style="width:100%;"><tr><td style="width:50%; text-align:right; vertical-align:top;"><div style="font-weight:900;">العضو القانوني</div><div>{عضو_قانوني2}</div><div style="margin-top:12px;">....................</div><div style="margin-top:20px;">تحرر في: {datetime.now().strftime('%d-%m-%Y')}</div></td><td style="width:50%; text-align:left; vertical-align:top;"><div style="font-weight:900;">مدير إدارة القضايا</div><div>{مدير_ادارة2}</div><div style="margin-top:12px;">....................</div></td></tr></table><div style="text-align:center; margin-top:20px;"><div style="font-weight:900; color:#dc3545;">مدير عام الإدارات القانونية</div><div>{مدير_عام2}</div><div style="margin-top:12px;">....................</div></div></div>"""
                full_html = header_html + f"<div class='table-container'>{html}</div>" + footer
                st.markdown(full_html, unsafe_allow_html=True)
                st.session_state.last_report_html = full_html; st.session_state.last_report_title = f"بيان_الاحكام_{region2}"; st.session_state.last_report_df = pd.DataFrame(df_data)
                html_export = get_export_html(full_html, st.session_state.last_report_title)
                c1,c2,c3 = st.columns(3)
                with c1: st.download_button("⬇️ PDF", data=html_export.encode('utf-8'), file_name=f"بيان_الاحكام_{region2}.html", mime="text/html", use_container_width=True, key="dl21")
                with c2: st.download_button("⬇️ Word", data=html_export.encode('utf-8'), file_name=f"بيان_الاحكام_{region2}.doc", mime="application/msword", use_container_width=True, key="dl22")
                with c3:
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                        df = pd.DataFrame(df_data)
                        df.to_excel(writer, index=False, sheet_name='Sheet1')
                        worksheet = writer.sheets['Sheet1']
                        for col in worksheet.columns: worksheet.column_dimensions[col[0].column_letter].width = 25
                        for cell in worksheet[1]: cell.alignment = cell.alignment.copy(wrap_text=True, horizontal='right')
                    st.download_button("⬇️ Excel", data=excel_buffer.getvalue(), file_name=f"بيان_الاحكام_{region2}.xlsx", use_container_width=True, key="dlx2")
                    # ================= القسم 2 من 2 =================
                # ========= تبويب 3: الاحصائيات =========
    with tab3:
        st.markdown("<h3 style='color:#B8860B; text-align:center; font-size:16px;'>📈 الإحصائيات العددية</h3>", unsafe_allow_html=True)
        st.markdown("<div style='background:#1E2A47; padding:12px; border-radius:8px; border:2px solid #B8860B; margin-bottom:12px'>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        with col1: stat_from = st.date_input("من تاريخ", key="s1")
        with col2: stat_to = st.date_input("حتى تاريخ", key="s2")
        st.markdown("</div>", unsafe_allow_html=True)
        if st.button("استخراج الإحصائيات", use_container_width=True, type="primary"):
            all_cases = my_cases
            متداولة = [c for c in all_cases if c.get('حالة') == 'متداولة' and c.get('تاريخ_جلسة') and stat_from <= datetime.strptime(c['تاريخ_جلسة'], '%Y-%m-%d').date() <= stat_to]
            احكام = [c for c in all_cases if c.get('حالة') == 'منتهية' and not c.get('تم_الحفظ_النهائي') and c.get('تاريخ_الحكم') and stat_from <= datetime.strptime(c['تاريخ_الحكم'], '%Y-%m-%d').date() <= stat_to]
            للصالح = [c for c in احكام if c.get('مسندة_ل_الحكم') == 'الصالح']
            للضد = [c for c in احكام if c.get('مسندة_ل_الحكم') == 'الضد']
            c1,c2,c3,c4 = st.columns(4)
            c1.metric("عدد القضايا المتداولة", len(متداولة))
            c2.metric("عدد الاحكام الصادرة", len(احكام))
            c3.metric("عدد الاحكام للصالح", len(للصالح))
            c4.metric("عدد الاحكام للضد", len(للضد))

    # ========= تبويب 4: البيان العددي =========
    with tab4:
        st.markdown("<h3 style='color:#B8860B; text-align:center; font-size:16px;'>📊 البيان العددي</h3>", unsafe_allow_html=True)
        st.markdown("<div style='background:#1E2A47; padding:12px; border-radius:8px; border:2px solid #B8860B; margin-bottom:12px'>", unsafe_allow_html=True)
        نوع_البيان_العددي = st.selectbox("نوع البيان العددي", ["جميع الدعاوى المتداولة", "الدعاوى المتداولة حسب موضوع الدعوى", "جميع الاحكام للصالح والضد", "الاحكام الصادرة للصالح", "الاحكام الصادرة للضد", "الاحكام حسب موضوع الدعوى"], key="no4_عددي")
        colA, colB, colC = st.columns(3)
        with colA: region_stat = st.text_input("ديوان عام منطقة", key="region_stat")
        with colB: مدير_عام_stat = st.text_input("اسم مدير عام الإدارات القانونية", key="modir_stat")
        with colC: مدير_ادارة_stat = st.text_input("اسم مدير إدارة القضايا", key="modir_idara_stat")
        عضو_قانوني_stat = st.text_input("اسم العضو القانوني", key="odo_stat")
        lawyer_stat = st.text_input("طرف الاستاذ/ المحامي", key="lawyer_stat")
        col1, col2 = st.columns(2)
        with col1: stat_from2 = st.date_input("من الفترة", key="s_from2")
        with col2: stat_to2 = st.date_input("حتى الفترة", key="s_to2")
        topic_stat = ""
        if "حسب موضوع" in نوع_البيان_العددي: topic_stat = st.text_input("موضوع الدعوى للفلترة", key="topic_stat")
        st.markdown("</div>", unsafe_allow_html=True)

        if st.button("🔍 عرض البيان العددي", use_container_width=True, type="primary", key="show_stat"):
            all_cases = my_cases
            العدد = 0; اسم_البيان = نوع_البيان_العددي
            if "الدعاوى المتداولة" in نوع_البيان_العددي:
                cases = [c for c in all_cases if str(c.get('حالة','')).strip() == 'متداولة']
                for c in cases:
                    if c.get('تاريخ_جلسة'):
                        try:
                            ت_جلسة = datetime.strptime(c['تاريخ_جلسة'], '%Y-%m-%d').date()
                            if stat_from2 <= ت_جلسة <= stat_to2:
                                if "حسب موضوع" in نوع_البيان_العددي and topic_stat:
                                    if topic_stat.lower() in str(c.get('موضوع','')).lower(): العدد += 1
                                elif "جميع" in نوع_البيان_العددي: العدد += 1
                        except: pass
            elif "الاحكام" in نوع_البيان_العددي:
                archive = [c for c in all_cases if c.get("حالة") == "منتهية" and not c.get("تم_الحفظ_النهائي")]
                for c in archive:
                    if c.get('تاريخ_الحكم'):
                        try:
                            ت_حكم = datetime.strptime(c['تاريخ_الحكم'], '%Y-%m-%d').date()
                            if stat_from2 <= ت_حكم <= stat_to2:
                                مسندة = c.get('مسندة_ل_الحكم')
                                if "حسب موضوع" in نوع_البيان_العددي and topic_stat:
                                    if topic_stat.lower() in str(c.get('موضوع','')).lower(): العدد += 1
                                elif "للصالح" in نوع_البيان_العددي and مسندة == 'الصالح': العدد += 1
                                elif "للضد" in نوع_البيان_العددي and مسندة == 'الضد': العدد += 1
                                elif "للصالح والضد" in نوع_البيان_العددي: العدد += 1
                        except: pass
            st.success(f"✅ العدد الإجمالي: {العدد}")
            title_stat = f"بيان عددي ب{اسم_البيان} خلال الفترة من {stat_from2.strftime('%d-%m-%Y')} إلى {stat_to2.strftime('%d-%m-%Y')} طرف الاستاذ/ {lawyer_stat} المحامي"
            header_html = report_header(region_stat, title_stat, مدير_عام_stat, مدير_ادارة_stat, عضو_قانوني_stat)
            ملاحظات = f"عن موضوع: {topic_stat}" if topic_stat else ""
            html = f"<table class='case-table'><tr><th>م</th><th>البيان</th><th>العدد</th><th>ملاحظات</th></tr>"
            html += f"<tr><td>1</td><td>{اسم_البيان}</td><td><b style='color:#dc3545; font-size:14px'>{العدد}</b></td><td>{ملاحظات}</td></tr>"
            html += "</table>"
            # ===== الفوتر المعدل: تحرر في بعد سطرين =====
            footer = f"""<div style="margin-top:25px; color:#B8860B; font-size:12px;"><p style="text-align:center; margin-bottom:20px; font-size:13px; font-weight:700;">تفضلوا بقبول وافر الاحترام والتقدير،</p><table style="width:100%;"><tr><td style="width:50%; text-align:right; vertical-align:top;"><div style="font-weight:900;">العضو القانوني</div><div>{عضو_قانوني_stat}</div><div style="margin-top:12px;">....................</div><div style="margin-top:20px;">تحرر في: {datetime.now().strftime('%d-%m-%Y')}</div></td><td style="width:50%; text-align:left; vertical-align:top;"><div style="font-weight:900;">مدير إدارة القضايا</div><div>{مدير_ادارة_stat}</div><div style="margin-top:12px;">....................</div></td></tr></table><div style="text-align:center; margin-top:20px;"><div style="font-weight:900; color:#dc3545;">مدير عام الإدارات القانونية</div><div>{مدير_عام_stat}</div><div style="margin-top:12px;">....................</div></div></div>"""
            full_html = header_html + f"<div class='table-container'>{html}</div>" + footer
            st.markdown(full_html, unsafe_allow_html=True)
            df_data = pd.DataFrame([{'م': 1, 'البيان': اسم_البيان, 'العدد': العدد, 'ملاحظات': ملاحظات}])
            st.session_state.last_report_html = full_html; st.session_state.last_report_title = f"بيان_عددي_{region_stat}"; st.session_state.last_report_df = df_data
            html_export = get_export_html(full_html, st.session_state.last_report_title)
            c1,c2,c3 = st.columns(3)
            with c1: st.download_button("⬇️ PDF", data=html_export.encode('utf-8'), file_name=f"بيان_عددي_{region_stat}.html", mime="text/html", use_container_width=True, key="dl_stat1")
            with c2: st.download_button("⬇️ Word", data=html_export.encode('utf-8'), file_name=f"بيان_عددي_{region_stat}.doc", mime="application/msword", use_container_width=True, key="dl_stat2")
            with c3:
                excel_buffer = io.BytesIO()
                with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                    df = df_data
                    df.to_excel(writer, index=False, sheet_name='Sheet1')
                    worksheet = writer.sheets['Sheet1']
                    for col in worksheet.columns: worksheet.column_dimensions[col[0].column_letter].width = 25
                    for cell in worksheet[1]: cell.alignment = cell.alignment.copy(wrap_text=True, horizontal='right')
                st.download_button("⬇️ Excel", data=excel_buffer.getvalue(), file_name=f"بيان_عددي_{region_stat}.xlsx", use_container_width=True, key="dl_stat3")
# ========================= نهاية الجزء 
import json
import os
from datetime import datetime, date


# ===== إدارة اليافطات: تخزين دائم بدون حذف تلقائي =====
BANNERS_FILE = os.path.join(LOCAL_DATA_DIR, "banners_v2.json")

def load_banners():
    banners = load_json_persistent("banners", BANNERS_FILE, [])
    if not isinstance(banners, list):
        return []
    changed = False
    normalized = []
    for b in banners:
        if not isinstance(b, dict):
            continue
        b = dict(b)
        if not b.get("id"):
            b["id"] = secrets.token_hex(8)
            changed = True
        if not b.get("expire_at") and b.get("expire"):
            b["expire_at"] = str(b.get("expire"))
            changed = True
        b.setdefault("audience", "الكل")
        b.setdefault("visible_to", [])
        normalized.append(b)
    if changed:
        save_json_persistent("banners", BANNERS_FILE, normalized)
    return normalized

def save_banners(banners):
    save_json_persistent("banners", BANNERS_FILE, banners)

def show_banners():
    """عرض اليافطات أعلى الصفحة. لا نحذف اليافطات المنتهية من البيانات."""
    if not st.session_state.get("user"):
        return
    username = st.session_state.user.get("username", "")
    now = datetime.now()
    banners = load_banners()
    visible = []
    for b in banners:
        raw = b.get("expire_at") or b.get("expire")
        try:
            expire_dt = datetime.fromisoformat(str(raw))
        except Exception:
            continue
        if expire_dt <= now:
            continue
        audience = b.get("audience", "الكل")
        if audience in ("الكل", "كل الأعضاء") or username in b.get("visible_to", []):
            visible.append(b)

    for b in visible:
        st.markdown(
            f"<div style='background:linear-gradient(90deg,{b.get('color','#D4AF37')},#ffffff22);"
            "color:#fff;padding:16px;border-radius:16px;margin:0 0 14px;"
            "font-weight:900;font-size:20px;text-align:center;border:2px solid "
            f"{b.get('color','#D4AF37')};box-shadow:0 8px 22px rgba(0,0,0,.25);'>📢 {b.get('text','')}</div>",
            unsafe_allow_html=True
        )

# ===== تشغيل الدخول والصفحات الإدارية الموحدة =====
if st.session_state.get("page") == "login":
    login_page()
elif st.session_state.get("page") == "extract_member":
    if st.session_state.get("user", {}).get("role") == "admin":
        extract_member_page()
elif st.session_state.get("page") == "لوحة تحكمى":
    if st.session_state.get("user", {}).get("role") in ("member", "manager", "general_manager", "leader1", "leader2", "leader3"):
        personal_control_page()
    else:
        st.session_state.page = "login"
        st.rerun()
elif st.session_state.get("page") == "set_password":
    set_password_page()
elif st.session_state.get("page") == "change_password":
    change_password_page()

# ===== تشغيل لوحة التحكم الموحدة =====
if st.session_state.get("page") == "لوحة التحكم":
    if st.session_state.get("user", {}).get("role") == "admin":
        admin_control_page()
    else:
        st.session_state.page = "login"
        st.rerun()
elif st.session_state.get("page") == "الرسائل":
    if st.session_state.get("user", {}).get("role") in ("admin", "member", "manager", "general_manager", "leader1", "leader2", "leader3"):
        messaging_page()
    else:
        st.session_state.page = "الرئيسية"
        st.rerun()
